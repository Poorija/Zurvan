import sys
import os
if os.geteuid() == 0:
    os.environ["QTWEBENGINE_CHROMIUM_FLAGS"] = "--no-sandbox"
    if "--no-sandbox" not in sys.argv:
        sys.argv.append("--no-sandbox")
import logging
import time
from threading import Event, Lock
import queue
import copy
import socket
import glob
import random
import csv
import platform
import psutil
import ipaddress
from PyQt6.QtCore import PYQT_VERSION_STR
import subprocess
import numpy as np
import json
import urllib.request
import tempfile
import webbrowser
import shutil
import signal
import uuid
import sqlite3
import gzip
from datetime import datetime
import hashlib
import pyotp
import base64
import requests

from app_lock import AppLockDialog

try:
    from lxml import etree
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False
    etree = None
    logging.warning("Optional XML reporting dependency not found. Please run 'pip install lxml'")

try:
    from geoip import geolite2
    GEOLITE_AVAILABLE = True
except ImportError:
    GEOLITE_AVAILABLE = False
    logging.warning("Optional offline GeoIP dependency not found. Please run 'pip install python-geoip-geolite2'")

import re
from theme_support import apply_stylesheet, list_themes
from PyQt6.QtGui import QActionGroup, QPixmap, QImage, QPalette, QPainter, QPainterPath
from PyQt6.QtSvg import QSvgRenderer

def create_themed_icon(icon_path, color_str):
    """Loads an SVG, intelligently replaces its color, and returns a QIcon."""
    try:
        with open(icon_path, 'r', encoding='utf-8') as f:
            svg_data = f.read()

        # First, try to replace a stroke color in a style block (for paper-airplane.svg)
        themed_svg_data, count = re.subn(r'stroke:#[0-9a-fA-F]{6}', f'stroke:{color_str}', svg_data)

        # If no stroke was found in a style, fall back to injecting a fill attribute (for gear.svg)
        if count == 0 and '<svg' in themed_svg_data:
            themed_svg_data = themed_svg_data.replace('<svg', f'<svg fill="{color_str}"')

        image = QImage.fromData(themed_svg_data.encode('utf-8'))
        pixmap = QPixmap.fromImage(image)
        return QIcon(pixmap)
    except Exception as e:
        logging.warning(f"Could not create themed icon for {icon_path}: {e}")
        # On failure, attempt to load the original icon directly.
        # This can still fail if the path is wrong, but it's a better fallback.
        return QIcon(icon_path)

def get_vendor(mac_address):
    """Retrieves the vendor for a given MAC address from an online API."""
    if not mac_address or mac_address == "N/A":
        return "N/A"
    try:
        # Use a timeout to prevent the application from hanging on network issues
        with urllib.request.urlopen(f"https://api.macvendors.com/{mac_address}", timeout=3) as url:
            data = url.read().decode()
            return data
    except Exception as e:
        logging.warning(f"Could not retrieve vendor for MAC {mac_address}: {e}")
        return "Unknown Vendor"

def _get_random_ip():
    """Generates a random, non-private IP address."""
    while True:
        ip = ".".join(str(random.randint(1, 223)) for _ in range(4))
        if not (ip.startswith('10.') or ip.startswith('192.168.') or (ip.startswith('172.') and 16 <= int(ip.split('.')[1]) <= 31)):
             return ip

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QStatusBar, QMenuBar, QTabWidget, QWidget,
    QVBoxLayout, QLabel, QDockWidget, QPlainTextEdit, QPushButton, QHBoxLayout,
    QTreeWidget, QTreeWidgetItem, QSplitter, QFileDialog, QMessageBox, QComboBox,
    QListWidget, QListWidgetItem, QScrollArea, QLineEdit, QCheckBox, QFrame, QMenu, QTextEdit, QGroupBox,
    QProgressBar, QTextBrowser, QRadioButton, QButtonGroup, QFormLayout, QGridLayout, QDialog,
    QHeaderView, QInputDialog, QGraphicsOpacityEffect, QStackedWidget, QToolButton, QTableView, QDateEdit, QSpinBox
)
from PyQt6.QtCore import QPropertyAnimation, QEasingCurve, QParallelAnimationGroup, QSequentialAnimationGroup
from ai_tab import AIAssistantTab, AISettingsDialog, AIGuideDialog
from login import LoginDialog
from admin_panel import AdminPanelDialog
from user_profile import UserProfileDialog
from offline_cve_manager import OfflineCveManagerWidget
from speed_test_tab import SpeedTestTab
from ssh_manager_tab import SshManagerTab
from PyQt6.QtCore import QObject, pyqtSignal, Qt, QThread, QTimer, QPropertyAnimation, QEasingCurve, QParallelAnimationGroup, QSequentialAnimationGroup, QPoint, QSize, QAbstractTableModel, QDate, QEvent
from PyQt6.QtGui import QAction, QIcon, QFont, QTextCursor, QActionGroup, QColor

class ActivityMonitorEventFilter(QObject):
    """An event filter that monitors for user activity and emits a signal."""
    user_active = pyqtSignal()

    def eventFilter(self, obj, event):
        # List of events that indicate user activity
        activity_events = [
            QEvent.Type.KeyPress,
            QEvent.Type.KeyRelease,
            QEvent.Type.MouseButtonPress,
            QEvent.Type.MouseButtonRelease,
            QEvent.Type.MouseButtonDblClick,
            QEvent.Type.MouseMove,
            QEvent.Type.Wheel,
        ]
        if event.type() in activity_events:
            self.user_active.emit()
        # Pass the event on
        return super().eventFilter(obj, event)


def sniffer_process_target(queue, iface, bpf_filter):
    """
    This function runs in a separate process. It sniffs packets and puts them
    into a multiprocessing.Queue. This completely isolates the blocking
    sniff() call from the main GUI application.
    """
    try:
        # The packet handler now simply puts the raw packet into the queue
        def packet_handler(packet):
            queue.put(bytes(packet))

        # We don't need a stop_filter anymore, as the process will be terminated directly.
        sniff(prn=packet_handler, iface=iface, filter=bpf_filter, store=False)
    except Exception as e:
        logging.error(f"Critical error in sniffer process: {e}", exc_info=True)


class KrackScanThread(QThread):
    vulnerability_detected = pyqtSignal(str, str) # bssid, client_mac

    def __init__(self, iface, parent=None):
        super().__init__(parent)
        self.iface = iface
        self.stop_event = Event()
        self.eapol_db = {} # { (bssid, client_mac): { replay_counter: count } }

    def _packet_handler(self, pkt):
        if not pkt.haslayer(EAPOL) or not pkt.haslayer(Dot11):
            return

        # Check if frame is going from AP to client (To DS=0, From DS=1)
        if pkt.FCfield & 0x3 != 1:
            return

        try:
            # Key Information field is a good indicator for Message 3
            key_info = pkt[EAPOL].key_info
            # Message 3: Pairwise, Install, Ack, MIC
            # Install = bit 6 (0x40), Ack = bit 7 (0x80), MIC = bit 8 (0x100)
            is_msg3 = (key_info & 0x1c0) == 0x1c0

            if is_msg3:
                bssid = pkt.addr2
                client_mac = pkt.addr1
                replay_counter = pkt[EAPOL].replay_counter

                key = (bssid, client_mac)

                if key not in self.eapol_db:
                    self.eapol_db[key] = {}

                if replay_counter not in self.eapol_db[key]:
                    self.eapol_db[key][replay_counter] = 1
                else:
                    # If we see the same replay counter again, it's a retransmission
                    self.eapol_db[key][replay_counter] += 1
                    if self.eapol_db[key][replay_counter] == 2:
                        logging.info(f"KRACK vulnerability detected! BSSID: {bssid}, Client: {client_mac}")
                        self.vulnerability_detected.emit(bssid, client_mac)
                        # Reset counter to avoid flooding with signals for the same retransmission
                        self.eapol_db[key][replay_counter] = 0


        except (IndexError, AttributeError) as e:
            logging.debug(f"Error processing EAPOL packet for KRACK scan: {e}")

    def run(self):
        logging.info(f"KRACK scanner started on interface {self.iface}")
        while not self.stop_event.is_set():
            try:
                sniff(iface=self.iface, prn=self._packet_handler, filter="ether proto 0x888e", timeout=1)
            except Exception as e:
                logging.error(f"Error in KRACK sniffer loop: {e}", exc_info=True)
                time.sleep(1)

    def stop(self):
        self.stop_event.set()


class AircrackThread(QThread):
    """A thread to run the aircrack-ng process and emit its output."""
    output_received = pyqtSignal(str)
    finished_signal = pyqtSignal(int)

    def __init__(self, pcap_file, wordlist, parent=None, threads=1):
        super().__init__(parent)
        self.pcap_file = pcap_file
        self.wordlist = wordlist
        self.threads = threads
        self.process = None

    def run(self):
        command = ["aircrack-ng", "-w", self.wordlist, "-p", str(self.threads), self.pcap_file]
        try:
            self.process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1)
            for line in iter(self.process.stdout.readline, ''):
                self.output_received.emit(line.strip())
            self.process.stdout.close()
            return_code = self.process.wait()
            self.finished_signal.emit(return_code)
        except FileNotFoundError:
            self.output_received.emit("ERROR: 'aircrack-ng' command not found. Please ensure it is installed and in your system's PATH.")
            self.finished_signal.emit(-1)
        except Exception as e:
            self.output_received.emit(f"An unexpected error occurred: {e}")
            self.finished_signal.emit(-1)

    def stop(self):
        if self.process and self.process.poll() is None:
            self.process.terminate()
            self.process.wait()
            logging.info("Aircrack-ng process terminated.")

try:
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    import docx
except ImportError:
    logging.warning("Optional PDF/DOCX export dependencies not found. Please run 'pip install reportlab python-docx'")

try:
    import pyqtgraph as pg
    PYQTGRAPH_AVAILABLE = True
except ImportError:
    PYQTGRAPH_AVAILABLE = False
    pg = None # Define pg as None to prevent other errors if it's referenced
    logging.warning("Optional graphing dependency not found. Please run 'pip install pyqtgraph'")


try:
    import GPUtil
except Exception as e:
    GPUtil = None
    logging.warning(f"Optional GPU monitoring dependency unavailable: {e}")

# --- Scapy Imports ---
try:
    from scapy.all import *
    from scapy.utils import hexdump
    from scapy.layers.dns import DNS, DNSQR
    from scapy.layers.dot11 import Dot11, Dot11Beacon, Dot11Elt, Dot11Deauth, RadioTap
    conf.verb = 0
except ImportError:
    logging.critical("Scapy is not installed.")

# --- Constants ---
AVAILABLE_PROTOCOLS = {"Ethernet": Ether, "ARP": ARP, "IP": IP, "IPv6": IPv6, "TCP": TCP, "UDP": UDP, "ICMP": ICMP, "DNS": DNS, "Raw": Raw}
PACKET_TEMPLATES = {
    "ICMP Ping (google.com)": [IP(dst="8.8.8.8"), ICMP()],
    "DNS Query (google.com)": [IP(dst="8.8.8.8"), UDP(dport=53), DNS(rd=1, qd=DNSQR(qname="google.com"))],
    "TCP SYN (localhost:80)": [IP(dst="127.0.0.1"), TCP(dport=80, flags="S")],
    "ARP Request (who-has 192.168.1.1)": [Ether(dst="ff:ff:ff:ff:ff:ff"), ARP(pdst="192.168.1.1")],
    "NTP Query (pool.ntp.org)": [IP(dst="pool.ntp.org"), UDP(sport=123, dport=123), NTP()],
    "SNMP GetRequest (public)": [IP(dst="127.0.0.1"), UDP(), SNMP(community="public", PDU=SNMPget(varbindlist=[SNMPvarbind(oid=ASN1_OID('1.3.6.1.2.1.1.1.0'))]))]
}
FIREWALL_PROBES = {
    "Standard SYN Scan (Top Ports)": [(lambda t: IP(dst=t)/TCP(dport=p, flags="S"), f"TCP SYN to port {p}") for p in [21, 22, 25, 53, 80, 110, 143, 443, 445, 3389, 8080]],
    "Stealthy Scans (FIN, Xmas, Null)": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="F"), f"FIN Scan to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="FPU"), f"Xmas Scan to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags=""), f"Null Scan to port {p}") for p in [80, 443]
    ],
    "ACK Scan (Firewall Detection)": [(lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="A"), f"ACK Scan to port {p}") for p in [22, 80, 443]],
    "Source Port Evasion (DNS)": [(lambda t, p=p: IP(dst=t)/TCP(sport=53, dport=p, flags="S"), f"SYN from port 53 to {p}") for p in [80, 443, 8080]],
    "Fragmented SYN Scan": [(lambda t, p=p: fragment(IP(dst=t)/TCP(dport=p, flags="S")), f"Fragmented SYN to port {p}") for p in [80, 443]],
    "TCP Options Probes (WScale, TS)": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="S", options=[('WScale', 10), ('Timestamp', (12345, 0))]), f"SYN+WScale+TS to port {p}") for p in [80, 443]
    ],
    "ECN Flag Probes": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="SE"), f"SYN+ECE to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="SC"), f"SYN+CWR to port {p}") for p in [80, 443]
    ],
    "HTTP Payload Probe": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="PA")/Raw(load="GET / HTTP/1.0\r\n\r\n"), f"HTTP GET probe to port {p}") for p in [80, 8080, 443]
    ],
    "Common UDP Probes": [(lambda t, p=p: IP(dst=t)/UDP(dport=p), f"UDP Probe to port {p}") for p in [53, 123, 161]],
    "ICMP Probes (Advanced)": [
        (lambda t: IP(dst=t)/ICMP(type=ty), f"ICMP Echo Request (Type 8)") for ty in [8]
    ] + [
        (lambda t: IP(dst=t)/ICMP(type=ty), f"ICMP Timestamp Request (Type 13)") for ty in [13]
    ] + [
        (lambda t: IP(dst=t)/ICMP(type=ty), f"ICMP Address Mask Request (Type 17)") for ty in [17]
    ],
    "Advanced Evasion Techniques": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="S", options=[('MSS', 10)]), f"SYN with tiny MSS to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t, flags="DF", frag=0)/TCP(dport=p, flags="S", seq=RandInt())/("X"*32), f"SYN with 'Don't Fragment' and data to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="S", chksum=0xf00d), f"SYN with bad checksum to port {p}") for p in [80, 443]
    ]
}
SCAN_TYPES = ["TCP SYN Scan", "TCP FIN Scan", "TCP Xmas Scan", "TCP Null Scan", "TCP ACK Scan", "UDP Scan"]
COMMON_FILTERS = [
    "", "tcp", "udp", "arp", "icmp",
    "port 80", "port 443", "udp port 53", "tcp port 22",
    "host 8.8.8.8", "net 192.168.1.0/24", "vlan"
]

COMMUNITY_TOOLS = {
    "Interpreters and REPLs": [
        ("scapy-console", "https://github.com/gpotter2/scapy-console", "A Scapy console with many other tools and features."),
        ("Scapy REPL", "https://github.com/GabrielCama/scapy-repl", "An interactive Scapy REPL with customized commands.")
    ],
    "Networking": [
        ("bettercap", "https://github.com/bettercap/bettercap", "A powerful, flexible and portable tool for network attacks and monitoring."),
        ("Routersploit", "https://github.com/threat9/routersploit", "An open-source exploitation framework dedicated to embedded devices."),
        ("Batfish", "https://www.batfish.org/", "A network configuration analysis tool for validating and verifying network designs.")
    ],
    "Network Scanners & Analyzers": [
        ("Wireshark", "https://www.wireshark.org/", "The world's foremost and widely-used network protocol analyzer."),
        ("Nmap", "https://nmap.org/", "The Network Mapper - a free and open source utility for network discovery and security auditing."),
        ("Naabu", "https://github.com/projectdiscovery/naabu", "A modern fast port scanner that pairs well with Nmap for follow-up service detection."),
        ("Zeek", "https://zeek.org/", "A powerful network analysis framework that is much different from a typical IDS."),
        ("BruteShark", "https://github.com/odedshimon/BruteShark", "An open-source, cross-platform network forensic analysis tool (NFAT).")
    ],
    "Wireless": [
        ("Kismet", "https://www.kismetwireless.net/", "A wireless network detector, sniffer, and intrusion detection system."),
        ("Airgeddon", "https://github.com/v1s1t0r1sh3r3/airgeddon", "A multi-use bash script for Linux systems to audit wireless networks."),
        ("wifiphisher", "https://github.com/wifiphisher/wifisher", "A rogue Access Point framework for conducting red team engagements or Wi-Fi security testing."),
        ("Wifite2", "https://github.com/derv82/wifite2", "A complete rewrite of the popular wireless network auditing tool, wifite.")
    ],
    "Password Cracking": [
        ("John the Ripper", "https://www.openwall.com/john/", "A fast password cracker, available for many operating systems."),
        ("Hashcat", "https://hashcat.net/hashcat/", "The world's fastest and most advanced password recovery utility."),
        ("hcxtools", "https://github.com/ZerBea/hcxtools", "Tools to convert Wi-Fi captures into hash formats for Hashcat or John.")
    ],
    "Web & API Security": [
        ("Katana", "https://github.com/projectdiscovery/katana", "A modern crawling and spidering framework for web attack-surface discovery."),
        ("OWASP Amass", "https://github.com/owasp-amass/amass", "A widely used attack surface mapping and external asset discovery framework."),
        ("Feroxbuster", "https://github.com/epi052/feroxbuster", "A fast content discovery tool for recursive web enumeration."),
        ("reNgine", "https://github.com/yogeshojha/rengine", "An automated reconnaissance framework for web applications."),
        ("Astra", "https://github.com/flipkart-incubator/Astra", "Automated Security Testing For REST APIs.")
    ],
}

class CrunchDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Crunch Wordlist Generator")
        layout = QVBoxLayout(self)
        form_layout = QFormLayout()

        self.min_len = QLineEdit("8")
        self.max_len = QLineEdit("8")
        self.charset = QLineEdit("abcdefghijklmnopqrstuvwxyz0123456789")
        self.output_file = QLineEdit()
        self.output_file.setReadOnly(True)

        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(self.browse_output)

        form_layout.addRow("Min Length:", self.min_len)
        form_layout.addRow("Max Length:", self.max_len)
        form_layout.addRow("Character Set:", self.charset)

        output_layout = QHBoxLayout()
        output_layout.addWidget(self.output_file)
        output_layout.addWidget(browse_btn)
        form_layout.addRow("Output File:", output_layout)

        layout.addLayout(form_layout)

        self.generate_button = QPushButton("Generate")
        self.generate_button.clicked.connect(self.accept)
        layout.addWidget(self.generate_button)

    def browse_output(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Wordlist", "", "Text Files (*.txt)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.output_file.setText(file_path)

    def get_values(self):
        return {
            "min": self.min_len.text(),
            "max": self.max_len.text(),
            "charset": self.charset.text(),
            "outfile": self.output_file.text()
        }

# --- Logging and Threads ---
class QtLogHandler(logging.Handler, QObject):
    """A custom logging handler that emits a Qt signal for each log record."""
    log_updated = pyqtSignal(str)
    def __init__(self): super().__init__(); QObject.__init__(self)
    def emit(self, record): self.log_updated.emit(self.format(record))

class SnifferThread(QThread):
    """
    This QThread does not sniff itself. Instead, it manages a separate
    multiprocessing.Process for sniffing to prevent the GUI from freezing.
    It communicates with the main thread exclusively via thread-safe Qt signals
    that carry raw bytes, not complex objects.
    """
    packet_bytes_received = pyqtSignal(bytes)

    def __init__(self, iface, bpf_filter, parent=None):
        super().__init__(parent)
        self.iface = iface
        self.bpf_filter = bpf_filter
        self.process = None
        self.queue = None
        self.stop_event = Event()

    def run(self):
        from multiprocessing import Process, Queue
        self.queue = Queue()
        self.process = Process(
            target=sniffer_process_target,
            args=(self.queue, self.iface, self.bpf_filter)
        )
        self.process.start()
        logging.info(f"Sniffer process started with PID: {self.process.pid}")

        while not self.stop_event.is_set():
            try:
                # Use a timeout on the queue to remain responsive
                pkt_bytes = self.queue.get(timeout=0.5)
                # Emit the raw bytes. Reconstruction will happen in the main thread.
                self.packet_bytes_received.emit(pkt_bytes)
            except queue.Empty:
                continue
            except Exception as e:
                logging.error(f"Error in SnifferThread queue loop: {e}")

        logging.info("SnifferThread manager loop stopped.")


    def stop(self):
        logging.info("Stopping sniffer manager thread and process...")
        self.stop_event.set()
        if self.process and self.process.is_alive():
            logging.info(f"Terminating sniffer process {self.process.pid}...")
            self.process.terminate()
            self.process.join(timeout=2) # Wait for the process to terminate
            if self.process.is_alive():
                logging.warning(f"Sniffer process {self.process.pid} did not terminate gracefully, killing.")
                self.process.kill()
            logging.info("Sniffer process stopped.")

class ChannelHopperThread(QThread):
    """A thread to automatically hop Wi-Fi channels on Linux for scanning."""
    def __init__(self, iface):
        super().__init__()
        self.iface = iface
        self.stop_event = Event()
    def run(self):
        if sys.platform != "linux":
            logging.warning("Channel hopping is only supported on Linux.")
            return
        logging.info(f"Channel hopper started for interface {self.iface}")
        channels = [1, 6, 11, 2, 7, 3, 8, 4, 9, 5, 10]
        while not self.stop_event.is_set():
            for ch in channels:
                if self.stop_event.is_set(): break
                try:
                    os.system(f"iwconfig {self.iface} channel {ch}")
                    time.sleep(0.5)
                except Exception as e:
                    logging.error(f"Failed to hop channel: {e}")
                    break
        logging.info("Channel hopper stopped.")
    def stop(self): self.stop_event.set()

class WorkerThread(QThread):
    """A generic QThread to run any function in the background."""
    def __init__(self, target, args=()): super().__init__(); self.target = target; self.args = args
    def run(self): self.target(*self.args)

class ResourceMonitorThread(QThread):
    """A thread that monitors and emits system resource usage statistics."""
    stats_updated = pyqtSignal(dict)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.stop_event = Event()
        self.is_paused = False
        self.interval = 1 # default interval

    def run(self):
        """The main loop for monitoring resources."""
        psutil.cpu_percent() # Initial call to prevent first reading from being 0.0
        last_disk_io = psutil.disk_io_counters()
        last_net_io = psutil.net_io_counters()

        while not self.stop_event.is_set():
            if self.is_paused:
                time.sleep(1)
                continue

            time.sleep(self.interval)

            if self.stop_event.is_set():
                break

            cpu_percent = psutil.cpu_percent()
            ram_percent = psutil.virtual_memory().percent

            # GPU Stats
            gpu_percent = 0
            if GPUtil:
                try:
                    gpus = GPUtil.getGPUs()
                    if gpus:
                        gpu = gpus[0] # Use the first GPU
                        gpu_percent = gpu.load * 100
                except Exception as e:
                    logging.debug(f"Could not retrieve GPU stats: {e}")


            disk_io = psutil.disk_io_counters()
            read_mb_s = (disk_io.read_bytes - last_disk_io.read_bytes) / (1024**2) / self.interval
            write_mb_s = (disk_io.write_bytes - last_disk_io.write_bytes) / (1024**2) / self.interval
            last_disk_io = disk_io

            net_io = psutil.net_io_counters()
            sent_kb_s = (net_io.bytes_sent - last_net_io.bytes_sent) / 1024 / self.interval
            recv_kb_s = (net_io.bytes_recv - last_net_io.bytes_recv) / 1024 / self.interval
            last_net_io = net_io

            stats = {
                "cpu_percent": cpu_percent,
                "ram_percent": ram_percent,
                "gpu_percent": gpu_percent,
                "disk_str": f"{read_mb_s:.2f}/{write_mb_s:.2f} MB/s",
                "net_str": f"{sent_kb_s:.2f}/{recv_kb_s:.2f} KB/s"
            }
            self.stats_updated.emit(stats)

    def set_interval(self, interval):
        self.interval = interval
        self.is_paused = False

    def pause(self):
        self.is_paused = True

    def stop(self):
        self.stop_event.set()

class HandshakeSnifferThread(QThread):
    """A specialized thread to capture WPA 4-way handshakes."""
    handshake_captured = pyqtSignal(str, str) # BSSID, file_path
    log_message = pyqtSignal(str)

    def __init__(self, iface, bssid, parent=None):
        super().__init__(parent)
        self.iface = iface
        self.bssid = bssid
        self.packets = []
        self.stop_event = Event()

    def run(self):
        self.log_message.emit(f"Starting handshake capture for BSSID: {self.bssid} on {self.iface}")
        try:
            sniff(iface=self.iface, prn=self._packet_handler, stop_filter=lambda p: self.stop_event.is_set(), filter="ether proto 0x888e")
        except Exception as e:
            self.log_message.emit(f"Handshake sniffer error: {e}")
        self.log_message.emit("Handshake sniffer stopped.")

    def _packet_handler(self, pkt):
        self.packets.append(pkt)
        # Simple check: once we have >= 4 EAPOL packets, save and stop.
        # A more robust implementation would check the actual handshake sequence.
        if len(self.packets) >= 4:
            self.log_message.emit("Potential handshake captured (4 EAPOL packets). Saving to file.")
            file_path = f"handshake_{self.bssid.replace(':', '')}.pcap"
            wrpcap(file_path, self.packets)
            self.handshake_captured.emit(self.bssid, file_path)
            self.stop()

    def stop(self):
        self.stop_event.set()

class IpFetchThread(QThread):
    """A thread to fetch the public IP address, optionally through a proxy."""
    ip_fetched = pyqtSignal(str, str, str)  # Signal emits (ip_type, ip_address, country_code)

    def __init__(self, ip_type, use_proxy=False, parent=None):
        super().__init__(parent)
        self.ip_type = ip_type
        self.use_proxy = use_proxy
        self.parent_app = parent

    def run(self):
        ip_address = "Error"
        country_code = "N/A"
        try:
            proxies = {}
            if self.use_proxy:
                # This logic checks for any active proxy/VPN and gets the appropriate proxy settings
                if self.parent_app.v2ray_process and self.parent_app.v2ray_process.poll() is None:
                     proxies = {"http": "socks5://127.0.0.1:10808", "https": "socks5://127.0.0.1:10808"}
                elif self.parent_app.tor_proxy_check.isChecked() and self.parent_app._check_tor_proxy():
                     proxies = {"http": "socks5://127.0.0.1:9050", "https": "socks5://127.0.0.1:9050"}
                # Add other VPN/proxy checks here in the future (OpenVPN, WireGuard, etc.)

            # Use ip-api.com as it provides country code
            response = requests.get("http://ip-api.com/json", proxies=proxies, timeout=10)
            response.raise_for_status()
            data = response.json()
            ip_address = data.get("query", "Parse Error")
            country_code = data.get("countryCode", "N/A")
        except requests.exceptions.ProxyError as e:
            logging.warning(f"Proxy error fetching IP for {self.ip_type}: {e}")
            ip_address = "Proxy Error"
        except requests.exceptions.Timeout as e:
            logging.warning(f"Timeout fetching IP for {self.ip_type}: {e}")
            ip_address = "Timeout"
        except requests.exceptions.ConnectionError as e:
            logging.warning(f"Connection error fetching IP for {self.ip_type}: {e}")
            ip_address = "Connection Error"
        except requests.exceptions.RequestException as e:
            logging.warning(f"Generic request error fetching IP for {self.ip_type}: {e}")
            ip_address = "Fetch Error"
        except Exception as e:
            logging.error(f"Unexpected error fetching IP for {self.ip_type}: {e}", exc_info=True)
            ip_address = "App Error"

        self.ip_fetched.emit(self.ip_type, ip_address, country_code)


if PYQTGRAPH_AVAILABLE:
    class ResourceGraph(pg.PlotWidget):
        """A custom PlotWidget for displaying a scrolling resource graph."""
        def __init__(self, parent=None, title="", color='c', text_color=(221, 221, 221)):
            super().__init__(parent)
            self.setMouseEnabled(x=False, y=False)
            self.setMenuEnabled(False)
            self.getPlotItem().hideAxis('bottom')
            self.getPlotItem().hideAxis('left')
            self.setBackground(background=(40, 44, 52)) # Default to dark theme background
            self.setRange(yRange=(0, 100), padding=0)

            self.data = np.zeros(60) # 60 data points for a 1-minute history at 1s refresh
            self.curve = self.plot(self.data, pen=pg.mkPen(color, width=2))

            self.text = pg.TextItem(text="", color=text_color, anchor=(0.5, 0.5))
            self.text.setPos(30, 50) # Position it in the middle of the graph
            self.addItem(self.text)


        def update_data(self, new_value):
            """Shifts the data and adds a new value to the end."""
            self.data[:-1] = self.data[1:]
            self.data[-1] = new_value
            self.curve.setData(self.data)
            self.text.setText(f"{new_value:.0f}%")
else:
    # If pyqtgraph is not available, create a dummy widget to avoid crashing.
    class ResourceGraph(QWidget):
        def __init__(self, parent=None, title="", color='c', text_color=(221, 221, 221)):
            super().__init__(parent)
            layout = QVBoxLayout(self)
            label = QLabel("Graphs disabled\n(pyqtgraph not installed)")
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            label.setStyleSheet("color: #888;")
            layout.addWidget(label)
            self.setMinimumHeight(60)
            # Make the placeholder visible
            self.setStyleSheet("background-color: #2d313a; border: 1px solid #444;")

        def update_data(self, new_value):
            """Dummy method, does nothing."""
            pass

class SubdomainResultsDialog(QDialog):
    """A dialog to show a list of found subdomains with an export option."""
    def __init__(self, subdomains, domain, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Subdomain Scan Results for {domain}")
        self.setMinimumSize(500, 400)
        self.parent = parent # To access the export handler
        self.domain = domain # Store domain for context

        layout = QVBoxLayout(self)

        summary_label = QLabel(f"<b>Found {len(subdomains)} unique subdomains.</b>")
        layout.addWidget(summary_label)

        self.tree = QTreeWidget()
        self.tree.setColumnCount(1)
        self.tree.setHeaderLabels(["Subdomain"])
        for sub in subdomains:
            self.tree.addTopLevelItem(QTreeWidgetItem([sub]))
        self.tree.resizeColumnToContents(0)
        layout.addWidget(self.tree)

        button_layout = QHBoxLayout()
        export_button = self.parent._create_export_button(self.tree)
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(lambda: self.parent._send_to_ai_analyst("subdomain", self.tree, self.domain))
        button_layout.addWidget(export_button)
        button_layout.addWidget(analyze_button)
        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)

        layout.addLayout(button_layout)

class NmapSummaryDialog(QDialog):
    """A dialog to show a summary of Nmap scan results from XML."""
    def __init__(self, xml_data, target_context, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Nmap Scan Summary")
        self.setMinimumSize(700, 500)
        self.xml_data = xml_data
        self.target_context = target_context
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Host / Details", "Port", "Service", "Version"])
        self.tree.header().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.tree.header().setStretchLastSection(False)
        layout.addWidget(self.tree)

        self.parse_and_populate(xml_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("nmap", self.xml_data, self.target_context)
            self.accept() # Close dialog after sending

    def parse_and_populate(self, xml_data):
        if not LXML_AVAILABLE:
            self.tree.addTopLevelItem(QTreeWidgetItem(["LXML library not installed."]))
            return
        if not xml_data:
            self.tree.addTopLevelItem(QTreeWidgetItem(["No XML data to parse."]))
            return

        try:
            parser = etree.XMLParser(recover=True, no_network=True, dtd_validation=False)
            root = etree.fromstring(xml_data.encode('utf-8'), parser=parser)

            for host in root.findall('host'):
                if host.find('status').get('state') != 'up':
                    continue

                address = host.find('address').get('addr')
                hostname_elem = host.find('hostnames/hostname')
                hostname = hostname_elem.get('name') if hostname_elem is not None else ""

                host_text = f"{address} ({hostname})" if hostname else address
                host_item = QTreeWidgetItem([host_text])
                host_item.setExpanded(True)
                self.tree.addTopLevelItem(host_item)

                ports_elem = host.find('ports')
                if ports_elem is None:
                    continue

                for port in ports_elem.findall('port'):
                    if port.find('state').get('state') == 'open':
                        port_id = port.get('portid')
                        protocol = port.get('protocol')

                        service_elem = port.find('service')
                        service = service_elem.get('name', '') if service_elem is not None else ''
                        version_parts = []
                        if service_elem is not None:
                            if service_elem.get('product'): version_parts.append(service_elem.get('product'))
                            if service_elem.get('version'): version_parts.append(service_elem.get('version'))
                        version = " ".join(version_parts)

                        port_item = QTreeWidgetItem(["", f"{port_id}/{protocol}", service, version])
                        host_item.addChild(port_item)

        except Exception as e:
            logging.error(f"Failed to parse Nmap XML for summary: {e}", exc_info=True)
            self.tree.addTopLevelItem(QTreeWidgetItem(["Error parsing XML data."]))

class HttpxResultsDialog(QDialog):
    def __init__(self, json_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("httpx Probe Results")
        self.setMinimumSize(800, 500)
        self.json_data = json_data
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        # Define columns based on common httpx JSON output
        self.tree.setColumnCount(5)
        self.tree.setHeaderLabels(["URL", "Status Code", "Title", "Web Server", "Technologies"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            # The AI can analyze the raw JSON data
            self.parent.ai_assistant_tab.send_to_analyst("httpx", self.json_data, "httpx probe results")
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            # httpx outputs JSON objects separated by newlines
            results = [json.loads(line) for line in json_data.strip().split('\n') if line]
            for res in results:
                url = res.get('url', '')
                status_code = str(res.get('status_code', ''))
                title = res.get('title', '')
                web_server = res.get('webserver', '')
                tech = ", ".join(res.get('tech', []))

                item = QTreeWidgetItem([url, status_code, title, web_server, tech])
                self.tree.addTopLevelItem(item)
        except json.JSONDecodeError:
            # Handle case where output is not JSON
            item = QTreeWidgetItem(["Error parsing JSON output. Displaying raw data in console."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing httpx JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class DirsearchResultsDialog(QDialog):
    def __init__(self, json_data, target_context, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"dirsearch Results for {target_context}")
        self.setMinimumSize(800, 500)
        self.json_data = json_data
        self.target_context = target_context
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Path", "Status Code", "Content-Length", "Redirect"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("dirsearch", self.json_data, self.target_context)
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            # dirsearch report is a dictionary where keys are hostnames
            results = json.loads(json_data)
            for host, findings in results.items():
                host_item = QTreeWidgetItem([f"Host: {host}"])
                self.tree.addTopLevelItem(host_item)
                host_item.setExpanded(True)
                for finding in findings:
                    path = finding.get('path', '')
                    status = str(finding.get('status', ''))
                    length = str(finding.get('content-length', ''))
                    redirect = finding.get('redirect', '')

                    child_item = QTreeWidgetItem([path, status, length, redirect])
                    host_item.addChild(child_item)
        except json.JSONDecodeError:
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing dirsearch JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class FfufResultsDialog(QDialog):
    def __init__(self, json_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("ffuf Scan Results")
        self.setMinimumSize(800, 500)
        self.json_data = json_data
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["URL", "Status", "Length", "Words"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("ffuf", self.json_data, "ffuf scan results")
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            results = json.loads(json_data).get('results', [])
            for res in results:
                url = res.get('url', '')
                status = str(res.get('status', ''))
                length = str(res.get('length', ''))
                words = str(res.get('words', ''))

                item = QTreeWidgetItem([url, status, length, words])
                self.tree.addTopLevelItem(item)
        except (json.JSONDecodeError, AttributeError):
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing ffuf JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class NucleiResultsDialog(QDialog):
    def __init__(self, json_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Nuclei Scan Results")
        self.setMinimumSize(900, 600)
        self.json_data = json_data
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(5)
        self.tree.setHeaderLabels(["Template ID", "Name", "Severity", "Host", "Matched At"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("nuclei", self.json_data, "Nuclei scan results")
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            # Nuclei outputs JSON objects separated by newlines
            results = [json.loads(line) for line in json_data.strip().split('\n') if line]
            for res in results:
                template_id = res.get('template-id', '')
                name = res.get('info', {}).get('name', '')
                severity = res.get('info', {}).get('severity', '')
                host = res.get('host', '')
                matched_at = res.get('matched-at', '')

                item = QTreeWidgetItem([template_id, name, severity, host, matched_at])
                self.tree.addTopLevelItem(item)

                # Add extracted results as children for more detail
                if 'extracted-results' in res:
                    for i, extracted in enumerate(res['extracted-results']):
                        child_item = QTreeWidgetItem([f"  - Extracted {i+1}", str(extracted)])
                        item.addChild(child_item)

                item.setExpanded(True)

        except json.JSONDecodeError:
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing Nuclei JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class TruffleHogResultsDialog(QDialog):
    def __init__(self, json_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("TruffleHog Scan Results")
        self.setMinimumSize(900, 600)
        self.json_data = json_data
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Detector", "Decoder", "File", "Raw Secret"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("trufflehog", self.json_data, "TruffleHog scan results")
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            # TruffleHog outputs JSON objects separated by newlines
            results = [json.loads(line) for line in json_data.strip().split('\n') if line]
            for res in results:
                detector = res.get('DetectorType', '')
                decoder = res.get('DecoderType', '')
                file = res.get('File', '')
                raw = res.get('Raw', '')

                item = QTreeWidgetItem([detector, decoder, file, raw])
                self.tree.addTopLevelItem(item)
        except json.JSONDecodeError:
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing TruffleHog JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class Enum4LinuxNGResultsDialog(QDialog):
    def __init__(self, json_data, target_context, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"enum4linux-ng Results for {target_context}")
        self.setMinimumSize(800, 600)
        self.json_data = json_data
        self.target_context = target_context
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(2)
        self.tree.setHeaderLabels(["Finding", "Details"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("enum4linux-ng", self.json_data, self.target_context)
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            results = json.loads(json_data)
            # The JSON is a list of dictionaries, each representing a finding
            for finding in results:
                method = finding.get('method', 'N/A')
                item = QTreeWidgetItem([method])
                self.tree.addTopLevelItem(item)

                # Add all other keys as children
                for key, value in finding.items():
                    if key != 'method':
                        child_item = QTreeWidgetItem([f"  {key}", str(value)])
                        item.addChild(child_item)
                item.setExpanded(True)
        except json.JSONDecodeError:
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing enum4linux-ng JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class CveTableModel(QAbstractTableModel):
    """A table model for displaying CVE data from the database with pagination."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.cves = []
        self.headers = ["CVE ID", "Description", "CVSSv3 Score", "Published Date"]

    def rowCount(self, parent):
        return len(self.cves)

    def columnCount(self, parent):
        return len(self.headers)

    def data(self, index, role):
        if not index.isValid() or role != Qt.ItemDataRole.DisplayRole:
            return None

        row_data = self.cves[index.row()]
        column = index.column()

        if column == 0:
            return row_data['cve_id']
        elif column == 1:
            return row_data['description']
        elif column == 2:
            return str(row_data['cvss_v3_score'] or 'N/A')
        elif column == 3:
            return row_data['published_date']
        return None

    def headerData(self, section, orientation, role):
        if role == Qt.ItemDataRole.DisplayRole and orientation == Qt.Orientation.Horizontal:
            return self.headers[section]
        return None

    def load_data(self, filter_text, sort_column, sort_order, page, page_size, cvss_min=None, cvss_max=None, date_from=None, date_to=None):
        self.beginResetModel()
        self.cves = database.get_cve_data(filter_text, sort_column, sort_order, page, page_size, cvss_min, cvss_max, date_from, date_to)
        self.endResetModel()

class CveViewerWidget(QWidget):
    """A widget for viewing, searching, and paginating the offline CVE database."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.current_page = 0
        self.page_size = 20
        self.total_records = 0
        self.filter_text = ""

        main_layout = QVBoxLayout(self)

        # --- Top Search Bar ---
        top_controls_layout = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search CVE descriptions...")
        self.search_input.returnPressed.connect(self.search_cves)
        top_controls_layout.addWidget(self.search_input)

        search_btn = QPushButton("Search")
        search_btn.clicked.connect(self.search_cves)
        top_controls_layout.addWidget(search_btn)
        main_layout.addLayout(top_controls_layout)

        # --- Advanced Filters ---
        self.filters_box = QGroupBox("Advanced Filters")
        self.filters_box.setCheckable(True)
        self.filters_box.setChecked(False)
        filters_layout = QFormLayout(self.filters_box)

        # CVSS Score Filter
        self.cvss_score_combo = QComboBox()
        self.cvss_score_combo.addItems([
            "Any",
            "Critical (9.0-10.0)",
            "High (7.0-8.9)",
            "Medium (4.0-6.9)",
            "Low (0.1-3.9)"
        ])
        filters_layout.addRow("Minimum CVSSv3 Score:", self.cvss_score_combo)

        # Date Range Filter
        date_layout = QHBoxLayout()
        self.date_from_edit = QDateEdit()
        self.date_from_edit.setCalendarPopup(True)
        self.date_from_edit.setDate(QDate.currentDate().addYears(-5)) # Default to 5 years ago
        self.date_to_edit = QDateEdit()
        self.date_to_edit.setCalendarPopup(True)
        self.date_to_edit.setDate(QDate.currentDate())
        date_layout.addWidget(QLabel("From:"))
        date_layout.addWidget(self.date_from_edit)
        date_layout.addWidget(QLabel("To:"))
        date_layout.addWidget(self.date_to_edit)
        filters_layout.addRow("Published Date:", date_layout)

        main_layout.addWidget(self.filters_box)

        # Connect signals for advanced filters
        self.filters_box.toggled.connect(self.search_cves)
        self.cvss_score_combo.currentTextChanged.connect(self.search_cves)
        self.date_from_edit.dateChanged.connect(self.search_cves)
        self.date_to_edit.dateChanged.connect(self.search_cves)

        # --- Table View ---
        self.table_view = QTableView()
        self.table_model = CveTableModel(self)
        self.table_view.setModel(self.table_model)
        self.table_view.setSortingEnabled(True)
        self.table_view.horizontalHeader().setSortIndicator(3, Qt.SortOrder.DescendingOrder) # Default sort by date
        self.table_view.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.table_view.horizontalHeader().setStretchLastSection(True)
        self.table_view.resizeColumnsToContents()
        self.table_view.horizontalHeader().sortIndicatorChanged.connect(self.sort_changed)
        main_layout.addWidget(self.table_view)

        # --- Pagination ---
        pagination_layout = QHBoxLayout()
        self.prev_btn = QPushButton("<< Previous")
        self.prev_btn.clicked.connect(self.prev_page)
        self.next_btn = QPushButton("Next >>")
        self.next_btn.clicked.connect(self.next_page)
        self.page_label = QLabel("Page 1 / 1")
        pagination_layout.addWidget(self.prev_btn)
        pagination_layout.addStretch()
        pagination_layout.addWidget(self.page_label)
        pagination_layout.addStretch()
        pagination_layout.addWidget(self.next_btn)
        main_layout.addLayout(pagination_layout)

        self.update_data()

    def update_data(self):
        sort_column_index = self.table_view.horizontalHeader().sortIndicatorSection()
        sort_order = self.table_view.horizontalHeader().sortIndicatorOrder()
        sort_column = self.table_model.headers[sort_column_index].lower().replace(" ", "_")
        sort_order_str = "DESC" if sort_order == Qt.SortOrder.DescendingOrder else "ASC"

        # Get advanced filter values
        cvss_min, cvss_max = None, None
        date_from, date_to = None, None

        if self.filters_box.isChecked():
            cvss_text = self.cvss_score_combo.currentText()
            if "Critical" in cvss_text: cvss_min, cvss_max = 9.0, 10.0
            elif "High" in cvss_text: cvss_min, cvss_max = 7.0, 8.9
            elif "Medium" in cvss_text: cvss_min, cvss_max = 4.0, 6.9
            elif "Low" in cvss_text: cvss_min, cvss_max = 0.1, 3.9

            date_from = self.date_from_edit.date().toString("yyyy-MM-dd")
            date_to = self.date_to_edit.date().toString("yyyy-MM-dd")

        self.total_records = database.get_cve_total_count(self.filter_text, cvss_min, cvss_max, date_from, date_to)
        self.table_model.load_data(self.filter_text, sort_column, sort_order_str, self.current_page, self.page_size, cvss_min, cvss_max, date_from, date_to)
        self.update_pagination_controls()

    def update_pagination_controls(self):
        total_pages = (self.total_records + self.page_size - 1) // self.page_size
        self.page_label.setText(f"Page {self.current_page + 1} / {max(1, total_pages)}")
        self.prev_btn.setEnabled(self.current_page > 0)
        self.next_btn.setEnabled((self.current_page + 1) * self.page_size < self.total_records)

    def search_cves(self):
        self.filter_text = self.search_input.text()
        self.current_page = 0
        self.update_data()

    def sort_changed(self, logicalIndex, order):
        self.current_page = 0
        self.update_data()

    def prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self.update_data()

    def next_page(self):
        if (self.current_page + 1) * self.page_size < self.total_records:
            self.current_page += 1
            self.update_data()

class SnortGuideDialog(QDialog):
    """A dialog that shows a detailed, user-friendly guide for installing Snort."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Snort Installation Guide")
        self.setMinimumSize(800, 600)

        layout = QVBoxLayout(self)
        text_browser = QTextBrowser()
        text_browser.setOpenExternalLinks(True)

        guide_html = """
        <html><head><style>
            body { font-family: sans-serif; line-height: 1.6; }
            h1, h2, h3 { color: #4a90e2; }
            code { background-color: #2d313a; padding: 2px 5px; border-radius: 4px; font-family: "Courier New", monospace; }
            a { color: #8be9fd; }
            .note { border-left: 5px solid #ffcc00; padding: 1em; background-color: #3a3a3a; }
        </style></head><body>
            <h1>Snort Installation Guide</h1>
            <p>Snort is a powerful Intrusion Detection System, but its setup can be complex. This guide provides a simplified walkthrough for common operating systems.</p>

            <h2>Linux (Debian/Ubuntu)</h2>
            <p>This is the most straightforward method.</p>
            <ol>
                <li><b>Install Snort:</b><br>
                    <code>sudo apt update && sudo apt install snort -y</code>
                </li>
                <li><b>Configure Network Interface:</b> Open the main configuration file with <code>sudo nano /etc/snort/snort.conf</code>. Find the line that starts with <code>ipvar HOME_NET</code> and change it to match your local network (e.g., <code>ipvar HOME_NET 192.168.1.0/24</code>).</li>
                <li><b>Validate Configuration:</b> Run the following command to test your configuration file for errors:<br>
                    <code>sudo snort -T -c /etc/snort/snort.conf</code><br>
                    If you see "Snort successfully validated the configuration!", you are ready.
                </li>
                <li><b>Set Paths in Zurvan:</b> In the Snort tab, set the "Config File" path to <code>/etc/snort/snort.conf</code> and the "Rules File" to <code>/etc/snort/rules/local.rules</code>.</li>
            </ol>

            <h2>Windows</h2>
            <p>Installing on Windows requires more manual steps.</p>
            <ol>
                <li><b>Download Snort:</b> Go to the <a href="https://www.snort.org/downloads">Snort downloads page</a> and download the installer for Windows. You will need to create a free account.</li>
                <li><b>Install Npcap:</b> Download and install <a href="https://npcap.com/#download">Npcap</a>, which is required for packet capture on Windows. During installation, be sure to check the box for "Install Npcap in WinPcap API-compatible Mode".</li>
                <li><b>Run the Snort Installer:</b> Run the downloaded Snort installer and accept the default settings. It will typically install to <code>C:\\Snort</code>.</li>
                <li><b>Download Rules:</b> From the Snort downloads page, get the "Community Rules" (<code>community-rules.tar.gz</code>).</li>
                <li><b>Extract Rules:</b> Using a tool like 7-Zip, extract the contents of the <code>community-rules.tar.gz</code> file. Copy the files from the extracted <code>community-rules</code> folder into <code>C:\\Snort\\rules</code>.</li>
                <li><b>Configure snort.conf:</b>
                    <ul>
                        <li>Navigate to <code>C:\\Snort\\etc</code> and open <code>snort.conf</code> in a text editor like Notepad++.</li>
                        <li>Find <code>ipvar HOME_NET any</code> and change it to your network (e.g., <code>ipvar HOME_NET 192.168.1.0/24</code>).</li>
                        <li>Find the rule path section. You must change the paths from the Linux format (e.g., <code>var RULE_PATH ../rules</code>) to the correct Windows format. Replace all rule path lines with the following (adjusting <code>C:\\Snort</code> if you installed elsewhere):<br>
<pre><code>var RULE_PATH C:\\Snort\\rules
var SO_RULE_PATH C:\\Snort\\so_rules
var PREPROC_RULE_PATH C:\\Snort\\preproc_rules
var WHITE_LIST_PATH C:\\Snort\\rules
var BLACK_LIST_PATH C:\\Snort\\rules</code></pre>
                        </li>
                        <li>Find and comment out the line <code>include $RULE_PATH/protocol-voip.rules</code> by adding a '#' at the beginning. This rule is known to cause issues on Windows.</li>
                        <li>Create a new file named <code>local.rules</code> inside <code>C:\\Snort\\rules</code>.</li>
                    </ul>
                </li>
                 <li><b>Set Paths in Zurvan:</b> In the Snort tab, set the "Config File" path to <code>C:\\Snort\\etc\\snort.conf</code> and the "Rules File" to <code>C:\\Snort\\rules\\local.rules</code>.</li>
            </ol>
            <p class="note"><b>Note:</b> You may need to run Zurvan as an Administrator on Windows for Snort to have the necessary permissions to capture network traffic.</p>
        </body></html>
        """
        text_browser.setHtml(guide_html)
        layout.addWidget(text_browser)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        layout.addWidget(ok_button, 0, Qt.AlignmentFlag.AlignRight)

class TorGuideDialog(QDialog):
    """A dialog that shows a detailed, user-friendly guide for installing Tor."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Tor Installation Guide")
        self.setMinimumSize(800, 600)

        layout = QVBoxLayout(self)
        text_browser = QTextBrowser()
        text_browser.setOpenExternalLinks(True)

        guide_html = """
        <html><head><style>
            body { font-family: sans-serif; line-height: 1.6; }
            h1, h2, h3 { color: #4a90e2; }
            code { background-color: #2d313a; padding: 2px 5px; border-radius: 4px; font-family: "Courier New", monospace; }
            a { color: #8be9fd; }
            .note { border-left: 5px solid #ffcc00; padding: 1em; background-color: #3a3a3a; }
        </style></head><body>
            <h1>Tor Proxy Installation Guide</h1>
            <p>To use this feature, you need to have Tor running and listening on the default SOCKS proxy port (9050). The easiest way to do this is by installing and running the <b>Tor Browser</b>.</p>

            <h2>1. Download and Install Tor Browser</h2>
            <p>Visit the official Tor Project website to download the browser for your operating system:</p>
            <p><a href="https://www.torproject.org/download/">https://www.torproject.org/download/</a></p>
            <ul>
                <li><b>Windows/macOS:</b> Run the downloaded installer and follow the on-screen instructions.</li>
                <li><b>Linux:</b> Extract the downloaded archive and run the <code>start-tor-browser.desktop</code> script.</li>
            </ul>

            <h2>2. Run Tor Browser</h2>
            <p>Simply start the Tor Browser. Once it successfully connects to the Tor network, it will automatically start a SOCKS proxy on <code>127.0.0.1:9050</code> in the background. <b>You must leave Tor Browser running while you use this feature in Zurvan.</b></p>

            <h2>3. (Optional) For Advanced Linux Users: Install Tor Service</h2>
            <p>If you don't want to run the full Tor Browser, you can install the Tor service:</p>
            <ul>
                <li><b>Debian/Ubuntu:</b> <code>sudo apt update && sudo apt install tor</code></li>
                <li><b>Fedora/CentOS:</b> <code>sudo dnf install tor</code></li>
            </ul>
            <p>After installation, the Tor service should start automatically. You can check its status with <code>sudo systemctl status tor</code>.</p>

            <p class="note"><b>Note:</b> For most tools to be routed through Tor, Zurvan uses the <code>proxychains-ng</code> command. Please ensure it is installed on your system (e.g., <code>sudo apt install proxychains-ng</code>).</p>
        </body></html>
        """
        text_browser.setHtml(guide_html)
        layout.addWidget(text_browser)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        layout.addWidget(ok_button, 0, Qt.AlignmentFlag.AlignRight)

class DnsReconResultsDialog(QDialog):
    def __init__(self, json_data, target_context, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"dnsrecon Results for {target_context}")
        self.setMinimumSize(800, 600)
        self.json_data = json_data
        self.target_context = target_context
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Type", "Target", "Address", "Name"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("dnsrecon", self.json_data, self.target_context)
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            results = json.loads(json_data)
            # The JSON is a list of dictionaries
            for res in results:
                rec_type = res.get('type', 'N/A')
                target = res.get('target', 'N/A')
                address = res.get('address', 'N/A')
                name = res.get('name', 'N/A')

                item = QTreeWidgetItem([rec_type, target, address, name])
                self.tree.addTopLevelItem(item)
        except json.JSONDecodeError:
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing dnsrecon JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class SherlockResultsDialog(QDialog):
    def __init__(self, csv_data, target_context, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Sherlock Results for {target_context}")
        self.setMinimumSize(800, 600)
        self.csv_data = csv_data
        self.target_context = target_context
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Username", "Service Name", "URL", "Status"])
        layout.addWidget(self.tree)

        self.parse_and_populate(csv_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("sherlock", self.csv_data, self.target_context)
            self.accept()

    def parse_and_populate(self, csv_data):
        try:
            # Use Python's built-in csv module to parse the data
            reader = csv.reader(csv_data.strip().splitlines())
            header = next(reader) # Skip header row
            for row in reader:
                # Assuming standard sherlock csv format: username,name,url,status
                if len(row) >= 4:
                    item = QTreeWidgetItem(row)
                    self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing Sherlock CSV: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

# --- Main Application ---
class Zurvan(QMainWindow):
    """The main application window, holding all UI elements and logic."""
    cve_import_status_changed = pyqtSignal(str)
    cve_import_buttons_enabled = pyqtSignal(bool)

    def __init__(self):
        """Initializes the main window, UI components, and internal state."""
        super().__init__()
        self.setWindowTitle("Zurvan - Comprehensive AI-Powered Security Platform")

        # Define script directory as an instance variable for consistent path resolution
        self.script_dir = os.path.dirname(os.path.realpath(__file__))

        # Set window icon using the new helper method
        self.setWindowIcon(QIcon(self.icon_path("Zurvan-mono.png")))
        # Use dynamic layout instead of fixed geometry, start maximized for responsiveness
        self.showMaximized()

        self.current_user = None
        self.packets_data = []; self.sniffer_thread = None; self.channel_hopper = None
        self.packet_layers = []; self.current_field_widgets = []; self.tcp_flag_vars = {}
        self.tool_results_queue = Queue()
        self.is_tool_running = False
        self.loaded_flood_packet = None
        self.found_networks = {}
        self.active_threads = []
        self.thread_finish_lock = Lock()
        self.finished_thread_count = 0
        self.tool_stop_event = Event()
        self.arp_spoof_current_victim = None
        self.arp_spoof_current_target = None
        self.resource_monitor_thread = None
        self.nmap_last_xml = None
        self.nmap_xml_temp_file = None
        self.aircrack_thread = None
        self.ps_thread_lock = Lock()
        self.ps_finished_threads = 0
        self.bf_ssid_list = []
        self.krack_thread = None
        self.sniffer_packet_buffer = []
        self.sniffer_buffer_lock = Lock()
        self.super_scan_active = False
        self.lab_test_chain = []
        self.threat_intel_loaded = False
        self.history_loaded = False
        self.recent_threats_data = []
        self.threats_current_page = 0
        self.threats_items_per_page = 30
        self.reporting_cve_manager = None
        self.threat_intel_cve_manager = None
        self.auto_lock_timer = None
        self.activity_monitor = None
        self.app_lock_timeout_minutes = 15 # Default
        self.journal_refresh_timer = None
        self.theharvester_controls = {}
        self.v2ray_process = None
        self.l2tp_process = None
        self.l2tp_config_files = {}
        self.wireguard_process = None
        self.openvpn_process = None
        self.sudo_cancel_handlers = {}
        self.real_ip_address = None
        self.proxy_ip_address = None
        self.proxy_rotation_timer = None
        self.current_proxy_index = -1
        self.whatweb_process = None
        self.is_test_mode = False
        self.doh_resolver_url = None
        # self.tool_config_widgets is no longer used.

        self.nmap_script_presets = {
            # --- Comprehensive Scans ---
            "Super Complete Scan": ("vulners,vuln,exploit,dos,http-vuln-cve*,smb-vuln-*,http-enum,smb-os-discovery,smb-enum-shares,dns-brute,ssl-heartbleed,http-shellshock", "", "A highly intrusive and comprehensive scan combining vulnerability, exploit, CVE, and enumeration scripts."),
            "Intrusive Vulnerability Scan": ("vulners,vuln", "", "Runs all scripts in the 'vuln' category and the Vulners script to check for known CVEs."),
            "Full Service Enumeration": ("-sV,-sC,banner,http-enum,smb-enum-shares,smb-enum-users,dns-brute", "", "Intense scan combining version detection, default scripts, and enumeration for HTTP, SMB, and DNS."),
            "Safe Network Discovery": ("safe,discovery,broadcast-dns-service-discovery,dns-brute,hostmap-bfk", "", "A non-intrusive scan to discover hosts, services, and basic information without causing disruption."),

            # --- Web & HTTP Focused ---
            "Web - Full Enumeration": ("http-enum,http-title,http-headers,http-sitemap-generator,http-robots.txt,http-git", "", "Crawls a web server to discover directories, titles, headers, sitemaps, robots.txt, and exposed git repos."),
            "Web - Vulnerability Scan": ("http-vuln-cve*,http-shellshock,http-sql-injection,http-xssed,http-csrf", "", "Scans a web server for common vulnerabilities including CVEs, Shellshock, SQLi, XSS, and CSRF."),
            "Web - Exposed Panels & Files": ("http-config-backup,http-wp-backup,http-frontpage-login,http-drupal-enum-users", "", "Searches for exposed configuration files, backups, and administrative panels."),

            # --- SMB (Windows Shares) Focused ---
            "SMB - Full Audit": ("smb-os-discovery,smb-enum-shares,smb-enum-users,smb-enum-processes,smb-security-mode,smb-protocols,smb-vuln-*", "", "Performs a comprehensive security audit of SMB services, including enumeration and vulnerability checks."),
            "SMB - Basic Enumeration": ("smb-os-discovery,smb-enum-shares,smb-enum-users", "", "Gathers basic information about an SMB server, including OS, shares, and users."),
            "SMB - Vulnerability Check (EternalBlue, etc.)": ("smb-vuln-ms17-010,smb-vuln-conficker,smb-vuln-ms08-067,smb-double-pulsar-backdoor", "", "Checks for major historical SMB vulnerabilities like EternalBlue, Conficker, and MS08-067."),

            # --- SSL/TLS Focused ---
            "SSL/TLS - Security Audit": ("ssl-cert,ssl-enum-ciphers,ssl-heartbleed,ssl-poodle,sslv2-drown,tls-ticketbleed", "", "Audits SSL/TLS configuration for weak ciphers and major vulnerabilities like Heartbleed, POODLE, and DROWN."),

            # --- Service & Protocol Specific ---
            "FTP - Security Check": ("ftp-anon,ftp-vsftpd-backdoor,ftp-vuln-cve2010-4221", "", "Checks for anonymous FTP login and known FTP server vulnerabilities."),
            "SMTP - User Enumeration": ("smtp-enum-users,smtp-commands", "", "Attempts to enumerate valid users on an SMTP server and discover supported commands."),
            "SNMP - Brute Force": ("snmp-brute", "", "Attempts to brute-force SNMP community strings to gain access."),
            "SSH - Security Info": ("ssh-auth-methods,ssh2-enum-algos,ssh-hostkey", "", "Gathers information about an SSH server's supported authentication methods, algorithms, and its host key."),
            "RDP - Vulnerability Check": ("rdp-vuln-ms12-020,rdp-enum-encryption", "", "Checks for the critical MS12-020 RDP vulnerability and enumerates encryption levels."),
            "Database - Info Gathering": ("mysql-info,mongodb-info,redis-info", "", "Gathers basic information and statistics from common database servers."),

            # --- Miscellaneous & Recon ---
            "DNS - Subdomain Discovery": ("dns-brute,dns-zone-transfer", "", "Attempts to discover subdomains via brute-forcing and zone transfer requests."),
            "Firewall - Evasion Test": ("firewall-bypass", "", "Attempts to discover firewall rules by using common application protocols (FTP, etc)."),
            "General Reconnaissance": ("whois-domain,traceroute-geolocation,banner", "", "Performs basic reconnaissance including WHOIS lookups, traceroutes, and banner grabbing."),
        }

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QVBoxLayout(self.central_widget)

        self._create_resource_bar()
        self.menu_bar = QMenuBar(self)
        self.setMenuBar(self.menu_bar)
        # The menu bar will be populated after login by calling _update_menu_bar()
        self._create_status_bar()

        # Create the header bar first, as the main tabs may need to connect to widgets inside it (e.g., theme_combo).
        self._create_header_bar()
        self._create_main_tabs()

        # Config widgets are now created inside their respective tool tabs.

        self._create_log_panel(); self._setup_logging()
        self.tab_widget.currentChanged.connect(self._on_main_tab_changed)

        # --- Connect Signals for Synced Widgets ---
        if self.reporting_cve_manager:
            self.reporting_cve_manager.start_import.connect(self._start_cve_import)
            self.cve_import_status_changed.connect(self.reporting_cve_manager.set_status)
            self.cve_import_buttons_enabled.connect(self.reporting_cve_manager.set_buttons_enabled)

        if self.threat_intel_cve_manager:
            self.threat_intel_cve_manager.start_import.connect(self._start_cve_import)
            self.cve_import_status_changed.connect(self.threat_intel_cve_manager.set_status)
            self.cve_import_buttons_enabled.connect(self.threat_intel_cve_manager.set_buttons_enabled)

        self._setup_result_handlers()
        self.results_processor = QTimer(self); self.results_processor.timeout.connect(self._process_tool_results); self.results_processor.start(100)

        # Setup timer for batching sniffer UI updates
        self.sniffer_ui_update_timer = QTimer(self)
        self.sniffer_ui_update_timer.timeout.connect(self._update_sniffer_display)
        self.sniffer_ui_update_timer.start(500) # Update every 500ms

        # Setup timer for the clock
        self.clock_timer = QTimer(self)
        self.clock_timer.timeout.connect(self._update_clock)
        self.clock_timer.start(1000) # 1000 ms = 1 second
        self._update_clock() # Initial call to show time immediately

        # Start the resource monitor
        self.resource_monitor_thread = ResourceMonitorThread(self)
        self.resource_monitor_thread.stats_updated.connect(self._update_resource_stats)
        self.resource_monitor_thread.start()

        # Setup timer for public IP display
        self.ip_check_timer = QTimer(self)
        self.ip_check_timer.timeout.connect(self._update_public_ips)
        self._update_public_ips() # Initial call
        self._handle_ip_refresh_change("60s")

        self._update_tool_targets() # Initial population after all widgets are created
        self._set_user_avatar()
        self._load_v2ray_servers()
        # Auto-detect cores
        self._auto_detect_cores()
        self.v2ray_core_combo.currentTextChanged.connect(self._auto_detect_cores)
        logging.info("Zurvan application initialized.")

    def _handle_ip_refresh_change(self, text):
        """Updates the public IP refresh timer interval."""
        if text == "Off":
            self.ip_check_timer.stop()
            logging.info("Public IP auto-refresh turned off.")
        else:
            interval_ms = int(text.replace('s', '')) * 1000
            self.ip_check_timer.start(interval_ms)
            logging.info(f"Public IP refresh interval set to {text}.")

    def _auto_detect_cores(self, selected_core=None):
        """Auto-detects the path to the selected core executable."""
        if not selected_core:
            selected_core = self.v2ray_core_combo.currentText()

        # Step 1: Check system PATH
        core_executable = shutil.which(selected_core) or shutil.which(f"{selected_core}.exe")

        # Step 2: If not found, check local tools directory relative to the script
        local_path = os.path.join(self.script_dir, "tools", selected_core, selected_core)
        if sys.platform == "win32":
            local_path += ".exe"

        if os.path.isfile(local_path):
            core_executable = os.path.abspath(local_path)

        if core_executable:
            self.v2ray_core_edit.setText(core_executable)
        else:
            self.v2ray_core_edit.clear()  # Clear if not found

    def icon_path(self, icon_name):
        """Constructs an absolute path to an icon file."""
        return os.path.join(self.script_dir, "icons", icon_name)

    def _get_tool_path(self, tool_name, relative_path=None):
        """
        Resolves the path to an external tool.
        1. Checks system PATH (using shutil.which).
        2. Checks the local 'tools/' directory (relative to the script).
        """
        # 1. Check system PATH first
        system_path = shutil.which(tool_name)
        if system_path:
            return system_path

        # 2. Check local tools directory if a relative path is provided
        if relative_path:
            local_path = os.path.join(self.script_dir, "tools", relative_path)
            if os.path.exists(local_path):
                # Ensure it's executable if it's a file
                if os.path.isfile(local_path) and not os.access(local_path, os.X_OK):
                     try:
                         os.chmod(local_path, 0o755)
                     except Exception as e:
                         logging.warning(f"Could not set execute permission on {local_path}: {e}")
                return os.path.abspath(local_path)

        return None

    def _find_zap_executable(self):
        """Finds a usable ZAP launcher from PATH or common macOS install locations."""
        candidates = [
            shutil.which("zap.sh"),
            shutil.which("zap"),
            "/Applications/ZAP.app/Contents/MacOS/zap.sh",
            "/Applications/OWASP ZAP.app/Contents/MacOS/zap.sh",
        ]

        candidates.extend(
            glob.glob("/opt/homebrew/Caskroom/zap/*/ZAP.app/Contents/MacOS/zap.sh")
        )

        for candidate in candidates:
            if candidate and os.path.exists(candidate):
                return candidate

        return None

    def _set_user_avatar(self):
        """Sets the user profile button icon based on the user's avatar."""
        if not self.current_user:
            return

        avatar_data = self.current_user['avatar']
        if avatar_data:
            pixmap = QPixmap()
            pixmap.loadFromData(avatar_data)
            # Create a circular pixmap
            circular_pixmap = self.create_circular_pixmap(pixmap, self.user_profile_button.iconSize())
            self.user_profile_button.setIcon(QIcon(circular_pixmap))
        else:
            # Fallback to a default icon if no avatar is set
            self.user_profile_button.setIcon(QIcon(self.icon_path("Zurvan-mono.png")))

    def create_circular_pixmap(self, source_pixmap, size):
        """Creates a circular pixmap from a source pixmap."""
        if source_pixmap.isNull():
            return QPixmap()

        # Scale the pixmap to the target size, keeping aspect ratio
        scaled_pixmap = source_pixmap.scaled(size, Qt.AspectRatioMode.KeepAspectRatioByExpanding, Qt.TransformationMode.SmoothTransformation)

        # Create a new pixmap with a transparent background
        circular_pixmap = QPixmap(size)
        circular_pixmap.fill(Qt.GlobalColor.transparent)

        # Use QPainter to draw the scaled pixmap in a circle
        painter = QPainter(circular_pixmap)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)

        # Create a circular clipping path
        path = QPainterPath()
        path.addEllipse(0, 0, size.width(), size.height())
        painter.setClipPath(path)

        # Draw the scaled pixmap onto the circular pixmap
        # Center the image
        x = (size.width() - scaled_pixmap.width()) / 2
        y = (size.height() - scaled_pixmap.height()) / 2
        painter.drawPixmap(int(x), int(y), scaled_pixmap)

        painter.end()

        return circular_pixmap

    def _show_user_menu(self):
        """Shows a context menu for the user profile button."""
        menu = QMenu(self)
        profile_action = menu.addAction("Profile...")
        logout_action = menu.addAction("Logout")

        action = menu.exec(self.user_profile_button.mapToGlobal(QPoint(0, self.user_profile_button.height())))

        if action == profile_action:
            self._show_user_profile()
        elif action == logout_action:
            self._logout()

    def _show_user_profile(self):
        """Opens the user profile dialog."""
        if not self.current_user:
            return

        # Re-fetch user data to ensure it's up-to-date
        self.current_user = database.get_user_by_id(self.current_user['id'])
        dialog = UserProfileDialog(self.current_user, self)
        if dialog.exec():
            # If changes were saved, re-fetch user data and update avatar
            self.current_user = database.get_user_by_id(self.current_user['id'])
            self._set_user_avatar()

    def _logout(self):
        """Logs the current user out and shows the login screen."""
        reply = QMessageBox.question(self, 'Logout Confirmation',
                                     "Are you sure you want to log out?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='Authentication',
                    action='User Logout',
                    target=self.current_user['username']
                )
            # Stop background threads before restarting
            logging.info("User confirmed logout. Stopping background threads.")
            if self.sniffer_thread and self.sniffer_thread.isRunning(): self.sniffer_thread.stop()
            if self.channel_hopper and self.channel_hopper.isRunning(): self.channel_hopper.stop()
            if self.resource_monitor_thread and self.resource_monitor_thread.isRunning():
                self.resource_monitor_thread.stop()

            # A bit of a hacky way to restart, but effective for a desktop app
            python = sys.executable
            os.execl(python, python, *sys.argv)

    def _update_menu_bar(self):
        """Creates or updates the main menu bar based on the current user's role."""
        self.menu_bar.clear()

        # --- File Menu ---
        file_menu = self.menu_bar.addMenu("&File")
        file_menu.addAction("&Save Captured Packets", self.save_packets)
        file_menu.addAction("&Load Packets from File", self.load_packets)
        file_menu.addSeparator()
        file_menu.addAction("&Exit", self.close)

        # --- Admin Menu (Conditional) ---
        # Use .get() for safer dictionary access, in case current_user is None
        logging.info(f"Updating menu bar for user: {self.current_user}")
        if self.current_user and self.current_user['is_admin']:
            admin_menu = self.menu_bar.addMenu("&Admin")
            admin_menu.addAction(QIcon(self.icon_path("Zurvan-mono.png")), "Admin Panel...", self._show_admin_panel)

        # --- Help Menu ---
        help_menu = self.menu_bar.addMenu("&Help")
        help_menu.addAction("Community Forums & Tools...", self._show_community_dialog)
        help_menu.addSeparator()
        help_menu.addAction("&About Zurvan", self._show_about_dialog)
        help_menu.addSeparator()
        help_menu.addAction("&AI Settings...", self._show_ai_settings_dialog)
        help_menu.addAction("AI Guide", self._show_ai_guide_dialog)

    def _show_community_dialog(self):
        """Shows a dialog with links to useful cybersecurity communities and resources."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Community Forums & Resources")
        dialog.setMinimumSize(600, 500)
        layout = QVBoxLayout(dialog)

        text_browser = QTextBrowser()
        text_browser.setOpenExternalLinks(True)

        html_content = """
        <html><head><style>
            body { font-family: sans-serif; line-height: 1.5; }
            h2 { color: #4a90e2; border-bottom: 1px solid #444; padding-bottom: 5px; }
            ul { list-style-type: none; padding-left: 0; }
            li { margin-bottom: 10px; }
            a { color: #8be9fd; text-decoration: none; }
            a:hover { text-decoration: underline; }
            b { color: #e0e0e0; }
        </style></head><body>
            <h1>Cybersecurity Communities & Resources</h1>

            <h2>News & Analysis</h2>
            <ul>
                <li><a href="https://thehackernews.com/"><b>The Hacker News:</b></a> Daily source for cybersecurity news and analysis.</li>
                <li><a href="https://www.darkreading.com/"><b>Dark Reading:</b></a> News and commentary for security professionals.</li>
                <li><a href="https://krebsonsecurity.com/"><b>Krebs on Security:</b></a> In-depth security investigations.</li>
                <li><a href="https://www.wired.com/category/security/"><b>WIRED Security:</b></a> Mainstream tech news with a security focus.</li>
            </ul>

            <h2>Forums & Discussion</h2>
            <ul>
                <li><a href="https://www.reddit.com/r/netsec/"><b>r/netsec (Reddit):</b></a> Large community for technical security discussions.</li>
                <li><a href="https://www.reddit.com/r/blueteamsec/"><b>r/blueteamsec (Reddit):</b></a> Defensive security and threat intelligence.</li>
                <li><a href="https://0x00sec.org/"><b>0x00sec:</b></a> Forum for hacking, reverse engineering, and malware analysis.</li>
                <li><a href="https://stackoverflow.com/questions/tagged/security"><b>Stack Overflow (Security):</b></a> Q&A for security programming questions.</li>
            </ul>

            <h2>Vulnerability Databases & Exploits</h2>
            <ul>
                <li><a href="https://nvd.nist.gov/"><b>NVD:</b></a> National Vulnerability Database.</li>
                <li><a href="https://www.exploit-db.com/"><b>Exploit-DB:</b></a> Archive of public exploits and shellcode.</li>
                <li><a href="https://packetstormsecurity.com/"><b>Packet Storm:</b></a> Up-to-date information about new vulnerabilities.</li>
                <li><a href="https://github.com/trickest/cve"><b>Trickest CVE Repository:</b></a> PoCs for recently published CVEs.</li>
            </ul>

            <h2>Learning & Resources</h2>
            <ul>
                <li><a href="https://tryhackme.com/"><b>TryHackMe:</b></a> Gamified platform for learning cybersecurity.</li>
                <li><a href="https://www.hackthebox.com/"><b>Hack The Box:</b></a> Online platform to test and advance your pentesting skills.</li>
                <li><a href="https://book.hacktricks.xyz/"><b>HackTricks:</b></a> A comprehensive repository of pentesting tricks and techniques.</li>
                <li><a href="https://owasp.org/"><b>OWASP:</b></a> The Open Web Application Security Project.</li>
            </ul>
        </body></html>
        """
        text_browser.setHtml(html_content)
        layout.addWidget(text_browser)

        ok_button = QPushButton("Close")
        ok_button.clicked.connect(dialog.accept)
        layout.addWidget(ok_button, 0, Qt.AlignmentFlag.AlignRight)

        dialog.exec()

    def _show_admin_panel(self):
        """Shows the admin panel dialog."""
        admin_dialog = AdminPanelDialog(self)
        admin_dialog.exec()

    def _show_ai_settings_dialog(self):
        """Shows the AI settings dialog."""
        dialog = AISettingsDialog(self)
        dialog.exec()

    def _show_ai_guide_dialog(self):
        """Shows the AI features user guide."""
        dialog = AIGuideDialog(self)
        dialog.exec()

    def _show_snort_guide(self):
        """Shows the Snort installation guide dialog."""
        dialog = SnortGuideDialog(self)
        dialog.exec()

    def _check_tor_proxy(self):
        """Checks if a SOCKS proxy is available on the default Tor port."""
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.settimeout(1) # 1-second timeout
                s.connect(("127.0.0.1", 9050))
            logging.info("Successfully connected to Tor SOCKS proxy on 127.0.0.1:9050.")
            return True
        except (socket.timeout, ConnectionRefusedError):
            logging.warning("Tor SOCKS proxy not found on 127.0.0.1:9050.")
            return False

    def _update_tool_availability(self):
        """Disables or enables tools based on Tor compatibility."""
        is_tor_enabled = self.tor_proxy_check.isChecked()

        # Tools incompatible with Tor (raw sockets, special network handling)
        scapy_tools = [
            self.scan_button, self.arp_scan_button, self.ps_start_button,
            self.trace_button, self.flood_button, self.fw_test_button,
            self.wifi_scan_button, self.deauth_button, self.bf_start_button,
            self.wpa_capture_btn, self.krack_start_btn
        ]

        cli_tools_incompatible = [
            self.masscan_controls['start_btn'], # Masscan uses its own raw socket implementation
        ]

        # Disable incompatible tools if Tor is on
        for widget in scapy_tools + cli_tools_incompatible:
             if hasattr(self, widget.objectName()): # Check if widget has been initialized
                widget.setEnabled(not is_tor_enabled)
                widget.setToolTip("This tool is not compatible with the Tor proxy." if is_tor_enabled else "")

    def _handle_tor_toggle(self, checked):
        """Handles the logic when the Tor proxy checkbox is toggled."""
        if checked:
            if not self._check_tor_proxy():
                self.tor_status_label.setText("Tor Status: <font color='red'>Inactive</font>")
                self.tor_proxy_check.setChecked(False) # Uncheck the box if connection fails

                # Show the guide
                guide = TorGuideDialog(self)
                guide.exec()
                return # Stop further processing
            else:
                self.tor_status_label.setText("Tor Status: <font color='green'>Active</font>")
        else:
            self.tor_status_label.setText("Tor Status: [?]")

        self._update_tool_availability()

    def get_ai_settings(self):
        """
        Loads AI settings from the JSON file and returns a dictionary
        containing the active provider's details (endpoint, model, api_key).
        """
        settings_file = "ai_settings.json"
        try:
            if not os.path.exists(settings_file):
                # Show settings dialog if no config exists
                if self._show_ai_settings_dialog() == QDialog.DialogCode.Rejected:
                    return None # User cancelled

            with open(settings_file, 'r') as f:
                settings = json.load(f)

            active_provider_name = settings.get("active_provider")
            active_model_name = settings.get("active_model")

            if not active_provider_name or not active_model_name:
                self.ai_assistant_tab.handle_ai_error("No active AI model selected. Please click the settings icon to choose one.")
                return None

            provider_details = {}
            if active_provider_name == "local_ai":
                local_settings = settings.get("local_ai", {})
                provider_details = {
                    "provider": "local_ai",
                    "endpoint": local_settings.get("endpoint"),
                    "model": active_model_name,
                    "api_key": None
                }
            else: # It's an online service
                online_settings = settings.get("online_ai", {})
                provider_data = online_settings.get(active_provider_name, {})
                api_key = provider_data.get("api_key")

                endpoint = ""
                if active_provider_name == "OpenAI":
                    endpoint = "https://api.openai.com/v1/chat/completions"
                # ... (add other online providers here)

                provider_details = {
                    "provider": active_provider_name,
                    "endpoint": endpoint,
                    "model": active_model_name,
                    "api_key": api_key
                }

            if not provider_details.get("endpoint"):
                 self.ai_assistant_tab.handle_ai_error(f"Endpoint for '{active_provider_name}' is missing or not supported yet.")
                 return None

            return provider_details

        except (IOError, json.JSONDecodeError) as e:
            self.ai_assistant_tab.handle_ai_error(f"Error loading AI settings: {e}")
            return None
        except Exception as e:
            self.ai_assistant_tab.handle_ai_error(f"An unexpected error occurred while getting AI settings: {e}")
            return None


    def _show_about_dialog(self):
        dialog = QMessageBox(self)
        dialog.setWindowTitle("About Zurvan + AI")

        # Add the logo
        pixmap = QIcon(self.icon_path("Zurvan.png")).pixmap(80, 80)
        dialog.setIconPixmap(pixmap)

        about_text = """
        <b>Zurvan + AI v3.0</b>
        <p>A Comprehensive Security Testing Platform, powered by AI.</p>
        <p>This application provides an integrated suite of tools for network analysis, vulnerability scanning, and penetration testing, augmented by AI-powered analysis and guidance.</p>
        <br>
        <p><b>Developers:</b><br>
        Mohammadmahdi Farhadianfard<br>
        Alireza Aminian<br>
        <br>
        Contact: mohammadmahdi.farhadianfard@gmail.com
        </p>
        """
        dialog.setText(about_text)
        dialog.exec()

    def _create_status_bar(self):
        self.status_bar = QStatusBar(self); self.setStatusBar(self.status_bar); self.status_bar.showMessage("Ready")

    def _create_resource_bar(self):
        """Creates the top resource monitor bar."""
        resource_frame = QFrame(); resource_frame.setFrameShape(QFrame.Shape.StyledPanel)
        resource_layout = QHBoxLayout(resource_frame)
        resource_layout.setContentsMargins(5, 2, 5, 2)

        # Add Logo and Tooltip
        logo_label = QLabel()
        logo_pixmap = QIcon(self.icon_path("Zurvan-mono.png")).pixmap(40, 40)
        logo_label.setPixmap(logo_pixmap)
        logo_label.setToolTip("Zurvan made by Poorija, Email: mohammadmahdi.farhadianfard@gmail.com")
        resource_layout.addWidget(logo_label)
        resource_layout.addSpacing(15)

        resource_layout.addWidget(QLabel("<b>CPU:</b>"))
        self.cpu_graph = ResourceGraph(color='c')
        self.cpu_graph.setFixedHeight(60)
        self.cpu_graph.setMaximumWidth(250)
        resource_layout.addWidget(self.cpu_graph, 1) # Add stretch factor

        resource_layout.addWidget(QLabel("<b>RAM:</b>"))
        self.ram_graph = ResourceGraph(color='m')
        self.ram_graph.setFixedHeight(60)
        self.ram_graph.setMaximumWidth(250)
        resource_layout.addWidget(self.ram_graph, 1) # Add stretch factor

        if GPUtil:
            resource_layout.addWidget(QLabel("<b>GPU:</b>"))
            self.gpu_graph = ResourceGraph(color='y')
            self.gpu_graph.setFixedHeight(60)
            self.gpu_graph.setMaximumWidth(250)
            resource_layout.addWidget(self.gpu_graph, 1)

        resource_layout.addWidget(QLabel("<b>Disk R/W:</b>"))
        self.disk_label = QLabel("---/--- MB/s"); resource_layout.addWidget(self.disk_label)
        resource_layout.addStretch()

        resource_layout.addWidget(QLabel("<b>Net Sent/Recv:</b>"))
        self.net_label = QLabel("---/--- KB/s"); resource_layout.addWidget(self.net_label)
        resource_layout.addStretch()

        # --- Public IP Display ---
        resource_layout.addWidget(QLabel("<b>Real IP:</b>"))
        self.real_ip_flag_label = QLabel()
        self.real_ip_label = QLabel("?.?.?.?")
        self.real_ip_label.setToolTip("Your public IP address without any proxy or VPN.")
        resource_layout.addWidget(self.real_ip_flag_label)
        resource_layout.addWidget(self.real_ip_label)
        resource_layout.addStretch()

        resource_layout.addWidget(QLabel("<b>Proxy/VPN IP:</b>"))
        self.proxy_ip_flag_label = QLabel()
        self.proxy_ip_label = QLabel("?.?.?.?")
        self.proxy_ip_label.setToolTip("Your public IP address as seen through the active proxy or VPN.")
        resource_layout.addWidget(self.proxy_ip_flag_label)
        resource_layout.addWidget(self.proxy_ip_label)
        resource_layout.addStretch()

        self.time_label = QLabel("Loading...")
        self.time_label.setToolTip("Current system date, time, and timezone")
        self.time_label.setStyleSheet("font-weight: bold;")

        separator = QFrame()
        separator.setFrameShape(QFrame.Shape.VLine)
        separator.setFrameShadow(QFrame.Shadow.Sunken)
        resource_layout.addWidget(separator)

        resource_layout.addWidget(self.time_label)
        resource_layout.addStretch()


        # --- User Profile Button ---
        self.user_profile_button = QPushButton()
        self.user_profile_button.setFlat(True)
        self.user_profile_button.setIconSize(QSize(32, 32))
        self.user_profile_button.setToolTip("User Profile & Logout")
        resource_layout.addWidget(self.user_profile_button)

        self.main_layout.addWidget(resource_frame)
        self.user_profile_button.clicked.connect(self._show_user_menu)


    def _update_clock(self):
        """Updates the time label with the current time and timezone."""
        # Use a try-except block to handle potential time formatting issues on different OSes
        try:
            # This format includes Year-Month-Day, Hour:Minute:Second, and Timezone Name
            current_time = time.strftime("%Y-%m-%d %H:%M:%S %Z")
        except Exception as e:
            logging.warning(f"Could not format time with timezone: {e}")
            # Fallback to a simpler format if the timezone name (%Z) causes issues
            current_time = time.strftime("%H:%M:%S")
        self.time_label.setText(current_time)

    def _update_resource_stats(self, stats):
        """Updates the resource labels with new stats from the monitor thread."""
        self.cpu_graph.update_data(stats["cpu_percent"])
        self.ram_graph.update_data(stats["ram_percent"])
        if hasattr(self, 'gpu_graph'):
            self.gpu_graph.update_data(stats.get("gpu_percent", 0))
        self.disk_label.setText(stats["disk_str"])
        self.net_label.setText(stats["net_str"])

    def _update_public_ips(self):
        """Starts threads to fetch the real and proxy/VPN public IPs."""
        self.real_ip_label.setText("Checking...")
        self.proxy_ip_label.setText("Checking...")

        # Fetch real IP
        self.real_ip_thread = IpFetchThread("real", use_proxy=False, parent=self)
        self.real_ip_thread.ip_fetched.connect(self._update_ip_label)
        self.real_ip_thread.start()

        # Fetch IP through proxy/VPN
        self.proxy_ip_thread = IpFetchThread("proxy", use_proxy=True, parent=self)
        self.proxy_ip_thread.ip_fetched.connect(self._update_ip_label)
        self.proxy_ip_thread.start()

    def _get_flag_icon(self, country_code):
        """
        Gets a QIcon for a country flag from the local 'icons/flags' directory.
        """
        if not country_code or country_code == "N/A":
            return QIcon()

        # Flags are now stored locally as SVG, which needs to be rendered to a QPixmap
        flags_dir = os.path.join(self.script_dir, "icons", "flags")
        flag_path = os.path.join(flags_dir, f"{country_code.lower()}.svg")

        if not os.path.exists(flag_path):
            logging.warning(f"Flag icon not found for country code: {country_code} at {flag_path}")
            return QIcon()

        # Render SVG to a QPixmap to be used in a QIcon
        try:
            # The flags are small, so a 16x16 rendering is sufficient for the UI
            renderer = QSvgRenderer(flag_path)
            pixmap = QPixmap(16, 16)
            pixmap.fill(Qt.GlobalColor.transparent)
            painter = QPainter(pixmap)
            renderer.render(painter)
            painter.end()
            return QIcon(pixmap)
        except Exception as e:
            logging.error(f"Could not render SVG flag for {country_code}: {e}")
            return QIcon()

    def _update_ip_label(self, ip_type, ip_address, country_code):
        """Updates the appropriate IP label and flag with the fetched address."""
        flag_icon = self._get_flag_icon(country_code)

        if ip_type == "real":
            self.real_ip_address = ip_address # Store in attribute
            self.real_ip_label.setText(f"<font color='cyan'>{ip_address}</font>")
            self.real_ip_flag_label.setPixmap(flag_icon.pixmap(16, 16))
        elif ip_type == "proxy":
            self.proxy_ip_address = ip_address # Store in attribute
            # Change color based on whether the IP is different from the real one
            if ip_address != "N/A" and ip_address != self.real_ip_address:
                color = 'lightgreen'
            else:
                color = 'orange'
            self.proxy_ip_label.setText(f"<font color='{color}'>{ip_address}</font>")
            self.proxy_ip_flag_label.setPixmap(flag_icon.pixmap(16, 16))

    def _handle_refresh_interval_change(self, text):
        """Updates the resource monitor's refresh interval."""
        if not self.resource_monitor_thread:
            return

        if text == "Off":
            self.resource_monitor_thread.pause()
        else:
            interval = int(text.replace('s', ''))
            self.resource_monitor_thread.set_interval(interval)

    def _setup_app_lock_monitor(self):
        """Initializes and starts the activity monitor and auto-lock timer."""
        # Setup timer
        self.auto_lock_timer = QTimer(self)
        self.auto_lock_timer.setSingleShot(True)
        self.auto_lock_timer.timeout.connect(self._show_lock_screen)

        # Setup event filter
        self.activity_monitor = ActivityMonitorEventFilter()
        self.activity_monitor.user_active.connect(self._reset_auto_lock_timer)
        QApplication.instance().installEventFilter(self.activity_monitor)

        logging.info("App Lock activity monitor installed.")
        self._reset_auto_lock_timer()

    def _load_and_apply_user_settings(self):
        """Loads user-specific settings from DB and applies them to the UI."""
        if not self.current_user:
            return

        user_data = database.get_user_by_id(self.current_user['id'])
        if not user_data:
            return

        # App Lock Timeout
        self.app_lock_timeout_minutes = user_data.get('app_lock_timeout', 15)
        for action in self.app_lock_timeout_group.actions():
            if action.data() == self.app_lock_timeout_minutes:
                action.setChecked(True)
                break
        self._reset_auto_lock_timer() # Start timer with loaded value

        # App Lock Unlock Method
        unlock_method = user_data.get('app_unlock_method', 'password')
        for action in self.app_unlock_method_group.actions():
            if action.text().lower().startswith(unlock_method):
                action.setChecked(True)
                break

    def _reset_auto_lock_timer(self):
        """Stops and restarts the auto-lock timer if a timeout is set."""
        if self.auto_lock_timer and self.app_lock_timeout_minutes > 0:
            self.auto_lock_timer.start(self.app_lock_timeout_minutes * 60 * 1000) # Convert minutes to ms

    def _handle_auto_lock_change(self, action):
        """Handles when the user changes the auto-lock timeout."""
        self.app_lock_timeout_minutes = action.data()
        if self.app_lock_timeout_minutes == 0:
            self.auto_lock_timer.stop()
            logging.info("Auto-lock timer disabled.")
        else:
            self._reset_auto_lock_timer()
            logging.info(f"Auto-lock timeout set to {self.app_lock_timeout_minutes} minutes.")

        # Save to DB
        if self.current_user:
            current_method_action = self.app_unlock_method_group.checkedAction()
            current_method = "pin" if current_method_action and "pin" in current_method_action.text().lower() else "password"
            database.update_user_app_lock_settings(self.current_user['id'], self.app_lock_timeout_minutes, current_method)

    def _handle_unlock_method_change(self, action):
        """Handles when the user changes the unlock method."""
        method = "pin" if "pin" in action.text().lower() else "password"

        if self.current_user and method == "pin":
            # Check if a PIN is actually set before allowing the change
            user_data = database.get_user_by_id(self.current_user['id'])
            if not user_data.get('pin_hash'):
                QMessageBox.warning(self, "PIN Not Set",
                                    "You have not configured a PIN for your account.\n\n"
                                    "Please go to Profile -> Set PIN to create one before enabling this feature.")
                # Revert the selection back to password
                for act in self.app_unlock_method_group.actions():
                    if "password" in act.text().lower():
                        # block signals to prevent recursion
                        act.blockSignals(True)
                        act.setChecked(True)
                        act.blockSignals(False)
                        break
                return # Stop processing

        logging.info(f"App unlock method changed to: {method}")
        # Save to DB
        if self.current_user:
            database.update_user_app_lock_settings(self.current_user['id'], self.app_lock_timeout_minutes, method)

    def _show_lock_screen(self):
        """Applies a blur effect and displays the application lock screen."""
        if not self.current_user:
            return

        # Log the lock event
        database.log_activity(
            user_id=self.current_user['id'],
            category='Authentication',
            action='App Locked',
            target=self.current_user['username']
        )

        # Apply blur effect
        blur_effect = QGraphicsOpacityEffect(self.central_widget)
        self.central_widget.setGraphicsEffect(blur_effect)

        # Animate the blur effect for a smoother transition
        self.blur_animation = QPropertyAnimation(blur_effect, b"opacity")
        self.blur_animation.setDuration(300)
        self.blur_animation.setStartValue(1.0)
        self.blur_animation.setEndValue(0.01) # 99% blur
        self.blur_animation.setEasingCurve(QEasingCurve.Type.InOutQuad)
        self.blur_animation.start()

        def verification_callback(username, password, otp):
            """
            Centralized verification logic for the lock screen.
            Handles password, PIN, and OTP verification based on which field has data.
            """
            user_record = database.get_user_by_username_or_email(username)
            if not user_record:
                return False

            # First, check if the account is currently locked.
            if user_record['lockout_until'] and datetime.fromisoformat(user_record['lockout_until']) > datetime.now():
                lockout_end_time = datetime.fromisoformat(user_record['lockout_until'])
                remaining = lockout_end_time - datetime.now()
                remaining_minutes = max(0, remaining.seconds // 60)
                remaining_seconds = max(0, remaining.seconds % 60)
                QMessageBox.warning(self, "Account Locked", f"This account is temporarily locked. Try again in {remaining_minutes:02d}:{remaining_seconds:02d}")
                return False

            # --- OTP Verification ---
            if otp:
                if 'otp_secret' not in user_record.keys() or not user_record['otp_secret']:
                    QMessageBox.warning(self, "Authentication Failed", "OTP is not enabled for this account.")
                    return False # No failed attempt registered as it's a config issue.

                totp = pyotp.TOTP(user_record['otp_secret'])
                if totp.verify(otp):
                    database.clear_login_attempts(user_record['id'])
                    database.log_activity(
                        user_id=user_record['id'],
                        category='Authentication',
                        action='App Unlocked',
                        target=user_record['username'],
                        details="Unlocked with OTP"
                    )
                    self._reset_auto_lock_timer()
                    return True
                else:
                    QMessageBox.warning(self, "Authentication Failed", "Invalid One-Time Password.")
                    database.register_failed_login_attempt(username)
                    return False

            # --- Password / PIN Verification ---
            if password:
                # 1. Check if it's the main password.
                # We use a special check that doesn't register a failed attempt yet.
                is_password_correct = database.verify_password_only(user_record['id'], password)
                if is_password_correct:
                    database.clear_login_attempts(user_record['id'])
                    database.log_activity(
                        user_id=user_record['id'],
                        category='Authentication',
                        action='App Unlocked',
                        target=user_record['username'],
                        details="Unlocked with Password"
                    )
                    self._reset_auto_lock_timer()
                    return True

                # 2. If not the password, check if it's the PIN.
                is_pin_correct = False
                if 'pin_hash' in user_record.keys() and user_record['pin_hash']:
                    is_pin_correct = (database.hash_password(password) == user_record['pin_hash'])

                if is_pin_correct:
                    database.clear_login_attempts(user_record['id'])
                    database.log_activity(
                        user_id=user_record['id'],
                        category='Authentication',
                        action='App Unlocked',
                        target=user_record['username'],
                        details="Unlocked with PIN"
                    )
                    self._reset_auto_lock_timer()
                    return True

                # 3. If neither password nor PIN is correct, now we register a single failed attempt.
                QMessageBox.warning(self, "Authentication Failed", "Invalid Password or PIN.")
                database.register_failed_login_attempt(username)
                return False

            # Should not be reached if called from the dialog, but as a fallback.
            return False

        lock_dialog = AppLockDialog(verification_callback, self)

        # This makes the lock dialog semi-transparent to see the blur effect
        lock_dialog.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)

        lock_dialog.exec()

        # Remove blur effect after unlocking
        self.central_widget.setGraphicsEffect(None)

    def _run_command_with_sudo_prompt(self, command, callback, cancel_handler_key):
        """
        Modified for Root Access: Bypass password prompt and execute callback directly.
        Since the app runs as root, we pass an empty string for sudo_password.
        """
        # از آنجا که برنامه دسترسی روت دارد، نیازی به پرسیدن پسورد نیست.
        # تابع callback بلافاصله با پسورد خالی فراخوانی می‌شود.
        # توابع اجرایی (مانند execute_command_thread) وقتی پسورد خالی باشد، دستور را بدون sudo اجرا می‌کنند
        # که برای کاربری root صحیح است.
        callback(sudo_password="")

    def _create_header_bar(self):
        """Creates the top header bar with interface and theme selectors."""
        header_frame = QWidget()
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(0, 5, 0, 5)

        # Interface Selector
        header_layout.addWidget(QLabel("Network Interface:"))
        try:
            ifaces = ["Automatic"] + [iface.name for iface in get_working_ifaces()]
        except Exception as e:
            logging.error(f"Could not get network interfaces: {e}", exc_info=True)
            ifaces = ["Automatic"]
        self.iface_combo = QComboBox()
        self.iface_combo.addItems(ifaces)
        self.iface_combo.currentTextChanged.connect(self._update_tool_targets)
        header_layout.addWidget(self.iface_combo)

        header_layout.addStretch()

        # --- App Lock Button ---
        self.app_lock_button = QToolButton()
        self.app_lock_button.setText("App Lock")
        self.app_lock_button.setIcon(QIcon(self.icon_path('lock.svg')))
        self.app_lock_button.setPopupMode(QToolButton.ToolButtonPopupMode.MenuButtonPopup)
        self.app_lock_button.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextBesideIcon)

        lock_menu = QMenu(self)
        lock_now_action = lock_menu.addAction("Lock Now")
        lock_now_action.triggered.connect(self._show_lock_screen)
        lock_menu.addSeparator()

        auto_lock_menu = lock_menu.addMenu("Auto-lock delay")
        self.app_lock_timeout_group = QActionGroup(self)
        self.app_lock_timeout_group.setExclusive(True)

        timeouts = {"5 Minutes": 5, "15 Minutes": 15, "30 Minutes": 30, "1 Hour": 60, "Disabled": 0}
        for text, minutes in timeouts.items():
            action = QAction(text, self, checkable=True)
            action.setData(minutes)
            self.app_lock_timeout_group.addAction(action)
            auto_lock_menu.addAction(action)
            if minutes == 15: # Default
                action.setChecked(True)

        lock_menu.addSeparator()
        unlock_method_menu = lock_menu.addMenu("Unlock Method")
        self.app_unlock_method_group = QActionGroup(self)
        self.app_unlock_method_group.setExclusive(True)

        pass_action = QAction("Use Password", self, checkable=True)
        pass_action.setChecked(True)
        pin_action = QAction("Use PIN", self, checkable=True)

        self.app_unlock_method_group.addAction(pass_action)
        self.app_unlock_method_group.addAction(pin_action)
        unlock_method_menu.addAction(pass_action)
        unlock_method_menu.addAction(pin_action)

        self.app_lock_button.setMenu(lock_menu)
        self.app_lock_button.setDefaultAction(lock_now_action)
        header_layout.addWidget(self.app_lock_button)

        # Theme Switcher
        header_layout.addWidget(QLabel("Theme:"))
        self.theme_combo = QComboBox()
        self.theme_combo.addItems([theme.replace('.xml', '') for theme in list_themes()])
        self.theme_combo.textActivated.connect(self._handle_theme_change)
        header_layout.addWidget(self.theme_combo)

        self.main_layout.addWidget(header_frame)

        # Connect App Lock signals
        self.app_lock_timeout_group.triggered.connect(self._handle_auto_lock_change)
        self.app_unlock_method_group.triggered.connect(self._handle_unlock_method_change)

    def _setup_app_lock_monitor(self):
        """Initializes and starts the activity monitor and auto-lock timer."""
        # Setup timer
        self.auto_lock_timer = QTimer(self)
        self.auto_lock_timer.setSingleShot(True)
        self.auto_lock_timer.timeout.connect(self._show_lock_screen)

        # Setup event filter
        self.activity_monitor = ActivityMonitorEventFilter()
        self.activity_monitor.user_active.connect(self._reset_auto_lock_timer)
        QApplication.instance().installEventFilter(self.activity_monitor)

        logging.info("App Lock activity monitor installed.")
        self._reset_auto_lock_timer()

    def _load_and_apply_user_settings(self):
        """Loads user-specific settings from DB and applies them to the UI."""
        if not self.current_user:
            return

        user_data = database.get_user_by_id(self.current_user['id'])
        if not user_data:
            return

        # App Lock Timeout
        self.app_lock_timeout_minutes = user_data.get('app_lock_timeout', 15)
        for action in self.app_lock_timeout_group.actions():
            if action.data() == self.app_lock_timeout_minutes:
                action.setChecked(True)
                break
        self._reset_auto_lock_timer() # Start timer with loaded value

        # App Lock Unlock Method
        unlock_method = user_data.get('app_unlock_method', 'password')
        for action in self.app_unlock_method_group.actions():
            if action.text().lower().startswith(unlock_method):
                action.setChecked(True)
                break

    def _reset_auto_lock_timer(self):
        """Stops and restarts the auto-lock timer if a timeout is set."""
        if self.auto_lock_timer and self.app_lock_timeout_minutes > 0:
            self.auto_lock_timer.start(self.app_lock_timeout_minutes * 60 * 1000) # Convert minutes to ms

    def _handle_auto_lock_change(self, action):
        """Handles when the user changes the auto-lock timeout."""
        self.app_lock_timeout_minutes = action.data()
        if self.app_lock_timeout_minutes == 0:
            self.auto_lock_timer.stop()
            logging.info("Auto-lock timer disabled.")
        else:
            self._reset_auto_lock_timer()
            logging.info(f"Auto-lock timeout set to {self.app_lock_timeout_minutes} minutes.")

        # Save to DB
        if self.current_user:
            current_method_action = self.app_unlock_method_group.checkedAction()
            current_method = "pin" if current_method_action and "pin" in current_method_action.text().lower() else "password"
            database.update_user_app_lock_settings(self.current_user['id'], self.app_lock_timeout_minutes, current_method)

    def _handle_unlock_method_change(self, action):
        """Handles when the user changes the unlock method."""
        method = "pin" if "pin" in action.text().lower() else "password"
        logging.info(f"App unlock method changed to: {method}")
        # Save to DB
        if self.current_user:
            database.update_user_app_lock_settings(self.current_user['id'], self.app_lock_timeout_minutes, method)

    def _handle_theme_change(self, theme_name):
        theme_file = f"{theme_name}.xml"
        invert_secondary = "light" in theme_name

        # This dictionary must be kept in sync with the one in login.py and main()
        extra_qss = {
            'QGroupBox': {
                'border': '1px solid #444;',
                'border-radius': '16px',
                'margin-top': '10px',
                'padding-top': '15px',
            },
            'QGroupBox::title': {
                'subcontrol-origin': 'margin',
                'subcontrol-position': 'top left',
                'padding': '0 10px',
            },
            'QTabWidget::pane': {
                'border-top': '1px solid #444;',
                'margin-top': '-1px',
                'border-radius': '12px',
            },
            'QFrame': {
                'border-radius': '12px',
            },
            'QPushButton': {
                'border-radius': '12px',
                'padding': '8px 16px',
            },
            'QLineEdit': {
                'border-radius': '12px',
                'padding': '6px 12px',
            },
            'QComboBox': {
                'border-radius': '12px',
                'padding': '6px 12px',
            },
            'QTextEdit': {
                'border-radius': '12px',
                'padding': '8px',
            },
            'QPlainTextEdit': {
                'border-radius': '12px',
                'padding': '8px',
            },
            'QListWidget': {
                'border-radius': '12px',
            },
            'QTreeWidget': {
                'border-radius': '12px',
            },
            'QPushButton:hover': {
                'background-color': '{{primaryColor}}',
                'color': '{{secondaryDarkColor}}',
                'border': '2px solid {{primaryLightColor}}',
            },
            'QPushButton:pressed': {
                'background-color': '{{secondaryColor}}',
                'color': '{{primaryColor}}',
                'border': '2px solid {{primaryColor}}',
            }
        }

        apply_stylesheet(QApplication.instance(), theme=theme_file, invert_secondary=invert_secondary, extra=extra_qss)

        # After applying the stylesheet, notify the AI tab to update its themed icons
        if hasattr(self, 'ai_assistant_tab'):
            self.ai_assistant_tab.update_theme()

    def get_selected_iface(self):
        iface = self.iface_combo.currentText()
        return iface if iface != "Automatic" else None

    def _create_main_tabs(self):
        """Creates the main QTabWidget and adds all the tool tabs."""
        self.tab_widget = QTabWidget()
        self.tab_widget.setStyleSheet("""
            QTabWidget::tab-bar {
                alignment: center;
            }
            QTabBar::tab:!selected:!last {
                border-right: 1px solid #444;
            }
        """)
        self.main_layout.addWidget(self.tab_widget)
        self.tab_widget.addTab(self._create_sniffer_tab(), QIcon(self.icon_path("search.svg")), "Packet Sniffer")
        self.tab_widget.addTab(self._create_crafter_tab(), QIcon(self.icon_path("edit-3.svg")), "Packet Crafter")
        self.tab_widget.addTab(self._create_tools_tab(), QIcon(self.icon_path("tool.svg")), "Network Tools")
        self.tab_widget.addTab(self._create_advanced_tools_tab(), QIcon(self.icon_path("shield.svg")), "Advanced Tools")
        self.tab_widget.addTab(self._create_osint_tab(), QIcon(self.icon_path("search.svg")), "OSINT")
        self.tab_widget.addTab(self._create_wireless_tools_tab(), QIcon(self.icon_path("wifi.svg")), "Wireless Tools")
        self.tab_widget.addTab(self._create_reporting_tab(), QIcon(self.icon_path("file-text.svg")), "Reporting")
        self.tab_widget.addTab(self._create_lab_tab(), QIcon(self.icon_path("layers.svg")), "LAB")

        self.ai_assistant_tab = AIAssistantTab(self)
        self.tab_widget.addTab(self.ai_assistant_tab, QIcon(self.icon_path("terminal.svg")), "AI Assistant")

        self.tab_widget.addTab(self._create_threat_intelligence_tab(), QIcon(self.icon_path("database.svg")), "Threat Intelligence")
        self.tab_widget.addTab(self._create_activity_journal_tab(), QIcon(self.icon_path("book-open.svg")), "Activity Journal")

        self.tab_widget.addTab(self._create_proxy_tab(), QIcon(self.icon_path("shield.svg")), "Proxy")
        self.tab_widget.addTab(self._create_system_info_tab(), QIcon(self.icon_path("info.svg")), "System Info")

    def _create_proxy_tab(self):
        """Creates the tab for managing proxies, VPNs, and Tor."""
        proxy_tabs = QTabWidget()

        proxy_tabs.addTab(self._create_proxy_manager_tab(), "Proxy Manager")
        proxy_tabs.addTab(self._create_v2ray_manager_tab(), "V2Ray Manager")
        proxy_tabs.addTab(self._create_tor_manager_tab(), "Tor Manager")
        proxy_tabs.addTab(self._create_l2tp_client_tab(), "L2TP/IPsec Client")
        proxy_tabs.addTab(self._create_openvpn_client_tab(), "OpenVPN Client")
        proxy_tabs.addTab(self._create_wireguard_client_tab(), "WireGuard Client")
        proxy_tabs.addTab(self._create_tun_system_tab(), "TUN System")
        proxy_tabs.addTab(self._create_anonymity_tools_tab(), "Anonymity & Evasion")

        return proxy_tabs

    def _create_proxy_manager_tab(self):
        """Creates the UI for the Proxy Manager."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Proxy List ---
        list_box = QGroupBox("Proxy List")
        list_layout = QVBoxLayout(list_box)
        self.proxy_list_widget = QListWidget()
        list_layout.addWidget(self.proxy_list_widget)

        # --- List Controls ---
        list_controls_layout = QHBoxLayout()
        add_proxy_btn = QPushButton("Add Proxy")
        add_proxy_btn.clicked.connect(self._add_proxy)
        remove_proxy_btn = QPushButton("Remove Selected")
        remove_proxy_btn.clicked.connect(self._remove_proxy)
        load_from_file_btn = QPushButton("Load from File")
        load_from_file_btn.clicked.connect(self._load_proxies_from_file)
        list_controls_layout.addWidget(add_proxy_btn)
        list_controls_layout.addWidget(remove_proxy_btn)
        list_controls_layout.addWidget(load_from_file_btn)
        list_layout.addLayout(list_controls_layout)
        layout.addWidget(list_box)

        # --- Proxychains Settings ---
        proxychains_box = QGroupBox("Proxychains-ng Settings")
        proxychains_layout = QFormLayout(proxychains_box)
        self.proxy_chain_type_combo = QComboBox()
        self.proxy_chain_type_combo.addItems(["strict_chain", "random_chain", "round_robin_chain", "dynamic_chain"])
        self.proxy_chain_type_combo.setToolTip("Set the chaining type for proxychains.\n- strict_chain: All proxies in the list are used in order.\n- random_chain: Each connection is made through a random proxy from the list.\n- dynamic_chain: Similar to strict_chain, but allows fallback to other proxies if one fails.")
        proxychains_layout.addRow("Chain Type:", self.proxy_chain_type_combo)
        self.proxy_dns_check = QCheckBox("Proxy DNS Requests")
        self.proxy_dns_check.setChecked(True)
        self.proxy_dns_check.setToolTip("Enable DNS-level proxying for requests.")
        proxychains_layout.addRow(self.proxy_dns_check)

        self.chain_len_edit = QLineEdit()
        self.chain_len_edit.setPlaceholderText("e.g., 2 (for random_chain)")
        self.chain_len_edit.setToolTip("In 'random_chain' mode, specify how many proxies from the list to use in the chain.")
        proxychains_layout.addRow("Random Chain Length:", self.chain_len_edit)

        self.tcp_read_timeout_edit = QLineEdit("15000")
        proxychains_layout.addRow("TCP Read Timeout (ms):", self.tcp_read_timeout_edit)
        self.tcp_connect_timeout_edit = QLineEdit("8000")
        proxychains_layout.addRow("TCP Connect Timeout (ms):", self.tcp_connect_timeout_edit)
        layout.addWidget(proxychains_box)

        # --- Status & Testing ---
        status_layout = QHBoxLayout()
        self.proxy_status_label = QLabel("Status: Inactive")
        test_proxies_btn = QPushButton("Test All Proxies")
        test_proxies_btn.clicked.connect(self._test_proxies)
        status_layout.addWidget(self.proxy_status_label)
        status_layout.addStretch()
        status_layout.addWidget(test_proxies_btn)
        layout.addLayout(status_layout)

        # --- Proxy Rotation ---
        self.rotation_box = QGroupBox("Proxy Rotation")
        self.rotation_box.setObjectName("rotation_box")
        self.rotation_box.setCheckable(True)
        self.rotation_box.setChecked(False)
        rotation_layout = QFormLayout(self.rotation_box)

        self.rotation_type_group = QButtonGroup(widget)
        self.fixed_interval_radio = QRadioButton("Fixed Interval")
        self.random_interval_radio = QRadioButton("Random Interval")
        self.fixed_interval_radio.setChecked(True)
        self.rotation_type_group.addButton(self.fixed_interval_radio)
        self.rotation_type_group.addButton(self.random_interval_radio)

        rotation_type_layout = QHBoxLayout()
        rotation_type_layout.addWidget(self.fixed_interval_radio)
        rotation_type_layout.addWidget(self.random_interval_radio)
        rotation_layout.addRow("Rotation Type:", rotation_type_layout)

        self.fixed_interval_edit = QLineEdit("60")
        rotation_layout.addRow("Interval (seconds):", self.fixed_interval_edit)

        self.random_interval_min_edit = QLineEdit("30")
        self.random_interval_max_edit = QLineEdit("120")
        random_layout = QHBoxLayout()
        random_layout.addWidget(QLabel("Min:"))
        random_layout.addWidget(self.random_interval_min_edit)
        random_layout.addWidget(QLabel("Max:"))
        random_layout.addWidget(self.random_interval_max_edit)
        self.random_interval_widget = QWidget()
        self.random_interval_widget.setLayout(random_layout)
        rotation_layout.addRow("Random Interval (s):", self.random_interval_widget)

        def toggle_rotation_inputs():
            is_fixed = self.fixed_interval_radio.isChecked()
            self.fixed_interval_edit.setVisible(is_fixed)
            rotation_layout.labelForField(self.fixed_interval_edit).setVisible(is_fixed)
            self.random_interval_widget.setVisible(not is_fixed)
            rotation_layout.labelForField(self.random_interval_widget).setVisible(not is_fixed)

        self.fixed_interval_radio.toggled.connect(toggle_rotation_inputs)
        self.random_interval_radio.toggled.connect(toggle_rotation_inputs)
        toggle_rotation_inputs() # Set initial state

        # --- Timer setup and connections ---
        self.proxy_rotation_timer = QTimer(self)
        self.proxy_rotation_timer.timeout.connect(self._rotate_proxy)
        self.rotation_box.toggled.connect(self._handle_rotation_toggle)
        self.fixed_interval_radio.toggled.connect(self._restart_rotation_logic)
        self.random_interval_radio.toggled.connect(self._restart_rotation_logic)
        self.fixed_interval_edit.textChanged.connect(self._update_rotation_timer)

        layout.addWidget(self.rotation_box)


        # --- IP Lookup ---
        ip_lookup_box = QGroupBox("IP Address Lookup")
        ip_lookup_layout = QVBoxLayout(ip_lookup_box)
        ip_lookup_controls_layout = QHBoxLayout()
        self.ip_lookup_edit = QLineEdit()
        self.ip_lookup_edit.setPlaceholderText("Enter IP address or domain...")
        ip_lookup_controls_layout.addWidget(self.ip_lookup_edit)
        ip_lookup_btn = QPushButton("Lookup")
        ip_lookup_btn.clicked.connect(self._lookup_ip_address)
        ip_lookup_controls_layout.addWidget(ip_lookup_btn)
        ip_lookup_layout.addLayout(ip_lookup_controls_layout)
        self.ip_lookup_results_browser = QTextBrowser()
        ip_lookup_layout.addWidget(self.ip_lookup_results_browser)
        layout.addWidget(ip_lookup_box)

        layout.addStretch()
        return widget

    def _add_proxy(self):
        """Opens a dialog to add a new proxy to the list."""
        proxy, ok = QInputDialog.getText(self, "Add Proxy", "Enter proxy address (e.g., socks5://127.0.0.1:9050):")
        if ok and proxy:
            self.proxy_list_widget.addItem(proxy)

    def _remove_proxy(self):
        """Removes the selected proxy from the list."""
        for item in self.proxy_list_widget.selectedItems():
            self.proxy_list_widget.takeItem(self.proxy_list_widget.row(item))

    def _load_proxies_from_file(self):
        """Loads a list of proxies from a text file."""
        file_path, _ = QFileDialog.getOpenFileName(self, "Load Proxies", "", "Text Files (*.txt);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            try:
                with open(file_path, 'r') as f:
                    proxies = [line.strip() for line in f if line.strip()]
                self.proxy_list_widget.addItems(proxies)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load proxies from file: {e}")

    def _test_proxies(self):
        """Tests the connectivity of all proxies in the list in a background thread."""
        if self.proxy_list_widget.count() == 0:
            QMessageBox.information(self, "No Proxies", "There are no proxies in the list to test.")
            return
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        self.is_tool_running = True
        self.proxy_status_label.setText("Status: Testing...")
        self.worker = WorkerThread(self._proxy_test_thread)
        self.active_threads.append(self.worker)
        self.worker.start()

    def _proxy_test_thread(self):
        """Worker thread to test proxies."""
        q = self.tool_results_queue
        proxies = [self.proxy_list_widget.item(i).text() for i in range(self.proxy_list_widget.count())]
        for i, proxy_str in enumerate(proxies):
            proxies_dict = {
                "http": proxy_str,
                "https": proxy_str,
            }
            try:
                # Test by trying to fetch a known lightweight page
                response = requests.get("http://httpbin.org/get", proxies=proxies_dict, timeout=5)
                if response.status_code == 200:
                    q.put(('proxy_test_result', i, True))
                else:
                    q.put(('proxy_test_result', i, False))
            except Exception:
                q.put(('proxy_test_result', i, False))
        q.put(('tool_finished', 'proxy_test'))

    def _handle_proxy_test_result(self, index, success):
        """Updates the UI with the result of a single proxy test."""
        item = self.proxy_list_widget.item(index)
        if item:
            item.setForeground(QColor("green") if success else QColor("red"))

    def _handle_rotation_toggle(self, checked):
        """Starts or stops the proxy rotation timer."""
        if checked:
            if self.proxy_list_widget.count() == 0:
                QMessageBox.warning(self, "No Proxies", "Cannot start rotation with an empty proxy list.")
                self.rotation_box.setChecked(False)
                return

            if self.current_proxy_index == -1:
                self.current_proxy_index = 0
            self._update_active_proxy_highlight()

            if self.fixed_interval_radio.isChecked():
                self._update_rotation_timer()
            else:
                self.proxy_rotation_timer.stop()
                self._schedule_random_rotation()
        else:
            self.proxy_rotation_timer.stop()
            self.current_proxy_index = -1
            self._update_active_proxy_highlight()
            logging.info("Proxy rotation stopped.")

    def _restart_rotation_logic(self, checked=False):
        """Stops and starts the rotation logic to apply new settings."""
        if self.rotation_box.isChecked():
            self._handle_rotation_toggle(False)
            self._handle_rotation_toggle(True)

    def _update_rotation_timer(self):
        """Updates the timer's interval based on UI settings. Only for fixed interval mode."""
        if not hasattr(self, 'rotation_box') or not self.rotation_box.isChecked() or not self.fixed_interval_radio.isChecked():
            if hasattr(self, 'proxy_rotation_timer'): self.proxy_rotation_timer.stop()
            return
        try:
            interval_ms = int(self.fixed_interval_edit.text()) * 1000
            if interval_ms > 0:
                self.proxy_rotation_timer.setInterval(interval_ms)
                if not self.proxy_rotation_timer.isActive():
                    self.proxy_rotation_timer.start()
                logging.info(f"Proxy rotation interval set to fixed {interval_ms}ms.")
            else:
                self.proxy_rotation_timer.stop()
        except ValueError:
            self.proxy_rotation_timer.stop()

    def _schedule_random_rotation(self):
        """Calculates a random interval and sets a single-shot timer."""
        if not hasattr(self, 'rotation_box') or not self.rotation_box.isChecked():
            return
        try:
            min_s = int(self.random_interval_min_edit.text())
            max_s = int(self.random_interval_max_edit.text())
            if min_s > 0 and max_s >= min_s:
                interval_s = random.randint(min_s, max_s)
                interval_ms = interval_s * 1000
                logging.info(f"Next random rotation in {interval_s} seconds.")
                QTimer.singleShot(interval_ms, self._rotate_proxy)
        except ValueError:
            logging.warning("Invalid random interval values.")

    def _rotate_proxy(self):
        """Selects the next proxy in the list and handles timer logic."""
        if not self.rotation_box.isChecked() or self.proxy_list_widget.count() == 0:
            self.proxy_rotation_timer.stop()
            return

        self.current_proxy_index = (self.current_proxy_index + 1) % self.proxy_list_widget.count()
        self._update_active_proxy_highlight()
        logging.info(f"Rotated to proxy at index {self.current_proxy_index}")

        if self.random_interval_radio.isChecked():
            self.proxy_rotation_timer.stop()
            self._schedule_random_rotation()

    def _update_active_proxy_highlight(self):
        """Highlights the currently active proxy in the list."""
        for i in range(self.proxy_list_widget.count()):
            item = self.proxy_list_widget.item(i)
            font = item.font()
            font.setBold(i == self.current_proxy_index)
            item.setFont(font)

    def _lookup_ip_address(self):
        """Starts the worker thread to look up an IP address or domain."""
        target = self.ip_lookup_edit.text().strip()
        if not target:
            QMessageBox.warning(self, "Input Error", "Please enter an IP address or domain to look up.")
            return
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        self.is_tool_running = True
        self.ip_lookup_results_browser.setText(f"Looking up {target}...")
        self.worker = WorkerThread(self._ip_lookup_thread, args=(target,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _ip_lookup_thread(self, target):
        """Worker thread to fetch geolocation data, trying offline DB first."""
        q = self.tool_results_queue
        try:
            # Try offline lookup first if available
            if GEOLITE_AVAILABLE:
                match = geolite2.lookup(target)
                if match:
                    data = {
                        "source": "Offline (GeoLite2)",
                        "ip": match.ip,
                        "country_code": match.country,
                        "continent": match.continent,
                        "timezone": match.timezone,
                        "subdivisions": list(match.subdivisions),
                    }
                    q.put(('ip_lookup_result', data))
                    return # Exit after successful offline lookup

            # Fallback to online API
            url = f"http://ip-api.com/json/{target}"
            with urllib.request.urlopen(url, timeout=10) as response:
                if response.status == 200:
                    data = json.loads(response.read().decode('utf-8'))
                    if data.get("status") == "success":
                        data["source"] = "Online (ip-api.com)"
                        q.put(('ip_lookup_result', data))
                    else:
                        raise Exception(data.get("message", "Unknown API error"))
                else:
                    raise Exception(f"API request failed with status code {response.status}")
        except Exception as e:
            q.put(('error', 'IP Lookup Error', str(e)))
        finally:
            q.put(('tool_finished', 'ip_lookup', target, "IP lookup complete."))

    def _handle_ip_lookup_result(self, data):
        """Formats and displays the IP lookup results in the QTextBrowser."""
        html = "<h3>IP Lookup Results</h3>"
        html += "<table border='1' style='border-collapse: collapse; width: 100%;'>"
        for key, value in data.items():
            html += f"<tr><td style='padding: 4px; font-weight: bold;'>{key.capitalize()}</td><td style='padding: 4px;'>{value}</td></tr>"
        html += "</table>"
        self.ip_lookup_results_browser.setHtml(html)

    def _create_anonymity_tools_tab(self):
        """Creates the UI for various anonymity and evasion tools."""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(20)

        # --- MAC Address Spoofing ---
        mac_box = QGroupBox("MAC Address Spoofing")
        mac_layout = QFormLayout(mac_box)
        self.mac_iface_combo = QComboBox()
        # Populate with interfaces that likely have a MAC address
        try:
            self.mac_iface_combo.addItems([i for i, addrs in psutil.net_if_addrs().items() if any(addr.family in [psutil.AF_LINK, getattr(socket, 'AF_PACKET', -1)] for addr in addrs)])
        except Exception as e:
            logging.warning(f"Could not list interfaces for MAC spoofing: {e}")

        self.original_mac_label = QLabel("Original MAC: N/A")
        mac_layout.addRow("Interface:", self.mac_iface_combo)
        mac_layout.addRow(self.original_mac_label)

        new_mac_layout = QHBoxLayout()
        self.new_mac_edit = QLineEdit()
        self.new_mac_edit.setPlaceholderText("e.g., 00:11:22:33:44:55")
        new_mac_layout.addWidget(self.new_mac_edit)
        random_mac_btn = QPushButton("Random")
        new_mac_layout.addWidget(random_mac_btn)
        mac_layout.addRow("New MAC:", new_mac_layout)

        mac_buttons = QHBoxLayout()
        self.apply_mac_btn = QPushButton("Apply MAC")
        self.reset_mac_btn = QPushButton("Reset to Original")
        mac_buttons.addWidget(self.apply_mac_btn)
        mac_buttons.addWidget(self.reset_mac_btn)
        mac_layout.addRow(mac_buttons)
        layout.addWidget(mac_box)

        # --- DNS Changer ---
        dns_box = QGroupBox("DNS Changer")
        dns_layout = QFormLayout(dns_box)
        self.dns_provider_combo = QComboBox()
        self.dns_provider_combo.addItems(["Cloudflare (1.1.1.1)", "Google (8.8.8.8)", "Quad9 (9.9.9.9)", "Custom"])
        dns_layout.addRow("DNS Provider:", self.dns_provider_combo)
        self.custom_dns_edit = QLineEdit()
        self.custom_dns_edit.setPlaceholderText("e.g., 208.67.222.222")
        dns_layout.addRow("Custom DNS:", self.custom_dns_edit)
        self.apply_dns_btn = QPushButton("Apply DNS")
        self.reset_dns_btn = QPushButton("Reset DNS")
        dns_buttons = QHBoxLayout()
        dns_buttons.addWidget(self.apply_dns_btn)
        dns_buttons.addWidget(self.reset_dns_btn)
        dns_layout.addRow(dns_buttons)
        layout.addWidget(dns_box)

        # --- MTU Tuner ---
        mtu_box = QGroupBox("MTU Tuner")
        mtu_layout = QFormLayout(mtu_box)
        self.mtu_iface_combo = QComboBox()
        try:
            self.mtu_iface_combo.addItems(list(psutil.net_if_addrs().keys()))
        except Exception: pass
        self.current_mtu_label = QLabel("Current MTU: N/A")
        mtu_layout.addRow("Interface:", self.mtu_iface_combo)
        mtu_layout.addRow(self.current_mtu_label)
        self.new_mtu_edit = QLineEdit("1500")
        mtu_layout.addRow("New MTU:", self.new_mtu_edit)
        self.apply_mtu_btn = QPushButton("Apply MTU")
        mtu_layout.addRow(self.apply_mtu_btn)
        layout.addWidget(mtu_box)

        # --- Hostname Changer ---
        hostname_box = QGroupBox("Hostname Changer")
        hostname_layout = QFormLayout(hostname_box)
        self.current_hostname_label = QLabel(f"Current Hostname: {socket.gethostname()}")
        hostname_layout.addRow(self.current_hostname_label)
        self.new_hostname_edit = QLineEdit()
        hostname_layout.addRow("New Hostname:", self.new_hostname_edit)
        self.apply_hostname_btn = QPushButton("Apply Hostname (Temporary)")
        hostname_layout.addRow(self.apply_hostname_btn)
        layout.addWidget(hostname_box)

        # --- Secure DNS (DoH/DoT) ---
        secure_dns_box = QGroupBox("Secure DNS (DoH/DoT)")
        secure_dns_layout = QFormLayout(secure_dns_box)
        self.doh_provider_combo = QComboBox()
        self.doh_provider_combo.addItems(["System Default", "Cloudflare (DoH)", "Google (DoH)", "Quad9 (DoH)"])
        secure_dns_layout.addRow("DNS Resolver:", self.doh_provider_combo)
        self.doh_status_label = QLabel("Status: Using System Default DNS")
        secure_dns_layout.addRow(self.doh_status_label)
        layout.addWidget(secure_dns_box)


        layout.addStretch()

        # --- Connections ---
        self.doh_provider_combo.currentTextChanged.connect(self._set_secure_dns_resolver)
        self.mac_iface_combo.currentTextChanged.connect(self._update_original_mac)
        random_mac_btn.clicked.connect(self._generate_random_mac)
        self.apply_mac_btn.clicked.connect(self._apply_mac_address)
        self.reset_mac_btn.clicked.connect(self._reset_mac_address)
        self.dns_provider_combo.currentTextChanged.connect(self._handle_dns_selection)
        self.apply_dns_btn.clicked.connect(self._apply_dns)
        self.reset_dns_btn.clicked.connect(self._reset_dns)
        self.mtu_iface_combo.currentTextChanged.connect(self._update_current_mtu)
        self.apply_mtu_btn.clicked.connect(self._apply_mtu)
        self.apply_hostname_btn.clicked.connect(self._apply_hostname)

        self._update_original_mac(self.mac_iface_combo.currentText())
        self._update_current_mtu(self.mtu_iface_combo.currentText())
        self._handle_dns_selection(self.dns_provider_combo.currentText())


        return widget

    def _set_secure_dns_resolver(self, provider_text):
        """Sets the active DoH/DoT resolver URL based on the user's selection."""
        provider_map = {
            "Cloudflare (DoH)": "https://cloudflare-dns.com/dns-query",
            "Google (DoH)": "https://dns.google/dns-query",
            "Quad9 (DoH)": "https://dns.quad9.net/dns-query",
        }
        self.doh_resolver_url = provider_map.get(provider_text)
        if self.doh_resolver_url:
            self.doh_status_label.setText(f"Status: Active ({provider_text})")
        else:
            self.doh_status_label.setText("Status: Using System Default DNS")

    def _update_original_mac(self, iface):
        """Reads and displays the original MAC address for the selected interface."""
        if not iface:
            return
        try:
            addrs = psutil.net_if_addrs().get(iface, [])
            for addr in addrs:
                if addr.family == psutil.AF_LINK or addr.family == getattr(socket, 'AF_PACKET', -1):
                    self.original_mac_label.setText(f"Original MAC: {addr.address}")
                    return
            self.original_mac_label.setText("Original MAC: N/A")
        except Exception as e:
            self.original_mac_label.setText("Original MAC: Error")
            logging.error(f"Could not get MAC for {iface}: {e}")

    def _generate_random_mac(self):
        """Generates a random MAC address and puts it in the input field."""
        mac = "02:%02x:%02x:%02x:%02x:%02x" % (
            random.randint(0, 255),
            random.randint(0, 255),
            random.randint(0, 255),
            random.randint(0, 255),
            random.randint(0, 255),
        )
        self.new_mac_edit.setText(mac)

    def _apply_mac_address(self, sudo_password=None):
        """Applies the new MAC address using ip link commands."""
        iface = self.mac_iface_combo.currentText()
        new_mac = self.new_mac_edit.text().strip()
        if not iface or not new_mac:
            QMessageBox.warning(self, "Input Error", "Interface and new MAC address are required.")
            return

        if not sudo_password and sys.platform != "win32":
            self._run_command_with_sudo_prompt([], self._apply_mac_address, 'apply_mac')
            return

        commands = [
            ["ip", "link", "set", "dev", iface, "down"],
            ["ip", "link", "set", "dev", iface, "address", new_mac],
            ["ip", "link", "set", "dev", iface, "up"]
        ]
        self.worker = WorkerThread(self._execute_sudo_commands_thread, args=(commands, sudo_password, "MAC Change", "MAC address changed successfully."))
        self.worker.start()

    def _reset_mac_address(self, sudo_password=None):
        """Resets the MAC address to its original value."""
        iface = self.mac_iface_combo.currentText()
        original_mac = self.original_mac_label.text().replace("Original MAC: ", "").strip()
        if not iface or not original_mac or original_mac in ["N/A", "Error"]:
            QMessageBox.warning(self, "Input Error", "Could not determine original MAC address.")
            return

        if not sudo_password and sys.platform != "win32":
            self._run_command_with_sudo_prompt([], self._reset_mac_address, 'reset_mac')
            return

        commands = [
            ["ip", "link", "set", "dev", iface, "down"],
            ["ip", "link", "set", "dev", iface, "address", original_mac],
            ["ip", "link", "set", "dev", iface, "up"]
        ]
        self.worker = WorkerThread(self._execute_sudo_commands_thread, args=(commands, sudo_password, "MAC Reset", "MAC address reset successfully."))
        self.worker.start()

    def _handle_dns_selection(self, selection):
        """Shows or hides the custom DNS input field."""
        self.custom_dns_edit.setVisible(selection == "Custom")

    def _apply_dns(self, sudo_password=None):
        """Applies the selected DNS servers by modifying /etc/resolv.conf."""
        selection = self.dns_provider_combo.currentText()
        dns_map = {
            "Cloudflare (1.1.1.1)": "nameserver 1.1.1.1\\nnameserver 1.0.0.1",
            "Google (8.8.8.8)": "nameserver 8.8.8.8\\nnameserver 8.8.4.4",
            "Quad9 (9.9.9.9)": "nameserver 9.9.9.9\\nnameserver 149.112.112.112",
        }
        if selection == "Custom":
            custom_dns = self.custom_dns_edit.text().strip()
            if not custom_dns:
                QMessageBox.warning(self, "Input Error", "Please enter a custom DNS server.")
                return
            resolv_content = f"nameserver {custom_dns}"
        else:
            resolv_content = dns_map[selection]

        if not sudo_password and sys.platform != "win32":
            self._run_command_with_sudo_prompt([], self._apply_dns, 'apply_dns')
            return

        # Backup the original resolv.conf
        if not os.path.exists("/etc/resolv.conf.bak"):
            self.worker = WorkerThread(self._execute_sudo_commands_thread, args=([["cp", "/etc/resolv.conf", "/etc/resolv.conf.bak"]], sudo_password, "DNS Backup", "DNS config backed up."))
            self.worker.start()
            self.worker.wait() # Wait for backup to finish

        command = f"echo -e '{resolv_content}' > /etc/resolv.conf"
        self.worker = WorkerThread(self._execute_sudo_shell_command_thread, args=(command, sudo_password, "DNS Change", "DNS servers changed successfully."))
        self.worker.start()

    def _reset_dns(self, sudo_password=None):
        """Restores the original /etc/resolv.conf from backup."""
        if not os.path.exists("/etc/resolv.conf.bak"):
            QMessageBox.information(self, "No Backup", "No backup of /etc/resolv.conf was found.")
            return

        if not sudo_password and sys.platform != "win32":
            self._run_command_with_sudo_prompt([], self._reset_dns, 'reset_dns')
            return

        command = [["mv", "/etc/resolv.conf.bak", "/etc/resolv.conf"]]
        self.worker = WorkerThread(self._execute_sudo_commands_thread, args=(command, sudo_password, "DNS Reset", "DNS servers restored successfully."))
        self.worker.start()

    def _update_current_mtu(self, iface):
        """Reads and displays the current MTU for the selected interface."""
        if not iface: return
        try:
            # This is a bit of a platform-specific way to do it.
            if sys.platform == "linux":
                with open(f'/sys/class/net/{iface}/mtu', 'r') as f:
                    mtu = f.read().strip()
                self.current_mtu_label.setText(f"Current MTU: {mtu}")
            else: # Fallback for other OSes
                self.current_mtu_label.setText("Current MTU: (Not available)")
        except Exception as e:
            self.current_mtu_label.setText("Current MTU: Error")
            logging.error(f"Could not get MTU for {iface}: {e}")

    def _apply_mtu(self, sudo_password=None):
        """Applies the new MTU setting."""
        iface = self.mtu_iface_combo.currentText()
        mtu = self.new_mtu_edit.text().strip()
        if not iface or not mtu:
            QMessageBox.warning(self, "Input Error", "Interface and MTU value are required.")
            return

        if not sudo_password and sys.platform != "win32":
            self._run_command_with_sudo_prompt([], self._apply_mtu, 'apply_mtu')
            return

        command = [["ip", "link", "set", "dev", iface, "mtu", mtu]]
        self.worker = WorkerThread(self._execute_sudo_commands_thread, args=(command, sudo_password, "MTU Change", f"MTU for {iface} set to {mtu}."))
        self.worker.start()

    def _apply_hostname(self, sudo_password=None):
        """Temporarily changes the system's hostname."""
        new_hostname = self.new_hostname_edit.text().strip()
        if not new_hostname:
            QMessageBox.warning(self, "Input Error", "New hostname cannot be empty.")
            return

        if not sudo_password and sys.platform != "win32":
            self._run_command_with_sudo_prompt([], self._apply_hostname, 'apply_hostname')
            return

        command = [["hostname", new_hostname]]
        self.worker = WorkerThread(self._execute_sudo_commands_thread, args=(command, sudo_password, "Hostname Change", "Hostname temporarily changed."))
        self.worker.start()

    def _execute_sudo_commands_thread(self, commands, sudo_password, tool_name, success_message):
        """Worker thread to execute a list of commands with sudo."""
        q = self.tool_results_queue
        try:
            for cmd in commands:
                full_command = ["sudo", "-S"] + cmd
                process = subprocess.Popen(full_command, stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
                stdout, _ = process.communicate(input=sudo_password + '\n')
                if process.returncode != 0:
                    raise Exception(f"Command failed: {' '.join(cmd)}\n{stdout}")
            q.put(('tool_finished', tool_name, "Success", success_message))
        except Exception as e:
            q.put(('error', f"{tool_name} Error", str(e)))

    def _execute_sudo_shell_command_thread(self, command, sudo_password, tool_name, success_message):
        """Worker thread to execute a shell command string with sudo."""
        q = self.tool_results_queue
        try:
            # Use sudo -S sh -c to handle commands with redirection
            full_command = ["sudo", "-S", "sh", "-c", command]
            process = subprocess.Popen(full_command, stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
            stdout, _ = process.communicate(input=sudo_password + '\n')
            if process.returncode != 0:
                raise Exception(f"Command failed: {command}\n{stdout}")
            q.put(('tool_finished', tool_name, "Success", success_message))
        except Exception as e:
            q.put(('error', f"{tool_name} Error", str(e)))

    def _resolve_targets_string(self, targets_string):
        """
        Resolves hostnames within a space-separated string of targets.
        Keeps IP addresses, CIDR ranges, and other special targets as-is.
        """
        if not self.doh_resolver_url:
            return targets_string

        parts = targets_string.split()
        resolved_parts = []
        resolution_map = {} # To show the user what was resolved

        for part in parts:
            try:
                # Check if it's an IP, network, or something that shouldn't be resolved.
                # strict=False allows for hosts in network definitions (e.g., 192.168.1.1/24)
                ipaddress.ip_network(part, strict=False)
                resolved_parts.append(part)
            except ValueError:
                # It's likely a hostname, so try to resolve it
                resolved_ip = self._resolve_hostname(part)
                resolved_parts.append(resolved_ip)
                if resolved_ip != part: # Only add to map if a change occurred
                    resolution_map[part] = resolved_ip

        # Log the resolutions to the appropriate console
        if resolution_map:
            # This is a bit of a hack to find the current visible console.
            # A better long-term solution might involve a dedicated logging area for this.
            current_tab = self.tab_widget.currentWidget()
            console = current_tab.findChild(QPlainTextEdit)
            if console:
                 console.insertPlainText(f"# Secure DNS Resolutions: {resolution_map}\n")

        return " ".join(resolved_parts)

    def _resolve_hostname(self, hostname):
        """Resolves a single hostname to an IP address, using the secure resolver if configured."""
        if not self.doh_resolver_url:
            return hostname

        # Quick check to avoid trying to resolve IPs or invalid hostnames
        try:
            ipaddress.ip_address(hostname)
            return hostname # It's already an IP
        except ValueError:
            pass # It's not an IP, so proceed with resolution

        try:
            import dns.message
            import dns.query
            import dns.rdatatype
            q = dns.message.make_query(hostname, dns.rdatatype.A)
            r = dns.query.https(q, self.doh_resolver_url, timeout=5)
            if r.answer:
                # Return the first A record found
                for answer in r.answer:
                    if answer.rdtype == dns.rdatatype.A:
                        resolved_ip = answer[0].to_text()
                        logging.info(f"DoH resolved '{hostname}' to '{resolved_ip}'")
                        return resolved_ip
            logging.warning(f"DoH resolution for '{hostname}' found no A record.")
            return hostname # Fallback to original hostname if no answer
        except Exception as e:
            logging.error(f"DoH resolution for '{hostname}' failed: {e}")
            self.tool_results_queue.put(('error', 'DNS Error', f"Secure DNS lookup for '{hostname}' failed: {e}"))
            return hostname # Fallback on error


    def _create_l2tp_client_tab(self):
        """Creates the UI for the L2TP/IPsec VPN client."""
        widget = QWidget()
        layout = QFormLayout(widget)

        self.l2tp_server_edit = QLineEdit()
        self.l2tp_server_edit.setPlaceholderText("e.g., vpn.example.com")
        layout.addRow("Gateway:", self.l2tp_server_edit)

        self.l2tp_user_edit = QLineEdit()
        layout.addRow("Username:", self.l2tp_user_edit)

        self.l2tp_pass_edit = QLineEdit()
        self.l2tp_pass_edit.setEchoMode(QLineEdit.EchoMode.Password)
        layout.addRow("Password:", self.l2tp_pass_edit)

        self.l2tp_psk_edit = QLineEdit()
        self.l2tp_psk_edit.setEchoMode(QLineEdit.EchoMode.Password)
        layout.addRow("IPsec PSK:", self.l2tp_psk_edit)

        self.l2tp_connect_btn = QPushButton("Connect")
        self.l2tp_status_label = QLabel("Status: Disconnected")
        layout.addRow(self.l2tp_connect_btn)
        layout.addRow(self.l2tp_status_label)

        self.l2tp_connect_btn.clicked.connect(lambda: self._toggle_l2tp_connection())

        return widget

    def _disconnect_l2tp(self, sudo_password=None):
        """Stops the L2TP/IPsec connection and cleans up configs."""
        # This is a simplified cleanup. A robust implementation would ensure
        # the processes are truly stopped before removing files.
        if not sudo_password and sys.platform != "win32":
            self._run_command_with_sudo_prompt(None, self._disconnect_l2tp, 'l2tp_vpn_disconnect')
            return

        self.l2tp_status_label.setText("Status: Disconnecting...")
        QApplication.processEvents()

        try:
            # Bring down the IPsec tunnel first
            self._execute_l2tp_command(["ipsec", "auto", "--down", "L2TP-PSK-Zurvan"], sudo_password, "ipsec_down_disconnect")
            # Stop the xl2tpd daemon
            self._execute_l2tp_command(["systemctl", "stop", "xl2tpd"], sudo_password, "xl2tpd_stop_disconnect")
            # Clean up the configuration files
            self._cleanup_l2tp_configs(sudo_password)

            self.l2tp_status_label.setText("Status: Disconnected")
            self.l2tp_connect_btn.setText("Connect")
            self.l2tp_process = None # Signal that the process is stopped
        except Exception as e:
            self.l2tp_status_label.setText(f"Status: <font color='red'>Disconnect Error</font>")
            logging.error(f"Error during L2TP disconnect: {e}", exc_info=True)


    def _cleanup_l2tp_configs(self, sudo_password):
        """Removes the temporary L2TP/IPsec configuration files using sudo."""
        if not hasattr(self, 'l2tp_config_files'):
            return

        commands = []
        for path in self.l2tp_config_files.values():
            commands.append(["rm", "-f", path])

        # Also try to remove the include from the main ipsec.conf
        commands.append(["sed", "-i", f"'/include {self.l2tp_config_files['ipsec.conf']}/d'", "/etc/ipsec.conf"])

        for cmd in commands:
            try:
                self._execute_l2tp_command(cmd, sudo_password, "cleanup")
            except Exception as e:
                logging.warning(f"Failed to cleanup L2TP config file {cmd[-1]}: {e}")

    def _execute_l2tp_command(self, command, sudo_password, step):
        """Executes a single command for the L2TP connection process."""
        q = self.tool_results_queue
        full_command = ["sudo", "-S"] + command

        try:
            process = subprocess.Popen(full_command, stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
            if sudo_password:
                stdout, _ = process.communicate(input=sudo_password + '\n')
            else:
                 stdout, _ = process.communicate()

            if process.returncode != 0:
                raise Exception(f"Command failed with code {process.returncode}:\n{stdout}")

            q.put(('l2tp_step_finished', step, sudo_password))

        except Exception as e:
            q.put(('error', 'L2TP/IPsec Error', str(e)))
            q.put(('l2tp_status', f"Status: <font color='red'>Error at step: {step}</font>"))

    def _create_tun_system_tab(self):
        """Creates the UI for the TUN System."""
        widget = QWidget()
        layout = QFormLayout(widget)

        self.tun_ip_edit = QLineEdit("10.0.0.1")
        layout.addRow("TUN IP Address:", self.tun_ip_edit)

        self.tun_netmask_edit = QLineEdit("255.255.255.0")
        layout.addRow("TUN Netmask:", self.tun_netmask_edit)

        self.tun_create_btn = QPushButton("Create TUN Device")
        self.tun_status_label = QLabel("Status: Inactive")
        layout.addRow(self.tun_create_btn)
        layout.addRow(self.tun_status_label)

        self.tun_create_btn.clicked.connect(self._toggle_tun_device)

        return widget

    def _toggle_tun_device(self):
        """Creates or destroys the TUN device."""
        if hasattr(self, 'tun_fd') and self.tun_fd:
            self._destroy_tun_device()
        else:
            self._create_tun_device()

    def _create_tun_device(self):
        """Creates and configures a TUN device."""
        if not self._require_root("TUN System"):
            return

        try:
            import fcntl
            import struct

            TUNSETIFF = 0x400454ca
            IFF_TUN = 0x0001
            IFF_NO_PI = 0x1000

            # Create the TUN device
            self.tun_fd = os.open("/dev/net/tun", os.O_RDWR)
            ifr = struct.pack('16sH', b'tun%d', IFF_TUN | IFF_NO_PI)
            self.ifname = fcntl.ioctl(self.tun_fd, TUNSETIFF, ifr)
            self.ifname = self.ifname.decode('utf-8').strip('\x00')

            self.tun_status_label.setText(f"Status: Created {self.ifname}")

            # Configure the TUN device
            ip_addr = self.tun_ip_edit.text()
            netmask = self.tun_netmask_edit.text()
            subprocess.check_call(f"ip addr add {ip_addr}/{netmask} dev {self.ifname}", shell=True)
            subprocess.check_call(f"ip link set dev {self.ifname} up", shell=True)

            self.tun_create_btn.setText("Destroy TUN Device")
            self.tun_status_label.setText(f"Status: {self.ifname} is up")

        except Exception as e:
            self.tun_status_label.setText(f"Status: <font color='red'>Error: {e}</font>")
            if hasattr(self, 'tun_fd') and self.tun_fd:
                os.close(self.tun_fd)
                self.tun_fd = None

    def _destroy_tun_device(self):
        """Destroys the TUN device."""
        if hasattr(self, 'tun_fd') and self.tun_fd:
            try:
                os.close(self.tun_fd)
                self.tun_fd = None
                self.tun_create_btn.setText("Create TUN Device")
                self.tun_status_label.setText("Status: Inactive")
            except Exception as e:
                self.tun_status_label.setText(f"Status: <font color='red'>Error: {e}</font>")

    def _handle_l2tp_step(self, step, sudo_password):
        """Handles the multi-step connection process for L2TP."""
        if step == "ipsec_restart":
            self.l2tp_status_label.setText("Status: IPsec service restarted. Bringing up tunnel...")
            up_cmd = ["ipsec", "auto", "--up", "L2TP-PSK-Zurvan"]
            self.worker = WorkerThread(self._execute_l2tp_command, args=(up_cmd, sudo_password, "ipsec_up"))
            self.worker.start()

        elif step == "ipsec_up":
            self.l2tp_status_label.setText("Status: IPsec tunnel established. Starting L2TP daemon...")
            # Now start xl2tpd in the background
            try:
                # We need to run xl2tpd with sudo, but Popen needs careful handling
                # This is a simplified approach; a production app might use a helper script
                # to manage the daemon.
                xl2tpd_cmd = ["sudo", "-S", "xl2tpd", "-c", self.l2tp_config_files['xl2tpd.conf'], "-D"]
                self.l2tp_process = subprocess.Popen(xl2tpd_cmd, stdin=subprocess.PIPE, text=True)
                self.l2tp_process.stdin.write(sudo_password + '\n')
                self.l2tp_process.stdin.flush()

                # A small delay to allow the daemon to start
                time.sleep(2)

                if self.l2tp_process.poll() is None:
                    self.l2tp_status_label.setText("Status: <font color='green'>Connected</font>")
                    self.l2tp_connect_btn.setText("Disconnect")
                else:
                    raise Exception("xl2tpd daemon failed to start.")

            except Exception as e:
                self.tool_results_queue.put(('error', 'L2TP/IPsec Error', f"Failed to start xl2tpd: {e}"))
                self.l2tp_status_label.setText(f"Status: <font color='red'>Error starting L2TP daemon</font>")

        elif step == "ipsec_down":
             self.l2tp_status_label.setText("Status: Disconnected")
             self.l2tp_connect_btn.setText("Connect")

    def _create_openvpn_client_tab(self):
        """Creates the UI for the OpenVPN client."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_box = QGroupBox("Configuration")
        form_layout = QFormLayout(config_box)

        config_layout = QHBoxLayout()
        self.openvpn_config_edit = QLineEdit()
        self.openvpn_config_edit.setPlaceholderText("Path to .ovpn file...")
        config_layout.addWidget(self.openvpn_config_edit)
        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(lambda: self._browse_file_for_lineedit(self.openvpn_config_edit, "Select OpenVPN Config File", "OpenVPN Files (*.ovpn)"))
        config_layout.addWidget(browse_btn)
        form_layout.addRow("Config File:", config_layout)

        self.openvpn_connect_btn = QPushButton("Connect")
        self.openvpn_status_label = QLabel("Status: Disconnected")
        form_layout.addRow(self.openvpn_connect_btn)
        form_layout.addRow(self.openvpn_status_label)
        layout.addWidget(config_box)

        log_box = QGroupBox("OpenVPN Logs")
        log_layout = QVBoxLayout(log_box)
        self.openvpn_output_console = QPlainTextEdit()
        self.openvpn_output_console.setReadOnly(True)
        self.openvpn_output_console.setFont(QFont("Courier New", 10))
        log_layout.addWidget(self.openvpn_output_console)
        layout.addWidget(log_box)

        self.openvpn_connect_btn.clicked.connect(self._toggle_openvpn_connection)

        return widget

    def _create_wireguard_client_tab(self):
        """Creates the UI for the WireGuard client."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_box = QGroupBox("Configuration")
        form_layout = QFormLayout(config_box)


        config_layout = QHBoxLayout()
        self.wireguard_config_edit = QLineEdit()
        self.wireguard_config_edit.setPlaceholderText("Path to .conf file...")
        config_layout.addWidget(self.wireguard_config_edit)
        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(lambda: self._browse_file_for_lineedit(self.wireguard_config_edit, "Select WireGuard Config File", "Config Files (*.conf)"))
        config_layout.addWidget(browse_btn)
        form_layout.addRow("Config File:", config_layout)

        self.wireguard_connect_btn = QPushButton("Connect")
        self.wireguard_status_label = QLabel("Status: Disconnected")
        form_layout.addRow(self.wireguard_connect_btn)
        form_layout.addRow(self.wireguard_status_label)
        layout.addWidget(config_box)

        log_box = QGroupBox("WireGuard Logs")
        log_layout = QVBoxLayout(log_box)
        self.wireguard_output_console = QPlainTextEdit()
        self.wireguard_output_console.setReadOnly(True)
        self.wireguard_output_console.setFont(QFont("Courier New", 10))
        log_layout.addWidget(self.wireguard_output_console)
        layout.addWidget(log_box)

        self.wireguard_connect_btn.clicked.connect(self._toggle_wireguard_connection)

        return widget

    def _create_v2ray_manager_tab(self):
        """Creates the UI for the V2Ray Manager, inspired by v2rayN."""
        widget = QWidget()
        layout = QHBoxLayout(widget)

        # --- Left Panel: Server List ---
        left_panel = QGroupBox("Servers")
        left_layout = QVBoxLayout(left_panel)

        self.v2ray_server_list = QTreeWidget()
        self.v2ray_server_list.setColumnCount(5)
        self.v2ray_server_list.setHeaderLabels(["Alias", "Address", "Port", "Type", "Latency"])
        left_layout.addWidget(self.v2ray_server_list)
        layout.addWidget(left_panel, 2) # Give it more stretch

        # --- Right Panel: Controls & Settings ---
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)

        # Connection Controls
        connection_box = QGroupBox("Connection")
        connection_layout = QFormLayout(connection_box)
        self.v2ray_connect_btn = QPushButton("Connect")
        self.v2ray_connect_btn.clicked.connect(self._toggle_v2ray_connection)
        self.v2ray_status_label = QLabel("Status: <font color='red'>Disconnected</font>")
        connection_layout.addRow(self.v2ray_connect_btn)
        connection_layout.addRow(self.v2ray_status_label)
        right_layout.addWidget(connection_box)

        # Server Management
        server_mg_box = QGroupBox("Server Management")
        server_mg_layout = QVBoxLayout(server_mg_box)
        add_server_btn = QPushButton("Add Server (Manual)")
        add_server_btn.setToolTip("Manually add a server from a Vmess link or config details.")
        add_server_btn.clicked.connect(self._add_v2ray_server)
        remove_server_btn = QPushButton("Remove Selected Server")
        remove_server_btn.setToolTip("Remove the currently selected server from the list.")
        remove_server_btn.clicked.connect(self._remove_v2ray_server)
        test_server_btn = QPushButton("Test Selected Server")
        test_server_btn.setToolTip("Test the latency of the selected server.")
        test_server_btn.clicked.connect(self._test_v2ray_connection)
        server_mg_layout.addWidget(add_server_btn)
        server_mg_layout.addWidget(remove_server_btn)
        server_mg_layout.addWidget(test_server_btn)
        right_layout.addWidget(server_mg_box)

        # Subscription Management
        sub_mg_box = QGroupBox("Subscriptions")
        sub_mg_layout = QFormLayout(sub_mg_box)
        self.v2ray_sub_url_edit = QLineEdit()
        self.v2ray_sub_url_edit.setPlaceholderText("Enter subscription URL...")
        add_sub_btn = QPushButton("Add/Update Subscription")
        add_sub_btn.setToolTip("Add a new subscription or update an existing one from the URL.")
        add_sub_btn.clicked.connect(self._update_v2ray_subscription)
        sub_mg_layout.addRow("URL:", self.v2ray_sub_url_edit)
        sub_mg_layout.addRow(add_sub_btn)
        right_layout.addWidget(sub_mg_box)

        # --- Core Selection and Tuning ---
        core_box = QGroupBox("Core and Tuning")
        core_layout = QFormLayout(core_box)
        self.v2ray_core_combo = QComboBox()
        self.v2ray_core_combo.addItems(["v2ray", "xray", "mihomo", "sing-box"])
        self.v2ray_core_combo.setToolTip("Select the core executable to use.")
        core_layout.addRow("Core Engine:", self.v2ray_core_combo)

        core_file_layout = QHBoxLayout()
        self.v2ray_core_edit = QLineEdit()
        self.v2ray_core_edit.setPlaceholderText("Path to selected core executable...")
        core_file_layout.addWidget(self.v2ray_core_edit)
        browse_core_btn = QPushButton("Browse...")
        browse_core_btn.clicked.connect(self._browse_v2ray_core)
        core_file_layout.addWidget(browse_core_btn)
        download_core_btn = QPushButton("Download Core")
        download_core_btn.setToolTip("Download the latest Xray core automatically.")
        download_core_btn.clicked.connect(self._download_v2ray_core)
        core_file_layout.addWidget(download_core_btn)
        core_layout.addRow("Core Path:", core_file_layout)

        # --- Tuning ---
        self.v2ray_mux_check = QCheckBox("Enable Mux (Multiplexing)")
        self.v2ray_mux_check.setChecked(True)
        core_layout.addRow(self.v2ray_mux_check)
        self.v2ray_mux_concurrency_edit = QLineEdit("8")
        self.v2ray_mux_concurrency_label = QLabel("Concurrency:")
        core_layout.addRow(self.v2ray_mux_concurrency_label, self.v2ray_mux_concurrency_edit)
        self.v2ray_mux_check.toggled.connect(self.v2ray_mux_concurrency_label.setVisible)
        self.v2ray_mux_check.toggled.connect(self.v2ray_mux_concurrency_edit.setVisible)

        self.v2ray_tune_check = QCheckBox("Enable Performance Tuning")
        self.v2ray_tune_check.setToolTip("Enables experimental performance and anti-censorship tuning options like TCP Fast Open and UDP fragmentation.")
        core_layout.addRow(self.v2ray_tune_check)

        right_layout.addWidget(core_box)

        right_layout.addStretch()
        layout.addWidget(right_panel, 1)

        return widget

    def _browse_v2ray_core(self):
        """Opens a file dialog to select the V2Ray core executable."""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select V2Ray/Xray Executable", "", "All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.v2ray_core_edit.setText(file_path)

    def _download_v2ray_core(self):
        """Downloads the latest Xray-core for the current OS and places it in tools/xray/."""
        import platform
        import zipfile

        system = platform.system().lower()
        machine = platform.machine().lower()

        # Map machine types
        arch_map = {
            'x86_64': '64',
            'amd64': '64',
            'i386': '32',
            'i686': '32',
            'arm64': 'arm64-v8a',
            'aarch64': 'arm64-v8a',
            'armv7l': 'arm32-v7a',
        }
        arch = arch_map.get(machine, '64')

        # --- مهم: Xray همه چیز را به صورت ZIP منتشر می‌کند ---
        if system == 'windows':
            asset_name = f"Xray-windows-{arch}.zip"
        elif system == 'darwin':
            asset_name = f"Xray-macos-{arch}.zip"
        else:  # Linux
            asset_name = f"Xray-linux-{arch}.zip"

        try:
            # Get latest release info
            release_url = "https://api.github.com/repos/XTLS/Xray-core/releases/latest"
            with urllib.request.urlopen(release_url) as resp:
                release_data = json.loads(resp.read().decode())
            assets = release_data['assets']
            download_url = None
            for asset in assets:
                if asset['name'] == asset_name:
                    download_url = asset['browser_download_url']
                    break
            if not download_url:
                QMessageBox.critical(self, "Download Error", f"No matching asset found for {asset_name}")
                return

            # Prepare destination
            xray_dir = os.path.join(self.script_dir, "tools", "xray")
            os.makedirs(xray_dir, exist_ok=True)
            archive_path = os.path.join(xray_dir, "xray_core.zip")
            exe_path = os.path.join(xray_dir, "xray" + (".exe" if system == "windows" else ""))

            # Download
            self.status_bar.showMessage("Downloading Xray core...")
            with urllib.request.urlopen(download_url) as response, open(archive_path, 'wb') as out_file:
                out_file.write(response.read())

            # Extract (همه چیز ZIP است)
            with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                zip_ref.extractall(xray_dir)

            # Make executable on Unix
            if system != 'windows':
                os.chmod(exe_path, 0o755)

            # Update UI
            self.v2ray_core_edit.setText(os.path.abspath(exe_path))
            self.v2ray_core_combo.setCurrentText("xray")
            self.status_bar.showMessage("Xray core downloaded and ready!", 5000)
            QMessageBox.information(self, "Success", f"Xray core downloaded to:\n{exe_path}")

        except Exception as e:
            logging.error(f"Failed to download Xray core: {e}", exc_info=True)
            QMessageBox.critical(self, "Download Error", f"Failed to download core:\n{str(e)}")
            self.status_bar.showMessage("Download failed.", 5000)

    def _toggle_v2ray_connection(self):
        """Starts or stops the V2Ray process based on the selected server."""
        if self.v2ray_process and self.v2ray_process.poll() is None:
            # --- Disconnect Logic ---
            self.v2ray_process.terminate()
            try:
                self.v2ray_process.wait(timeout=3)
            except subprocess.TimeoutExpired:
                self.v2ray_process.kill()
                self.v2ray_process.wait()
            self.v2ray_process = None
            self.v2ray_status_label.setText("Status: <font color='red'>Disconnected</font>")
            self.v2ray_connect_btn.setText("Connect")
            return

        # --- Connect Logic ---
        selected_item = self.v2ray_server_list.currentItem()
        if not selected_item:
            QMessageBox.warning(self, "No Server Selected", "Please select a server from the list to connect.")
            return

        # The full config is stored as user data in the item
        server_config = selected_item.data(0, Qt.ItemDataRole.UserRole)
        if not server_config:
            QMessageBox.critical(self, "Error", "Selected server has no configuration associated with it.")
            return

        core_engine = self.v2ray_core_combo.currentText()
        core_executable = self.v2ray_core_edit.text()
        if not core_executable or not os.path.isfile(core_executable):
            QMessageBox.critical(self, f"{core_engine.capitalize()} Error", f"'{core_executable}' not found. Please select a valid core executable.")
            return

        # --- Validate the configuration before attempting to run ---
        config_path = None
        try:
            # Add a log object to the config for debugging
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".log", encoding='utf-8') as tmp_log:
                server_config["log"] = {"loglevel": "warning", "log_path": tmp_log.name}

            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_conf:
                json.dump(server_config, tmp_conf)
                config_path = tmp_conf.name

            # Use 'run -test' to validate the config file
            test_command = [core_executable, "run", "-test", "-c", config_path]
            process = subprocess.run(test_command, capture_output=True, text=True, timeout=10)

            # For xray, a successful test exits with code 0 and prints "Configuration OK." to stderr.
            if process.returncode != 0 or "Configuration OK." not in process.stderr:
                error_output = process.stdout + process.stderr
                QMessageBox.critical(self, "Invalid Configuration", f"The generated configuration for this server is invalid and cannot be started.\n\nError:\n{error_output}")
                return

        except Exception as e:
            QMessageBox.critical(self, "Config Error", f"Could not create or validate temporary config file: {e}")
            return
        finally:
            # Clean up the temp file if validation failed, but keep it for the run command if it succeeded
            if 'process' in locals() and process.returncode != 0 and config_path and os.path.exists(config_path):
                 os.remove(config_path)


        # --- Build and run the final command ---
        command = [core_executable, "run", "-c", config_path]
        try:
            # Capture stderr to provide feedback on failure
            self.v2ray_process = subprocess.Popen(command, stderr=subprocess.PIPE, stdout=subprocess.DEVNULL, text=True, encoding='utf-8', errors='replace')
            # Give it a moment to see if it fails immediately
            time.sleep(1)
            if self.v2ray_process.poll() is not None:
                error_output = self.v2ray_process.stderr.read()
                raise Exception(f"Core process terminated unexpectedly.\n\nError:\n{error_output}")

            self.v2ray_status_label.setText(f"Status: <font color='green'>Connected</font> to {selected_item.text(0)}")
            self.v2ray_connect_btn.setText("Disconnect")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to start {core_engine}: {e}")
            self.v2ray_process = None
            if config_path and os.path.exists(config_path):
                os.remove(config_path) # Clean up on failure

    def _update_v2ray_subscription(self):
        """Fetches and parses a V2Ray subscription link to populate the server list."""
        sub_url = self.v2ray_sub_url_edit.text().strip()
        if not sub_url:
            QMessageBox.warning(self, "Input Error", "Please enter a subscription URL.")
            return

        try:
            self.v2ray_status_label.setText("Status: <font color='orange'>Updating...</font>")
            QApplication.processEvents() # Force UI update

            # Use requests for better handling of headers and potential issues
            import requests
            response = requests.get(sub_url, timeout=10)
            response.raise_for_status() # Raise an exception for bad status codes

            # The subscription content is base64 encoded
            decoded_content = base64.b64decode(response.text).decode('utf-8')

            # Clear existing list before adding new servers
            self.v2ray_server_list.clear()

            server_links = decoded_content.strip().split('\n')
            for link in server_links:
                link = link.strip()
                if link.startswith("vmess://") or link.startswith("vless://"):
                    try:
                        is_vless = link.startswith("vless://")
                        if is_vless:
                            server_data = self._parse_vless_link(link)
                        else:
                            server_data = self._parse_vmess_link(link)

                        if not server_data:
                            continue

                        client_config = self._create_v2ray_client_config(server_data, is_vless=is_vless)

                        item = QTreeWidgetItem([
                            server_data.get('ps', 'N/A'),
                            server_data.get('add', 'N/A'),
                            str(server_data.get('port', 'N/A')),
                            "vless" if is_vless else "vmess"
                        ])
                        item.setData(0, Qt.ItemDataRole.UserRole, client_config)
                        self.v2ray_server_list.addTopLevelItem(item)

                    except Exception as e:
                        logging.warning(f"Failed to parse V2Ray link '{link[:30]}...': {e}")

            self.v2ray_server_list.resizeColumnToContents(0)
            self.v2ray_server_list.resizeColumnToContents(1)
            self.v2ray_status_label.setText("Status: <font color='green'>Subscription Updated</font>")
            self._save_v2ray_servers()

        except Exception as e:
            self.v2ray_status_label.setText("Status: <font color='red'>Update Failed</font>")
            QMessageBox.critical(self, "Subscription Error", f"Failed to update subscription:\n{e}")
            logging.error(f"Failed to update V2Ray subscription: {e}", exc_info=True)

    def _parse_vmess_link(self, link):
        """Parses a vmess:// link and returns a dictionary of its components."""
        if not link.startswith("vmess://"):
            return None

        # Decode the base64 part of the link
        encoded_part = link[len("vmess://"):]
        decoded_json = base64.b64decode(encoded_part).decode('utf-8')
        return json.loads(decoded_json)

    def _parse_vless_link(self, link):
        """Parses a vless:// link and returns a dictionary of its components."""
        if not link.startswith("vless://"):
            return None

        # VLESS format: vless://uuid@host:port?params...#alias

        # Separate the alias
        parts = link.split("#", 1)
        alias = urllib.parse.unquote(parts[1]) if len(parts) > 1 else "N/A"

        # Parse the main part
        main_part = parts[0][len("vless://"):]
        uri = urllib.parse.urlparse(f"vless://{main_part}")

        params = urllib.parse.parse_qs(uri.query)

        vless_data = {
            "ps": alias,
            "id": uri.username,
            "add": uri.hostname,
            "port": uri.port,
            "net": params.get("type", ["tcp"])[0],
            "tls": params.get("security", [""])[0],
            "path": urllib.parse.unquote(params.get("path", ["/"])[0]),
            "host": params.get("host", [uri.hostname])[0],
            "sni": params.get("sni", [uri.hostname])[0],
            "alpn": params.get("alpn", [""])[0],
            "fp": params.get("fp", [""])[0],
            "allowInsecure": params.get("allowInsecure", ["0"])[0] == "1",
        }
        return vless_data

    def _create_v2ray_client_config(self, server_data, is_vless=False):
        """Creates a full V2Ray client JSON configuration from parsed server data."""
        # This is a template for a basic V2Ray client config
        # It sets up a SOCKS5 inbound on port 10808 and routes traffic through the specified outbound

        outbound_protocol = "vless" if is_vless else "vmess"

        # User object is different for vless vs vmess
        if is_vless:
            user_object = {
                "id": server_data.get("id", ""),
                "encryption": "none",
                "flow": ""
            }
        else: # VMess
            user_object = {
                "id": server_data.get("id", ""),
                "alterId": int(server_data.get("aid", 0)),
                "security": "auto"
            }

        config = {
            "inbounds": [{
                "port": 10808,  # Local SOCKS port
                "listen": "127.0.0.1",
                "protocol": "socks",
                "settings": { "auth": "noauth", "udp": True }
            }],
            "outbounds": [{
                "protocol": outbound_protocol,
                "settings": {
                    "vnext": [{
                        "address": server_data.get("add", ""),
                        "port": int(server_data.get("port", 443)),
                        "users": [user_object]
                    }]
                },
                "streamSettings": {
                    "network": server_data.get("net", "tcp"),
                    "security": server_data.get("tls", ""),
                },
                "mux": {
                    "enabled": self.v2ray_mux_check.isChecked(),
                    "concurrency": int(self.v2ray_mux_concurrency_edit.text()) if self.v2ray_mux_check.isChecked() else -1
                }
            }]
        }

        # Add network-specific settings (ws, tcp, etc.)
        net_type = server_data.get("net", "tcp")
        if net_type == "ws":
            config["outbounds"][0]["streamSettings"]["wsSettings"] = {
                "path": server_data.get("path", "/"),
                "headers": { "Host": server_data.get("host", "") }
            }

        # Add TLS settings if applicable
        if server_data.get("tls", "") in ["tls", "reality"]:
            tls_settings = {
                "serverName": server_data.get("sni", server_data.get("add", "")),
                "allowInsecure": server_data.get("allowInsecure", False)
            }
            if server_data.get("fp"):
                tls_settings["fingerprint"] = server_data.get("fp")
            if server_data.get("alpn"):
                # ALPN can be a comma-separated string
                tls_settings["alpn"] = [h.strip() for h in server_data["alpn"].split(',')]

            config["outbounds"][0]["streamSettings"]["tlsSettings"] = tls_settings

        # Add experimental "Tune" settings if enabled
        if self.v2ray_tune_check.isChecked():
            config["outbounds"][0]["settings"]["tuning"] = {
                "tcpFastOpen": True,
                "udpFragment": True
            }

        return config

    def _add_v2ray_server_from_link(self, link):
        """Adds a V2Ray server from a vmess/vless link string."""
        if not (link.startswith("vmess://") or link.startswith("vless://")):
            return

        try:
            is_vless = link.startswith("vless://")
            if is_vless:
                server_data = self._parse_vless_link(link)
            else:
                server_data = self._parse_vmess_link(link)

            if not server_data:
                raise ValueError("Parsed data is empty.")

            client_config = self._create_v2ray_client_config(server_data, is_vless=is_vless)

            item = QTreeWidgetItem([
                server_data.get('ps', 'N/A'),
                server_data.get('add', 'N/A'),
                str(server_data.get('port', 'N/A')),
                "vless" if is_vless else "vmess"
            ])
            item.setData(0, Qt.ItemDataRole.UserRole, client_config)
            self.v2ray_server_list.addTopLevelItem(item)
            self._save_v2ray_servers()
        except Exception as e:
            # Log the error instead of showing a dialog for automated use
            logging.error(f"Failed to parse the provided V2Ray link: {e}")

    def _add_v2ray_server(self):
        """Opens a dialog to manually add a V2Ray server from a vmess link."""
        vmess_link, ok = QInputDialog.getMultiLineText(self, "Add V2Ray Server", "Paste vmess:// or vless:// link here:")
        if ok and vmess_link.strip():
            self._add_v2ray_server_from_link(vmess_link.strip())

    def _remove_v2ray_server(self):
        """Removes the currently selected server from the V2Ray list."""
        selected_items = self.v2ray_server_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "No Server Selected", "Please select a server to remove.")
            return

        reply = QMessageBox.question(self, "Confirm Deletion", "Are you sure you want to remove the selected server?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            for item in selected_items:
                self.v2ray_server_list.takeTopLevelItem(self.v2ray_server_list.indexOfTopLevelItem(item))
            self._save_v2ray_servers()

    def _save_v2ray_servers(self):
        """Saves the current list of V2Ray servers to a JSON file."""
        servers_data = []
        for i in range(self.v2ray_server_list.topLevelItemCount()):
            item = self.v2ray_server_list.topLevelItem(i)
            config = item.data(0, Qt.ItemDataRole.UserRole)
            if config:
                servers_data.append(config)

        try:
            with open("v2ray_servers.json", "w") as f:
                json.dump(servers_data, f, indent=4)
            logging.info(f"Saved {len(servers_data)} V2Ray servers to v2ray_servers.json")
        except Exception as e:
            logging.error(f"Failed to save V2Ray servers: {e}", exc_info=True)
            # Optionally, inform the user
            # self.v2ray_status_label.setText("<font color='red'>Error saving servers!</font>")


    def _load_v2ray_servers(self):
        """Loads V2Ray servers from the JSON file on startup."""
        try:
            if not os.path.exists("v2ray_servers.json"):
                return # No file to load

            with open("v2ray_servers.json", "r") as f:
                servers_data = json.load(f)

            self.v2ray_server_list.clear()
            for config in servers_data:
                # We need to extract display info from the config
                outbound_settings = config.get("outbounds", [{}])[0].get("settings", {})
                vnext = outbound_settings.get("vnext", [{}])[0]
                # The alias can be in different places depending on the link type
                # A more robust way is to check the user object
                user_object = vnext.get("users", [{}])[0]
                alias = user_object.get("ps", "N/A")

                item = QTreeWidgetItem([
                    alias,
                    vnext.get("address", "N/A"),
                    str(vnext.get("port", "N/A")),
                    config.get("outbounds", [{}])[0].get("protocol", "N/A"),
                    "N/A" # Latency column
                ])
                item.setData(0, Qt.ItemDataRole.UserRole, config)
                self.v2ray_server_list.addTopLevelItem(item)

            logging.info(f"Loaded {len(servers_data)} V2Ray servers from v2ray_servers.json")

        except Exception as e:
            logging.error(f"Failed to load V2Ray servers: {e}", exc_info=True)
            QMessageBox.critical(self, "Load Error", "Failed to load V2Ray server configurations. The file might be corrupted.")

    def _test_v2ray_connection(self):
        """Starts a background thread to test the latency of the selected V2Ray server."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        selected_item = self.v2ray_server_list.currentItem()
        if not selected_item:
            QMessageBox.warning(self, "No Server Selected", "Please select a server to test.")
            return

        self.is_tool_running = True
        # Store original text and update UI before starting the thread
        original_text = selected_item.text(0)
        selected_item.setText(0, f"[Testing...] {original_text}")
        selected_item.setText(4, "...") # Latency column

        self.worker = WorkerThread(self._v2ray_test_thread, args=(selected_item,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _v2ray_test_thread(self, item):
        """
        Worker thread to test a single V2Ray connection.
        It starts an Xray instance, polls for the proxy port, and measures latency.
        """
        q = self.tool_results_queue
        # The original text might contain "[Testing...]". We need the clean alias.
        original_item_text_with_status = item.text(0)
        clean_alias = original_item_text_with_status.replace("[Testing...] ", "")
        server_config = item.data(0, Qt.ItemDataRole.UserRole)
        v2ray_executable = self.v2ray_core_edit.text().strip()

        if not v2ray_executable or not os.path.isfile(v2ray_executable):
            q.put(('v2ray_test_result', clean_alias, "<font color='red'>Core not found</font>"))
            q.put(('tool_finished', 'v2ray_test'))
            return

        process = None
        temp_config_path = None
        error_msg = "Test failed"

        try:
            # 1. Create temporary config file
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_conf:
                json.dump(server_config, tmp_conf, indent=2)
                temp_config_path = tmp_conf.name

            # 2. Start the xray 'run' process
            logging.info(f"Starting V2Ray test for '{clean_alias}' with config: {temp_config_path}")
            command = [v2ray_executable, "run", "-c", temp_config_path]
            # Capture stderr to provide feedback on failure
            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding='utf-8', errors='replace')

            # 3. Poll for the local SOCKS proxy to be ready
            proxy_ready = False
            for _ in range(20):  # Try for 10 seconds (20 * 0.5s)
                if self.tool_stop_event.is_set(): raise Exception("Test cancelled by user")
                try:
                    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                        s.settimeout(0.5)
                        s.connect(("127.0.0.1", 10808))
                    proxy_ready = True
                    break
                except (socket.timeout, ConnectionRefusedError):
                    time.sleep(0.5)

            if not proxy_ready:
                raise Exception("Proxy failed to start within 10 seconds.")

            # 4. Perform latency test
            proxies = {"http": "socks5://127.0.0.1:10808", "https": "socks5://127.0.0.1:10808"}
            # This is a standard URL used by Google for connectivity checks, it's lightweight.
            test_url = "http://www.gstatic.com/generate_204"

            start_time = time.time()
            response = requests.get(test_url, proxies=proxies, timeout=8)
            response.raise_for_status() # Raise an exception for non-2xx status codes

            end_time = time.time()
            duration_ms = (end_time - start_time) * 1000

            result_text = f"<font color='green'>{duration_ms:.0f} ms</font>"
            q.put(('v2ray_test_result', clean_alias, result_text))

        except requests.exceptions.Timeout:
            error_msg = "Timeout"
            q.put(('v2ray_test_result', clean_alias, f"<font color='red'>{error_msg}</font>"))
        except requests.exceptions.ProxyError:
            error_msg = "Proxy Error"
            q.put(('v2ray_test_result', clean_alias, f"<font color='red'>{error_msg}</font>"))
        except Exception as e:
            error_msg = str(e)
            logging.error(f"V2Ray test failed for '{clean_alias}': {error_msg}", exc_info=True)
            q.put(('v2ray_test_result', clean_alias, f"<font color='red'>Failed</font>"))
        finally:
            # 5. Cleanup
            if process:
                process.terminate()
                try:
                    process.wait(timeout=3)
                except subprocess.TimeoutExpired:
                    process.kill()
            if temp_config_path and os.path.exists(temp_config_path):
                try:
                    os.remove(temp_config_path)
                except OSError:
                    pass
            q.put(('tool_finished', 'v2ray_test'))

    def _write_proxy_to_conf(self, conf_file, proxy_str):
        """Helper to parse and write a single proxy line to the config file."""
        try:
            # Use regex to handle various formats like socks5://user:pass@host:port
            match = re.match(r"(\w+)://(?:[^@]+@)?([^:]+):(\d+)", proxy_str)
            if match:
                proto, host, port = match.groups()
                # proxychains format is "type host port"
                conf_file.write(f"{proto} {host} {port}\n")
            else: # Fallback for simple "host:port"
                parts = proxy_str.split(':')
                if len(parts) == 2:
                    # Assume socks5 as default, could be http as well but socks5 is more common for this app
                    conf_file.write(f"socks5 {parts[0]} {parts[1]}\n")
        except Exception as e:
            logging.warning(f"Could not parse proxy string for proxychains.conf: '{proxy_str}': {e}")

    def _create_proxychains_conf(self):
        """Creates a temporary proxychains configuration file and returns its path."""
        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".conf", encoding='utf-8') as tmp_conf:
                is_rotation_active = self.rotation_box.isChecked() and self.current_proxy_index != -1
                chain_type = self.proxy_chain_type_combo.currentText()

                # Always write the selected chain type. If rotation is on, we only list one proxy,
                # so the chain type doesn't have a practical effect but is still required syntax.
                tmp_conf.write(f"{chain_type}\n")

                if chain_type == "random_chain" and self.chain_len_edit.text().strip():
                    tmp_conf.write(f"chain_len = {self.chain_len_edit.text().strip()}\n")

                if self.proxy_dns_check.isChecked():
                    tmp_conf.write("proxy_dns\n")

                # Add timeouts
                if self.tcp_read_timeout_edit.text().strip():
                    tmp_conf.write(f"tcp_read_time_out {self.tcp_read_timeout_edit.text().strip()}\n")
                if self.tcp_connect_timeout_edit.text().strip():
                    tmp_conf.write(f"tcp_connect_time_out {self.tcp_connect_timeout_edit.text().strip()}\n")


                tmp_conf.write("\n[ProxyList]\n")

                # If rotation is active, only write the current proxy
                if is_rotation_active:
                    if self.proxy_list_widget.count() > self.current_proxy_index:
                        proxy_str = self.proxy_list_widget.item(self.current_proxy_index).text()
                        self._write_proxy_to_conf(tmp_conf, proxy_str)
                else:
                    # Otherwise, write all proxies from the list
                    if self.tor_proxy_check.isChecked():
                        tmp_conf.write("socks5 127.0.0.1 9050\n")

                    for i in range(self.proxy_list_widget.count()):
                        proxy_str = self.proxy_list_widget.item(i).text()
                        self._write_proxy_to_conf(tmp_conf, proxy_str)

                return tmp_conf.name
        except Exception as e:
            logging.error(f"Failed to create proxychains config: {e}")
            return None

    def _create_tor_manager_tab(self):
        """Creates the UI for the Tor Manager."""
        widget = QWidget()
        layout = QFormLayout(widget)

        self.tor_proxy_check = QCheckBox("Use Tor Proxy")
        self.tor_proxy_check.setToolTip("Route all compatible tool traffic through a local Tor SOCKS proxy (127.0.0.1:9050).")
        self.tor_status_label = QLabel("Tor Status: [?]")

        layout.addRow(self.tor_proxy_check, self.tor_status_label)

        tor_guide_btn = QPushButton("Tor Installation Guide")
        layout.addRow(tor_guide_btn)

        # --- Connections ---
        self.tor_proxy_check.toggled.connect(self._handle_tor_toggle)
        tor_guide_btn.clicked.connect(self._show_tor_guide_dialog)

        return widget

    def _show_tor_guide_dialog(self):
        """Shows the Tor installation guide dialog."""
        dialog = TorGuideDialog(self)
        dialog.exec()

    def _create_threat_intelligence_tab(self):
        """Creates the tab container for the Threat Intelligence tools."""
        threat_tabs = QTabWidget()
        threat_tabs.addTab(self._create_recent_threats_tab(), "Recent Threats")
        threat_tabs.addTab(self._create_exploit_db_search_tab(), "Exploit-DB Search")

        # Offline database tools
        offline_db_tab = QTabWidget()
        self.threat_intel_cve_manager = OfflineCveManagerWidget(self)
        offline_db_tab.addTab(self.threat_intel_cve_manager, "Import/Update")
        offline_db_tab.addTab(CveViewerWidget(self), "View Database")
        threat_tabs.addTab(offline_db_tab, "Offline CVE Database")

        return threat_tabs

    def _create_recent_threats_tab(self):
        """Creates the UI for the Recent Threats CVE viewer."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Controls ---
        controls_layout = QHBoxLayout()
        self.fetch_threats_btn = QPushButton("Fetch Latest CVEs")
        controls_layout.addWidget(self.fetch_threats_btn)
        controls_layout.addStretch()
        layout.addLayout(controls_layout)

        # --- Splitter for Table and Details ---
        splitter = QSplitter(Qt.Orientation.Vertical)

        # Top part with table and pagination
        top_widget = QWidget()
        top_layout = QVBoxLayout(top_widget)
        top_layout.setContentsMargins(0,0,0,0)

        self.threats_table = QTreeWidget()
        self.threats_table.setColumnCount(4)
        self.threats_table.setHeaderLabels(["CVE ID", "Published", "CVSS", "Summary"])
        self.threats_table.header().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.threats_table.header().setStretchLastSection(True)
        self.threats_table.header().resizeSection(0, 150) # CVE ID
        self.threats_table.header().resizeSection(1, 150) # Published
        self.threats_table.header().resizeSection(2, 50)  # CVSS
        top_layout.addWidget(self.threats_table)

        # Pagination controls
        pagination_layout = QHBoxLayout()
        self.threats_prev_btn = QPushButton("<< Previous")
        self.threats_next_btn = QPushButton("Next >>")
        self.threats_page_label = QLabel("Page 1 / 1")

        items_per_page_combo = QComboBox()
        items_per_page_combo.addItems(["15", "30", "50", "100"])
        items_per_page_combo.setCurrentText(str(self.threats_items_per_page))

        pagination_layout.addWidget(self.threats_prev_btn)
        pagination_layout.addStretch()
        pagination_layout.addWidget(self.threats_page_label)
        pagination_layout.addWidget(QLabel("Items per page:"))
        pagination_layout.addWidget(items_per_page_combo)
        pagination_layout.addStretch()
        pagination_layout.addWidget(self.threats_next_btn)
        top_layout.addLayout(pagination_layout)

        splitter.addWidget(top_widget)

        # Bottom part for details
        self.threat_details_view = QTextBrowser()
        self.threat_details_view.setOpenExternalLinks(True)
        splitter.addWidget(self.threat_details_view)

        splitter.setSizes([500, 200])
        layout.addWidget(splitter)

        # --- Connections ---
        self.fetch_threats_btn.clicked.connect(self._start_fetch_threats)
        self.threats_table.currentItemChanged.connect(self._display_threat_details)
        self.threats_prev_btn.clicked.connect(self._go_to_previous_threats_page)
        self.threats_next_btn.clicked.connect(self._go_to_next_threats_page)
        items_per_page_combo.currentTextChanged.connect(self._change_threats_per_page)

        return widget

    def _create_activity_journal_tab(self):
        """Creates the UI for the Activity Journal tab with advanced filtering."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Filter Controls ---
        filters_box = QGroupBox("Filters")
        filters_layout = QGridLayout(filters_box)

        self.journal_search_input = QLineEdit()
        self.journal_search_input.setPlaceholderText("Search actions, targets, details...")
        filters_layout.addWidget(QLabel("Search:"), 0, 0)
        filters_layout.addWidget(self.journal_search_input, 0, 1, 1, 3)

        self.journal_category_combo = QComboBox()
        self.journal_category_combo.addItems([
            "All", "Tool Execution", "Authentication", "UI Interaction",
            "File Operation", "System", "Admin Action", "Offensive",
            "Packet Operations", "Miscellaneous"
        ])
        filters_layout.addWidget(QLabel("Category:"), 1, 0)
        filters_layout.addWidget(self.journal_category_combo, 1, 1)

        self.journal_severity_combo = QComboBox()
        self.journal_severity_combo.addItems(["All", "Low", "Medium", "High", "Critical"])
        filters_layout.addWidget(QLabel("Severity:"), 1, 2)
        filters_layout.addWidget(self.journal_severity_combo, 1, 3)

        self.journal_result_combo = QComboBox()
        self.journal_result_combo.addItems(["All", "Success", "Failure", "N/A"])
        filters_layout.addWidget(QLabel("Result:"), 1, 4)
        filters_layout.addWidget(self.journal_result_combo, 1, 5)


        # Admin-only user filter
        self.journal_user_combo = QComboBox()
        self.journal_user_label = QLabel("User:")
        if self.current_user and self.current_user['is_admin']:
            filters_layout.addWidget(self.journal_user_label, 2, 0)
            filters_layout.addWidget(self.journal_user_combo, 2, 1)
        else:
            self.journal_user_label.hide()
            self.journal_user_combo.hide()

        self.journal_date_from = QDateEdit(QDate.currentDate().addMonths(-1))
        self.journal_date_from.setCalendarPopup(True)
        filters_layout.addWidget(QLabel("From:"), 2, 2)
        filters_layout.addWidget(self.journal_date_from, 2, 3)

        self.journal_date_to = QDateEdit(QDate.currentDate())
        self.journal_date_to.setCalendarPopup(True)
        filters_layout.addWidget(QLabel("To:"), 2, 4)
        filters_layout.addWidget(self.journal_date_to, 2, 5)

        self.journal_apply_filters_btn = QPushButton(QIcon(self.icon_path("search.svg")), "Apply Filters")
        filters_layout.addWidget(self.journal_apply_filters_btn, 0, 4, 1, 2)

        self.journal_refresh_btn = QPushButton(QIcon(self.icon_path("refresh-cw.svg")), "Refresh")
        filters_layout.addWidget(self.journal_refresh_btn, 0, 6, 1, 1)

        self.journal_auto_refresh_check = QCheckBox("Auto-Refresh")
        self.journal_auto_refresh_check.setToolTip("Automatically refresh the journal every 30 seconds.")
        filters_layout.addWidget(self.journal_auto_refresh_check, 0, 7, 1, 1)


        layout.addWidget(filters_box)

        # --- Journal Tree ---
        self.journal_tree = QTreeWidget()
        self.journal_tree.setColumnCount(8)
        self.journal_tree.setHeaderLabels(["Timestamp", "User", "Category", "Action", "Severity", "Result", "Target", "Details"])
        header = self.journal_tree.header()
        header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        header.setStretchLastSection(True)
        header.resizeSection(0, 160); header.resizeSection(1, 100); header.resizeSection(2, 120)
        header.resizeSection(3, 150); header.resizeSection(4, 80); header.resizeSection(5, 80)
        header.resizeSection(6, 150)
        layout.addWidget(self.journal_tree)

        # --- Connect signals ---
        self.journal_apply_filters_btn.clicked.connect(self._populate_activity_journal_tab)
        self.journal_refresh_btn.clicked.connect(self._populate_activity_journal_tab)
        self.journal_search_input.returnPressed.connect(self._populate_activity_journal_tab)
        self.journal_auto_refresh_check.toggled.connect(self._handle_journal_auto_refresh)


        return widget

    def _handle_journal_auto_refresh(self, checked):
        """Starts or stops the timer for auto-refreshing the activity journal."""
        if checked:
            if not self.journal_refresh_timer:
                self.journal_refresh_timer = QTimer(self)
                self.journal_refresh_timer.timeout.connect(self._populate_activity_journal_tab)
            # Start timer only if the tab is currently visible
            if self.tab_widget.tabText(self.tab_widget.currentIndex()) == "Activity Journal":
                self.journal_refresh_timer.start(30000) # 30 seconds
                logging.info("Activity Journal auto-refresh started.")
        else:
            if self.journal_refresh_timer and self.journal_refresh_timer.isActive():
                self.journal_refresh_timer.stop()
                logging.info("Activity Journal auto-refresh stopped.")

    def _populate_activity_journal_tab(self):
        """Fetches and displays the activity log based on the selected filters."""
        if not self.current_user:
            return

        self.journal_tree.clear()

        search_term = self.journal_search_input.text()
        category = self.journal_category_combo.currentText()
        severity = self.journal_severity_combo.currentText()
        result = self.journal_result_combo.currentText()
        date_from = self.journal_date_from.date().toString("yyyy-MM-dd")
        date_to = self.journal_date_to.date().toString("yyyy-MM-dd")

        user_id_filter = None
        if self.current_user['is_admin']:
            # Populate user dropdown, preserving selection
            current_user_selection = self.journal_user_combo.currentText()
            self.journal_user_combo.blockSignals(True)
            self.journal_user_combo.clear()
            self.journal_user_combo.addItem("All Users", userData=None)
            all_users = database.get_all_users()
            for user in all_users:
                self.journal_user_combo.addItem(user['username'], userData=user['id'])
            self.journal_user_combo.setCurrentText(current_user_selection)
            self.journal_user_combo.blockSignals(False)

            user_id_filter = self.journal_user_combo.currentData()
        else:
            # Non-admins can only see their own activity
            user_id_filter = self.current_user['id']

        try:
            log_entries = database.get_activity_log(
                user_id=user_id_filter,
                search_term=search_term,
                category_filter=category,
                severity_filter=severity,
                result_filter=result,
                date_from=date_from,
                date_to=date_to
            )

            if not log_entries:
                self.journal_tree.addTopLevelItem(QTreeWidgetItem(["No matching activity found."]))
                return

            for record in log_entries:
                try:
                    severity = record['severity']
                except (IndexError, KeyError):
                    severity = 'N/A'
                try:
                    result = record['result']
                except (IndexError, KeyError):
                    result = 'N/A'

                item = QTreeWidgetItem([
                    str(record['timestamp']),
                    str(record['username']),
                    str(record['category']),
                    str(record['action']),
                    str(severity),
                    str(result),
                    str(record['target']),
                    str(record['details'])
                ])
                self._set_journal_item_color(item, severity)
                self.journal_tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Failed to populate activity journal: {e}", exc_info=True)
            QMessageBox.critical(self, "Journal Error", f"Could not load activity journal: {e}")

    def _set_journal_item_color(self, item, severity):
        """Sets the background color of a journal item based on its severity."""
        if not severity:
            return

        # Define colors. These could be adapted to the current theme later.
        colors = {
            "Low": QColor("#3c5c4c"),
            "Medium": QColor("#7a6332"),
            "High": QColor("#8b4513"),
            "Critical": QColor("#8b0000"),
        }

        color = colors.get(severity)
        if color:
            for i in range(item.columnCount()):
                item.setBackground(i, color)

    def _on_main_tab_changed(self, index):
        """Handler for when the main tab is changed, to auto-load data."""
        try:
            # Apply fade-in animation to the newly selected tab
            current_widget = self.tab_widget.widget(index)
            if current_widget:
                self.tab_opacity_effect = QGraphicsOpacityEffect(current_widget)
                current_widget.setGraphicsEffect(self.tab_opacity_effect)
                self.tab_animation = QPropertyAnimation(self.tab_opacity_effect, b"opacity")
                self.tab_animation.setDuration(300)
                self.tab_animation.setStartValue(0.0)
                self.tab_animation.setEndValue(1.0)
                self.tab_animation.setEasingCurve(QEasingCurve.Type.InOutQuad)
                self.tab_animation.start()

            # Stop the timer if it's running, regardless of which tab was previously selected.
            # The logic below will restart it if the new tab is the correct one.
            if self.journal_refresh_timer and self.journal_refresh_timer.isActive():
                self.journal_refresh_timer.stop()
                logging.info("Activity Journal auto-refresh timer stopped due to tab change.")

            tab_text = self.tab_widget.tabText(index)

            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='UI Interaction',
                    action='Tab Switched',
                    target=tab_text,
                    details=f"Switched to tab index {index}"
                )

            if tab_text == "Threat Intelligence" and not self.threat_intel_loaded:
                logging.info("Threat Intelligence tab opened for the first time. Auto-fetching CVEs.")
                self.threat_intel_loaded = True
                self._start_fetch_threats()
            elif tab_text == "Activity Journal":
                # Start timer if auto-refresh is checked and the tab is now visible
                if self.journal_auto_refresh_check.isChecked() and self.journal_refresh_timer:
                    self.journal_refresh_timer.start(30000)
                    logging.info("Activity Journal tab became visible, starting auto-refresh timer.")
                # Always populate on first load
                if not self.history_loaded:
                    logging.info("Activity Journal tab opened for the first time. Auto-populating.")
                    self.history_loaded = True
                    self._populate_activity_journal_tab()

        except Exception as e:
            logging.error(f"Error in _on_main_tab_changed: {e}")

    def _start_fetch_threats(self):
        """Initiates the background thread to fetch recent CVEs."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        self.is_tool_running = True
        self.fetch_threats_btn.setEnabled(False)
        self.status_bar.showMessage("Fetching latest CVEs from cve.circl.lu...")
        self.threats_table.clear()
        self.threat_details_view.clear()

        self.worker = WorkerThread(self._recent_threats_thread)
        self.active_threads.append(self.worker)
        self.worker.start()

    def _recent_threats_thread(self):
        """Worker thread to fetch the last 30 CVEs."""
        q = self.tool_results_queue
        data = [] # Default to an empty list
        try:
            url = "https://cve.circl.lu/api/last"
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:109.0) Gecko/20100101 Firefox/119.0'}
            req = urllib.request.Request(url, headers=headers)

            with urllib.request.urlopen(req, timeout=15) as response:
                if response.status == 200:
                    raw_data = response.read().decode('utf-8')
                    if raw_data:
                        data = json.loads(raw_data)
                    else:
                        logging.warning("Recent threats API returned empty response.")
                else:
                    raise Exception(f"API request failed with status code {response.status}")
        except Exception as e:
            logging.error(f"Failed to fetch recent threats: {e}", exc_info=True)
            q.put(('error', 'Threat Intelligence Error', f"Could not fetch data: {e}"))
        finally:
            # Always put the result, even if it's an empty list.
            # The UI handler will be responsible for displaying the "No results" message.
            q.put(('recent_threats_result', data))
            q.put(('tool_finished', 'fetch_threats', "cve.circl.lu", f"Fetched {len(data)} CVEs"))

    def _handle_recent_threats_result(self, data):
        """Handles the successful result from the threat fetcher thread."""
        self.recent_threats_data = data
        self.threats_current_page = 0
        self._update_threats_display()
        if data:
            self.status_bar.showMessage(f"Successfully fetched {len(data)} recent CVEs.", 5000)
        else:
            self.status_bar.showMessage("Could not fetch recent threats or none were found.", 5000)

    def _update_threats_display(self):
        """Updates the threat table and pagination controls based on the current state."""
        self.threats_table.clear()

        start_index = self.threats_current_page * self.threats_items_per_page
        end_index = start_index + self.threats_items_per_page

        page_data = self.recent_threats_data[start_index:end_index]

        for cve in page_data:
            cve_id = cve.get('id', 'N/A')
            published = cve.get('Published', 'N/A')
            cvss = str(cve.get('cvss', 'N/A'))
            summary = cve.get('summary', '')

            item = QTreeWidgetItem([cve_id, published, cvss, summary])
            item.setData(0, Qt.ItemDataRole.UserRole, cve) # Store full data
            self.threats_table.addTopLevelItem(item)

        total_pages = (len(self.recent_threats_data) + self.threats_items_per_page - 1) // self.threats_items_per_page
        self.threats_page_label.setText(f"Page {self.threats_current_page + 1} / {total_pages}")
        self.threats_prev_btn.setEnabled(self.threats_current_page > 0)
        self.threats_next_btn.setEnabled(end_index < len(self.recent_threats_data))

    def _display_threat_details(self, current_item, previous_item):
        """Displays full details for the selected CVE."""
        if not current_item:
            return
        cve_data = current_item.data(0, Qt.ItemDataRole.UserRole)
        if not cve_data:
            return

        html = f"<h3>{cve_data.get('id', 'N/A')}</h3>"
        html += f"<p><b>Summary:</b><br>{cve_data.get('summary', 'No summary available.')}</p>"
        html += f"<b>Published:</b> {cve_data.get('Published', 'N/A')}<br>"
        html += f"<b>Modified:</b> {cve_data.get('Modified', 'N/A')}<br>"
        html += f"<b>CVSS Score:</b> {cve_data.get('cvss', 'N/A')}<br>"

        refs = cve_data.get('references', [])
        if refs:
            html += "<p><b>References:</b><ul>"
            for ref in refs:
                html += f'<li><a href="{ref}">{ref}</a></li>'
            html += "</ul></p>"

        self.threat_details_view.setHtml(html)

    def _go_to_next_threats_page(self):
        self.threats_current_page += 1
        self._update_threats_display()

    def _go_to_previous_threats_page(self):
        self.threats_current_page -= 1
        self._update_threats_display()

    def _change_threats_per_page(self, text):
        self.threats_items_per_page = int(text)
        self.threats_current_page = 0 # Reset to first page
        self._update_threats_display()

    def _create_exploit_db_search_tab(self):
        """Creates the UI for the Exploit-DB Search tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- API Key Input ---
        api_key_layout = QHBoxLayout()
        api_key_layout.addWidget(QLabel("Vulners API Key:"))
        self.exploitdb_api_key_input = QLineEdit()
        self.exploitdb_api_key_input.setPlaceholderText("Get a free key from vulners.com")
        self.exploitdb_api_key_input.setEchoMode(QLineEdit.EchoMode.Password)
        api_key_layout.addWidget(self.exploitdb_api_key_input)
        save_api_key_btn = QPushButton("Save Key")
        api_key_layout.addWidget(save_api_key_btn)
        layout.addLayout(api_key_layout)

        # --- Search Controls ---
        controls_layout = QHBoxLayout()
        controls_layout.addWidget(QLabel("Search Exploits:"))
        self.exploitdb_search_input = QLineEdit()
        self.exploitdb_search_input.setPlaceholderText("Enter software name, version, etc. (e.g., 'wordpress 4.7.0')")
        controls_layout.addWidget(self.exploitdb_search_input)
        self.exploitdb_search_button = QPushButton("Search")
        controls_layout.addWidget(self.exploitdb_search_button)
        layout.addLayout(controls_layout)

        # --- Results Display ---
        self.exploitdb_results_table = QTreeWidget()
        self.exploitdb_results_table.setColumnCount(3)
        self.exploitdb_results_table.setHeaderLabels(["ID", "Title", "URL"])
        self.exploitdb_results_table.header().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.exploitdb_results_table.header().setStretchLastSection(True)
        layout.addWidget(self.exploitdb_results_table)

        # --- Connections ---
        save_api_key_btn.clicked.connect(self.save_vulners_api_key)
        self.exploitdb_search_button.clicked.connect(self.start_exploit_search)
        self.exploitdb_results_table.itemDoubleClicked.connect(self.open_exploit_url)

        # Load saved API key on startup
        self.load_vulners_api_key()

        return widget

    def save_vulners_api_key(self):
        """Saves the Vulners API key to a file."""
        api_key = self.exploitdb_api_key_input.text()
        if not api_key:
            QMessageBox.warning(self, "Input Error", "Please enter an API key to save.")
            return
        try:
            with open("vulners_api.key", "w") as f:
                f.write(api_key)
            QMessageBox.information(self, "Success", "Vulners API key saved successfully.")
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not save API key: {e}")

    def load_vulners_api_key(self):
        """Loads the Vulners API key from a file."""
        try:
            with open("vulners_api.key", "r") as f:
                api_key = f.read().strip()
                self.exploitdb_api_key_input.setText(api_key)
                logging.info("Loaded Vulners API key.")
        except FileNotFoundError:
            pass # It's okay if the file doesn't exist yet
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not load API key: {e}")

    def start_exploit_search(self):
        """Starts the Exploit-DB search worker thread."""
        if not shutil.which("getsploit"):
            QMessageBox.critical(self, "GetSploit Error", "'getsploit' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        query = self.exploitdb_search_input.text()
        api_key = self.exploitdb_api_key_input.text()

        if not query:
            QMessageBox.critical(self, "Input Error", "Please provide a search query.")
            return
        if not api_key:
            QMessageBox.critical(self, "API Key Error", "Vulners API key is required for searching exploits.")
            return

        self.is_tool_running = True
        self.exploitdb_search_button.setEnabled(False)
        self.exploitdb_results_table.clear()
        self.status_bar.showMessage(f"Searching for exploits related to '{query}'...")

        self.worker = WorkerThread(self._exploit_search_thread, args=(query, api_key))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _exploit_search_thread(self, query, api_key):
        """Worker thread to search for exploits using getsploit."""
        q = self.tool_results_queue
        command = ["getsploit", "--api", api_key, query]

        try:
            # Use CREATE_NO_WINDOW flag on Windows to hide the console
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            output, _ = process.communicate()

            if process.returncode != 0:
                raise Exception(output)

            # Parse the table-formatted output
            lines = output.strip().split('\n')
            # Find the header line to start parsing from
            header_index = -1
            for i, line in enumerate(lines):
                if 'ID' in line and 'Exploit Title' in line and 'URL' in line:
                    header_index = i
                    break

            if header_index == -1:
                q.put(('exploit_search_status', "No results found or could not parse output."))
                return

            results = []
            # Start from 2 lines after the header to skip the header and the '======' line
            for line in lines[header_index + 2:]:
                if line.startswith('+--'): # End of table
                    break
                parts = [p.strip() for p in line.split('|') if p.strip()]
                if len(parts) >= 3:
                    results.append(parts)

            q.put(('exploit_search_results', results))
            q.put(('exploit_search_status', f"Found {len(results)} exploits."))

        except FileNotFoundError:
            q.put(('error', 'GetSploit Error', "'getsploit' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"getsploit search failed: {e}", exc_info=True)
            q.put(('error', 'Exploit Search Error', str(e)))
        finally:
            q.put(('tool_finished', 'exploit_search', query, f"Found {len(results)} exploits."))

    def open_exploit_url(self, item, column):
        """Opens the selected exploit URL in the default web browser."""
        url = item.text(2) # URL is in the 3rd column
        if url and url.startswith("http"):
            webbrowser.open(url)
        else:
            QMessageBox.warning(self, "Invalid URL", f"The selected item does not have a valid URL: {url}")

    def _create_log_panel(self):
        """Creates the dockable logging panel at the bottom of the window."""
        log_dock_widget = QDockWidget("Live Log", self)
        log_dock_widget.setAllowedAreas(Qt.DockWidgetArea.BottomDockWidgetArea)
        self.log_console = QPlainTextEdit(); self.log_console.setReadOnly(True)
        log_dock_widget.setWidget(self.log_console)
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, log_dock_widget)

    def _create_sniffer_tab(self):
        """Creates the UI for the Packet Sniffer tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Create the results widget first
        self.packet_list_widget = QTreeWidget()
        self.packet_list_widget.setColumnCount(6)
        self.packet_list_widget.setHeaderLabels(["No.", "Time", "Source", "Destination", "Protocol", "Length"])
        # Make header columns resizable and stretch the last section
        header = self.packet_list_widget.header()
        header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        header.setStretchLastSection(True)

        # --- Control Panel ---
        control_panel = QFrame()
        control_panel.setObjectName("controlPanel")
        control_panel.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; }")
        control_layout = QHBoxLayout(control_panel)
        control_layout.setContentsMargins(10, 10, 10, 10)
        control_layout.setSpacing(10)

        self.start_sniff_btn = QPushButton(QIcon(self.icon_path("search.svg")), " Start Sniffing")
        self.stop_sniff_btn = QPushButton(QIcon(self.icon_path("square.svg")), " Stop Sniffing"); self.stop_sniff_btn.setEnabled(False)
        self.clear_sniff_btn = QPushButton("Clear")
        export_btn = self._create_export_button(self.packet_list_widget)

        control_layout.addWidget(self.start_sniff_btn)
        control_layout.addWidget(self.stop_sniff_btn)
        control_layout.addWidget(self.clear_sniff_btn)
        control_layout.addWidget(export_btn)
        control_layout.addStretch(1)

        control_layout.addWidget(QLabel("BPF Filter:"))
        self.filter_input = QLineEdit()
        self.filter_input.setPlaceholderText("e.g., tcp and port 80")
        control_layout.addWidget(self.filter_input, 2) # Give filter more stretch

        control_layout.addWidget(QLabel("Common:"))
        self.common_filter_combo = QComboBox()
        self.common_filter_combo.addItems(COMMON_FILTERS)
        self.common_filter_combo.textActivated.connect(self.filter_input.setText)
        control_layout.addWidget(self.common_filter_combo)
        layout.addWidget(control_panel)

        # Main splitter for top (list) and bottom (details)
        main_splitter = QSplitter(Qt.Orientation.Vertical)
        main_splitter.addWidget(self.packet_list_widget)

        # Bottom splitter for details tree and hex view
        bottom_splitter = QSplitter(Qt.Orientation.Vertical)

        self.packet_details_tree = QTreeWidget()
        self.packet_details_tree.setHeaderLabels(["Field", "Value"])
        # Make header columns resizable and stretch the last section
        details_header = self.packet_details_tree.header()
        details_header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        details_header.setStretchLastSection(True)
        bottom_splitter.addWidget(self.packet_details_tree)

        self.packet_hex_view = QTextBrowser()
        self.packet_hex_view.setReadOnly(True)
        # Use a monospaced font for the hex view for proper alignment
        self.packet_hex_view.setFont(QFont("Courier New", 10))
        bottom_splitter.addWidget(self.packet_hex_view)

        bottom_splitter.setSizes([200, 100]) # Initial sizes for tree and hex view

        main_splitter.addWidget(bottom_splitter)
        main_splitter.setSizes([400, 300]) # Initial sizes for packet list and details area
        layout.addWidget(main_splitter)

        # Connect signals
        self.start_sniff_btn.clicked.connect(self.start_sniffing)
        self.stop_sniff_btn.clicked.connect(self.stop_sniffing)
        self.clear_sniff_btn.clicked.connect(self.clear_sniffer_display)
        self.packet_list_widget.currentItemChanged.connect(self.display_packet_details)
        return widget

    def _setup_logging(self):
        """Configures the logging system to output to a file and the UI panel."""
        root_logger = logging.getLogger()
        for handler in root_logger.handlers[:]: root_logger.removeHandler(handler)
        file_handler = logging.FileHandler('zurvan.log', mode='w')
        formatter = logging.Formatter('%(asctime)s - %(threadName)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        root_logger.addHandler(file_handler)
        self.qt_handler = QtLogHandler()
        self.qt_handler.log_updated.connect(self.log_console.appendPlainText)
        self.qt_handler.setFormatter(formatter)
        root_logger.addHandler(self.qt_handler)
        root_logger.setLevel(logging.INFO)

    def _create_sniffer_tab(self):
        """Creates the UI for the Packet Sniffer tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Create the results widget first
        self.packet_list_widget = QTreeWidget()
        self.packet_list_widget.setColumnCount(6)
        self.packet_list_widget.setHeaderLabels(["No.", "Time", "Source", "Destination", "Protocol", "Length"])
        # Make header columns resizable and stretch the last section
        header = self.packet_list_widget.header()
        header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        header.setStretchLastSection(True)

        # --- Control Panel ---
        control_panel = QFrame()
        control_panel.setObjectName("controlPanel")
        control_panel.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; }")
        control_layout = QHBoxLayout(control_panel)
        control_layout.setContentsMargins(10, 10, 10, 10)
        control_layout.setSpacing(10)

        self.start_sniff_btn = QPushButton(QIcon(self.icon_path("search.svg")), " Start Sniffing")
        self.stop_sniff_btn = QPushButton(QIcon(self.icon_path("square.svg")), " Stop Sniffing"); self.stop_sniff_btn.setEnabled(False)
        self.clear_sniff_btn = QPushButton("Clear")
        export_btn = self._create_export_button(self.packet_list_widget)

        control_layout.addWidget(self.start_sniff_btn)
        control_layout.addWidget(self.stop_sniff_btn)
        control_layout.addWidget(self.clear_sniff_btn)
        control_layout.addWidget(export_btn)
        control_layout.addStretch(1)

        control_layout.addWidget(QLabel("BPF Filter:"))
        self.filter_input = QLineEdit()
        self.filter_input.setPlaceholderText("e.g., tcp and port 80")
        control_layout.addWidget(self.filter_input, 2) # Give filter more stretch

        control_layout.addWidget(QLabel("Common:"))
        self.common_filter_combo = QComboBox()
        self.common_filter_combo.addItems(COMMON_FILTERS)
        self.common_filter_combo.textActivated.connect(self.filter_input.setText)
        control_layout.addWidget(self.common_filter_combo)
        layout.addWidget(control_panel)

        # Main splitter for top (list) and bottom (details)
        main_splitter = QSplitter(Qt.Orientation.Vertical)
        main_splitter.addWidget(self.packet_list_widget)

        # Bottom splitter for details tree and hex view
        bottom_splitter = QSplitter(Qt.Orientation.Vertical)

        self.packet_details_tree = QTreeWidget()
        self.packet_details_tree.setHeaderLabels(["Field", "Value"])
        # Make header columns resizable and stretch the last section
        details_header = self.packet_details_tree.header()
        details_header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        details_header.setStretchLastSection(True)
        bottom_splitter.addWidget(self.packet_details_tree)

        self.packet_hex_view = QTextBrowser()
        self.packet_hex_view.setReadOnly(True)
        # Use a monospaced font for the hex view for proper alignment
        self.packet_hex_view.setFont(QFont("Courier New", 10))
        bottom_splitter.addWidget(self.packet_hex_view)

        bottom_splitter.setSizes([200, 100]) # Initial sizes for tree and hex view

        main_splitter.addWidget(bottom_splitter)
        main_splitter.setSizes([400, 300]) # Initial sizes for packet list and details area
        layout.addWidget(main_splitter)

        # Connect signals
        self.start_sniff_btn.clicked.connect(self.start_sniffing)
        self.stop_sniff_btn.clicked.connect(self.stop_sniffing)
        self.clear_sniff_btn.clicked.connect(self.clear_sniffer_display)
        self.packet_list_widget.currentItemChanged.connect(self.display_packet_details)
        return widget

    def _create_crafter_tab(self):
        """Creates the UI for the Packet Crafter tab."""
        widget = QWidget(); main_layout = QVBoxLayout(widget)
        top_splitter = QSplitter(Qt.Orientation.Horizontal); main_layout.addWidget(top_splitter)
        left_panel = QWidget(); left_layout = QVBoxLayout(left_panel); top_splitter.addWidget(left_panel)
        controls_layout = QHBoxLayout()
        self.proto_to_add = QComboBox(); self.proto_to_add.addItems(AVAILABLE_PROTOCOLS.keys())
        add_btn = QPushButton("Add"); remove_btn = QPushButton("Remove");
        controls_layout.addWidget(self.proto_to_add); controls_layout.addWidget(add_btn); controls_layout.addWidget(remove_btn)
        left_layout.addLayout(controls_layout)

        layer_actions_layout = QHBoxLayout()
        fuzz_btn = QPushButton("Fuzz/Unfuzz Selected Layer"); layer_actions_layout.addWidget(fuzz_btn)
        templates_btn = QPushButton("Templates"); layer_actions_layout.addWidget(templates_btn)
        left_layout.addLayout(layer_actions_layout)

        self.layer_list_widget = QListWidget(); left_layout.addWidget(self.layer_list_widget)
        left_layout.addWidget(QLabel("Packet Summary:")); self.crafter_summary = QPlainTextEdit(); self.crafter_summary.setReadOnly(True); left_layout.addWidget(self.crafter_summary)
        right_panel = QWidget(); right_layout = QVBoxLayout(right_panel)
        right_layout.addWidget(QLabel("Layer Fields")); self.scroll_area = QScrollArea(); self.scroll_area.setWidgetResizable(True)
        self.fields_widget = QWidget(); self.fields_layout = QVBoxLayout(self.fields_widget); self.scroll_area.setWidget(self.fields_widget)
        right_layout.addWidget(self.scroll_area); top_splitter.addWidget(right_panel); top_splitter.setSizes([300, 400])
        send_frame = QFrame(); send_frame.setFrameShape(QFrame.Shape.StyledPanel); main_layout.addWidget(send_frame)
        send_layout = QVBoxLayout(send_frame)
        send_controls_layout = QHBoxLayout()
        send_controls_layout.addWidget(QLabel("Count:")); self.send_count_edit = QLineEdit("1"); send_controls_layout.addWidget(self.send_count_edit)
        send_controls_layout.addWidget(QLabel("Interval:")); self.send_interval_edit = QLineEdit("0.1"); send_controls_layout.addWidget(self.send_interval_edit)
        self.send_btn = QPushButton("Send Packet(s)")
        self.send_cancel_btn = QPushButton("Cancel"); self.send_cancel_btn.setEnabled(False)
        send_controls_layout.addWidget(self.send_btn)
        send_controls_layout.addWidget(self.send_cancel_btn)
        send_layout.addLayout(send_controls_layout)
        self.send_results_widget = QTreeWidget(); self.send_results_widget.setColumnCount(3); self.send_results_widget.setHeaderLabels(["No.", "Sent", "Received"])
        send_layout.addWidget(self.send_results_widget)
        send_layout.addWidget(self._create_export_button(self.send_results_widget))
        add_btn.clicked.connect(self.crafter_add_layer); remove_btn.clicked.connect(self.crafter_remove_layer)
        self.layer_list_widget.currentRowChanged.connect(self.crafter_display_layer_fields)
        templates_menu = QMenu(self)
        for name in PACKET_TEMPLATES.keys():
            action = QAction(name, self); action.triggered.connect(lambda checked, n=name: self.crafter_load_template(n)); templates_menu.addAction(action)
        templates_btn.setMenu(templates_menu)
        self.send_btn.clicked.connect(self.crafter_send_packet)
        self.send_cancel_btn.clicked.connect(self.cancel_tool)
        fuzz_btn.clicked.connect(self.crafter_toggle_fuzz_layer)
        return widget

    def _create_nmap_scanner_tool(self):
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Configurable options widget ---
        config_widget, self.nmap_controls = self._create_nmap_config_widget()
        main_layout.addWidget(config_widget)

        # Connect signals now that we have the controls dictionary
        controls = self.nmap_controls
        controls['start_btn'].clicked.connect(self.start_nmap_scan)
        controls['cancel_btn'].clicked.connect(self.cancel_tool)
        controls['report_btn'].clicked.connect(self.generate_nmap_report)
        controls['all_ports_btn'].clicked.connect(self._nmap_set_all_ports)
        controls['super_complete_btn'].clicked.connect(self._nmap_toggle_super_complete)
        controls['preset_combo'].textActivated.connect(self._handle_nmap_preset_selected)
        controls['a_check'].toggled.connect(self._nmap_on_aggressive_toggled)
        controls['scan_type_combo'].currentTextChanged.connect(self._nmap_on_ping_scan_toggled)

        # --- Output Console ---
        self.nmap_output_console = QPlainTextEdit()
        self.nmap_output_console.setReadOnly(True)
        self.nmap_output_console.setFont(QFont("Courier New", 10))
        self.nmap_output_console.setPlaceholderText("Nmap command output will be displayed here...")
        main_layout.addWidget(self.nmap_output_console, 1) # Give it stretch factor

        return widget

    def _create_nmap_config_widget(self):
        """Creates a reusable, self-contained widget with all of Nmap's configuration options, now with more tabs."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        controls = {}

        # --- Top Controls (Target, Ports, Start/Stop) ---
        top_controls = QFrame()
        top_controls.setObjectName("controlPanel")
        top_controls.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        top_layout = QGridLayout(top_controls)

        top_layout.addWidget(QLabel("Target(s):"), 0, 0)
        controls['target_edit'] = QLineEdit("localhost")
        controls['target_edit'].setToolTip("Enter hosts, separated by spaces (e.g., scanme.nmap.org, 192.168.1.0/24).")
        top_layout.addWidget(controls['target_edit'], 0, 1, 1, 3)

        ports_layout = QHBoxLayout()
        ports_layout.addWidget(QLabel("Ports:"))
        controls['ports_edit'] = QLineEdit()
        controls['ports_edit'].setToolTip("e.g., 22,80,443 or 1-1024. Leave blank for default (top 1000).")
        ports_layout.addWidget(controls['ports_edit'])
        controls['top_ports_check'] = QCheckBox("--top-ports")
        ports_layout.addWidget(controls['top_ports_check'])
        controls['top_ports_edit'] = QLineEdit("100")
        ports_layout.addWidget(controls['top_ports_edit'])
        top_layout.addLayout(ports_layout, 1, 0, 1, 4)


        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['cancel_btn'] = QPushButton("Cancel"); controls['cancel_btn'].setEnabled(False)
        controls['report_btn'] = QPushButton("Generate HTML Report"); controls['report_btn'].setEnabled(False)
        top_layout.addWidget(controls['start_btn'], 0, 4)
        top_layout.addWidget(controls['cancel_btn'], 1, 4)
        top_layout.addWidget(controls['report_btn'], 0, 5, 2, 1)

        presets_layout = QVBoxLayout()
        controls['all_ports_btn'] = QPushButton("All Ports")
        controls['super_complete_btn'] = QPushButton("Super Complete Scan")
        presets_layout.addWidget(controls['all_ports_btn'])
        presets_layout.addWidget(controls['super_complete_btn'])
        top_layout.addLayout(presets_layout, 0, 6, 2, 1)
        main_layout.addWidget(top_controls)

        # --- Tabbed Interface for Options ---
        nmap_tabs = QTabWidget()
        main_layout.addWidget(nmap_tabs)

        # --- Tab 1: Scan & Detection ---
        scan_tab = QWidget()
        scan_layout = QFormLayout(scan_tab)
        controls['scan_type_combo'] = QComboBox()
        controls['scan_type_combo'].addItems([
            "SYN Stealth Scan (-sS)", "TCP Connect Scan (-sT)", "UDP Scan (-sU)",
            "FIN Scan (-sF)", "Xmas Scan (-sX)", "Null Scan (-sN)",
            "SCTP INIT Scan (-sY)", "SCTP COOKIE-ECHO Scan (-sZ)"
        ])
        scan_layout.addRow("Scan Type:", controls['scan_type_combo'])

        controls['detection_box'] = QGroupBox("Detection")
        detection_grid = QGridLayout(controls['detection_box'])
        controls['sv_check'] = QCheckBox("Service/Version (-sV)")
        controls['o_check'] = QCheckBox("OS Detection (-O)")
        controls['sc_check'] = QCheckBox("Default Scripts (-sC)")
        controls['a_check'] = QCheckBox("Aggressive Scan (-A)")
        controls['reason_check'] = QCheckBox("Show Reason (--reason)")
        controls['osscan_guess_check'] = QCheckBox("Guess OS (--osscan-guess)")
        controls['version_light_check'] = QCheckBox("Light Version Scan (--version-light)")
        controls['version_all_check'] = QCheckBox("Intense Version Scan (--version-all)")


        detection_grid.addWidget(controls['sv_check'], 0, 0)
        detection_grid.addWidget(controls['o_check'], 0, 1)
        detection_grid.addWidget(controls['sc_check'], 1, 0)
        detection_grid.addWidget(controls['a_check'], 1, 1)
        detection_grid.addWidget(controls['reason_check'], 2, 0)
        detection_grid.addWidget(controls['osscan_guess_check'], 2, 1)
        detection_grid.addWidget(controls['version_light_check'], 3, 0)
        detection_grid.addWidget(controls['version_all_check'], 3, 1)


        controls['version_intensity_edit'] = QLineEdit("7")
        detection_grid.addWidget(QLabel("Version Intensity (0-9):"), 4, 0)
        detection_grid.addWidget(controls['version_intensity_edit'], 4, 1)
        scan_layout.addRow(controls['detection_box'])
        nmap_tabs.addTab(scan_tab, "Scan & Detection")

        # --- Tab 2: Host Discovery ---
        discovery_tab = QWidget()
        discovery_layout = QFormLayout(discovery_tab)
        controls['ping_scan_check'] = QCheckBox("Ping Scan Only (-sn)")
        controls['ping_scan_check'].setToolTip("Disables port scanning, only discovers online hosts.")
        discovery_layout.addRow(controls['ping_scan_check'])

        controls['pn_check'] = QCheckBox("No Ping (-Pn)")
        controls['pn_check'].setToolTip("Treat all hosts as online, skipping discovery.")
        discovery_layout.addRow("Skip Discovery:", controls['pn_check'])

        discovery_layout.addRow(QFrame(frameShape=QFrame.Shape.HLine))
        discovery_layout.addRow(QLabel("<b>Discovery Probes:</b>"))
        probes_grid = QGridLayout()
        controls['ps_check'] = QCheckBox("TCP SYN Ping (-PS)"); probes_grid.addWidget(controls['ps_check'], 0, 0)
        controls['pa_check'] = QCheckBox("TCP ACK Ping (-PA)"); probes_grid.addWidget(controls['pa_check'], 0, 1)
        controls['pu_check'] = QCheckBox("UDP Ping (-PU)"); probes_grid.addWidget(controls['pu_check'], 1, 0)
        controls['pr_check'] = QCheckBox("ARP Ping (-PR)"); probes_grid.addWidget(controls['pr_check'], 1, 1)
        controls['pe_check'] = QCheckBox("ICMP Echo Ping (-PE)"); probes_grid.addWidget(controls['pe_check'], 2, 0)
        controls['pp_check'] = QCheckBox("ICMP Timestamp Ping (-PP)"); probes_grid.addWidget(controls['pp_check'], 2, 1)
        controls['disable_arp_ping_check'] = QCheckBox("Disable ARP Ping (--disable-arp-ping)"); probes_grid.addWidget(controls['disable_arp_ping_check'], 3, 0, 1, 2)
        discovery_layout.addRow(probes_grid)
        nmap_tabs.addTab(discovery_tab, "Host Discovery")

        # --- Tab 3: Performance ---
        perf_tab = QWidget()
        perf_layout = QFormLayout(perf_tab)
        controls['timing_combo'] = QComboBox()
        controls['timing_combo'].addItems(["T0", "T1", "T2", "T3", "T4", "T5"])
        controls['timing_combo'].setCurrentIndex(3)
        perf_layout.addRow("Timing Template (-T):", controls['timing_combo'])

        rate_layout = QHBoxLayout()
        rate_layout.addWidget(QLabel("Min:"))
        controls['min_rate_edit'] = QLineEdit(); rate_layout.addWidget(controls['min_rate_edit'])
        rate_layout.addWidget(QLabel("Max:"))
        controls['max_rate_edit'] = QLineEdit(); rate_layout.addWidget(controls['max_rate_edit'])
        perf_layout.addRow("Packet Rate (--min-rate):", rate_layout)

        parallel_layout = QHBoxLayout()
        parallel_layout.addWidget(QLabel("Min:"))
        controls['min_parallel_edit'] = QLineEdit(); parallel_layout.addWidget(controls['min_parallel_edit'])
        parallel_layout.addWidget(QLabel("Max:"))
        controls['max_parallel_edit'] = QLineEdit(); parallel_layout.addWidget(controls['max_parallel_edit'])
        perf_layout.addRow("Parallelism (--min-parallelism):", parallel_layout)

        hostgroup_layout = QHBoxLayout()
        hostgroup_layout.addWidget(QLabel("Min:"))
        controls['min_hostgroup_edit'] = QLineEdit(); hostgroup_layout.addWidget(controls['min_hostgroup_edit'])
        hostgroup_layout.addWidget(QLabel("Max:"))
        controls['max_hostgroup_edit'] = QLineEdit(); hostgroup_layout.addWidget(controls['max_hostgroup_edit'])
        perf_layout.addRow("Host Group Size (--min-hostgroup):", hostgroup_layout)


        controls['scan_delay_edit'] = QLineEdit(); controls['scan_delay_edit'].setPlaceholderText("e.g., 500ms, 1s")
        perf_layout.addRow("Scan Delay (--scan-delay):", controls['scan_delay_edit'])
        controls['max_retries_edit'] = QLineEdit(); controls['max_retries_edit'].setPlaceholderText("e.g., 3")
        perf_layout.addRow("Max Retries (--max-retries):", controls['max_retries_edit'])
        controls['host_timeout_edit'] = QLineEdit(); controls['host_timeout_edit'].setPlaceholderText("e.g., 10m")
        perf_layout.addRow("Host Timeout (--host-timeout):", controls['host_timeout_edit'])
        nmap_tabs.addTab(perf_tab, "Performance")


        # --- Tab 4: Evasion ---
        evasion_tab = QWidget()
        evasion_layout = QFormLayout(evasion_tab)
        controls['frag_check'] = QCheckBox("Fragment packets (-f)")
        evasion_layout.addRow(controls['frag_check'])
        controls['mtu_edit'] = QLineEdit(); controls['mtu_edit'].setPlaceholderText("e.g., 8, 16, 24")
        evasion_layout.addRow("Set MTU (--mtu):", controls['mtu_edit'])
        controls['badsum_check'] = QCheckBox("Use bad checksums (--badsum)")
        evasion_layout.addRow(controls['badsum_check'])
        controls['source_port_edit'] = QLineEdit(); controls['source_port_edit'].setPlaceholderText("e.g., 53, 88")
        evasion_layout.addRow("Source Port (-g):", controls['source_port_edit'])
        controls['source_host_edit'] = QLineEdit(); controls['source_host_edit'].setPlaceholderText("e.g., 10.10.10.10")
        evasion_layout.addRow("Source Host (-S):", controls['source_host_edit'])
        controls['decoys_edit'] = QLineEdit(); controls['decoys_edit'].setPlaceholderText("e.g., ME,decoy1,decoy2")
        evasion_layout.addRow("Decoys (-D):", controls['decoys_edit'])
        controls['zombie_edit'] = QLineEdit(); controls['zombie_edit'].setPlaceholderText("e.g., idle.host.com")
        evasion_layout.addRow("Idle/Zombie Scan (-sI):", controls['zombie_edit'])
        controls['data_len_edit'] = QLineEdit(); controls['data_len_edit'].setPlaceholderText("e.g., 16")
        evasion_layout.addRow("Append Data (--data-length):", controls['data_len_edit'])
        controls['spoof_mac_edit'] = QLineEdit(); controls['spoof_mac_edit'].setPlaceholderText("e.g., 0, Cisco, Apple, or MAC")
        evasion_layout.addRow("Spoof MAC (--spoof-mac):", controls['spoof_mac_edit'])

        evasion_layout.addRow(QFrame(frameShape=QFrame.Shape.HLine))

        controls['data_string_edit'] = QLineEdit(); controls['data_string_edit'].setPlaceholderText("e.g., 'Scan from public network'")
        evasion_layout.addRow("Data String (--data-string):", controls['data_string_edit'])

        controls['ip_options_edit'] = QLineEdit(); controls['ip_options_edit'].setPlaceholderText("e.g., 'S 1.2.3.4'")
        evasion_layout.addRow("IP Options (--ip-options):", controls['ip_options_edit'])

        controls['ttl_edit'] = QLineEdit(); controls['ttl_edit'].setPlaceholderText("e.g., 64")
        evasion_layout.addRow("Time To Live (--ttl):", controls['ttl_edit'])

        controls['randomize_hosts_check'] = QCheckBox("Randomize hosts (--randomize-hosts)")
        evasion_layout.addRow(controls['randomize_hosts_check'])

        nmap_tabs.addTab(evasion_tab, "Evasion")

        # --- Tab 5: NSE ---
        nse_tab = QWidget()
        nse_layout = QFormLayout(nse_tab)
        controls['preset_combo'] = QComboBox();
        controls['preset_combo'].addItems(["-- Select a Preset --"] + list(self.nmap_script_presets.keys()))
        nse_layout.addRow("Script Presets:", controls['preset_combo'])
        controls['nse_vuln_check'] = QCheckBox("vuln"); controls['nse_discovery_check'] = QCheckBox("discovery"); controls['nse_safe_check'] = QCheckBox("safe")
        category_layout = QHBoxLayout(); category_layout.addWidget(controls['nse_vuln_check']); category_layout.addWidget(controls['nse_discovery_check']); category_layout.addWidget(controls['nse_safe_check'])
        nse_layout.addRow("Categories:", category_layout)
        controls['custom_script_edit'] = QLineEdit(); controls['custom_script_edit'].setPlaceholderText("e.g., http-title,smb-os-discovery")
        nse_layout.addRow("Custom Scripts:", controls['custom_script_edit'])
        controls['script_args_edit'] = QLineEdit(); controls['script_args_edit'].setPlaceholderText("e.g., http.useragent=MyAgent")
        nse_layout.addRow("Script Arguments:", controls['script_args_edit'])
        controls['script_desc_label'] = QLabel("Description: --"); controls['script_desc_label'].setWordWrap(True)
        nse_layout.addRow(controls['script_desc_label'])
        nmap_tabs.addTab(nse_tab, "NSE")

        # --- Tab 6: Output & Misc ---
        output_tab = QWidget()
        output_layout = QFormLayout(output_tab)
        controls['output_normal_edit'] = QLineEdit()
        output_layout.addRow("Normal Output (-oN):", controls['output_normal_edit'])
        controls['output_grepable_edit'] = QLineEdit()
        output_layout.addRow("Grepable Output (-oG):", controls['output_grepable_edit'])
        controls['resume_edit'] = QLineEdit()
        output_layout.addRow("Resume From (-oN):", controls['resume_edit'])
        controls['append_output_check'] = QCheckBox("Append to output files (--append-output)")
        output_layout.addRow(controls['append_output_check'])

        output_layout.addRow(QLabel("XML Output (-oX) is always enabled for reporting."))

        output_layout.addRow(QFrame(frameShape=QFrame.Shape.HLine))
        controls['v_check'] = QCheckBox("Verbose (-v)")
        output_layout.addRow(controls['v_check'])
        controls['packet_trace_check'] = QCheckBox("Packet Trace (--packet-trace)")
        output_layout.addRow(controls['packet_trace_check'])
        controls['traceroute_check'] = QCheckBox("Traceroute (--traceroute)")
        output_layout.addRow(controls['traceroute_check'])
        controls['ipv6_check'] = QCheckBox("Enable IPv6 (-6)")
        output_layout.addRow(controls['ipv6_check'])

        nmap_tabs.addTab(output_tab, "Output & Misc")

        # --- UI Logic Connections ---
        def toggle_scan_options(is_ping_scan):
            for i in range(nmap_tabs.count()):
                tab_text = nmap_tabs.tabText(i)
                if tab_text not in ["Host Discovery", "Output & Misc"]:
                    nmap_tabs.widget(i).setEnabled(not is_ping_scan)
            controls['ports_edit'].setEnabled(not is_ping_scan)
            controls['top_ports_check'].setEnabled(not is_ping_scan)
            controls['top_ports_edit'].setEnabled(not is_ping_scan)

        def toggle_port_edits(is_top_ports):
            controls['ports_edit'].setDisabled(is_top_ports)
            controls['top_ports_edit'].setEnabled(is_top_ports)

        controls['ping_scan_check'].toggled.connect(toggle_scan_options)
        controls['top_ports_check'].toggled.connect(toggle_port_edits)
        toggle_port_edits(False)

        return widget, controls

    def _nmap_on_aggressive_toggled(self, checked):
        controls = self.nmap_controls
        controls['sv_check'].setDisabled(checked)
        controls['o_check'].setDisabled(checked)
        controls['sc_check'].setDisabled(checked)
        controls['traceroute_check'].setDisabled(checked)

    def _nmap_on_ping_scan_toggled(self, text):
        controls = self.nmap_controls
        is_ping_scan = (text == "Ping Scan (-sn)")
        for w_key in ['detection_box', 'misc_box', 'ports_edit', 'timing_box', 'nse_box']:
             controls[w_key].setDisabled(is_ping_scan)

    def _nmap_set_all_ports(self):
        """Sets the Nmap port text field to scan all ports."""
        self.nmap_controls['ports_edit'].setText("1-65535")

    def _handle_nmap_preset_selected(self, preset_name):
        """Populates the script fields based on the selected Nmap preset."""
        controls = self.nmap_controls
        if preset_name == "-- Select a Preset --":
            controls['custom_script_edit'].clear()
            controls['script_args_edit'].clear()
            controls['script_desc_label'].setText("Description: --")
            return

        scripts, args, desc = self.nmap_script_presets.get(preset_name, ("", "", "No description available."))
        controls['custom_script_edit'].setText(scripts)
        controls['script_args_edit'].setText(args)
        controls['script_desc_label'].setText(f"Description: {desc}")

    def _nmap_toggle_super_complete(self):
        """Toggles the 'Super Complete Scan' preset."""
        # This now works for both the main tab and the LAB tab if the button is clicked there
        # by finding which set of controls is currently visible.
        controls = self.nmap_controls if self.nmap_controls['start_btn'].isVisible() else self.lab_tool_configs["Nmap Scan"]['controls']
        output_console = self.nmap_output_console if self.nmap_controls['start_btn'].isVisible() else self.lab_output_console

        # Determine if the preset is currently active for this specific UI context
        is_active = controls.get('_super_scan_active', False)

        if not is_active:
            controls['ports_edit'].setText("1-65535")
            controls['scan_type_combo'].setCurrentText("SYN Stealth Scan (-sS)")
            controls['timing_combo'].setCurrentText("T4 (Aggressive)")
            controls['a_check'].setChecked(True)
            controls['v_check'].setChecked(True)

            target = controls['target_edit'].text() or "[target]"
            command, _, _ = self._build_nmap_command(controls, target)
            output_console.clear()
            output_console.setPlainText(f"# Preset command preview:\n$ {' '.join(command)}")
            QMessageBox.information(self, "Preset Loaded", "Super Complete Scan options have been set.\nClick 'Start Scan' to run, or click the preset button again to cancel.")

            controls['super_complete_btn'].setText("Cancel Super Scan")
            controls['_super_scan_active'] = True # Store state within the controls dict
        else:
            controls['ports_edit'].setText("")
            controls['scan_type_combo'].setCurrentIndex(0)
            controls['timing_combo'].setCurrentIndex(3)
            controls['a_check'].setChecked(False)
            controls['v_check'].setChecked(False)

            controls['super_complete_btn'].setText("Super Complete Scan")
            output_console.clear()
            controls['_super_scan_active'] = False

    def _get_control_value(self, controls, key, widget_type):
        """A helper to get a value from either a dictionary of widgets or a dictionary of values."""
        # This check is fragile. A better approach would be to pass a flag, but this works for now.
        is_lab_run = controls and not isinstance(next(iter(controls.values()), None), QWidget)

        if is_lab_run:
            if widget_type == 'check': return controls.get(key, False)
            # For lab runs, the value is already a string or bool
            return controls.get(key, "")
        else: # It's a dict of widgets
            if not controls or key not in controls: return None
            widget = controls[key]
            if widget_type == 'check': return widget.isChecked()
            if widget_type == 'combo': return widget.currentText()
            if widget_type == 'text': return widget.text() # Return raw text, strip later
            if widget_type == 'enabled': return widget.isEnabled()
        return None

    def _build_nmap_script_args(self, controls):
        """Builds the --script and --script-args parts of the nmap command, accepting a controls dict or widget dict."""
        script_parts = []

        if self._get_control_value(controls, 'sc_check', 'check'): script_parts.append("default")
        if self._get_control_value(controls, 'nse_vuln_check', 'check'): script_parts.append("vuln")
        if self._get_control_value(controls, 'nse_discovery_check', 'check'): script_parts.append("discovery")
        if self._get_control_value(controls, 'nse_safe_check', 'check'): script_parts.append("safe")

        custom_scripts = self._get_control_value(controls, 'custom_script_edit', 'text')
        if custom_scripts and custom_scripts.strip():
            script_parts.append(custom_scripts.strip())

        command_args = []
        if script_parts:
            unique_scripts = sorted(list(set(script_parts)))
            command_args.extend(["--script", ",".join(unique_scripts)])

        script_args = self._get_control_value(controls, 'script_args_edit', 'text')
        if script_args and script_args.strip():
            command_args.extend(["--script-args", script_args.strip()])

        return command_args

    def _build_gobuster_command(self, controls):
        """Builds the Gobuster command list from a dictionary of controls or widgets."""
        url = self._get_control_value(controls, 'url_edit', 'text')
        wordlist = self._get_control_value(controls, 'wordlist_edit', 'text')

        if not url or not url.strip() or not wordlist or not wordlist.strip():
            return None, None, "Target URL and Wordlist are required for Gobuster."

        target_for_log = url.strip()
        command = ["gobuster", "dir", "-u", target_for_log, "-w", wordlist.strip()]

        if threads := self._get_control_value(controls, 'threads_edit', 'text'):
            if threads.strip(): command.extend(["-t", threads.strip()])
        if extensions := self._get_control_value(controls, 'extensions_edit', 'text'):
            if extensions.strip(): command.extend(["-x", extensions.strip()])
        if status_codes := self._get_control_value(controls, 'status_codes_edit', 'text'):
            if status_codes.strip(): command.extend(["-s", status_codes.strip()])
        if blacklist_codes := self._get_control_value(controls, 'status_codes_blacklist_edit', 'text'):
            if blacklist_codes.strip(): command.extend(["-b", blacklist_codes.strip()])
        if self._get_control_value(controls, 'add_slash_check', 'check'):
            command.append("-f")
        if self._get_control_value(controls, 'follow_redirect_check', 'check'):
            command.append("-r")
        if useragent := self._get_control_value(controls, 'useragent_edit', 'text'):
            if useragent.strip(): command.extend(["-a", useragent.strip()])
        if self._get_control_value(controls, 'random_agent_check', 'check'):
            command.append("--random-agent")
        if cookies := self._get_control_value(controls, 'cookies_edit', 'text'):
            if cookies.strip(): command.extend(["-c", cookies.strip()])
        if proxy := self._get_control_value(controls, 'proxy_edit', 'text'):
            if proxy.strip(): command.extend(["--proxy", proxy.strip()])
        if timeout := self._get_control_value(controls, 'timeout_edit', 'text'):
            if timeout.strip(): command.extend(["--timeout", timeout.strip()])
        if username := self._get_control_value(controls, 'username_edit', 'text'):
            if username.strip(): command.extend(["-U", username.strip()])
        if password := self._get_control_value(controls, 'password_edit', 'text'):
            if password: command.extend(["-P", password])
        if output_file := self._get_control_value(controls, 'output_file_edit', 'text'):
            if output_file.strip(): command.extend(["-o", output_file.strip()])
        if self._get_control_value(controls, 'no_progress_check', 'check'):
            command.append("-z")
        if self._get_control_value(controls, 'quiet_check', 'check'):
            command.append("-q")
        if self._get_control_value(controls, 'expanded_check', 'check'):
            command.append("-e")

        return command, target_for_log, None

    def _build_nikto_command(self, controls):
        """Builds the Nikto command list from a dictionary of controls or widgets."""
        target = self._get_control_value(controls, 'target_edit', 'text')
        if not target or not target.strip():
            return None, None, "A target is required for Nikto."
        target = target.strip()

        tool_path = self._get_tool_path("nikto", os.path.join("nikto", "bin", "nikto"))
        if not tool_path:
            tool_path = self._get_tool_path("nikto", os.path.join("nikto", "program", "nikto.pl"))
        if not tool_path:
            tool_path = self._get_tool_path("nikto.pl", os.path.join("nikto", "nikto.pl"))
        if not tool_path:
            return None, None, "Nikto not found."

        if tool_path.endswith(".pl"):
            command = ["perl", tool_path, "-host", target]
        else:
            command = [tool_path, "-host", target]

        if port := self._get_control_value(controls, 'port_edit', 'text'):
            if port.strip(): command.extend(["-port", port.strip()])
        if self._get_control_value(controls, 'ssl_check', 'check'):
            command.append("-ssl")
        if vhost := self._get_control_value(controls, 'vhost_edit', 'text'):
            if vhost.strip(): command.extend(["-vhost", vhost.strip()])
        if (tuning_text := self._get_control_value(controls, 'tuning_combo', 'combo')) != "Default":
            command.extend(["-Tuning", tuning_text.split(" ")[0]])
        if (mutate_text := self._get_control_value(controls, 'mutate_combo', 'combo')) != "None":
            command.extend(["-mutate", mutate_text.split(" ")[0]])
        if plugins := self._get_control_value(controls, 'plugins_edit', 'text'):
            if plugins.strip(): command.extend(["-Plugins", plugins.strip()])
        if cgidirs := self._get_control_value(controls, 'cgidirs_edit', 'text'):
            if cgidirs.strip(): command.extend(["-Cgidirs", cgidirs.strip()])
        if (evasion_text := self._get_control_value(controls, 'evasion_combo', 'combo')) != "None":
            command.extend(["-evasion", evasion_text.split(" ")[0]])
        if (timeout := self._get_control_value(controls, 'timeout_edit', 'text')) != "10":
            if timeout.strip(): command.extend(["-timeout", timeout.strip()])
        if maxtime := self._get_control_value(controls, 'maxtime_edit', 'text'):
            if maxtime.strip(): command.extend(["-maxtime", maxtime.strip()])
        if pause := self._get_control_value(controls, 'pause_edit', 'text'):
            if pause.strip(): command.extend(["-Pause", pause.strip()])
        if auth_id := self._get_control_value(controls, 'id_edit', 'text'):
            if auth_id.strip(): command.extend(["-id", auth_id.strip()])
        if root := self._get_control_value(controls, 'root_edit', 'text'):
            if root.strip(): command.extend(["-root", root.strip()])
        if self._get_control_value(controls, 'proxy_check', 'check'):
            command.append("-useproxy")
        if output_file := self._get_control_value(controls, 'output_file_edit', 'text'):
            if output_file.strip():
                output_format = self._get_control_value(controls, 'format_combo', 'combo')
                command.extend(["-o", output_file.strip(), "-Format", output_format])
        if save_dir := self._get_control_value(controls, 'save_dir_edit', 'text'):
            if save_dir.strip(): command.extend(["-Save", save_dir.strip()])
        if extra_opts := self._get_control_value(controls, 'extra_opts_edit', 'text'):
            if extra_opts.strip(): command.extend(extra_opts.strip().split())

        return command, target, None

    def _build_whatweb_command(self, controls):
        """Builds the WhatWeb command list from a dictionary of controls or widgets."""
        target = self._get_control_value(controls, 'target_edit', 'text')
        if not target or not target.strip():
            return None, None, "A target is required."

        tool_path = self._get_tool_path("whatweb", os.path.join("WhatWeb", "bin", "whatweb"))
        if not tool_path:
            tool_path = self._get_tool_path("whatweb", os.path.join("WhatWeb", "whatweb"))
        if not tool_path:
            return None, None, "WhatWeb not found."

        resolved_target = self._resolve_targets_string(target.strip())
        command = [tool_path]

        aggression_text = self._get_control_value(controls, 'aggression_combo', 'combo')
        aggression_level = aggression_text.split(" ")[0] if aggression_text else "1"
        command.extend(["-a", aggression_level])

        if self._get_control_value(controls, 'verbose_check', 'check'):
            command.append("-v")

        if extra_opts := self._get_control_value(controls, 'extra_opts_edit', 'text'):
            if extra_opts.strip():
                command.extend(extra_opts.strip().split())

        command.extend(resolved_target.split())
        return command, resolved_target, None

    def _build_fierce_command(self, controls):
        """Builds the fierce command list from a dictionary of controls or widgets."""
        domain = self._get_control_value(controls, 'domain_edit', 'text')
        if not domain or not domain.strip():
            return None, None, "A domain is required for fierce."
        domain = domain.strip()

        tool_path = self._get_tool_path("fierce", "fierce")
        if not tool_path:
             return None, None, "fierce not found."

        command = [tool_path, "--domain", domain]

        if self._get_control_value(controls, 'connect_check', 'check'):
            command.append("--connect")
        if self._get_control_value(controls, 'wide_check', 'check'):
            command.append("--wide")
        if traverse := self._get_control_value(controls, 'traverse_edit', 'text'):
            if traverse.strip():
                command.extend(["--traverse", traverse.strip()])
        if delay := self._get_control_value(controls, 'delay_edit', 'text'):
            if delay.strip():
                command.extend(["--delay", delay.strip()])

        return command, domain, None

    def _build_dnsrecon_command(self, controls):
        """Builds the dnsrecon command list from a dictionary of controls or widgets."""
        domain = self._get_control_value(controls, 'domain_edit', 'text')
        if not domain or not domain.strip():
            return None, None, "A domain is required for dnsrecon."
        domain = domain.strip()

        tool_path = self._get_tool_path("dnsrecon", "dnsrecon")
        if not tool_path:
             return None, None, "dnsrecon not found."

        command = [tool_path, "-d", domain]

        scan_type = self._get_control_value(controls, 'scan_type_combo', 'combo')
        if scan_type:
            command.extend(["-t", scan_type])

        if scan_type == 'brt':
            if dictionary := self._get_control_value(controls, 'dict_edit', 'text'):
                if dictionary.strip():
                    command.extend(["-D", dictionary.strip()])

        if json_output := self._get_control_value(controls, 'json_output_edit', 'text'):
            if json_output.strip():
                command.extend(["--json", json_output.strip()])

        return command, domain, None

    def _build_enum4linux_ng_command(self, controls):
        """Builds the enum4linux-ng command list from a dictionary of controls or widgets."""
        target = self._get_control_value(controls, 'target_edit', 'text')
        if not target or not target.strip():
            return None, None, "A target host is required for enum4linux-ng."
        target = target.strip()

        tool_path = self._get_tool_path("enum4linux-ng", "enum4linux-ng")
        if not tool_path:
             # Try .py fallback if installed via pip or repo
             tool_path = self._get_tool_path("enum4linux-ng.py", os.path.join("enum4linux-ng", "enum4linux-ng.py"))
             if tool_path and tool_path.endswith(".py"):
                 command = ["python3", tool_path]
             elif tool_path:
                 command = [tool_path]
             else:
                 return None, None, "enum4linux-ng not found."
        else:
             command = [tool_path]

        if self._get_control_value(controls, 'all_check', 'check'):
            command.append("-A")
        else:
            if self._get_control_value(controls, 'users_check', 'check'): command.append("-U")
            if self._get_control_value(controls, 'groups_check', 'check'): command.append("-G")
            if self._get_control_value(controls, 'shares_check', 'check'): command.append("-S")
            if self._get_control_value(controls, 'policy_check', 'check'): command.append("-P")
            if self._get_control_value(controls, 'os_check', 'check'): command.append("-O")

        if user := self._get_control_value(controls, 'user_edit', 'text'):
            if user.strip():
                command.extend(["-u", user.strip()])
        if pwd := self._get_control_value(controls, 'pass_edit', 'text'):
            if pwd: # Don't strip password
                command.extend(["-p", pwd])

        command.append(target)
        return command, target, None

    def _build_ffuf_command(self, controls):
        """Builds the ffuf command list from a dictionary of controls or widgets."""
        url = self._get_control_value(controls, 'url_edit', 'text')
        wordlist = self._get_control_value(controls, 'wordlist_edit', 'text')

        if not url or not url.strip() or not wordlist or not wordlist.strip():
            return None, None, "Target URL and Wordlist are required for ffuf."

        if "FUZZ" not in url:
            return None, None, "Target URL must contain the 'FUZZ' keyword."

        tool_path = self._get_tool_path("ffuf", "ffuf")
        if not tool_path:
             return None, None, "ffuf not found."

        target_for_log = url.strip()
        command = [tool_path, "-u", target_for_log, "-w", wordlist.strip()]

        if extensions := self._get_control_value(controls, 'extensions_edit', 'text'):
            if extensions.strip():
                command.extend(["-e", extensions.strip()])
        if threads := self._get_control_value(controls, 'threads_edit', 'text'):
            if threads.strip():
                command.extend(["-t", threads.strip()])
        if method := self._get_control_value(controls, 'method_edit', 'text'):
            if method.strip():
                command.extend(["-X", method.strip()])
        if mc := self._get_control_value(controls, 'match_codes_edit', 'text'):
            if mc.strip():
                command.extend(["-mc", mc.strip()])
        if fc := self._get_control_value(controls, 'filter_codes_edit', 'text'):
            if fc.strip():
                command.extend(["-fc", fc.strip()])

        return command, target_for_log, None

    def _build_dirsearch_command(self, controls):
        """Builds the dirsearch command list from a dictionary of controls or widgets."""
        url = self._get_control_value(controls, 'url_edit', 'text')
        wordlist = self._get_control_value(controls, 'wordlist_edit', 'text')

        if not url or not url.strip() or not wordlist or not wordlist.strip():
            return None, None, "Target URL and Wordlist are required for dirsearch."

        tool_path = self._get_tool_path("dirsearch", "dirsearch")
        if not tool_path:
             # Try .py fallback
             tool_path = self._get_tool_path("dirsearch.py", os.path.join("dirsearch", "dirsearch.py"))
             if tool_path and tool_path.endswith(".py"):
                 command = ["python3", tool_path]
             elif tool_path:
                 command = [tool_path]
             else:
                 return None, None, "dirsearch not found."
        else:
             command = [tool_path]

        target_for_log = url.strip()
        command.extend(["-u", target_for_log, "-w", wordlist.strip()])

        # Scan Tab
        if extensions := self._get_control_value(controls, 'extensions_edit', 'text'):
            if extensions.strip(): command.extend(["-e", extensions.strip()])
        if threads := self._get_control_value(controls, 'threads_edit', 'text'):
            if threads.strip(): command.extend(["-t", threads.strip()])
        if self._get_control_value(controls, 'recursive_check', 'check'):
            command.append("-r")

        # Request Tab
        if method := self._get_control_value(controls, 'method_edit', 'text'):
            if method.strip().upper() != 'GET': command.extend(["-m", method.strip().upper()])
        if headers := self._get_control_value(controls, 'headers_edit', 'text'): # QTextEdit
             if headers.strip():
                 for header in headers.strip().split('\n'):
                     if header.strip(): command.extend(["-H", header.strip()])
        if timeout := self._get_control_value(controls, 'timeout_edit', 'text'):
            if timeout.strip(): command.extend(["--timeout", timeout.strip()])

        # Filter Tab
        if include_status := self._get_control_value(controls, 'include_status_edit', 'text'):
            if include_status.strip(): command.extend(["-i", include_status.strip()])
        if exclude_status := self._get_control_value(controls, 'exclude_status_edit', 'text'):
            if exclude_status.strip(): command.extend(["-x", exclude_status.strip()])

        # Output Tab (user-specified)
        if output_file := self._get_control_value(controls, 'output_edit', 'text'):
            if output_file.strip():
                command.extend(["--json-report", output_file.strip()])

        command.append("--no-color")
        return command, target_for_log, None

    def _build_rustscan_command(self, controls):
        """Builds the RustScan command list from a dictionary of controls or widgets."""
        targets = self._get_control_value(controls, 'targets_edit', 'text')
        if not targets or not targets.strip():
            return None, None, "At least one target is required for RustScan."
        targets = targets.strip()
        target_for_log = targets

        tool_path = self._get_tool_path("rustscan", "rustscan")
        if not tool_path:
             return None, None, "RustScan not found."

        command = [tool_path, "-a", targets]

        if ports := self._get_control_value(controls, 'ports_edit', 'text'):
            if ports.strip():
                command.extend(["-p", ports.strip()])
        if batch_size := self._get_control_value(controls, 'batch_size_edit', 'text'):
            if batch_size.strip():
                command.extend(["-b", batch_size.strip()])
        if timeout := self._get_control_value(controls, 'timeout_edit', 'text'):
            if timeout.strip():
                command.extend(["-T", timeout.strip()])

        if self._get_control_value(controls, 'quiet_check', 'check'):
            command.append("-q")
        else:
            if nmap_args := self._get_control_value(controls, 'nmap_args_edit', 'text'):
                if nmap_args.strip():
                    command.append("--")
                    command.extend(nmap_args.strip().split())

        return command, target_for_log, None

    def _build_httpx_command(self, controls):
        """Builds the httpx command list from a dictionary of controls or widgets."""
        target_list = self._get_control_value(controls, 'target_list_edit', 'text')
        if not target_list or not target_list.strip():
            return None, None, "A target list file is required for httpx."
        target_list = target_list.strip()
        target_for_log = target_list

        tool_path = self._get_tool_path("httpx", "httpx")
        if not tool_path:
             return None, None, "httpx not found."

        command = [tool_path, "-l", target_list, "-silent"]

        probe_flags = {
            'probe_status_code': '-status-code',
            'probe_title': '-title',
            'probe_tech_detect': '-tech-detect',
            'probe_web_server': '-web-server',
            'probe_cdn': '-cdn',
            'probe_jarm': '-jarm',
        }
        for control_name, flag in probe_flags.items():
            if self._get_control_value(controls, control_name, 'check'):
                command.append(flag)

        if ports := self._get_control_value(controls, 'ports_edit', 'text'):
            if ports.strip():
                command.extend(["-ports", ports.strip()])

        if self._get_control_value(controls, 'json_output_check', 'check'):
            command.append("-json")

        return command, target_for_log, None

    def _build_subfinder_command(self, controls):
        """Builds the Subfinder command list from a dictionary of controls or widgets."""
        domain = self._get_control_value(controls, 'domain_edit', 'text')
        if not domain or not domain.strip():
            return None, None, "A domain is required for Subfinder."
        domain = domain.strip()
        target_for_log = domain

        tool_path = self._get_tool_path("subfinder", "subfinder")
        if not tool_path:
             return None, None, "subfinder not found."

        command = [tool_path, "-d", domain, "-silent"]

        if self._get_control_value(controls, 'recursive_check', 'check'):
            command.append("-recursive")
        if self._get_control_value(controls, 'all_sources_check', 'check'):
            command.append("-all")
        if output_file := self._get_control_value(controls, 'output_edit', 'text'):
             if output_file.strip():
                command.extend(["-o", output_file.strip()])

        return command, target_for_log, None

    def _build_sublist3r_command(self, controls):
        """Builds the Sublist3r command list from a dictionary of controls or widgets."""
        domain = self._get_control_value(controls, 'domain_edit', 'text')
        if not domain or not domain.strip():
            return None, None, "A domain is required for Sublist3r."
        domain = domain.strip()
        # The target for logging is the domain itself.
        target_for_log = domain

        tool_path = self._get_tool_path("sublist3r.py", os.path.join("sublist3r", "sublist3r.py"))
        if tool_path:
             command = ["python3", tool_path, "-d", domain]
        else:
             return None, None, "Sublist3r script not found."

        return command, target_for_log, None

    def _build_masscan_command(self, controls):
        """Builds the Masscan command list from a dictionary of controls or widgets."""
        target = self._get_control_value(controls, 'target_edit', 'text')
        if not target or not target.strip():
            return None, None, "Target is required for Masscan."

        tool_path = self._get_tool_path("masscan", "masscan")
        if not tool_path:
             return None, None, "Masscan not found."

        command = [tool_path, target.strip()]

        ports = self._get_control_value(controls, 'ports_edit', 'text')
        if ports:
            command.extend(["-p", ports.strip()])
        else:
            command.extend(["-p", "0-65535"]) # Default to all ports if not specified

        rate = self._get_control_value(controls, 'rate_edit', 'text')
        if rate:
            command.extend(["--rate", rate.strip()])

        return command, target.strip(), None

    def _build_nmap_command(self, controls, target):
        """Builds the Nmap command list from a dictionary of controls or widgets."""
        if not target or not target.strip():
            return None, None, "Target is required for Nmap scan."

        target_for_log = target.strip()

        tool_path = self._get_tool_path("nmap", "nmap")
        if not tool_path:
             return None, None, "Nmap not found."
        command = [tool_path]

        # --- Host Discovery Tab ---
        if self._get_control_value(controls, 'ping_scan_check', 'check'): command.append("-sn")
        if self._get_control_value(controls, 'pn_check', 'check'): command.append("-Pn")
        if self._get_control_value(controls, 'ps_check', 'check'): command.append("-PS")
        if self._get_control_value(controls, 'pa_check', 'check'): command.append("-PA")
        if self._get_control_value(controls, 'pu_check', 'check'): command.append("-PU")
        if self._get_control_value(controls, 'pr_check', 'check'): command.append("-PR")
        if self._get_control_value(controls, 'pe_check', 'check'): command.append("-PE")
        if self._get_control_value(controls, 'pp_check', 'check'): command.append("-PP")
        if self._get_control_value(controls, 'disable_arp_ping_check', 'check'): command.append("--disable-arp-ping")


        # --- Scan & Detection Tab ---
        scan_type_text = self._get_control_value(controls, 'scan_type_combo', 'combo')
        if scan_type_text: command.append(scan_type_text.split(" ")[-1].strip("()"))

        if self._get_control_value(controls, 'a_check', 'check'):
            command.append("-A")
        else:
            if self._get_control_value(controls, 'sv_check', 'check'): command.append("-sV")
            if self._get_control_value(controls, 'o_check', 'check'): command.append("-O")
            if self._get_control_value(controls, 'sc_check', 'check'): command.append("-sC")

        if self._get_control_value(controls, 'reason_check', 'check'): command.append("--reason")
        if self._get_control_value(controls, 'osscan_guess_check', 'check'): command.append("--osscan-guess")
        if self._get_control_value(controls, 'version_light_check', 'check'): command.append("--version-light")
        if self._get_control_value(controls, 'version_all_check', 'check'): command.append("--version-all")
        if intensity := self._get_control_value(controls, 'version_intensity_edit', 'text'):
            if intensity.strip(): command.extend(["--version-intensity", intensity.strip()])


        # --- Performance Tab ---
        timing_text = self._get_control_value(controls, 'timing_combo', 'combo')
        if timing_text and len(timing_text) > 1: command.append(f"-T{timing_text[1]}")
        if min_rate := self._get_control_value(controls, 'min_rate_edit', 'text'):
            if min_rate.strip(): command.extend(["--min-rate", min_rate.strip()])
        if max_rate := self._get_control_value(controls, 'max_rate_edit', 'text'):
            if max_rate.strip(): command.extend(["--max-rate", max_rate.strip()])
        if min_parallel := self._get_control_value(controls, 'min_parallel_edit', 'text'):
            if min_parallel.strip(): command.extend(["--min-parallelism", min_parallel.strip()])
        if max_parallel := self._get_control_value(controls, 'max_parallel_edit', 'text'):
            if max_parallel.strip(): command.extend(["--max-parallelism", max_parallel.strip()])
        if min_hostgroup := self._get_control_value(controls, 'min_hostgroup_edit', 'text'):
            if min_hostgroup.strip(): command.extend(["--min-hostgroup", min_hostgroup.strip()])
        if max_hostgroup := self._get_control_value(controls, 'max_hostgroup_edit', 'text'):
            if max_hostgroup.strip(): command.extend(["--max-hostgroup", max_hostgroup.strip()])
        if scan_delay := self._get_control_value(controls, 'scan_delay_edit', 'text'):
            if scan_delay.strip(): command.extend(["--scan-delay", scan_delay.strip()])
        if max_retries := self._get_control_value(controls, 'max_retries_edit', 'text'):
            if max_retries.strip(): command.extend(["--max-retries", max_retries.strip()])
        if host_timeout := self._get_control_value(controls, 'host_timeout_edit', 'text'):
            if host_timeout.strip(): command.extend(["--host-timeout", host_timeout.strip()])

        # --- Evasion Tab ---
        if self._get_control_value(controls, 'frag_check', 'check'): command.append("-f")
        if mtu := self._get_control_value(controls, 'mtu_edit', 'text'):
            if mtu.strip(): command.extend(["--mtu", mtu.strip()])
        if self._get_control_value(controls, 'badsum_check', 'check'): command.append("--badsum")
        if source_port := self._get_control_value(controls, 'source_port_edit', 'text'):
            if source_port.strip(): command.extend(["-g", source_port.strip()])
        if source_host := self._get_control_value(controls, 'source_host_edit', 'text'):
            if source_host.strip(): command.extend(["-S", source_host.strip()])
        if decoys := self._get_control_value(controls, 'decoys_edit', 'text'):
            if decoys.strip(): command.extend(["-D", decoys.strip()])
        if zombie := self._get_control_value(controls, 'zombie_edit', 'text'):
            if zombie.strip(): command.extend(["-sI", zombie.strip()])
        if data_len := self._get_control_value(controls, 'data_len_edit', 'text'):
            if data_len.strip(): command.extend(["--data-length", data_len.strip()])
        if spoof_mac := self._get_control_value(controls, 'spoof_mac_edit', 'text'):
            if spoof_mac.strip(): command.extend(["--spoof-mac", spoof_mac.strip()])
        if data_string := self._get_control_value(controls, 'data_string_edit', 'text'):
            if data_string.strip(): command.extend(["--data-string", data_string.strip()])
        if ip_options := self._get_control_value(controls, 'ip_options_edit', 'text'):
            if ip_options.strip(): command.extend(["--ip-options", ip_options.strip()])
        if ttl := self._get_control_value(controls, 'ttl_edit', 'text'):
            if ttl.strip(): command.extend(["--ttl", ttl.strip()])
        if self._get_control_value(controls, 'randomize_hosts_check', 'check'):
            command.append("--randomize-hosts")

        # --- NSE Tab ---
        command.extend(self._build_nmap_script_args(controls))

        # --- Output & Misc Tab ---
        if output_normal := self._get_control_value(controls, 'output_normal_edit', 'text'):
            if output_normal.strip(): command.extend(["-oN", output_normal.strip()])
        if output_grepable := self._get_control_value(controls, 'output_grepable_edit', 'text'):
            if output_grepable.strip(): command.extend(["-oG", output_grepable.strip()])
        if resume_file := self._get_control_value(controls, 'resume_edit', 'text'):
            if resume_file.strip(): command.extend(["--resume", resume_file.strip()])
        if self._get_control_value(controls, 'append_output_check', 'check'): command.append("--append-output")
        if self._get_control_value(controls, 'v_check', 'check'): command.append("-v")
        if self._get_control_value(controls, 'packet_trace_check', 'check'): command.append("--packet-trace")
        if self._get_control_value(controls, 'traceroute_check', 'check'): command.append("--traceroute")
        if self._get_control_value(controls, 'ipv6_check', 'check'): command.append("-6")


        # --- Ports (from top controls) ---
        if "-sn" not in command:
            if self._get_control_value(controls, 'top_ports_check', 'check'):
                if top_ports_val := self._get_control_value(controls, 'top_ports_edit', 'text'):
                    if top_ports_val.strip():
                        command.extend(["--top-ports", top_ports_val.strip()])
            elif ports := self._get_control_value(controls, 'ports_edit', 'text'):
                if ports.strip():
                    command.extend(["-p", ports.strip()])

        # --- Target ---
        command.append(target_for_log)

        return command, target_for_log, None

    def _execute_lab_tool(self, tool_name, controls, output_widget):
        """
        A generic, asynchronous executor for any command-line tool that has a
        config widget and an output console. It starts the tool in a WorkerThread.
        This is used by the main UI tabs, NOT by the LAB chain executor.
        """
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        # --- Command Building ---
        command = []
        target_for_log = ""

        if tool_name == "Nmap Scan":
            target_for_log = controls['target_edit'].text()
            if not target_for_log:
                QMessageBox.critical(self, "Input Error", "Please provide a target for the Nmap scan.")
                return
            command, _, error = self._build_nmap_command(controls, target_for_log)
            if error:
                QMessageBox.critical(self, "Input Error", error)
                return
        else:
            QMessageBox.critical(self, "Error", f"Execution logic for '{tool_name}' is not implemented yet.")
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, f"{tool_name} Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        # --- UI State Update ---
        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        if 'report_btn' in controls:
            controls['report_btn'].setEnabled(False)
        self.tool_stop_event.clear()
        output_widget.clear()

        # --- Thread Execution ---
        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, f'{tool_name.lower().replace(" ", "_")}_scan', target_for_log, output_widget
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_phoneinfoga_tool(self):
        """Creates the UI for the PhoneInfoga OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>PhoneInfoga - Phone Number Scanner</b></font>
        <p>One of the most advanced tools to scan phone numbers using only free resources.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.phoneinfoga_controls = self._create_phoneinfoga_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.phoneinfoga_controls['start_btn'])
        buttons_layout.addWidget(self.phoneinfoga_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.phoneinfoga_output_console = QPlainTextEdit()
        self.phoneinfoga_output_console.setReadOnly(True)
        self.phoneinfoga_output_console.setFont(QFont("Courier New", 10))
        self.phoneinfoga_output_console.setPlaceholderText("PhoneInfoga output will be displayed here...")
        main_layout.addWidget(self.phoneinfoga_output_console, 1)

        self.phoneinfoga_controls['start_btn'].clicked.connect(self.start_phoneinfoga_scan)
        self.phoneinfoga_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_phoneinfoga_config_widget(self):
        """Creates a reusable, self-contained widget for PhoneInfoga's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['number_edit'] = QLineEdit()
        controls['number_edit'].setPlaceholderText("e.g., +15551234567")
        layout.addRow("Phone Number (-n):", controls['number_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_phoneinfoga_command(self, controls):
        """Builds the PhoneInfoga command list from a dictionary of controls or widgets."""
        command = ["./tools/PhoneInfoga/phoneinfoga", "scan"]

        number = self._get_control_value(controls, 'number_edit', 'text')
        if not number or not number.strip():
            return None, None, "A phone number is required."
        command.extend(["-n", number.strip()])
        target_for_log = number.strip()

        return command, target_for_log, None

    def start_phoneinfoga_scan(self):
        """Starts the PhoneInfoga scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.phoneinfoga_controls
        command, target_for_log, error = self._build_phoneinfoga_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        phoneinfoga_path = os.path.join(self.script_dir, "tools", "PhoneInfoga", "phoneinfoga")
        if not os.path.exists(phoneinfoga_path):
            QMessageBox.critical(self, "PhoneInfoga Error", f"'phoneinfoga' executable not found at {phoneinfoga_path}. Please build it first.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.phoneinfoga_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'phoneinfoga_scan', target_for_log, self.phoneinfoga_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_metagoofil_tool(self):
        """Creates the UI for the Metagoofil OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Metagoofil - Metadata Collector</b></font>
        <p>An information gathering tool designed for extracting metadata of public documents (pdf,doc,xls,ppt,etc) available on target websites.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.metagoofil_controls = self._create_metagoofil_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.metagoofil_controls['start_btn'])
        buttons_layout.addWidget(self.metagoofil_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.metagoofil_output_console = QPlainTextEdit()
        self.metagoofil_output_console.setReadOnly(True)
        self.metagoofil_output_console.setFont(QFont("Courier New", 10))
        self.metagoofil_output_console.setPlaceholderText("Metagoofil output will be displayed here...")
        main_layout.addWidget(self.metagoofil_output_console, 1)

        self.metagoofil_controls['start_btn'].clicked.connect(self.start_metagoofil_scan)
        self.metagoofil_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_metagoofil_config_widget(self):
        """Creates a reusable, self-contained widget for Metagoofil's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['domain_edit'] = QLineEdit("example.com")
        layout.addRow("Domain (-d):", controls['domain_edit'])

        controls['file_types_edit'] = QLineEdit("pdf,doc,xls")
        layout.addRow("File Types (-t):", controls['file_types_edit'])

        controls['search_max_edit'] = QLineEdit("100")
        layout.addRow("Search Max (-l):", controls['search_max_edit'])

        controls['download_limit_edit'] = QLineEdit("100")
        layout.addRow("Download Limit (-n):", controls['download_limit_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Metagoofil Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_metagoofil_command(self, controls):
        """Builds the Metagoofil command list from a dictionary of controls or widgets."""
        command = ["python3", "tools/metagoofil/metagoofil.py"]

        domain = self._get_control_value(controls, 'domain_edit', 'text')
        if not domain or not domain.strip():
            return None, None, "A domain is required."
        command.extend(["-d", domain.strip()])
        target_for_log = domain.strip()

        file_types = self._get_control_value(controls, 'file_types_edit', 'text')
        if not file_types or not file_types.strip():
            return None, None, "File types are required."
        command.extend(["-t", file_types.strip()])

        if search_max := self._get_control_value(controls, 'search_max_edit', 'text'):
            command.extend(["-l", search_max.strip()])

        if download_limit := self._get_control_value(controls, 'download_limit_edit', 'text'):
            command.extend(["-n", download_limit.strip()])

        return command, target_for_log, None

    def start_metagoofil_scan(self):
        """Starts the Metagoofil scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.metagoofil_controls
        command, target_for_log, error = self._build_metagoofil_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        metagoofil_path = os.path.join(self.script_dir, "tools", "metagoofil", "metagoofil.py")
        if not os.path.exists(metagoofil_path):
            QMessageBox.critical(self, "Metagoofil Error", f"'metagoofil.py' not found at {metagoofil_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.metagoofil_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'metagoofil_scan', target_for_log, self.metagoofil_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_social_analyzer_tool(self):
        """Creates the UI for the Social Analyzer OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Social Analyzer - Social Media Finder</b></font>
        <p>Analyzes & finds social media profiles across 300+ platforms. It's a CLI tool for analyzing and finding a person's profile across 300+ social media websites.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.social_analyzer_controls = self._create_social_analyzer_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.social_analyzer_controls['start_btn'])
        buttons_layout.addWidget(self.social_analyzer_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.social_analyzer_output_console = QPlainTextEdit()
        self.social_analyzer_output_console.setReadOnly(True)
        self.social_analyzer_output_console.setFont(QFont("Courier New", 10))
        self.social_analyzer_output_console.setPlaceholderText("Social Analyzer output will be displayed here...")
        main_layout.addWidget(self.social_analyzer_output_console, 1)

        self.social_analyzer_controls['start_btn'].clicked.connect(self.start_social_analyzer_scan)
        self.social_analyzer_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_social_analyzer_config_widget(self):
        """Creates a reusable, self-contained widget for Social Analyzer's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['username_edit'] = QLineEdit()
        controls['username_edit'].setPlaceholderText("Enter a single username to search")
        layout.addRow("Username:", controls['username_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_social_analyzer_command(self, controls):
        """Builds the Social Analyzer command list from a dictionary of controls or widgets."""
        command = ["python3", "tools/social-analyzer/main.py", "--cli", "--silent"]

        username = self._get_control_value(controls, 'username_edit', 'text')
        if not username or not username.strip():
            return None, None, "A username is required."
        command.extend(["--username", username.strip()])
        target_for_log = username.strip()

        return command, target_for_log, None

    def start_social_analyzer_scan(self):
        """Starts the Social Analyzer scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.social_analyzer_controls
        command, target_for_log, error = self._build_social_analyzer_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        social_analyzer_path = os.path.join(self.script_dir, "tools", "social-analyzer", "main.py")
        if not os.path.exists(social_analyzer_path):
            QMessageBox.critical(self, "Social Analyzer Error", f"'main.py' not found at {social_analyzer_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.social_analyzer_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'social_analyzer_scan', target_for_log, self.social_analyzer_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_photon_tool(self):
        """Creates the UI for the Photon OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Photon - OSINT & Crawler</b></font>
        <p>An incredibly fast crawler designed for OSINT. Extracts URLs, files, intel & endpoints from a target.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.photon_controls = self._create_photon_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.photon_controls['start_btn'])
        buttons_layout.addWidget(self.photon_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.photon_output_console = QPlainTextEdit()
        self.photon_output_console.setReadOnly(True)
        self.photon_output_console.setFont(QFont("Courier New", 10))
        self.photon_output_console.setPlaceholderText("Photon output will be displayed here...")
        main_layout.addWidget(self.photon_output_console, 1)

        self.photon_controls['start_btn'].clicked.connect(self.start_photon_scan)
        self.photon_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_photon_config_widget(self):
        """Creates a reusable, self-contained widget for Photon's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['url_edit'] = QLineEdit("http://example.com")
        layout.addRow("Root URL (-u):", controls['url_edit'])

        controls['level_edit'] = QLineEdit("2")
        controls['level_edit'].setToolTip("Levels to crawl (-l).")
        layout.addRow("Crawl Level (-l):", controls['level_edit'])

        controls['threads_edit'] = QLineEdit("10")
        controls['threads_edit'].setToolTip("Number of threads (-t).")
        layout.addRow("Threads (-t):", controls['threads_edit'])

        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("e.g., results/")
        controls['output_edit'].setToolTip("Output directory (-o).")
        layout.addRow("Output Directory (-o):", controls['output_edit'])

        # Checkboxes for boolean flags
        checkbox_layout = QGridLayout()
        controls['dns_check'] = QCheckBox("DNS (--dns)")
        controls['keys_check'] = QCheckBox("Secret Keys (--keys)")
        controls['wayback_check'] = QCheckBox("Wayback Machine (--wayback)")
        controls['only_urls_check'] = QCheckBox("Only URLs (--only-urls)")
        checkbox_layout.addWidget(controls['dns_check'], 0, 0)
        checkbox_layout.addWidget(controls['keys_check'], 0, 1)
        checkbox_layout.addWidget(controls['wayback_check'], 1, 0)
        checkbox_layout.addWidget(controls['only_urls_check'], 1, 1)
        layout.addRow(checkbox_layout)

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Photon Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_photon_command(self, controls):
        """Builds the Photon command list from a dictionary of controls or widgets."""
        command = ["python3", "tools/Photon/photon.py"]

        url = self._get_control_value(controls, 'url_edit', 'text')
        if not url or not url.strip():
            return None, None, "A root URL is required."
        command.extend(["-u", url.strip()])
        target_for_log = url.strip()

        if level := self._get_control_value(controls, 'level_edit', 'text'):
            command.extend(["-l", level.strip()])
        if threads := self._get_control_value(controls, 'threads_edit', 'text'):
            command.extend(["-t", threads.strip()])
        if output := self._get_control_value(controls, 'output_edit', 'text'):
            command.extend(["-o", output.strip()])

        if self._get_control_value(controls, 'dns_check', 'check'): command.append("--dns")
        if self._get_control_value(controls, 'keys_check', 'check'): command.append("--keys")
        if self._get_control_value(controls, 'wayback_check', 'check'): command.append("--wayback")
        if self._get_control_value(controls, 'only_urls_check', 'check'): command.append("--only-urls")

        return command, target_for_log, None

    def start_photon_scan(self):
        """Starts the Photon scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.photon_controls
        command, target_for_log, error = self._build_photon_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        photon_path = os.path.join(self.script_dir, "tools", "Photon", "photon.py")
        if not os.path.exists(photon_path):
            QMessageBox.critical(self, "Photon Error", f"'photon.py' not found at {photon_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.photon_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'photon_scan', target_for_log, self.photon_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_infoga_tool(self):
        """Creates the UI for the Infoga OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Infoga - Email OSINT</b></font>
        <p>A tool for gathering email account information from various public sources.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.infoga_controls = self._create_infoga_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.infoga_controls['start_btn'])
        buttons_layout.addWidget(self.infoga_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.infoga_output_console = QPlainTextEdit()
        self.infoga_output_console.setReadOnly(True)
        self.infoga_output_console.setFont(QFont("Courier New", 10))
        self.infoga_output_console.setPlaceholderText("Infoga output will be displayed here...")
        main_layout.addWidget(self.infoga_output_console, 1)

        self.infoga_controls['start_btn'].clicked.connect(self.start_infoga_scan)
        self.infoga_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_infoga_config_widget(self):
        """Creates a reusable, self-contained widget for Infoga's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['domain_edit'] = QLineEdit()
        controls['domain_edit'].setPlaceholderText("Enter a domain to search (e.g., example.com)")
        layout.addRow("Domain:", controls['domain_edit'])

        controls['source_edit'] = QLineEdit("all")
        controls['source_edit'].setToolTip("Data source: all, google, twitter, github, etc.")
        layout.addRow("Source:", controls['source_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_infoga_command(self, controls):
        """Builds the Infoga command list from a dictionary of controls or widgets."""
        command = ["python3", "tools/Infoga/infoga.py"]

        domain = self._get_control_value(controls, 'domain_edit', 'text')
        if not domain or not domain.strip():
            return None, None, "A domain is required."
        command.extend(["--domain", domain.strip()])
        target_for_log = domain.strip()

        source = self._get_control_value(controls, 'source_edit', 'text')
        if source and source.strip():
            command.extend(["--source", source.strip()])

        return command, target_for_log, None

    def start_infoga_scan(self):
        """Starts the Infoga scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.infoga_controls
        command, target_for_log, error = self._build_infoga_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        infoga_path = os.path.join(self.script_dir, "tools", "Infoga", "infoga.py")
        if not os.path.exists(infoga_path):
            QMessageBox.critical(self, "Infoga Error", f"'infoga.py' not found at {infoga_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.infoga_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'infoga_scan', target_for_log, self.infoga_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_sherlock_config_widget(self):
        """Creates a reusable, self-contained widget with Sherlock's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['usernames_edit'] = QLineEdit()
        controls['usernames_edit'].setToolTip("One or more usernames to check, separated by spaces.")
        main_layout.addRow("Usernames:", controls['usernames_edit'])

        controls['timeout_edit'] = QLineEdit("60")
        controls['timeout_edit'].setToolTip("Timeout in seconds for each request.")
        main_layout.addRow("Timeout (--timeout):", controls['timeout_edit'])

        output_file_layout = QHBoxLayout()
        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("Optional: path to save text report...")
        output_file_layout.addWidget(controls['output_edit'])
        browse_output_btn = QPushButton("Browse...")
        browse_output_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_edit'], "Save Sherlock Report"))
        output_file_layout.addWidget(browse_output_btn)
        main_layout.addRow("Output File (-o):", output_file_layout)

        controls['csv_check'] = QCheckBox("Export as CSV")
        main_layout.addRow("--csv:", controls['csv_check'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Hunt Usernames")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_sherlock_command(self, controls, csv_temp_dir=None):
        """Builds the Sherlock command list from a dictionary of controls or widgets."""
        sherlock_executable = os.path.join(self.script_dir, "tools", "Sherlock", "sherlock", "sherlock.py")
        command = ["python3", sherlock_executable, "--no-color"]

        usernames = self._get_control_value(controls, 'usernames_edit', 'text')
        if not usernames or not usernames.strip():
            return None, None, "At least one username is required."
        target_for_log = usernames.strip()

        if timeout := self._get_control_value(controls, 'timeout_edit', 'text'):
            if timeout.strip():
                command.extend(["--timeout", timeout.strip()])

        if output_file := self._get_control_value(controls, 'output_edit', 'text'):
             if output_file.strip():
                command.extend(["-o", output_file.strip()])

        if csv_temp_dir:
            command.extend(["-fo", csv_temp_dir, "--csv"])
        elif self._get_control_value(controls, 'csv_check', 'check'):
            command.append("--csv")

        command.extend(target_for_log.split())

        return command, target_for_log, None

    def start_sherlock_scan(self):
        """Starts the Sherlock scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.sherlock_controls
        command, target_for_log, error = self._build_sherlock_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        sherlock_path = os.path.join(self.script_dir, "tools", "Sherlock", "sherlock", "sherlock.py")
        if not os.path.exists(sherlock_path):
            QMessageBox.critical(self, "Sherlock Error", f"'sherlock.py' not found at {sherlock_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.sherlock_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'sherlock_scan', target_for_log, self.sherlock_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_spiderfoot_tool(self):
        """Creates the UI for the Spiderfoot OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.spiderfoot_controls = self._create_spiderfoot_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.spiderfoot_controls['start_btn'])
        buttons_layout.addWidget(self.spiderfoot_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.spiderfoot_output_console = QPlainTextEdit()
        self.spiderfoot_output_console.setReadOnly(True)
        self.spiderfoot_output_console.setFont(QFont("Courier New", 10))
        self.spiderfoot_output_console.setPlaceholderText("Spiderfoot output will be displayed here...")
        main_layout.addWidget(self.spiderfoot_output_console, 1)

        self.spiderfoot_controls['start_btn'].clicked.connect(self.start_spiderfoot_scan)
        self.spiderfoot_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_spiderfoot_config_widget(self):
        """Creates a reusable, self-contained widget with Spiderfoot's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['target_edit'] = QLineEdit("example.com")
        main_layout.addRow("Target (-s):", controls['target_edit'])

        controls['types_edit'] = QLineEdit("EMAILADDR,DNS_MX,PHONE_NUMBER")
        controls['types_edit'].setToolTip("Comma-separated list of data types to collect (e.g., EMAILADDR, PHONE_NUMBER).")
        main_layout.addRow("Scan Types (-t):", controls['types_edit'])

        controls['modules_edit'] = QLineEdit()
        controls['modules_edit'].setPlaceholderText("Optional: e.g., sfp_dns,sfp_email")
        controls['modules_edit'].setToolTip("Comma-separated list of specific modules to run (-m).")
        main_layout.addRow("Modules (-m):", controls['modules_edit'])

        controls['silent_check'] = QCheckBox("Silent Output")
        controls['silent_check'].setToolTip("Only report errors (-q).")
        main_layout.addRow("-q:", controls['silent_check'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_spiderfoot_command(self, controls):
        """Builds the Spiderfoot command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("spiderfoot-cli", "spiderfoot-cli")
        if not tool_path:
             # Try python script fallback
             tool_path = self._get_tool_path("sf.py", os.path.join("spiderfoot", "sf.py"))
             if tool_path and tool_path.endswith(".py"):
                 command = [sys.executable or "python3", tool_path]
             elif tool_path:
                 command = [tool_path]
             else:
                 return None, None, "Spiderfoot not found."
        else:
             command = [tool_path]

        target = self._get_control_value(controls, 'target_edit', 'text')
        if not target or not target.strip():
            return None, None, "A target is required."
        target = target.strip()
        target_for_log = target

        command.extend(["-s", target])

        if types := self._get_control_value(controls, 'types_edit', 'text'):
            if types.strip():
                command.extend(["-t", types.strip()])
        if modules := self._get_control_value(controls, 'modules_edit', 'text'):
            if modules.strip():
                command.extend(["-m", modules.strip()])
        if self._get_control_value(controls, 'silent_check', 'check'):
            command.append("-q")

        command.append("-n") # Disable history logging for non-interactive use

        return command, target_for_log, None

    def start_spiderfoot_scan(self):
        """Starts the Spiderfoot scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.spiderfoot_controls
        command, target_for_log, error = self._build_spiderfoot_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        spiderfoot_path = os.path.join(self.script_dir, "tools", "spiderfoot", "sf.py")
        if not os.path.exists(spiderfoot_path):
            QMessageBox.critical(self, "Spiderfoot Error", f"'sf.py' not found at {spiderfoot_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.spiderfoot_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'spiderfoot_scan', target_for_log, self.spiderfoot_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def start_nmap_scan(self, sudo_password=None):
        """Wrapper to start the Nmap scan from the main tools tab."""
        if self.is_tool_running and not sudo_password:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return
        controls = self.nmap_controls
        target = controls['target_edit'].text()
        if not target:
            QMessageBox.critical(self, "Input Error", "Please provide a target for the Nmap scan.")
            return

        # Resolve hostnames in the target string before building the command
        resolved_target = self._resolve_targets_string(target)
        command, target_for_log, error = self._build_nmap_command(controls, resolved_target)
        if error:
            QMessageBox.critical(self, "Input Error", error)
            return
        if not shutil.which(command[0]):
            QMessageBox.critical(self, "Nmap Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        # Check if root is required
        requires_root = any(flag in command for flag in ['-sS', '-O', '-A', '--traceroute'])
        if requires_root and not sudo_password and sys.platform != "win32":
            if 'nmap_scan' not in self.sudo_cancel_handlers:
                self.sudo_cancel_handlers['nmap_scan'] = lambda: (
                    controls['start_btn'].setEnabled(True),
                    controls['cancel_btn'].setEnabled(False)
                )
            self._run_command_with_sudo_prompt(command, self.start_nmap_scan, 'nmap_scan')
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        if 'report_btn' in controls:
            controls['report_btn'].setEnabled(False)
        self.tool_stop_event.clear()
        self.nmap_output_console.clear()
        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'nmap_scan', target_for_log, self.nmap_output_console, sudo_password
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def generate_nmap_report(self):
        """Saves the Nmap XML and generates a styled HTML report using lxml."""
        if not self.nmap_last_xml:
            QMessageBox.information(self, "No Data", "Please run an Nmap scan first to generate data for the report.")
            return

        if not LXML_AVAILABLE:
            QMessageBox.critical(self, "Dependency Error", "The 'lxml' library is required for HTML report generation. Please install it using 'pip install lxml'.")
            return

        save_path, _ = QFileDialog.getSaveFileName(self, "Save Nmap HTML Report", "nmap_report.html", "HTML Files (*.html);;XML Files (*.xml)", options=QFileDialog.Option.DontUseNativeDialog)
        if not save_path:
            return

        try:
            # Always save the raw XML data first
            if save_path.endswith('.html'):
                xml_path = os.path.splitext(save_path)[0] + ".xml"
            else:
                xml_path = save_path

            with open(xml_path, 'w', encoding='utf-8') as f:
                f.write(self.nmap_last_xml)

            # If the user wants HTML, perform the transformation
            if save_path.endswith('.html'):
                # Use a parser that can recover from errors, which can happen with Nmap's XML
                parser = etree.XMLParser(recover=True)
                xml_doc = etree.fromstring(self.nmap_last_xml.encode('utf-8'), parser=parser)

                # Check for the stylesheet file
                xsl_path = "nmap-bootstrap.xsl"
                if not os.path.exists(xsl_path):
                    QMessageBox.critical(self, "File Not Found", f"Stylesheet '{xsl_path}' not found. Make sure it is in the same directory as the application.")
                    return

                xsl_doc = etree.parse(xsl_path)
                transform = etree.XSLT(xsl_doc)
                html_doc = transform(xml_doc)

                with open(save_path, 'wb') as f:
                    f.write(etree.tostring(html_doc, pretty_print=True))

            QMessageBox.information(self, "Report Saved", f"Report successfully saved to:\n{os.path.realpath(save_path)}")

        except Exception as e:
            logging.error(f"Failed to generate or save Nmap report: {e}", exc_info=True)
            QMessageBox.critical(self, "Report Generation Error", f"An unexpected error occurred:\n{e}")


    def _create_subfinder_tool(self):
        """Creates the UI for the Subfinder tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.subfinder_controls = self._create_subfinder_config_widget()
        layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.subfinder_controls['start_btn'])
        buttons_layout.addWidget(self.subfinder_controls['cancel_btn'])
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.subfinder_output = QPlainTextEdit()
        self.subfinder_output.setReadOnly(True)
        self.subfinder_output.setFont(QFont("Courier New", 10))
        self.subfinder_output.setPlaceholderText("Subfinder output will be displayed here...")
        layout.addWidget(self.subfinder_output, 1)

        self.subfinder_controls['start_btn'].clicked.connect(self.start_subfinder_scan)
        self.subfinder_controls['cancel_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_subfinder_config_widget(self):
        """Creates a reusable, self-contained widget for the Subfinder scanner's configuration."""
        widget = QFrame()
        widget.setObjectName("controlPanel")
        widget.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        layout = QFormLayout(widget)

        controls = {}

        controls['domain_edit'] = QLineEdit("example.com")
        controls['domain_edit'].setToolTip("Enter the target domain to enumerate subdomains for (-d).")
        layout.addRow("Domain:", controls['domain_edit'])

        controls['recursive_check'] = QCheckBox("Recursive Scan")
        controls['recursive_check'].setToolTip("Enable recursive subdomain discovery (-recursive).")
        layout.addRow(controls['recursive_check'])

        controls['all_sources_check'] = QCheckBox("Use All Sources")
        controls['all_sources_check'].setToolTip("Use all available sources for enumeration (-all).")
        layout.addRow(controls['all_sources_check'])

        output_file_layout = QHBoxLayout()
        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("Optional: path to save results...")
        output_file_layout.addWidget(controls['output_edit'])
        browse_output_btn = QPushButton("Browse...")
        browse_output_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_edit'], "Save Subfinder Results"))
        output_file_layout.addWidget(browse_output_btn)
        layout.addRow("Output File (-o):", output_file_layout)

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['cancel_btn'] = QPushButton("Cancel")
        controls['cancel_btn'].setEnabled(False)

        return widget, controls

    def start_subfinder_scan(self):
        """Starts the Subfinder scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.subfinder_controls
        original_domain = self._get_control_value(controls, 'domain_edit', 'text')
        # Subfinder is a DNS tool, so resolving before running defeats its purpose.
        # We will NOT resolve the target for this tool.
        command, target_for_log, error = self._build_subfinder_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "Subfinder Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.subfinder_output.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'subfinder_scan', target_for_log, self.subfinder_output
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_httpx_tool(self):
        """Creates the UI for the httpx tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.httpx_controls = self._create_httpx_config_widget()
        layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.httpx_controls['start_btn'])
        buttons_layout.addWidget(self.httpx_controls['cancel_btn'])
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.httpx_output = QPlainTextEdit()
        self.httpx_output.setReadOnly(True)
        self.httpx_output.setFont(QFont("Courier New", 10))
        self.httpx_output.setPlaceholderText("httpx output will be displayed here...")
        layout.addWidget(self.httpx_output, 1)

        self.httpx_controls['start_btn'].clicked.connect(self.start_httpx_scan)
        self.httpx_controls['cancel_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_httpx_config_widget(self):
        """Creates a reusable, self-contained widget for the httpx scanner's configuration."""
        widget = QFrame()
        widget.setObjectName("controlPanel")
        widget.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        main_layout = QVBoxLayout(widget)
        controls = {}

        # --- Top Row: Target List ---
        top_layout = QFormLayout()
        target_file_layout = QHBoxLayout()
        controls['target_list_edit'] = QLineEdit()
        controls['target_list_edit'].setPlaceholderText("Path to a file with hosts/URLs (one per line)...")
        controls['target_list_edit'].setToolTip("A file containing a list of targets to probe (-l).")
        target_file_layout.addWidget(controls['target_list_edit'])
        browse_target_btn = QPushButton("Browse...")
        browse_target_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['target_list_edit'], "Select Target List File"))
        target_file_layout.addWidget(browse_target_btn)
        top_layout.addRow("Target List (-l):", target_file_layout)
        main_layout.addLayout(top_layout)

        # --- Middle Row: Probes ---
        probes_box = QGroupBox("Probes to Run")
        probes_layout = QGridLayout(probes_box)
        controls['probe_status_code'] = QCheckBox("Status Code"); controls['probe_status_code'].setChecked(True)
        controls['probe_title'] = QCheckBox("Title"); controls['probe_title'].setChecked(True)
        controls['probe_tech_detect'] = QCheckBox("Tech Detect")
        controls['probe_web_server'] = QCheckBox("Web Server")
        controls['probe_cdn'] = QCheckBox("CDN")
        controls['probe_jarm'] = QCheckBox("JARM Hash")
        probes_layout.addWidget(controls['probe_status_code'], 0, 0)
        probes_layout.addWidget(controls['probe_title'], 0, 1)
        probes_layout.addWidget(controls['probe_tech_detect'], 0, 2)
        probes_layout.addWidget(controls['probe_web_server'], 1, 0)
        probes_layout.addWidget(controls['probe_cdn'], 1, 1)
        probes_layout.addWidget(controls['probe_jarm'], 1, 2)
        main_layout.addWidget(probes_box)

        # --- Bottom Row: Other Options ---
        bottom_layout = QFormLayout()
        controls['ports_edit'] = QLineEdit()
        controls['ports_edit'].setPlaceholderText("e.g., 80,443,8080 (optional)")
        controls['ports_edit'].setToolTip("Comma-separated list of ports to scan (-ports).")
        bottom_layout.addRow("Ports (-ports):", controls['ports_edit'])

        controls['json_output_check'] = QCheckBox("JSON Output (-json)")
        controls['json_output_check'].setToolTip("Output results in JSON format.")
        bottom_layout.addRow(controls['json_output_check'])

        main_layout.addLayout(bottom_layout)

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Probing")
        controls['cancel_btn'] = QPushButton("Cancel"); controls['cancel_btn'].setEnabled(False)

        return widget, controls

    def start_httpx_scan(self):
        """Starts the httpx scan worker thread."""
        # httpx resolves hostnames internally, and pre-resolving a list file
        # is complex. We will skip DoH for this tool for now.
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.httpx_controls
        command, target_for_log, error = self._build_httpx_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "httpx Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.httpx_output.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'httpx_scan', target_for_log, self.httpx_output
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_rustscan_tool(self):
        """Creates the UI for the RustScan tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.rustscan_controls = self._create_rustscan_config_widget()
        layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.rustscan_controls['start_btn'])
        buttons_layout.addWidget(self.rustscan_controls['cancel_btn'])
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.rustscan_output = QPlainTextEdit()
        self.rustscan_output.setReadOnly(True)
        self.rustscan_output.setFont(QFont("Courier New", 10))
        self.rustscan_output.setPlaceholderText("RustScan and Nmap output will be displayed here...")
        layout.addWidget(self.rustscan_output, 1)

        self.rustscan_controls['start_btn'].clicked.connect(self.start_rustscan_scan)
        self.rustscan_controls['cancel_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_rustscan_config_widget(self):
        """Creates a reusable, self-contained widget for the RustScan tool's configuration."""
        widget = QFrame()
        widget.setObjectName("controlPanel")
        widget.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        layout = QFormLayout(widget)

        controls = {}

        controls['targets_edit'] = QLineEdit("localhost")
        controls['targets_edit'].setToolTip("A single target or a comma-separated list of targets (e.g., 127.0.0.1, scanme.nmap.org).")
        layout.addRow("Targets (-a):", controls['targets_edit'])

        controls['ports_edit'] = QLineEdit("1-1000")
        controls['ports_edit'].setToolTip("A comma-separated list of ports or a range (e.g., 22,80,443 or 1-1024).")
        layout.addRow("Ports (-p):", controls['ports_edit'])

        controls['batch_size_edit'] = QLineEdit("4500")
        controls['batch_size_edit'].setToolTip("The number of ports to scan at once (-b).")
        layout.addRow("Batch Size (-b):", controls['batch_size_edit'])

        controls['timeout_edit'] = QLineEdit("1500")
        controls['timeout_edit'].setToolTip("The timeout in milliseconds for each port (-T).")
        layout.addRow("Timeout (-T):", controls['timeout_edit'])

        controls['nmap_args_edit'] = QLineEdit("-sV -sC -A")
        controls['nmap_args_edit'].setToolTip("Arguments to pass to Nmap after the port scan (e.g., -sV -A).")
        layout.addRow("Nmap Args:", controls['nmap_args_edit'])

        controls['quiet_check'] = QCheckBox("Quiet Mode (No Nmap)")
        controls['quiet_check'].setToolTip("Only output open ports, do not run Nmap (-q).")
        layout.addRow(controls['quiet_check'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['cancel_btn'] = QPushButton("Cancel"); controls['cancel_btn'].setEnabled(False)

        return widget, controls

    def start_rustscan_scan(self, sudo_password=None):
        """Starts the RustScan worker thread, prompting for sudo if necessary."""
        if self.is_tool_running and not sudo_password:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.rustscan_controls
        original_targets = self._get_control_value(controls, 'targets_edit', 'text')
        resolved_targets = self._resolve_targets_string(original_targets)
        controls['targets_edit'].setText(resolved_targets)
        command, target_for_log, error = self._build_rustscan_command(controls)
        controls['targets_edit'].setText(original_targets)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "RustScan Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        # Rustscan needs root for raw socket access.
        if not sudo_password and sys.platform != "win32":
            if 'rustscan_scan' not in self.sudo_cancel_handlers:
                self.sudo_cancel_handlers['rustscan_scan'] = lambda: (
                    controls['start_btn'].setEnabled(True),
                    controls['cancel_btn'].setEnabled(False)
                )
            self._run_command_with_sudo_prompt(command, self.start_rustscan_scan, 'rustscan_scan')
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.rustscan_output.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'rustscan_scan', target_for_log, self.rustscan_output, sudo_password
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_dirsearch_tool(self):
        """Creates the UI for the dirsearch tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.dirsearch_controls = self._create_dirsearch_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.dirsearch_controls['start_btn'])
        buttons_layout.addWidget(self.dirsearch_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.dirsearch_output_console = QPlainTextEdit()
        self.dirsearch_output_console.setReadOnly(True)
        self.dirsearch_output_console.setFont(QFont("Courier New", 10))
        self.dirsearch_output_console.setPlaceholderText("dirsearch output will be displayed here...")
        main_layout.addWidget(self.dirsearch_output_console, 1)

        self.dirsearch_controls['start_btn'].clicked.connect(self.start_dirsearch_scan)
        self.dirsearch_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_dirsearch_config_widget(self):
        """Creates a reusable, self-contained widget with all of dirsearch's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)
        controls = {}

        # --- Top-level URL and Wordlist ---
        top_frame = QFrame()
        top_layout = QFormLayout(top_frame)
        controls['url_edit'] = QLineEdit("http://localhost")
        controls['url_edit'].setToolTip("The full URL of the target application to scan (-u).")
        top_layout.addRow("Target URL (-u):", controls['url_edit'])

        wordlist_layout = QHBoxLayout()
        controls['wordlist_edit'] = QLineEdit()
        controls['wordlist_edit'].setPlaceholderText("Path to wordlist file (required)...")
        wordlist_layout.addWidget(controls['wordlist_edit'])
        browse_wordlist_btn = QPushButton("Browse...")
        browse_wordlist_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['wordlist_edit'], "Select Wordlist File"))
        wordlist_layout.addWidget(browse_wordlist_btn)
        top_layout.addRow("Wordlist (-w):", wordlist_layout)
        main_layout.addWidget(top_frame)


        # --- Tabbed Interface for Options ---
        tabs = QTabWidget()
        main_layout.addWidget(tabs)

        # --- Scan Tab ---
        scan_tab = QWidget()
        scan_layout = QFormLayout(scan_tab)
        controls['extensions_edit'] = QLineEdit("php,html,txt")
        controls['extensions_edit'].setToolTip("Comma-separated list of file extensions to append (-e).")
        scan_layout.addRow("Extensions (-e):", controls['extensions_edit'])
        controls['threads_edit'] = QLineEdit("25")
        controls['threads_edit'].setToolTip("Number of concurrent threads to use (-t).")
        scan_layout.addRow("Threads (-t):", controls['threads_edit'])
        controls['recursive_check'] = QCheckBox("Recursive Scan (-r)")
        scan_layout.addRow(controls['recursive_check'])
        tabs.addTab(scan_tab, "Scan")

        # --- Request Tab ---
        req_tab = QWidget()
        req_layout = QFormLayout(req_tab)
        controls['method_edit'] = QLineEdit("GET")
        req_layout.addRow("HTTP Method (-m):", controls['method_edit'])
        controls['headers_edit'] = QTextEdit()
        controls['headers_edit'].setPlaceholderText("e.g.,\nUser-Agent: MyAgent\nReferer: http://example.com")
        req_layout.addRow("Headers (-H):", controls['headers_edit'])
        controls['timeout_edit'] = QLineEdit("10")
        req_layout.addRow("Timeout (--timeout):", controls['timeout_edit'])
        tabs.addTab(req_tab, "Request")

        # --- Filter Tab ---
        filter_tab = QWidget()
        filter_layout = QFormLayout(filter_tab)
        controls['include_status_edit'] = QLineEdit("200-299,301,302,307,401,403")
        filter_layout.addRow("Include Status Codes (-i):", controls['include_status_edit'])
        controls['exclude_status_edit'] = QLineEdit("404,500-599")
        filter_layout.addRow("Exclude Status Codes (-x):", controls['exclude_status_edit'])
        tabs.addTab(filter_tab, "Filters")

        # --- Output Tab ---
        output_tab = QWidget()
        output_layout = QFormLayout(output_tab)
        output_file_layout = QHBoxLayout()
        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("Optional: path to save json report...")
        output_file_layout.addWidget(controls['output_edit'])
        browse_output_btn = QPushButton("Browse...")
        browse_output_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_edit'], "Save dirsearch Report"))
        output_file_layout.addWidget(browse_output_btn)
        output_layout.addRow("JSON Report (--json-report):", output_file_layout)
        tabs.addTab(output_tab, "Output")


        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_dirsearch_scan(self):
        """Starts the dirsearch scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.dirsearch_controls
        # Resolve the URL before building the command
        original_url = self._get_control_value(controls, 'url_edit', 'text')
        resolved_url = self._resolve_targets_string(original_url)
        # Temporarily update the control's value for the build function
        controls['url_edit'].setText(resolved_url)
        command, target_for_log, error = self._build_dirsearch_command(controls)
        # Restore the original value in the UI
        controls['url_edit'].setText(original_url)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "dirsearch Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        # Create a temporary file for the JSON report for parsing, regardless of user setting
        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                self.dirsearch_json_temp_file = tmp_json.name
            # Remove any user-specified report path and add our temp one
            command = [arg for arg in command if not arg.startswith('--json-report')]
            command.extend(["--json-report", self.dirsearch_json_temp_file])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary file for dirsearch report: {e}")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.dirsearch_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'dirsearch_scan', target_for_log, self.dirsearch_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_ffuf_tool(self):
        """Creates the UI for the ffuf tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.ffuf_controls = self._create_ffuf_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.ffuf_controls['start_btn'])
        buttons_layout.addWidget(self.ffuf_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.ffuf_output_console = QPlainTextEdit()
        self.ffuf_output_console.setReadOnly(True)
        self.ffuf_output_console.setFont(QFont("Courier New", 10))
        self.ffuf_output_console.setPlaceholderText("ffuf output will be displayed here...")
        main_layout.addWidget(self.ffuf_output_console, 1)

        self.ffuf_controls['start_btn'].clicked.connect(self.start_ffuf_scan)
        self.ffuf_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_ffuf_config_widget(self):
        """Creates a reusable, self-contained widget with all of ffuf's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['url_edit'] = QLineEdit("http://localhost/FUZZ")
        controls['url_edit'].setToolTip("The full URL to fuzz. Must contain the 'FUZZ' keyword.")
        main_layout.addRow("Target URL (-u):", controls['url_edit'])

        wordlist_layout = QHBoxLayout()
        controls['wordlist_edit'] = QLineEdit()
        controls['wordlist_edit'].setPlaceholderText("Path to wordlist file (required)...")
        wordlist_layout.addWidget(controls['wordlist_edit'])
        browse_wordlist_btn = QPushButton("Browse...")
        browse_wordlist_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['wordlist_edit'], "Select Wordlist File"))
        wordlist_layout.addWidget(browse_wordlist_btn)
        main_layout.addRow("Wordlist (-w):", wordlist_layout)

        controls['extensions_edit'] = QLineEdit(".php,.html,.txt")
        controls['extensions_edit'].setToolTip("Comma-separated list of file extensions to append (-e).")
        main_layout.addRow("Extensions (-e):", controls['extensions_edit'])

        controls['threads_edit'] = QLineEdit("40")
        controls['threads_edit'].setToolTip("Number of concurrent threads to use (-t).")
        main_layout.addRow("Threads (-t):", controls['threads_edit'])

        controls['method_edit'] = QLineEdit("GET")
        controls['method_edit'].setToolTip("HTTP method to use (-X).")
        main_layout.addRow("HTTP Method (-X):", controls['method_edit'])

        controls['match_codes_edit'] = QLineEdit("200,204,301,302,307,401,403,405")
        controls['match_codes_edit'].setToolTip("Match HTTP status codes (-mc).")
        main_layout.addRow("Match Codes (-mc):", controls['match_codes_edit'])

        controls['filter_codes_edit'] = QLineEdit("404")
        controls['filter_codes_edit'].setToolTip("Filter HTTP status codes (-fc).")
        main_layout.addRow("Filter Codes (-fc):", controls['filter_codes_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_ffuf_scan(self):
        """Starts the ffuf scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.ffuf_controls
        original_url = self._get_control_value(controls, 'url_edit', 'text')
        resolved_url = self._resolve_targets_string(original_url)
        controls['url_edit'].setText(resolved_url)
        command, target_for_log, error = self._build_ffuf_command(controls)
        controls['url_edit'].setText(original_url)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "ffuf Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        # Always use a temporary file for JSON output for parsing
        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                self.ffuf_json_temp_file = tmp_json.name
            command.extend(["-o", self.ffuf_json_temp_file, "-of", "json"])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary file for ffuf report: {e}")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.ffuf_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'ffuf_scan', target_for_log, self.ffuf_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_enum4linux_ng_tool(self):
        """Creates the UI for the enum4linux-ng tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.enum4linux_ng_controls = self._create_enum4linux_ng_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.enum4linux_ng_controls['start_btn'])
        buttons_layout.addWidget(self.enum4linux_ng_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.enum4linux_ng_output_console = QPlainTextEdit()
        self.enum4linux_ng_output_console.setReadOnly(True)
        self.enum4linux_ng_output_console.setFont(QFont("Courier New", 10))
        self.enum4linux_ng_output_console.setPlaceholderText("enum4linux-ng output will be displayed here...")
        main_layout.addWidget(self.enum4linux_ng_output_console, 1)

        self.enum4linux_ng_controls['start_btn'].clicked.connect(self.start_enum4linux_ng_scan)
        self.enum4linux_ng_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_enum4linux_ng_config_widget(self):
        """Creates a reusable, self-contained widget with enum4linux-ng's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        controls = {}

        # --- Target and Auth ---
        top_box = QGroupBox("Target & Authentication")
        top_layout = QFormLayout(top_box)
        controls['target_edit'] = QLineEdit()
        top_layout.addRow("Target Host:", controls['target_edit'])
        controls['user_edit'] = QLineEdit()
        top_layout.addRow("Username (-u):", controls['user_edit'])
        controls['pass_edit'] = QLineEdit(); controls['pass_edit'].setEchoMode(QLineEdit.EchoMode.Password)
        top_layout.addRow("Password (-p):", controls['pass_edit'])
        main_layout.addWidget(top_box)

        # --- Enumeration Options ---
        enum_box = QGroupBox("Enumeration Options")
        enum_layout = QGridLayout(enum_box)
        controls['all_check'] = QCheckBox("All Simple Enum (-A)"); controls['all_check'].setChecked(True)
        controls['users_check'] = QCheckBox("Users (-U)")
        controls['groups_check'] = QCheckBox("Groups (-G)")
        controls['shares_check'] = QCheckBox("Shares (-S)")
        controls['policy_check'] = QCheckBox("Password Policy (-P)")
        controls['os_check'] = QCheckBox("OS Info (-O)")
        enum_layout.addWidget(controls['all_check'], 0, 0)
        enum_layout.addWidget(controls['users_check'], 1, 0)
        enum_layout.addWidget(controls['groups_check'], 1, 1)
        enum_layout.addWidget(controls['shares_check'], 2, 0)
        enum_layout.addWidget(controls['policy_check'], 2, 1)
        enum_layout.addWidget(controls['os_check'], 3, 0)
        main_layout.addWidget(enum_box)

        # --- UI Logic ---
        def toggle_enum_options(checked):
            for key in ['users_check', 'groups_check', 'shares_check', 'policy_check', 'os_check']:
                controls[key].setDisabled(checked)
        controls['all_check'].toggled.connect(toggle_enum_options)
        toggle_enum_options(True)

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_enum4linux_ng_scan(self):
        """Starts the enum4linux-ng scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.enum4linux_ng_controls
        original_target = self._get_control_value(controls, 'target_edit', 'text')
        resolved_target = self._resolve_targets_string(original_target)
        controls['target_edit'].setText(resolved_target)
        command, target_for_log, error = self._build_enum4linux_ng_command(controls)
        controls['target_edit'].setText(original_target)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "enum4linux-ng Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                self.enum4linux_ng_json_temp_file = tmp_json.name
            command.insert(1, "-oJ")
            command.insert(2, self.enum4linux_ng_json_temp_file)
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary file for enum4linux-ng report: {e}")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.enum4linux_ng_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'enum4linux_ng_scan', target_for_log, self.enum4linux_ng_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_dnsrecon_tool(self):
        """Creates the UI for the dnsrecon tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.dnsrecon_controls = self._create_dnsrecon_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.dnsrecon_controls['start_btn'])
        buttons_layout.addWidget(self.dnsrecon_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.dnsrecon_output_console = QPlainTextEdit()
        self.dnsrecon_output_console.setReadOnly(True)
        self.dnsrecon_output_console.setFont(QFont("Courier New", 10))
        self.dnsrecon_output_console.setPlaceholderText("dnsrecon output will be displayed here...")
        main_layout.addWidget(self.dnsrecon_output_console, 1)

        self.dnsrecon_controls['start_btn'].clicked.connect(self.start_dnsrecon_scan)
        self.dnsrecon_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_dnsrecon_config_widget(self):
        """Creates a reusable, self-contained widget with dnsrecon's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['domain_edit'] = QLineEdit("example.com")
        main_layout.addRow("Domain (-d):", controls['domain_edit'])

        controls['scan_type_combo'] = QComboBox()
        controls['scan_type_combo'].addItems(["std", "axfr", "brt", "srv", "zonewalk"])
        controls['scan_type_combo'].setToolTip("Select the enumeration type (-t).")
        main_layout.addRow("Scan Type (-t):", controls['scan_type_combo'])

        wordlist_layout = QHBoxLayout()
        controls['dict_edit'] = QLineEdit()
        controls['dict_edit'].setPlaceholderText("Path to wordlist for 'brt' scan...")
        wordlist_layout.addWidget(controls['dict_edit'])
        browse_dict_btn = QPushButton("Browse...")
        browse_dict_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['dict_edit'], "Select Dictionary File"))
        wordlist_layout.addWidget(browse_dict_btn)
        main_layout.addRow("Dictionary (-D):", wordlist_layout)

        controls['json_output_edit'] = QLineEdit()
        controls['json_output_edit'].setPlaceholderText("Optional: path to save JSON report...")
        main_layout.addRow("JSON Output (--json):", controls['json_output_edit'])

        # UI Logic
        def toggle_dict_visibility(text):
            controls['dict_edit'].setVisible(text == 'brt')
            browse_dict_btn.setVisible(text == 'brt')
            # Also find and hide the label
            label = main_layout.labelForField(wordlist_layout)
            if label: label.setVisible(text == 'brt')

        controls['scan_type_combo'].currentTextChanged.connect(toggle_dict_visibility)
        toggle_dict_visibility(controls['scan_type_combo'].currentText())

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_dnsrecon_scan(self):
        """Starts the dnsrecon scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.dnsrecon_controls
        command, target_for_log, error = self._build_dnsrecon_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "dnsrecon Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        # Always use a temporary file for JSON output for parsing
        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                self.dnsrecon_json_temp_file = tmp_json.name
            # Remove any user-specified report path and add our temp one
            command = [arg for arg in command if not arg.startswith('--json')]
            command.extend(["--json", self.dnsrecon_json_temp_file])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary file for dnsrecon report: {e}")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.dnsrecon_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'dnsrecon_scan', target_for_log, self.dnsrecon_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_fierce_tool(self):
        """Creates the UI for the fierce DNS scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.fierce_controls = self._create_fierce_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.fierce_controls['start_btn'])
        buttons_layout.addWidget(self.fierce_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.fierce_output_console = QPlainTextEdit()
        self.fierce_output_console.setReadOnly(True)
        self.fierce_output_console.setFont(QFont("Courier New", 10))
        self.fierce_output_console.setPlaceholderText("fierce output will be displayed here...")
        main_layout.addWidget(self.fierce_output_console, 1)

        self.fierce_controls['start_btn'].clicked.connect(self.start_fierce_scan)
        self.fierce_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_fierce_config_widget(self):
        """Creates a reusable, self-contained widget with fierce's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['domain_edit'] = QLineEdit("example.com")
        main_layout.addRow("Domain (--domain):", controls['domain_edit'])

        controls['connect_check'] = QCheckBox("Attempt HTTP Connections")
        controls['connect_check'].setToolTip("Attempt to connect to any non-RFC1918 hosts found via HTTP.")
        main_layout.addRow("--connect", controls['connect_check'])

        controls['wide_check'] = QCheckBox("Wide Scan")
        controls['wide_check'].setToolTip("Scan the entire Class C of any discovered records.")
        main_layout.addRow("--wide", controls['wide_check'])

        controls['traverse_edit'] = QLineEdit("5")
        controls['traverse_edit'].setToolTip("Scan a number of IPs above and below discovered hosts.")
        main_layout.addRow("--traverse", controls['traverse_edit'])

        controls['delay_edit'] = QLineEdit("1")
        controls['delay_edit'].setToolTip("Delay in seconds between lookups.")
        main_layout.addRow("--delay", controls['delay_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_fierce_scan(self):
        """Starts the fierce scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.fierce_controls
        command, target_for_log, error = self._build_fierce_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "fierce Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.fierce_output_console.clear()

        self.worker = WorkerThread(self._fierce_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _fierce_thread(self, command, domain):
        """Worker thread for running the fierce command."""
        q = self.tool_results_queue
        logging.info(f"Starting fierce with command: {' '.join(command)}")
        q.put(('fierce_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.fierce_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('fierce_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('fierce_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'fierce Error', "'fierce' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"fierce thread error: {e}", exc_info=True)
            q.put(('error', 'fierce Error', str(e)))
        finally:
            q.put(('tool_finished', 'fierce_scan', domain, "".join(full_output)))
            with self.thread_finish_lock:
                self.fierce_process = None
            logging.info("fierce scan thread finished.")

    def _handle_fierce_output(self, line):
        self.fierce_output_console.insertPlainText(line)
        self.fierce_output_console.verticalScrollBar().setValue(self.fierce_output_console.verticalScrollBar().maximum())



    def _create_ssh_manager_tab(self):
        """Creates the UI for the SSH Manager tool by instantiating the SshManagerTab widget."""
        # SshManagerTab is a QWidget that contains all its own logic.
        return SshManagerTab(self.icon_path)

    def _create_speed_test_tab(self):
        """Creates the UI for the Speed Test tool."""
        # SpeedTestTab is a QWidget that contains all its own logic.
        return SpeedTestTab(self)

    def _create_tools_tab(self,p=None):
        """Creates the tab container for the standard network tools."""
        tools_tabs = QTabWidget()
        tools_tabs.addTab(self._create_nmap_scanner_tool(), "Nmap Scan")
        tools_tabs.addTab(self._create_subdomain_scanner_tool(), "Subdomain Scanner (Sublist3r)")
        tools_tabs.addTab(self._create_subfinder_tool(), "Subdomain Scanner (Subfinder)")
        tools_tabs.addTab(self._create_httpx_tool(), "httpx Probe")
        tools_tabs.addTab(self._create_rustscan_tool(), "RustScan")
        tools_tabs.addTab(self._create_dirsearch_tool(), "dirsearch")
        tools_tabs.addTab(self._create_ffuf_tool(), "ffuf")
        tools_tabs.addTab(self._create_enum4linux_ng_tool(), "enum4linux-ng")
        tools_tabs.addTab(self._create_dnsrecon_tool(), "dnsrecon")
        tools_tabs.addTab(self._create_fierce_tool(), "fierce")
        tools_tabs.addTab(self._create_nikto_scanner_tool(), "Nikto Scan")
        tools_tabs.addTab(self._create_gobuster_tool(), "Gobuster")
        tools_tabs.addTab(self._create_whatweb_tool(), "WhatWeb")
        tools_tabs.addTab(self._create_masscan_tool(), "Masscan")
        tools_tabs.addTab(self._create_port_scanner_tool(), "Port Scanner (Scapy)")
        tools_tabs.addTab(self._create_arp_scan_tool(), "ARP Scan (Scapy)")
        tools_tabs.addTab(self._create_arp_scan_cli_tool(), "ARP Scan (CLI)")
        tools_tabs.addTab(self._create_speed_test_tab(), "Speed Test")
        tools_tabs.addTab(self._create_ssh_manager_tab(), "SSH Manager")
        tools_tabs.addTab(self._create_ping_sweep_tool(), "Ping Sweep")
        tools_tabs.addTab(self._create_traceroute_tool(), "Traceroute")
        return tools_tabs

    def _create_subdomain_scanner_tool(self):
        """Creates the UI for the Sublist3r Subdomain Scanner tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.subdomain_controls = self._create_subdomain_scanner_config_widget()
        layout.addWidget(config_widget)

        # Connect signals
        self.subdomain_controls['start_btn'].clicked.connect(self.start_sublist3r_scan)
        self.subdomain_controls['cancel_btn'].clicked.connect(self.cancel_tool)

        # --- Output Console ---
        self.sublist3r_output = QPlainTextEdit()
        self.sublist3r_output.setReadOnly(True)
        self.sublist3r_output.setFont(QFont("Courier New", 10))
        self.sublist3r_output.setPlaceholderText("Sublist3r output will be displayed here...")
        layout.addWidget(self.sublist3r_output, 1)

        return widget

    def _create_subdomain_scanner_config_widget(self):
        """Creates a reusable, self-contained widget for the Subdomain scanner's configuration."""
        widget = QFrame()
        widget.setObjectName("controlPanel")
        widget.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        controls_layout = QHBoxLayout(widget)

        controls = {}

        controls_layout.addWidget(QLabel("Domain:"))
        controls['domain_edit'] = QLineEdit("example.com")
        controls['domain_edit'].setToolTip("Enter the target domain to enumerate subdomains for (e.g., example.com).")
        controls_layout.addWidget(controls['domain_edit'], 1) # Add stretch

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['cancel_btn'] = QPushButton("Cancel")
        controls['cancel_btn'].setEnabled(False)

        controls_layout.addWidget(controls['start_btn'])
        controls_layout.addWidget(controls['cancel_btn'])

        return widget, controls

    def start_sublist3r_scan(self):
        """Starts the Sublist3r scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.subdomain_controls
        command, target_for_log, error = self._build_sublist3r_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        sublist3r_path = os.path.join("tools", "sublist3r", "sublist3r.py")
        if not os.path.exists(sublist3r_path):
            QMessageBox.critical(self, "Tool Error", f"'sublist3r.py' not found at {sublist3r_path}")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.sublist3r_output.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'sublist3r_scan', target_for_log, self.sublist3r_output
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_nikto_scanner_tool(self):
        """Creates the UI for the Nikto Web Scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.nikto_controls = self._create_nikto_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.nikto_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.nikto_output_console = QPlainTextEdit()
        self.nikto_output_console.setReadOnly(True)
        self.nikto_output_console.setFont(QFont("Courier New", 10))
        self.nikto_output_console.setPlaceholderText("Nikto output will be displayed here...")
        main_layout.addWidget(self.nikto_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_nikto_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_nikto_config_widget(self):
        """Creates a reusable, self-contained widget with all of Nikto's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)

        controls = {}

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Nikto Web Scanner</b></font>
        <p>This tool runs the Nikto web server scanner. Use the tabs below to configure the scan.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        nikto_tabs = QTabWidget()
        main_layout.addWidget(nikto_tabs)

        target_tab = QWidget()
        target_layout = QFormLayout(target_tab)
        controls['target_edit'] = QLineEdit("localhost"); controls['target_edit'].setToolTip("Enter the target host, IP, or CIDR range to scan.")
        target_layout.addRow("Target Host:", controls['target_edit'])
        controls['port_edit'] = QLineEdit("80"); controls['port_edit'].setToolTip("Specify the port(s) to scan. Can be a single port, a comma-separated list (80,443), or a range (80-90).")
        target_layout.addRow("Target Port:", controls['port_edit'])
        controls['ssl_check'] = QCheckBox("Force SSL Mode"); controls['ssl_check'].setToolTip("Force SSL mode on the target port(s) (-ssl). Nikto automatically uses SSL on port 443, but this forces it for other non-standard SSL ports.")
        target_layout.addRow(controls['ssl_check'])
        controls['vhost_edit'] = QLineEdit(); controls['vhost_edit'].setToolTip("Specify the virtual host to use in the HTTP Host header. Useful for testing multiple sites on one IP.")
        target_layout.addRow("Virtual Host:", controls['vhost_edit'])
        nikto_tabs.addTab(target_tab, "Target")

        scan_tab = QWidget()
        scan_layout = QFormLayout(scan_tab)
        controls['tuning_combo'] = QComboBox(); controls['tuning_combo'].setToolTip("Select a scan tuning profile (-Tuning) to focus on specific types of tests.\n'x' can be used to reverse the logic (e.g., -Tuning x12 will exclude 'Interesting File' and 'Misconfiguration').")
        controls['tuning_combo'].addItems(["Default", "0 - File Upload", "1 - Interesting File", "2 - Misconfiguration", "3 - Information Disclosure", "4 - Injection (XSS/Script/HTML)", "5 - Remote File Retrieval", "6 - Denial of Service", "8 - Command Execution", "9 - SQL Injection", "a - Auth Bypass", "b - Software ID", "c - Remote Source Inclusion", "x - Reverse Tuning"])
        scan_layout.addRow("Tuning Profile:", controls['tuning_combo'])
        controls['mutate_combo'] = QComboBox(); controls['mutate_combo'].setToolTip("Perform mutation tests (-mutate) to guess additional file and directory names based on known ones found during the scan.")
        controls['mutate_combo'].addItems(["None", "1 - Test files with root dirs", "2 - Guess password files", "3 - Enumerate users via Apache", "4 - Enumerate users via cgiwrap", "5 - Brute force sub-domains", "6 - Guess directory names"])
        scan_layout.addRow("Mutate:", controls['mutate_combo'])
        controls['plugins_edit'] = QLineEdit(); controls['plugins_edit'].setToolTip("Select specific plugins to run, separated by commas (e.g., apache_users,cgi).\nUse 'list' to see all available plugins in the console output.")
        scan_layout.addRow("Plugins:", controls['plugins_edit'])
        controls['cgidirs_edit'] = QLineEdit(); controls['cgidirs_edit'].setToolTip("Scan these CGI directories. Common values are 'all', 'none', or a specific path like '/cgi-bin/'.")
        scan_layout.addRow("CGI Dirs:", controls['cgidirs_edit'])
        nikto_tabs.addTab(scan_tab, "Scan")

        evasion_tab = QWidget()
        evasion_layout = QFormLayout(evasion_tab)
        controls['evasion_combo'] = QComboBox(); controls['evasion_combo'].setToolTip("Select an IDS evasion technique (-evasion).\nThese techniques attempt to bypass Intrusion Detection Systems by encoding or formatting requests in non-standard ways. Use with caution.")
        controls['evasion_combo'].addItems(["None", "1 - Random URI encoding", "2 - Directory self-reference (/./)", "3 - Premature URL ending", "4 - Prepend long random string", "5 - Fake parameter", "6 - TAB as request spacer", "7 - Change case of URL", "8 - Use Windows directory separator (\\)", "A - Use carriage return (0x0d)", "B - Use binary value 0x0b"])
        evasion_layout.addRow("Evasion Technique:", controls['evasion_combo'])
        nikto_tabs.addTab(evasion_tab, "Evasion")

        config_tab = QWidget()
        config_layout = QFormLayout(config_tab)
        controls['timeout_edit'] = QLineEdit("10"); controls['timeout_edit'].setToolTip("Set the timeout in seconds for each individual HTTP request (default is 10).")
        config_layout.addRow("Timeout (s):", controls['timeout_edit'])
        controls['maxtime_edit'] = QLineEdit(); controls['maxtime_edit'].setToolTip("Set the maximum total testing time for the entire scan per host (e.g., 1h, 60m, 3600s).")
        config_layout.addRow("Max Time:", controls['maxtime_edit'])
        controls['pause_edit'] = QLineEdit(); controls['pause_edit'].setToolTip("Pause in seconds between each test (HTTP request). Useful for reducing scan speed.")
        config_layout.addRow("Pause (s):", controls['pause_edit'])
        controls['id_edit'] = QLineEdit(); controls['id_edit'].setToolTip("Provide HTTP Basic authentication credentials in the format 'id:password' or 'id:password:realm'.")
        config_layout.addRow("Auth (id:pass):", controls['id_edit'])
        controls['root_edit'] = QLineEdit(); controls['root_edit'].setToolTip("Prepend a value to the beginning of every request URI. Useful if the web application is in a subdirectory (e.g., /app).")
        config_layout.addRow("Root Directory:", controls['root_edit'])
        controls['proxy_check'] = QCheckBox("Use proxy from nikto.conf")
        config_layout.addRow(controls['proxy_check'])
        nikto_tabs.addTab(config_tab, "Config")

        output_tab = QWidget()
        output_layout = QFormLayout(output_tab)
        output_file_layout = QHBoxLayout()
        controls['output_file_edit'] = QLineEdit(); controls['output_file_edit'].setPlaceholderText("Optional: path to save report...")
        output_file_layout.addWidget(controls['output_file_edit'])
        browse_out_btn = QPushButton("Browse..."); browse_out_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_file_edit'], "Save Nikto Report", "HTML Files (*.html);;CSV Files (*.csv);;Text Files (*.txt);;XML Files (*.xml);;All Files (*)"))
        output_file_layout.addWidget(browse_out_btn)
        output_layout.addRow("Output File:", output_file_layout)
        controls['format_combo'] = QComboBox(); controls['format_combo'].addItems(["html", "csv", "txt", "xml", "nbe"]); controls['format_combo'].setToolTip("Select the report format (requires an output file to be set).")
        output_layout.addRow("Report Format:", controls['format_combo'])
        save_dir_layout = QHBoxLayout()
        controls['save_dir_edit'] = QLineEdit(); controls['save_dir_edit'].setPlaceholderText("Optional: directory to save positive responses...")
        save_dir_layout.addWidget(controls['save_dir_edit'])
        browse_save_btn = QPushButton("Browse..."); browse_save_btn.clicked.connect(lambda: self._browse_dir_for_lineedit(controls['save_dir_edit'], "Select Directory to Save Responses"))
        save_dir_layout.addWidget(browse_save_btn)
        output_layout.addRow("Save Directory:", save_dir_layout)
        nikto_tabs.addTab(output_tab, "Output")

        controls['extra_opts_edit'] = QLineEdit(); controls['extra_opts_edit'].setToolTip("Enter any additional, space-separated Nikto flags here. These will be appended directly to the command.")
        main_layout.addWidget(QLabel("Additional Raw Options:"))
        main_layout.addWidget(controls['extra_opts_edit'])

        # Add buttons to controls dict for external connection
        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Nikto Scan")
        controls['stop_btn'] = QPushButton("Stop Nikto"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_nikto_scan(self):
        """Starts the Nikto scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.nikto_controls
        original_target = self._get_control_value(controls, 'target_edit', 'text')
        resolved_target = self._resolve_targets_string(original_target)
        controls['target_edit'].setText(resolved_target)
        command, target_for_log, error = self._build_nikto_command(controls)
        controls['target_edit'].setText(original_target)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "Nikto Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.nikto_output_console.clear()

        self.worker = WorkerThread(self._nikto_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_gobuster_tool(self):
        """Creates the UI for the Gobuster tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.gobuster_controls = self._create_gobuster_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.gobuster_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.gobuster_output_console = QPlainTextEdit()
        self.gobuster_output_console.setReadOnly(True)
        self.gobuster_output_console.setFont(QFont("Courier New", 10))
        self.gobuster_output_console.setPlaceholderText("Gobuster output will be displayed here...")
        main_layout.addWidget(self.gobuster_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_gobuster_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_gobuster_config_widget(self):
        """Creates a reusable, self-contained widget with all of Gobuster's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)

        controls = {}

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Gobuster Directory/File Brute-forcer</b></font>
        <p>This tool uses Gobuster to discover hidden directories and files on web servers. Configure the scan using the tabs below.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        gobuster_tabs = QTabWidget()
        main_layout.addWidget(gobuster_tabs)

        scan_tab = QWidget()
        scan_layout = QFormLayout(scan_tab)
        controls['url_edit'] = QLineEdit("http://localhost"); controls['url_edit'].setToolTip("The full URL of the target application to scan, including the scheme (http/https).")
        scan_layout.addRow("Target URL:", controls['url_edit'])

        wordlist_layout = QHBoxLayout()
        controls['wordlist_edit'] = QLineEdit(); controls['wordlist_edit'].setPlaceholderText("Path to wordlist file (required)..."); controls['wordlist_edit'].setToolTip("The path to the wordlist file to use for brute-forcing directories and files.")
        wordlist_layout.addWidget(controls['wordlist_edit'])
        browse_wordlist_btn = QPushButton("Browse..."); browse_wordlist_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['wordlist_edit'], "Select Wordlist File"))
        wordlist_layout.addWidget(browse_wordlist_btn)
        scan_layout.addRow("Wordlist:", wordlist_layout)

        controls['threads_edit'] = QLineEdit("10"); controls['threads_edit'].setToolTip("Number of concurrent threads to use for the scan. Higher numbers are faster but can overwhelm the target.")
        scan_layout.addRow("Threads:", controls['threads_edit'])
        controls['extensions_edit'] = QLineEdit("php,html,txt"); controls['extensions_edit'].setToolTip("A comma-separated list of file extensions to append to each word in the wordlist (e.g., php,html,txt) (-x).")
        scan_layout.addRow("Extensions (-x):", controls['extensions_edit'])
        controls['status_codes_edit'] = QLineEdit("200,204,301,302,307"); controls['status_codes_edit'].setToolTip("A comma-separated list of HTTP status codes to treat as valid and display in the results (e.g., 200,204,301) (-s).")
        scan_layout.addRow("Status Codes (-s):", controls['status_codes_edit'])
        controls['status_codes_blacklist_edit'] = QLineEdit("404"); controls['status_codes_blacklist_edit'].setToolTip("A comma-separated list of status codes to hide from the output (e.g., 403,404) (-b). This takes precedence over the positive status code list.")
        scan_layout.addRow("Blacklist Status Codes (-b):", controls['status_codes_blacklist_edit'])

        checkbox_layout = QHBoxLayout()
        controls['add_slash_check'] = QCheckBox("Add Slash"); controls['add_slash_check'].setToolTip("Append a forward slash to each directory request (-f). Useful for specifically finding directories.")
        checkbox_layout.addWidget(controls['add_slash_check'])
        controls['follow_redirect_check'] = QCheckBox("Follow Redirect"); controls['follow_redirect_check'].setToolTip("Follow HTTP redirects to their final destination (-r).")
        checkbox_layout.addWidget(controls['follow_redirect_check'])
        scan_layout.addRow(checkbox_layout)
        gobuster_tabs.addTab(scan_tab, "Scan Options")

        request_tab = QWidget()
        request_layout = QFormLayout(request_tab)
        controls['useragent_edit'] = QLineEdit(); controls['useragent_edit'].setToolTip("Set a custom User-Agent string for all requests. Can be used to impersonate different browsers.")
        request_layout.addRow("User-Agent:", controls['useragent_edit'])
        controls['random_agent_check'] = QCheckBox("Use Random User-Agent")
        request_layout.addRow(controls['random_agent_check'])
        controls['cookies_edit'] = QLineEdit(); controls['cookies_edit'].setToolTip("Set cookies for the request. The format is 'name=value; name2=value2'.")
        request_layout.addRow("Cookies:", controls['cookies_edit'])
        controls['proxy_edit'] = QLineEdit(); controls['proxy_edit'].setToolTip("Proxy server to use for requests (e.g., http://127.0.0.1:8080, socks5://127.0.0.1:9050).")
        request_layout.addRow("Proxy:", controls['proxy_edit'])
        controls['timeout_edit'] = QLineEdit("10s"); controls['timeout_edit'].setToolTip("Timeout for each individual HTTP request (e.g., 10s, 1m, 500ms).")
        request_layout.addRow("Timeout:", controls['timeout_edit'])
        controls['username_edit'] = QLineEdit()
        request_layout.addRow("Username (Basic Auth):", controls['username_edit'])
        controls['password_edit'] = QLineEdit(); controls['password_edit'].setEchoMode(QLineEdit.EchoMode.Password)
        request_layout.addRow("Password (Basic Auth):", controls['password_edit'])
        gobuster_tabs.addTab(request_tab, "Request")

        output_tab = QWidget()
        output_layout = QFormLayout(output_tab)
        output_file_layout = QHBoxLayout()
        controls['output_file_edit'] = QLineEdit(); controls['output_file_edit'].setPlaceholderText("Optional: path to save output file...")
        output_file_layout.addWidget(controls['output_file_edit'])
        browse_out_btn = QPushButton("Browse..."); browse_out_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_file_edit'], "Save Gobuster Output"))
        output_file_layout.addWidget(browse_out_btn)
        output_layout.addRow("Output File:", output_file_layout)

        output_checkbox_layout = QHBoxLayout()
        controls['no_progress_check'] = QCheckBox("No Progress"); controls['no_progress_check'].setToolTip("Don't display the real-time progress bar during the scan (-z).")
        output_checkbox_layout.addWidget(controls['no_progress_check'])
        controls['quiet_check'] = QCheckBox("Quiet"); controls['quiet_check'].setToolTip("Don't print the startup banner and other non-result information (-q).")
        output_checkbox_layout.addWidget(controls['quiet_check'])
        controls['expanded_check'] = QCheckBox("Expanded View"); controls['expanded_check'].setToolTip("Print the full URL for each result, not just the relative path (-e).")
        output_checkbox_layout.addWidget(controls['expanded_check'])
        output_layout.addRow(output_checkbox_layout)
        gobuster_tabs.addTab(output_tab, "Output")

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Gobuster Scan")
        controls['stop_btn'] = QPushButton("Stop Gobuster"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _create_whatweb_config_widget(self):
        """Creates a reusable, self-contained widget with WhatWeb's configuration options."""
        widget = QGroupBox("Scan Options")
        # In a real app, you might want a QWidget as the base and put the QGroupBox inside it
        # but for this structure, returning the groupbox is fine.
        controls_layout = QFormLayout(widget)
        controls = {}

        controls['target_edit'] = QLineEdit("http://localhost")
        controls['target_edit'].setToolTip("Enter one or more targets to scan, separated by spaces.\nCan be URLs, hostnames, or IP ranges.")
        controls_layout.addRow("Target(s):", controls['target_edit'])

        controls['aggression_combo'] = QComboBox()
        controls['aggression_combo'].addItems(["1 - Stealthy", "3 - Aggressive", "4 - Heavy"])
        controls['aggression_combo'].setToolTip("Set the aggression level (-a).\n- 1 (Stealthy): Light and fast, makes few requests.\n- 3 (Aggressive): Makes more requests, may trigger alerts.\n- 4 (Heavy): Very noisy, runs every single plugin.")
        controls_layout.addRow("Aggression Level (-a):", controls['aggression_combo'])

        controls['verbose_check'] = QCheckBox("Enable Verbose Output (-v)")
        controls['verbose_check'].setToolTip("Enable verbose output (-v).\nShows more detail during the scan, including which plugins are running.")
        controls_layout.addRow(controls['verbose_check'])

        controls['extra_opts_edit'] = QLineEdit()
        controls['extra_opts_edit'].setToolTip("Enter any additional, space-separated WhatWeb flags here. These will be appended directly to the command.")
        controls_layout.addRow("Additional Raw Options:", controls['extra_opts_edit'])

        # These are needed for the main tab to function correctly, but might be ignored by the LAB
        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start WhatWeb Scan")
        controls['stop_btn'] = QPushButton("Stop WhatWeb"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _create_whatweb_tool(self):
        """Creates the UI for the WhatWeb tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>WhatWeb - Web Scanner</b></font>
        <p>This tool identifies technologies used on websites, including content management systems (CMS), blogging platforms, visitor statistics/analytics packages, JavaScript libraries, web servers, and embedded devices.</p>
        """)
        instructions.setFixedHeight(100)
        main_layout.addWidget(instructions)

        # --- Controls ---
        controls_frame, self.whatweb_controls = self._create_whatweb_config_widget()
        main_layout.addWidget(controls_frame)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.whatweb_controls['start_btn'])
        buttons_layout.addWidget(self.whatweb_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.whatweb_output_console = QPlainTextEdit()
        self.whatweb_output_console.setReadOnly(True)
        self.whatweb_output_console.setFont(QFont("Courier New", 10))
        self.whatweb_output_console.setPlaceholderText("WhatWeb output will be displayed here...")
        main_layout.addWidget(self.whatweb_output_console, 1)

        self.whatweb_controls['start_btn'].clicked.connect(self.start_whatweb_scan)
        self.whatweb_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def start_whatweb_scan(self):
        """Starts the WhatWeb scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.whatweb_controls
        command, target_for_log, error = self._build_whatweb_command(controls)

        if error:
            QMessageBox.critical(self, "WhatWeb Error", error)
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.whatweb_output_console.clear()

        self.worker = WorkerThread(self._whatweb_thread, args=(command,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def start_gobuster_scan(self):
        """Starts the Gobuster scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.gobuster_controls
        original_url = self._get_control_value(controls, 'url_edit', 'text')
        resolved_url = self._resolve_targets_string(original_url)
        controls['url_edit'].setText(resolved_url)
        command, target_for_log, error = self._build_gobuster_command(controls)
        controls['url_edit'].setText(original_url)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "Gobuster Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.gobuster_output_console.clear()

        self.worker = WorkerThread(self._gobuster_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_whatweb_config_widget(self):
        """Creates a reusable, self-contained widget with WhatWeb's configuration options."""
        widget = QGroupBox("Scan Options")
        # In a real app, you might want a QWidget as the base and put the QGroupBox inside it
        # but for this structure, returning the groupbox is fine.
        controls_layout = QFormLayout(widget)
        controls = {}

        controls['target_edit'] = QLineEdit("http://localhost")
        controls['target_edit'].setToolTip("Enter one or more targets to scan, separated by spaces.\nCan be URLs, hostnames, or IP ranges.")
        controls_layout.addRow("Target(s):", controls['target_edit'])

        controls['aggression_combo'] = QComboBox()
        controls['aggression_combo'].addItems(["1 - Stealthy", "3 - Aggressive", "4 - Heavy"])
        controls['aggression_combo'].setToolTip("Set the aggression level (-a).\n- 1 (Stealthy): Light and fast, makes few requests.\n- 3 (Aggressive): Makes more requests, may trigger alerts.\n- 4 (Heavy): Very noisy, runs every single plugin.")
        controls_layout.addRow("Aggression Level (-a):", controls['aggression_combo'])

        controls['verbose_check'] = QCheckBox("Enable Verbose Output (-v)")
        controls['verbose_check'].setToolTip("Enable verbose output (-v).\nShows more detail during the scan, including which plugins are running.")
        controls_layout.addRow(controls['verbose_check'])

        controls['extra_opts_edit'] = QLineEdit()
        controls['extra_opts_edit'].setToolTip("Enter any additional, space-separated WhatWeb flags here. These will be appended directly to the command.")
        controls_layout.addRow("Additional Raw Options:", controls['extra_opts_edit'])

        # These are needed for the main tab to function correctly, but might be ignored by the LAB
        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start WhatWeb Scan")
        controls['stop_btn'] = QPushButton("Stop WhatWeb"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _create_whatweb_tool(self):
        """Creates the UI for the WhatWeb tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>WhatWeb - Web Scanner</b></font>
        <p>This tool identifies technologies used on websites, including content management systems (CMS), blogging platforms, visitor statistics/analytics packages, JavaScript libraries, web servers, and embedded devices.</p>
        """)
        instructions.setFixedHeight(100)
        main_layout.addWidget(instructions)

        # --- Controls ---
        controls_frame, self.whatweb_controls = self._create_whatweb_config_widget()
        main_layout.addWidget(controls_frame)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.whatweb_controls['start_btn'])
        buttons_layout.addWidget(self.whatweb_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.whatweb_output_console = QPlainTextEdit()
        self.whatweb_output_console.setReadOnly(True)
        self.whatweb_output_console.setFont(QFont("Courier New", 10))
        self.whatweb_output_console.setPlaceholderText("WhatWeb output will be displayed here...")
        main_layout.addWidget(self.whatweb_output_console, 1)

        self.whatweb_controls['start_btn'].clicked.connect(self.start_whatweb_scan)
        self.whatweb_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def start_whatweb_scan(self):
        """Starts the WhatWeb scan worker thread."""
        controls = self.whatweb_controls
        if not shutil.which("whatweb"):
            QMessageBox.critical(self, "WhatWeb Error", "'whatweb' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = controls['target_edit'].text().strip()
        if not target:
            QMessageBox.critical(self, "Input Error", "A target is required.")
            return

        resolved_target = self._resolve_targets_string(target)

        command = ["whatweb"]

        # Aggression
        aggression_level = controls['aggression_combo'].currentText().split(" ")[0]
        command.extend(["-a", aggression_level])

        # Verbose
        if controls['verbose_check'].isChecked():
            command.append("-v")

        # Additional Options
        if extra_opts := controls['extra_opts_edit'].text().strip():
            command.extend(extra_opts.split())

        # Target(s) must be last
        command.extend(resolved_target.split())

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.whatweb_output_console.clear()

        self.worker = WorkerThread(self._whatweb_thread, args=(command,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_sqlmap_config_widget(self):
        """Creates a reusable, self-contained widget with SQLMap's core configuration options."""
        widget = QGroupBox("SQLMap Core Options")
        layout = QFormLayout(widget)
        controls = {}

        # --- Target ---
        target_layout = QHBoxLayout()
        controls['url_edit'] = QLineEdit()
        controls['url_edit'].setPlaceholderText('e.g., "http://testphp.vulnweb.com/listproducts.php?cat=1"')
        target_layout.addWidget(controls['url_edit'])
        target_layout.addWidget(QLabel(" OR "))
        controls['reqfile_edit'] = QLineEdit()
        controls['reqfile_edit'].setPlaceholderText("Load from request file...")
        browse_req_btn = QPushButton("Browse...")
        browse_req_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['reqfile_edit'], "Select Request File"))
        target_layout.addWidget(controls['reqfile_edit'])
        target_layout.addWidget(browse_req_btn)
        layout.addRow("Target URL (-u) / File (-r):", target_layout)

        # --- Core Options ---
        controls['level_combo'] = QComboBox(); controls['level_combo'].addItems(["1","2","3","4","5"]);
        controls['level_combo'].setCurrentIndex(0)
        layout.addRow("Level (1-5):", controls['level_combo'])

        controls['risk_combo'] = QComboBox(); controls['risk_combo'].addItems(["1","2","3"]);
        controls['risk_combo'].setCurrentIndex(0)
        layout.addRow("Risk (1-3):", controls['risk_combo'])

        # --- Enumeration ---
        enum_box = QGroupBox("Enumeration")
        enum_box.setCheckable(True)
        enum_box.setChecked(False)
        enum_layout = QGridLayout(enum_box)
        controls['enum_banner_check'] = QCheckBox("Banner (-b)")
        controls['enum_dbs_check'] = QCheckBox("Databases (--dbs)")
        controls['enum_tables_check'] = QCheckBox("Tables (--tables)")
        controls['enum_columns_check'] = QCheckBox("Columns (--columns)")
        controls['enum_dump_check'] = QCheckBox("Dump (--dump)")
        enum_layout.addWidget(controls['enum_banner_check'], 0, 0)
        enum_layout.addWidget(controls['enum_dbs_check'], 0, 1)
        enum_layout.addWidget(controls['enum_tables_check'], 1, 0)
        enum_layout.addWidget(controls['enum_columns_check'], 1, 1)
        enum_layout.addWidget(controls['enum_dump_check'], 2, 0)
        layout.addRow(enum_box)

        # --- Action Buttons ---
        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start SQLMap Scan")
        controls['stop_btn'] = QPushButton("Stop SQLMap"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _create_sqlmap_tool(self):
        """Creates the UI for the SQLMap tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.sqlmap_controls = self._create_sqlmap_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.sqlmap_controls['start_btn'])
        buttons_layout.addWidget(self.sqlmap_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.sqlmap_output_console = QPlainTextEdit()
        self.sqlmap_output_console.setReadOnly(True)
        self.sqlmap_output_console.setFont(QFont("Courier New", 10))
        self.sqlmap_output_console.setPlaceholderText("SQLMap output will be displayed here...")
        main_layout.addWidget(self.sqlmap_output_console, 1)

        self.sqlmap_controls['start_btn'].clicked.connect(self.start_sqlmap_scan)
        self.sqlmap_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_theharvester_tool(self):
        """Creates the UI for theHarvester OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>theHarvester - OSINT Tool</b></font>
        <p>This tool gathers emails, subdomains, hosts, employee names, open ports and banners from different public sources like search engines, PGP key servers and SHODAN database.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.theharvester_controls = self._create_theharvester_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.theharvester_controls['start_btn'])
        buttons_layout.addWidget(self.theharvester_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.theharvester_output_console = QPlainTextEdit()
        self.theharvester_output_console.setReadOnly(True)
        self.theharvester_output_console.setFont(QFont("Courier New", 10))
        self.theharvester_output_console.setPlaceholderText("theHarvester output will be displayed here...")
        main_layout.addWidget(self.theharvester_output_console, 1)

        self.theharvester_controls['start_btn'].clicked.connect(self.start_theharvester_scan)
        self.theharvester_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_theharvester_config_widget(self):
        """Creates a reusable, self-contained widget for theHarvester's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['domain_edit'] = QLineEdit("example.com")
        layout.addRow("Domain (-d):", controls['domain_edit'])

        controls['limit_edit'] = QLineEdit("500")
        layout.addRow("Result Limit (-l):", controls['limit_edit'])

        controls['source_edit'] = QLineEdit("google,bing")
        controls['source_edit'].setToolTip("Comma-separated list of data sources (e.g., google,bing,shodan).")
        layout.addRow("Sources (-b):", controls['source_edit'])

        controls['dns_check'] = QCheckBox("DNS Brute-force (-c)")
        controls['dns_check'].setToolTip("Enable DNS brute-forcing for subdomains.")
        layout.addRow(controls['dns_check'])

        controls['dns_lookup_check'] = QCheckBox("DNS Lookup (-n)")
        controls['dns_lookup_check'].setToolTip("Enable DNS lookup for found hosts.")
        layout.addRow(controls['dns_lookup_check'])

        controls['dns_server_edit'] = QLineEdit()
        controls['dns_server_edit'].setPlaceholderText("Optional: e.g., 8.8.8.8")
        layout.addRow("DNS Server (-e):", controls['dns_server_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Harvesting")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_theharvester_command(self, controls):
        """Builds the theHarvester command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("theHarvester", os.path.join("theHarvester", "bin", "theHarvester"))
        if not tool_path:
            tool_path = self._get_tool_path("theharvester", os.path.join("theHarvester", "bin", "theHarvester"))
        if not tool_path:
            return None, None, "theHarvester not found."

        command = [tool_path]

        domain = self._get_control_value(controls, 'domain_edit', 'text')
        if not domain or not domain.strip():
            return None, None, "A domain is required."
        command.extend(["-d", domain.strip()])
        target_for_log = domain.strip()

        limit = self._get_control_value(controls, 'limit_edit', 'text')
        if limit and limit.strip():
            command.extend(["-l", limit.strip()])

        source = self._get_control_value(controls, 'source_edit', 'text')
        if source and source.strip():
            command.extend(["-b", source.strip()])

        dns_server = self._get_control_value(controls, 'dns_server_edit', 'text')
        if dns_server and dns_server.strip():
            command.extend(["-e", dns_server.strip()])

        if self._get_control_value(controls, 'dns_check', 'check'):
            command.append("-c")

        if self._get_control_value(controls, 'dns_lookup_check', 'check'):
            command.append("-n")

        return command, target_for_log, None

    def start_theharvester_scan(self):
        """Starts the theHarvester scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.theharvester_controls
        command, target_for_log, error = self._build_theharvester_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.theharvester_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'theharvester_scan', target_for_log, self.theharvester_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _handle_theharvester_output(self, line):
        self.theharvester_output_console.insertPlainText(line)
        self.theharvester_output_console.verticalScrollBar().setValue(self.theharvester_output_console.verticalScrollBar().maximum())

    def _build_sqlmap_command(self, controls):
        """Builds the SQLMap command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("sqlmap", "sqlmap")
        if not tool_path:
             # Try .py fallback
             tool_path = self._get_tool_path("sqlmap.py", os.path.join("sqlmap", "sqlmap.py"))
             if tool_path and tool_path.endswith(".py"):
                 command = ["python3", tool_path]
             elif tool_path:
                 command = [tool_path]
             else:
                 return None, None, "SQLMap not found."
        else:
             command = [tool_path]

        target_for_log = ""

        # --- Target ---
        url = self._get_control_value(controls, 'url_edit', 'text')
        reqfile = self._get_control_value(controls, 'reqfile_edit', 'text')
        if url and url.strip():
            target_for_log = url.strip()
            command.extend(["-u", target_for_log])
        elif reqfile and reqfile.strip():
            target_for_log = reqfile.strip()
            command.extend(["-r", target_for_log])
        else:
            return None, None, "A Target URL (-u) or Request File (-r) is required."

        # --- Core Options ---
        level = self._get_control_value(controls, 'level_combo', 'combo') or '1'
        command.extend(["--level", level])
        risk = self._get_control_value(controls, 'risk_combo', 'combo') or '1'
        command.extend(["--risk", risk])

        # --- Enumeration ---
        enum_flags = {
            'enum_banner_check': '-b',
            'enum_dbs_check': '--dbs',
            'enum_tables_check': '--tables',
            'enum_columns_check': '--columns',
            'enum_dump_check': '--dump'
        }
        # Check if the enumeration groupbox is checked
        if self._get_control_value(controls, 'enum_box', 'check'):
            for control_name, flag in enum_flags.items():
                if self._get_control_value(controls, control_name, 'check'):
                    command.append(flag)

        # --- Default flags for GUI operation ---
        command.extend(["--batch", "--answers", "quit=N"])

        return command, target_for_log, None

    def start_sqlmap_scan(self):
        """Starts the SQLMap scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.sqlmap_controls
        original_url = self._get_control_value(controls, 'url_edit', 'text')
        if original_url:
            resolved_url = self._resolve_targets_string(original_url)
            controls['url_edit'].setText(resolved_url)

        command, target_for_log, error = self._build_sqlmap_command(controls)

        if original_url:
            controls['url_edit'].setText(original_url)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "SQLMap Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.sqlmap_output_console.clear()

        self.worker = WorkerThread(self._sqlmap_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _sqlmap_thread(self, command, target):
        """Worker thread for running the sqlmap command."""
        q = self.tool_results_queue
        logging.info(f"Starting SQLMap with command: {' '.join(command)}")
        q.put(('sqlmap_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.sqlmap_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('sqlmap_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('sqlmap_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'SQLMap Error', "'sqlmap' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"SQLMap thread error: {e}", exc_info=True)
            q.put(('error', 'SQLMap Error', str(e)))
        finally:
            q.put(('tool_finished', 'sqlmap_scan', target, "".join(full_output)))
            with self.thread_finish_lock:
                self.sqlmap_process = None
            logging.info("SQLMap scan thread finished.")

    def _create_hashcat_config_widget(self):
        """Creates a reusable, self-contained widget with Hashcat's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)
        controls = {}

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Hashcat: The world's fastest password cracker.</b></font>
        <p><b>WARNING:</b> This is a powerful tool. Ensure you have permission to crack the provided hashes. This tool can be very resource-intensive.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Main Tab Widget for Options ---
        hashcat_tabs = QTabWidget()
        main_layout.addWidget(hashcat_tabs)

        # --- Main Config Tab ---
        config_tab = QWidget()
        config_layout = QFormLayout(config_tab)

        hashfile_layout = QHBoxLayout()
        controls['hashfile_edit'] = QLineEdit()
        controls['hashfile_edit'].setToolTip("The file containing the hashes to crack.")
        hashfile_layout.addWidget(controls['hashfile_edit'])
        browse_hash_btn = QPushButton("Browse...")
        browse_hash_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['hashfile_edit'], "Select Hash File"))
        hashfile_layout.addWidget(browse_hash_btn)
        config_layout.addRow("Hash File:", hashfile_layout)

        hash_type_layout = QVBoxLayout()
        controls['type_edit'] = QLineEdit()
        controls['type_edit'].setPlaceholderText("e.g., 0 for MD5, 1000 for NTLM")
        controls['type_edit'].setToolTip("The hash mode code (-m). Click the link below to find the correct mode for your hash type.")
        hash_type_label = QLabel("Hash Mode (-m):")
        hash_type_link = QLabel('<a href="https://hashcat.net/wiki/doku.php?id=example_hashes">Find Hash Mode</a>')
        hash_type_link.setOpenExternalLinks(True)
        hash_type_layout.addWidget(controls['type_edit'])
        hash_type_layout.addWidget(hash_type_link)
        config_layout.addRow(hash_type_label, hash_type_layout)

        controls['attack_mode_combo'] = QComboBox()
        controls['attack_mode_combo'].addItems([
            "0 - Straight (Dictionary)",
            "1 - Combination",
            "3 - Brute-force (Mask)",
            "6 - Hybrid (Wordlist + Mask)",
            "7 - Hybrid (Mask + Wordlist)"
        ])
        controls['attack_mode_combo'].setToolTip("The attack mode (-a) to use.\n- Straight: Dictionary attack.\n- Combination: Combines words from two dictionaries.\n- Brute-force: Tries all possible character combinations based on a mask.\n- Hybrid: Combines dictionary words with a mask.")
        config_layout.addRow("Attack Mode (-a):", controls['attack_mode_combo'])

        outfile_layout = QHBoxLayout()
        controls['outfile_edit'] = QLineEdit()
        controls['outfile_edit'].setToolTip("The file to save cracked hashes to (-o).")
        outfile_layout.addWidget(controls['outfile_edit'])
        browse_out_btn = QPushButton("Browse...")
        browse_out_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['outfile_edit'], "Save Output File", "Text Files (*.txt);;All Files (*)"))
        outfile_layout.addWidget(browse_out_btn)
        config_layout.addRow("Output File (-o):", outfile_layout)

        hashcat_tabs.addTab(config_tab, "Configuration")

        # --- Wordlist Tab ---
        controls['wordlist_tab'] = QWidget()
        wordlist_layout = QVBoxLayout(controls['wordlist_tab'])
        controls['wordlist_list'] = QListWidget()
        wordlist_layout.addWidget(controls['wordlist_list'])
        wordlist_buttons = QHBoxLayout()
        add_wordlist_btn = QPushButton("Add Wordlist(s)")
        add_wordlist_btn.clicked.connect(lambda: self._browse_files_for_listwidget(controls['wordlist_list'], "Select Wordlist(s)"))
        remove_wordlist_btn = QPushButton("Remove Selected")
        remove_wordlist_btn.clicked.connect(lambda: controls['wordlist_list'].takeItem(controls['wordlist_list'].currentRow()))
        wordlist_buttons.addWidget(add_wordlist_btn)
        wordlist_buttons.addWidget(remove_wordlist_btn)
        wordlist_layout.addLayout(wordlist_buttons)
        hashcat_tabs.addTab(controls['wordlist_tab'], "Wordlists")

        # --- Mask Tab ---
        controls['mask_tab'] = QWidget()
        mask_layout = QFormLayout(controls['mask_tab'])
        controls['mask_edit'] = QLineEdit()
        controls['mask_edit'].setPlaceholderText("e.g., ?l?l?l?l?l?l?l?l")
        controls['mask_edit'].setToolTip("The mask to use for brute-force or hybrid attacks. Click the link below for syntax details.")
        mask_layout.addRow("Mask:", controls['mask_edit'])
        mask_link = QLabel('<a href="https://hashcat.net/wiki/doku.php?id=mask_attack">Mask Attack Info</a>')
        mask_link.setOpenExternalLinks(True)
        mask_layout.addRow(mask_link)
        hashcat_tabs.addTab(controls['mask_tab'], "Mask")

        # --- Advanced Tab ---
        adv_tab = QWidget()
        adv_layout = QFormLayout(adv_tab)
        controls['force_check'] = QCheckBox("Ignore warnings")
        controls['force_check'].setToolTip("Ignore warnings and force the cracking session to start (--force).")
        adv_layout.addRow("Force:", controls['force_check'])
        controls['extra_opts_edit'] = QLineEdit()
        controls['extra_opts_edit'].setToolTip("Enter any additional, space-separated Hashcat flags here. These will be appended to the command.")
        adv_layout.addRow("Additional Options:", controls['extra_opts_edit'])
        adv_tab.setLayout(adv_layout)
        hashcat_tabs.addTab(adv_tab, "Advanced")

        # --- UI Logic ---
        def update_tabs(text):
            mode = int(text.split(" ")[0])
            controls['wordlist_tab'].setEnabled(mode in [0, 1, 6, 7])
            controls['mask_tab'].setEnabled(mode in [3, 6, 7])
        controls['attack_mode_combo'].currentTextChanged.connect(update_tabs)
        update_tabs(controls['attack_mode_combo'].currentText()) # Initial state

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("tool.svg")), " Start Hashcat")
        controls['stop_btn'] = QPushButton("Stop Hashcat"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _create_hashcat_tool(self):
        """Creates the UI for the Hashcat tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.hashcat_controls = self._create_hashcat_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons & Output ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.hashcat_controls['start_btn'])
        buttons_layout.addWidget(self.hashcat_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        self.hashcat_output_console = QPlainTextEdit()
        self.hashcat_output_console.setReadOnly(True)
        self.hashcat_output_console.setFont(QFont("Courier New", 10))
        main_layout.addWidget(self.hashcat_output_console, 1)

        self.hashcat_controls['start_btn'].clicked.connect(self.start_hashcat_scan)
        self.hashcat_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _build_hashcat_command(self, controls):
        """Builds the Hashcat command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("hashcat", "hashcat")
        if not tool_path:
             return None, None, "Hashcat not found."
        command = [tool_path]

        # --- Config Tab ---
        hashfile = self._get_control_value(controls, 'hashfile_edit', 'text')
        if not hashfile or not hashfile.strip():
            return None, None, "Hash file is required."
        hashfile = hashfile.strip()
        command.append(hashfile)
        target_for_log = hashfile

        hash_type = self._get_control_value(controls, 'type_edit', 'text')
        if not hash_type or not hash_type.strip():
            return None, None, "Hash mode (-m) is required."
        command.extend(["-m", hash_type.strip()])

        attack_mode_text = self._get_control_value(controls, 'attack_mode_combo', 'combo')
        attack_mode = attack_mode_text.split(" ")[0] if attack_mode_text else "0"
        command.extend(["-a", attack_mode])

        if outfile := self._get_control_value(controls, 'outfile_edit', 'text'):
            if outfile.strip(): command.extend(["-o", outfile.strip()])

        # --- Wordlist/Mask Tabs ---
        attack_mode_int = int(attack_mode)
        if attack_mode_int in [0, 1, 6, 7]: # Modes that use wordlists
            is_lab_run = not isinstance(controls.get('wordlist_list'), QWidget) if 'wordlist_list' in controls else True
            if is_lab_run:
                wordlists = controls.get('wordlist_list', [])
            else:
                wordlist_widget = controls['wordlist_list']
                wordlists = [wordlist_widget.item(i).text() for i in range(wordlist_widget.count())]

            if not wordlists:
                return None, None, "This attack mode requires at least one wordlist."
            command.extend(wordlists)

        if attack_mode_int in [3, 6, 7]: # Modes that use masks
            mask = self._get_control_value(controls, 'mask_edit', 'text')
            if not mask or not mask.strip():
                return None, None, "This attack mode requires a mask."
            command.append(mask.strip())

        # --- Advanced Tab ---
        if self._get_control_value(controls, 'force_check', 'check'):
            command.append("--force")
        if extra_opts := self._get_control_value(controls, 'extra_opts_edit', 'text'):
            if extra_opts.strip(): command.extend(extra_opts.strip().split())

        return command, target_for_log, None

    def start_hashcat_scan(self):
        """Starts the Hashcat worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.hashcat_controls
        command, target_for_log, error = self._build_hashcat_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "Hashcat Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.hashcat_output_console.clear()

        self.worker = WorkerThread(self._hashcat_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _hashcat_thread(self, command, hashfile):
        """Worker thread for running the hashcat command."""
        q = self.tool_results_queue
        logging.info(f"Starting Hashcat with command: {' '.join(command)}")
        q.put(('hashcat_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.hashcat_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('hashcat_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('hashcat_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Hashcat Error', "'hashcat' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Hashcat thread error: {e}", exc_info=True)
            q.put(('error', 'Hashcat Error', str(e)))
        finally:
            q.put(('tool_finished', 'hashcat_scan', hashfile, "".join(full_output)))
            with self.thread_finish_lock:
                self.hashcat_process = None
            logging.info("Hashcat scan thread finished.")

    def _create_nuclei_tool(self):
        """Creates the UI for the Nuclei Web Scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.nuclei_controls = self._create_nuclei_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.nuclei_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.nuclei_output_console = QPlainTextEdit()
        self.nuclei_output_console.setReadOnly(True)
        self.nuclei_output_console.setFont(QFont("Courier New", 10))
        self.nuclei_output_console.setPlaceholderText("Nuclei output will be displayed here...")
        main_layout.addWidget(self.nuclei_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_nuclei_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_nuclei_config_widget(self):
        """Creates a reusable, self-contained widget with Nuclei's configuration options."""
        widget = QGroupBox("Nuclei Configuration")
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(10, 10, 10, 10)

        controls = {}

        # --- Main Tab Widget for Options ---
        nuclei_tabs = QTabWidget()
        main_layout.addWidget(nuclei_tabs)

        # --- Main Tab Widget for Options ---
        nuclei_tabs = QTabWidget()
        main_layout.addWidget(nuclei_tabs)

        # --- Target Tab ---
        target_tab = QWidget()
        target_layout = QFormLayout(target_tab)

        target_file_layout = QHBoxLayout()
        controls['target_edit'] = QLineEdit()
        controls['target_edit'].setPlaceholderText("e.g., https://example.com or path to target list file...")
        controls['target_edit'].setToolTip("Enter a single target URL (-u) or path to a file with a list of targets (-l).")
        target_file_layout.addWidget(controls['target_edit'])
        browse_target_btn = QPushButton("Browse...")
        browse_target_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['target_edit'], "Select Target List File"))
        target_file_layout.addWidget(browse_target_btn)
        target_layout.addRow("Target URL/List (-u/-l):", target_file_layout)
        nuclei_tabs.addTab(target_tab, "Target")

        # --- Templates Tab ---
        templates_tab = QWidget()
        templates_layout = QFormLayout(templates_tab)

        template_file_layout = QHBoxLayout()
        controls['templates_edit'] = QLineEdit()
        controls['templates_edit'].setPlaceholderText("e.g., cves/, http/exposures/, path/to/custom/templates/")
        controls['templates_edit'].setToolTip("Comma-separated list of template directories or single template files to run (-t).")
        template_file_layout.addWidget(controls['templates_edit'])
        browse_template_btn = QPushButton("Browse...")
        browse_template_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['templates_edit'], "Select Template File or Directory"))
        template_file_layout.addWidget(browse_template_btn)
        templates_layout.addRow("Templates (-t):", template_file_layout)

        controls['severity_combo'] = QComboBox()
        controls['severity_combo'].addItems(["all", "info", "low", "medium", "high", "critical"])
        controls['severity_combo'].setToolTip("Filter templates by severity (-s).")
        templates_layout.addRow("Severity (-s):", controls['severity_combo'])

        controls['tags_edit'] = QLineEdit()
        controls['tags_edit'].setPlaceholderText("e.g., cve,rce,wordpress")
        controls['tags_edit'].setToolTip("Filter templates by tags (-tags).")
        templates_layout.addRow("Tags (-tags):", controls['tags_edit'])

        nuclei_tabs.addTab(templates_tab, "Templates")

        # --- Output & Config Tab ---
        config_tab = QWidget()
        config_layout = QFormLayout(config_tab)

        output_file_layout = QHBoxLayout()
        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("Optional: path to save report...")
        output_file_layout.addWidget(controls['output_edit'])
        browse_output_btn = QPushButton("Browse...")
        browse_output_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_edit'], "Save Nuclei Report"))
        output_file_layout.addWidget(browse_output_btn)
        config_layout.addRow("Output File (-o):", output_file_layout)

        controls['concurrency_edit'] = QLineEdit("25")
        controls['concurrency_edit'].setToolTip("Number of concurrent requests to send (-c).")
        config_layout.addRow("Concurrency (-c):", controls['concurrency_edit'])

        controls['ratelimit_edit'] = QLineEdit("150")
        controls['ratelimit_edit'].setToolTip("Requests per second (-rl).")
        config_layout.addRow("Rate Limit (-rl):", controls['ratelimit_edit'])

        controls['verbose_check'] = QCheckBox("Verbose Output (-v)")
        config_layout.addRow(controls['verbose_check'])

        nuclei_tabs.addTab(config_tab, "Configuration")

        # --- Action Buttons ---
        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Nuclei Scan")
        controls['stop_btn'] = QPushButton("Stop Nuclei"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_nuclei_command(self, controls):
        """Builds the Nuclei command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("nuclei", "nuclei")
        if not tool_path:
             return None, None, "Nuclei not found."
        command = [tool_path]

        target = self._get_control_value(controls, 'target_edit', 'text')
        if not target or not target.strip():
            return None, None, "A target or target list is required."
        target = target.strip()
        target_for_log = target

        # Determine if target is a file or a URL
        if os.path.exists(target):
            command.extend(["-l", target])
        else:
            command.extend(["-u", target])

        if templates := self._get_control_value(controls, 'templates_edit', 'text'):
            if templates.strip(): command.extend(["-t", templates.strip()])

        if (severity := self._get_control_value(controls, 'severity_combo', 'combo')) != "all":
            if severity: command.extend(["-s", severity])

        if tags := self._get_control_value(controls, 'tags_edit', 'text'):
            if tags.strip(): command.extend(["-tags", tags.strip()])

        if output_file := self._get_control_value(controls, 'output_edit', 'text'):
            if output_file.strip(): command.extend(["-o", output_file.strip()])

        if concurrency := self._get_control_value(controls, 'concurrency_edit', 'text'):
            if concurrency.strip(): command.extend(["-c", concurrency.strip()])

        if ratelimit := self._get_control_value(controls, 'ratelimit_edit', 'text'):
            if ratelimit.strip(): command.extend(["-rl", ratelimit.strip()])

        if self._get_control_value(controls, 'verbose_check', 'check'):
            command.append("-v")

        command.extend(["-nC", "-json"]) # Always disable color and enable JSON for parsing

        return command, target_for_log, None

    def start_nuclei_scan(self):
        """Starts the Nuclei scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.nuclei_controls
        command, target_for_log, error = self._build_nuclei_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "Nuclei Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.nuclei_output_console.clear()

        self.worker = WorkerThread(self._nuclei_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _nuclei_thread(self, command, target):
        """Worker thread for running the nuclei command."""
        q = self.tool_results_queue
        logging.info(f"Starting Nuclei with command: {' '.join(command)}")
        q.put(('nuclei_output', f"$ {' '.join(command)}\n\n"))
        json_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.nuclei_process = process

            full_output = []
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('nuclei_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                # Don't send JSON to the console, just capture it
                full_output.append(line)

            process.stdout.close()
            process.wait()

            json_data = "".join(full_output)
            q.put(('nuclei_output', json_data)) # Put the full JSON blob in the console
            if not self.tool_stop_event.is_set() and json_data.strip():
                q.put(('nuclei_results', json_data))


        except FileNotFoundError:
            q.put(('error', 'Nuclei Error', "'nuclei' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Nuclei thread error: {e}", exc_info=True)
            q.put(('error', 'Nuclei Error', str(e)))
        finally:
            q.put(('tool_finished', 'nuclei_scan', target, json_data))
            with self.thread_finish_lock:
                self.nuclei_process = None
            logging.info("Nuclei scan thread finished.")

    def _handle_nuclei_output(self, line):
        self.nuclei_output_console.insertPlainText(line)
        self.nuclei_output_console.verticalScrollBar().setValue(self.nuclei_output_console.verticalScrollBar().maximum())

    def _create_trufflehog_tool(self):
        """Creates the UI for the TruffleHog Secret Scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.trufflehog_controls = self._create_trufflehog_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.trufflehog_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.trufflehog_output_console = QPlainTextEdit()
        self.trufflehog_output_console.setReadOnly(True)
        self.trufflehog_output_console.setFont(QFont("Courier New", 10))
        self.trufflehog_output_console.setPlaceholderText("TruffleHog output will be displayed here...")
        main_layout.addWidget(self.trufflehog_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_trufflehog_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_trufflehog_config_widget(self):
        """Creates a reusable, self-contained widget with TruffleHog's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)

        controls = {}

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>TruffleHog - Secret Scanner</b></font>
        <p>This tool scans sources like git repositories, GitHub, and filesystems for leaked secrets.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Source Type and Target ---
        source_box = QGroupBox("Scan Target")
        source_layout = QVBoxLayout(source_box)

        controls['source_type_group'] = QButtonGroup(self)
        rb_layout = QHBoxLayout()
        controls['git_rb'] = QRadioButton("Git Repo"); controls['git_rb'].setChecked(True)
        controls['github_rb'] = QRadioButton("GitHub")
        controls['filesystem_rb'] = QRadioButton("Filesystem")
        controls['source_type_group'].addButton(controls['git_rb'])
        controls['source_type_group'].addButton(controls['github_rb'])
        controls['source_type_group'].addButton(controls['filesystem_rb'])
        rb_layout.addWidget(controls['git_rb']); rb_layout.addWidget(controls['github_rb']); rb_layout.addWidget(controls['filesystem_rb'])
        source_layout.addLayout(rb_layout)

        target_layout = QHBoxLayout()
        controls['target_edit'] = QLineEdit()
        controls['target_edit'].setPlaceholderText("Enter Git URL, GitHub repo, or filesystem path...")
        target_layout.addWidget(controls['target_edit'])
        controls['browse_btn'] = QPushButton("Browse...")
        controls['browse_btn'].clicked.connect(lambda: self._browse_dir_for_lineedit(controls['target_edit'], "Select Directory to Scan"))
        target_layout.addWidget(controls['browse_btn'])
        source_layout.addLayout(target_layout)

        main_layout.addWidget(source_box)

        # --- Options ---
        options_box = QGroupBox("Options")
        options_layout = QFormLayout(options_box)
        controls['only_verified_check'] = QCheckBox("Only Verified Results")
        controls['only_verified_check'].setToolTip("Only output secrets that have been successfully verified against their respective APIs.")
        options_layout.addRow(controls['only_verified_check'])
        main_layout.addWidget(options_box)

        # --- Action Buttons ---
        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        # UI Logic to show/hide browse button
        def toggle_browse_button():
            controls['browse_btn'].setVisible(controls['filesystem_rb'].isChecked())
        controls['source_type_group'].buttonClicked.connect(toggle_browse_button)
        toggle_browse_button() # Set initial state

        return widget, controls

    def _browse_files_for_listwidget(self, list_widget, dialog_title):
        file_paths, _ = QFileDialog.getOpenFileNames(self, dialog_title, "", "Text Files (*.txt);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_paths:
            list_widget.addItems(file_paths)

    def _browse_dir_for_lineedit(self, line_edit_widget, dialog_title):
        dir_path = QFileDialog.getExistingDirectory(self, dialog_title, options=QFileDialog.Option.DontUseNativeDialog)
        if dir_path:
            line_edit_widget.setText(dir_path)

    def _build_trufflehog_command(self, controls):
        """Builds the TruffleHog command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("trufflehog", "trufflehog")
        if not tool_path:
             return None, None, "TruffleHog not found."
        command = [tool_path]

        target = self._get_control_value(controls, 'target_edit', 'text')
        if not target or not target.strip():
            return None, None, "A target is required."
        target = target.strip()
        target_for_log = target

        source_type = ""
        if self._get_control_value(controls, 'git_rb', 'check'):
            source_type = "git"
        elif self._get_control_value(controls, 'github_rb', 'check'):
            source_type = "github"
        elif self._get_control_value(controls, 'filesystem_rb', 'check'):
            source_type = "filesystem"

        if not source_type:
             return None, None, "A source type (Git, GitHub, Filesystem) must be selected."

        command.extend([source_type, target])

        if self._get_control_value(controls, 'only_verified_check', 'check'):
            command.append("--only-verified")

        command.append("--json") # Always add --json for parsing

        return command, target_for_log, None

    def start_trufflehog_scan(self):
        """Starts the TruffleHog scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.trufflehog_controls
        command, target_for_log, error = self._build_trufflehog_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "TruffleHog Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.trufflehog_output_console.clear()

        self.worker = WorkerThread(self._trufflehog_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _trufflehog_thread(self, command, target):
        """Worker thread for running the trufflehog command."""
        q = self.tool_results_queue
        logging.info(f"Starting TruffleHog with command: {' '.join(command)}")
        q.put(('trufflehog_output', f"$ {' '.join(command)}\n\n"))
        json_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.trufflehog_process = process

            full_output = []
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('trufflehog_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                # Don't send JSON to the console, just capture it
                full_output.append(line)

            process.stdout.close()
            process.wait()

            json_data = "".join(full_output)
            q.put(('trufflehog_output', json_data))
            if not self.tool_stop_event.is_set() and json_data.strip():
                q.put(('trufflehog_results', json_data))

        except FileNotFoundError:
            q.put(('error', 'TruffleHog Error', "'trufflehog' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"TruffleHog thread error: {e}", exc_info=True)
            q.put(('error', 'TruffleHog Error', str(e)))
        finally:
            q.put(('tool_finished', 'trufflehog_scan', target, json_data))
            with self.thread_finish_lock:
                self.trufflehog_process = None
            logging.info("TruffleHog scan thread finished.")

    def _handle_trufflehog_output(self, line):
        self.trufflehog_output_console.insertPlainText(line)
        self.trufflehog_output_console.verticalScrollBar().setValue(self.trufflehog_output_console.verticalScrollBar().maximum())

    def _create_jtr_tool(self):
        """Creates the UI for the John the Ripper tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.jtr_controls = self._create_jtr_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.jtr_controls['start_btn'])
        buttons_layout.addWidget(self.jtr_controls['show_btn'])
        buttons_layout.addWidget(self.jtr_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.jtr_output_console = QPlainTextEdit()
        self.jtr_output_console.setReadOnly(True)
        self.jtr_output_console.setFont(QFont("Courier New", 10))
        self.jtr_output_console.setPlaceholderText("John the Ripper output will be displayed here...")
        main_layout.addWidget(self.jtr_output_console, 1)

        self.jtr_controls['start_btn'].clicked.connect(self.start_jtr_crack)
        self.jtr_controls['show_btn'].clicked.connect(self.show_jtr_cracked)
        self.jtr_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_jtr_config_widget(self):
        """Creates a reusable, self-contained widget with JTR's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        hash_file_layout = QHBoxLayout()
        controls['hash_file_edit'] = QLineEdit()
        controls['hash_file_edit'].setPlaceholderText("Path to hash file (required)...")
        hash_file_layout.addWidget(controls['hash_file_edit'])
        browse_hash_btn = QPushButton("Browse...")
        browse_hash_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['hash_file_edit'], "Select Hash File"))
        hash_file_layout.addWidget(browse_hash_btn)
        main_layout.addRow("Hash File:", hash_file_layout)

        wordlist_layout = QHBoxLayout()
        controls['wordlist_edit'] = QLineEdit()
        controls['wordlist_edit'].setPlaceholderText("Path to wordlist file (optional)...")
        wordlist_layout.addWidget(controls['wordlist_edit'])
        browse_wordlist_btn = QPushButton("Browse...")
        browse_wordlist_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['wordlist_edit'], "Select Wordlist File"))
        wordlist_layout.addWidget(browse_wordlist_btn)
        main_layout.addRow("Wordlist:", wordlist_layout)

        controls['format_edit'] = QLineEdit()
        controls['format_edit'].setPlaceholderText("e.g., raw-md5, nt, sha512crypt")
        controls['format_edit'].setToolTip("Specify the hash format (--format). Leave blank for auto-detection.")
        main_layout.addRow("Format:", controls['format_edit'])

        controls['rules_check'] = QCheckBox("Enable Word Mangling Rules")
        controls['rules_check'].setToolTip("Enable rules for wordlist mode (--rules).")
        main_layout.addRow(controls['rules_check'])

        controls['incremental_check'] = QCheckBox("Incremental Mode (Brute-force)")
        controls['incremental_check'].setToolTip("Enable incremental mode (--incremental). If wordlist is also specified, this will run after.")
        main_layout.addRow(controls['incremental_check'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("tool.svg")), " Start Cracking")
        controls['show_btn'] = QPushButton("Show Cracked")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_jtr_crack(self):
        """Starts the John the Ripper worker thread."""
        self._start_jtr_generic(crack_mode=True)

    def show_jtr_cracked(self):
        """Shows already cracked passwords."""
        self._start_jtr_generic(crack_mode=False)

    def _build_jtr_command(self, controls, crack_mode):
        """Builds the John the Ripper command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("john", "john")
        if not tool_path:
             return None, None, "John the Ripper not found."
        command = [tool_path]

        hash_file = self._get_control_value(controls, 'hash_file_edit', 'text')
        if not hash_file or not hash_file.strip():
            return None, None, "A hash file is required."
        hash_file = hash_file.strip()
        target_for_log = hash_file

        if crack_mode:
            wordlist = self._get_control_value(controls, 'wordlist_edit', 'text')
            if wordlist and wordlist.strip():
                command.extend([f"--wordlist={wordlist.strip()}"])
                if self._get_control_value(controls, 'rules_check', 'check'):
                    command.append("--rules")
            if self._get_control_value(controls, 'incremental_check', 'check'):
                command.append("--incremental")
        else: # Show mode
            command.append("--show")

        if format_type := self._get_control_value(controls, 'format_edit', 'text'):
            if format_type.strip():
                command.extend([f"--format={format_type.strip()}"])

        command.append(hash_file)

        return command, target_for_log, None

    def _start_jtr_generic(self, crack_mode):
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.jtr_controls
        command, target_for_log, error = self._build_jtr_command(controls, crack_mode)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "JTR Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['show_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.jtr_output_console.clear()

        self.worker = WorkerThread(self._jtr_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _jtr_thread(self, command, hash_file):
        """Worker thread for running the john command."""
        q = self.tool_results_queue
        logging.info(f"Starting JTR with command: {' '.join(command)}")
        q.put(('jtr_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.jtr_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('jtr_output', "\n\n--- Canceled By User ---\n"))
                    break
                q.put(('jtr_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'JTR Error', "'john' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"JTR thread error: {e}", exc_info=True)
            q.put(('error', 'JTR Error', str(e)))
        finally:
            q.put(('tool_finished', 'jtr_scan', hash_file, "".join(full_output)))
            with self.thread_finish_lock:
                self.jtr_process = None
            logging.info("JTR scan thread finished.")

    def _handle_jtr_output(self, line):
        self.jtr_output_console.insertPlainText(line)
        self.jtr_output_console.verticalScrollBar().setValue(self.jtr_output_console.verticalScrollBar().maximum())


    def _create_hydra_config_widget(self):
        """Creates a reusable, self-contained widget with Hydra's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['target_edit'] = QLineEdit("localhost")
        main_layout.addRow("Target Host:", controls['target_edit'])

        controls['service_edit'] = QLineEdit("ssh")
        controls['service_edit'].setToolTip("The service to attack (e.g., ssh, ftp, smb, rdp).")
        main_layout.addRow("Service:", controls['service_edit'])

        user_layout = QHBoxLayout()
        controls['user_edit'] = QLineEdit("root")
        controls['user_edit'].setToolTip("A single username to test (-l).")
        user_layout.addWidget(controls['user_edit'])
        user_layout.addWidget(QLabel("OR"))
        controls['user_list_edit'] = QLineEdit()
        controls['user_list_edit'].setToolTip("Path to a file containing a list of usernames (-L).")
        user_layout.addWidget(controls['user_list_edit'])
        browse_user_btn = QPushButton("Browse...")
        browse_user_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['user_list_edit'], "Select User List"))
        user_layout.addWidget(browse_user_btn)
        main_layout.addRow("Username (-l) / List (-L):", user_layout)

        pass_layout = QHBoxLayout()
        controls['pass_edit'] = QLineEdit()
        controls['pass_edit'].setToolTip("A single password to test (-p).")
        pass_layout.addWidget(controls['pass_edit'])
        pass_layout.addWidget(QLabel("OR"))
        controls['pass_list_edit'] = QLineEdit()
        controls['pass_list_edit'].setToolTip("Path to a file containing a list of passwords (-P).")
        pass_layout.addWidget(controls['pass_list_edit'])
        browse_pass_btn = QPushButton("Browse...")
        browse_pass_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['pass_list_edit'], "Select Password List"))
        pass_layout.addWidget(browse_pass_btn)
        main_layout.addRow("Password (-p) / List (-P):", pass_layout)

        controls['tasks_edit'] = QLineEdit("16")
        controls['tasks_edit'].setToolTip("Number of parallel tasks/threads (-t).")
        main_layout.addRow("Tasks (-t):", controls['tasks_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("tool.svg")), " Start Attack")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_hydra_command(self, controls):
        """Builds the Hydra command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("hydra", "hydra")
        if not tool_path:
             return None, None, "Hydra not found."
        command = [tool_path]

        target = self._get_control_value(controls, 'target_edit', 'text')
        service = self._get_control_value(controls, 'service_edit', 'text')

        if not target or not target.strip() or not service or not service.strip():
            return None, None, "Target and Service are required."

        target = target.strip()
        service = service.strip()
        target_for_log = f"{service}://{target}"

        user = self._get_control_value(controls, 'user_edit', 'text')
        user_list = self._get_control_value(controls, 'user_list_edit', 'text')
        if user and user.strip():
            command.extend(["-l", user.strip()])
        elif user_list and user_list.strip():
            command.extend(["-L", user_list.strip()])
        else:
            return None, None, "A username or user list is required."

        pwd = self._get_control_value(controls, 'pass_edit', 'text')
        pass_list = self._get_control_value(controls, 'pass_list_edit', 'text')
        if pwd: # Do not strip password
            command.extend(["-p", pwd])
        elif pass_list and pass_list.strip():
            command.extend(["-P", pass_list.strip()])
        else:
            return None, None, "A password or password list is required."

        if tasks := self._get_control_value(controls, 'tasks_edit', 'text'):
            if tasks.strip():
                command.extend(["-t", tasks.strip()])

        command.append(target_for_log)

        return command, target_for_log, None

    def start_hydra_attack(self):
        """Starts the Hydra attack worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.hydra_controls
        original_target = self._get_control_value(controls, 'target_edit', 'text')
        resolved_target = self._resolve_targets_string(original_target)
        controls['target_edit'].setText(resolved_target)
        command, target_for_log, error = self._build_hydra_command(controls)
        controls['target_edit'].setText(original_target)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "Hydra Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.hydra_output_console.clear()

        self.worker = WorkerThread(self._hydra_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _hydra_thread(self, command, target):
        """Worker thread for running the hydra command."""
        q = self.tool_results_queue
        logging.info(f"Starting Hydra with command: {' '.join(command)}")
        q.put(('hydra_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.hydra_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('hydra_output', "\n\n--- Canceled By User ---\n"))
                    break
                q.put(('hydra_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Hydra Error', "'hydra' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Hydra thread error: {e}", exc_info=True)
            q.put(('error', 'Hydra Error', str(e)))
        finally:
            q.put(('tool_finished', 'hydra_scan', target, "".join(full_output)))
            with self.thread_finish_lock:
                self.hydra_process = None
            logging.info("Hydra scan thread finished.")

    def _handle_hydra_output(self, line):
        self.hydra_output_console.insertPlainText(line)
        self.hydra_output_console.verticalScrollBar().setValue(self.hydra_output_console.verticalScrollBar().maximum())

    def _create_firewall_evasion_tab(self):
        """Creates the UI for the Firewall Evasion tab, using a simplified Nmap scanner."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Nmap Firewall Evasion Scanner</b></font>
        <p>This tool uses various Nmap techniques to attempt to bypass firewall and IDS/IPS rules. These scans can be "noisy" and may be detected.</p>
        <p><b>Note:</b> A SYN Stealth Scan (-sS) is used by default, which requires root privileges.</p>
        """)
        instructions.setFixedHeight(100)
        main_layout.addWidget(instructions)


        # Simplified config widget
        config_widget, self.nmap_evasion_controls = self._create_nmap_evasion_config_widget()
        main_layout.addWidget(config_widget)

        # Output Console
        self.nmap_evasion_output_console = QPlainTextEdit()
        self.nmap_evasion_output_console.setReadOnly(True)
        self.nmap_evasion_output_console.setFont(QFont("Courier New", 10))
        self.nmap_evasion_output_console.setPlaceholderText("Nmap evasion scan output will be displayed here...")
        main_layout.addWidget(self.nmap_evasion_output_console, 1)

        return widget

    def _create_nmap_evasion_config_widget(self):
        """Creates the UI controls for the simplified Nmap Evasion Scanner."""
        widget = QGroupBox("Evasion Scan Options")
        main_layout = QFormLayout(widget)
        controls = {}

        # --- Core Controls ---
        controls['target_edit'] = QLineEdit("localhost")
        main_layout.addRow("Target(s):", controls['target_edit'])

        controls['ports_edit'] = QLineEdit("80,443")
        main_layout.addRow("Port(s):", controls['ports_edit'])

        # --- Evasion Techniques ---
        evasion_box = QGroupBox("Techniques")
        evasion_layout = QGridLayout(evasion_box)

        controls['frag_check'] = QCheckBox("Fragment packets (-f)")
        evasion_layout.addWidget(controls['frag_check'], 0, 0)

        controls['badsum_check'] = QCheckBox("Use bad checksums (--badsum)")
        evasion_layout.addWidget(controls['badsum_check'], 0, 1)

        decoys_layout = QHBoxLayout()
        decoys_layout.addWidget(QLabel("Decoys (-D):"))
        controls['decoys_edit'] = QLineEdit()
        controls['decoys_edit'].setPlaceholderText("e.g., ME,decoy1,decoy2")
        decoys_layout.addWidget(controls['decoys_edit'])
        evasion_layout.addLayout(decoys_layout, 1, 0, 1, 2)

        source_port_layout = QHBoxLayout()
        source_port_layout.addWidget(QLabel("Source Port (-g):"))
        controls['source_port_edit'] = QLineEdit()
        controls['source_port_edit'].setPlaceholderText("e.g., 53, 88")
        source_port_layout.addWidget(controls['source_port_edit'])
        evasion_layout.addLayout(source_port_layout, 2, 0, 1, 2)

        main_layout.addRow(evasion_box)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        controls['start_btn'] = QPushButton(QIcon(self.icon_path("shield.svg")), " Start Evasion Scan")
        controls['cancel_btn'] = QPushButton("Cancel"); controls['cancel_btn'].setEnabled(False)
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['cancel_btn'])
        main_layout.addRow(buttons_layout)

        # --- Connections ---
        controls['start_btn'].clicked.connect(self.start_nmap_evasion_scan)
        controls['cancel_btn'].clicked.connect(self.cancel_tool)

        return widget, controls

    def start_nmap_evasion_scan(self, sudo_password=None):
        """Starts the Nmap Evasion scan worker thread."""
        if self.is_tool_running and not sudo_password:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.nmap_evasion_controls
        command, target_for_log, error = self._build_nmap_evasion_command(controls)
        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "Nmap Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        # Evasion scans with -sS require root
        if not sudo_password and sys.platform != "win32":
            if 'nmap_evasion_scan' not in self.sudo_cancel_handlers:
                self.sudo_cancel_handlers['nmap_evasion_scan'] = lambda: (
                    controls['start_btn'].setEnabled(True),
                    controls['cancel_btn'].setEnabled(False)
                )
            self._run_command_with_sudo_prompt(command, self.start_nmap_evasion_scan, 'nmap_evasion_scan')
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.nmap_evasion_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'nmap_evasion_scan', target_for_log, self.nmap_evasion_output_console, sudo_password
        ))
        self.active_threads.append(self.worker)
        self.worker.start()


    def _build_nmap_evasion_command(self, controls):
        """Builds a simplified Nmap command focused on evasion techniques."""
        target = self._get_control_value(controls, 'target_edit', 'text')
        ports = self._get_control_value(controls, 'ports_edit', 'text')
        if not target or not target.strip():
            return None, None, "Target is required."
        if not ports or not ports.strip():
            return None, None, "Ports are required."

        target_for_log = target.strip()

        tool_path = self._get_tool_path("nmap", "nmap")
        if not tool_path:
             return None, None, "Nmap not found."

        command = [tool_path, "-sS", "-p", ports.strip()] # Default to SYN scan

        if self._get_control_value(controls, 'frag_check', 'check'): command.append("-f")
        if self._get_control_value(controls, 'badsum_check', 'check'): command.append("--badsum")
        if decoys := self._get_control_value(controls, 'decoys_edit', 'text'):
            if decoys.strip(): command.extend(["-D", decoys.strip()])
        if source_port := self._get_control_value(controls, 'source_port_edit', 'text'):
            if source_port.strip(): command.extend(["-g", source_port.strip()])

        command.append(target_for_log)
        return command, target_for_log, None

    def _create_nuclei_tool(self):
        """Creates the UI for the Nuclei Web Scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.nuclei_controls = self._create_nuclei_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.nuclei_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.nuclei_output_console = QPlainTextEdit()
        self.nuclei_output_console.setReadOnly(True)
        self.nuclei_output_console.setFont(QFont("Courier New", 10))
        self.nuclei_output_console.setPlaceholderText("Nuclei output will be displayed here...")
        main_layout.addWidget(self.nuclei_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_nuclei_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_trufflehog_tool(self):
        """Creates the UI for the TruffleHog Secret Scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.trufflehog_controls = self._create_trufflehog_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.trufflehog_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.trufflehog_output_console = QPlainTextEdit()
        self.trufflehog_output_console.setReadOnly(True)
        self.trufflehog_output_console.setFont(QFont("Courier New", 10))
        self.trufflehog_output_console.setPlaceholderText("TruffleHog output will be displayed here...")
        main_layout.addWidget(self.trufflehog_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_trufflehog_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_rustscan_tool(self):
        """Creates the UI for the RustScan tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.rustscan_controls = self._create_rustscan_config_widget()
        layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.rustscan_controls['start_btn'])
        buttons_layout.addWidget(self.rustscan_controls['cancel_btn'])
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.rustscan_output = QPlainTextEdit()
        self.rustscan_output.setReadOnly(True)
        self.rustscan_output.setFont(QFont("Courier New", 10))
        self.rustscan_output.setPlaceholderText("RustScan and Nmap output will be displayed here...")
        layout.addWidget(self.rustscan_output, 1)

        self.rustscan_controls['start_btn'].clicked.connect(self.start_rustscan_scan)
        self.rustscan_controls['cancel_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_enum4linux_ng_tool(self):
        """Creates the UI for the enum4linux-ng tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.enum4linux_ng_controls = self._create_enum4linux_ng_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.enum4linux_ng_controls['start_btn'])
        buttons_layout.addWidget(self.enum4linux_ng_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.enum4linux_ng_output_console = QPlainTextEdit()
        self.enum4linux_ng_output_console.setReadOnly(True)
        self.enum4linux_ng_output_console.setFont(QFont("Courier New", 10))
        self.enum4linux_ng_output_console.setPlaceholderText("enum4linux-ng output will be displayed here...")
        main_layout.addWidget(self.enum4linux_ng_output_console, 1)

        self.enum4linux_ng_controls['start_btn'].clicked.connect(self.start_enum4linux_ng_scan)
        self.enum4linux_ng_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_sherlock_tool(self):
        """Creates the UI for the Sherlock username scanner."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.sherlock_controls = self._create_sherlock_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.sherlock_controls['start_btn'])
        buttons_layout.addWidget(self.sherlock_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.sherlock_output_console = QPlainTextEdit()
        self.sherlock_output_console.setReadOnly(True)
        self.sherlock_output_console.setFont(QFont("Courier New", 10))
        self.sherlock_output_console.setPlaceholderText("Sherlock output will be displayed here...")
        main_layout.addWidget(self.sherlock_output_console, 1)

        self.sherlock_controls['start_btn'].clicked.connect(self.start_sherlock_scan)
        self.sherlock_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_sherlock_config_widget(self):
        """Creates a reusable, self-contained widget with Sherlock's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['usernames_edit'] = QLineEdit()
        controls['usernames_edit'].setToolTip("One or more usernames to check, separated by spaces.")
        main_layout.addRow("Usernames:", controls['usernames_edit'])

        controls['timeout_edit'] = QLineEdit("60")
        controls['timeout_edit'].setToolTip("Timeout in seconds for each request.")
        main_layout.addRow("Timeout (--timeout):", controls['timeout_edit'])

        output_file_layout = QHBoxLayout()
        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("Optional: path to save text report...")
        output_file_layout.addWidget(controls['output_edit'])
        browse_output_btn = QPushButton("Browse...")
        browse_output_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_edit'], "Save Sherlock Report"))
        output_file_layout.addWidget(browse_output_btn)
        main_layout.addRow("Output File (-o):", output_file_layout)

        controls['csv_check'] = QCheckBox("Export as CSV")
        main_layout.addRow("--csv:", controls['csv_check'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Hunt Usernames")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_sherlock_command(self, controls):
        """Builds the Sherlock command list from a dictionary of controls or widgets."""
        command = ["sherlock", "--no-color"]

        usernames = self._get_control_value(controls, 'usernames_edit', 'text')
        if not usernames or not usernames.strip():
            return None, None, "At least one username is required."
        target_for_log = usernames.strip()

        if timeout := self._get_control_value(controls, 'timeout_edit', 'text'):
            if timeout.strip():
                command.extend(["--timeout", timeout.strip()])

        # This is for the user-specified output file.
        if output_file := self._get_control_value(controls, 'output_edit', 'text'):
             if output_file.strip():
                command.extend(["-o", output_file.strip()])

        command.extend(target_for_log.split())

        return command, target_for_log, None

    def start_sherlock_scan(self):
        """Starts the Sherlock scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.sherlock_controls
        command, target_for_log, error = self._build_sherlock_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "Sherlock Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        # Create a temporary directory for sherlock to save its files for parsing
        try:
            self.sherlock_temp_dir = tempfile.mkdtemp()
            # The command is modified to include the temp folder for CSV output
            command.extend(["-fo", self.sherlock_temp_dir, "--csv"])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary directory for Sherlock report: {e}")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.sherlock_output_console.clear()

        self.worker = WorkerThread(self._sherlock_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _sherlock_thread(self, command, usernames):
        """Worker thread for running the sherlock command."""
        q = self.tool_results_queue
        logging.info(f"Starting Sherlock with command: {' '.join(command)}")
        q.put(('sherlock_output', f"$ {' '.join(command)}\n\n"))
        csv_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.sherlock_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('sherlock_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('sherlock_output', line))

            process.stdout.close()
            process.wait()

            if not self.tool_stop_event.is_set():
                try:
                    # Find the first CSV file in the temp directory
                    for filename in os.listdir(self.sherlock_temp_dir):
                        if filename.endswith(".csv"):
                            with open(os.path.join(self.sherlock_temp_dir, filename), 'r', encoding='utf-8') as f:
                                csv_data = f.read()
                            break # Just read the first one for now
                    if csv_data:
                        q.put(('sherlock_results', csv_data, usernames))
                except Exception as e:
                    logging.error(f"Could not read Sherlock CSV report: {e}")
                finally:
                    shutil.rmtree(self.sherlock_temp_dir)
                    self.sherlock_temp_dir = None


        except FileNotFoundError:
            q.put(('error', 'Sherlock Error', "'sherlock' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Sherlock thread error: {e}", exc_info=True)
            q.put(('error', 'Sherlock Error', str(e)))
        finally:
            q.put(('tool_finished', 'sherlock_scan', usernames, csv_data))
            with self.thread_finish_lock:
                self.sherlock_process = None
            logging.info("Sherlock scan thread finished.")

    def _handle_sherlock_output(self, line):
        self.sherlock_output_console.insertPlainText(line)
        self.sherlock_output_console.verticalScrollBar().setValue(self.sherlock_output_console.verticalScrollBar().maximum())

    def _create_spiderfoot_tool(self):
        """Creates the UI for the Spiderfoot OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.spiderfoot_controls = self._create_spiderfoot_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.spiderfoot_controls['start_btn'])
        buttons_layout.addWidget(self.spiderfoot_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.spiderfoot_output_console = QPlainTextEdit()
        self.spiderfoot_output_console.setReadOnly(True)
        self.spiderfoot_output_console.setFont(QFont("Courier New", 10))
        self.spiderfoot_output_console.setPlaceholderText("Spiderfoot output will be displayed here...")
        main_layout.addWidget(self.spiderfoot_output_console, 1)

        self.spiderfoot_controls['start_btn'].clicked.connect(self.start_spiderfoot_scan)
        self.spiderfoot_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_spiderfoot_config_widget(self):
        """Creates a reusable, self-contained widget with Spiderfoot's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['target_edit'] = QLineEdit("example.com")
        main_layout.addRow("Target (-s):", controls['target_edit'])

        controls['types_edit'] = QLineEdit("EMAILADDR,DNS_MX,PHONE_NUMBER")
        controls['types_edit'].setToolTip("Comma-separated list of data types to collect (e.g., EMAILADDR, PHONE_NUMBER).")
        main_layout.addRow("Scan Types (-t):", controls['types_edit'])

        controls['modules_edit'] = QLineEdit()
        controls['modules_edit'].setPlaceholderText("Optional: e.g., sfp_dns,sfp_email")
        controls['modules_edit'].setToolTip("Comma-separated list of specific modules to run (-m).")
        main_layout.addRow("Modules (-m):", controls['modules_edit'])

        controls['silent_check'] = QCheckBox("Silent Output")
        controls['silent_check'].setToolTip("Only report errors (-q).")
        main_layout.addRow("-q:", controls['silent_check'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_spiderfoot_command(self, controls):
        """Builds the Spiderfoot command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("spiderfoot-cli", "spiderfoot-cli")
        if not tool_path:
             # Try python script fallback
             tool_path = self._get_tool_path("sf.py", os.path.join("spiderfoot", "sf.py"))
             if tool_path and tool_path.endswith(".py"):
                 command = ["python3", tool_path]
             elif tool_path:
                 command = [tool_path]
             else:
                 return None, None, "Spiderfoot not found."
        else:
             command = [tool_path]

        target = self._get_control_value(controls, 'target_edit', 'text')
        if not target or not target.strip():
            return None, None, "A target is required."
        target = target.strip()
        target_for_log = target

        command.extend(["-s", target])

        if types := self._get_control_value(controls, 'types_edit', 'text'):
            if types.strip():
                command.extend(["-t", types.strip()])
        if modules := self._get_control_value(controls, 'modules_edit', 'text'):
            if modules.strip():
                command.extend(["-m", modules.strip()])
        if self._get_control_value(controls, 'silent_check', 'check'):
            command.append("-q")

        command.append("-n") # Disable history logging for non-interactive use

        return command, target_for_log, None

    def start_spiderfoot_scan(self):
        """Starts the Spiderfoot scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.spiderfoot_controls
        command, target_for_log, error = self._build_spiderfoot_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.spiderfoot_output_console.clear()

        self.worker = WorkerThread(self._spiderfoot_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _spiderfoot_thread(self, command, target):
        """Worker thread for running the spiderfoot-cli command."""
        q = self.tool_results_queue
        logging.info(f"Starting Spiderfoot with command: {' '.join(command)}")
        q.put(('spiderfoot_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.spiderfoot_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('spiderfoot_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('spiderfoot_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Spiderfoot Error', "'spiderfoot-cli' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Spiderfoot thread error: {e}", exc_info=True)
            q.put(('error', 'Spiderfoot Error', str(e)))
        finally:
            q.put(('tool_finished', 'spiderfoot_scan', target, "".join(full_output)))
            with self.thread_finish_lock:
                self.spiderfoot_process = None
            logging.info("Spiderfoot scan thread finished.")

    def _handle_spiderfoot_output(self, line):
        self.spiderfoot_output_console.insertPlainText(line)
        self.spiderfoot_output_console.verticalScrollBar().setValue(self.spiderfoot_output_console.verticalScrollBar().maximum())

    def _create_zap_scanner_tool(self):
        """Creates the UI for the OWASP ZAP scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>OWASP ZAP - Web Application Scanner</b></font>
        <p>This tool automates OWASP ZAP to perform security scanning on web applications using a generated configuration file.</p>
        <p><b>Usage:</b> Configure the target and scan options below. ZAP will be run in headless mode, and a report will be generated in the specified location.</p>
        """)
        instructions.setFixedHeight(100)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.zap_controls = self._create_zap_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.zap_controls['start_btn'])
        buttons_layout.addWidget(self.zap_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.zap_output_console = QPlainTextEdit()
        self.zap_output_console.setReadOnly(True)
        self.zap_output_console.setFont(QFont("Courier New", 10))
        self.zap_output_console.setPlaceholderText("OWASP ZAP output will be displayed here...")
        main_layout.addWidget(self.zap_output_console, 1)

        self.zap_controls['start_btn'].clicked.connect(self.start_zap_scan)
        self.zap_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_zap_config_widget(self):
        """Creates a reusable, self-contained widget for ZAP's configuration."""
        widget = QGroupBox("ZAP Automation Configuration")
        layout = QFormLayout(widget)
        controls = {}

        controls['target_url_edit'] = QLineEdit("https://public-firing-range.appspot.com")
        layout.addRow("Target URL:", controls['target_url_edit'])

        controls['context_name_edit'] = QLineEdit("Default Context")
        layout.addRow("Context Name:", controls['context_name_edit'])

        scan_box = QGroupBox("Scan Types")
        scan_layout = QVBoxLayout(scan_box)
        controls['spider_check'] = QCheckBox("Traditional Spider"); controls['spider_check'].setChecked(True)
        controls['spider_ajax_check'] = QCheckBox("AJAX Spider")
        controls['active_scan_check'] = QCheckBox("Active Scan"); controls['active_scan_check'].setChecked(True)
        scan_layout.addWidget(controls['spider_check'])
        scan_layout.addWidget(controls['spider_ajax_check'])
        scan_layout.addWidget(controls['active_scan_check'])
        layout.addRow(scan_box)

        report_box = QGroupBox("Report Generation")
        report_layout = QFormLayout(report_box)
        controls['report_format_combo'] = QComboBox()
        controls['report_format_combo'].addItems(["html", "md", "xml", "json"])
        report_layout.addRow("Report Format:", controls['report_format_combo'])

        report_file_layout = QHBoxLayout()
        controls['report_file_edit'] = QLineEdit()
        controls['report_file_edit'].setPlaceholderText("Path to save report (required)...")
        report_file_layout.addWidget(controls['report_file_edit'])
        browse_report_btn = QPushButton("Browse...")
        browse_report_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['report_file_edit'], "Save ZAP Report"))
        report_file_layout.addWidget(browse_report_btn)
        report_layout.addRow("Report File:", report_file_layout)
        layout.addRow(report_box)

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("shield.svg")), " Start ZAP Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _generate_zap_yaml(self, controls):
        """Generates the ZAP Automation Framework YAML file from UI controls."""
        target_url = self._get_control_value(controls, 'target_url_edit', 'text')
        context_name = self._get_control_value(controls, 'context_name_edit', 'text')
        report_format = self._get_control_value(controls, 'report_format_combo', 'combo')
        report_file = self._get_control_value(controls, 'report_file_edit', 'text')

        if not target_url or not report_file:
            return None, "Target URL and Report File are required."

        # Base YAML structure
        yaml_content = {
            'env': {
                'contexts': [{
                    'name': context_name,
                    'urls': [target_url]
                }],
                'parameters': {
                    'failOnError': True,
                    'failOnWarning': False,
                    'progressToStdout': True
                }
            },
            'jobs': []
        }

        # Add jobs based on checkboxes
        if self._get_control_value(controls, 'spider_check', 'check'):
            yaml_content['jobs'].append({'type': 'spider', 'name': 'spider', 'parameters': {'context': context_name}})
        if self._get_control_value(controls, 'spider_ajax_check', 'check'):
            yaml_content['jobs'].append({'type': 'spiderAjax', 'name': 'spiderAjax', 'parameters': {'context': context_name, 'url': target_url}})
        if self._get_control_value(controls, 'active_scan_check', 'check'):
            yaml_content['jobs'].append({'type': 'activeScan', 'name': 'activeScan', 'parameters': {'context': context_name}})

        # Add report job
        yaml_content['jobs'].append({
            'type': 'report',
            'name': 'report',
            'parameters': {
                'template': report_format,
                'reportDir': os.path.dirname(report_file),
                'reportFile': os.path.basename(report_file)
            }
        })

        try:
            # Using a temporary file for the YAML config
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".yaml", encoding='utf-8') as tmp_yaml:
                import yaml # Import locally to avoid startup dependency
                yaml.dump(yaml_content, tmp_yaml, sort_keys=False)
                return tmp_yaml.name, None
        except ImportError:
            return None, "PyYAML is not installed. Please run 'pip install pyyaml'."
        except Exception as e:
            return None, f"Failed to generate YAML file: {e}"

    def start_zap_scan(self):
        """Starts the ZAP scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        zap_path = self._find_zap_executable()
        if not zap_path:
            QMessageBox.critical(self, "ZAP Error", "OWASP ZAP executable was not found. Install ZAP or place its launcher in PATH.")
            return

        controls = self.zap_controls
        yaml_path, error = self._generate_zap_yaml(controls)

        if error:
            QMessageBox.critical(self, "ZAP Config Error", error)
            return

        target_url = self._get_control_value(controls, 'target_url_edit', 'text')
        command = [zap_path, "-cmd", "-autorun", yaml_path]

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.zap_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'zap_scan', target_url, self.zap_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _handle_zap_output(self, line):
        self.zap_output_console.insertPlainText(line)
        self.zap_output_console.verticalScrollBar().setValue(self.zap_output_console.verticalScrollBar().maximum())

    def _create_metasploit_tool(self):
        """Creates the UI for the Metasploit auxiliary scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Metasploit Auxiliary Scanner</b></font>
        <p>This tool runs Metasploit Framework auxiliary modules non-interactively. Specify a module and its options to start a scan.</p>
        <p><b>Note:</b> This requires Metasploit Framework to be installed and `msfconsole` to be in your system's PATH.</p>
        """)
        instructions.setFixedHeight(100)
        main_layout.addWidget(instructions)

        config_widget, self.metasploit_controls = self._create_metasploit_config_widget()
        main_layout.addWidget(config_widget)

        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.metasploit_controls['start_btn'])
        buttons_layout.addWidget(self.metasploit_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        self.metasploit_output_console = QPlainTextEdit()
        self.metasploit_output_console.setReadOnly(True)
        self.metasploit_output_console.setFont(QFont("Courier New", 10))
        self.metasploit_output_console.setPlaceholderText("Metasploit output will be displayed here...")
        main_layout.addWidget(self.metasploit_output_console, 1)

        self.metasploit_controls['start_btn'].clicked.connect(self.start_metasploit_scan)
        self.metasploit_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_metasploit_config_widget(self):
        """Creates a reusable, self-contained widget for Metasploit's configuration."""
        widget = QGroupBox("Metasploit Scan Configuration")
        layout = QFormLayout(widget)
        controls = {}

        controls['module_edit'] = QLineEdit("auxiliary/scanner/portscan/tcp")
        controls['module_edit'].setToolTip("The full path of the Metasploit auxiliary module to use.")
        layout.addRow("Module:", controls['module_edit'])

        controls['rhosts_edit'] = QLineEdit()
        controls['rhosts_edit'].setToolTip("The target host(s) or CIDR range for the scan (RHOSTS).")
        layout.addRow("Target Hosts (RHOSTS):", controls['rhosts_edit'])

        controls['rport_edit'] = QLineEdit("80")
        controls['rport_edit'].setToolTip("The target port(s) for the scan (RPORT/RPORTs).")
        layout.addRow("Target Port (RPORT):", controls['rport_edit'])

        controls['threads_edit'] = QLineEdit("10")
        controls['threads_edit'].setToolTip("Number of concurrent threads to use.")
        layout.addRow("Threads:", controls['threads_edit'])

        controls['extra_opts_edit'] = QLineEdit()
        controls['extra_opts_edit'].setPlaceholderText("e.g., SET Proxies http:127.0.0.1:8080, SET ShowProgress false")
        controls['extra_opts_edit'].setToolTip("Additional options to set, separated by commas.\nFormat: 'SET key value, SET key2 value2'")
        layout.addRow("Extra Options:", controls['extra_opts_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("tool.svg")), " Start Metasploit Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_metasploit_command(self, controls):
        """Builds the msfconsole command list from a dictionary of controls or widgets."""
        module = self._get_control_value(controls, 'module_edit', 'text')
        rhosts = self._get_control_value(controls, 'rhosts_edit', 'text')

        if not module or not rhosts:
            return None, None, "Module and Target Hosts (RHOSTS) are required."

        tool_path = self._get_tool_path("msfconsole", "msfconsole")
        if not tool_path:
             return None, None, "Metasploit Framework (msfconsole) not found."

        commands = f"use {module.strip()};"
        commands += f"set RHOSTS {rhosts.strip()};"
        if rport := self._get_control_value(controls, 'rport_edit', 'text'):
            if rport.strip():
                commands += f"set RPORT {rport.strip()};"
        if threads := self._get_control_value(controls, 'threads_edit', 'text'):
            if threads.strip():
                commands += f"set THREADS {threads.strip()};"

        if extra_opts := self._get_control_value(controls, 'extra_opts_edit', 'text'):
            if extra_opts.strip():
                # Split by comma, then add semicolon
                opts_list = [opt.strip() for opt in extra_opts.split(',')]
                commands += "".join([f"{opt};" for opt in opts_list])

        commands += "run; exit"

        msf_command = ["msfconsole", "-q", "-x", commands]
        return msf_command, rhosts.strip(), None

    def start_metasploit_scan(self):
        """Starts the Metasploit scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        if not shutil.which("msfconsole"):
            QMessageBox.critical(self, "Metasploit Error", "'msfconsole' command not found. Please ensure it is installed and in your system's PATH.")
            return

        controls = self.metasploit_controls
        command, target_for_log, error = self._build_metasploit_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.metasploit_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'metasploit_scan', target_for_log, self.metasploit_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _handle_metasploit_output(self, line):
        self.metasploit_output_console.insertPlainText(line)
        self.metasploit_output_console.verticalScrollBar().setValue(self.metasploit_output_console.verticalScrollBar().maximum())

    def _create_hydra_tool(self):
        """Creates the UI for the Hydra network logon cracker."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.hydra_controls = self._create_hydra_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.hydra_controls['start_btn'])
        buttons_layout.addWidget(self.hydra_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.hydra_output_console = QPlainTextEdit()
        self.hydra_output_console.setReadOnly(True)
        self.hydra_output_console.setFont(QFont("Courier New", 10))
        self.hydra_output_console.setPlaceholderText("Hydra output will be displayed here...")
        main_layout.addWidget(self.hydra_output_console, 1)

        self.hydra_controls['start_btn'].clicked.connect(self.start_hydra_attack)
        self.hydra_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_masscan_config_widget(self):
        """Creates a reusable, self-contained widget with Masscan's configuration options."""
        widget = QGroupBox("Scan Options")
        controls_layout = QFormLayout(widget)
        controls = {}

        controls['target_edit'] = QLineEdit("0.0.0.0/0")
        controls['target_edit'].setToolTip("Enter target IP ranges (e.g., 10.0.0.0/8, 192.168.0.1-192.168.0.254).")
        controls_layout.addRow("Target(s):", controls['target_edit'])

        ports_layout = QHBoxLayout()
        controls['ports_edit'] = QLineEdit("0-65535")
        controls['ports_edit'].setToolTip("Specify ports to scan (e.g., 80,443, 0-65535).")
        ports_layout.addWidget(controls['ports_edit'])
        common_ports_btn = QPushButton("Common Ports")
        common_ports_btn.clicked.connect(lambda: controls['ports_edit'].setText("21,22,23,25,53,80,110,111,135,139,143,443,445,993,995,1723,3306,3389,5900,8080"))
        ports_layout.addWidget(common_ports_btn)
        controls_layout.addRow("Ports:", ports_layout)

        controls['rate_edit'] = QLineEdit("1000")
        controls['rate_edit'].setToolTip("Set the transmission rate in packets/second.")
        controls_layout.addRow("Rate (--rate):", controls['rate_edit'])

        outfile_layout = QHBoxLayout()
        controls['outfile_edit'] = QLineEdit()
        controls['outfile_edit'].setPlaceholderText("Optional: path to save report...")
        outfile_layout.addWidget(controls['outfile_edit'])
        browse_out_btn = QPushButton("Browse...")
        browse_out_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['outfile_edit'], "Save Masscan Report", "JSON files (*.json);;All Files (*)"))
        outfile_layout.addWidget(browse_out_btn)
        controls_layout.addRow("Output File (-oJ):", outfile_layout)

        controls['extra_opts_edit'] = QLineEdit()
        controls['extra_opts_edit'].setToolTip("Enter any additional, space-separated Masscan flags here.")
        controls_layout.addRow("Additional Options:", controls['extra_opts_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Masscan")
        controls['stop_btn'] = QPushButton("Stop Masscan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _create_masscan_tool(self):
        """Creates the UI for the Masscan tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Masscan: The Mass IP Port Scanner</b></font>
        <p>This is an Internet-scale port scanner, capable of scanning the entire Internet in under 5 minutes.</p>
        <p><b>WARNING:</b> Scanning at high rates can cause network disruption and may be detected by network administrators. Use responsibly.</p>
        """)
        instructions.setFixedHeight(100)
        main_layout.addWidget(instructions)

        # --- Controls ---
        controls_frame, self.masscan_controls = self._create_masscan_config_widget()
        main_layout.addWidget(controls_frame)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.masscan_controls['start_btn'])
        buttons_layout.addWidget(self.masscan_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.masscan_output_console = QPlainTextEdit()
        self.masscan_output_console.setReadOnly(True)
        self.masscan_output_console.setFont(QFont("Courier New", 10))
        self.masscan_output_console.setPlaceholderText("Masscan output will be displayed here... (Note: Masscan primarily outputs to stderr)")
        main_layout.addWidget(self.masscan_output_console, 1)

        self.masscan_controls['start_btn'].clicked.connect(self.start_masscan_scan)
        self.masscan_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def start_masscan_scan(self, sudo_password=None):
        """Starts the Masscan worker thread, prompting for sudo if necessary."""
        controls = self.masscan_controls
        if not shutil.which("masscan"):
            QMessageBox.critical(self, "Masscan Error", "'masscan' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running and not sudo_password:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = controls['target_edit'].text().strip()
        ports = controls['ports_edit'].text().strip()
        rate = controls['rate_edit'].text().strip()

        if not target or not ports or not rate:
            QMessageBox.critical(self, "Input Error", "Target, Ports, and Rate are required.")
            return

        resolved_target = self._resolve_targets_string(target)

        command = ["masscan", resolved_target, "-p", ports, "--rate", rate]

        if outfile := controls['outfile_edit'].text().strip():
            command.extend(["-oJ", outfile])

        if extra_opts := controls['extra_opts_edit'].text().strip():
            command.extend(extra_opts.split())

        # Masscan always needs root on non-Windows systems.
        if not sudo_password and sys.platform != "win32":
            # Add a specific check to the sudo prompter to re-enable the right buttons on cancel
            if 'masscan_scan' not in self.sudo_cancel_handlers:
                self.sudo_cancel_handlers['masscan_scan'] = lambda: (
                    controls['start_btn'].setEnabled(True),
                    controls['stop_btn'].setEnabled(False)
                )
            self._run_command_with_sudo_prompt(command, self.start_masscan_scan, 'masscan_scan')
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.masscan_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'masscan_scan', target, self.masscan_output_console, sudo_password
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _masscan_thread(self, command, target):
        """Worker thread for running the masscan command."""
        q = self.tool_results_queue
        logging.info(f"Starting Masscan with command: {' '.join(command)}")
        q.put(('masscan_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            # Masscan outputs progress to stderr, so we need to capture both
            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.masscan_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('masscan_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('masscan_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Masscan Error', "'masscan' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Masscan thread error: {e}", exc_info=True)
            q.put(('error', 'Masscan Error', str(e)))
        finally:
            q.put(('tool_finished', 'masscan_scan', target, "".join(full_output)))
            with self.thread_finish_lock:
                self.masscan_process = None
            logging.info("Masscan scan thread finished.")

    def _gobuster_thread(self, command, url):
        """Worker thread for running the gobuster command."""
        q = self.tool_results_queue
        logging.info(f"Starting Gobuster with command: {' '.join(command)}")
        q.put(('gobuster_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.gobuster_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('gobuster_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('gobuster_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Gobuster Error', "'gobuster' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Gobuster thread error: {e}", exc_info=True)
            q.put(('error', 'Gobuster Error', str(e)))
        finally:
            q.put(('tool_finished', 'gobuster_scan', url, "".join(full_output)))
            with self.thread_finish_lock:
                self.gobuster_process = None
            logging.info("Gobuster scan thread finished.")

    def _nikto_thread(self, command, target):
        """Worker thread for running the nikto command."""
        q = self.tool_results_queue
        logging.info(f"Starting Nikto with command: {' '.join(command)}")
        q.put(('nikto_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.nikto_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('nikto_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('nikto_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Nikto Error', "'nikto' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Nikto thread error: {e}", exc_info=True)
            q.put(('error', 'Nikto Error', str(e)))
        finally:
            q.put(('tool_finished', 'nikto_scan', target, "".join(full_output)))
            with self.thread_finish_lock:
                self.nikto_process = None
            logging.info("Nikto scan thread finished.")

    def _create_traceroute_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        self.trace_tree = QTreeWidget(); self.trace_tree.setColumnCount(4); self.trace_tree.setHeaderLabels(["Hop", "IP Address", "Host Name", "Time (ms)"])

        controls = QHBoxLayout()
        controls.addWidget(QLabel("Target:")); self.trace_target = QLineEdit("google.com"); controls.addWidget(self.trace_target)
        self.trace_button = QPushButton("Trace"); controls.addWidget(self.trace_button)
        self.trace_cancel_button = QPushButton("Cancel"); self.trace_cancel_button.setEnabled(False); controls.addWidget(self.trace_cancel_button)
        self.trace_status = QLabel(""); controls.addWidget(self.trace_status); controls.addStretch()

        layout.addLayout(controls)
        layout.addWidget(self.trace_tree)
        layout.addWidget(self._create_export_button(self.trace_tree))
        self.trace_button.clicked.connect(self.start_traceroute)
        self.trace_cancel_button.clicked.connect(self.cancel_tool)
        return widget

    def _update_tcp_scan_options_visibility(self, checked):
        """Shows or hides the TCP scan mode dropdown based on protocol selection."""
        is_tcp_selected = self.scan_proto_tcp_radio.isChecked() or self.scan_proto_both_radio.isChecked()
        self.tcp_scan_type_label.setVisible(is_tcp_selected)
        self.tcp_scan_type_combo.setVisible(is_tcp_selected)

    def _create_port_scanner_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget); controls = QFrame(); clayout = QVBoxLayout(controls)
        row1 = QHBoxLayout(); row1.addWidget(QLabel("Target:")); self.scan_target = QLineEdit("127.0.0.1"); self.scan_target.setToolTip("The IP address of the target machine.")
        row1.addWidget(self.scan_target)
        clayout.addLayout(row1)
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Ports:")); self.scan_ports = QLineEdit("22,80,443"); self.scan_ports.setToolTip("A comma-separated list of ports or port ranges (e.g., 22,80,100-200).")
        row2.addWidget(self.scan_ports)
        all_ports_btn = QPushButton("All"); all_ports_btn.setToolTip("Set the port range to all 65535 ports.")
        all_ports_btn.clicked.connect(lambda: self.scan_ports.setText("1-65535")); row2.addWidget(all_ports_btn)
        clayout.addLayout(row2)

        # Row 3: Protocol Type Radio Buttons
        row3 = QHBoxLayout()
        row3.addWidget(QLabel("Protocol:"))
        self.scan_proto_tcp_radio = QRadioButton("TCP"); self.scan_proto_tcp_radio.setChecked(True)
        self.scan_proto_udp_radio = QRadioButton("UDP")
        self.scan_proto_both_radio = QRadioButton("Both")
        self.scan_proto_group = QButtonGroup(self)
        self.scan_proto_group.addButton(self.scan_proto_tcp_radio)
        self.scan_proto_group.addButton(self.scan_proto_udp_radio)
        self.scan_proto_group.addButton(self.scan_proto_both_radio)
        row3.addWidget(self.scan_proto_tcp_radio)
        row3.addWidget(self.scan_proto_udp_radio)
        row3.addWidget(self.scan_proto_both_radio)
        row3.addStretch()
        clayout.addLayout(row3)

        # Row 4: Advanced TCP Scan Options
        row4 = QHBoxLayout()
        self.tcp_scan_type_label = QLabel("TCP Scan Mode:")
        row4.addWidget(self.tcp_scan_type_label)
        self.tcp_scan_type_combo = QComboBox()
        self.tcp_scan_type_combo.addItems(["SYN Scan", "FIN Scan", "Xmas Scan", "Null Scan", "ACK Scan"])
        self.tcp_scan_type_combo.setToolTip("Select the type of TCP scan to perform for firewall evasion.")
        row4.addWidget(self.tcp_scan_type_combo)
        row4.addStretch()
        self.scan_frag_check = QCheckBox("Use Fragments"); self.scan_frag_check.setToolTip("Send fragmented packets to potentially evade simple firewalls.")
        row4.addWidget(self.scan_frag_check)
        clayout.addLayout(row4)

        # Connect signals for UI logic
        self.scan_proto_tcp_radio.toggled.connect(self._update_tcp_scan_options_visibility)
        self.scan_proto_udp_radio.toggled.connect(self._update_tcp_scan_options_visibility)
        self.scan_proto_both_radio.toggled.connect(self._update_tcp_scan_options_visibility)

        scan_buttons_layout = QHBoxLayout()
        self.scan_button = QPushButton("Scan"); self.scan_button.setToolTip("Start the port scan.")
        scan_buttons_layout.addWidget(self.scan_button)
        self.scan_cancel_button = QPushButton("Cancel"); self.scan_cancel_button.setEnabled(False); self.scan_cancel_button.setToolTip("Stop the current scan.")
        scan_buttons_layout.addWidget(self.scan_cancel_button)
        clayout.addLayout(scan_buttons_layout)
        self.scan_status = QLabel(""); clayout.addWidget(self.scan_status)
        layout.addWidget(controls)
        self.scan_tree = QTreeWidget(); self.scan_tree.setColumnCount(3); self.scan_tree.setHeaderLabels(["Port", "State", "Service"])
        layout.addWidget(self.scan_tree)
        layout.addWidget(self._create_export_button(self.scan_tree))
        self.scan_button.clicked.connect(self.start_port_scan)
        self.scan_cancel_button.clicked.connect(self.cancel_tool)

        self._update_tcp_scan_options_visibility(True) # Initial state
        return widget

    def _create_arp_scan_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget)
        self.arp_tree=QTreeWidget(); self.arp_tree.setColumnCount(3); self.arp_tree.setHeaderLabels(["IP Address","MAC Address", "Status"])

        controls = QHBoxLayout()
        controls.addWidget(QLabel("Target Network:")); self.arp_target=QLineEdit("192.168.1.0/24"); controls.addWidget(self.arp_target)
        self.arp_scan_button=QPushButton("Scan"); controls.addWidget(self.arp_scan_button)
        self.arp_status=QLabel(""); controls.addWidget(self.arp_status); controls.addStretch()

        layout.addLayout(controls)
        layout.addWidget(self.arp_tree)
        layout.addWidget(self._create_export_button(self.arp_tree))
        self.arp_scan_button.clicked.connect(self.start_arp_scan)
        return widget

    def _create_ping_sweep_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Main controls
        controls_frame = QFrame()
        controls_frame.setFrameShape(QFrame.Shape.StyledPanel)
        controls = QVBoxLayout(controls_frame)

        row1 = QHBoxLayout()
        row1.addWidget(QLabel("Target Network (CIDR):"))
        self.ps_target_edit = QLineEdit("192.168.1.0/24")
        self.ps_target_edit.setToolTip("The target network range in CIDR notation (e.g., 192.168.1.0/24).")
        row1.addWidget(self.ps_target_edit)
        controls.addLayout(row1)

        # Options Box
        options_layout = QFormLayout()
        options_layout.setContentsMargins(5, 10, 5, 10)

        # Probe Type
        self.ps_probe_type_combo = QComboBox()
        self.ps_probe_type_combo.addItems(["ICMP Echo", "TCP SYN", "TCP ACK", "UDP Probe"])
        options_layout.addRow("Probe Type:", self.ps_probe_type_combo)

        # Ports
        self.ps_ports_label = QLabel("Target Port(s):")
        self.ps_ports_edit = QLineEdit("80,443,8080")
        self.ps_ports_edit.setToolTip("Comma-separated list of ports for TCP/UDP probes.")
        options_layout.addRow(self.ps_ports_label, self.ps_ports_edit)

        # Timeout
        self.ps_timeout_edit = QLineEdit("1")
        self.ps_timeout_edit.setToolTip("Timeout in seconds for each probe.")
        options_layout.addRow("Timeout (s):", self.ps_timeout_edit)

        # Threads
        self.ps_threads_edit = QLineEdit("10")
        self.ps_threads_edit.setToolTip("Number of concurrent threads to use for scanning.")
        options_layout.addRow("Threads:", self.ps_threads_edit)
        controls.addLayout(options_layout)
        layout.addWidget(controls_frame)

        buttons_layout = QHBoxLayout()
        self.ps_start_button = QPushButton("Start Sweep")
        buttons_layout.addWidget(self.ps_start_button)
        self.ps_cancel_button = QPushButton("Cancel")
        self.ps_cancel_button.setEnabled(False)
        buttons_layout.addWidget(self.ps_cancel_button)
        layout.addLayout(buttons_layout)

        self.ps_status_label = QLabel("Status: Idle")
        layout.addWidget(self.ps_status_label)

        self.ps_tree = QTreeWidget()
        self.ps_tree.setColumnCount(2)
        self.ps_tree.setHeaderLabels(["IP Address", "Status"])
        layout.addWidget(self.ps_tree)
        layout.addWidget(self._create_export_button(self.ps_tree))

        # --- Connections and Logic ---
        def toggle_ports_visibility(text):
            is_tcp_or_udp = "TCP" in text or "UDP" in text
            self.ps_ports_label.setVisible(is_tcp_or_udp)
            self.ps_ports_edit.setVisible(is_tcp_or_udp)

        self.ps_probe_type_combo.currentTextChanged.connect(toggle_ports_visibility)
        # Set initial state
        is_tcp_or_udp_initial = "TCP" in self.ps_probe_type_combo.currentText() or "UDP" in self.ps_probe_type_combo.currentText()
        self.ps_ports_label.setVisible(is_tcp_or_udp_initial)
        self.ps_ports_edit.setVisible(is_tcp_or_udp_initial)


        self.ps_start_button.clicked.connect(self.start_ping_sweep)
        self.ps_cancel_button.clicked.connect(self.cancel_tool)

        return widget

    def _create_osint_tab(self):
        """Creates the tab container for OSINT tools."""
        osint_tabs = QTabWidget()
        osint_tabs.addTab(self._create_theharvester_tool(), "theHarvester")
        osint_tabs.addTab(self._create_sherlock_tool(), "Sherlock")
        osint_tabs.addTab(self._create_spiderfoot_tool(), "Spiderfoot")
        osint_tabs.addTab(self._create_magma_tool(), "Magma OSINT")
        osint_tabs.addTab(self._create_photon_tool(), "Photon")
        osint_tabs.addTab(self._create_phoneinfoga_tool(), "PhoneInfoga")
        osint_tabs.addTab(self._create_metagoofil_tool(), "Metagoofil")
        osint_tabs.addTab(self._create_social_analyzer_tool(), "Social Analyzer")
        osint_tabs.addTab(self._create_infoga_tool(), "Infoga")
        osint_tabs.addTab(self._create_maigret_tool(), "Maigret")
        osint_tabs.addTab(self._create_holehe_tool(), "Holehe")
        osint_tabs.addTab(self._create_ghunt_tool(), "GHunt")
        osint_tabs.addTab(self._create_shodan_tool(), "Shodan")
        osint_tabs.addTab(self._create_recon_ng_tool(), "Recon-ng")
        return osint_tabs

    def _create_shodan_tool(self):
        """Creates the UI for the Shodan OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Shodan - The search engine for everything on the internet</b></font>
        <p>Shodan is a search engine that lets the user find specific types of computers connected to the internet using a variety of filters.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.shodan_controls = self._create_shodan_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.shodan_controls['start_btn'])
        buttons_layout.addWidget(self.shodan_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.shodan_output_console = QPlainTextEdit()
        self.shodan_output_console.setReadOnly(True)
        self.shodan_output_console.setFont(QFont("Courier New", 10))
        self.shodan_output_console.setPlaceholderText("Shodan output will be displayed here...")
        main_layout.addWidget(self.shodan_output_console, 1)

        self.shodan_controls['start_btn'].clicked.connect(self.start_shodan_scan)
        self.shodan_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_shodan_config_widget(self):
        """Creates a reusable, self-contained widget for Shodan's configuration options."""
        widget = QGroupBox("Shodan Options")
        main_layout = QVBoxLayout(widget)
        controls = {}

        # API Key (common to all tabs)
        api_layout = QHBoxLayout()
        api_layout.addWidget(QLabel("API Key:"))
        controls['api_key_edit'] = QLineEdit()
        controls['api_key_edit'].setPlaceholderText("Enter your Shodan API key")
        api_layout.addWidget(controls['api_key_edit'])
        main_layout.addLayout(api_layout)

        shodan_tabs = QTabWidget()
        main_layout.addWidget(shodan_tabs)

        # --- Search Tab ---
        search_tab = QWidget()
        search_layout = QFormLayout(search_tab)
        controls['search_query_edit'] = QLineEdit()
        search_layout.addRow("Search Query:", controls['search_query_edit'])
        controls['search_limit_edit'] = QLineEdit("100")
        search_layout.addRow("Limit:", controls['search_limit_edit'])
        controls['search_fields_edit'] = QLineEdit("ip_str,port,org,hostnames")
        search_layout.addRow("Fields:", controls['search_fields_edit'])

        shodan_tabs.addTab(search_tab, "Search")
        controls['search_tab'] = search_tab

        # --- Host Tab ---
        host_tab = QWidget()
        host_layout = QFormLayout(host_tab)
        controls['host_ip_edit'] = QLineEdit()
        host_layout.addRow("Host IP:", controls['host_ip_edit'])
        shodan_tabs.addTab(host_tab, "Host")
        controls['host_tab'] = host_tab

        # --- Count Tab ---
        count_tab = QWidget()
        count_layout = QFormLayout(count_tab)
        controls['count_query_edit'] = QLineEdit()
        count_layout.addRow("Count Query:", controls['count_query_edit'])
        shodan_tabs.addTab(count_tab, "Count")
        controls['count_tab'] = count_tab

        # --- Domain Tab ---
        domain_tab = QWidget()
        domain_layout = QFormLayout(domain_tab)
        controls['domain_edit'] = QLineEdit()
        domain_layout.addRow("Domain:", controls['domain_edit'])
        shodan_tabs.addTab(domain_tab, "Domain")
        controls['domain_tab'] = domain_tab

        # --- Honeyscore Tab ---
        honeyscore_tab = QWidget()
        honeyscore_layout = QFormLayout(honeyscore_tab)
        controls['honeyscore_ip_edit'] = QLineEdit()
        honeyscore_layout.addRow("IP Address:", controls['honeyscore_ip_edit'])
        shodan_tabs.addTab(honeyscore_tab, "Honeyscore")
        controls['honeyscore_tab'] = honeyscore_tab

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Run Command")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        # Add the tab widget to controls to access it later
        controls['tabs'] = shodan_tabs

        return widget, controls

    def _build_shodan_command(self, controls):
        """Builds the Shodan command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("shodan", "shodan")
        if not tool_path:
             return None, None, "Shodan CLI not found.", None
        command = [tool_path]

        api_key = self._get_control_value(controls, 'api_key_edit', 'text')
        # We don't need to add the key to the command, as we will use the SHODAN_API_KEY env var
        # But we need to check it exists
        if not api_key or not api_key.strip():
            return None, None, "An API key is required.", None

        # Determine active tab
        active_tab = controls['tabs'].currentWidget()
        target_for_log = ""

        if active_tab == controls['search_tab']:
            command.append("search")
            query = self._get_control_value(controls, 'search_query_edit', 'text')
            if not query: return None, None, "A search query is required.", None
            target_for_log = query
            command.append(query)
            if limit := self._get_control_value(controls, 'search_limit_edit', 'text'):
                command.extend(["--limit", limit])
            if fields := self._get_control_value(controls, 'search_fields_edit', 'text'):
                command.extend(["--fields", fields])

        elif active_tab == controls['host_tab']:
            command.append("host")
            ip = self._get_control_value(controls, 'host_ip_edit', 'text')
            if not ip: return None, None, "A host IP is required.", None
            target_for_log = ip
            command.append(ip)

        elif active_tab == controls['count_tab']:
            command.append("count")
            query = self._get_control_value(controls, 'count_query_edit', 'text')
            if not query: return None, None, "A count query is required.", None
            target_for_log = query
            command.append(query)

        elif active_tab == controls['domain_tab']:
            command.append("domain")
            domain = self._get_control_value(controls, 'domain_edit', 'text')
            if not domain: return None, None, "A domain is required.", None
            target_for_log = domain
            command.append(domain)

        elif active_tab == controls['honeyscore_tab']:
            command.append("honeyscore")
            ip = self._get_control_value(controls, 'honeyscore_ip_edit', 'text')
            if not ip: return None, None, "An IP address is required for honeyscore.", None
            target_for_log = ip
            command.append(ip)
        else:
            return None, None, "Could not determine active Shodan command.", None

        return command, target_for_log, None, {'SHODAN_API_KEY': api_key.strip()}

    def start_shodan_scan(self):
        """Starts the Shodan scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.shodan_controls
        command, target_for_log, error, env = self._build_shodan_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which("shodan"):
            QMessageBox.critical(self, "Shodan Error", "'shodan' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.shodan_output_console.clear()

        # The generic command executor doesn't support env vars, so I need a custom thread call
        self.worker = WorkerThread(self._execute_shodan_command_thread, args=(
            command, 'shodan_scan', target_for_log, self.shodan_output_console, env
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _execute_shodan_command_thread(self, command, tool_name, target, output_widget, env):
        q = self.tool_results_queue
        logging.info(f"Starting {tool_name} with command: {' '.join(command)}")
        q.put((f'{tool_name}_output', f"$ {' '.join(command)}\n\n"))
        full_output = []
        success = False
        try:
            # Create a copy of the current environment and update it with the API key
            process_env = os.environ.copy()
            if env:
                process_env.update(env)

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                                     text=True, bufsize=1, encoding='utf-8', errors='replace', env=process_env)
            with self.thread_finish_lock:
                setattr(self, f'{tool_name}_process', process)

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put((f'{tool_name}_output', "\n\n--- Canceled By User ---\n"))
                    break
                q.put((f'{tool_name}_output', line))
                full_output.append(line)

            process.stdout.close()
            retcode = process.wait()
            success = (retcode == 0)
        except FileNotFoundError:
            q.put(('error', f'{tool_name.capitalize()} Error', f"'{command[0]}' command not found."))
        except Exception as e:
            logging.error(f"{tool_name} thread error: {e}", exc_info=True)
            q.put(('error', f'{tool_name.capitalize()} Error', str(e)))
        finally:
            q.put(('tool_finished', tool_name, target, "".join(full_output)))
            with self.thread_finish_lock:
                setattr(self, f'{tool_name}_process', None)

    def _create_maigret_tool(self):
        """Creates the UI for the Maigret OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Maigret - Username Checker</b></font>
        <p>Collect a dossier on a person by username only, checking for accounts on a huge number of sites and gathering all the available information from web pages.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.maigret_controls = self._create_maigret_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.maigret_controls['start_btn'])
        buttons_layout.addWidget(self.maigret_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.maigret_output_console = QPlainTextEdit()
        self.maigret_output_console.setReadOnly(True)
        self.maigret_output_console.setFont(QFont("Courier New", 10))
        self.maigret_output_console.setPlaceholderText("Maigret output will be displayed here...")
        main_layout.addWidget(self.maigret_output_console, 1)

        self.maigret_controls['start_btn'].clicked.connect(self.start_maigret_scan)
        self.maigret_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_maigret_config_widget(self):
        """Creates a reusable, self-contained widget for Maigret's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['username_edit'] = QLineEdit()
        controls['username_edit'].setPlaceholderText("Enter one or more usernames separated by space")
        layout.addRow("Usernames:", controls['username_edit'])

        controls['tags_edit'] = QLineEdit()
        controls['tags_edit'].setPlaceholderText("e.g., photo, dating, sport")
        layout.addRow("Tags:", controls['tags_edit'])

        sites_layout = QHBoxLayout()
        controls['all_sites_check'] = QCheckBox("Use all sites for scan")
        controls['top_sites_check'] = QCheckBox("Top 500 sites")
        sites_layout.addWidget(controls['all_sites_check'])
        sites_layout.addWidget(controls['top_sites_check'])
        layout.addRow(sites_layout)

        # More options
        options_layout = QHBoxLayout()
        controls['no_recursion_check'] = QCheckBox("No Recursion")
        controls['no_extracting_check'] = QCheckBox("No Extracting")
        controls['no_confirm_check'] = QCheckBox("No Confirm")
        options_layout.addWidget(controls['no_recursion_check'])
        options_layout.addWidget(controls['no_extracting_check'])
        options_layout.addWidget(controls['no_confirm_check'])
        layout.addRow(options_layout)

        # Timeout and ignore IDs
        timeout_layout = QHBoxLayout()
        controls['timeout_label'] = QLabel("Timeout:")
        controls['timeout_spin'] = QSpinBox()
        controls['timeout_spin'].setRange(0, 3600)
        controls['timeout_spin'].setValue(60)
        controls['ignore_ids_check'] = QCheckBox("Ignore IDs")
        timeout_layout.addWidget(controls['timeout_label'])
        timeout_layout.addWidget(controls['timeout_spin'])
        timeout_layout.addWidget(controls['ignore_ids_check'])
        layout.addRow(timeout_layout)

        # More checkboxes
        more_checks_layout = QHBoxLayout()
        controls['ids_search_check'] = QCheckBox("IDS Search")
        controls['parse_usernames_check'] = QCheckBox("Parse Usernames")
        controls['use_disabled_sites_check'] = QCheckBox("Use Disabled Sites")
        more_checks_layout.addWidget(controls['ids_search_check'])
        more_checks_layout.addWidget(controls['parse_usernames_check'])
        more_checks_layout.addWidget(controls['use_disabled_sites_check'])
        layout.addRow(more_checks_layout)

        report_box = QGroupBox("Reports")
        report_layout = QHBoxLayout(report_box)
        controls['pdf_check'] = QCheckBox("PDF")
        controls['html_check'] = QCheckBox("HTML")
        controls['txt_check'] = QCheckBox("TXT")
        controls['csv_check'] = QCheckBox("CSV")
        report_layout.addWidget(controls['pdf_check'])
        report_layout.addWidget(controls['html_check'])
        report_layout.addWidget(controls['txt_check'])
        report_layout.addWidget(controls['csv_check'])
        layout.addRow(report_box)

        controls['output_folder_edit'] = QLineEdit("maigret_reports")
        layout.addRow("Output Folder:", controls['output_folder_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_maigret_command(self, controls):
        """Builds the Maigret command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("maigret", "maigret")
        if not tool_path:
             return None, None, "Maigret not found."
        command = [tool_path]

        usernames = self._get_control_value(controls, 'username_edit', 'text')
        if not usernames or not usernames.strip():
            return None, None, "At least one username is required."

        target_for_log = usernames.strip()
        command.extend(target_for_log.split())

        if tags := self._get_control_value(controls, 'tags_edit', 'text'):
            command.extend(["--tags", tags.strip()])

        if self._get_control_value(controls, 'all_sites_check', 'check'):
            command.append("-a")

        if self._get_control_value(controls, 'top_sites_check', 'check'):
            command.append("--top-sites")

        if self._get_control_value(controls, 'no_recursion_check', 'check'):
            command.append("--no-recursion")

        if self._get_control_value(controls, 'no_extracting_check', 'check'):
            command.append("--no-extracting")

        if self._get_control_value(controls, 'no_confirm_check', 'check'):
            command.append("--no-confirm")

        timeout = self._get_control_value(controls, 'timeout_spin', 'value')
        if timeout and timeout != 60:
            command.extend(["--timeout", str(timeout)])

        if self._get_control_value(controls, 'ignore_ids_check', 'check'):
            command.append("--ignore-ids")

        if self._get_control_value(controls, 'ids_search_check', 'check'):
            command.append("--ids-search")

        if self._get_control_value(controls, 'parse_usernames_check', 'check'):
            command.append("--parse-usernames")

        if self._get_control_value(controls, 'use_disabled_sites_check', 'check'):
            command.append("--use-disabled-sites")

        if self._get_control_value(controls, 'pdf_check', 'check'):
            command.append("-P")
        if self._get_control_value(controls, 'html_check', 'check'):
            command.append("-H")
        if self._get_control_value(controls, 'txt_check', 'check'):
            command.append("-T")
        if self._get_control_value(controls, 'csv_check', 'check'):
            command.append("-C")

        if output_folder := self._get_control_value(controls, 'output_folder_edit', 'text'):
            command.extend(["-fo", output_folder.strip()])

        return command, target_for_log, None

    def start_maigret_scan(self):
        """Starts the Maigret scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.maigret_controls
        command, target_for_log, error = self._build_maigret_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.maigret_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'maigret_scan', target_for_log, self.maigret_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_holehe_tool(self):
        """Creates the UI for the Holehe OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Holehe - Email Checker</b></font>
        <p>Check if an email is used on different sites like twitter, instagram, and more. Holehe will retrieve information on sites with the forgotten password function.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.holehe_controls = self._create_holehe_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.holehe_controls['start_btn'])
        buttons_layout.addWidget(self.holehe_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.holehe_output_console = QPlainTextEdit()
        self.holehe_output_console.setReadOnly(True)
        self.holehe_output_console.setFont(QFont("Courier New", 10))
        self.holehe_output_console.setPlaceholderText("Holehe output will be displayed here...")
        main_layout.addWidget(self.holehe_output_console, 1)

        self.holehe_controls['start_btn'].clicked.connect(self.start_holehe_scan)
        self.holehe_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_holehe_config_widget(self):
        """Creates a reusable, self-contained widget for Holehe's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['email_edit'] = QLineEdit()
        controls['email_edit'].setPlaceholderText("Enter a single email address")
        layout.addRow("Email Address:", controls['email_edit'])

        controls['only_used_check'] = QCheckBox("Show only used accounts")
        layout.addRow("--only-used:", controls['only_used_check'])

        controls['no_color_check'] = QCheckBox("Disable colorized output")
        layout.addRow("--no-color:", controls['no_color_check'])

        controls['no_password_recovery_check'] = QCheckBox("No Password Recovery")
        layout.addRow("--no-password-recovery:", controls['no_password_recovery_check'])

        controls['csv_check'] = QCheckBox("Create CSV")
        layout.addRow("--csv:", controls['csv_check'])

        timeout_layout = QHBoxLayout()
        controls['timeout_label'] = QLabel("Timeout:")
        controls['timeout_spin'] = QSpinBox()
        controls['timeout_spin'].setRange(1, 300)
        controls['timeout_spin'].setValue(10)
        timeout_layout.addWidget(controls['timeout_label'])
        timeout_layout.addWidget(controls['timeout_spin'])
        layout.addRow(timeout_layout)

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_holehe_command(self, controls):
        """Builds the Holehe command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("holehe", "holehe")
        if not tool_path:
             return None, None, "Holehe not found."
        command = [tool_path]

        email = self._get_control_value(controls, 'email_edit', 'text')
        if not email or not email.strip():
            return None, None, "An email address is required."

        target_for_log = email.strip()
        command.append(target_for_log)

        if self._get_control_value(controls, 'only_used_check', 'check'):
            command.append("--only-used")

        if self._get_control_value(controls, 'no_color_check', 'check'):
            command.append("--no-color")

        if self._get_control_value(controls, 'no_password_recovery_check', 'check'):
            command.append("--no-password-recovery")

        if self._get_control_value(controls, 'csv_check', 'check'):
            command.append("--csv")

        timeout_value = self._get_control_value(controls, 'timeout_spin', 'value')
        if timeout_value is not None:
            command.extend(["--timeout", str(timeout_value)])

        return command, target_for_log, None

    def start_holehe_scan(self):
        """Starts the Holehe scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.holehe_controls
        command, target_for_log, error = self._build_holehe_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which("holehe"):
            QMessageBox.critical(self, "Holehe Error", "'holehe' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.holehe_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'holehe_scan', target_for_log, self.holehe_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_ghunt_tool(self):
        """Creates the UI for the GHunt OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>GHunt - Offensive Google Framework</b></font>
        <p>An offensive Google framework that can be used to get information about Google accounts and more.</p>
        <p><b>Note:</b> GHunt requires authentication. Run <code>ghunt login</code> in your terminal first.</p>
        """)
        instructions.setFixedHeight(100)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.ghunt_controls = self._create_ghunt_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.ghunt_controls['start_btn'])
        buttons_layout.addWidget(self.ghunt_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.ghunt_output_console = QPlainTextEdit()
        self.ghunt_output_console.setReadOnly(True)
        self.ghunt_output_console.setFont(QFont("Courier New", 10))
        self.ghunt_output_console.setPlaceholderText("GHunt output will be displayed here...")
        main_layout.addWidget(self.ghunt_output_console, 1)

        self.ghunt_controls['start_btn'].clicked.connect(self.start_ghunt_scan)
        self.ghunt_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_ghunt_config_widget(self):
        """Creates a reusable, self-contained widget for GHunt's configuration options."""
        widget = QGroupBox("GHunt Modules")
        main_layout = QVBoxLayout(widget)
        controls = {}

        ghunt_tabs = QTabWidget()
        main_layout.addWidget(ghunt_tabs)
        controls['tabs'] = ghunt_tabs

        # --- Email Tab ---
        email_tab = QWidget()
        email_layout = QFormLayout(email_tab)
        controls['email_address_edit'] = QLineEdit()
        email_layout.addRow("Email Address:", controls['email_address_edit'])
        ghunt_tabs.addTab(email_tab, "Email")
        controls['email_tab'] = email_tab

        # --- Gaia Tab ---
        gaia_tab = QWidget()
        gaia_layout = QFormLayout(gaia_tab)
        controls['gaia_id_edit'] = QLineEdit()
        gaia_layout.addRow("Gaia ID:", controls['gaia_id_edit'])
        ghunt_tabs.addTab(gaia_tab, "Gaia")
        controls['gaia_tab'] = gaia_tab

        # --- Drive Tab ---
        drive_tab = QWidget()
        drive_layout = QFormLayout(drive_tab)
        controls['drive_file_id_edit'] = QLineEdit()
        drive_layout.addRow("File/Folder ID:", controls['drive_file_id_edit'])
        ghunt_tabs.addTab(drive_tab, "Drive")
        controls['drive_tab'] = drive_tab

        # --- Geolocate Tab ---
        geolocate_tab = QWidget()
        geolocate_layout = QFormLayout(geolocate_tab)
        controls['geolocate_bssid_edit'] = QLineEdit()
        geolocate_layout.addRow("BSSID (-b):", controls['geolocate_bssid_edit'])
        controls['geolocate_file_edit'] = QLineEdit()
        geolocate_layout.addRow("File (-f):", controls['geolocate_file_edit'])
        ghunt_tabs.addTab(geolocate_tab, "Geolocate")
        controls['geolocate_tab'] = geolocate_tab

        # --- Spider DAL Tab ---
        spider_tab = QWidget()
        spider_layout = QFormLayout(spider_tab)
        controls['spider_package_edit'] = QLineEdit()
        spider_layout.addRow("Package (-p):", controls['spider_package_edit'])
        controls['spider_fingerprint_edit'] = QLineEdit()
        spider_layout.addRow("Fingerprint (-f):", controls['spider_fingerprint_edit'])
        controls['spider_url_edit'] = QLineEdit()
        spider_layout.addRow("URL (-u):", controls['spider_url_edit'])
        controls['spider_strict_check'] = QCheckBox("Strict mode (-s)")
        spider_layout.addRow(controls['spider_strict_check'])
        ghunt_tabs.addTab(spider_tab, "Spider DAL")
        controls['spider_tab'] = spider_tab

        # --- Common Options ---
        common_options = QGroupBox("Common Options")
        common_layout = QFormLayout(common_options)
        controls['json_output_edit'] = QLineEdit()
        controls['json_output_edit'].setPlaceholderText("Optional: path/to/output.json")
        common_layout.addRow("JSON Output File:", controls['json_output_edit'])
        main_layout.addWidget(common_options)

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_ghunt_command(self, controls):
        """Builds the GHunt command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("ghunt", "ghunt")
        if not tool_path:
             return None, None, "GHunt not found."
        command = [tool_path]

        active_tab = controls['tabs'].currentWidget()
        target_for_log = ""

        if active_tab == controls['email_tab']:
            command.append("email")
            email = self._get_control_value(controls, 'email_address_edit', 'text')
            if not email: return None, None, "An email address is required."
            target_for_log = email
            command.append(email)
        elif active_tab == controls['gaia_tab']:
            command.append("gaia")
            gaia_id = self._get_control_value(controls, 'gaia_id_edit', 'text')
            if not gaia_id: return None, None, "A Gaia ID is required."
            target_for_log = gaia_id
            command.append(gaia_id)
        elif active_tab == controls['drive_tab']:
            command.append("drive")
            file_id = self._get_control_value(controls, 'drive_file_id_edit', 'text')
            if not file_id: return None, None, "A Drive file/folder ID is required."
            target_for_log = file_id
            command.append(file_id)
        elif active_tab == controls['geolocate_tab']:
            command.append("geolocate")
            bssid = self._get_control_value(controls, 'geolocate_bssid_edit', 'text')
            file = self._get_control_value(controls, 'geolocate_file_edit', 'text')
            if bssid:
                command.extend(["-b", bssid])
                target_for_log = bssid
            elif file:
                command.extend(["-f", file])
                target_for_log = file
            else:
                return None, None, "A BSSID or a file is required for geolocation."
        elif active_tab == controls['spider_tab']:
            command.append("spiderdal")
            package = self._get_control_value(controls, 'spider_package_edit', 'text')
            fingerprint = self._get_control_value(controls, 'spider_fingerprint_edit', 'text')
            url = self._get_control_value(controls, 'spider_url_edit', 'text')
            if package:
                command.extend(["-p", package])
                target_for_log = package
            elif fingerprint:
                command.extend(["-f", fingerprint])
                target_for_log = "fingerprint"
            elif url:
                command.extend(["-u", url])
                target_for_log = url
            else:
                return None, None, "A package, fingerprint, or URL is required for Spider DAL."
            if self._get_control_value(controls, 'spider_strict_check', 'check'):
                command.append("-s")
        else:
            return None, None, "Could not determine active GHunt command."

        if json_output := self._get_control_value(controls, 'json_output_edit', 'text'):
            command.extend(["--json", json_output.strip()])

        return command, target_for_log, None

    def start_ghunt_scan(self):
        """Starts the GHunt scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.ghunt_controls
        command, target_for_log, error = self._build_ghunt_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which("ghunt"):
            QMessageBox.critical(self, "GHunt Error", "'ghunt' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.ghunt_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'ghunt_scan', target_for_log, self.ghunt_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_recon_ng_tool(self):
        """Creates the UI for the Recon-ng OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Recon-ng - Web Reconnaissance Framework</b></font>
        <p>Recon-ng is a powerful tool for web-based reconnaissance with a modular framework. You can automate a series of commands by writing a script below.</p>
        <p><b>Example Script:</b><br>
        <code>workspaces create example.com<br>
        modules load recon/domains-hosts/bing_domain_web<br>
        options set SOURCE example.com<br>
        run<br>
        modules load recon/hosts-hosts/resolve<br>
        run<br>
        show hosts</code></p>
        """)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.recon_ng_controls = self._create_recon_ng_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.recon_ng_controls['start_btn'])
        buttons_layout.addWidget(self.recon_ng_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.recon_ng_output_console = QPlainTextEdit()
        self.recon_ng_output_console.setReadOnly(True)
        self.recon_ng_output_console.setFont(QFont("Courier New", 10))
        self.recon_ng_output_console.setPlaceholderText("Recon-ng output will be displayed here...")
        main_layout.addWidget(self.recon_ng_output_console, 1)

        self.recon_ng_controls['start_btn'].clicked.connect(self.start_recon_ng_scan)
        self.recon_ng_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_recon_ng_config_widget(self):
        """Creates a reusable, self-contained widget for Recon-ng's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['workspace_edit'] = QLineEdit("default")
        controls['workspace_edit'].setToolTip("The workspace to use for this scan (-w).")
        layout.addRow("Workspace (-w):", controls['workspace_edit'])

        controls['script_edit'] = QTextEdit()
        controls['script_edit'].setFont(QFont("Courier New", 10))
        controls['script_edit'].setPlaceholderText("Enter Recon-ng commands, one per line...")
        controls['script_edit'].setToolTip("A script of commands that will be executed non-interactively (-r).")
        layout.addRow("Script (-r):", controls['script_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_recon_ng_command(self, controls):
        """Builds the Recon-ng command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("recon-ng", "recon-ng")
        if not tool_path:
             return None, None, "Recon-ng not found."
        command = [tool_path]

        workspace = self._get_control_value(controls, 'workspace_edit', 'text')
        if not workspace or not workspace.strip():
            return None, None, "A workspace name is required."
        command.extend(["-w", workspace.strip()])
        target_for_log = workspace.strip()

        script_content = self._get_control_value(controls, 'script_edit', 'toPlainText')
        if not script_content or not script_content.strip():
            return None, None, "A script is required to run Recon-ng in a non-interactive mode."

        # Write the script content to a temporary file
        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".rc", encoding='utf-8') as tmp_rc:
                tmp_rc.write(script_content)
                self.recon_ng_rc_file = tmp_rc.name
            command.extend(["-r", self.recon_ng_rc_file])
        except Exception as e:
            return None, None, f"Could not create temporary script file: {e}"

        return command, target_for_log, None

    def start_recon_ng_scan(self):
        """Starts the Recon-ng scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.recon_ng_controls
        command, target_for_log, error = self._build_recon_ng_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            # Clean up the temp file if it was created before the error
            if hasattr(self, 'recon_ng_rc_file') and self.recon_ng_rc_file:
                os.remove(self.recon_ng_rc_file)
                self.recon_ng_rc_file = None
            return

        if not shutil.which("recon-ng"):
            QMessageBox.critical(self, "Recon-ng Error", "'recon-ng' command not found. Please ensure it is installed and in your system's PATH.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.recon_ng_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'recon_ng_scan', target_for_log, self.recon_ng_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_infoga_tool(self):
        """Creates the UI for the Infoga OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Infoga - Email OSINT</b></font>
        <p>A tool for gathering email account information from various public sources.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.infoga_controls = self._create_infoga_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.infoga_controls['start_btn'])
        buttons_layout.addWidget(self.infoga_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.infoga_output_console = QPlainTextEdit()
        self.infoga_output_console.setReadOnly(True)
        self.infoga_output_console.setFont(QFont("Courier New", 10))
        self.infoga_output_console.setPlaceholderText("Infoga output will be displayed here...")
        main_layout.addWidget(self.infoga_output_console, 1)

        self.infoga_controls['start_btn'].clicked.connect(self.start_infoga_scan)
        self.infoga_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_infoga_config_widget(self):
        """Creates a reusable, self-contained widget for Infoga's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['domain_edit'] = QLineEdit()
        controls['domain_edit'].setPlaceholderText("Enter a domain to search (e.g., example.com)")
        layout.addRow("Domain:", controls['domain_edit'])

        controls['source_edit'] = QLineEdit("all")
        controls['source_edit'].setToolTip("Data source: all, google, twitter, github, etc.")
        layout.addRow("Source:", controls['source_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_infoga_command(self, controls):
        """Builds the Infoga command list from a dictionary of controls or widgets."""
        command = ["python3", "tools/Infoga/infoga.py"]

        domain = self._get_control_value(controls, 'domain_edit', 'text')
        if not domain or not domain.strip():
            return None, None, "A domain is required."
        command.extend(["--domain", domain.strip()])
        target_for_log = domain.strip()

        source = self._get_control_value(controls, 'source_edit', 'text')
        if source and source.strip():
            command.extend(["--source", source.strip()])

        return command, target_for_log, None

    def start_infoga_scan(self):
        """Starts the Infoga scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.infoga_controls
        command, target_for_log, error = self._build_infoga_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        infoga_path = os.path.join(self.script_dir, "tools", "Infoga", "infoga.py")
        if not os.path.exists(infoga_path):
            QMessageBox.critical(self, "Infoga Error", f"'infoga.py' not found at {infoga_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.infoga_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'infoga_scan', target_for_log, self.infoga_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_social_analyzer_tool(self):
        """Creates the UI for the Social Analyzer OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Social Analyzer - Social Media Finder</b></font>
        <p>Analyzes & finds social media profiles across 300+ platforms. It's a CLI tool for analyzing and finding a person's profile across 300+ social media websites.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.social_analyzer_controls = self._create_social_analyzer_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.social_analyzer_controls['start_btn'])
        buttons_layout.addWidget(self.social_analyzer_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.social_analyzer_output_console = QPlainTextEdit()
        self.social_analyzer_output_console.setReadOnly(True)
        self.social_analyzer_output_console.setFont(QFont("Courier New", 10))
        self.social_analyzer_output_console.setPlaceholderText("Social Analyzer output will be displayed here...")
        main_layout.addWidget(self.social_analyzer_output_console, 1)

        self.social_analyzer_controls['start_btn'].clicked.connect(self.start_social_analyzer_scan)
        self.social_analyzer_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_social_analyzer_config_widget(self):
        """Creates a reusable, self-contained widget for Social Analyzer's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['username_edit'] = QLineEdit()
        controls['username_edit'].setPlaceholderText("Enter a single username to search")
        layout.addRow("Username:", controls['username_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_social_analyzer_command(self, controls):
        """Builds the Social Analyzer command list from a dictionary of controls or widgets."""
        command = ["python3", "tools/social-analyzer/main.py", "--cli", "--silent"]

        username = self._get_control_value(controls, 'username_edit', 'text')
        if not username or not username.strip():
            return None, None, "A username is required."
        command.extend(["--username", username.strip()])
        target_for_log = username.strip()

        return command, target_for_log, None

    def start_social_analyzer_scan(self):
        """Starts the Social Analyzer scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.social_analyzer_controls
        command, target_for_log, error = self._build_social_analyzer_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        social_analyzer_path = os.path.join(self.script_dir, "tools", "social-analyzer", "main.py")
        if not os.path.exists(social_analyzer_path):
            QMessageBox.critical(self, "Social Analyzer Error", f"'main.py' not found at {social_analyzer_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.social_analyzer_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'social_analyzer_scan', target_for_log, self.social_analyzer_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_metagoofil_tool(self):
        """Creates the UI for the Metagoofil OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Metagoofil - Metadata Collector</b></font>
        <p>An information gathering tool designed for extracting metadata of public documents (pdf,doc,xls,ppt,etc) available on target websites.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.metagoofil_controls = self._create_metagoofil_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.metagoofil_controls['start_btn'])
        buttons_layout.addWidget(self.metagoofil_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.metagoofil_output_console = QPlainTextEdit()
        self.metagoofil_output_console.setReadOnly(True)
        self.metagoofil_output_console.setFont(QFont("Courier New", 10))
        self.metagoofil_output_console.setPlaceholderText("Metagoofil output will be displayed here...")
        main_layout.addWidget(self.metagoofil_output_console, 1)

        self.metagoofil_controls['start_btn'].clicked.connect(self.start_metagoofil_scan)
        self.metagoofil_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_metagoofil_config_widget(self):
        """Creates a reusable, self-contained widget for Metagoofil's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['domain_edit'] = QLineEdit("example.com")
        layout.addRow("Domain (-d):", controls['domain_edit'])

        controls['file_types_edit'] = QLineEdit("pdf,doc,xls")
        layout.addRow("File Types (-t):", controls['file_types_edit'])

        controls['search_max_edit'] = QLineEdit("100")
        layout.addRow("Search Max (-l):", controls['search_max_edit'])

        controls['download_limit_edit'] = QLineEdit("100")
        layout.addRow("Download Limit (-n):", controls['download_limit_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Metagoofil Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_metagoofil_command(self, controls):
        """Builds the Metagoofil command list from a dictionary of controls or widgets."""
        command = ["python3", "tools/metagoofil/metagoofil.py"]

        domain = self._get_control_value(controls, 'domain_edit', 'text')
        if not domain or not domain.strip():
            return None, None, "A domain is required."
        command.extend(["-d", domain.strip()])
        target_for_log = domain.strip()

        file_types = self._get_control_value(controls, 'file_types_edit', 'text')
        if not file_types or not file_types.strip():
            return None, None, "File types are required."
        command.extend(["-t", file_types.strip()])

        if search_max := self._get_control_value(controls, 'search_max_edit', 'text'):
            command.extend(["-l", search_max.strip()])

        if download_limit := self._get_control_value(controls, 'download_limit_edit', 'text'):
            command.extend(["-n", download_limit.strip()])

        return command, target_for_log, None

    def start_metagoofil_scan(self):
        """Starts the Metagoofil scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.metagoofil_controls
        command, target_for_log, error = self._build_metagoofil_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        metagoofil_path = os.path.join(self.script_dir, "tools", "metagoofil", "metagoofil.py")
        if not os.path.exists(metagoofil_path):
            QMessageBox.critical(self, "Metagoofil Error", f"'metagoofil.py' not found at {metagoofil_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.metagoofil_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'metagoofil_scan', target_for_log, self.metagoofil_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_phoneinfoga_tool(self):
        """Creates the UI for the PhoneInfoga OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>PhoneInfoga - Phone Number Scanner</b></font>
        <p>One of the most advanced tools to scan phone numbers using only free resources.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.phoneinfoga_controls = self._create_phoneinfoga_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.phoneinfoga_controls['start_btn'])
        buttons_layout.addWidget(self.phoneinfoga_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.phoneinfoga_output_console = QPlainTextEdit()
        self.phoneinfoga_output_console.setReadOnly(True)
        self.phoneinfoga_output_console.setFont(QFont("Courier New", 10))
        self.phoneinfoga_output_console.setPlaceholderText("PhoneInfoga output will be displayed here...")
        main_layout.addWidget(self.phoneinfoga_output_console, 1)

        self.phoneinfoga_controls['start_btn'].clicked.connect(self.start_phoneinfoga_scan)
        self.phoneinfoga_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_phoneinfoga_config_widget(self):
        """Creates a reusable, self-contained widget for PhoneInfoga's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['number_edit'] = QLineEdit()
        controls['number_edit'].setPlaceholderText("e.g., +15551234567")
        layout.addRow("Phone Number (-n):", controls['number_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_phoneinfoga_command(self, controls):
        """Builds the PhoneInfoga command list from a dictionary of controls or widgets."""
        command = ["./tools/PhoneInfoga/bin/phoneinfoga", "scan"]

        number = self._get_control_value(controls, 'number_edit', 'text')
        if not number or not number.strip():
            return None, None, "A phone number is required."
        command.extend(["-n", number.strip()])
        target_for_log = number.strip()

        return command, target_for_log, None

    def start_phoneinfoga_scan(self):
        """Starts the PhoneInfoga scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.phoneinfoga_controls
        command, target_for_log, error = self._build_phoneinfoga_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        phoneinfoga_path = os.path.join(self.script_dir, "tools", "PhoneInfoga", "bin", "phoneinfoga")
        if not os.path.exists(phoneinfoga_path):
            QMessageBox.critical(self, "PhoneInfoga Error", f"'phoneinfoga' executable not found at {phoneinfoga_path}. Please build it first.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.phoneinfoga_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'phoneinfoga_scan', target_for_log, self.phoneinfoga_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_photon_tool(self):
        """Creates the UI for the Photon OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Photon - OSINT & Crawler</b></font>
        <p>An incredibly fast crawler designed for OSINT. Extracts URLs, files, intel & endpoints from a target.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.photon_controls = self._create_photon_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.photon_controls['start_btn'])
        buttons_layout.addWidget(self.photon_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.photon_output_console = QPlainTextEdit()
        self.photon_output_console.setReadOnly(True)
        self.photon_output_console.setFont(QFont("Courier New", 10))
        self.photon_output_console.setPlaceholderText("Photon output will be displayed here...")
        main_layout.addWidget(self.photon_output_console, 1)

        self.photon_controls['start_btn'].clicked.connect(self.start_photon_scan)
        self.photon_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_photon_config_widget(self):
        """Creates a reusable, self-contained widget for Photon's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['url_edit'] = QLineEdit("http://example.com")
        layout.addRow("Root URL (-u):", controls['url_edit'])

        controls['level_edit'] = QLineEdit("2")
        controls['level_edit'].setToolTip("Levels to crawl (-l).")
        layout.addRow("Crawl Level (-l):", controls['level_edit'])

        controls['threads_edit'] = QLineEdit("10")
        controls['threads_edit'].setToolTip("Number of threads (-t).")
        layout.addRow("Threads (-t):", controls['threads_edit'])

        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("e.g., results/")
        controls['output_edit'].setToolTip("Output directory (-o).")
        layout.addRow("Output Directory (-o):", controls['output_edit'])

        # Checkboxes for boolean flags
        checkbox_layout = QGridLayout()
        controls['dns_check'] = QCheckBox("DNS (--dns)")
        controls['keys_check'] = QCheckBox("Secret Keys (--keys)")
        controls['wayback_check'] = QCheckBox("Wayback Machine (--wayback)")
        controls['only_urls_check'] = QCheckBox("Only URLs (--only-urls)")
        checkbox_layout.addWidget(controls['dns_check'], 0, 0)
        checkbox_layout.addWidget(controls['keys_check'], 0, 1)
        checkbox_layout.addWidget(controls['wayback_check'], 1, 0)
        checkbox_layout.addWidget(controls['only_urls_check'], 1, 1)
        layout.addRow(checkbox_layout)

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Photon Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_photon_command(self, controls):
        """Builds the Photon command list from a dictionary of controls or widgets."""
        command = ["python3", "tools/Photon/photon.py"]

        url = self._get_control_value(controls, 'url_edit', 'text')
        if not url or not url.strip():
            return None, None, "A root URL is required."
        command.extend(["-u", url.strip()])
        target_for_log = url.strip()

        if level := self._get_control_value(controls, 'level_edit', 'text'):
            command.extend(["-l", level.strip()])
        if threads := self._get_control_value(controls, 'threads_edit', 'text'):
            command.extend(["-t", threads.strip()])
        if output := self._get_control_value(controls, 'output_edit', 'text'):
            command.extend(["-o", output.strip()])

        if self._get_control_value(controls, 'dns_check', 'check'): command.append("--dns")
        if self._get_control_value(controls, 'keys_check', 'check'): command.append("--keys")
        if self._get_control_value(controls, 'wayback_check', 'check'): command.append("--wayback")
        if self._get_control_value(controls, 'only_urls_check', 'check'): command.append("--only-urls")

        return command, target_for_log, None

    def start_photon_scan(self):
        """Starts the Photon scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.photon_controls
        command, target_for_log, error = self._build_photon_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        photon_path = os.path.join(self.script_dir, "tools", "Photon", "photon.py")
        if not os.path.exists(photon_path):
            QMessageBox.critical(self, "Photon Error", f"'photon.py' not found at {photon_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.photon_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'photon_scan', target_for_log, self.photon_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_magma_tool(self):
        """Creates the UI for the Magma OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Magma OSINT</b></font>
        <p>This tool searches Google for a given query and analyzes the resulting URLs for related links and data.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        controls_frame = QGroupBox("Scan Options")
        layout = QFormLayout(controls_frame)
        self.magma_controls = {}

        self.magma_controls['query_edit'] = QLineEdit()
        self.magma_controls['query_edit'].setPlaceholderText("Enter a search query (e.g., a name, company, or domain)")
        layout.addRow("Query:", self.magma_controls['query_edit'])

        self.magma_controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        self.magma_controls['stop_btn'] = QPushButton("Stop"); self.magma_controls['stop_btn'].setEnabled(False)
        main_layout.addWidget(controls_frame)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.magma_controls['start_btn'])
        buttons_layout.addWidget(self.magma_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.magma_output_console = QPlainTextEdit()
        self.magma_output_console.setReadOnly(True)
        self.magma_output_console.setFont(QFont("Courier New", 10))
        self.magma_output_console.setPlaceholderText("Magma OSINT output will be displayed here...")
        main_layout.addWidget(self.magma_output_console, 1)

        self.magma_controls['start_btn'].clicked.connect(self.start_magma_scan)
        self.magma_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _build_magma_command(self, controls):
        """Builds the Magma command list and prepares the input."""
        tool_path = self._get_tool_path("osint.py", os.path.join("Magma-Osint", "osint.py"))

        if tool_path:
             command = ["python3", tool_path]
        else:
             return None, None, None, "Magma-Osint script (osint.py) not found."

        query = self._get_control_value(controls, 'query_edit', 'text')
        if not query or not query.strip():
            return None, None, None, "A query is required."

        target_for_log = query.strip()
        stdin_input = f"{query.strip()}\n"

        return command, target_for_log, stdin_input, None

    def start_magma_scan(self):
        """Starts the Magma OSINT scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.magma_controls
        command, target_for_log, stdin_input, error = self._build_magma_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        magma_path = os.path.join(self.script_dir, "tools", "Magma-Osint", "osint.py")
        if not os.path.exists(magma_path):
            QMessageBox.critical(self, "Magma Error", f"'osint.py' not found at {magma_path}.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.magma_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread_with_input, args=(
            command, 'magma_scan', target_for_log, self.magma_output_console, stdin_input
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _execute_command_thread_with_input(self, command, tool_name, target, output_widget, stdin_input):
        """
        A generic worker thread for running external command-line tools that require stdin.
        """
        q = self.tool_results_queue
        logging.info(f"Starting {tool_name} with command: {' '.join(command)}")
        q.put((f'{tool_name}_output', f"$ {' '.join(command)}\n"))
        full_output = []
        success = False
        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                                     stdin=subprocess.PIPE, text=True, bufsize=1,
                                     startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                setattr(self, f'{tool_name}_process', process)

            # Write the input to the process's stdin
            if stdin_input:
                process.stdin.write(stdin_input)
                process.stdin.close()

            # Read output line by line
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put((f'{tool_name}_output', "\n\n--- Canceled By User ---\n"))
                    break
                q.put((f'{tool_name}_output', line))
                full_output.append(line)

            process.stdout.close()
            retcode = process.wait()
            success = (retcode == 0)

        except FileNotFoundError:
            q.put(('error', f'{tool_name.capitalize()} Error', f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"{tool_name} thread error: {e}", exc_info=True)
            q.put(('error', f'{tool_name.capitalize()} Error', str(e)))
        finally:
            # Clean up the temporary script file for Recon-ng
            if tool_name == 'recon_ng_scan' and hasattr(self, 'recon_ng_rc_file') and self.recon_ng_rc_file:
                try:
                    os.remove(self.recon_ng_rc_file)
                    logging.info(f"Removed temporary recon-ng script: {self.recon_ng_rc_file}")
                    self.recon_ng_rc_file = None
                except OSError as e:
                    logging.error(f"Error removing recon-ng temp file: {e}")

            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='Tool Execution',
                    action=f"Executed {tool_name.replace('_', ' ').title()}",
                    target=target,
                    details=f"Command: {' '.join(command)}",
                    result='Success' if success else 'Failure'
                )
            q.put(('tool_finished', tool_name, target, "".join(full_output)))
            with self.thread_finish_lock:
                setattr(self, f'{tool_name}_process', None)
            logging.info(f"{tool_name} scan thread finished.")

    def _create_advanced_tools_tab(self, p=None):
        """Creates the tab container for advanced tools."""
        adv_tabs = QTabWidget()
        adv_tabs.addTab(self._create_flooder_tool(), "Packet Flooder")
        adv_tabs.addTab(self._create_firewall_tester_tool(), "Firewall Tester")
        adv_tabs.addTab(self._create_arp_spoofer_tool(), "ARP Spoofer")
        adv_tabs.addTab(self._create_snort_ids_tool(), "Snort IDS")
        adv_tabs.addTab(self._create_sqlmap_tool(), "SQLMap")
        adv_tabs.addTab(self._create_hashcat_tool(), "Hashcat")
        adv_tabs.addTab(self._create_nuclei_tool(), "Nuclei Scanner")
        adv_tabs.addTab(self._create_trufflehog_tool(), "TruffleHog Scanner")
        adv_tabs.addTab(self._create_jtr_tool(), "John the Ripper")
        adv_tabs.addTab(self._create_hydra_tool(), "Hydra")
        adv_tabs.addTab(self._create_zap_scanner_tool(), "OWASP ZAP")
        adv_tabs.addTab(self._create_metasploit_tool(), "Metasploit Scanner")
        adv_tabs.addTab(self._create_firewall_evasion_tab(), "Firewall Evasion")
        return adv_tabs

    def _create_flooder_tool(self):
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Attack Template Box ---
        template_box = QGroupBox("Attack Configuration")
        template_layout = QFormLayout(template_box)

        self.flood_template_combo = QComboBox()
        self.flood_template_combo.addItems(["Custom (from Crafter)", "TCP SYN Flood", "UDP Flood", "ICMP Echo Flood"])
        template_layout.addRow("Template:", self.flood_template_combo)

        self.flood_target_label = QLabel("Target IP:")
        self.flood_target_edit = QLineEdit("127.0.0.1")
        template_layout.addRow(self.flood_target_label, self.flood_target_edit)

        self.flood_ports_label = QLabel("Target Port(s):")
        self.flood_ports_edit = QLineEdit("80")
        self.flood_ports_edit.setToolTip("A single port for the flood attack.")
        template_layout.addRow(self.flood_ports_label, self.flood_ports_edit)

        self.flood_rand_src_ip_check = QCheckBox()
        self.flood_rand_src_ip_check.setToolTip("Randomize the source IP address for each packet.")
        template_layout.addRow("Randomize Source IP:", self.flood_rand_src_ip_check)
        main_layout.addWidget(template_box)

        # --- Custom Packet Box (for loading) ---
        packet_frame = QGroupBox("Custom Packet Loader")
        packet_layout = QVBoxLayout(packet_frame)
        self.flood_packet_label = QLabel("Packet to send: (Load from Crafter)")
        packet_layout.addWidget(self.flood_packet_label)
        load_btn = QPushButton("Load Packet from Crafter")
        load_btn.clicked.connect(self.load_flood_packet)
        packet_layout.addWidget(load_btn)
        main_layout.addWidget(packet_frame)

        # --- Flood Parameters ---
        controls_frame = QGroupBox("Flood Parameters")
        controls_layout = QFormLayout(controls_frame)
        self.flood_count = QLineEdit("1000")
        self.flood_count.setToolTip("The total number of packets to send.")
        controls_layout.addRow("Count:", self.flood_count)
        self.flood_interval = QLineEdit("0.01")
        self.flood_interval.setToolTip("The time interval (in seconds) between sending each packet.")
        controls_layout.addRow("Interval:", self.flood_interval)
        self.flood_threads = QLineEdit("4")
        self.flood_threads.setToolTip("The number of parallel threads to use for sending packets.")
        controls_layout.addRow("Threads:", self.flood_threads)
        main_layout.addWidget(controls_frame)

        # --- Action Buttons ---
        flood_buttons_layout = QHBoxLayout()
        self.flood_button = QPushButton("Start Flood")
        self.flood_button.setToolTip("Start the packet flood. Warning: This can cause network disruption.")
        flood_buttons_layout.addWidget(self.flood_button)
        self.stop_flood_button = QPushButton("Stop Flood")
        self.stop_flood_button.setEnabled(False)
        self.stop_flood_button.setToolTip("Stop the ongoing flood.")
        flood_buttons_layout.addWidget(self.stop_flood_button)
        main_layout.addLayout(flood_buttons_layout)

        self.flood_status = QLabel("")
        main_layout.addWidget(self.flood_status)
        main_layout.addStretch()

        # --- UI Logic ---
        def update_template_ui(text):
            is_custom = (text == "Custom (from Crafter)")
            is_icmp = (text == "ICMP Echo Flood")

            self.flood_target_label.setVisible(not is_custom)
            self.flood_target_edit.setVisible(not is_custom)
            self.flood_ports_label.setVisible(not is_custom and not is_icmp)
            self.flood_ports_edit.setVisible(not is_custom and not is_icmp)
            self.flood_rand_src_ip_check.setEnabled(not is_custom)
            packet_frame.setVisible(is_custom)

        self.flood_template_combo.currentTextChanged.connect(update_template_ui)
        update_template_ui(self.flood_template_combo.currentText()) # Initial state

        self.flood_button.clicked.connect(self.start_flood)
        self.stop_flood_button.clicked.connect(self.cancel_tool)

        return widget

    def _create_firewall_tester_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget); controls = QHBoxLayout()
        controls.addWidget(QLabel("Target:")); self.fw_target=QLineEdit("127.0.0.1"); controls.addWidget(self.fw_target)
        controls.addWidget(QLabel("Probe Set:")); self.fw_probe_set=QComboBox(); self.fw_probe_set.addItems(FIREWALL_PROBES.keys()); controls.addWidget(self.fw_probe_set)
        self.fw_test_button=QPushButton("Start Test"); controls.addWidget(self.fw_test_button)
        self.fw_cancel_button = QPushButton("Cancel"); self.fw_cancel_button.setEnabled(False); controls.addWidget(self.fw_cancel_button)
        self.fw_status=QLabel(""); controls.addWidget(self.fw_status); controls.addStretch()
        layout.addLayout(controls)
        self.fw_tree=QTreeWidget(); self.fw_tree.setColumnCount(3); self.fw_tree.setHeaderLabels(["Probe Description","Packet Summary","Result"])
        layout.addWidget(self.fw_tree)
        layout.addWidget(self._create_export_button(self.fw_tree))
        self.fw_test_button.clicked.connect(self.start_firewall_test)
        self.fw_cancel_button.clicked.connect(self.cancel_tool)
        return widget

    def _update_tcp_scan_options_visibility(self, checked):
        """Shows or hides the TCP scan mode dropdown based on protocol selection."""
        is_tcp_selected = self.scan_proto_tcp_radio.isChecked() or self.scan_proto_udp_radio.isChecked()
        self.tcp_scan_type_label.setVisible(is_tcp_selected)
        self.tcp_scan_type_combo.setVisible(is_tcp_selected)

    def _create_arp_spoofer_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Ethical Warning
        warning_label = QTextEdit()
        warning_label.setReadOnly(True)
        warning_label.setStyleSheet("background-color: #4c2222; color: #f0f0f0; border: 1px solid #993333;")
        warning_label.setHtml("""
        <font color='#ffcc00'><b>WARNING & ETHICAL NOTICE:</b></font>
        <p>ARP Spoofing is a powerful technique that can intercept and modify network traffic (Man-in-the-Middle attack). Using this tool on networks you do not own or have explicit, written permission to test is <b>illegal</b> and unethical.</p>
        <p>This tool is for educational and authorized security testing purposes only. The developer assumes no liability for misuse.</p>
        """)
        layout.addWidget(warning_label)

        # Controls
        controls = QFrame()
        controls.setFrameShape(QFrame.Shape.StyledPanel)
        clayout = QVBoxLayout(controls)

        # Target Inputs
        row1 = QHBoxLayout()
        row1.addWidget(QLabel("Victim IP:"))
        self.arp_spoof_victim_ip = QLineEdit()
        self.arp_spoof_victim_ip.setPlaceholderText("e.g., 192.168.1.10")
        self.arp_spoof_victim_ip.setToolTip("The IP address of the target (victim) machine on the local network.")
        row1.addWidget(self.arp_spoof_victim_ip)
        clayout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Target IP (Gateway):"))
        self.arp_spoof_target_ip = QLineEdit()
        self.arp_spoof_target_ip.setPlaceholderText("e.g., 192.168.1.1")
        self.arp_spoof_target_ip.setToolTip("The IP address of the machine you want to impersonate (usually the gateway).")
        row2.addWidget(self.arp_spoof_target_ip)
        clayout.addLayout(row2)

        # Buttons
        buttons_layout = QHBoxLayout()
        self.arp_spoof_start_btn = QPushButton("Start Spoofing")
        self.arp_spoof_start_btn.setToolTip("Begin sending malicious ARP packets to poison the cache of the victim and target.")
        buttons_layout.addWidget(self.arp_spoof_start_btn)
        self.arp_spoof_stop_btn = QPushButton("Stop Spoofing")
        self.arp_spoof_stop_btn.setEnabled(False)
        self.arp_spoof_stop_btn.setToolTip("Stop the attack and send corrective ARP packets to restore the network.")
        buttons_layout.addWidget(self.arp_spoof_stop_btn)
        clayout.addLayout(buttons_layout)

        # Status Label
        self.arp_spoof_status = QLabel("Status: Idle")
        clayout.addWidget(self.arp_spoof_status)

        layout.addWidget(controls)
        layout.addStretch()

        self.arp_spoof_start_btn.clicked.connect(self.start_arp_spoof)
        self.arp_spoof_stop_btn.clicked.connect(self.stop_arp_spoof)

        return widget

    def _create_snort_ids_tool(self):
        """Creates the UI for the Snort IDS tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Snort - Intrusion Detection System</b></font>
        <p>Run the Snort IDS to monitor network traffic in real-time and detect potential threats based on a defined set of rules. Ensure your selected interface is correct.</p>
        """)
        instructions.setFixedHeight(80)
        layout.addWidget(instructions)

        # --- Main Controls ---
        main_controls_frame = QGroupBox("Main Configuration")
        main_controls_layout = QFormLayout(main_controls_frame)
        self.snort_controls = {}

        # Config File
        conf_layout = QHBoxLayout()
        self.snort_controls['conf_edit'] = QLineEdit()
        conf_layout.addWidget(self.snort_controls['conf_edit'])
        browse_conf_btn = QPushButton("Browse...")
        browse_conf_btn.clicked.connect(lambda: self._browse_file_for_lineedit(self.snort_controls['conf_edit'], "Select snort.conf"))
        conf_layout.addWidget(browse_conf_btn)
        main_controls_layout.addRow("Config File (-c):", conf_layout)

        # Rules File (for editing convenience)
        rules_layout = QHBoxLayout()
        self.snort_controls['rules_edit'] = QLineEdit()
        rules_layout.addWidget(self.snort_controls['rules_edit'])
        browse_rules_btn = QPushButton("Browse...")
        browse_rules_btn.clicked.connect(lambda: self._browse_file_for_lineedit(self.snort_controls['rules_edit'], "Select local.rules"))
        rules_layout.addWidget(browse_rules_btn)
        main_controls_layout.addRow("Rules File (for editing):", rules_layout)
        layout.addWidget(main_controls_frame)

        # --- Tabs for Advanced Options ---
        options_tabs = QTabWidget()

        # Logging Tab
        logging_tab = QWidget()
        logging_layout = QFormLayout(logging_tab)
        self.snort_controls['alert_mode_combo'] = QComboBox()
        self.snort_controls['alert_mode_combo'].addItems(["console", "cmg", "fast", "full", "none"])
        logging_layout.addRow("Alert Mode (-A):", self.snort_controls['alert_mode_combo'])

        log_dir_layout = QHBoxLayout()
        self.snort_controls['log_dir_edit'] = QLineEdit()
        log_dir_layout.addWidget(self.snort_controls['log_dir_edit'])
        browse_log_dir_btn = QPushButton("Browse...")
        browse_log_dir_btn.clicked.connect(lambda: self._browse_dir_for_lineedit(self.snort_controls['log_dir_edit'], "Select Log Directory"))
        log_dir_layout.addWidget(browse_log_dir_btn)
        logging_layout.addRow("Log Directory (-l):", log_dir_layout)

        self.snort_controls['binary_log_check'] = QCheckBox("Log packets in pcap format (-b)")
        logging_layout.addRow(self.snort_controls['binary_log_check'])

        self.snort_controls['no_log_check'] = QCheckBox("Disable Logging (-N)")
        logging_layout.addRow(self.snort_controls['no_log_check'])

        options_tabs.addTab(logging_tab, "Logging & Alerting")

        # Behavior Tab
        behavior_tab = QWidget()
        behavior_layout = QFormLayout(behavior_tab)
        self.snort_controls['packet_count_edit'] = QLineEdit()
        self.snort_controls['packet_count_edit'].setPlaceholderText("Process all packets")
        behavior_layout.addRow("Packet Count (-n):", self.snort_controls['packet_count_edit'])

        self.snort_controls['checksum_mode_combo'] = QComboBox()
        self.snort_controls['checksum_mode_combo'].addItems(["all", "noip", "notcp", "noudp", "noicmp", "none"])
        behavior_layout.addRow("Checksum Mode (-k):", self.snort_controls['checksum_mode_combo'])

        self.snort_controls['verbose_check'] = QCheckBox("Verbose (-v)")
        self.snort_controls['quiet_check'] = QCheckBox("Quiet (-q)")
        self.snort_controls['daemon_check'] = QCheckBox("Run as Daemon (-D)")
        behavior_layout.addRow(self.snort_controls['verbose_check'])
        behavior_layout.addRow(self.snort_controls['quiet_check'])
        behavior_layout.addRow(self.snort_controls['daemon_check'])
        options_tabs.addTab(behavior_tab, "Behavior")

        # Packet Dump Tab
        dump_tab = QWidget()
        dump_layout = QFormLayout(dump_tab)
        self.snort_controls['dump_app_check'] = QCheckBox("Dump Application Layer (-d)")
        self.snort_controls['dump_link_check'] = QCheckBox("Dump Link Layer (-e)")
        self.snort_controls['dump_hex_check'] = QCheckBox("Dump Packets in Hex (-X)")
        dump_layout.addRow(self.snort_controls['dump_app_check'])
        dump_layout.addRow(self.snort_controls['dump_link_check'])
        dump_layout.addRow(self.snort_controls['dump_hex_check'])
        options_tabs.addTab(dump_tab, "Packet Dump")

        layout.addWidget(options_tabs)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        self.snort_start_btn = QPushButton(QIcon(self.icon_path("shield.svg")), " Start Snort")
        self.snort_stop_btn = QPushButton("Stop Snort"); self.snort_stop_btn.setEnabled(False)
        self.snort_guide_btn = QPushButton(QIcon(self.icon_path("help-circle.svg")), " Installation Guide")
        buttons_layout.addWidget(self.snort_start_btn)
        buttons_layout.addWidget(self.snort_stop_btn)
        buttons_layout.addStretch()
        buttons_layout.addWidget(self.snort_guide_btn)
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.snort_output_console = QPlainTextEdit()
        self.snort_output_console.setReadOnly(True)
        self.snort_output_console.setFont(QFont("Courier New", 10))
        self.snort_output_console.setPlaceholderText("Snort alerts will be displayed here...")
        layout.addWidget(self.snort_output_console, 1)

        # --- Check for Snort and enable/disable ---
        snort_path = shutil.which("snort")
        if not snort_path:
            widget.setEnabled(False)
            error_label = QLabel("<b>Snort command not found in system PATH. Please install Snort to use this feature.</b>")
            error_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            error_label.setStyleSheet("color: #ffcc00;")
            layout.insertWidget(0, error_label)

        # Auto-detect common paths
        if sys.platform == "linux" and os.path.exists("/etc/snort/snort.conf"):
            self.snort_controls['conf_edit'].setText("/etc/snort/snort.conf")
            self.snort_controls['rules_edit'].setText("/etc/snort/rules/local.rules")
        elif sys.platform == "win32" and snort_path:
            snort_dir = os.path.dirname(snort_path)
            win_conf_path = os.path.join(snort_dir, "..", "etc", "snort.conf")
            if os.path.exists(win_conf_path):
                 self.snort_controls['conf_edit'].setText(os.path.normpath(win_conf_path))
                 self.snort_controls['rules_edit'].setText(os.path.normpath(os.path.join(snort_dir, "..", "rules", "local.rules")))


        # --- Connections ---
        self.snort_start_btn.clicked.connect(self.start_snort)
        self.snort_stop_btn.clicked.connect(self.cancel_tool)
        self.snort_guide_btn.clicked.connect(self._show_snort_guide)

        return widget

    def start_snort(self, sudo_password=None):
        """Builds the Snort command from the UI and starts the worker thread."""
        if self.is_tool_running and not sudo_password:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.snort_controls
        conf_file = controls['conf_edit'].text().strip()
        if not conf_file:
            QMessageBox.critical(self, "Input Error", "A Snort configuration file (-c) is required.")
            return
        if not os.path.exists(conf_file):
            QMessageBox.critical(self, "File Error", f"Configuration file not found:\n{conf_file}")
            return

        command = ["snort", "-c", conf_file]

        # Interface
        iface = self.get_selected_iface()
        if iface:
            command.extend(["-i", iface])
        else:
            QMessageBox.warning(self, "Interface Warning", "No interface selected. Snort may default to the first available, which might not be what you want.")

        # Logging & Alerting Tab
        command.extend(["-A", controls['alert_mode_combo'].currentText()])
        if log_dir := controls['log_dir_edit'].text().strip():
            command.extend(["-l", log_dir])
        if controls['binary_log_check'].isChecked():
            command.append("-b")
        if controls['no_log_check'].isChecked():
            command.append("-N")

        # Behavior Tab
        if count := controls['packet_count_edit'].text().strip():
            command.extend(["-n", count])
        command.extend(["-k", controls['checksum_mode_combo'].currentText()])
        if controls['verbose_check'].isChecked(): command.append("-v")
        if controls['quiet_check'].isChecked(): command.append("-q")
        if controls['daemon_check'].isChecked(): command.append("-D")

        # Packet Dump Tab
        if controls['dump_app_check'].isChecked(): command.append("-d")
        if controls['dump_link_check'].isChecked(): command.append("-e")
        if controls['dump_hex_check'].isChecked(): command.append("-X")

        # Sudo prompt if needed
        if not sudo_password and sys.platform != "win32":
            if 'snort_ids' not in self.sudo_cancel_handlers:
                self.sudo_cancel_handlers['snort_ids'] = lambda: (
                    self.snort_start_btn.setEnabled(True),
                    self.snort_stop_btn.setEnabled(False)
                )
            self._run_command_with_sudo_prompt(command, self.start_snort, 'snort_ids')
            return

        self.is_tool_running = True
        self.snort_start_btn.setEnabled(False)
        self.snort_stop_btn.setEnabled(True)
        self.tool_stop_event.clear()
        self.snort_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'snort_ids', conf_file, self.snort_output_console, sudo_password
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _handle_snort_output(self, line):
        self.snort_output_console.insertPlainText(line)
        self.snort_output_console.verticalScrollBar().setValue(self.snort_output_console.verticalScrollBar().maximum())

    def _create_system_info_tab(self):
        """Creates the System Info tab with a redesigned, more modern layout."""
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setStyleSheet("QScrollArea { border: none; }") # Remove scroll area border

        main_widget = QWidget()
        main_layout = QVBoxLayout(main_widget)
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(20, 20, 20, 20)
        scroll_area.setWidget(main_widget)

        # --- Helper for creating styled GroupBoxes ---
        def create_info_box(title, icon_name=None):
            box = QGroupBox()
            box.setStyleSheet("""
                QGroupBox {
                    border: 1px solid #444;
                    border-radius: 8px;
                    margin-top: 10px;
                }
            """)

            box_layout = QVBoxLayout(box)
            box_layout.setSpacing(10)

            # Custom Title Layout
            title_layout = QHBoxLayout()
            if icon_name:
                icon_label = QLabel()
                icon_path = os.path.join(os.path.dirname(__file__), "icons", icon_name)
                icon_label.setPixmap(QIcon(icon_path).pixmap(22, 22))
                title_layout.addWidget(icon_label)

            title_label = QLabel(title)
            title_label.setStyleSheet("font-size: 14px; font-weight: bold;")
            title_layout.addWidget(title_label)
            title_layout.addStretch()
            box_layout.addLayout(title_layout)

            # Content Layout
            content_layout = QFormLayout()
            content_layout.setSpacing(10)
            content_layout.setContentsMargins(15, 10, 15, 15)
            box_layout.addLayout(content_layout)

            return box, content_layout

        # --- Top Row: System, CPU, Memory ---
        top_row_layout = QHBoxLayout()
        top_row_layout.setSpacing(20)

        # System Info Box
        sys_box, sys_layout = create_info_box("System", "info.svg")
        sys_layout.addRow("OS:", QLabel(f"{platform.system()} {platform.release()}"))
        sys_layout.addRow("Architecture:", QLabel(platform.machine()))
        sys_layout.addRow("Hostname:", QLabel(platform.node()))
        sys_layout.addRow("Python Version:", QLabel(platform.python_version()))
        top_row_layout.addWidget(sys_box)

        # CPU Info Box
        cpu_box, cpu_layout = create_info_box("CPU", "layers.svg")
        try:
            cpu_freq = psutil.cpu_freq()
            freq_str = f"{cpu_freq.current:.2f} Mhz (Max: {cpu_freq.max:.2f} Mhz)" if cpu_freq else "N/A"
        except Exception:
            freq_str = "N/A (Permission Denied)"
        cpu_layout.addRow("Frequency:", QLabel(freq_str))
        cpu_layout.addRow("Physical Cores:", QLabel(str(psutil.cpu_count(logical=False))))
        cpu_layout.addRow("Logical Cores:", QLabel(str(psutil.cpu_count(logical=True))))
        top_row_layout.addWidget(cpu_box)

        # Memory Info Box
        mem_box, mem_layout = create_info_box("Memory", "database.svg")
        mem = psutil.virtual_memory()
        swap = psutil.swap_memory()
        mem_layout.addRow("Total RAM:", QLabel(f"{mem.total / (1024**3):.2f} GB"))
        mem_layout.addRow("Available RAM:", QLabel(f"{mem.available / (1024**3):.2f} GB"))
        mem_layout.addRow("Swap Total:", QLabel(f"{swap.total / (1024**3):.2f} GB"))
        mem_layout.addRow("Swap Used:", QLabel(f"{swap.used / (1024**3):.2f} GB ({swap.percent}%)"))
        top_row_layout.addWidget(mem_box)

        main_layout.addLayout(top_row_layout)

        # --- Second Row: Libraries and GPU ---
        second_row_layout = QHBoxLayout()
        second_row_layout.setSpacing(20)

        # Library Versions Box
        try:
            scapy_version = scapy.VERSION
        except AttributeError:
            scapy_version = "Unknown"
        lib_box, lib_layout = create_info_box("Library Versions", "file-text.svg")
        lib_layout.addRow("Scapy:", QLabel(scapy_version))
        lib_layout.addRow("PyQt6:", QLabel(PYQT_VERSION_STR))
        lib_layout.addRow("psutil:", QLabel(psutil.__version__))
        if GPUtil:
            lib_layout.addRow("GPUtil:", QLabel(getattr(GPUtil, '__version__', 'N/A')))
        second_row_layout.addWidget(lib_box)

        # GPU Info Box
        if GPUtil:
            gpu_box, gpu_layout = create_info_box("GPU Information", "tool.svg")
            try:
                gpus = GPUtil.getGPUs()
                if not gpus:
                    gpu_layout.addRow(QLabel("No NVIDIA GPU detected."))
                else:
                    for i, gpu in enumerate(gpus):
                        gpu_layout.addRow(f"GPU {i} Name:", QLabel(gpu.name))
                        gpu_layout.addRow("  - Driver:", QLabel(gpu.driver))
                        gpu_layout.addRow("  - Memory:", QLabel(f"{gpu.memoryUsed}MB / {gpu.memoryTotal}MB"))
            except Exception as e:
                gpu_layout.addRow(QLabel(f"Could not retrieve GPU info: {e}"))
            second_row_layout.addWidget(gpu_box)

        second_row_layout.addStretch()
        main_layout.addLayout(second_row_layout)

        # --- Disk Partitions Box ---
        disk_box, disk_layout = create_info_box("Disk Partitions", "folder.svg")
        disk_content_widget = QWidget() # Use a simple widget for the grid
        disk_grid_layout = QGridLayout(disk_content_widget)
        try:
            partitions = psutil.disk_partitions()
            if not partitions:
                disk_grid_layout.addWidget(QLabel("No disk partitions found."), 0, 0)
            else:
                row, col = 0, 0
                for part in partitions:
                    try:
                        usage = psutil.disk_usage(part.mountpoint)
                        part_label = QLabel(f"<b>{part.device}</b> on {part.mountpoint} ({part.fstype})<br>"
                                          f"&nbsp;&nbsp;Total: {usage.total / (1024**3):.2f} GB, "
                                          f"Used: {usage.used / (1024**3):.2f} GB ({usage.percent}%)")
                        disk_grid_layout.addWidget(part_label, row, col)
                        col += 1
                        if col >= 2: # 2 columns
                            col = 0
                            row += 1
                    except Exception:
                        continue # Skip inaccessible drives
        except Exception as e:
            disk_grid_layout.addWidget(QLabel(f"Could not retrieve disk partitions: {e}"), 0, 0)
        disk_layout.addRow(disk_content_widget)
        main_layout.addWidget(disk_box)


        # --- Network Interfaces Box ---
        net_box, net_layout = create_info_box("Network Interfaces", "wifi.svg")
        net_content_widget = QWidget()
        net_main_layout = QVBoxLayout(net_content_widget)
        try:
            ifaddrs = psutil.net_if_addrs()
            if not ifaddrs:
                net_main_layout.addWidget(QLabel("No network interfaces found."))
            else:
                for iface, addrs in sorted(ifaddrs.items()):
                    # Skip loopback interfaces unless they have a non-standard address
                    is_loopback = 'loopback' in iface.lower() or iface.startswith('lo')
                    if is_loopback and all(addr.address in ['127.0.0.1', '::1'] for addr in addrs):
                        continue

                    iface_box = QGroupBox(iface)
                    iface_box.setStyleSheet("QGroupBox { border: 1px solid #555; margin-top: 5px; }")
                    iface_layout = QFormLayout(iface_box)

                    addr_map = {'ipv4': [], 'ipv6': [], 'mac': 'N/A'}
                    for addr in addrs:
                        if addr.family == socket.AF_INET:
                            addr_map['ipv4'].append(addr.address)
                        elif addr.family == socket.AF_INET6:
                            # Filter out link-local addresses for cleaner display
                            if not addr.address.startswith('fe80::'):
                                addr_map['ipv6'].append(addr.address)
                        # This logic correctly handles cross-platform MAC address retrieval
                        elif hasattr(psutil, 'AF_LINK') and addr.family == psutil.AF_LINK:
                            addr_map['mac'] = addr.address
                        elif hasattr(socket, 'AF_PACKET') and addr.family == socket.AF_PACKET:
                             addr_map['mac'] = addr.address

                    # Join multiple IPs, or display N/A if none were found
                    ipv4_str = ", ".join(addr_map['ipv4']) or "N/A"
                    ipv6_str = ", ".join(addr_map['ipv6']) or "N/A"

                    iface_layout.addRow(QLabel("<b>IPv4 Address:</b>"), QLabel(ipv4_str))
                    iface_layout.addRow(QLabel("<b>IPv6 Address:</b>"), QLabel(ipv6_str))
                    iface_layout.addRow(QLabel("<b>MAC Address:</b>"), QLabel(addr_map['mac']))

                    net_main_layout.addWidget(iface_box)
        except Exception as e:
            logging.error(f"Could not retrieve network interfaces: {e}", exc_info=True)
            net_main_layout.addWidget(QLabel(f"Could not retrieve interfaces: {e}"))

        net_layout.addRow(net_content_widget)
        main_layout.addWidget(net_box)

        # --- Refresh Intervals Box ---
        refresh_box, refresh_layout = create_info_box("Refresh Intervals", "refresh-cw.svg")

        self.sys_refresh_combo = QComboBox()
        self.sys_refresh_combo.addItems(["1s", "2s", "5s", "10s", "Off"])
        self.sys_refresh_combo.setToolTip("Set the refresh interval for the CPU, RAM, and Disk monitors.")
        self.sys_refresh_combo.textActivated.connect(self._handle_refresh_interval_change)
        refresh_layout.addRow("System Monitor:", self.sys_refresh_combo)

        self.ip_refresh_combo = QComboBox()
        self.ip_refresh_combo.addItems(["5s", "10s", "30s", "60s", "Off"])
        self.ip_refresh_combo.setToolTip("Set the refresh interval for the public IP display.")
        self.ip_refresh_combo.textActivated.connect(self._handle_ip_refresh_change)
        refresh_layout.addRow("Public IP Display:", self.ip_refresh_combo)

        main_layout.addWidget(refresh_box)

        main_layout.addStretch() # Push everything to the top
        return scroll_area

    def _create_wireless_tools_tab(self, p=None):
        """Creates the tab container for 802.11 wireless tools."""
        wireless_tabs = QTabWidget()
        wireless_tabs.addTab(self._create_wifi_scanner_tool(), "Wi-Fi Scanner")
        wireless_tabs.addTab(self._create_deauth_tool(), "Deauthentication Tool")
        wireless_tabs.addTab(self._create_beacon_flood_tool(), "Beacon Flood")
        wireless_tabs.addTab(self._create_wpa_crack_tool(), "WPA Handshake Tool")
        wireless_tabs.addTab(self._create_krack_scanner_tool(), "KRACK Scanner")
        wireless_tabs.addTab(self._create_wifite_tool(), "Wifite Auditor")
        return wireless_tabs

    def _require_root(self, tool_name):
        """
        Checks for root privileges and shows a message box if not available. Returns True if root, False otherwise.

        Note on Scapy/Root: Unlike external command-line tools which can be individually prefixed with 'sudo'
        and receive a password via stdin, Scapy's functions (like sniff, srp) are called directly within
        the running Python process. Elevating privileges for these specific function calls on-the-fly is not
        feasible without a major architectural change (e.g., moving all Scapy calls to separate helper scripts).
        Therefore, this function takes the pragmatic approach of checking for root and, if not present, informing
        the user of the requirement to restart the whole application with sudo. This prevents crashes and clearly
        communicates the necessary action to the user.
        """
        if sys.platform == "win32":
            # On Windows, admin is handled differently and often not required for Scapy networking
            return True
        try:
            if os.geteuid() == 0:
                return True
        except AttributeError:
            # os.geteuid() doesn't exist on Windows, so we can consider it "ok"
            return True

        QMessageBox.critical(self, "Root Privileges Required",
            f"The '{tool_name}' tool uses low-level networking functions that require root privileges.\n\n"
            "To use this tool, please restart Zurvan from your terminal using 'sudo':\n"
            "$ sudo python3 zurvan.py")
        return False

    def _create_krack_scanner_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>KRACK Vulnerability Scanner</b></font>
        <p>This tool passively detects networks vulnerable to Key Reinstallation Attacks (KRACK). It works by listening for retransmitted EAPOL Message 3 packets during a 4-way handshake.</p>
        <p><b>Usage:</b></p>
        <ol>
            <li>Ensure your wireless card is in <b>Monitor Mode</b> and select it at the top.</li>
            <li>Click "Start Scan". The tool will listen indefinitely.</li>
            <li>To trigger a handshake, you can use the Deauthentication Tool to briefly disconnect a client, forcing it to reconnect.</li>
            <li>Any vulnerable networks detected will appear in the results table below.</li>
        </ol>
        """)
        layout.addWidget(instructions)

        controls = QHBoxLayout()
        self.krack_start_btn = QPushButton("Start Scan")
        self.krack_stop_btn = QPushButton("Stop Scan"); self.krack_stop_btn.setEnabled(False)
        controls.addWidget(self.krack_start_btn)
        controls.addWidget(self.krack_stop_btn)
        layout.addLayout(controls)

        self.krack_results_tree = QTreeWidget()
        self.krack_results_tree.setColumnCount(3)
        self.krack_results_tree.setHeaderLabels(["BSSID (AP)", "Client MAC", "Time Detected"])
        layout.addWidget(self.krack_results_tree)

        self.krack_start_btn.clicked.connect(self.start_krack_scan)
        self.krack_stop_btn.clicked.connect(self.stop_krack_scan)

        return widget

    def _create_wifite_config_widget(self):
        """Creates a reusable, self-contained widget with Wifite's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)
        controls = {}

        # --- Main Tab Widget for Options ---
        wifite_tabs = QTabWidget()
        main_layout.addWidget(wifite_tabs)

        # --- Settings Tab ---
        settings_tab = QWidget()
        settings_layout = QFormLayout(settings_tab)

        controls['verbose_check'] = QCheckBox("Show more options & outputs (-v)")
        settings_layout.addRow("Verbose:", controls['verbose_check'])

        controls['channel_edit'] = QLineEdit()
        controls['channel_edit'].setPlaceholderText("e.g., 1,3-6 (default: all 2Ghz)")
        settings_layout.addRow("Scan Channel (-c):", controls['channel_edit'])

        controls['infinite_check'] = QCheckBox("Enable infinite attack mode (-inf)")
        settings_layout.addRow("Infinite Mode:", controls['infinite_check'])

        controls['random_mac_check'] = QCheckBox("Randomize MAC address (-mac)")
        settings_layout.addRow("Random MAC:", controls['random_mac_check'])

        controls['scan_time_edit'] = QLineEdit()
        controls['scan_time_edit'].setPlaceholderText("Attack all targets after X seconds")
        settings_layout.addRow("Pillage Time (-p):", controls['scan_time_edit'])

        controls['kill_check'] = QCheckBox("Kill conflicting processes (--kill)")
        settings_layout.addRow("Kill Processes:", controls['kill_check'])

        controls['power_edit'] = QLineEdit()
        controls['power_edit'].setPlaceholderText("Minimum signal strength")
        settings_layout.addRow("Min Power (--power):", controls['power_edit'])

        controls['skip_crack_check'] = QCheckBox("Skip cracking handshakes (--skip-crack)")
        settings_layout.addRow("Skip Cracking:", controls['skip_crack_check'])

        controls['first_edit'] = QLineEdit()
        controls['first_edit'].setPlaceholderText("e.g., 5 (attacks first 5 targets)")
        settings_layout.addRow("Attack First X (--first):", controls['first_edit'])

        controls['ignore_cracked_check'] = QCheckBox("Hide previously-cracked targets (-ic)")
        settings_layout.addRow("Ignore Cracked:", controls['ignore_cracked_check'])

        controls['clients_only_check'] = QCheckBox("Only show targets with clients (--clients-only)")
        settings_layout.addRow("Clients Only:", controls['clients_only_check'])

        controls['nodeauths_check'] = QCheckBox("Passive mode, no deauths (--nodeauths)")
        settings_layout.addRow("Passive Mode:", controls['nodeauths_check'])

        controls['daemon_check'] = QCheckBox("Put device back in managed mode after quitting (--daemon)")
        settings_layout.addRow("Daemon Mode:", controls['daemon_check'])

        wifite_tabs.addTab(settings_tab, "Settings")

        # --- WEP Tab ---
        wep_tab = QWidget()
        wep_layout = QFormLayout(wep_tab)
        controls['wep_only_check'] = QCheckBox("Show only WEP networks (--wep)")
        wep_layout.addRow("WEP Only:", controls['wep_only_check'])
        controls['require_fakeauth_check'] = QCheckBox("Fail attacks if fake-auth fails (--require-fakeauth)")
        wep_layout.addRow("Require Fake Auth:", controls['require_fakeauth_check'])
        controls['keep_ivs_check'] = QCheckBox("Retain and reuse .IVS files (--keep-ivs)")
        wep_layout.addRow("Keep IVS files:", controls['keep_ivs_check'])
        wifite_tabs.addTab(wep_tab, "WEP")

        # --- WPA Tab ---
        wpa_tab = QWidget()
        wpa_layout = QFormLayout(wpa_tab)
        controls['wpa_only_check'] = QCheckBox("Show only WPA networks (--wpa)")
        wpa_layout.addRow("WPA Only:", controls['wpa_only_check'])
        controls['new_hs_check'] = QCheckBox("Capture new handshakes, ignore existing (--new-hs)")
        wpa_layout.addRow("New Handshakes Only:", controls['new_hs_check'])

        dict_layout = QHBoxLayout()
        controls['dict_edit'] = QLineEdit()
        controls['dict_edit'].setPlaceholderText("Path to wordlist file...")
        dict_layout.addWidget(controls['dict_edit'])
        browse_dict_btn = QPushButton("Browse...")
        browse_dict_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['dict_edit'], "Select Dictionary File"))
        dict_layout.addWidget(browse_dict_btn)
        wpa_layout.addRow("Password Dictionary (--dict):", dict_layout)
        wifite_tabs.addTab(wpa_tab, "WPA")

        # --- WPS Tab ---
        wps_tab = QWidget()
        wps_layout = QFormLayout(wps_tab)
        controls['wps_networks_check'] = QCheckBox("Show only WPS-enabled networks (--wps)")
        wps_layout.addRow("WPS Networks Only:", controls['wps_networks_check'])
        controls['wps_only_check'] = QCheckBox("Only use WPS PIN & Pixie-Dust attacks (--wps-only)")
        wps_layout.addRow("WPS Attacks Only:", controls['wps_only_check'])

        wps_tool_group = QButtonGroup(self)
        rb_layout = QHBoxLayout()
        controls['reaver_rb'] = QRadioButton("Reaver"); controls['reaver_rb'].setChecked(True)
        controls['bully_rb'] = QRadioButton("Bully")
        wps_tool_group.addButton(controls['reaver_rb'])
        wps_tool_group.addButton(controls['bully_rb'])
        rb_layout.addWidget(controls['reaver_rb']); rb_layout.addWidget(controls['bully_rb'])
        wps_layout.addRow("WPS Tool:", rb_layout)

        controls['ignore_locks_check'] = QCheckBox("Do not stop WPS PIN attack if AP becomes locked (--ignore-locks)")
        wps_layout.addRow("Ignore Locks:", controls['ignore_locks_check'])
        wifite_tabs.addTab(wps_tab, "WPS")

        # --- PMKID Tab ---
        pmkid_tab = QWidget()
        pmkid_layout = QFormLayout(pmkid_tab)
        controls['pmkid_only_check'] = QCheckBox("Only use PMKID capture (--pmkid)")
        pmkid_layout.addRow("PMKID Only:", controls['pmkid_only_check'])
        controls['no_pmkid_check'] = QCheckBox("Don't use PMKID capture (--no-pmkid)")
        pmkid_layout.addRow("No PMKID:", controls['no_pmkid_check'])
        controls['pmkid_timeout_edit'] = QLineEdit("300")
        pmkid_layout.addRow("PMKID Timeout (--pmkid-timeout):", controls['pmkid_timeout_edit'])
        wifite_tabs.addTab(pmkid_tab, "PMKID")

        # --- Action Buttons ---
        controls['start_btn'] = QPushButton(QIcon(self.icon_path("wifi.svg")), " Start Wifite Scan")
        controls['stop_btn'] = QPushButton("Stop Wifite"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _create_wifite_tool(self):
        """Creates the UI for the Wifite automated wireless auditor."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Instructions and Warning ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Wifite Automated Auditor</b></font>
        <p>This tool runs the Wifite script to automatically audit WPA/WPA2/WPS encrypted networks.</p>
        <p><b>WARNING:</b> This is an active attack tool. You MUST have explicit permission to test the target network. Ensure your wireless card is in <b>Monitor Mode</b> and select it from the main interface dropdown.</p>
        """)
        layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.wifite_controls = self._create_wifite_config_widget()
        layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.wifite_controls['start_btn'])
        buttons_layout.addWidget(self.wifite_controls['stop_btn'])
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.wifite_output_console = QPlainTextEdit()
        self.wifite_output_console.setReadOnly(True)
        self.wifite_output_console.setFont(QFont("Courier New", 10))
        self.wifite_output_console.setPlaceholderText("Wifite output will be displayed here...")
        layout.addWidget(self.wifite_output_console, 1) # Add stretch factor

        self.wifite_controls['start_btn'].clicked.connect(self.start_wifite_scan)
        self.wifite_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _build_wifite_command(self, controls):
        """Builds the Wifite command list from a dictionary of controls or widgets."""
        iface = self.get_selected_iface()
        if not iface:
            return None, None, "A monitor-mode wireless interface must be selected."

        tool_path = self._get_tool_path("wifite", "wifite")
        if not tool_path:
             # Try Python script fallback if Wifite is installed as a script
             tool_path = self._get_tool_path("wifite.py", os.path.join("wifite2", "Wifite.py"))
             if tool_path and tool_path.endswith(".py"):
                 command = ["python3", tool_path, "-i", iface]
             elif tool_path:
                  command = [tool_path, "-i", iface]
             else:
                 return None, None, "Wifite not found."
        else:
             command = [tool_path, "-i", iface]

        target_for_log = f"all on {iface}"

        # --- Settings Tab ---
        if self._get_control_value(controls, 'verbose_check', 'check'): command.append("-v")
        if channel := self._get_control_value(controls, 'channel_edit', 'text'): command.extend(["-c", channel.strip()])
        if self._get_control_value(controls, 'infinite_check', 'check'): command.append("-inf")
        if self._get_control_value(controls, 'random_mac_check', 'check'): command.append("-mac")
        if scan_time := self._get_control_value(controls, 'scan_time_edit', 'text'): command.extend(["-p", scan_time.strip()])
        if self._get_control_value(controls, 'kill_check', 'check'): command.append("--kill")
        if power := self._get_control_value(controls, 'power_edit', 'text'): command.extend(["--power", power.strip()])
        if self._get_control_value(controls, 'skip_crack_check', 'check'): command.append("--skip-crack")
        if first := self._get_control_value(controls, 'first_edit', 'text'): command.extend(["--first", first.strip()])
        if self._get_control_value(controls, 'ignore_cracked_check', 'check'): command.append("-ic")
        if self._get_control_value(controls, 'clients_only_check', 'check'): command.append("--clients-only")
        if self._get_control_value(controls, 'nodeauths_check', 'check'): command.append("--nodeauths")
        if self._get_control_value(controls, 'daemon_check', 'check'): command.append("--daemon")

        # --- WEP Tab ---
        if self._get_control_value(controls, 'wep_only_check', 'check'): command.append("--wep")
        if self._get_control_value(controls, 'require_fakeauth_check', 'check'): command.append("--require-fakeauth")
        if self._get_control_value(controls, 'keep_ivs_check', 'check'): command.append("--keep-ivs")

        # --- WPA Tab ---
        if self._get_control_value(controls, 'wpa_only_check', 'check'): command.append("--wpa")
        if self._get_control_value(controls, 'new_hs_check', 'check'): command.append("--new-hs")
        if dictionary := self._get_control_value(controls, 'dict_edit', 'text'): command.extend(["--dict", dictionary.strip()])

        # --- WPS Tab ---
        if self._get_control_value(controls, 'wps_networks_check', 'check'): command.append("--wps")
        if self._get_control_value(controls, 'wps_only_check', 'check'): command.append("--wps-only")
        if self._get_control_value(controls, 'bully_rb', 'check'): command.append("--bully")
        if self._get_control_value(controls, 'reaver_rb', 'check'): command.append("--reaver")
        if self._get_control_value(controls, 'ignore_locks_check', 'check'): command.append("--ignore-locks")

        # --- PMKID Tab ---
        if self._get_control_value(controls, 'pmkid_only_check', 'check'): command.append("--pmkid")
        if self._get_control_value(controls, 'no_pmkid_check', 'check'): command.append("--no-pmkid")
        if timeout := self._get_control_value(controls, 'pmkid_timeout_edit', 'text'): command.extend(["--pmkid-timeout", timeout.strip()])

        return command, target_for_log, None

    def start_wifite_scan(self, sudo_password=None):
        """Starts the Wifite scan worker thread, prompting for sudo if necessary."""
        if self.is_tool_running and not sudo_password:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.wifite_controls
        command, target_for_log, error = self._build_wifite_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "Wifite Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        # Wifite always needs root on non-Windows systems.
        if not sudo_password and sys.platform != "win32":
            if 'wifite_scan' not in self.sudo_cancel_handlers:
                self.sudo_cancel_handlers['wifite_scan'] = lambda: (
                    controls['start_btn'].setEnabled(True),
                    controls['stop_btn'].setEnabled(False)
                )
            self._run_command_with_sudo_prompt(command, self.start_wifite_scan, 'wifite_scan')
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.wifite_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'wifite_scan', target_for_log, self.wifite_output_console, sudo_password
        ))
        self.active_threads.append(self.worker)
        self.worker.start()


    def _create_wifi_scanner_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget)
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setStyleSheet("background-color: #3c3c3c; color: #f0f0f0; border: 1px solid #555;")
        instructions.setHtml("""
        <font color='#ffcc00'><b>WARNING:</b> Wireless tools require the selected interface to be in <b>Monitor Mode</b>.</font>
        <p>Zurvan cannot enable this mode for you. You must do it manually before scanning.</p>
        <p><b>Example for Linux (using airmon-ng):</b></p>
        <ol>
            <li>Find your interface: <code>iwconfig</code> (e.g., wlan0)</li>
            <li>Start monitor mode: <code>sudo airmon-ng start wlan0</code></li>
            <li>A new interface (e.g., wlan0mon) will be created.</li>
            <li><b>Select the new monitor interface (e.g., wlan0mon) from the dropdown at the top of the Zurvan window.</b></li>
        </ol>
        """)
        layout.addWidget(instructions)
        controls = QHBoxLayout()
        self.wifi_scan_button = QPushButton("Scan for Wi-Fi Networks")
        self.wifi_scan_button.setToolTip("Scans for nearby Wi-Fi networks.\nThe selected interface must be in monitor mode.")
        controls.addWidget(self.wifi_scan_button)
        self.wifi_scan_stop_button = QPushButton("Stop Scan")
        self.wifi_scan_stop_button.setToolTip("Stops the current Wi-Fi scan.")
        self.wifi_scan_stop_button.setEnabled(False)
        controls.addWidget(self.wifi_scan_stop_button)
        self.wifi_scan_status = QLabel(""); controls.addWidget(self.wifi_scan_status); controls.addStretch()
        layout.addLayout(controls)
        self.wifi_tree = QTreeWidget(); self.wifi_tree.setColumnCount(4); self.wifi_tree.setHeaderLabels(["SSID", "BSSID", "Channel", "Signal"])
        layout.addWidget(self.wifi_tree)
        layout.addWidget(self._create_export_button(self.wifi_tree))
        self.wifi_scan_button.clicked.connect(self.start_wifi_scan)
        self.wifi_scan_stop_button.clicked.connect(self.stop_wifi_scan)
        return widget

    def _create_wpa_crack_tool(self):
        """Creates the UI for the WPA Handshake and Cracking tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Handshake Capture Section ---
        capture_box = QGroupBox("WPA Handshake Capture")
        capture_layout = QVBoxLayout(capture_box)

        target_layout = QHBoxLayout()
        target_layout.addWidget(QLabel("Target BSSID:"))
        self.wpa_target_combo = QComboBox(); self.wpa_target_combo.setToolTip("Select a target network from the list discovered by the Wi-Fi Scanner.")
        target_layout.addWidget(self.wpa_target_combo)
        refresh_btn = QPushButton("Refresh List"); refresh_btn.setToolTip("Update the list of targets from the Wi-Fi Scanner tab.")
        refresh_btn.clicked.connect(self._refresh_wpa_targets)
        target_layout.addWidget(refresh_btn)
        capture_layout.addLayout(target_layout)

        capture_controls = QHBoxLayout()
        self.wpa_capture_btn = QPushButton("Start Handshake Capture"); self.wpa_capture_btn.setToolTip("Begin sniffing for a WPA handshake from the selected target.")
        capture_controls.addWidget(self.wpa_capture_btn)
        self.wpa_deauth_client_btn = QPushButton("Deauth Client to Speed Up"); self.wpa_deauth_client_btn.setToolTip("Send deauthentication packets to the network to encourage a client to reconnect, speeding up handshake capture.")
        capture_controls.addWidget(self.wpa_deauth_client_btn)
        capture_layout.addLayout(capture_controls)

        self.wpa_capture_status = QLabel("Status: Idle")
        capture_layout.addWidget(self.wpa_capture_status)
        layout.addWidget(capture_box)

        # --- Hash Cracker Section ---
        cracker_box = QGroupBox("WPA Hash Cracker")
        cracker_layout = QVBoxLayout(cracker_box)

        pcap_layout = QHBoxLayout()
        pcap_layout.addWidget(QLabel("Handshake File (.pcap):"))
        self.wpa_pcap_edit = QLineEdit(); self.wpa_pcap_edit.setPlaceholderText("Path to .pcap file containing the handshake...")
        self.wpa_pcap_edit.setToolTip("The .pcap file containing the captured WPA handshake.")
        pcap_layout.addWidget(self.wpa_pcap_edit)
        pcap_browse_btn = QPushButton("Browse...")
        pcap_browse_btn.setToolTip("Browse for a .pcap file containing a WPA handshake.")
        pcap_browse_btn.clicked.connect(lambda: self._browse_file_for_lineedit(self.wpa_pcap_edit, "Select Handshake File", "Pcap Files (*.pcap *.pcapng);;All Files (*)"))
        pcap_layout.addWidget(pcap_browse_btn)
        cracker_layout.addLayout(pcap_layout)

        wordlist_layout = QHBoxLayout()
        wordlist_layout.addWidget(QLabel("Wordlist File:"))
        self.wpa_wordlist_edit = QLineEdit(); self.wpa_wordlist_edit.setPlaceholderText("Path to wordlist file (or leave blank for default)...")
        self.wpa_wordlist_edit.setToolTip("The wordlist file to use for the dictionary attack. If left blank, a small internal list will be used.")
        wordlist_layout.addWidget(self.wpa_wordlist_edit)
        wordlist_browse_btn = QPushButton("Browse...")
        wordlist_browse_btn.setToolTip("Browse for a wordlist file (.txt).")
        wordlist_browse_btn.clicked.connect(lambda: self._browse_file_for_lineedit(self.wpa_wordlist_edit, "Select Wordlist File", "Text Files (*.txt);;All Files (*)"))
        wordlist_layout.addWidget(wordlist_browse_btn)
        crunch_btn = QPushButton("Generate...")
        crunch_btn.setToolTip("Generate a custom wordlist using Crunch (must be installed).")
        crunch_btn.clicked.connect(self.open_crunch_generator)
        wordlist_layout.addWidget(crunch_btn)
        cracker_layout.addLayout(wordlist_layout)

        cpu_layout = QHBoxLayout()
        cpu_layout.addWidget(QLabel("CPU Threads:"))
        self.wpa_threads_edit = QLineEdit("1"); self.wpa_threads_edit.setToolTip("Number of CPU threads for aircrack-ng to use.")
        cpu_layout.addWidget(self.wpa_threads_edit)
        cpu_layout.addStretch()
        cracker_layout.addLayout(cpu_layout)

        self.wpa_crack_btn = QPushButton("Start Cracking"); self.wpa_crack_btn.setToolTip("Begin the cracking process using aircrack-ng.")
        cracker_layout.addWidget(self.wpa_crack_btn)

        self.wpa_crack_output = QPlainTextEdit(); self.wpa_crack_output.setReadOnly(True)
        self.wpa_crack_output.setPlaceholderText("Aircrack-ng output will be shown here...")
        cracker_layout.addWidget(self.wpa_crack_output)
        layout.addWidget(cracker_box)

        self.wpa_capture_btn.clicked.connect(self.start_handshake_capture)
        self.wpa_deauth_client_btn.clicked.connect(self.deauth_for_handshake)
        self.wpa_crack_btn.clicked.connect(self.start_wpa_crack)

        return widget

    def _refresh_wpa_targets(self):
        self.wpa_target_combo.clear()
        if not self.found_networks:
            QMessageBox.information(self, "No Networks", "No networks found. Please run the Wi-Fi Scanner first.")
            return
        for bssid, info in self.found_networks.items():
            ssid = info[0]
            self.wpa_target_combo.addItem(f"{ssid} ({bssid})", bssid)

    def _browse_save_file_for_lineedit(self, line_edit_widget, dialog_title, filter="All Files (*)"):
        file_path, _ = QFileDialog.getSaveFileName(self, dialog_title, "", filter, options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            line_edit_widget.setText(file_path)

    def _browse_file_for_lineedit(self, line_edit_widget, dialog_title, filter="All Files (*)"):
        file_path, _ = QFileDialog.getOpenFileName(self, dialog_title, "", filter, options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            line_edit_widget.setText(file_path)

    def start_wpa_crack(self):
        if self.aircrack_thread and self.aircrack_thread.isRunning():
            self.aircrack_thread.stop()
            return

        pcap_file = self.wpa_pcap_edit.text()
        wordlist = self.wpa_wordlist_edit.text()
        try:
            threads = int(self.wpa_threads_edit.text())
        except ValueError:
            QMessageBox.warning(self, "Input Error", "CPU threads must be a valid number.")
            return

        if not pcap_file:
            QMessageBox.warning(self, "Input Error", "Please provide a handshake file.")
            return
        if not os.path.exists(pcap_file):
            QMessageBox.warning(self, "File Error", f"Pcap file not found:\n{pcap_file}")
            return

        if not wordlist:
            QMessageBox.warning(self, "Input Error", "Please provide a wordlist file.")
            return

        if not os.path.exists(wordlist):
            QMessageBox.warning(self, "File Error", f"Wordlist file not found:\n{wordlist}")
            return

        self.wpa_crack_output.clear()
        self.wpa_crack_btn.setText("Stop Cracking")
        self.aircrack_thread = AircrackThread(pcap_file, wordlist, self, threads)
        self.aircrack_thread.output_received.connect(self._process_aircrack_output)
        self.aircrack_thread.finished_signal.connect(self._on_aircrack_finished)
        self.aircrack_thread.start()

    def _process_aircrack_output(self, line):
        self.wpa_crack_output.appendPlainText(line)
        if "KEY FOUND!" in line:
            self.wpa_crack_output.appendPlainText("\n\n---> PASSWORD FOUND! <---")
            self.aircrack_thread.stop()

    def _on_aircrack_finished(self, return_code):
        self.wpa_crack_btn.setText("Start Cracking")
        self.wpa_crack_output.appendPlainText(f"\n--- Process finished with exit code {return_code} ---")

    def open_crunch_generator(self):
        dialog = CrunchDialog(self)
        if dialog.exec():
            values = dialog.get_values()
            min_len, max_len, charset, outfile = values["min"], values["max"], values["charset"], values["outfile"]

            if not all([min_len, max_len, charset, outfile]):
                QMessageBox.warning(self, "Input Error", "All fields are required to generate a wordlist.")
                return

            command = ["crunch", min_len, max_len, charset, "-o", outfile]

            try:
                self.wpa_crack_output.appendPlainText(f"Starting crunch: {' '.join(command)}")

                def run_crunch():
                    try:
                        # Use CREATE_NO_WINDOW flag on Windows to hide the console
                        startupinfo = None
                        if sys.platform == "win32":
                            startupinfo = subprocess.STARTUPINFO()
                            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

                        process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, startupinfo=startupinfo)
                        for line in iter(process.stdout.readline, ''):
                            logging.info(f"[crunch] {line.strip()}")
                        process.wait()
                        self.tool_results_queue.put(('crunch_finished', outfile, process.returncode))
                    except FileNotFoundError:
                        self.tool_results_queue.put(('error', 'Crunch Error', "'crunch' command not found. Please ensure it is installed and in your system's PATH."))
                    except Exception as e:
                        self.tool_results_queue.put(('error', 'Crunch Error', str(e)))

                self.worker = WorkerThread(target=run_crunch)
                self.worker.start()

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to start crunch: {e}")

    def start_handshake_capture(self):
        if not self._require_root("WPA Handshake Capture"):
            return
        if self.is_tool_running:
            self.stop_handshake_capture()
            return

        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Interface Error", "Please select a monitor-mode interface.")
            return

        bssid = self.wpa_target_combo.currentData()
        if not bssid:
            QMessageBox.warning(self, "Target Error", "Please select a target network.")
            return

        self.is_tool_running = True
        self.wpa_capture_btn.setText("Stop Capture")
        self.handshake_sniffer_thread = HandshakeSnifferThread(iface, bssid)
        self.handshake_sniffer_thread.log_message.connect(lambda msg: self.wpa_capture_status.setText(f"Status: {msg}"))
        self.handshake_sniffer_thread.handshake_captured.connect(self._on_handshake_captured)
        self.handshake_sniffer_thread.start()

    def stop_handshake_capture(self):
        if self.handshake_sniffer_thread and self.handshake_sniffer_thread.isRunning():
            self.handshake_sniffer_thread.stop()
            self.handshake_sniffer_thread.wait()
        self.is_tool_running = False
        self.wpa_capture_btn.setText("Start Handshake Capture")
        self.wpa_capture_status.setText("Status: Idle")

    def _on_handshake_captured(self, bssid, file_path):
        self.wpa_capture_status.setText(f"Status: Handshake for {bssid} captured and saved to {file_path}!")
        self.stop_handshake_capture()
        QMessageBox.information(self, "Success", f"Handshake captured and saved to {file_path}")
        self.wpa_pcap_edit.setText(file_path)

    def deauth_for_handshake(self):
        if not self._require_root("Deauthentication"):
            return
        bssid = self.wpa_target_combo.currentData()
        if not bssid:
            QMessageBox.warning(self, "Target Error", "Please select a target network to deauthenticate.")
            return
        args = (bssid, "ff:ff:ff:ff:ff:ff", 5)
        self.worker = WorkerThread(self._deauth_thread, args=args)
        self.worker.start()
        QMessageBox.information(self, "Deauth Sent", f"Sent 5 deauth packets to the network {bssid} to encourage re-association.")

    def _create_deauth_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget)
        warning_label = QLabel("WARNING: Sending deauthentication packets can disrupt networks you do not own. Use responsibly and only on your own network for testing purposes.")
        warning_label.setStyleSheet("color: #ffcc00;")
        layout.addWidget(warning_label)
        controls = QFrame(); clayout = QVBoxLayout(controls)
        row1 = QHBoxLayout()
        row1.addWidget(QLabel("AP BSSID (MAC):"))
        self.deauth_bssid = QLineEdit("ff:ff:ff:ff:ff:ff")
        self.deauth_bssid.setToolTip("The MAC address (BSSID) of the target Access Point.")
        row1.addWidget(self.deauth_bssid)
        clayout.addLayout(row1)
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Client MAC:"))
        self.deauth_client = QLineEdit("ff:ff:ff:ff:ff:ff")
        self.deauth_client.setToolTip("The MAC address of the client to deauthenticate.\nUse 'ff:ff:ff:ff:ff:ff' to deauthenticate all clients.")
        row2.addWidget(self.deauth_client)
        clayout.addLayout(row2)
        row3 = QHBoxLayout()
        row3.addWidget(QLabel("Count:"))
        self.deauth_count = QLineEdit("10")
        self.deauth_count.setToolTip("The number of deauthentication packets to send.")
        row3.addWidget(self.deauth_count)
        clayout.addLayout(row3)
        self.deauth_button = QPushButton(QIcon("icons/user-minus.svg"), " Send Deauth Packets")
        self.deauth_button.setToolTip("Start sending deauthentication packets.\nWARNING: This will disrupt the target's connection.")
        clayout.addWidget(self.deauth_button)
        self.deauth_status = QLabel(""); clayout.addWidget(self.deauth_status)
        layout.addWidget(controls)
        self.deauth_button.clicked.connect(self.start_deauth)
        return widget

    def _create_beacon_flood_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        warning_label = QLabel("WARNING: Flooding the air with beacon frames can disrupt Wi-Fi networks in the area. Use this tool responsibly and only for legitimate testing purposes.")
        warning_label.setStyleSheet("color: #ffcc00;")
        warning_label.setWordWrap(True)
        layout.addWidget(warning_label)

        controls = QFrame()
        controls.setFrameShape(QFrame.Shape.StyledPanel)
        clayout = QFormLayout(controls)

        # SSID controls
        ssid_layout = QHBoxLayout()
        self.bf_ssid_edit = QLineEdit("TestNet")
        self.bf_ssid_edit.setToolTip("A single SSID, or load multiple from a file.")
        ssid_layout.addWidget(self.bf_ssid_edit)
        self.bf_ssid_from_file_btn = QPushButton("Load from File")
        self.bf_ssid_from_file_btn.setToolTip("Load a list of SSIDs from a .txt file (one per line).")
        ssid_layout.addWidget(self.bf_ssid_from_file_btn)
        clayout.addRow("SSID(s):", ssid_layout)

        self.bf_bssid_edit = QLineEdit("random")
        self.bf_bssid_edit.setToolTip("The BSSID (MAC address) of the fake AP. 'random' will generate a new MAC for each packet.")
        clayout.addRow("BSSID:", self.bf_bssid_edit)

        # Encryption
        self.bf_enc_combo = QComboBox()
        self.bf_enc_combo.addItems(["Open", "WEP", "WPA2-PSK", "WPA3-SAE"])
        self.bf_enc_combo.setToolTip("Select the advertised encryption type for the fake network(s).")
        clayout.addRow("Encryption:", self.bf_enc_combo)

        # Channel
        self.bf_channel_edit = QLineEdit("1")
        self.bf_channel_edit.setToolTip("The 802.11 channel to broadcast the beacons on.")
        clayout.addRow("Channel:", self.bf_channel_edit)

        self.bf_count_edit = QLineEdit("1000")
        self.bf_count_edit.setToolTip("The number of beacon frames to send. Use '0' for an infinite flood.")
        clayout.addRow("Count:", self.bf_count_edit)

        self.bf_interval_edit = QLineEdit("0.1")
        self.bf_interval_edit.setToolTip("The time interval (in seconds) between sending each beacon frame.")
        clayout.addRow("Interval:", self.bf_interval_edit)

        layout.addWidget(controls)

        buttons_layout = QHBoxLayout()
        self.bf_start_button = QPushButton("Start Beacon Flood")
        self.bf_start_button.setToolTip("Begin sending fake beacon frames.")
        buttons_layout.addWidget(self.bf_start_button)

        self.bf_stop_button = QPushButton("Stop Flood")
        self.bf_stop_button.setEnabled(False)
        self.bf_stop_button.setToolTip("Stop the ongoing beacon flood.")
        buttons_layout.addWidget(self.bf_stop_button)
        layout.addLayout(buttons_layout)

        self.bf_status_label = QLabel("Status: Idle")
        layout.addWidget(self.bf_status_label)
        layout.addStretch()

        self.bf_start_button.clicked.connect(self.start_beacon_flood)
        self.bf_stop_button.clicked.connect(self.cancel_tool)
        self.bf_ssid_from_file_btn.clicked.connect(self.load_ssids_for_beacon_flood)

        return widget

    def load_ssids_for_beacon_flood(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select SSID List File", "", "Text Files (*.txt);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    self.bf_ssid_list = [line.strip() for line in f if line.strip()]
                if self.bf_ssid_list:
                    self.bf_ssid_edit.setText(f"Loaded {len(self.bf_ssid_list)} SSIDs from file")
                    self.bf_ssid_edit.setReadOnly(True)
                    logging.info(f"Loaded {len(self.bf_ssid_list)} SSIDs for beacon flood.")
                else:
                    self.bf_ssid_edit.setText("")
                    self.bf_ssid_edit.setReadOnly(False)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load SSID file: {e}")


    # --- Backend Methods: Sniffer ---
    def start_sniffing(self):
        """Starts the packet sniffer thread."""
        if not self._require_root("Packet Sniffer"):
            return
        self.start_sniff_btn.setEnabled(False)
        self.stop_sniff_btn.setEnabled(True)
        self.clear_sniffer_display()
        iface = self.get_selected_iface()
        bpf_filter = self.filter_input.text()

        if self.current_user:
            database.log_activity(
                self.current_user['id'], 'Packet Sniffer', 'Start Sniffing',
                f"Interface: {iface or 'default'}", f"BPF Filter: '{bpf_filter}'"
            )

        # Create the thread manager
        self.sniffer_thread = SnifferThread(iface=iface, bpf_filter=bpf_filter, parent=self)
        # Connect the new thread-safe signal to the reconstruction slot
        self.sniffer_thread.packet_bytes_received.connect(self._handle_packet_bytes)
        self.sniffer_thread.finished.connect(self._on_sniffer_finished)
        self.sniffer_thread.start()
        self.status_bar.showMessage(f"Sniffing on interface: {iface or 'default'}...")

    def _handle_packet_bytes(self, pkt_bytes):
        """Reconstructs a packet from bytes and adds it to a buffer for batch updating."""
        try:
            # Reconstruction is quick, so it's safe to do here.
            packet = Ether(pkt_bytes)
            with self.sniffer_buffer_lock:
                self.sniffer_packet_buffer.append(packet)
        except Exception as e:
            logging.error(f"Failed to reconstruct or buffer packet: {e}")

    def stop_sniffing(self):
        """Signals the packet sniffer thread to stop."""
        if self.sniffer_thread and self.sniffer_thread.isRunning():
            if self.current_user:
                database.log_activity(self.current_user['id'], 'Packet Sniffer', 'Stop Sniffing', self.sniffer_thread.iface or 'default')
            self.stop_sniff_btn.setEnabled(False) # Prevent multiple clicks
            self.status_bar.showMessage("Stopping sniffer...")
            self.sniffer_thread.stop()

    def _on_sniffer_finished(self):
        """Handles cleanup after the sniffer thread has terminated."""
        self.start_sniff_btn.setEnabled(True)
        # The stop button is already disabled in stop_sniffing, but let's ensure it here too
        self.stop_sniff_btn.setEnabled(False)
        self.status_bar.showMessage("Sniffing stopped.")
        self.sniffer_thread = None # Clear the reference to the finished thread

    def _update_sniffer_display(self):
        """Periodically called by a timer to batch-update the sniffer GUI."""
        with self.sniffer_buffer_lock:
            if not self.sniffer_packet_buffer:
                return
            # Quickly swap the buffer and release the lock
            packets_to_add = self.sniffer_packet_buffer
            self.sniffer_packet_buffer = []

        # Now process the packets without holding the lock
        items_to_add = []
        for packet in packets_to_add:
            self.packets_data.append(packet)
            n = len(self.packets_data)
            try:
                pt = f"{time.strftime('%H:%M:%S', time.localtime(packet.time))}.{int(packet.time * 1000) % 1000}"
                src = packet[IP].src if packet.haslayer(IP) else (packet[ARP].psrc if packet.haslayer(ARP) else "N/A")
                dst = packet[IP].dst if packet.haslayer(IP) else (packet[ARP].pdst if packet.haslayer(ARP) else "N/A")
                proto = packet.summary().split('/')[1].strip() if '/' in packet.summary() else "N/A"
                length = len(packet)
                item_data = [str(n), pt, src, dst, proto, str(length)]
            except Exception:
                item_data = [str(n), "Parse Error", "N/A", "N/A", "N/A", "N/A"]

            items_to_add.append(QTreeWidgetItem(item_data))

        self.packet_list_widget.addTopLevelItems(items_to_add)
        self.packet_list_widget.scrollToBottom()


    def add_packet_to_list(self, packet):
        """Callback function to add a sniffed packet to the UI list."""
        self.packets_data.append(packet); n = len(self.packets_data)
        try:
            pt = f"{time.strftime('%H:%M:%S', time.localtime(packet.time))}.{int(packet.time * 1000) % 1000}"
            src = packet[IP].src if packet.haslayer(IP) else (packet[ARP].psrc if packet.haslayer(ARP) else "N/A")
            dst = packet[IP].dst if packet.haslayer(IP) else (packet[ARP].pdst if packet.haslayer(ARP) else "N/A")
            proto = packet.summary().split('/')[1].strip() if '/' in packet.summary() else "N/A"
            length = len(packet)
        except Exception: pt, src, dst, proto, length = "Parse Error", "N/A", "N/A", "N/A", "N/A"
        item = QTreeWidgetItem([str(n), pt, src, dst, proto, str(length)]); self.packet_list_widget.addTopLevelItem(item); self.packet_list_widget.scrollToBottom()

    def display_packet_details(self, current_item, previous_item):
        """Displays the selected packet's details in the tree and hex views."""
        self.packet_details_tree.clear()
        self.packet_hex_view.clear()

        if not current_item:
            return

        try:
            packet_index = int(current_item.text(0)) - 1
            if not (0 <= packet_index < len(self.packets_data)):
                return

            packet = self.packets_data[packet_index]

            # Populate the hex view
            hex_dump = hexdump(packet, dump=True)
            self.packet_hex_view.setText(hex_dump)

            # Populate the details tree
            # We need to keep track of layer names to avoid duplicates from scapy's perspective
            layer_counts = {}
            current_layer = packet
            while current_layer:
                layer_name_raw = current_layer.name
                if layer_name_raw in layer_counts:
                    layer_counts[layer_name_raw] += 1
                    layer_name = f"{layer_name_raw} #{layer_counts[layer_name_raw]}"
                else:
                    layer_counts[layer_name_raw] = 1
                    layer_name = layer_name_raw

                layer_item = QTreeWidgetItem([layer_name])
                self.packet_details_tree.addTopLevelItem(layer_item)

                for field in current_layer.fields_desc:
                    field_name = field.name
                    try:
                        val = current_layer.getfieldval(field_name)
                        # i2repr is the standard Scapy way to get a display-friendly representation
                        display_value = field.i2repr(current_layer, val)
                    except Exception as e:
                        # Log the actual error for debugging, but still show a user-friendly message
                        logging.warning(f"Could not display field '{field_name}': {e}")
                        display_value = "Error reading value"

                    field_item = QTreeWidgetItem([field_name, display_value])
                    layer_item.addChild(field_item)

                layer_item.setExpanded(True)
                current_layer = current_layer.payload

            self.packet_details_tree.resizeColumnToContents(0)

        except (ValueError, IndexError):
            self.packet_details_tree.addTopLevelItem(QTreeWidgetItem(["Error displaying packet details."]))
        except Exception as e:
            logging.error(f"Unexpected error in display_packet_details: {e}", exc_info=True)
            self.packet_details_tree.addTopLevelItem(QTreeWidgetItem([f"Error: {e}"]))

    def clear_sniffer_display(self):
        self.packet_list_widget.clear(); self.packet_details_tree.clear(); self.packet_hex_view.clear(); self.packets_data.clear(); logging.info("Sniffer display cleared.")

    def save_packets(self):
        """Saves captured packets to a pcap file."""
        if not self.packets_data:
            QMessageBox.information(self, "Info", "There are no packets to save.")
            return
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Packets", "", "Pcap Files (*.pcap *.pcapng);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            try:
                wrpcap(file_path, self.packets_data)
                self.status_bar.showMessage(f"Saved {len(self.packets_data)} packets to {file_path}")
                if self.current_user:
                    database.log_activity(self.current_user['id'], 'File Operation', 'Save Packets', file_path, f"Saved {len(self.packets_data)} packets.")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save packets: {e}")

    def load_packets(self):
        """Loads packets from a pcap file into the sniffer view."""
        if self.packets_data and QMessageBox.question(self, "Confirm", "Clear captured packets?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.No:
            return
        self.clear_sniffer_display()
        file_path, _ = QFileDialog.getOpenFileName(self, "Load Packets", "", "Pcap Files (*.pcap *.pcapng);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            try:
                loaded_packets = rdpcap(file_path)
                for packet in loaded_packets:
                    self.add_packet_to_list(packet)
                self.status_bar.showMessage(f"Loaded {len(loaded_packets)} packets from {file_path}")
                if self.current_user:
                    database.log_activity(self.current_user['id'], 'File Operation', 'Load Packets', file_path, f"Loaded {len(loaded_packets)} packets.")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load packets: {e}")

    def crafter_add_layer(self):
        """Adds a new protocol layer to the packet being crafted."""
        proto_name = self.proto_to_add.currentText()
        if proto_name in AVAILABLE_PROTOCOLS:
            self.packet_layers.append(AVAILABLE_PROTOCOLS[proto_name]())
            self.crafter_rebuild_layer_list(); self.layer_list_widget.setCurrentRow(len(self.packet_layers) - 1)

    def crafter_remove_layer(self):
        """Removes the selected protocol layer from the packet."""
        if (row := self.layer_list_widget.currentRow()) >= 0:
            del self.packet_layers[row]; self.crafter_rebuild_layer_list(); self.crafter_clear_fields_display()

    def crafter_toggle_fuzz_layer(self):
        """Toggles fuzzing on the selected layer."""
        row = self.layer_list_widget.currentRow()
        if row < 0:
            QMessageBox.information(self, "Info", "Please select a layer to fuzz/unfuzz.")
            return

        layer = self.packet_layers[row]

        # Use hasattr to reliably check for fuzzed layers (duck typing)
        if hasattr(layer, 'obj'):
            # It's already fuzzed, so unfuzz it by replacing it with its original object
            self.packet_layers[row] = layer.obj
        else:
            # It's a normal layer, so wrap it with fuzz()
            self.packet_layers[row] = fuzz(layer)

        self.crafter_rebuild_layer_list()
        self.layer_list_widget.setCurrentRow(row)

    def crafter_rebuild_layer_list(self):
        """Updates the UI list of layers from the internal self.packet_layers."""
        self.layer_list_widget.clear()
        for i, layer in enumerate(self.packet_layers):
            if hasattr(layer, 'obj'):
                # Display fuzzed layers differently
                self.layer_list_widget.addItem(f"{i}: Fuzzed({layer.obj.name})")
            else:
                self.layer_list_widget.addItem(f"{i}: {layer.name}")
        self.crafter_update_packet_summary()

    def crafter_load_template(self, name):
        if self.packet_layers and QMessageBox.question(self, "Confirm", "Clear current packet stack?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.No: return
        self.packet_layers = [copy.deepcopy(l) for l in PACKET_TEMPLATES[name]]
        self.crafter_rebuild_layer_list()
        if self.packet_layers: self.layer_list_widget.setCurrentRow(0)

    def crafter_clear_fields_display(self):
        for widget in self.current_field_widgets: widget.deleteLater()
        self.current_field_widgets = []

    def crafter_display_layer_fields(self, row):
        self.crafter_clear_fields_display()
        if not (0 <= row < len(self.packet_layers)): return

        layer = self.packet_layers[row]

        if hasattr(layer, 'obj'):
            self.scroll_area.setEnabled(False)
            label = QLabel("Fields are not editable for fuzzed layers.")
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.fields_layout.addWidget(label)
            self.current_field_widgets.append(label)
            return

        self.scroll_area.setEnabled(True)
        for field in layer.fields_desc:
            container = QWidget(); hbox = QHBoxLayout(container); hbox.setContentsMargins(0,0,0,0)
            hbox.addWidget(QLabel(f"{field.name}:"))
            if isinstance(layer, TCP) and field.name == "flags":
                flags_widget = QWidget(); flags_layout = QHBoxLayout(flags_widget)
                self.tcp_flag_vars = {}
                for flag in "FSRPAUEC":
                    var = QCheckBox(flag); self.tcp_flag_vars[flag] = var
                    if flag in str(layer.flags): var.setChecked(True)
                    var.stateChanged.connect(lambda state, l=layer: self.crafter_update_tcp_flags(l))
                    flags_layout.addWidget(var)
                hbox.addWidget(flags_widget)
            else:
                le = QLineEdit(str(getattr(layer, field.name, ''))); le.editingFinished.connect(lambda l=layer, f=field.name, w=le: self.crafter_update_field(l, f, w.text()))
                hbox.addWidget(le)
            self.fields_layout.addWidget(container); self.current_field_widgets.append(container)

    def crafter_update_tcp_flags(self, layer):
        layer.flags = "".join([f for f, v in self.tcp_flag_vars.items() if v.isChecked()])
        self.crafter_update_packet_summary()

    def crafter_update_field(self, layer, field_name, text):
        try: setattr(layer, field_name, text)
        except: pass
        self.crafter_update_packet_summary()

    def build_packet(self):
        if not self.packet_layers: return None

        # Avoid deepcopying fuzz objects, as it can cause crashes.
        layers = []
        for l in self.packet_layers:
            if hasattr(l, 'obj'):
                layers.append(l)  # Use the fuzz object directly
            else:
                layers.append(copy.deepcopy(l))  # Deepcopy standard layers

        if not layers: return None

        pkt = layers[0]
        for i in range(1, len(layers)):
            pkt /= layers[i]
        return pkt

    def crafter_update_packet_summary(self):
        try: pkt = self.build_packet(); summary = pkt.summary() if pkt else "No layers."
        except Exception as e: summary = f"Error: {e}"
        self.crafter_summary.setPlainText(summary)

    def crafter_send_packet(self):
        """Starts the thread to send the crafted packet(s)."""
        if not self._require_root("Packet Crafter"):
            return
        if not self.packet_layers:
            QMessageBox.critical(self, "Error", "No packet layers to build a packet from.")
            return
        try:
            count, interval = int(self.send_count_edit.text()), float(self.send_interval_edit.text())
        except ValueError:
            QMessageBox.critical(self, "Error", "Invalid count or interval.")
            return

        # Log the activity before starting the thread
        if self.current_user:
            pkt = self.build_packet()
            if pkt:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='Packet Crafter',
                    action='Send Packets',
                    target=pkt.dst if pkt.haslayer(IP) else 'N/A',
                    details=f"Summary: {pkt.summary()} | Count: {count}, Interval: {interval}s"
                )

        self.send_results_widget.clear()
        self.send_btn.setEnabled(False)
        self.send_cancel_btn.setEnabled(True)
        self.tool_stop_event.clear()
        self.worker = WorkerThread(self._send_thread, args=(count, interval)); self.worker.start()

    def _send_thread(self, c, i):
        iface = self.get_selected_iface()
        q = self.tool_results_queue
        try:
            ans_list = []
            unans_list = []
            for pkt_num in range(c):
                if self.tool_stop_event.is_set():
                    logging.info("Packet sending cancelled.")
                    break

                pkt = self.build_packet()
                if not pkt:
                    logging.error("Failed to build packet in send thread.")
                    break

                send_receive_func = srp1 if pkt.haslayer(Ether) else sr1
                reply = send_receive_func(pkt, timeout=2, iface=iface, verbose=0)
                if reply:
                    ans_list.append((pkt, reply))
                else:
                    unans_list.append(pkt)
                time.sleep(i)
            q.put(('send_results', ans_list, unans_list))
        except Exception as e:
            logging.error("Send packet failed", exc_info=True)
            q.put(('error', 'Send Error', str(e)))
        finally:
            q.put(('send_finished',))

    def _execute_command_thread(self, command, tool_name, target, output_widget, sudo_password=None):
        """
        A generic worker thread for running external command-line tools.
        It handles sudo password prompting and sends output to the specified widget.
        It also wraps commands with proxychains-ng if the Tor proxy is enabled.
        """
        q = self.tool_results_queue

        # --- Temp File Handling ---
        temp_file_path = None
        if tool_name == 'nmap_scan':
            try:
                with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".xml", encoding='utf-8') as tmp:
                    temp_file_path = tmp.name
                command.extend(["-oX", temp_file_path])
            except Exception as e:
                q.put(('error', 'File Error', f"Could not create temporary file for Nmap report: {e}"))
                q.put(('tool_finished', tool_name, target, "Error creating temp file"))
                return

        # Check if Tor proxy or Proxy Manager is enabled and wrap the command if necessary
        is_tor_enabled = self.tor_proxy_check.isChecked()
        is_proxy_manager_enabled = self.proxy_list_widget.count() > 0

        if (is_tor_enabled or is_proxy_manager_enabled) and shutil.which("proxychains-ng"):
            incompatible_tools = ['masscan_scan']
            if tool_name not in incompatible_tools:
                proxy_conf_path = self._create_proxychains_conf()
                if proxy_conf_path:
                    command = ["proxychains-ng", "-f", proxy_conf_path] + command
                else:
                    q.put(('error', 'Proxy Error', 'Failed to create a proxychains configuration file.'))
                    q.put(('tool_finished', tool_name, target, "Proxy config creation failed"))
                    return

        logging.info(f"Starting {tool_name} with command: {' '.join(command)}")
        q.put((f'{tool_name}_output', f"$ {' '.join(command)}\n\n"))
        full_output = []
        success = False
        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            # Prepend sudo if a password was provided
            if sudo_password:
                command = ["sudo", "-S"] + command

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                                     stdin=subprocess.PIPE, text=True, bufsize=1,
                                     startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                setattr(self, f'{tool_name}_process', process)

            if sudo_password:
                process.stdin.write(sudo_password + '\n')
                process.stdin.flush()

            # Read output line by line
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put((f'{tool_name}_output', "\n\n--- Canceled By User ---\n"))
                    break
                q.put((f'{tool_name}_output', line))
                full_output.append(line)

            process.stdout.close()
            retcode = process.wait()
            success = (retcode == 0)

            # Special handling for tools that produce structured output
            if not self.tool_stop_event.is_set():
                if tool_name == 'nmap_scan' and temp_file_path and os.path.exists(temp_file_path):
                    with open(temp_file_path, 'r', encoding='utf-8') as f:
                        xml_content = f.read()
                    if xml_content:
                        q.put(('nmap_xml_result', xml_content))
                    os.remove(temp_file_path)
                elif tool_name == 'dirsearch_scan' and hasattr(self, 'dirsearch_json_temp_file') and self.dirsearch_json_temp_file and os.path.exists(self.dirsearch_json_temp_file):
                    with open(self.dirsearch_json_temp_file, 'r', encoding='utf-8') as f:
                        json_data = f.read()
                    if json_data:
                        q.put(('dirsearch_results', json_data, target))
                    os.remove(self.dirsearch_json_temp_file)
                    self.dirsearch_json_temp_file = None
                elif tool_name == 'ffuf_scan' and hasattr(self, 'ffuf_json_temp_file') and self.ffuf_json_temp_file and os.path.exists(self.ffuf_json_temp_file):
                    with open(self.ffuf_json_temp_file, 'r', encoding='utf-8') as f:
                        json_data = f.read()
                    if json_data:
                        q.put(('ffuf_results', json_data))
                    os.remove(self.ffuf_json_temp_file)
                    self.ffuf_json_temp_file = None
                elif tool_name == 'enum4linux_ng_scan' and hasattr(self, 'enum4linux_ng_json_temp_file') and self.enum4linux_ng_json_temp_file and os.path.exists(self.enum4linux_ng_json_temp_file):
                    with open(self.enum4linux_ng_json_temp_file, 'r', encoding='utf-8') as f:
                        json_data = f.read()
                    if json_data:
                        q.put(('enum4linux_ng_results', json_data, target))
                    os.remove(self.enum4linux_ng_json_temp_file)
                    self.enum4linux_ng_json_temp_file = None
                elif tool_name == 'dnsrecon_scan' and hasattr(self, 'dnsrecon_json_temp_file') and self.dnsrecon_json_temp_file and os.path.exists(self.dnsrecon_json_temp_file):
                    with open(self.dnsrecon_json_temp_file, 'r', encoding='utf-8') as f:
                        json_data = f.read()
                    if json_data:
                        q.put(('dnsrecon_results', json_data, target))
                    os.remove(self.dnsrecon_json_temp_file)
                    self.dnsrecon_json_temp_file = None
                elif tool_name == 'sherlock_scan' and hasattr(self, 'sherlock_temp_dir') and self.sherlock_temp_dir and os.path.exists(self.sherlock_temp_dir):
                    try:
                        # Find the first CSV file in the temp directory
                        for filename in os.listdir(self.sherlock_temp_dir):
                            if filename.endswith(".csv"):
                                with open(os.path.join(self.sherlock_temp_dir, filename), 'r', encoding='utf-8') as f:
                                    csv_data = f.read()
                                if csv_data:
                                    q.put(('sherlock_results', csv_data, target))
                                break # Just read the first one
                    except Exception as e:
                        logging.error(f"Could not read Sherlock CSV report: {e}")
                    finally:
                        shutil.rmtree(self.sherlock_temp_dir)
                        self.sherlock_temp_dir = None
                # Add other structured output handlers here in the future if needed

        except FileNotFoundError:
            q.put(('error', f'{tool_name.capitalize()} Error', f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"{tool_name} thread error: {e}", exc_info=True)
            q.put(('error', f'{tool_name.capitalize()} Error', str(e)))
        finally:
            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='Tool Execution',
                    action=f"Executed {tool_name.replace('_', ' ').title()}",
                    target=target,
                    details=f"Command: {' '.join(command)}",
                    result='Success' if success else 'Failure'
                )
            q.put(('tool_finished', tool_name, target, "".join(full_output)))
            with self.thread_finish_lock:
                setattr(self, f'{tool_name}_process', None)
            logging.info(f"{tool_name} scan thread finished.")

    def start_traceroute(self):
        """Starts the traceroute worker thread."""
        if not self._require_root("Traceroute"):
            return
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t = self.trace_target.text()
        if not t: QMessageBox.critical(self, "Error", "Please enter a target."); return
        self.trace_button.setEnabled(False)
        self.trace_cancel_button.setEnabled(True)
        self.is_tool_running = True
        self.tool_stop_event.clear()
        self.worker = WorkerThread(self._traceroute_thread, args=(t,)); self.worker.start()

    def _traceroute_thread(self,t):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"Traceroute thread started for target: {t} on iface: {iface}")
        results = []
        try:
            q.put(('trace_status', f"Resolving {t}..."))
            dest_ip = self._resolve_hostname(t)
            q.put(('trace_clear',));
            initial_msg = ("",f"Traceroute to {t} ({dest_ip})","","")
            q.put(('trace_result', initial_msg))
            results.append(initial_msg)

            for i in range(1,30):
                if self.tool_stop_event.is_set():
                    q.put(('trace_status', "Traceroute Canceled."))
                    break
                q.put(('trace_status',f"Sending probe to TTL {i}"))
                pkt=IP(dst=dest_ip,ttl=i)/UDP(dport=33434)
                st=time.time(); reply=sr1(pkt,timeout=2,iface=iface); rtt=(time.time()-st)*1000

                if reply is None:
                    result_tuple = (i,"* * *","Timeout","")
                else:
                    h_ip=reply.src
                    try: h_name,_,_=socket.gethostbyaddr(h_ip)
                    except socket.herror: h_name="Unknown"
                    result_tuple = (i,h_ip,h_name,f"{rtt:.2f}")

                q.put(('trace_result', result_tuple))
                results.append(result_tuple)
                if reply and (reply.type==3 or h_ip==dest_ip):
                    q.put(('trace_status',"Trace Complete."))
                    break
            else: q.put(('trace_status',"Trace Finished (Max hops reached)."))
        except Exception as e:
            logging.error("Exception in traceroute thread",exc_info=True)
            q.put(('error',"Traceroute Error",str(e)))
            success = False
        else:
            success = True
        finally:
            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='Tool Execution',
                    action='Executed Traceroute',
                    target=t,
                    details=f"Hops: {len(results)}",
                    result='Success' if success else 'Failure'
                )
            results_str = "\n".join([f"{hop} - {ip} - {name} - {rtt}ms" for hop, ip, name, rtt in results])
            q.put(('tool_finished','traceroute', t, results_str))
            logging.info("Traceroute thread finished.")

    def start_port_scan(self):
        """Starts the port scanner worker thread."""
        if not self._require_root("Scapy Port Scanner"):
            return
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t=self.scan_target.text(); ps=self.scan_ports.text(); use_frags=self.scan_frag_check.isChecked()

        scan_protocols = []
        if self.scan_proto_tcp_radio.isChecked(): scan_protocols.append("TCP")
        if self.scan_proto_udp_radio.isChecked(): scan_protocols.append("UDP")
        if self.scan_proto_both_radio.isChecked(): scan_protocols.extend(["TCP", "UDP"])

        tcp_scan_type = self.tcp_scan_type_combo.currentText() if self.tcp_scan_type_combo.isVisible() else "SYN Scan"

        if not t or not ps: QMessageBox.critical(self, "Error", "Target and ports required."); return
        try: ports=sorted(list(set(self._parse_ports(ps))))
        except ValueError: QMessageBox.critical(self, "Error","Invalid port format. Use '22, 80, 100-200'."); return

        self.scan_button.setEnabled(False)
        self.scan_cancel_button.setEnabled(True)
        self.is_tool_running=True
        self.tool_stop_event.clear()

        args = (t, ports, scan_protocols, tcp_scan_type, use_frags)
        self.worker = WorkerThread(self._port_scan_thread, args=args); self.worker.start()

    def _parse_ports(self,ps):
        ports=[]
        for part in ps.split(','):
            part=part.strip()
            if '-' in part: start,end=map(int,part.split('-')); ports.extend(range(start,end+1))
            else: ports.append(int(part))
        return ports

    def _port_scan_thread(self,t,ports,scan_protocols,tcp_scan_type,use_frags):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"Port scan started: T={t}, P={ports}, Protocols={scan_protocols}, TCP_Mode={tcp_scan_type}, Frags={use_frags}")
        scan_results = []
        try:
            target_ip = self._resolve_hostname(t)
            q.put(('scan_clear',))
            total_ports = len(ports) * len(scan_protocols)
            ports_scanned = 0

            tcp_scan_flags = {
                "SYN Scan": "S", "FIN Scan": "F", "Xmas Scan": "FPU",
                "Null Scan": "", "ACK Scan": "A"
            }

            for protocol in scan_protocols:
                if self.tool_stop_event.is_set(): break
                for port in ports:
                    if self.tool_stop_event.is_set(): break

                    ports_scanned += 1
                    status_msg = f"Scanning {t}:{port} ({protocol}"
                    if protocol == "TCP": status_msg += f"/{tcp_scan_type}"
                    status_msg += f") - {ports_scanned}/{total_ports}"
                    q.put(('scan_status', status_msg))

                    pkt = None
                    if protocol == "TCP":
                        flags = tcp_scan_flags.get(tcp_scan_type, "S")
                        pkt = IP(dst=t)/TCP(dport=port, flags=flags)
                    elif protocol == "UDP":
                        pkt = IP(dst=t)/UDP(dport=port)

                    if not pkt: continue

                    probes = fragment(pkt) if use_frags else [pkt]
                    # Only need one response, not for every fragment
                    resp=sr1(probes[0] if len(probes) == 1 else probes, timeout=1, iface=iface, verbose=0)
                    state = "No Response / Filtered"
                    if resp:
                        if resp.haslayer(TCP):
                            if resp.getlayer(TCP).flags == 0x12: state = "Open" # SYN-ACK
                            elif resp.getlayer(TCP).flags == 0x14: state = "Closed" # RST-ACK
                            elif resp.getlayer(TCP).flags == 0x4: state = "Unfiltered (RST)" # RST from ACK scan
                        elif resp.haslayer(UDP):
                            state = "Open | Filtered" # UDP is connectionless, open might not respond
                        elif resp.haslayer(ICMP) and resp.getlayer(ICMP).type == 3:
                            if resp.getlayer(ICMP).code in [1, 2, 3, 9, 10, 13]:
                                state = "Filtered"
                            else:
                                state = "Closed (ICMP)"

                    service = "Unknown"
                    if state.startswith("Open"):
                        try: service=socket.getservbyport(port, protocol.lower())
                        except OSError: pass

                    # Add to list for final popup and also to queue for live view
                    result_tuple = (f"{port}/{protocol.lower()}", state, service)
                    scan_results.append(result_tuple)
                    q.put(('scan_result', result_tuple))

            if self.tool_stop_event.is_set():
                q.put(('scan_status', "Scan Canceled."))
            else:
                q.put(('scan_status',"Scan Complete."))
                q.put(('show_port_scan_popup', scan_results, t)) # New message for popup
        except Exception as e:
            logging.error("Exception in port scan thread",exc_info=True)
            q.put(('error',"Scan Error",str(e)))
            success = False
        else:
            success = True
        finally:
            open_ports = [res[0] for res in scan_results if "Open" in res[1]]
            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='Tool Execution',
                    action='Executed Port Scan',
                    target=t,
                    details=f"Scanned {len(ports)} ports. Found {len(open_ports)} open.",
                    result='Success' if success else 'Failure'
                )
            results_str = "\n".join([f"{p} - {s} - {svc}" for p, s, svc in scan_results])
            q.put(('tool_finished','scanner', t, results_str));
            logging.info("Port scan thread finished.")

    def start_arp_scan(self):
        """Starts the ARP scan worker thread."""
        if not self._require_root("Scapy ARP Scan"):
            return
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t=self.arp_target.text()
        if not t: QMessageBox.critical(self, "Error", "Please enter a target network."); return
        self.arp_scan_button.setEnabled(False); self.is_tool_running=True
        self.worker = WorkerThread(self._arp_scan_thread, args=(t,)); self.worker.start()

    def _arp_scan_thread(self,t):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"ARP scan thread started for target: {t} on iface: {iface}")
        try:
            q.put(('arp_status', f"Scanning {t}...")); q.put(('arp_clear',))
            pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/ARP(pdst=t)
            ans,unans=srp(pkt,timeout=2,iface=iface,verbose=0)

            # Keep adding to tree for live results
            answered_results_for_tree = [{'ip': r.psrc, 'mac': r.hwsrc, 'status': 'Responded'} for s, r in ans]
            if answered_results_for_tree:
                q.put(('arp_results', answered_results_for_tree))

            # Now prepare results for popup
            popup_results = []
            q.put(('arp_status', f"Found {len(ans)} hosts. Resolving vendors..."))
            for i, (s, r) in enumerate(ans):
                q.put(('arp_status', f"Resolving vendor for {r.hwsrc} ({i+1}/{len(ans)})"))
                vendor = get_vendor(r.hwsrc)
                popup_results.append({'ip': r.psrc, 'mac': r.hwsrc, 'vendor': vendor})

            total_found = len(ans)
            q.put(('arp_status',f"Scan Complete. Found {total_found} active hosts."))
            q.put(('show_arp_scan_popup', popup_results, t)) # New message for popup

        except Exception as e:
            logging.error("Exception in ARP scan thread",exc_info=True)
            q.put(('error',"ARP Scan Error",str(e)))
            success = False
        else:
            success = True
        finally:
            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='Tool Execution',
                    action='Executed ARP Scan',
                    target=t,
                    details=f"Found {len(ans)} hosts.",
                    result='Success' if success else 'Failure'
                )
            results_str = "\n".join([f"{res['ip']} - {res['mac']} - {res['vendor']}" for res in popup_results])
            q.put(('tool_finished','arp_scan', t, results_str))
            logging.info("ARP scan thread finished.")

    def _create_arp_scan_cli_tool(self):
        """Creates the UI for the arp-scan CLI tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.arp_scan_cli_controls = self._create_arp_scan_cli_config_widget()
        layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.arp_scan_cli_controls['start_btn'])
        buttons_layout.addWidget(self.arp_scan_cli_controls['stop_btn'])
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.arp_scan_cli_output_console = QPlainTextEdit()
        self.arp_scan_cli_output_console.setReadOnly(True)
        self.arp_scan_cli_output_console.setFont(QFont("Courier New", 10))
        self.arp_scan_cli_output_console.setPlaceholderText("arp-scan output will be displayed here...")
        layout.addWidget(self.arp_scan_cli_output_console, 1)

        self.arp_scan_cli_controls['start_btn'].clicked.connect(self.start_arp_scan_cli)
        self.arp_scan_cli_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_arp_scan_cli_config_widget(self):
        """Creates a reusable, self-contained widget with arp-scan's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['localnet_check'] = QCheckBox("Scan Local Network (--localnet)")
        controls['localnet_check'].setChecked(True)
        controls['localnet_check'].setToolTip("Automatically scan the network of the selected interface.")
        main_layout.addRow(controls['localnet_check'])

        controls['target_edit'] = QLineEdit()
        controls['target_edit'].setPlaceholderText("e.g., 192.168.1.0/24 (optional if --localnet is checked)")
        controls['target_edit'].setToolTip("Specify a custom target network or host if not using --localnet.")
        main_layout.addRow("Custom Target:", controls['target_edit'])

        controls['verbose_check'] = QCheckBox("Verbose Output (-v)")
        main_layout.addRow("--verbose:", controls['verbose_check'])

        # UI Logic
        def toggle_target_edit(checked):
            controls['target_edit'].setDisabled(checked)
        controls['localnet_check'].toggled.connect(toggle_target_edit)
        toggle_target_edit(True) # Initial state

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_arp_scan_cli_command(self, controls):
        """Builds the arp-scan CLI command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("arp-scan", "arp-scan")
        if not tool_path:
             # Caller expects 4 values: command, target, error, env
             return None, None, "arp-scan not found."
        command = [tool_path]

        iface = self.get_selected_iface()
        if iface:
            command.extend(["--interface", iface])

        target_for_log = "--localnet"
        if self._get_control_value(controls, 'localnet_check', 'check'):
            command.append("--localnet")
        else:
            target = self._get_control_value(controls, 'target_edit', 'text')
            if not target or not target.strip():
                return None, None, "A custom target is required if --localnet is not checked."
            command.append(target.strip())
            target_for_log = target.strip()

        if self._get_control_value(controls, 'verbose_check', 'check'):
            command.append("--verbose")

        return command, target_for_log, None

    def start_arp_scan_cli(self, sudo_password=None):
        """Starts the arp-scan CLI worker thread, prompting for sudo if necessary."""
        if self.is_tool_running and not sudo_password:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.arp_scan_cli_controls
        # arp-scan works at L2, DNS resolution is not applicable.
        build_result = self._build_arp_scan_cli_command(controls)

        # Handle potential 4-tuple return (legacy or inconsistency)
        if len(build_result) == 4:
             command, target_for_log, error, _ = build_result
        else:
             command, target_for_log, error = build_result

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        if not shutil.which(command[0]):
            QMessageBox.critical(self, "arp-scan Error", f"'{command[0]}' command not found. Please ensure it is installed and in your system's PATH.")
            return

        # arp-scan always needs root on non-Windows systems.
        if not sudo_password and sys.platform != "win32":
            if 'arp_scan_cli_scan' not in self.sudo_cancel_handlers:
                self.sudo_cancel_handlers['arp_scan_cli_scan'] = lambda: (
                    controls['start_btn'].setEnabled(True),
                    controls['stop_btn'].setEnabled(False)
                )
            self._run_command_with_sudo_prompt(command, self.start_arp_scan_cli, 'arp_scan_cli_scan')
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.arp_scan_cli_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'arp_scan_cli_scan', target_for_log, self.arp_scan_cli_output_console, sudo_password
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _handle_arp_scan_cli_output(self, line):
        self.arp_scan_cli_output_console.insertPlainText(line)
        self.arp_scan_cli_output_console.verticalScrollBar().setValue(self.arp_scan_cli_output_console.verticalScrollBar().maximum())

    def start_ping_sweep(self):
        if not self._require_root("Ping Sweep"):
            return
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target_network = self.ps_target_edit.text()
        probe_type = self.ps_probe_type_combo.currentText()
        ports_str = self.ps_ports_edit.text()

        try:
            net = ipaddress.ip_network(target_network, strict=False)
            timeout = float(self.ps_timeout_edit.text())
            num_threads = int(self.ps_threads_edit.text())
            ports = [int(p.strip()) for p in ports_str.split(',')] if ports_str else []
        except ValueError as e:
            QMessageBox.critical(self, "Input Error", f"Invalid input: {e}")
            return

        if ("TCP" in probe_type or "UDP" in probe_type) and not ports:
            QMessageBox.critical(self, "Input Error", "Please specify at least one port for TCP/UDP probes.")
            return

        if self.current_user:
            database.log_activity(
                user_id=self.current_user['id'],
                category='Tool Execution',
                action='Start Ping Sweep',
                target=target_network,
                details=f"Probe: {probe_type}, Ports: {ports_str}, Threads: {num_threads}"
            )

        self.is_tool_running = True
        self.ps_start_button.setEnabled(False)
        self.ps_cancel_button.setEnabled(True)
        self.tool_stop_event.clear()
        self.ps_tree.clear()

        args = (net, probe_type, ports, timeout, num_threads)
        self.worker = WorkerThread(self._ping_sweep_thread, args=args)
        self.worker.start()

    def _ping_sweep_thread(self, net, probe_type, ports, timeout, num_threads):
        """Master thread that populates a queue and starts worker threads."""
        q = self.tool_results_queue
        logging.info(f"Ping sweep started for {net} with {probe_type} on ports {ports}")

        hosts_queue = queue.Queue()
        for host in net.hosts():
            hosts_queue.put(str(host))

        if hosts_queue.qsize() == 0:
            q.put(('ps_status', "Sweep Complete (No hosts in range)."))
            q.put(('tool_finished', 'ping_sweep'))
            return

        self.ps_finished_threads = 0
        self.active_threads = []

        for i in range(num_threads):
            worker = WorkerThread(target=self._ping_sweep_worker, args=(hosts_queue, probe_type, ports, timeout, num_threads))
            self.active_threads.append(worker)
            worker.start()

    def _ping_sweep_worker(self, hosts_queue, probe_type, ports, timeout, num_threads):
        """Worker function that each ping sweep thread executes."""
        q = self.tool_results_queue
        while not self.tool_stop_event.is_set():
            try:
                host_str = hosts_queue.get_nowait()
            except queue.Empty:
                break # Queue is empty, this thread is done

            q.put(('ps_status', f"Pinging {host_str}..."))

            reply = None
            try:
                target_ip = self._resolve_hostname(host_str)
                if probe_type == "ICMP Echo":
                    pkt = IP(dst=target_ip)/ICMP()
                    reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                elif probe_type == "TCP SYN":
                    for port in ports:
                        pkt = IP(dst=target_ip)/TCP(dport=port, flags="S")
                        reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                        if reply and reply.haslayer(TCP) and reply.getlayer(TCP).flags == 0x12: # SYN-ACK
                            break # Host is up, no need to check other ports
                elif probe_type == "TCP ACK":
                    for port in ports:
                        pkt = IP(dst=target_ip)/TCP(dport=port, flags="A")
                        reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                        if reply and reply.haslayer(TCP) and reply.getlayer(TCP).flags == 0x4: # RST
                            break # Host is up, no need to check other ports
                elif probe_type == "UDP Probe":
                    for port in ports:
                        pkt = IP(dst=target_ip)/UDP(dport=port)
                        reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                        if reply and reply.haslayer(ICMP) and reply.getlayer(ICMP).type == 3: # Dest Unreachable
                            break # Port is closed, but host is up.
            except Exception as e:
                logging.warning(f"Probe to {host_str} failed: {e}")


            if reply:
                q.put(('ps_result', (host_str, "Host is up")))

        # Signal that this worker is done
        q.put(('ps_worker_finished', num_threads))

    def load_flood_packet(self):
        packet=self.build_packet()
        if not packet: QMessageBox.critical(self, "Error", "Please craft a packet in the Packet Crafter tab first."); return
        self.loaded_flood_packet=packet
        self.flood_packet_label.setText(f"Loaded: {self.loaded_flood_packet.summary()}")
        logging.info(f"Loaded flood packet: {self.loaded_flood_packet.summary()}")

    def start_flood(self):
        """Starts the packet flooder worker threads."""
        if not self._require_root("Packet Flooder"):
            return
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        template = self.flood_template_combo.currentText()
        if template == "Custom (from Crafter)" and not self.loaded_flood_packet:
            QMessageBox.critical(self, "Error", "Please load a packet from the crafter first.")
            return

        warning_msg = "WARNING: This tool sends a high volume of packets and can disrupt network services. Only use this tool on networks you own or have explicit permission to test. Misuse of this tool may be illegal.\n\nDo you accept responsibility and wish to continue?"
        if not QMessageBox.question(self, "Ethical Use Warning", warning_msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.Yes:
            return

        try:
            count = int(self.flood_count.text())
            interval = float(self.flood_interval.text())
            num_threads = int(self.flood_threads.text())
            target_ip = self.flood_target_edit.text()
            target_port = int(self.flood_ports_edit.text()) if self.flood_ports_edit.text() else 80
            random_source = self.flood_rand_src_ip_check.isChecked()
        except ValueError:
            QMessageBox.critical(self, "Error", "Invalid count, interval, or thread number.")
            return

        if self.current_user:
            details = f"Template: {template}, Count: {count}, Interval: {interval}s, Threads: {num_threads}"
            database.log_activity(
                self.current_user['id'], 'Offensive', 'Start Packet Flood', target_ip, details
            )

        self.flood_button.setEnabled(False)
        self.stop_flood_button.setEnabled(True)
        self.is_tool_running = True
        self.finished_thread_count = 0
        self.active_threads = []
        self.tool_stop_event.clear()

        packets_per_thread = count // num_threads
        extra_packets = count % num_threads

        flood_params = {
            "template": template, "target_ip": target_ip, "target_port": target_port,
            "random_source": random_source, "custom_packet": self.loaded_flood_packet
        }

        for i in range(num_threads):
            count_for_this_thread = packets_per_thread + (1 if i < extra_packets else 0)
            if count_for_this_thread == 0:
                continue

            worker = WorkerThread(self._flood_thread, args=(flood_params, count_for_this_thread, interval, num_threads))
            self.active_threads.append(worker)
            worker.start()

    def _flood_thread(self, params, count, interval, total_threads):
        q = self.tool_results_queue
        iface = self.get_selected_iface()
        logging.info(f"Flood thread started. Params: {params}, Count: {count}")
        try:
            # Resolve hostname at the beginning of the thread
            target_ip_resolved = self._resolve_hostname(params["target_ip"])
            q.put(('flood_status', f"Flooding {target_ip_resolved} with {count} packets..."))
            send_func = sendp # Assume Layer 2 for templates for now

            for i in range(count):
                if self.tool_stop_event.is_set():
                    logging.info("Flood thread detected stop event.")
                    break

                pkt = None
                template = params["template"]

                if template == "Custom (from Crafter)":
                    pkt = params["custom_packet"].copy() # Use a copy to avoid modifying the original
                    if pkt.haslayer(IP):
                        pkt[IP].dst = self._resolve_hostname(pkt[IP].dst)
                    send_func = sendp if pkt.haslayer(Ether) else send
                else:
                    # On-the-fly packet creation for templates
                    src_ip = _get_random_ip() if params["random_source"] else "1.2.3.4" # Dummy IP if not random
                    target_port = params["target_port"]

                    if template == "TCP SYN Flood":
                        pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/IP(src=src_ip, dst=target_ip_resolved) / TCP(sport=RandShort(), dport=target_port, flags="S")
                    elif template == "UDP Flood":
                        pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/IP(src=src_ip, dst=target_ip_resolved) / UDP(sport=RandShort(), dport=target_port) / Raw(load=b"X"*1024)
                    elif template == "ICMP Echo Flood":
                        pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/IP(src=src_ip, dst=target_ip_resolved) / ICMP()

                if pkt:
                    send_func(pkt, iface=iface, verbose=0)

                time.sleep(interval)

        except Exception as e:
            logging.error("Exception in flood thread", exc_info=True)
            q.put(('error', "Flood Error", str(e)))
        finally:
            q.put(('flood_thread_finished', total_threads))
            logging.info("A flood thread finished.")

    def start_krack_scan(self):
        if not self._require_root("KRACK Scanner"):
            return
        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Interface Error", "Please select a monitor-mode interface.")
            return

        if self.current_user:
            database.log_activity(self.current_user['id'], 'Wireless', 'Start KRACK Scan', f"Interface: {iface}")

        self.krack_start_btn.setEnabled(False)
        self.krack_stop_btn.setEnabled(True)
        self.krack_results_tree.clear()

        self.krack_thread = KrackScanThread(iface, self)
        self.krack_thread.vulnerability_detected.connect(self.add_krack_result)
        self.krack_thread.start()

    def stop_krack_scan(self):
        if self.krack_thread and self.krack_thread.isRunning():
            if self.current_user:
                database.log_activity(self.current_user['id'], 'Wireless', 'Stop KRACK Scan', self.krack_thread.iface)
            self.krack_thread.stop()
            self.krack_thread.wait()
        self.krack_start_btn.setEnabled(True)
        self.krack_stop_btn.setEnabled(False)

    def add_krack_result(self, bssid, client_mac):
        # Avoid adding duplicates
        for i in range(self.krack_results_tree.topLevelItemCount()):
            item = self.krack_results_tree.topLevelItem(i)
            if item.text(0) == bssid and item.text(1) == client_mac:
                return # Already exists

        timestamp = time.strftime('%H:%M:%S')
        item = QTreeWidgetItem([bssid, client_mac, timestamp])
        self.krack_results_tree.addTopLevelItem(item)

    def start_firewall_test(self):
        """Starts the firewall testing worker thread."""
        if not self._require_root("Firewall Tester"):
            return
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t=self.fw_target.text(); ps_name=self.fw_probe_set.currentText()
        if not t: QMessageBox.critical(self, "Error", "Please enter a target."); return
        self.fw_test_button.setEnabled(False); self.is_tool_running=True
        self.worker = WorkerThread(self._firewall_test_thread, args=(t,ps_name)); self.worker.start()

    def _firewall_test_thread(self,t,ps_name):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"Firewall test thread started for target {t}, probe set {ps_name}")
        results = []
        try:
            target_ip = self._resolve_hostname(t)
            q.put(('fw_clear',)); q.put(('fw_status',f"Testing {ps_name}..."))
            probe_set = FIREWALL_PROBES[ps_name]
            for i, (pkt_builder, desc) in enumerate(probe_set):
                q.put(('fw_status',f"Sending probe {i+1}/{len(probe_set)}: {desc}"))

                pkt = pkt_builder(target_ip)
                pkt_summary = ""

                if isinstance(pkt, list): # It's a fragmented packet
                    pkt_summary = f"{len(pkt)} fragments"
                    ans, unans = sr(pkt, timeout=2, iface=iface, verbose=0)
                    resp = ans[0][1] if ans else None # Take the first response as representative
                else: # It's a single packet
                    pkt_summary = pkt.summary()
                    resp = sr1(pkt, timeout=2, iface=iface, verbose=0)

                result = "Responded" if resp is not None else "No Response / Blocked"
                result_tuple = (desc, pkt_summary, result)
                q.put(('fw_result', result_tuple))
                results.append(result_tuple)
            q.put(('fw_status',"Firewall Test Complete."))
        except Exception as e: logging.error("Exception in firewall test thread",exc_info=True); q.put(('error',"Firewall Test Error",str(e)))
        finally:
            results_str = "\n".join([f"{desc} - {summary} - {res}" for desc, summary, res in results])
            q.put(('tool_finished','fw_tester', t, results_str))
            logging.info("Firewall test thread finished.")

    def start_wifi_scan(self):
        """Starts the Wi-Fi scanner and channel hopper threads."""
        if not self._require_root("Wi-Fi Scanner"):
            return
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return

        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Warning", "Please select a wireless interface for scanning.")
            return

        if self.current_user:
            database.log_activity(self.current_user['id'], 'Wireless', 'Start Wi-Fi Scan', f"Interface: {iface}")

        self.wifi_scan_button.setEnabled(False)
        self.wifi_scan_stop_button.setEnabled(True)
        self.is_tool_running = True
        self.found_networks = {}
        self.wifi_tree.clear()

        self.sniffer_thread = SnifferThread(iface=iface, handler=self._wifi_scan_handler, bpf_filter="type mgt subtype beacon or type mgt subtype probe-resp")
        self.sniffer_thread.start()
        self.channel_hopper = ChannelHopperThread(iface)
        self.channel_hopper.start()

        self.tool_results_queue.put(('wifi_scan_status', 'Scanning... Press Stop to finish.'))

        # We can still have a timeout as a safeguard, but the user can now stop it.
        self.scan_timer = QTimer(self)
        self.scan_timer.setSingleShot(True)
        self.scan_timer.timeout.connect(self.stop_wifi_scan)
        self.scan_timer.start(30000) # 30 second safeguard timer

    def _wifi_scan_handler(self, pkt):
        if pkt.haslayer(Dot11Beacon) or pkt.haslayer(Dot11ProbeResp):
            bssid = pkt[Dot11].addr2
            if bssid not in self.found_networks:
                try: ssid = pkt[Dot11Elt].info.decode(errors="ignore")
                except: ssid = "<Hidden>"
                if not ssid: ssid = "<Hidden>"

                channel = "N/A"
                try:
                    elt = pkt.getlayer(Dot11Elt, ID=3)
                    if elt: channel = ord(elt.info)
                except: pass
                signal = "N/A"
                try: signal = pkt[RadioTap].dbm_antsignal
                except: pass
                self.found_networks[bssid] = (ssid, bssid, channel, signal)
                self.tool_results_queue.put(('wifi_scan_update', self.found_networks[bssid]))

    def stop_wifi_scan(self):
        if hasattr(self, 'scan_timer') and self.scan_timer.isActive():
            self.scan_timer.stop()

        if self.sniffer_thread and self.sniffer_thread.isRunning():
            self.sniffer_thread.stop()
            self.sniffer_thread.wait()
        if self.channel_hopper and self.channel_hopper.isRunning():
            self.channel_hopper.stop()
            self.channel_hopper.wait()

        self.tool_results_queue.put(('wifi_scan_status', 'Scan Finished.'))
        self.tool_results_queue.put(('tool_finished', 'wifi_scan'))

    def start_deauth(self):
        if not self._require_root("Deauthentication Tool"):
            return
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        bssid = self.deauth_bssid.text(); client = self.deauth_client.text()
        try: count = int(self.deauth_count.text())
        except ValueError: QMessageBox.critical(self, "Error", "Count must be an integer."); return
        warning_msg="This will send deauthentication packets which can disrupt a network. Are you sure you want to continue?"
        if QMessageBox.question(self, "Confirm Deauth", warning_msg) == QMessageBox.StandardButton.No: return
        self.deauth_button.setEnabled(False); self.is_tool_running = True
        args = (bssid, client, count)
        self.worker = WorkerThread(self._deauth_thread, args=args); self.worker.start()

    def _deauth_thread(self, bssid, client, count):
        q = self.tool_results_queue; iface = self.get_selected_iface()
        logging.info(f"Deauth thread started: BSSID={bssid}, Client={client}, Count={count}")
        try:
            pkt = RadioTap()/Dot11(type=0, subtype=12, addr1=client, addr2=bssid, addr3=bssid)/Dot11Deauth(reason=7)
            q.put(('deauth_status', f"Sending {count} deauth packets..."))
            sendp(pkt, iface=iface, count=count, inter=0.1, verbose=0)
            q.put(('deauth_status', "Deauth packets sent."))
        except Exception as e:
            logging.error("Exception in deauth thread",exc_info=True)
            q.put(('error',"Deauth Error",str(e)))
            success = False
        else:
            success = True
        finally:
            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='Tool Execution',
                    action='Executed Deauthentication Attack',
                    target=bssid,
                    details=f"Client: {client}, Count: {count}",
                    result='Success' if success else 'Failure'
                )
            q.put(('tool_finished','deauth', bssid, f"Sent {count} deauths to {client}"))
            logging.info("Deauth thread finished.")

    def start_beacon_flood(self):
        if not self._require_root("Beacon Flood"):
            return
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Interface Error", "Please select a monitor-mode interface.")
            return

        bssid = self.bf_bssid_edit.text()
        enc_type = self.bf_enc_combo.currentText()

        # Handle SSIDs
        ssids = []
        if self.bf_ssid_edit.isReadOnly(): # Loaded from file
            ssids = self.bf_ssid_list
        else:
            ssids = [self.bf_ssid_edit.text().strip()]

        if not ssids or not ssids[0]:
            QMessageBox.critical(self, "Input Error", "Please provide at least one SSID.")
            return

        try:
            count = int(self.bf_count_edit.text())
            interval = float(self.bf_interval_edit.text())
            channel = int(self.bf_channel_edit.text())
            if not (1 <= channel <= 14):
                raise ValueError("Channel must be between 1 and 14.")
        except ValueError as e:
            QMessageBox.critical(self, "Input Error", f"Invalid input for Count, Interval, or Channel: {e}")
            return

        self.is_tool_running = True
        self.bf_start_button.setEnabled(False)
        self.bf_stop_button.setEnabled(True)
        self.tool_stop_event.clear()

        args = (iface, ssids, bssid, count, interval, enc_type, channel)
        self.worker = WorkerThread(self._beacon_flood_thread, args=args)
        self.worker.start()

    def _build_beacon_frame(self, ssid, bssid, channel, enc_type):
        dot11 = Dot11(type=0, subtype=8, addr1='ff:ff:ff:ff:ff:ff', addr2=bssid, addr3=bssid)

        cap = 'ESS'
        if enc_type != "Open":
            cap += '+privacy'

        beacon = Dot11Beacon(cap=cap)
        essid = Dot11Elt(ID='SSID', info=ssid)
        ds_param = Dot11Elt(ID='DSset', info=chr(channel))

        frame = RadioTap() / dot11 / beacon / essid / ds_param

        if enc_type == "WEP":
            # WEP is signaled by the privacy bit in the capability field alone.
            pass
        elif enc_type == "WPA2-PSK":
            rsn_info = Dot11Elt(ID='RSNinfo', info=(
                b'\x01\x00'      # RSN Version 1
                b'\x00\x0f\xac\x04'  # Group Cipher Suite: AES (CCMP)
                b'\x01\x00'      # 1 Pairwise Cipher Suite
                b'\x00\x0f\xac\x04'  # AES (CCMP)
                b'\x01\x00'      # 1 Authentication Key Management Suite (AKM)
                b'\x00\x0f\xac\x02'  # PSK
                b'\x00\x00'      # RSN Capabilities
            ))
            frame /= rsn_info
        elif enc_type == "WPA3-SAE":
            rsn_info = Dot11Elt(ID='RSNinfo', info=(
                b'\x01\x00'      # RSN Version 1
                b'\x00\x0f\xac\x04'  # Group Cipher Suite: AES (CCMP)
                b'\x01\x00'      # 1 Pairwise Cipher Suite
                b'\x00\x0f\xac\x04'  # AES (CCMP)
                b'\x01\x00'      # 1 Authentication Key Management Suite (AKM)
                b'\x00\x0f\xac\x08'  # SAE
                b'\x8c\x00'      # RSN Capabilities (MFPC, MFPR)
            ))
            frame /= rsn_info

        return frame

    def _beacon_flood_thread(self, iface, ssids, bssid, count, interval, enc_type, channel):
        q = self.tool_results_queue
        logging.info(f"Beacon flood started: SSIDs={len(ssids)}, BSSID={bssid}, Count={count}, Enc={enc_type}")

        sent_count = 0
        ssid_index = 0
        infinite_mode = (count == 0)

        try:
            while not self.tool_stop_event.is_set():
                if not infinite_mode and sent_count >= count:
                    break

                current_bssid = RandMAC() if bssid.lower() == 'random' else bssid
                current_ssid = ssids[ssid_index]

                beacon_frame = self._build_beacon_frame(current_ssid, current_bssid, channel, enc_type)

                sendp(beacon_frame, iface=iface, verbose=0)
                sent_count += 1
                ssid_index = (ssid_index + 1) % len(ssids) # Cycle through SSIDs

                status_msg = f"Flooding {current_ssid}... (Packets sent: {sent_count})"
                if not infinite_mode:
                    status_msg += f" / {count}"
                q.put(('bf_status', status_msg))

                time.sleep(interval)

            if self.tool_stop_event.is_set():
                q.put(('bf_status', "Beacon flood canceled."))
            else:
                q.put(('bf_status', "Beacon flood complete."))

        except Exception as e:
            logging.error("Exception in beacon flood thread", exc_info=True)
            q.put(('error', "Beacon Flood Error", str(e)))
            success = False
        else:
            success = True
        finally:
            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='Tool Execution',
                    action='Executed Beacon Flood',
                    target=iface,
                    details=f"SSIDs: {len(ssids)}, BSSID: {bssid}, Encryption: {enc_type}",
                    result='Success' if success else 'Failure'
                )
            q.put(('tool_finished', 'beacon_flood', iface, f"Flooded {len(ssids)} SSIDs"))
            logging.info("Beacon flood thread finished.")

    def _arp_spoof_thread(self, victim_ip, target_ip):
        q = self.tool_results_queue
        iface = self.get_selected_iface()
        logging.info(f"ARP spoof thread started for Victim={victim_ip}, Target={target_ip}")

        try:
            q.put(('arp_spoof_status', "Resolving MAC addresses..."))
            victim_mac = getmacbyip(victim_ip)
            target_mac = getmacbyip(target_ip)

            if not victim_mac or not target_mac:
                raise Exception("Could not resolve MAC address for one or both targets. Are they online?")

            q.put(('arp_spoof_status', f"Victim: {victim_mac} | Target: {target_mac}"))
            logging.info(f"Resolved MACs -> Victim: {victim_mac}, Target: {target_mac}")

            victim_packet = Ether(dst=victim_mac)/ARP(op=2, pdst=victim_ip, hwdst=victim_mac, psrc=target_ip)
            target_packet = Ether(dst=target_mac)/ARP(op=2, pdst=target_ip, hwdst=target_mac, psrc=victim_ip)

            sent_count = 0
            while not self.tool_stop_event.is_set():
                sendp(victim_packet, iface=iface, verbose=0)
                sendp(target_packet, iface=iface, verbose=0)
                sent_count += 2
                q.put(('arp_spoof_status', f"Spoofing active... (Packets sent: {sent_count})"))
                time.sleep(2)

        except Exception as e:
            logging.error("Exception in ARP spoof thread", exc_info=True)
            q.put(('error', "ARP Spoof Error", str(e)))
            success = False
        else:
            # This part is only reached if the loop is broken by the stop event
            success = True
        finally:
            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='Tool Execution',
                    action='Executed ARP Spoof',
                    target=f"Victim: {victim_ip}",
                    details=f"Impersonating: {target_ip}",
                    result='Success' if success else 'Failure'
                )
            q.put(('tool_finished', 'arp_spoof', f"{victim_ip} -> {target_ip}", "ARP spoofing session"))
            logging.info("ARP spoof thread finished.")

    def start_arp_spoof(self):
        if not self._require_root("ARP Spoofer"):
            return
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        victim_ip = self.arp_spoof_victim_ip.text()
        target_ip = self.arp_spoof_target_ip.text()

        if not victim_ip or not target_ip:
            QMessageBox.critical(self, "Error", "Victim IP and Target IP are required.")
            return

        warning_msg = """
        <p>You are about to perform an ARP Spoofing attack. This will intercept traffic between the two targets and constitutes a Man-in-the-Middle attack.</p>
        <p>Ensure you have <b>explicit, written permission</b> to test on this network. Misuse of this tool is illegal.</p>
        <p><b>Do you accept full responsibility and wish to continue?</b></p>
        """
        if QMessageBox.question(self, "Ethical Use Confirmation", warning_msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.No:
            return

        if self.current_user:
            database.log_activity(
                self.current_user['id'], 'Offensive', 'Start ARP Spoof',
                f"Victim: {victim_ip}", f"Impersonating: {target_ip}"
            )

        self.arp_spoof_current_victim = victim_ip
        self.arp_spoof_current_target = target_ip

        self.is_tool_running = True
        self.arp_spoof_start_btn.setEnabled(False)
        self.arp_spoof_stop_btn.setEnabled(True)
        self.tool_stop_event.clear()

        args = (victim_ip, target_ip)
        self.worker = WorkerThread(self._arp_spoof_thread, args=args)
        self.worker.start()

    def stop_arp_spoof(self):
        if self.is_tool_running:
            logging.info("User requested to stop ARP spoofing.")
            if self.current_user:
                database.log_activity(
                    self.current_user['id'], 'Offensive', 'Stop ARP Spoof',
                    f"Victim: {self.arp_spoof_current_victim}", f"Target: {self.arp_spoof_current_target}"
                )
            self.arp_spoof_status.setText("Stopping...")
            self.tool_stop_event.set()

    def _restore_arp(self, victim_ip, target_ip):
        iface = self.get_selected_iface()
        logging.info(f"Attempting to restore ARP tables for {victim_ip} and {target_ip}")
        try:
            victim_mac = getmacbyip(victim_ip)
            target_mac = getmacbyip(target_ip)

            if not victim_mac or not target_mac:
                raise Exception("Could not resolve MACs for restoration. Manual correction may be needed.")

            # Create the legitimate ARP packets
            restore_victim_packet = Ether(dst=victim_mac)/ARP(op=2, pdst=victim_ip, hwdst=victim_mac, psrc=target_ip, hwsrc=target_mac)
            restore_target_packet = Ether(dst=target_mac)/ARP(op=2, pdst=target_ip, hwdst=target_mac, psrc=victim_ip, hwsrc=victim_mac)

            # Send them multiple times to ensure the cache is corrected
            sendp([restore_victim_packet, restore_target_packet], count=5, inter=0.2, iface=iface, verbose=0)

            logging.info("ARP restoration packets sent.")
            self.arp_spoof_status.setText("ARP tables restored. Attack stopped.")

        except Exception as e:
            logging.error(f"Failed to restore ARP tables: {e}", exc_info=True)
            QMessageBox.critical(self, "Restore Error", f"Could not restore ARP tables: {e}")

    def cancel_tool(self):
        if self.is_tool_running:
            logging.info("User requested to cancel the current tool.")
            self.tool_stop_event.set()

            # Special handling for subprocesses that need to be terminated directly
            # This is more robust than only relying on the thread's check loop.
            with self.thread_finish_lock:
                if hasattr(self, 'nmap_process') and self.nmap_process and self.nmap_process.poll() is None:
                    try:
                        self.nmap_process.terminate()
                        logging.info("Nmap process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Nmap process: {e}")

                if hasattr(self, 'arp_scan_cli_process') and self.arp_scan_cli_process and self.arp_scan_cli_process.poll() is None:
                    try:
                        self.arp_scan_cli_process.terminate()
                        logging.info("arp-scan process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating arp-scan process: {e}")
                if hasattr(self, 'spiderfoot_process') and self.spiderfoot_process and self.spiderfoot_process.poll() is None:
                    try:
                        self.spiderfoot_process.terminate()
                        logging.info("Spiderfoot process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Spiderfoot process: {e}")
                if hasattr(self, 'sherlock_process') and self.sherlock_process and self.sherlock_process.poll() is None:
                    try:
                        self.sherlock_process.terminate()
                        logging.info("Sherlock process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Sherlock process: {e}")
                if hasattr(self, 'fierce_process') and self.fierce_process and self.fierce_process.poll() is None:
                    try:
                        self.fierce_process.terminate()
                        logging.info("fierce process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating fierce process: {e}")
                if hasattr(self, 'dnsrecon_process') and self.dnsrecon_process and self.dnsrecon_process.poll() is None:
                    try:
                        self.dnsrecon_process.terminate()
                        logging.info("dnsrecon process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating dnsrecon process: {e}")
                if hasattr(self, 'enum4linux_ng_process') and self.enum4linux_ng_process and self.enum4linux_ng_process.poll() is None:
                    try:
                        self.enum4linux_ng_process.terminate()
                        logging.info("enum4linux-ng process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating enum4linux-ng process: {e}")
                if hasattr(self, 'hydra_process') and self.hydra_process and self.hydra_process.poll() is None:
                    try:
                        self.hydra_process.terminate()
                        logging.info("Hydra process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Hydra process: {e}")
                if hasattr(self, 'jtr_process') and self.jtr_process and self.jtr_process.poll() is None:
                    try:
                        self.jtr_process.terminate()
                        logging.info("JTR process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating JTR process: {e}")
                if hasattr(self, 'ffuf_process') and self.ffuf_process and self.ffuf_process.poll() is None:
                    try:
                        self.ffuf_process.terminate()
                        logging.info("ffuf process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating ffuf process: {e}")
                if hasattr(self, 'dirsearch_process') and self.dirsearch_process and self.dirsearch_process.poll() is None:
                    try:
                        self.dirsearch_process.terminate()
                        logging.info("dirsearch process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating dirsearch process: {e}")
                if hasattr(self, 'rustscan_process') and self.rustscan_process and self.rustscan_process.poll() is None:
                    try:
                        self.rustscan_process.terminate()
                        logging.info("RustScan process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating RustScan process: {e}")
                if hasattr(self, 'trufflehog_process') and self.trufflehog_process and self.trufflehog_process.poll() is None:
                    try:
                        self.trufflehog_process.terminate()
                        logging.info("TruffleHog process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating TruffleHog process: {e}")
                if hasattr(self, 'httpx_process') and self.httpx_process and self.httpx_process.poll() is None:
                    try:
                        self.httpx_process.terminate()
                        logging.info("httpx process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating httpx process: {e}")
                if hasattr(self, 'subfinder_process') and self.subfinder_process and self.subfinder_process.poll() is None:
                    try:
                        self.subfinder_process.terminate()
                        logging.info("Subfinder process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Subfinder process: {e}")

                if hasattr(self, 'sublist3r_process') and self.sublist3r_process and self.sublist3r_process.poll() is None:
                    try:
                        self.sublist3r_process.terminate()
                        logging.info("Sublist3r process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Sublist3r process: {e}")

                if hasattr(self, 'wifite_process') and self.wifite_process and self.wifite_process.poll() is None:
                    try:
                        os.kill(self.wifite_process.pid, signal.SIGTERM)
                        logging.info("Wifite process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating wifite process: {e}")

                if hasattr(self, 'nikto_process') and self.nikto_process and self.nikto_process.poll() is None:
                    try:
                        self.nikto_process.terminate()
                        logging.info("Nikto process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Nikto process: {e}")

                if hasattr(self, 'gobuster_process') and self.gobuster_process and self.gobuster_process.poll() is None:
                    try:
                        self.gobuster_process.terminate()
                        logging.info("Gobuster process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Gobuster process: {e}")

                if hasattr(self, 'sqlmap_process') and self.sqlmap_process and self.sqlmap_process.poll() is None:
                    try:
                        self.sqlmap_process.terminate()
                        logging.info("SQLMap process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating SQLMap process: {e}")

                if hasattr(self, 'whatweb_process') and self.whatweb_process and self.whatweb_process.poll() is None:
                    try:
                        self.whatweb_process.terminate()
                        logging.info("WhatWeb process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating WhatWeb process: {e}")

                if hasattr(self, 'hashcat_process') and self.hashcat_process and self.hashcat_process.poll() is None:
                    try:
                        self.hashcat_process.terminate()
                        logging.info("Hashcat process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Hashcat process: {e}")

                if hasattr(self, 'nuclei_process') and self.nuclei_process and self.nuclei_process.poll() is None:
                    try:
                        self.nuclei_process.terminate()
                        logging.info("Nuclei process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Nuclei process: {e}")
                if hasattr(self, 'masscan_process') and self.masscan_process and self.masscan_process.poll() is None:
                    try:
                        self.masscan_process.terminate()
                        logging.info("Masscan process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Masscan process: {e}")

                if hasattr(self, 'masscan_process') and self.masscan_process and self.masscan_process.poll() is None:
                    try:
                        self.masscan_process.terminate()
                        logging.info("Masscan process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Masscan process: {e}")

                if hasattr(self, 'masscan_process') and self.masscan_process and self.masscan_process.poll() is None:
                    try:
                        self.masscan_process.terminate()
                        logging.info("Masscan process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Masscan process: {e}")

    def _show_port_scan_summary_popup(self, results, target):
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Port Scan Results for {target}")
        dialog.setMinimumSize(400, 300)
        layout = QVBoxLayout(dialog)

        categorized = {"Open": [], "Open | Filtered": [], "Closed": [], "Filtered": [], "Unfiltered (RST)": [], "No Response / Filtered": []}
        for port, state, service in results:
            # Normalize states
            normalized_state = state
            if "No Response" in state:
                normalized_state = "No Response / Filtered"

            if normalized_state in categorized:
                categorized[normalized_state].append(f"{port} ({service})")
            else: # Fallback for any unexpected state
                if "Other" not in categorized: categorized["Other"] = []
                categorized["Other"].append(f"{port} ({state}, {service})")


        text_browser = QTextBrowser()
        text_browser.setOpenExternalLinks(False)
        html = f"<h1>Scan Report: {target}</h1>"
        # Display open ports first
        if categorized["Open"]:
            html += f"<h2>Open Ports ({len(categorized['Open'])})</h2>"
            html += "<ul>" + "".join(f"<li>{p}</li>" for p in sorted(categorized['Open'])) + "</ul>"

        for state, ports in categorized.items():
            if state != "Open" and ports:
                html += f"<h2>{state} ({len(ports)})</h2>"
                html += "<ul>" + "".join(f"<li>{p}</li>" for p in sorted(ports)) + "</ul>"

        text_browser.setHtml(html)
        layout.addWidget(text_browser)

        button_layout = QHBoxLayout()

        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(lambda: self._send_to_ai_analyst("port_scanner", results, context=target))
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(dialog.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

        dialog.exec()

    def _show_arp_scan_summary_popup(self, results, target):
        dialog = QDialog(self)
        dialog.setWindowTitle(f"ARP Scan Results for {target}")
        dialog.setMinimumSize(500, 400)
        layout = QVBoxLayout(dialog)

        summary_label = QLabel(f"<b>Found {len(results)} active hosts on network {target}.</b>")
        layout.addWidget(summary_label)

        tree = QTreeWidget()
        tree.setColumnCount(3)
        tree.setHeaderLabels(["IP Address", "MAC Address", "Vendor"])
        for res in results:
            item = QTreeWidgetItem([res['ip'], res['mac'], res['vendor']])
            tree.addTopLevelItem(item)
        tree.resizeColumnToContents(0)
        tree.resizeColumnToContents(1)
        layout.addWidget(tree)

        export_button = self._create_export_button(tree) # Reuse export functionality
        layout.addWidget(export_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(dialog.accept)
        layout.addWidget(ok_button)

        dialog.exec()

    def _process_tool_results(self):
        """Processes results from worker threads via a queue using a handler dictionary."""
        while not self.tool_results_queue.empty():
            msg = self.tool_results_queue.get()
            msg_type = msg[0]

            # Prioritize exact matches
            if msg_type in self.result_handlers:
                # Unpack arguments; msg[1:] creates a tuple of the remaining elements
                self.result_handlers[msg_type](*msg[1:])
                continue

            # Check for suffix-based dynamic handlers
            matched = False
            for suffix, handler in self.dynamic_handlers.items():
                if msg_type.endswith(suffix):
                    tool_name = msg_type.rsplit(suffix, 1)[0]
                    handler(tool_name, *msg[1:])
                    matched = True
                    break

            if not matched:
                logging.warning(f"No handler found for message type: {msg_type}")

    def _setup_result_handlers(self):
        """Initializes the dictionary mapping result queue messages to handler functions."""
        self.result_handlers = {
            # Exact message matches
            'send_results': self._handle_send_results,
            'send_finished': self._handle_send_finished,
            'tool_finished': self._handle_tool_finished,
            'report_finished': self._handle_report_finished,
            'flood_thread_finished': self._handle_flood_thread_finished,
            'ps_worker_finished': self._handle_ps_worker_finished,
            'crunch_finished': self._handle_crunch_finished,
            'show_port_scan_popup': self._show_port_scan_summary_popup,
            'show_arp_scan_popup': self._show_arp_scan_summary_popup,
            'arp_results': self._handle_arp_results,
            'v2ray_test_result': self._handle_v2ray_test_result,
            'ip_lookup_result': self._handle_ip_lookup_result,
            'l2tp_step_finished': self._handle_l2tp_step,
            'error': self._handle_error,
        }
        # Handlers for dynamic message types that end with a specific suffix
        self.dynamic_handlers = {
            '_output': self._handle_generic_output,
            'lab_status': self._handle_lab_status,
            'cve_search_status': self._handle_cve_search_status,
            'cve_result': self._handle_cve_result,
            'exploit_search_status': self._handle_exploit_search_status,
            'exploit_search_results': self._handle_exploit_search_results,
            '_status': self._handle_status_update,
            '_clear': self._handle_clear_update,
            '_result': self._handle_result_update,
            '_update': self._handle_result_update, # Catches 'wifi_scan_update'
            '_output': self._handle_generic_output,
        }
        self.result_handlers['nmap_xml_result'] = self._handle_nmap_xml_result
        self.result_handlers['sublist3r_results'] = self._show_subdomain_results_popup
        self.result_handlers['subdomain_results'] = self._show_subdomain_results_popup
        self.result_handlers['httpx_results'] = self._show_httpx_results_popup
        self.result_handlers['trufflehog_results'] = self._show_trufflehog_results_popup
        self.result_handlers['dirsearch_results'] = self._show_dirsearch_results_popup
        self.result_handlers['ffuf_results'] = self._show_ffuf_results_popup
        self.result_handlers['enum4linux_ng_results'] = self._show_enum4linux_ng_results_popup
        self.result_handlers['dnsrecon_results'] = self._show_dnsrecon_results_popup
        self.result_handlers['sherlock_results'] = self._show_sherlock_results_popup
        self.result_handlers['nuclei_results'] = self._show_nuclei_results_popup
        self.result_handlers['report_finding'] = self._handle_report_finding
        self.result_handlers['recent_threats_result'] = self._handle_recent_threats_result
        self.result_handlers['cve_import_finished'] = self._handle_cve_import_finished
        self.result_handlers['trufflehog_results'] = self._show_trufflehog_results_popup

    def _create_trufflehog_tool(self):
        """Creates the UI for the TruffleHog Secret Scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.trufflehog_controls = self._create_trufflehog_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.trufflehog_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.trufflehog_output_console = QPlainTextEdit()
        self.trufflehog_output_console.setReadOnly(True)
        self.trufflehog_output_console.setFont(QFont("Courier New", 10))
        self.trufflehog_output_console.setPlaceholderText("TruffleHog output will be displayed here...")
        main_layout.addWidget(self.trufflehog_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_trufflehog_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_trufflehog_config_widget(self):
        """Creates a reusable, self-contained widget with TruffleHog's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)

        controls = {}

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>TruffleHog - Secret Scanner</b></font>
        <p>This tool scans sources like git repositories, GitHub, and filesystems for leaked secrets.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Source Type and Target ---
        source_box = QGroupBox("Scan Target")
        source_layout = QVBoxLayout(source_box)

        controls['source_type_group'] = QButtonGroup(self)
        rb_layout = QHBoxLayout()
        controls['git_rb'] = QRadioButton("Git Repo"); controls['git_rb'].setChecked(True)
        controls['github_rb'] = QRadioButton("GitHub")
        controls['filesystem_rb'] = QRadioButton("Filesystem")
        controls['source_type_group'].addButton(controls['git_rb'])
        controls['source_type_group'].addButton(controls['github_rb'])
        controls['source_type_group'].addButton(controls['filesystem_rb'])
        rb_layout.addWidget(controls['git_rb']); rb_layout.addWidget(controls['github_rb']); rb_layout.addWidget(controls['filesystem_rb'])
        source_layout.addLayout(rb_layout)

        target_layout = QHBoxLayout()
        controls['target_edit'] = QLineEdit()
        controls['target_edit'].setPlaceholderText("Enter Git URL, GitHub repo, or filesystem path...")
        target_layout.addWidget(controls['target_edit'])
        controls['browse_btn'] = QPushButton("Browse...")
        controls['browse_btn'].clicked.connect(lambda: self._browse_dir_for_lineedit(controls['target_edit'], "Select Directory to Scan"))
        target_layout.addWidget(controls['browse_btn'])
        source_layout.addLayout(target_layout)

        main_layout.addWidget(source_box)

        # --- Options ---
        options_box = QGroupBox("Options")
        options_layout = QFormLayout(options_box)
        controls['only_verified_check'] = QCheckBox("Only Verified Results")
        controls['only_verified_check'].setToolTip("Only output secrets that have been successfully verified against their respective APIs.")
        options_layout.addRow(controls['only_verified_check'])
        main_layout.addWidget(options_box)

        # --- Action Buttons ---
        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        # UI Logic to show/hide browse button
        def toggle_browse_button():
            controls['browse_btn'].setVisible(controls['filesystem_rb'].isChecked())
        controls['source_type_group'].buttonClicked.connect(toggle_browse_button)
        toggle_browse_button() # Set initial state

        return widget, controls

    def _build_trufflehog_command(self, controls):
        """Builds the TruffleHog command list from a dictionary of controls or widgets."""
        tool_path = self._get_tool_path("trufflehog", os.path.join("trufflehog", "trufflehog"))
        if not tool_path:
            return None, None, "TruffleHog not found."
        command = [tool_path]

        target = self._get_control_value(controls, 'target_edit', 'text')
        if not target or not target.strip():
            return None, None, "A target is required."
        target = target.strip()
        target_for_log = target

        source_type = ""
        if self._get_control_value(controls, 'git_rb', 'check'):
            source_type = "git"
        elif self._get_control_value(controls, 'github_rb', 'check'):
            source_type = "github"
        elif self._get_control_value(controls, 'filesystem_rb', 'check'):
            source_type = "filesystem"

        if not source_type:
             return None, None, "A source type (Git, GitHub, Filesystem) must be selected."

        command.extend([source_type, "--repo", target])

        if self._get_control_value(controls, 'only_verified_check', 'check'):
            command.append("--only-verified")

        command.append("--json") # Always add --json for parsing

        return command, target_for_log, None

    def start_trufflehog_scan(self):
        """Starts the TruffleHog scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.trufflehog_controls
        command, target_for_log, error = self._build_trufflehog_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.trufflehog_output_console.clear()

        self.worker = WorkerThread(self._trufflehog_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _trufflehog_thread(self, command, target):
        """Worker thread for running the trufflehog command."""
        q = self.tool_results_queue
        logging.info(f"Starting TruffleHog with command: {' '.join(command)}")
        q.put(('trufflehog_output', f"$ {' '.join(command)}\n\n"))
        json_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.trufflehog_process = process

            full_output = []
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('trufflehog_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                full_output.append(line)

            process.stdout.close()
            process.wait()

            json_data = "".join(full_output)
            q.put(('trufflehog_output', json_data))
            if not self.tool_stop_event.is_set() and json_data.strip():
                q.put(('trufflehog_results', json_data))

        except FileNotFoundError:
            q.put(('error', 'TruffleHog Error', "'trufflehog' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"TruffleHog thread error: {e}", exc_info=True)
            q.put(('error', 'TruffleHog Error', str(e)))
        finally:
            q.put(('tool_finished', 'trufflehog_scan', target, json_data))
            with self.thread_finish_lock:
                self.trufflehog_process = None
            logging.info("TruffleHog scan thread finished.")

    def _handle_trufflehog_output(self, line):
        self.trufflehog_output_console.insertPlainText(line)
        self.trufflehog_output_console.verticalScrollBar().setValue(self.trufflehog_output_console.verticalScrollBar().maximum())

    def _show_trufflehog_results_popup(self, json_data):
        """Shows the results of a trufflehog scan in a dedicated dialog."""
        dialog = TruffleHogResultsDialog(json_data, self)
        dialog.exec()

    def _handle_cve_import_finished(self, success, total_records, error_message):
        """Shows a confirmation message after the CVE import process is complete."""
        if success:
            QMessageBox.information(self, "Import Successful",
                f"Successfully imported {total_records} total records into the offline CVE database.")
        else:
            QMessageBox.critical(self, "Import Failed",
                f"The CVE import process failed with an error:\n\n{error_message}\n\n"
                "Please check the log for more details.")
        # Reset the status label and re-enable buttons
        self.cve_import_status_changed.emit("Idle")
        self.cve_import_buttons_enabled.emit(True)

    def _handle_report_finished(self, success, message):
        """Handles the result of the report generation thread."""
        self.report_generate_html_btn.setEnabled(True)  # ✅ تصحیح شد
        if success:
            QMessageBox.information(self, "Report Generated", f"Report successfully saved to:\n{message}")
            self.status_bar.showMessage("Report generation successful.", 5000)
        else:
            QMessageBox.critical(self, "Report Generation Error", f"Failed to generate report:\n{message}")
            self.status_bar.showMessage("Report generation failed.", 5000)

    def _handle_report_finding(self, finding_data):
        """Adds a finding to the report tree. finding_data is a tuple."""
        # The data should be (host, port_service, finding, details)
        item = QTreeWidgetItem([str(col) for col in finding_data])
        self.report_findings_tree.addTopLevelItem(item)

    def _handle_aggregation(self):
        """Starts the background thread for aggregating and enriching tool results."""
        if not self.nmap_last_xml:
            QMessageBox.warning(self, "No Data", "No Nmap scan data is available to analyze. Please run an Nmap scan with XML output first.")
            return

        self.report_aggregate_btn.setEnabled(False)
        self.report_findings_tree.clear()
        self.status_bar.showMessage("Aggregating and enriching results...")

        self.worker = WorkerThread(self._aggregation_thread)
        self.active_threads.append(self.worker)
        self.worker.start()

    def _aggregation_thread(self):
        """Parses tool outputs, enriches data, and sends it to the reporting tab."""
        q = self.tool_results_queue
        use_offline_db = self.offline_cve_check.isChecked()
        try:
            if not LXML_AVAILABLE:
                q.put(('error', 'Dependency Error', "The 'lxml' library is required for report generation. Please install it."))
                return

            parser = etree.XMLParser(recover=True, no_network=True, dtd_validation=False)
            root = etree.fromstring(self.nmap_last_xml.encode('utf-8'), parser=parser)

            for host in root.findall('host'):
                if host.find('status').get('state') != 'up':
                    continue

                address = host.find('address').get('addr')
                ports_elem = host.find('ports')
                if ports_elem is None:
                    continue

                for port in ports_elem.findall('port'):
                    if port.find('state').get('state') != 'open':
                        continue

                    port_id = port.get('portid')
                    protocol = port.get('protocol')
                    port_service_str = f"{port_id}/{protocol}"

                    service_elem = port.find('service')
                    if service_elem is not None:
                        service_name = service_elem.get('name', 'N/A')
                        product = service_elem.get('product', '')
                        version = service_elem.get('version', '')

                        # Construct a search term, prioritizing product and version
                        search_term = f"{product} {version}".strip()
                        if not search_term:
                            search_term = service_name

                        full_service_str = f"{service_name} ({search_term})"

                        # Enrich data with CVEs and Exploits
                        if use_offline_db:
                            cve_details = self._query_local_cve_db(search_term)
                        else:
                            cve_details = self._query_cve_api(search_term)
                        exploit_details = self._query_searchsploit(search_term)

                        q.put(('report_finding', (address, port_service_str, full_service_str, f"{cve_details}\n{exploit_details}")))

        except Exception as e:
            logging.error(f"Error during result aggregation: {e}", exc_info=True)
            q.put(('error', 'Aggregation Error', f'An error occurred: {e}'))
        finally:
            q.put(('tool_finished', 'aggregation'))

    def _handle_generate_report(self):
        """Handles the report generation process by gathering data and starting a worker."""
        # 1. Gather all data from the UI
        report_data = {
            "client": self.report_client_name.text(),
            "dates": self.report_assessment_dates.text(),
            "objectives": self.report_objectives.toPlainText(),
            "in_scope": self.report_in_scope.toPlainText(),
            "out_of_scope": self.report_out_of_scope.toPlainText(),
            "summary": self.report_summary_text.toPlainText(),
            "findings": []
        }

        if self.report_findings_tree.topLevelItemCount() == 0:
            QMessageBox.warning(self, "No Data", "There are no findings to report. Please run the 'Aggregate & Enrich Results' tool first.")
            return

        for i in range(self.report_findings_tree.topLevelItemCount()):
            item = self.report_findings_tree.topLevelItem(i)
            report_data["findings"].append({
                "host": item.text(0),
                "service": item.text(1),
                "finding": item.text(2),
                "details": item.text(3)
            })

        # 2. Prompt user for save location
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Report", "Security_Report.html", "HTML Files (*.html)", options=QFileDialog.Option.DontUseNativeDialog)
        if not file_path:
            return

        template_name = self.report_template_combo.currentText()

        # 3. Start worker thread
        self.status_bar.showMessage("Generating report...")
        self.report_generate_html_btn.setEnabled(False)
        self.worker = WorkerThread(self._report_generation_thread, args=(report_data, file_path, template_name))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _show_cve_update_info(self):
        QMessageBox.information(self, "Manual CVE Database Update",
            "Automatic downloads from the National Vulnerability Database (NVD) are currently blocked by their security provider (Cloudflare).\n\n"
            "To update the offline database, please follow these steps:\n"
            "1. Go to the NVD Data Feeds page: https://nvd.nist.gov/vuln/data-feeds\n"
            "2. Download the desired JSON feeds (e.g., 'nvdcve-1.1-modified.json.gz', 'nvdcve-1.1-2023.json.gz', etc.).\n"
            "3. Return to this tab and click the 'Import NVD File' button.\n"
            "4. Select the downloaded `.json.gz` file(s) to import them into the local database.")

    def _start_cve_import(self):
        """Starts the background thread to import local CVE files."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        file_paths, _ = QFileDialog.getOpenFileNames(self, "Select NVD JSON Feed(s)", "", "NVD JSON Files (*.json.gz)", options=QFileDialog.Option.DontUseNativeDialog)
        if not file_paths:
            return

        self.is_tool_running = True
        self.cve_import_buttons_enabled.emit(False)
        self.status_bar.showMessage("Starting CVE database import...")
        self.cve_import_status_changed.emit("Starting import...")

        self.worker = WorkerThread(self._cve_import_thread, args=(file_paths,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _cve_import_thread(self, file_paths):
        """Worker thread to process local NVD data feeds into the SQLite DB."""
        q = self.tool_results_queue
        db_path = "cve.db"
        total_records = 0
        success = True
        error_message = ""

        try:
            # Ensure the CVE-specific table exists in the correct database file.
            database.create_cve_table()
            con = sqlite3.connect(db_path)
            cur = con.cursor()

            total_files = len(file_paths)
            for i, file_path in enumerate(file_paths):
                if self.tool_stop_event.is_set():
                    q.put(('cve_import_status', "Import cancelled."))
                    break

                filename = os.path.basename(file_path)
                q.put(('cve_import_status', f"Processing file {i+1}/{total_files}: {filename}..."))

                with gzip.open(file_path, 'rb') as f:
                    json_data = f.read()
                data = json.loads(json_data)

                cves_to_insert = []

                # Check for NVD JSON 2.0 format vs 1.1
                if 'vulnerabilities' in data:
                    # NVD 2.0 Format Parsing
                    for cve_item in data.get('vulnerabilities', []):
                        cve_data = cve_item.get('cve', {})
                        cve_id = cve_data.get('id', 'N/A')

                        description = "No description available."
                        for desc in cve_data.get('descriptions', []):
                            if desc.get('lang') == 'en':
                                description = desc.get('value', '')
                                break

                        cvss_v3_score = None
                        if 'cvssMetricV31' in cve_data.get('metrics', {}):
                            cvss_v3_score = cve_data['metrics']['cvssMetricV31'][0]['cvssData'].get('baseScore')

                        cvss_v2_score = None
                        if 'cvssMetricV2' in cve_data.get('metrics', {}):
                             cvss_v2_score = cve_data['metrics']['cvssMetricV2'][0]['cvssData'].get('baseScore')

                        published_date = cve_data.get('published', '')

                        # Keyword extraction for 2.0 is more complex, involving CPEs
                        keywords = set()
                        configs = cve_data.get('configurations', [])
                        for config in configs:
                            for node in config.get('nodes', []):
                                for cpe_match in node.get('cpeMatch', []):
                                    cpe_uri = cpe_match.get('criteria', '')
                                    parts = cpe_uri.split(':')
                                    if len(parts) > 4:
                                        keywords.add(parts[3]) # vendor
                                        keywords.add(parts[4]) # product
                        keywords_str = " ".join(sorted(list(keywords)))

                        cves_to_insert.append((cve_id, description, cvss_v3_score, cvss_v2_score, keywords_str, published_date))
                else:
                    # Fallback to NVD 1.1 Format Parsing
                    for cve_item in data.get('CVE_Items', []):
                        cve_id = cve_item['cve']['CVE_data_meta']['ID']
                        description = cve_item['cve']['description']['description_data'][0]['value']
                        keywords = set()
                        nodes = cve_item.get('configurations', {}).get('nodes', [])
                        for node in nodes:
                            cpe_matches = node.get('cpe_match', [])
                            for cpe_match in cpe_matches:
                                cpe_uri = cpe_match.get('cpe23Uri', '')
                                parts = cpe_uri.split(':')
                                if len(parts) > 4:
                                    keywords.add(parts[3])
                                    keywords.add(parts[4])
                        keywords_str = " ".join(sorted(list(keywords)))
                        metrics_v3 = cve_item.get('impact', {}).get('baseMetricV3', {})
                        cvss_v3_score = metrics_v3.get('cvssV3', {}).get('baseScore')
                        metrics_v2 = cve_item.get('impact', {}).get('baseMetricV2', {})
                        cvss_v2_score = metrics_v2.get('cvssV2', {}).get('baseScore')
                        published_date = cve_item.get('publishedDate', '')
                        cves_to_insert.append((cve_id, description, cvss_v3_score, cvss_v2_score, keywords_str, published_date))

                cur.executemany("INSERT OR REPLACE INTO vulnerabilities VALUES (?, ?, ?, ?, ?, ?)", cves_to_insert)
                con.commit()
                total_records += len(cves_to_insert)
                q.put(('cve_import_status', f"Finished {filename}. Imported {len(cves_to_insert)} records."))

        except Exception as e:
            success = False
            error_message = str(e)
            logging.error(f"Failed to import offline CVE DB: {e}", exc_info=True)
            q.put(('error', 'CVE DB Error', str(e)))
        finally:
            if 'con' in locals() and con:
                con.close()
            q.put(('cve_import_finished', success, total_records, error_message))
            q.put(('tool_finished', 'cve_import', ", ".join(file_paths), f"Imported {total_records} records."))

    def _query_local_cve_db(self, keyword):
        """Queries the local SQLite CVE database for a given keyword."""
        db_path = "cve.db"
        if not os.path.exists(db_path):
            return "Offline CVE database (cve.db) not found. Please update it first."

        try:
            con = sqlite3.connect(db_path)
            cur = con.cursor()

            search_terms = keyword.split()
            query = "SELECT cve_id, cvss_v3_score, description FROM vulnerabilities WHERE "
            query += " AND ".join(["keywords LIKE ?"] * len(search_terms))
            query += " ORDER BY cvss_v3_score DESC LIMIT 5"

            params = [f"%{term}%" for term in search_terms]

            cur.execute(query, params)
            results = cur.fetchall()
            con.close()

            if not results:
                return "No CVEs found in offline DB."

            output = ["--- CVEs (Offline) ---"]
            for cve_id, score, description in results:
                output.append(f"{cve_id} (Score: {score or 'N/A'}): {description[:100]}...")
            return "\n".join(output)

        except Exception as e:
            logging.error(f"Offline CVE DB query for '{keyword}' failed: {e}")
            return f"Offline CVE lookup failed: {e}"

    def _report_generation_thread(self, report_data, file_path, template_name):
        """Worker thread to generate the final HTML report."""
        q = self.tool_results_queue
        try:
            template_path = os.path.join("report_templates", template_name)
            with open(template_path, 'r', encoding='utf-8') as f:
                template_html = f.read()

            findings_html = ""
            for finding in report_data['findings']:
                findings_html += "<tr>\n"
                findings_html += f"    <td>{finding['host']}</td>\n"
                findings_html += f"    <td>{finding['service']}</td>\n"
                findings_html += f"    <td>{finding['finding']}</td>\n"
                findings_html += f"    <td><pre>{finding['details']}</pre></td>\n"
                findings_html += "</tr>\n"

            # Sanitize and format the data
            sanitized_data = {k: v.replace('<', '&lt;').replace('>', '&gt;') if isinstance(v, str) else v for k, v in report_data.items()}

            final_html = template_html.format(
                client=sanitized_data['client'],
                dates=sanitized_data['dates'],
                objectives=sanitized_data['objectives'],
                in_scope=sanitized_data['in_scope'],
                out_of_scope=sanitized_data['out_of_scope'],
                summary=sanitized_data['summary'].replace('\n', '<br>'),
                findings_loop=findings_html
            )

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(final_html)

            q.put(('report_finished', True, file_path))

        except Exception as e:
            logging.error(f"Error during report generation: {e}", exc_info=True)
            q.put(('report_finished', False, str(e)))

    def _gather_report_data(self):
        """Gathers all data from the reporting UI fields into a dictionary."""
        if self.report_findings_tree.topLevelItemCount() == 0:
            QMessageBox.warning(self, "No Data", "There are no findings to report. Please run the 'Aggregate & Enrich Results' tool first.")
            return None

        report_data = {
            "client_name": self.report_client_name.text(),
            "assessment_dates": self.report_assessment_dates.text(),
            "objectives": self.report_objectives.toPlainText(),
            "in_scope": self.report_in_scope.toPlainText(),
            "out_of_scope": self.report_out_of_scope.toPlainText(),
            "summary": self.report_summary_text.toPlainText(),
            "findings": []
        }

        for i in range(self.report_findings_tree.topLevelItemCount()):
            item = self.report_findings_tree.topLevelItem(i)
            report_data["findings"].append({
                "host": item.text(0),
                "service": item.text(1),
                "finding": item.text(2),
                "details": item.text(3)
            })
        return report_data

    def _handle_generate_doc_report(self, file_format):
        """Handles the generation of reports in different document formats."""
        report_data = self._gather_report_data()
        if not report_data:
            return

        file_path, _ = QFileDialog.getSaveFileName(self, f"Save Report as {file_format.upper()}", f"Security_Report.{file_format}", f"{file_format.upper()} Files (*.{file_format})", options=QFileDialog.Option.DontUseNativeDialog)
        if not file_path:
            return

        self.status_bar.showMessage(f"Generating {file_format.upper()} report...")

        try:
            if file_format == 'docx':
                self._export_report_to_docx(report_data, file_path)
            elif file_format == 'pdf':
                self._export_report_to_pdf(report_data, file_path)

            QMessageBox.information(self, "Success", f"Report successfully saved to:\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate {file_format.upper()} report: {e}")
        finally:
            self.status_bar.showMessage("Report generation finished.", 5000)

    def _export_report_to_docx(self, data, file_path):
        """Exports the structured report data to a DOCX file."""
        document = docx.Document()
        document.add_heading(f"Penetration Test Report: {data['client_name']}", 0)

        document.add_heading("Executive Summary", level=1)
        document.add_paragraph(data['summary'])

        document.add_heading("Scope", level=1)
        document.add_paragraph(f"Assessment Dates: {data['assessment_dates']}")
        document.add_heading("Objectives", level=2)
        document.add_paragraph(data['objectives'])
        document.add_heading("In-Scope Targets", level=2)
        document.add_paragraph(data['in_scope'])
        document.add_heading("Out-of-Scope Targets", level=2)
        document.add_paragraph(data['out_of_scope'])

        document.add_heading("Detailed Findings", level=1)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Host'
        hdr_cells[1].text = 'Service'
        hdr_cells[2].text = 'Finding'
        hdr_cells[3].text = 'Details'

        for finding in data['findings']:
            row_cells = table.add_row().cells
            row_cells[0].text = finding['host']
            row_cells[1].text = finding['service']
            row_cells[2].text = finding['finding']
            row_cells[3].text = finding['details']

        document.save(file_path)

    def _export_report_to_pdf(self, data, file_path):
        """Exports the structured report data to a PDF file."""
        doc = SimpleDocTemplate(file_path)
        styles = getSampleStyleSheet()
        elements = [Paragraph(f"Penetration Test Report: {data['client_name']}", styles['h1'])]

        elements.append(Paragraph("Executive Summary", styles['h2']))
        elements.append(Paragraph(data['summary'].replace('\n', '<br/>'), styles['BodyText']))

        elements.append(Paragraph("Scope", styles['h2']))
        elements.append(Paragraph(f"<b>Assessment Dates:</b> {data['assessment_dates']}", styles['BodyText']))
        elements.append(Paragraph("Objectives", styles['h3']))
        elements.append(Paragraph(data['objectives'].replace('\n', '<br/>'), styles['BodyText']))
        elements.append(Paragraph("In-Scope Targets", styles['h3']))
        elements.append(Paragraph(data['in_scope'].replace('\n', '<br/>'), styles['BodyText']))
        elements.append(Paragraph("Out-of-Scope Targets", styles['h3']))
        elements.append(Paragraph(data['out_of_scope'].replace('\n', '<br/>'), styles['BodyText']))

        elements.append(Paragraph("Detailed Findings", styles['h2']))

        table_data = [['Host', 'Service', 'Finding', 'Details']]
        for f in data['findings']:
            details_p = Paragraph(f['details'], styles['BodyText'])
            table_data.append([f['host'], f['service'], f['finding'], details_p])

        table = Table(table_data, colWidths=[1.5*inch, 1.5*inch, 2*inch, 2.5*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        elements.append(table)

        doc.build(elements)

    def _handle_ai_report_generation(self):
        """Gathers findings and instructions, and sends them to the AI for analysis."""
        report_data = self._gather_report_data()
        if not report_data:
            return

        persona = self.ai_persona_combo.currentText()
        custom_instructions = self.ai_instructions_edit.toPlainText()

        findings_text = ""
        for i, finding in enumerate(report_data['findings']):
            findings_text += f"Finding {i+1}:\n"
            findings_text += f"  Host: {finding['host']}\n"
            findings_text += f"  Service: {finding['service']}\n"
            findings_text += f"  Vulnerability: {finding['finding']}\n"
            findings_text += f"  Details: {finding['details']}\n\n"

        prompt = (
            f"You are an AI assistant. Please analyze the following penetration testing findings from the perspective of a **{persona}**. "
            "Your goal is to provide a detailed analysis and recommendations based on these findings.\n\n"
        )

        if custom_instructions:
            prompt += f"Please follow these custom instructions: '{custom_instructions}'\n\n"

        prompt += (
            "--- FINDINGS ---\n"
            f"{findings_text}"
            "--- END FINDINGS ---\n\n"
            "Please provide your analysis."
        )

        self.ai_assistant_tab.send_message(prompt)
        self.tab_widget.setCurrentWidget(self.ai_assistant_tab)
        QMessageBox.information(self, "AI Task Started", "The AI is analyzing the findings. You can see the results in the 'AI Assistant' tab.")

    def _query_cve_api(self, keyword):
        """Queries the NVD CVE API for a given keyword and returns a formatted string."""
        if not keyword:
            return "No service info to query CVEs."
        try:
            # URL-encode the keyword to handle spaces and special characters
            encoded_keyword = urllib.parse.quote(keyword)
            url = f"https://services.nvd.nist.gov/rest/json/cves/2.0?keywordSearch={encoded_keyword}&resultsPerPage=5"

            # Add a user-agent to be compliant with API usage policies
            req = urllib.request.Request(url, headers={'User-Agent': 'Zurvan/3.0'})

            with urllib.request.urlopen(req, timeout=10) as response:
                data = json.load(response)

            vulnerabilities = data.get('vulnerabilities', [])
            if not vulnerabilities:
                return "No CVEs found."

            output = ["--- CVEs ---"]
            for item in vulnerabilities:
                cve = item.get('cve', {})
                cve_id = cve.get('id', 'N/A')

                # Get description
                description = "No description available."
                for desc in cve.get('descriptions', []):
                    if desc.get('lang') == 'en':
                        description = desc.get('value', '')
                        break

                # Get CVSS V3 score if available, otherwise V2
                cvss_score = "N/A"
                if 'cvssMetricV31' in cve.get('metrics', {}):
                    cvss_score = cve['metrics']['cvssMetricV31'][0]['cvssData']['baseScore']
                elif 'cvssMetricV2' in cve.get('metrics', {}):
                    cvss_score = cve['metrics']['cvssMetricV2'][0]['cvssData']['baseScore']

                output.append(f"{cve_id} (Score: {cvss_score}): {description[:100]}...")
            return "\n".join(output)

        except Exception as e:
            logging.error(f"CVE API query for '{keyword}' failed: {e}")
            return f"CVE lookup failed: {e}"

    def _query_searchsploit(self, keyword):
        """Queries the local searchsploit database and returns a formatted string."""
        if not keyword:
            return "No service info to query exploits."

        if not shutil.which("searchsploit"):
            return "searchsploit command not found. Please install Exploit-DB."

        try:
            command = ["searchsploit", "--json", keyword]
            process = subprocess.run(command, capture_output=True, text=True, timeout=15)

            if process.returncode != 0:
                return f"Searchsploit error: {process.stderr}"

            data = json.loads(process.stdout)
            results = data.get('RESULTS_EXPLOIT', [])

            if not results:
                return "No exploits found."

            output = ["--- Exploits ---"]
            for item in results[:5]: # Limit to top 5 results
                title = item.get('Title', 'N/A')
                edb_id = item.get('EDB-ID', 'N/A')
                output.append(f"EDB-ID: {edb_id} - {title}")
            return "\n".join(output)

        except Exception as e:
            logging.error(f"Searchsploit query for '{keyword}' failed: {e}")
            return f"Exploit lookup failed: {e}"

    def _show_subdomain_results_popup(self, domain, subdomains):
        """Shows the results of a subdomain scan in a dedicated dialog."""
        if not subdomains:
            QMessageBox.information(self, "No Results", f"No subdomains were found for {domain}.")
            return
        dialog = SubdomainResultsDialog(subdomains, domain, self)
        dialog.exec()

    def _handle_wifite_output(self, line):
        self.wifite_output_console.insertPlainText(line)
        self.wifite_output_console.verticalScrollBar().setValue(self.wifite_output_console.verticalScrollBar().maximum())

    def _whatweb_thread(self, command):
        """Worker thread for running the whatweb command."""
        q = self.tool_results_queue
        target_for_log = command[-1] # Assume target is the last argument
        logging.info(f"Starting WhatWeb with command: {' '.join(command)}")
        q.put(('whatweb_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.whatweb_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('whatweb_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('whatweb_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'WhatWeb Error', "'whatweb' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"WhatWeb thread error: {e}", exc_info=True)
            q.put(('error', 'WhatWeb Error', str(e)))
        finally:
            q.put(('tool_finished', 'whatweb_scan', target_for_log, "".join(full_output)))
            with self.thread_finish_lock:
                self.whatweb_process = None
            logging.info("WhatWeb scan thread finished.")

    def _handle_whatweb_output(self, line):
        self.whatweb_output_console.insertPlainText(line)
        self.whatweb_output_console.verticalScrollBar().setValue(self.whatweb_output_console.verticalScrollBar().maximum())

    def _handle_nikto_output(self, line):
        self.nikto_output_console.insertPlainText(line)
        self.nikto_output_console.verticalScrollBar().setValue(self.nikto_output_console.verticalScrollBar().maximum())

    def _handle_gobuster_output(self, line):
        self.gobuster_output_console.insertPlainText(line)
        self.gobuster_output_console.verticalScrollBar().setValue(self.gobuster_output_console.verticalScrollBar().maximum())

    def _handle_sqlmap_output(self, line):
        self.sqlmap_output_console.insertPlainText(line)
        self.sqlmap_output_console.verticalScrollBar().setValue(self.sqlmap_output_console.verticalScrollBar().maximum())

    def _handle_whatweb_output(self, line):
        self.whatweb_output_console.insertPlainText(line)
        self.whatweb_output_console.verticalScrollBar().setValue(self.whatweb_output_console.verticalScrollBar().maximum())

    def _create_theharvester_tool(self):
        """Creates the UI for theHarvester OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Instructions ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>theHarvester - OSINT Tool</b></font>
        <p>This tool gathers emails, subdomains, hosts, employee names, open ports and banners from different public sources like search engines, PGP key servers and SHODAN database.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Controls ---
        config_widget, self.theharvester_controls = self._create_theharvester_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.theharvester_controls['start_btn'])
        buttons_layout.addWidget(self.theharvester_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.theharvester_output_console = QPlainTextEdit()
        self.theharvester_output_console.setReadOnly(True)
        self.theharvester_output_console.setFont(QFont("Courier New", 10))
        self.theharvester_output_console.setPlaceholderText("theHarvester output will be displayed here...")
        main_layout.addWidget(self.theharvester_output_console, 1)

        self.theharvester_controls['start_btn'].clicked.connect(self.start_theharvester_scan)
        self.theharvester_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_theharvester_config_widget(self):
        """Creates a reusable, self-contained widget for theHarvester's configuration options."""
        widget = QGroupBox("Scan Options")
        layout = QFormLayout(widget)
        controls = {}

        controls['domain_edit'] = QLineEdit("example.com")
        layout.addRow("Domain (-d):", controls['domain_edit'])

        controls['limit_edit'] = QLineEdit("500")
        layout.addRow("Result Limit (-l):", controls['limit_edit'])

        controls['source_edit'] = QLineEdit("google,bing")
        controls['source_edit'].setToolTip("Comma-separated list of data sources (e.g., google,bing,shodan).")
        layout.addRow("Sources (-b):", controls['source_edit'])

        controls['dns_check'] = QCheckBox("DNS Brute-force (-c)")
        controls['dns_check'].setToolTip("Enable DNS brute-forcing for subdomains.")
        layout.addRow(controls['dns_check'])

        controls['dns_lookup_check'] = QCheckBox("DNS Lookup (-n)")
        controls['dns_lookup_check'].setToolTip("Enable DNS lookup for found hosts.")
        layout.addRow(controls['dns_lookup_check'])

        controls['dns_server_edit'] = QLineEdit()
        controls['dns_server_edit'].setPlaceholderText("Optional: e.g., 8.8.8.8")
        layout.addRow("DNS Server (-e):", controls['dns_server_edit'])

        controls['start_btn'] = QPushButton(QIcon(self.icon_path("search.svg")), " Start Harvesting")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def _build_theharvester_command(self, controls):
        """Builds the theHarvester command list from a dictionary of controls or widgets."""
        harvester_path = os.path.join(self.script_dir, "tools", "theHarvester", "theHarvester.py")
        command = ["python3", harvester_path]

        domain = self._get_control_value(controls, 'domain_edit', 'text')
        if not domain or not domain.strip():
            return None, None, "A domain is required."
        command.extend(["-d", domain.strip()])
        target_for_log = domain.strip()

        limit = self._get_control_value(controls, 'limit_edit', 'text')
        if limit and limit.strip():
            command.extend(["-l", limit.strip()])

        source = self._get_control_value(controls, 'source_edit', 'text')
        if source and source.strip():
            command.extend(["-b", source.strip()])

        dns_server = self._get_control_value(controls, 'dns_server_edit', 'text')
        if dns_server and dns_server.strip():
            command.extend(["-e", dns_server.strip()])

        if self._get_control_value(controls, 'dns_check', 'check'):
            command.append("-c")

        if self._get_control_value(controls, 'dns_lookup_check', 'check'):
            command.append("-n")

        return command, target_for_log, None

    def start_theharvester_scan(self):
        """Starts the theHarvester scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        controls = self.theharvester_controls
        command, target_for_log, error = self._build_theharvester_command(controls)

        if error:
            QMessageBox.critical(self, "Input Error", error)
            return

        harvester_path = os.path.join(self.script_dir, "tools", "theHarvester", "theHarvester.py")
        if not os.path.exists(harvester_path):
            QMessageBox.critical(self, "theHarvester Error", f"'theHarvester.py' not found at {harvester_path}. Please ensure it is cloned correctly.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.theharvester_output_console.clear()

        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'theharvester_scan', target_for_log, self.theharvester_output_console
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _handle_theharvester_output(self, line):
        self.theharvester_output_console.insertPlainText(line)
        self.theharvester_output_console.verticalScrollBar().setValue(self.theharvester_output_console.verticalScrollBar().maximum())

    def _handle_hashcat_output(self, line):
        self.hashcat_output_console.insertPlainText(line)
        self.hashcat_output_console.verticalScrollBar().setValue(self.hashcat_output_console.verticalScrollBar().maximum())

    def _handle_masscan_output(self, line):
        self.masscan_output_console.insertPlainText(line)
        self.masscan_output_console.verticalScrollBar().setValue(self.masscan_output_console.verticalScrollBar().maximum())

    def _handle_nmap_output(self, line):
        self.nmap_output_console.insertPlainText(line)
        self.nmap_output_console.verticalScrollBar().setValue(self.nmap_output_console.verticalScrollBar().maximum())

    def _handle_nmap_xml_result(self, xml_content):
        """Stores the captured Nmap XML report and shows a summary dialog."""
        self.nmap_last_xml = xml_content
        logging.info(f"Captured Nmap XML report ({len(xml_content)} bytes).")
        self.status_bar.showMessage("Nmap scan complete. XML report captured.", 5000)
        self.nmap_controls['report_btn'].setEnabled(True)

        # Automatically show the summary dialog
        target_context = self.nmap_controls['target_edit'].text()
        summary_dialog = NmapSummaryDialog(xml_content, target_context, self)
        summary_dialog.exec()

    def _show_subdomain_results_popup(self, domain, subdomains):
        """Shows the results of a subdomain scan in a dedicated dialog."""
        if not subdomains:
            QMessageBox.information(self, "No Results", f"No subdomains were found for {domain}.")
            return
        dialog = SubdomainResultsDialog(subdomains, domain, self)
        dialog.exec()

    def _show_httpx_results_popup(self, json_data):
        """Shows the results of an httpx scan in a dedicated dialog."""
        dialog = HttpxResultsDialog(json_data, self)
        dialog.exec()

    def _show_dirsearch_results_popup(self, json_data, target_context):
        """Shows the results of a dirsearch scan in a dedicated dialog."""
        dialog = DirsearchResultsDialog(json_data, target_context, self)
        dialog.exec()

    def _show_ffuf_results_popup(self, json_data):
        """Shows the results of an ffuf scan in a dedicated dialog."""
        dialog = FfufResultsDialog(json_data, self)
        dialog.exec()

    def _show_nuclei_results_popup(self, json_data):
        """Shows the results of a nuclei scan in a dedicated dialog."""
        dialog = NucleiResultsDialog(json_data, self)
        dialog.exec()

    def _show_trufflehog_results_popup(self, json_data):
        """Shows the results of a trufflehog scan in a dedicated dialog."""
        dialog = TruffleHogResultsDialog(json_data, self)
        dialog.exec()

    def _show_enum4linux_ng_results_popup(self, json_data, target_context):
        """Shows the results of an enum4linux-ng scan in a dedicated dialog."""
        dialog = Enum4LinuxNGResultsDialog(json_data, target_context, self)
        dialog.exec()

    def _show_dnsrecon_results_popup(self, json_data, target_context):
        """Shows the results of a dnsrecon scan in a dedicated dialog."""
        dialog = DnsReconResultsDialog(json_data, target_context, self)
        dialog.exec()

    def _show_sherlock_results_popup(self, csv_data, target_context):
        """Shows the results of a sherlock scan in a dedicated dialog."""
        dialog = SherlockResultsDialog(csv_data, target_context, self)
        dialog.exec()

    def _handle_send_results(self, ans, unans):
        self.send_results_widget.clear()
        for i, (s, r) in enumerate(ans):
            self.send_results_widget.addTopLevelItem(QTreeWidgetItem([str(i+1), s.summary(), r.summary()]))
        start_num = len(ans)
        for i, s in enumerate(unans):
            self.send_results_widget.addTopLevelItem(QTreeWidgetItem([str(start_num+i+1), s.summary(), "No response"]))

    def _handle_send_finished(self):
        self.send_btn.setEnabled(True)
        self.send_cancel_btn.setEnabled(False)

    def _handle_status_update(self, tool_name, status_text):
        widgets = {'trace': self.trace_status, 'scan': self.scan_status, 'arp': self.arp_status,
                   'flood': self.flood_status, 'fw': self.fw_status, 'wifi_scan': self.wifi_scan_status,
                   'deauth': self.deauth_status, 'arp_spoof': self.arp_spoof_status,
                   'bf': self.bf_status_label, 'ps': self.ps_status_label, 'cve_import': self.cve_import_status_label}
        if tool_name in widgets:
            widgets[tool_name].setText(status_text)

    def _handle_speed_test_servers_result(self, servers):
        """Populates the server combo box with the fetched server list."""
        # This check is necessary because the user might have closed the app
        # while the thread was running.
        if not hasattr(self, 'speed_test_server_combo'):
            return

        self.speed_test_server_combo.clear()
        self.speed_test_server_combo.addItem("Auto Select", None)
        for server in servers:
            display_text = f"{server['id']}: {server['name']} ({server['location']}) [{server['distance']:.2f} km]"
            self.speed_test_server_combo.addItem(display_text, server['id'])

    def _handle_speed_test_output(self, line):
        """Parses a line of output from speedtest-cli and updates the UI."""
        if not hasattr(self, 'speed_test_graph'): # Check if UI is still there
            return

        try:
            data = json.loads(line)
            type = data.get("type")
            timestamp = datetime.fromisoformat(data.get("timestamp").replace("Z", "+00:00"))

            if type == "ping":
                ping = data.get("ping", {}).get("latency")
                if ping is not None:
                    self.speed_test_ping_value.setText(f"{ping:.2f} ms")
                    self.current_ping = ping


            elif type == "download" or type == "upload":
                speed_mbps = data.get(type, {}).get("bandwidth", 0) * 8 / 1_000_000
                is_download = type == "download"

                # Update live value
                (self.speed_test_download_value if is_download else self.speed_test_upload_value).setText(f"{speed_mbps:.2f} Mbps")

                # Add to data series for graphing
                (self.download_series if is_download else self.upload_series).append(speed_mbps)
                self.time_series.append(timestamp)

                # Update graph and stats
                self._update_speed_test_graph()
                self._update_speed_test_stats(is_download)


            elif type == "result":
                self.speed_test_start_btn.setText("Start Test")
                self.speed_test_start_btn.setEnabled(True)
                # You can add final processing here if needed, like saving the log URL
                # url = data.get("result", {}).get("url")


        except (json.JSONDecodeError, AttributeError):
            # Ignore lines that are not valid JSON or don't have expected structure
            pass


    def _update_speed_test_graph(self):
        """Updates the PyQtGraph widget with the current data."""
        # Limit the data points to the last N to keep the graph readable
        max_points = 100
        start_index = max(0, len(self.time_series) - max_points)

        # Convert timestamps to seconds since the first data point for plotting
        if self.time_series:
            time_data = [(t - self.time_series[0]).total_seconds() for t in self.time_series[start_index:]]

            self.download_curve.setData(time_data, self.download_series[start_index:])
            self.upload_curve.setData(time_data, self.upload_series[start_index:])


    def _update_speed_test_stats(self, is_download):
        """Updates the average, top, and low speed labels."""
        series = self.download_series if is_download else self.upload_series
        avg_label = self.speed_test_avg_download_value if is_download else self.speed_test_avg_upload_value
        top_label = self.speed_test_top_download_value if is_download else self.speed_test_top_upload_value
        low_label = self.speed_test_low_download_value if is_download else self.speed_test_low_upload_value

        if not series:
            return

        avg_speed = sum(series) / len(series)
        top_speed = max(series)
        low_speed = min(s for s in series if s > 0) # Avoid 0s before test starts

        avg_label.setText(f"{avg_speed:.2f} Mbps")
        top_label.setText(f"{top_speed:.2f} Mbps")
        low_label.setText(f"{low_speed:.2f} Mbps")

    def _toggle_speed_test(self):
        """Starts or stops the internet speed test."""
        if hasattr(self, 'speed_test_thread') and self.speed_test_thread.isRunning():
            self.speed_test_thread.stop()
            self.speed_test_start_btn.setText("Start Test")
            self.speed_test_start_btn.setEnabled(True)
            return

        # --- Reset UI and data for a new test ---
        self.speed_test_start_btn.setText("Stop Test")
        self.download_series.clear()
        self.upload_series.clear()
        self.time_series.clear()
        self.download_curve.clear()
        self.upload_curve.clear()

        # Reset labels
        self.speed_test_ping_value.setText("N/A")
        self.speed_test_download_value.setText("0.00 Mbps")
        self.speed_test_upload_value.setText("0.00 Mbps")
        self.speed_test_avg_download_value.setText("0.00 Mbps")
        self.speed_test_top_download_value.setText("0.00 Mbps")
        self.speed_test_low_download_value.setText("0.00 Mbps")
        self.speed_test_avg_upload_value.setText("0.00 Mbps")
        self.speed_test_top_upload_value.setText("0.00 Mbps")
        self.speed_test_low_upload_value.setText("0.00 Mbps")

        # Get selected server
        server_id = self.speed_test_server_combo.currentData()

        # --- Start the test thread ---
        self.speed_test_thread = SpeedTestThread(server_id)
        self.speed_test_thread.output_ready.connect(self._handle_speed_test_output)
        # Handle thread finishing (normally or stopped)
        self.speed_test_thread.finished.connect(lambda: self.speed_test_start_btn.setText("Start Test"))
        self.speed_test_thread.start()

    def _handle_v2ray_test_result(self, original_item_text, result_text):
        """Finds the server item and updates its text with the test result."""
        for i in range(self.v2ray_server_list.topLevelItemCount()):
            item = self.v2ray_server_list.topLevelItem(i)
            if item.text(0) == f"[Testing...] {original_item_text}":
                item.setText(0, original_item_text) # Restore original text
                item.setText(4, result_text) # Set latency
                break
            elif item.text(0) == original_item_text:
                item.setText(4, result_text) # Set latency
                break

    def _check_command_exists(self, command):
        """Checks if a command exists in the system's PATH."""
        if not shutil.which(command):
            QMessageBox.critical(self, "Command Not Found", f"The command '{command}' was not found. Please ensure it is installed and in your system's PATH.")
            return False
        return True

    def _toggle_l2tp_connection(self, sudo_password=None):
        """Handles the logic for connecting and disconnecting the L2TP/IPsec VPN."""
        if not self._check_command_exists("ipsec") or not self._check_command_exists("systemctl"):
            return
        if self.l2tp_process and self.l2tp_process.poll() is None:
            self._disconnect_l2tp(sudo_password)
            return

        server = self.l2tp_server_edit.text().strip()
        user = self.l2tp_user_edit.text().strip()
        password = self.l2tp_pass_edit.text() # No strip
        psk = self.l2tp_psk_edit.text() # No strip

        if not all([server, user, password, psk]):
            QMessageBox.warning(self, "Input Error", "Gateway, Username, Password, and IPsec PSK are all required.")
            return

        # Prompt for sudo password if not already provided
        if not sudo_password and sys.platform != "win32":
            self._run_command_with_sudo_prompt(None, self._toggle_l2tp_connection, 'l2tp_vpn')
            return

        self.l2tp_status_label.setText("Status: Configuring...")
        QApplication.processEvents()

        # --- Generate Config Files ---
        # Note: These paths are placeholders. The actual paths will be in /etc.
        self.l2tp_config_files = {
            "ipsec.conf": f"/etc/ipsec.d/l2tp-zurvan.conf",
            "ipsec.secrets": f"/etc/ipsec.d/l2tp-zurvan.secrets",
            "xl2tpd.conf": f"/etc/xl2tpd/l2tp-zurvan.conf",
            "options.l2tpd.client": f"/etc/ppp/options.l2tpd.zurvan"
        }

        ipsec_conf_content = f"""
conn L2TP-PSK-Zurvan
    authby=secret
    pfs=no
    auto=add
    keyingtries=3
    dpddelay=30
    dpdtimeout=120
    dpdaction=clear
    rekey=yes
    ike=aes256-sha1-modp1024!
    phase2alg=aes256-sha1
    type=transport
    left=%defaultroute
    leftprotoport=17/1701
    right={server}
    rightprotoport=17/1701
"""
        ipsec_secrets_content = f': PSK "{psk}"\n'
        xl2tpd_conf_content = f"""
[lac L2TP-VPN-Zurvan]
lns = {server}
ppp debug = yes
pppoptfile = {self.l2tp_config_files["options.l2tpd.client"]}
length bit = yes
"""
        ppp_options_content = f"""
ipcp-accept-local
ipcp-accept-remote
refuse-eap
require-mschap-v2
noccp
noauth
idle 1800
mtu 1410
mru 1410
defaultroute
usepeerdns
debug
connect-delay 5000
name {user}
password {password}
"""

        # --- Write Configs to Temp Files and then Sudo-Copy them ---
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                tmp_ipsec_conf = os.path.join(tmpdir, "l2tp-zurvan.conf")
                with open(tmp_ipsec_conf, "w") as f: f.write(ipsec_conf_content)

                tmp_ipsec_secrets = os.path.join(tmpdir, "l2tp-zurvan.secrets")
                with open(tmp_ipsec_secrets, "w") as f: f.write(ipsec_secrets_content)

                tmp_xl2tpd_conf = os.path.join(tmpdir, "l2tp-zurvan.conf")
                with open(tmp_xl2tpd_conf, "w") as f: f.write(xl2tpd_conf_content)

                tmp_ppp_options = os.path.join(tmpdir, "options.l2tpd.zurvan")
                with open(tmp_ppp_options, "w") as f: f.write(ppp_options_content)

                # Now, copy these files to their final destinations using sudo
                commands_to_run = [
                    ["cp", tmp_ipsec_conf, self.l2tp_config_files["ipsec.conf"]],
                    ["cp", tmp_ipsec_secrets, self.l2tp_config_files["ipsec.secrets"]],
                    ["chmod", "600", self.l2tp_config_files["ipsec.secrets"]],
                    ["mkdir", "-p", "/etc/xl2tpd"],
                    ["cp", tmp_xl2tpd_conf, self.l2tp_config_files["xl2tpd.conf"]],
                    ["mkdir", "-p", "/etc/ppp"],
                    ["cp", tmp_ppp_options, self.l2tp_config_files["options.l2tpd.client"]]
                ]

                for cmd in commands_to_run:
                    self._execute_l2tp_command(cmd, sudo_password, f"setup_{cmd[0]}")

            self.l2tp_status_label.setText("Status: Connecting...")

            # Restart strongswan to load new configs
            restart_cmd = ["systemctl", "restart", "strongswan-starter"]
            self.worker = WorkerThread(self._execute_l2tp_command, args=(restart_cmd, sudo_password, "ipsec_restart"))
            self.worker.start()

        except Exception as e:
            self.l2tp_status_label.setText(f"Status: <font color='red'>Error: {e}</font>")


    def _disconnect_l2tp(self, sudo_password=None):
        self.l2tp_status_label.setText("Status: Disconnecting...")

        # Stop xl2tpd
        if self.l2tp_process:
            self.l2tp_process.terminate()
            self.l2tp_process.wait()
            self.l2tp_process = None

        # Bring down IPsec connection
        ipsec_down_cmd = ["ipsec", "auto", "--down", "L2TP-PSK-Zurvan"]
        self.worker = WorkerThread(self._execute_l2tp_command, args=(ipsec_down_cmd, sudo_password, "ipsec_down"))
        self.worker.start()

        self._cleanup_l2tp_configs(sudo_password)

        self.l2tp_status_label.setText("Status: Disconnected")
        self.l2tp_connect_btn.setText("Connect")

    def _cleanup_l2tp_configs(self, sudo_password):
        """Removes the temporary L2TP/IPsec configuration files."""
        if not hasattr(self, 'l2tp_config_files'):
            return

        commands = []
        for path in self.l2tp_config_files.values():
            commands.append(["rm", "-f", path])

        # Also try to remove the include from the main ipsec.conf
        commands.append(["sed", "-i", f"'/include {self.l2tp_config_files['ipsec.conf']}/d'", "/etc/ipsec.conf"])

        for cmd in commands:
            try:
                self._execute_l2tp_command(cmd, sudo_password, "cleanup")
            except Exception as e:
                logging.warning(f"Failed to cleanup L2TP config file {cmd[-1]}: {e}")

    def _execute_l2tp_command(self, command, sudo_password, step):
        """Executes a single command for the L2TP connection process."""
        q = self.tool_results_queue
        full_command = ["sudo", "-S"] + command

        try:
            process = subprocess.Popen(full_command, stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
            if sudo_password:
                stdout, _ = process.communicate(input=sudo_password + '\n')
            else:
                 stdout, _ = process.communicate()

            if process.returncode != 0:
                raise Exception(f"Command failed with code {process.returncode}:\n{stdout}")

            q.put(('l2tp_step_finished', step, sudo_password))

        except Exception as e:
            q.put(('error', 'L2TP/IPsec Error', str(e)))
            q.put(('l2tp_status', f"Status: <font color='red'>Error at step: {step}</font>"))

    def _create_tun_system_tab(self):
        """Creates the UI for the TUN System."""
        widget = QWidget()
        layout = QFormLayout(widget)

        self.tun_ip_edit = QLineEdit("10.0.0.1")
        layout.addRow("TUN IP Address:", self.tun_ip_edit)

        self.tun_netmask_edit = QLineEdit("255.255.255.0")
        layout.addRow("TUN Netmask:", self.tun_netmask_edit)

        self.tun_create_btn = QPushButton("Create TUN Device")
        self.tun_status_label = QLabel("Status: Inactive")
        layout.addRow(self.tun_create_btn)
        layout.addRow(self.tun_status_label)

        self.tun_create_btn.clicked.connect(self._toggle_tun_device)

        return widget

    def _toggle_tun_device(self):
        """Creates or destroys the TUN device."""
        if hasattr(self, 'tun_fd') and self.tun_fd:
            self._destroy_tun_device()
        else:
            self._create_tun_device()

    def _create_tun_device(self):
        """Creates and configures a TUN device."""
        if not self._require_root("TUN System"):
            return

        try:
            import fcntl
            import struct

            TUNSETIFF = 0x400454ca
            IFF_TUN = 0x0001
            IFF_NO_PI = 0x1000

            # Create the TUN device
            self.tun_fd = os.open("/dev/net/tun", os.O_RDWR)
            ifr = struct.pack('16sH', b'tun%d', IFF_TUN | IFF_NO_PI)
            self.ifname = fcntl.ioctl(self.tun_fd, TUNSETIFF, ifr)
            self.ifname = self.ifname.decode('utf-8').strip('\x00')

            self.tun_status_label.setText(f"Status: Created {self.ifname}")

            # Configure the TUN device
            ip_addr = self.tun_ip_edit.text()
            netmask = self.tun_netmask_edit.text()
            subprocess.check_call(f"ip addr add {ip_addr}/{netmask} dev {self.ifname}", shell=True)
            subprocess.check_call(f"ip link set dev {self.ifname} up", shell=True)

            self.tun_create_btn.setText("Destroy TUN Device")
            self.tun_status_label.setText(f"Status: {self.ifname} is up")

        except Exception as e:
            self.tun_status_label.setText(f"Status: <font color='red'>Error: {e}</font>")
            if hasattr(self, 'tun_fd') and self.tun_fd:
                os.close(self.tun_fd)
                self.tun_fd = None

    def _destroy_tun_device(self):
        """Destroys the TUN device."""
        if hasattr(self, 'tun_fd') and self.tun_fd:
            try:
                os.close(self.tun_fd)
                self.tun_fd = None
                self.tun_create_btn.setText("Create TUN Device")
                self.tun_status_label.setText("Status: Inactive")
            except Exception as e:
                self.tun_status_label.setText(f"Status: <font color='red'>Error: {e}</font>")

    def _handle_l2tp_step(self, step, sudo_password):
        """Handles the multi-step connection process for L2TP."""
        if step == "ipsec_restart":
            self.l2tp_status_label.setText("Status: IPsec service restarted. Bringing up tunnel...")
            up_cmd = ["ipsec", "auto", "--up", "L2TP-PSK-Zurvan"]
            self.worker = WorkerThread(self._execute_l2tp_command, args=(up_cmd, sudo_password, "ipsec_up"))
            self.worker.start()

        elif step == "ipsec_up":
            self.l2tp_status_label.setText("Status: IPsec tunnel established. Starting L2TP daemon...")
            # Now start xl2tpd in the background
            try:
                # We need to run xl2tpd with sudo, but Popen needs careful handling
                # This is a simplified approach; a production app might use a helper script
                # to manage the daemon.
                xl2tpd_cmd = ["sudo", "-S", "xl2tpd", "-c", self.l2tp_config_files['xl2tpd.conf'], "-D"]
                self.l2tp_process = subprocess.Popen(xl2tpd_cmd, stdin=subprocess.PIPE, text=True)
                self.l2tp_process.stdin.write(sudo_password + '\n')
                self.l2tp_process.stdin.flush()

                # A small delay to allow the daemon to start
                time.sleep(2)

                if self.l2tp_process.poll() is None:
                    self.l2tp_status_label.setText("Status: <font color='green'>Connected</font>")
                    self.l2tp_connect_btn.setText("Disconnect")
                else:
                    raise Exception("xl2tpd daemon failed to start.")

            except Exception as e:
                self.tool_results_queue.put(('error', 'L2TP/IPsec Error', f"Failed to start xl2tpd: {e}"))
                self.l2tp_status_label.setText(f"Status: <font color='red'>Error starting L2TP daemon</font>")

        elif step == "ipsec_down":
             self.l2tp_status_label.setText("Status: Disconnected")
             self.l2tp_connect_btn.setText("Connect")

    def _toggle_wireguard_connection(self, sudo_password=None):
        """Handles the logic for connecting and disconnecting the WireGuard VPN."""
        if not self._check_command_exists("wg-quick"):
            return

        # Use the generic _execute_command_thread's process tracking
        if hasattr(self, 'wireguard_log_process') and self.wireguard_log_process and self.wireguard_log_process.poll() is None:
            self._disconnect_wireguard(sudo_password)
            return

        config_file = self.wireguard_config_edit.text().strip()
        if not config_file or not os.path.exists(config_file):
            QMessageBox.warning(self, "Input Error", "A valid WireGuard configuration file is required.")
            return

        if not sudo_password and sys.platform != "win32":
            self._run_command_with_sudo_prompt(None, self._toggle_wireguard_connection, 'wireguard_vpn')
            return

        self.wireguard_status_label.setText("Status: Connecting...")
        self.wireguard_output_console.clear()
        QApplication.processEvents()

        command = ["wg-quick", "up", config_file]
        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'wireguard_log', f"WireGuard session for {os.path.basename(config_file)}", self.wireguard_output_console, sudo_password
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

        self.wireguard_status_label.setText("Status: <font color='orange'>Process Started</font>")
        self.wireguard_connect_btn.setText("Disconnect")


    def _disconnect_wireguard(self, sudo_password=None):
        """Stops the WireGuard connection."""
        self.cancel_tool() # Kills the 'up' process

        config_file = self.wireguard_config_edit.text().strip()
        if not config_file: return # Can't disconnect if there's no file path

        if not sudo_password and sys.platform != "win32":
            # Don't need a cancel handler for a quick disconnect command
            self._run_command_with_sudo_prompt(None, self._disconnect_wireguard, 'wireguard_vpn_disconnect')
            return

        # Run the 'down' command in a separate, short-lived thread
        command = ["wg-quick", "down", config_file]
        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'wireguard_log', "wg-quick down", self.wireguard_output_console, sudo_password
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

        self.wireguard_status_label.setText("Status: Disconnected")
        self.wireguard_connect_btn.setText("Connect")

    def _toggle_openvpn_connection(self, sudo_password=None):
        """Handles the logic for connecting and disconnecting the OpenVPN client."""
        if not self._check_command_exists("openvpn"):
            return
        if self.openvpn_process and self.openvpn_process.poll() is None:
            self._disconnect_openvpn()
            return

        config_file = self.openvpn_config_edit.text().strip()
        if not config_file or not os.path.exists(config_file):
            QMessageBox.warning(self, "Input Error", "A valid OpenVPN configuration file is required.")
            return

        if not sudo_password and sys.platform != "win32":
            self._run_command_with_sudo_prompt(None, self._toggle_openvpn_connection, 'openvpn_vpn')
            return

        self.openvpn_status_label.setText("Status: Connecting...")
        self.openvpn_output_console.clear()
        QApplication.processEvents()

        # We will run the process in a thread to monitor its output
        command = ["openvpn", "--config", config_file, "--log-append", "/dev/stdout"] # Force output to stdout
        self.worker = WorkerThread(self._execute_command_thread, args=(
            command, 'openvpn_log', f"OpenVPN session for {os.path.basename(config_file)}", self.openvpn_output_console, sudo_password
        ))
        self.active_threads.append(self.worker)
        self.worker.start()

        # Update UI immediately, success/failure will be seen in logs
        self.openvpn_status_label.setText("Status: <font color='orange'>Process Started</font>")
        self.openvpn_connect_btn.setText("Disconnect")


    def _disconnect_openvpn(self):
        """Stops the OpenVPN process."""
        # The generic cancel_tool will find and kill the process named 'openvpn_log_process'
        self.cancel_tool()
        self.openvpn_status_label.setText("Status: Disconnected")
        self.openvpn_connect_btn.setText("Connect")
        self.openvpn_output_console.appendPlainText("\n--- Connection Terminated by User ---")

    def _handle_clear_update(self, tool_name):
        widgets = {'trace': self.trace_tree, 'scan': self.scan_tree, 'arp': self.arp_tree,
                   'fw': self.fw_tree, 'wifi_scan': self.wifi_tree}
        if tool_name in widgets:
            widgets[tool_name].clear()

    def _handle_result_update(self, tool_name, result_data):
        widgets = {'trace': self.trace_tree, 'scan': self.scan_tree, 'fw': self.fw_tree,
                   'wifi_scan': self.wifi_tree, 'ps': self.ps_tree}
        if tool_name in widgets:
            widgets[tool_name].addTopLevelItem(QTreeWidgetItem([str(x) for x in result_data]))

    def _handle_generic_output(self, tool_name, line):
        """A generic handler for tool output that finds the correct console widget."""
        console_widgets = {
            'nmap_scan': self.nmap_output_console,
            'sublist3r_scan': self.sublist3r_output,
            'subfinder_scan': self.subfinder_output,
            'httpx_scan': self.httpx_output,
            'rustscan_scan': self.rustscan_output,
            'dirsearch_scan': self.dirsearch_output_console,
            'ffuf_scan': self.ffuf_output_console,
            'enum4linux_ng_scan': self.enum4linux_ng_output_console,
            'dnsrecon_scan': self.dnsrecon_output_console,
            'fierce_scan': self.fierce_output_console,
            'nikto_scan': self.nikto_output_console,
            'gobuster_scan': self.gobuster_output_console,
            'whatweb_scan': self.whatweb_output_console,
            'theharvester_scan': self.theharvester_output_console,
            'masscan_scan': self.masscan_output_console,
            'sqlmap_scan': self.sqlmap_output_console,
            'hashcat_scan': self.hashcat_output_console,
            'nuclei_scan': self.nuclei_output_console,
            'trufflehog_scan': self.trufflehog_output_console,
            'jtr_scan': self.jtr_output_console,
            'hydra_scan': self.hydra_output_console,
            'sherlock_scan': self.sherlock_output_console,
            'spiderfoot_scan': self.spiderfoot_output_console,
            'arp_scan_cli_scan': self.arp_scan_cli_output_console,
            'wifite_scan': self.wifite_output_console,
            'snort_ids': self.snort_output_console,
            'nmap_evasion_scan': self.nmap_evasion_output_console,
        }
        if tool_name in console_widgets:
            widget = console_widgets[tool_name]
            widget.insertPlainText(line)
            widget.verticalScrollBar().setValue(widget.verticalScrollBar().maximum())
        else:
            logging.warning(f"No console widget found for tool output: {tool_name}")

    def _handle_arp_results(self, results):
        for res in results:
            self.arp_tree.addTopLevelItem(QTreeWidgetItem([res['ip'], res['mac'], res['status']]))

    def _handle_crunch_finished(self, outfile, returncode):
        if returncode == 0:
            self.wpa_crack_output.appendPlainText(f"Crunch finished successfully. Wordlist saved to:\n{outfile}")
            self.wpa_wordlist_edit.setText(outfile)
        else:
            self.wpa_crack_output.appendPlainText(f"Crunch finished with an error (code: {returncode}). Check zurvan.log for details.")

    def _handle_ps_worker_finished(self, total_threads):
        with self.ps_thread_lock:
            self.ps_finished_threads += 1
            if self.ps_finished_threads >= total_threads:
                if self.tool_stop_event.is_set():
                    self.ps_status_label.setText("Ping sweep canceled.")
                else:
                    self.ps_status_label.setText("Ping sweep complete.")

                target = self.ps_target_edit.text()
                results = f"Found {self.ps_tree.topLevelItemCount()} hosts."
                self.tool_results_queue.put(('tool_finished', 'ping_sweep', target, results))

    def _handle_flood_thread_finished(self, total_threads):
        with self.thread_finish_lock:
            self.finished_thread_count += 1
            if self.finished_thread_count >= total_threads:
                self.is_tool_running = False
                self.flood_button.setEnabled(True)
                self.stop_flood_button.setEnabled(False)
                if self.tool_stop_event.is_set():
                    self.flood_status.setText("Flood Canceled.")
                else:
                    self.flood_status.setText("Flood complete.")

    def _handle_tool_finished(self, tool, target=None, results=""):
        """
        Handles the 'tool_finished' signal from any worker thread.
        Resets UI state and logs the test to the activity log.
        """
        # Log the action to the activity log table
        if self.current_user and target is not None:
            tool_categories = {
                'nmap_scan': 'Network Scan', 'scanner': 'Network Scan', 'arp_scan': 'Network Scan',
                'ping_sweep': 'Network Scan', 'traceroute': 'Network Scan', 'masscan_scan': 'Network Scan',
                'arp_scan_cli_scan': 'Network Scan',
                'sublist3r_scan': 'Reconnaissance', 'subfinder_scan': 'Reconnaissance', 'dnsrecon_scan': 'Reconnaissance',
                'fierce_scan': 'Reconnaissance', 'sherlock_scan': 'OSINT', 'spiderfoot_scan': 'OSINT', 'trufflehog_scan': 'OSINT',
                'dirsearch_scan': 'Web Scan', 'ffuf_scan': 'Web Scan', 'nikto_scan': 'Web Scan',
                'gobuster_scan': 'Web Scan', 'whatweb_scan': 'Web Scan', 'httpx_scan': 'Web Scan',
                'flooder': 'Offensive', 'firewall_tester': 'Offensive', 'arp_spoof': 'Offensive',
                'deauth': 'Offensive', 'beacon_flood': 'Offensive', 'sqlmap_scan': 'Offensive',
                'hashcat_scan': 'Password Cracking', 'jtr_scan': 'Password Cracking', 'hydra_scan': 'Password Cracking',
                'wifite_scan': 'Wireless', 'wpa_crack': 'Wireless', 'krack_scanner': 'Wireless', 'wifi_scan': 'Wireless',
                'fetch_threats': 'Threat Intelligence', 'exploit_search': 'Threat Intelligence', 'cve_import': 'Threat Intelligence',
                'aggregation': 'Reporting', 'report_generation': 'Reporting',
                'send': 'Packet Operations', 'cve_db_update': 'System', 'tool_finished': 'System'
            }
            category = tool_categories.get(tool, 'Miscellaneous')
            action = tool.replace('_', ' ').title()

            if not isinstance(results, str):
                results = json.dumps(results, indent=2)

            if len(results) > 10000:
                results = results[:10000] + "\n... (truncated)"

            database.log_activity(
                user_id=self.current_user['id'],
                category=category,
                action=action,
                target=target,
                details=results
            )

        # Special handling for tools that need their output parsed before UI is updated
        if tool == 'sublist3r_scan' and results:
            subdomains = re.findall(r'^[a-zA-Z0-9.\-]+\.' + re.escape(target), results, re.MULTILINE)
            unique_subdomains = sorted(list(set(subdomains)))
            self.tool_results_queue.put(('sublist3r_results', target, unique_subdomains))
        elif tool == 'subfinder_scan' and results:
            # The target is the domain for subfinder
            subdomains = [line for line in results.splitlines() if target in line and ' ' not in line and not line.startswith('$')]
            unique_subdomains = sorted(list(set(subdomains)))
            # The popup is generic for any subdomain tool
            self.tool_results_queue.put(('subdomain_results', target, unique_subdomains))
        elif tool == 'httpx_scan' and results:
            # The result handler for httpx will attempt to parse JSON
            self.tool_results_queue.put(('httpx_results', results))
        elif tool == 'nuclei_scan' and results:
            self.tool_results_queue.put(('nuclei_results', results))
        elif tool == 'trufflehog_scan' and results:
            self.tool_results_queue.put(('trufflehog_results', results))

        if tool == 'aggregation':
            self.report_aggregate_btn.setEnabled(True)
            self.status_bar.showMessage("Result aggregation and enrichment complete.", 5000)
            return

        if tool == 'cve_db_update':
            self.update_cve_db_btn.setEnabled(True)
            self.status_bar.showMessage("CVE DB update finished.", 5000)
            self.is_tool_running = False # Explicitly set this as it's a special case
            return

        self.is_tool_running = False
        buttons = {'traceroute': self.trace_button, 'scanner': self.scan_button, 'arp_scan': self.arp_scan_button,
                   'flooder': self.flood_button, 'fw_tester': self.fw_test_button, 'wifi_scan': self.wifi_scan_button,
                   'deauth': self.deauth_button, 'arp_spoof': self.arp_spoof_start_btn,
                   'beacon_flood': self.bf_start_button, 'ping_sweep': self.ps_start_button, 'nmap_scan': self.nmap_controls['start_btn'],
                   'sublist3r_scan': self.subdomain_controls['start_btn'], 'subfinder_scan': self.subfinder_controls['start_btn'], 'httpx_scan': self.httpx_controls['start_btn'], 'trufflehog_scan': self.trufflehog_controls['start_btn'], 'rustscan_scan': self.rustscan_controls['start_btn'], 'dirsearch_scan': self.dirsearch_controls['start_btn'], 'ffuf_scan': self.ffuf_controls['start_btn'], 'jtr_scan': self.jtr_controls['start_btn'], 'hydra_scan': self.hydra_controls['start_btn'], 'enum4linux_ng_scan': self.enum4linux_ng_controls['start_btn'], 'dnsrecon_scan': self.dnsrecon_controls['start_btn'], 'fierce_scan': self.fierce_controls['start_btn'], 'sherlock_scan': self.sherlock_controls['start_btn'], 'spiderfoot_scan': self.spiderfoot_controls['start_btn'], 'arp_scan_cli_scan': self.arp_scan_cli_controls['start_btn'], 'wifite_scan': self.wifite_controls['start_btn'], 'nikto_scan': self.nikto_controls['start_btn'], 'gobuster_scan': self.gobuster_controls['start_btn'], 'sqlmap_scan': self.sqlmap_controls['start_btn'], 'whatweb_scan': self.whatweb_controls['start_btn'], 'hashcat_scan': self.hashcat_controls['start_btn'], 'masscan_scan': self.masscan_controls['start_btn'], 'nuclei_scan': self.nuclei_controls['start_btn'],
                   'exploit_search': self.exploitdb_search_button, 'fetch_threats': self.fetch_threats_btn}
        cancel_buttons = {'scanner': self.scan_cancel_button, 'flooder': self.stop_flood_button,
                          'arp_spoof': self.arp_spoof_stop_btn, 'beacon_flood': self.bf_stop_button,
                          'ping_sweep': self.ps_cancel_button, 'fw_tester': self.fw_cancel_button,
                          'traceroute': self.trace_cancel_button, 'wifi_scan': self.wifi_scan_stop_button, 'nmap_scan': self.nmap_controls['cancel_btn'],
                          'sublist3r_scan': self.subdomain_controls['cancel_btn'], 'subfinder_scan': self.subfinder_controls['cancel_btn'], 'httpx_scan': self.httpx_controls['cancel_btn'], 'trufflehog_scan': self.trufflehog_controls['stop_btn'], 'rustscan_scan': self.rustscan_controls['cancel_btn'], 'dirsearch_scan': self.dirsearch_controls['stop_btn'], 'ffuf_scan': self.ffuf_controls['stop_btn'], 'jtr_scan': self.jtr_controls['stop_btn'], 'hydra_scan': self.hydra_controls['stop_btn'], 'enum4linux_ng_scan': self.enum4linux_ng_controls['stop_btn'], 'dnsrecon_scan': self.dnsrecon_controls['stop_btn'], 'fierce_scan': self.fierce_controls['stop_btn'], 'sherlock_scan': self.sherlock_controls['stop_btn'], 'spiderfoot_scan': self.spiderfoot_controls['stop_btn'], 'arp_scan_cli_scan': self.arp_scan_cli_controls['stop_btn'], 'wifite_scan': self.wifite_controls['stop_btn'], 'nikto_scan': self.nikto_controls['stop_btn'], 'gobuster_scan': self.gobuster_controls['stop_btn'], 'sqlmap_scan': self.sqlmap_controls['stop_btn'], 'whatweb_scan': self.whatweb_controls['stop_btn'], 'hashcat_scan': self.hashcat_controls['stop_btn'], 'masscan_scan': self.masscan_controls['stop_btn'], 'nuclei_scan': self.nuclei_controls['stop_btn']}

        if tool == 'lab_chain':
            self.lab_run_chain_btn.setEnabled(True)
            self.status_bar.showMessage("LAB chain finished.", 5000)
            return

        if tool == 'jtr_scan':
            self.jtr_controls['start_btn'].setEnabled(True)
            self.jtr_controls['show_btn'].setEnabled(True)
            self.jtr_controls['stop_btn'].setEnabled(False)
            return

        if tool == 'arp_spoof':
            if self.arp_spoof_current_victim and self.arp_spoof_current_target:
                self._restore_arp(self.arp_spoof_current_victim, self.arp_spoof_current_target)
                self.arp_spoof_current_victim = None
                self.arp_spoof_current_target = None

        if tool in buttons:
            buttons[tool].setEnabled(True)
        if tool in cancel_buttons:
            cancel_buttons[tool].setEnabled(False)

        if self.tool_stop_event.is_set():
            status_labels = {'scanner': self.scan_status, 'traceroute': self.trace_status}
            if tool in status_labels:
                status_labels[tool].setText("Canceled by user.")

    def _handle_cve_search_status(self, status_text):
        self.status_bar.showMessage(status_text, 5000)

    def _handle_cve_result(self, result_data, cve_object):
        """Adds a CVE result to the table and stores the full object."""
        item = QTreeWidgetItem(result_data)
        item.setData(0, Qt.ItemDataRole.UserRole, cve_object) # Store the full object
        self.cve_results_table.addTopLevelItem(item)

    def _handle_exploit_search_status(self, status_text):
        self.status_bar.showMessage(status_text, 5000)

    def _handle_exploit_search_results(self, results):
        """Adds exploit search results to the table."""
        self.exploitdb_results_table.clear()
        for result in results:
            item = QTreeWidgetItem(result)
            self.exploitdb_results_table.addTopLevelItem(item)

    def _handle_lab_status(self, status_text):
        self.status_bar.showMessage(status_text, 0) # 0 means it stays until changed

    def _handle_error(self, title, text):
        QMessageBox.critical(self, title, text)

    def _create_export_button(self, source_widget):
        button = QPushButton("Export Results")
        button.setToolTip("Export the results to a file (CSV, HTML, PDF, DOCX).")
        button.clicked.connect(lambda: self._handle_export(source_widget))
        return button

    def _get_export_source_name(self, widget):
        """Returns a user-friendly name for a given source widget for logging purposes."""
        if widget is self.packet_list_widget:
            return "Packet Sniffer"
        if widget is self.send_results_widget:
            return "Packet Crafter"
        if widget is self.trace_tree:
            return "Traceroute"
        if widget is self.scan_tree:
            return "Port Scanner (Scapy)"
        if widget is self.arp_tree:
            return "ARP Scan (Scapy)"
        if widget is self.ps_tree:
            return "Ping Sweep"
        if widget is self.wifi_tree:
            return "Wi-Fi Scanner"
        if widget is self.fw_tree:
            return "Firewall Tester"
        # Fallback for other potential widgets
        return "Unknown View"

    def _handle_export(self, source_widget):
        if source_widget.topLevelItemCount() == 0:
            QMessageBox.information(self, "No Data", "There is no data to export.")
            return

        formats = "HTML (*.html);;CSV (*.csv);;PDF (*.pdf);;Word Document (*.docx)"
        file_path, selected_format = QFileDialog.getSaveFileName(self, "Export Results", "", formats, options=QFileDialog.Option.DontUseNativeDialog)

        if not file_path:
            return

        export_type = "Unknown"
        try:
            if 'html' in selected_format:
                export_type = "HTML"
                self._export_to_html(source_widget, file_path)
            elif 'csv' in selected_format:
                export_type = "CSV"
                self._export_to_csv(source_widget, file_path)
            elif 'pdf' in selected_format:
                export_type = "PDF"
                self._export_to_pdf(source_widget, file_path)
            elif 'docx' in selected_format:
                export_type = "DOCX"
                self._export_to_docx(source_widget, file_path)
            else:
                QMessageBox.warning(self, "Unsupported Format", "Selected file format is not supported.")
                return

            self.status_bar.showMessage(f"Successfully exported results to {file_path}")
            if self.current_user:
                source_name = self._get_export_source_name(source_widget)
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='File Operation',
                    action=f'Export Results ({export_type})',
                    target=file_path,
                    details=f"Exported {source_widget.topLevelItemCount()} rows from '{source_name}'."
                )
        except NameError:
            logging.error("Export failed due to missing optional dependencies.", exc_info=True)
            QMessageBox.critical(self, "Dependency Error", "Optional libraries for PDF/DOCX export are not installed.\nPlease run: pip install reportlab python-docx")
        except Exception as e:
            logging.error(f"Failed to export results: {e}", exc_info=True)
            QMessageBox.critical(self, "Export Error", f"An error occurred during export:\n{e}")

    def _export_to_csv(self, tree_widget, file_path):
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]
            writer.writerow(header)
            for i in range(tree_widget.topLevelItemCount()):
                item = tree_widget.topLevelItem(i)
                row = [item.text(j) for j in range(tree_widget.columnCount())]
                writer.writerow(row)

    def _export_to_html(self, tree_widget, file_path):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("<html><head><title>Exported Results</title>")
            f.write("<style>body { font-family: sans-serif; } table { border-collapse: collapse; width: 100%; }")
            f.write("th, td { border: 1px solid #dddddd; text-align: left; padding: 8px; }")
            f.write("tr:nth-child(even) { background-color: #f2f2f2; }</style></head><body>")
            f.write("<h2>Exported Results</h2><table><tr>")
            header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]
            for h in header:
                f.write(f"<th>{h}</th>")
            f.write("</tr>")
            for i in range(tree_widget.topLevelItemCount()):
                f.write("<tr>")
                item = tree_widget.topLevelItem(i)
                for j in range(tree_widget.columnCount()):
                    f.write(f"<td>{item.text(j)}</td>")
                f.write("</tr>")
            f.write("</table></body></html>")

    def _export_to_pdf(self, tree_widget, file_path):
        doc = SimpleDocTemplate(file_path)
        elements = []
        styles = getSampleStyleSheet()
        elements.append(Paragraph("Exported Results", styles['h1']))

        header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]
        data = [header]
        for i in range(tree_widget.topLevelItemCount()):
            row = [tree_widget.topLevelItem(i).text(j) for j in range(tree_widget.columnCount())]
            data.append(row)

        table = Table(data)
        style = TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ])
        table.setStyle(style)
        elements.append(table)
        doc.build(elements)

    def _export_to_docx(self, tree_widget, file_path):
        document = docx.Document()
        document.add_heading('Exported Results', 0)

        header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]

        table = document.add_table(rows=1, cols=len(header))
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(header):
            hdr_cells[i].text = h

        for i in range(tree_widget.topLevelItemCount()):
            row_cells = table.add_row().cells
            item = tree_widget.topLevelItem(i)
            for j in range(tree_widget.columnCount()):
                row_cells[j].text = item.text(j)

        document.save(file_path)

    def _update_tool_targets(self):
        """Automatically updates tool target fields based on the selected interface."""
        iface_name = self.get_selected_iface()

        network_cidr = "192.168.1.0/24" # Default fallback
        if iface_name and iface_name != "Automatic":
            try:
                addrs = psutil.net_if_addrs().get(iface_name, [])
                for addr in addrs:
                    if addr.family == socket.AF_INET:
                        ip = addr.address
                        netmask = addr.netmask
                        if ip and netmask:
                            # Use ipaddress module to calculate network CIDR
                            host_iface = ipaddress.IPv4Interface(f"{ip}/{netmask}")
                            network_cidr = host_iface.network.with_prefixlen
                            logging.info(f"Updated tool targets for interface {iface_name} to {network_cidr}")
                            break # Found the IPv4 addr, no need to continue
            except Exception as e:
                logging.error(f"Could not auto-populate tool targets for {iface_name}: {e}")
                # Keep the default fallback

        # Update all relevant tool fields
        if hasattr(self, 'arp_target'):
            self.arp_target.setText(network_cidr)
        if hasattr(self, 'ps_target_edit'):
            self.ps_target_edit.setText(network_cidr)

    def take_screenshot_and_exit(self, path):
        """Takes a screenshot of the main window and then exits the application."""
        try:
            # Switch to the correct tab before taking the screenshot
            self.tab_widget.setCurrentIndex(2) # Network Tools
            # Find the SSH Manager tab within the Network Tools tab widget
            network_tools_tab_widget = self.tab_widget.widget(2)
            if isinstance(network_tools_tab_widget, QTabWidget):
                for i in range(network_tools_tab_widget.count()):
                    if network_tools_tab_widget.tabText(i) == "SSH Manager":
                        network_tools_tab_widget.setCurrentIndex(i)
                        break

            # Give it a moment to switch tabs
            QTimer.singleShot(1500, lambda: self._capture_and_exit(path))
        except Exception as e:
            logging.error(f"Failed to switch tabs for screenshot: {e}")
            self.close()

    def _capture_and_exit(self, path):
        try:
            screenshot = self.grab()
            screenshot.save(path, 'png')
            logging.info(f"Screenshot saved to {path}")
        except Exception as e:
            logging.error(f"Failed to take screenshot: {e}")
        finally:
            self.close()

    def closeEvent(self, event):
        """Shows a confirmation dialog and ensures background threads are stopped on exit."""
        logging.info("closeEvent triggered!")
        # In test mode, bypass dialogs and perform a clean shutdown
        if self.is_test_mode:
            logging.info("Test mode detected, performing clean shutdown.")
            self._clean_shutdown()
            event.accept()
            return

        # Bypass confirmation in screenshot mode
        if hasattr(sys, 'argv') and '--screenshot' in sys.argv:
            self._clean_shutdown()
            event.accept()
            return

        reply = QMessageBox.question(self, 'Exit Confirmation',
                                     "Are you sure you want to exit?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            if self.current_user:
                database.log_activity(
                    user_id=self.current_user['id'],
                    category='System',
                    action='Application Shutdown',
                    target=f"v3.0 on {platform.system()}",
                    details=f"User '{self.current_user['username']}' closed the application."
                )
            self._clean_shutdown()
            event.accept()
        else:
            logging.info("User canceled exit.")
            event.ignore()

    def _clean_shutdown(self):
        """Stops all running background threads and timers."""
        logging.info("Performing clean shutdown. Stopping threads and timers.")

        # Explicitly shut down logging before Qt objects are destroyed
        logging.shutdown()

        # Set stop event for all threads that use it
        self.tool_stop_event.set()

        # Stop all timers to prevent them from firing during shutdown
        if hasattr(self, 'results_processor') and self.results_processor and self.results_processor.isActive(): self.results_processor.stop()
        if hasattr(self, 'sniffer_ui_update_timer') and self.sniffer_ui_update_timer and self.sniffer_ui_update_timer.isActive(): self.sniffer_ui_update_timer.stop()
        if hasattr(self, 'clock_timer') and self.clock_timer and self.clock_timer.isActive(): self.clock_timer.stop()
        if hasattr(self, 'ip_check_timer') and self.ip_check_timer and self.ip_check_timer.isActive(): self.ip_check_timer.stop()
        if hasattr(self, 'auto_lock_timer') and self.auto_lock_timer and self.auto_lock_timer.isActive(): self.auto_lock_timer.stop()
        if hasattr(self, 'journal_refresh_timer') and self.journal_refresh_timer and self.journal_refresh_timer.isActive(): self.journal_refresh_timer.stop()
        if hasattr(self, 'proxy_rotation_timer') and self.proxy_rotation_timer and self.proxy_rotation_timer.isActive(): self.proxy_rotation_timer.stop()

        # Wait indefinitely for all active worker threads to finish
        for thread in self.active_threads[:]:
            if thread.isRunning():
                thread.wait()

        # Stop other persistent threads
        if self.sniffer_thread and self.sniffer_thread.isRunning():
            self.sniffer_thread.stop()
            self.sniffer_thread.wait()
        if self.channel_hopper and self.channel_hopper.isRunning():
            self.channel_hopper.stop()
            self.channel_hopper.wait()
        if self.resource_monitor_thread and self.resource_monitor_thread.isRunning():
            self.resource_monitor_thread.stop()
            self.resource_monitor_thread.wait()

        # Wait for IP fetcher threads
        if hasattr(self, 'real_ip_thread') and self.real_ip_thread.isRunning(): self.real_ip_thread.wait()
        if hasattr(self, 'proxy_ip_thread') and self.proxy_ip_thread.isRunning(): self.proxy_ip_thread.wait()

        # Terminate any remaining subprocesses
        if self.v2ray_process and self.v2ray_process.poll() is None:
            self.v2ray_process.terminate()
            self.v2ray_process.wait()

        logging.info("All threads and timers stopped. Zurvan application closing.")

    def _create_reporting_tab(self):
        """Creates the UI for the Reporting & Analysis tab."""
        widget = QWidget()
        main_layout = QHBoxLayout(widget)
        main_splitter = QSplitter(Qt.Orientation.Horizontal)

        # --- Left Panel: Configuration & Narrative ---
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(0,0,0,0)

        # Rules of Engagement Box
        roe_box = QGroupBox("Rules of Engagement (ROE)")
        roe_layout = QFormLayout(roe_box)
        self.report_client_name = QLineEdit()
        self.report_assessment_dates = QLineEdit()
        self.report_objectives = QTextEdit()
        self.report_objectives.setPlaceholderText(
"""Example:
- Objective 1: Determine the ability of a threat actor to compromise critical customer transactional data.
- Objective 2: Evaluate the integrity of the customer's order database.
- Objective 3: Assess the effectiveness of Incident Response procedures."""
        )
        self.report_in_scope = QTextEdit()
        self.report_in_scope.setPlaceholderText(
"""Example:

--- Authorized Target Space ---
- IP Range(s): 10.10.12.0/24, 10.10.13.0/24
- Domains: *.example.com
- URLs: https://www.example.com/login
- Network Segments: Corporate LAN, Guest Wi-Fi

--- Authorized Hosts ---
- All hosts not expressly restricted."""
        )
        self.report_out_of_scope = QTextEdit()
        self.report_out_of_scope.setPlaceholderText(
"""Example:

--- Explicit Restrictions ---
- No Denial of Service (DoS) attacks.
- No testing outside of business hours (9am-5pm Local Time).
- Social engineering of staff is not permitted.

--- Restricted IP Addresses ---
- 10.10.10.0/24 (HR Department)
- 10.10.11.0/24 (Accounting)

--- Restricted Hosts ---
- CRITICAL_DB_SERVER_01"""
        )
        roe_layout.addRow("Client Name:", self.report_client_name)
        roe_layout.addRow("Assessment Dates:", self.report_assessment_dates)
        roe_layout.addRow("Objectives:", self.report_objectives)
        roe_layout.addRow("In-Scope Targets:", self.report_in_scope)
        roe_layout.addRow("Out-of-Scope & Restrictions:", self.report_out_of_scope)
        left_layout.addWidget(roe_box)

        # Executive Summary Box
        summary_box = QGroupBox("Executive Summary")
        summary_layout = QVBoxLayout(summary_box)
        self.report_summary_text = QTextEdit()
        self.report_summary_text.setPlaceholderText("Write a high-level summary of the assessment's findings and recommendations for a non-technical audience, or generate one with AI.")

        ai_summary_btn = QPushButton(QIcon("icons/terminal.svg"), " Generate Summary with AI")
        ai_summary_btn.clicked.connect(self._handle_ai_summary_generation)

        summary_layout.addWidget(self.report_summary_text)
        summary_layout.addWidget(ai_summary_btn)
        left_layout.addWidget(summary_box)

        left_panel.setLayout(left_layout)

        # --- Right Panel: Findings & Generation ---
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(0,0,0,0)

        # Findings Box
        findings_box = QGroupBox("Aggregated Findings")
        findings_layout = QVBoxLayout(findings_box)

        aggregation_controls = QHBoxLayout()
        self.report_aggregate_btn = QPushButton(QIcon("icons/search.svg"), " Aggregate & Enrich Results")
        self.report_aggregate_btn.setToolTip("Scan the results from all tool outputs in the current session and enrich them with CVE and Exploit-DB information.")
        self.report_aggregate_btn.clicked.connect(self._handle_aggregation)
        aggregation_controls.addWidget(self.report_aggregate_btn)
        aggregation_controls.addStretch()
        self.offline_cve_check = QCheckBox("Use offline CVE_DB for enrichment")
        self.offline_cve_check.setToolTip("Use a local copy of CVE data for enrichment. Requires initial download via the manager below.")
        aggregation_controls.addWidget(self.offline_cve_check)
        findings_layout.addLayout(aggregation_controls)

        self.report_findings_tree = QTreeWidget()
        self.report_findings_tree.setColumnCount(4)
        self.report_findings_tree.setHeaderLabels(["Host", "Port/Service", "Vulnerability/Finding", "Details (CVE/Exploit)"])
        self.report_findings_tree.header().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.report_findings_tree.header().setStretchLastSection(True)
        findings_layout.addWidget(self.report_findings_tree)
        right_layout.addWidget(findings_box)

        # --- CVE DB Manager ---
        self.reporting_cve_manager = OfflineCveManagerWidget(self)
        right_layout.addWidget(self.reporting_cve_manager)

        # Generation Box
        generation_box = QGroupBox("Report Generation")
        generation_layout = QVBoxLayout(generation_box)

        # --- AI Generation ---
        ai_box = QGroupBox("AI-Powered Generation")
        ai_layout = QFormLayout(ai_box)
        self.ai_persona_combo = QComboBox()
        self.ai_persona_combo.addItems([
            "Default", "Cybersecurity Analyst", "Penetration Tester", "Incident Responder",
            "Malware Analyst", "Threat Hunter", "Security Auditor", "Network Engineer",
            "Technical Manager", "C-Suite Executive (CSO/CISO)", "Lead Developer (DevSecOps)"
        ])
        self.ai_persona_combo.setToolTip("Select a persona for the AI to adopt when generating report sections.")
        ai_layout.addRow("AI Persona:", self.ai_persona_combo)

        self.ai_instructions_edit = QTextEdit()
        self.ai_instructions_edit.setPlaceholderText("Optional: Provide custom instructions for the AI. For example, 'Focus on the financial impact of the SQL injection vulnerability'.")
        self.ai_instructions_edit.setFixedHeight(80)
        ai_layout.addRow("AI Instructions:", self.ai_instructions_edit)

        generation_layout.addWidget(ai_box)

        # --- Template Examples ---
        template_box = QGroupBox("Template Examples")
        template_layout = QHBoxLayout(template_box)
        self.load_roe_btn = QPushButton("Load ROE Template")
        self.load_roe_btn.setToolTip("Populate the ROE fields with a standard template.")
        self.load_example_report_btn = QPushButton("Load Example Report")
        self.load_example_report_btn.setToolTip("Populate the main fields with an example report for guidance.")
        template_layout.addWidget(self.load_roe_btn)
        template_layout.addWidget(self.load_example_report_btn)
        generation_layout.addWidget(template_box)

        # --- Final Report Generation ---
        final_report_box = QGroupBox("Final Export")
        final_report_layout = QFormLayout(final_report_box)

        self.report_template_combo = QComboBox()
        try:
            templates = [f for f in os.listdir("report_templates") if f.endswith('.html')]
            self.report_template_combo.addItems(templates)
        except FileNotFoundError:
            logging.error("report_templates directory not found. Report generation may fail.")
            self.report_template_combo.addItem("default_report.html")
        self.report_template_combo.addItem("professional_blue.html")

        final_report_layout.addRow("HTML Template:", self.report_template_combo)

        self.report_generate_html_btn = QPushButton(QIcon("icons/file-text.svg"), "Generate Final HTML Report")
        self.report_generate_html_btn.setToolTip("Compile all the information above into a final HTML report document.")
        self.report_generate_html_btn.clicked.connect(self._handle_generate_report) # Connect to existing handler for now

        self.report_generate_doc_btn = QToolButton()
        self.report_generate_doc_btn.setText("Generate Document")
        self.report_generate_doc_btn.setIcon(QIcon("icons/file.svg"))
        self.report_generate_doc_btn.setPopupMode(QToolButton.ToolButtonPopupMode.MenuButtonPopup)
        self.report_generate_doc_btn.setToolTip("Generate the report in various document formats (e.g., DOCX, PDF).")

        doc_menu = QMenu(self)
        doc_menu.addAction("Export as DOCX", lambda: self._handle_generate_doc_report('docx'))
        doc_menu.addAction("Export as PDF", lambda: self._handle_generate_doc_report('pdf'))
        self.report_generate_doc_btn.setMenu(doc_menu)

        self.report_generate_ai_btn_2 = QPushButton(QIcon("icons/terminal.svg"), "Generate with AI")
        self.report_generate_ai_btn_2.setToolTip("Use the AI Assistant to generate sections of the report based on findings.")
        self.report_generate_ai_btn_2.clicked.connect(self._handle_ai_report_generation)

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.report_generate_html_btn)
        button_layout.addWidget(self.report_generate_ai_btn_2)
        button_layout.addWidget(self.report_generate_doc_btn)
        final_report_layout.addRow(button_layout)

        generation_layout.addWidget(final_report_box)
        right_layout.addWidget(generation_box)

        right_panel.setLayout(right_layout)

        # --- Add panels to splitter ---
        main_splitter.addWidget(left_panel)
        main_splitter.addWidget(right_panel)
        main_splitter.setSizes([400, 600]) # Initial sizing
        main_layout.addWidget(main_splitter)

        # --- Connections for new buttons ---
        self.load_roe_btn.clicked.connect(self._load_roe_template)
        self.load_example_report_btn.clicked.connect(self._load_example_report)

        return widget

    def _load_roe_template(self):
        """Populates the ROE fields with a standardized template."""
        if self.report_objectives.toPlainText() or self.report_in_scope.toPlainText() or self.report_out_of_scope.toPlainText():
            reply = QMessageBox.question(self, "Confirm Overwrite",
                                         "This will overwrite the current ROE fields. Are you sure?",
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                         QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.No:
                return

        self.report_objectives.setPlainText(
"""- Objective 1: Identify and assess vulnerabilities in the external network perimeter.
- Objective 2: Determine if a threat actor can gain unauthorized access to internal systems.
- Objective 3: Evaluate the effectiveness of current security controls and monitoring."""
        )
        self.report_in_scope.setPlainText(
"""--- Authorized Target Space ---
- IP Range(s): [e.g., 1.2.3.0/24]
- Domains: *.example-company.com
- Applications: Main Web Application (www.example-company.com)

--- Authorized Hosts ---
- All hosts within the specified IP range unless explicitly restricted."""
        )
        self.report_out_of_scope.setPlainText(
"""--- Explicit Restrictions ---
- No intentional Denial of Service (DoS) or activities likely to cause service degradation.
- Testing is restricted to the hours of 10:00 PM to 6:00 AM UTC.
- Social engineering of company staff is not permitted.

--- Restricted Hosts ---
- payroll.example-company.com
- hr.example-company.com"""
        )
        QMessageBox.information(self, "Template Loaded", "ROE template has been loaded into the fields.")

    def _load_example_report(self):
        """Populates the main report fields with example content."""
        if self.report_summary_text.toPlainText():
            reply = QMessageBox.question(self, "Confirm Overwrite",
                                         "This will overwrite the Executive Summary. Are you sure?",
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                         QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.No:
                return

        self.report_client_name.setText("Example-Company Inc.")
        self.report_assessment_dates.setText(f"{datetime.now().year}-01-01 to {datetime.now().year}-01-07")
        self.report_summary_text.setPlainText(
"""This report details the findings of a penetration test conducted against Example-Company's external network from [Start Date] to [End Date]. The primary objective was to identify vulnerabilities that could be exploited by an external attacker to compromise the network and access sensitive data.

During the assessment, several high-risk vulnerabilities were discovered, including an unpatched web server susceptible to remote code execution and a misconfigured firewall allowing unauthorized access to an internal development server. These findings indicate a significant risk of a full network compromise.

Key recommendations include immediately patching the vulnerable web server, correcting the firewall misconfiguration, and implementing a comprehensive vulnerability management program. A detailed breakdown of all findings and their technical remediation steps is provided in the body of this report."""
        )
        QMessageBox.information(self, "Example Loaded", "Example report summary has been loaded.")

    def _handle_ai_summary_generation(self):
        """Gathers findings, sends them to the AI, and sets a callback to populate the summary."""
        if self.report_findings_tree.topLevelItemCount() == 0:
            QMessageBox.warning(self, "No Data", "There are no findings to summarize. Please run the 'Aggregate & Enrich Results' tool first.")
            return

        findings_text = ""
        for i in range(self.report_findings_tree.topLevelItemCount()):
            item = self.report_findings_tree.topLevelItem(i)
            host = item.text(0)
            service = item.text(1)
            finding = item.text(2)
            details = item.text(3)
            findings_text += f"- Host: {host}, Service: {service}, Finding: {finding}\n  Details: {details}\n\n"

        prompt = (
            "Based on the following list of penetration testing findings, please write a concise executive summary "
            "suitable for a non-technical audience. Focus on the overall risk posture, key areas of weakness, "
            "and high-level recommendations. The summary should be a few paragraphs long.\n\n"
            f"--- FINDINGS ---\n{findings_text}--- END FINDINGS ---"
        )

        def _populate_summary(generated_text):
            self.report_summary_text.setPlainText(generated_text)
            QMessageBox.information(self, "Success", "AI-generated summary has been populated.")

        self.ai_assistant_tab.set_completion_callback(_populate_summary)
        self.ai_assistant_tab.send_message(prompt)
        self.tab_widget.setCurrentWidget(self.ai_assistant_tab)
        QMessageBox.information(self, "AI Task Started", "The AI is generating the summary. You will be notified upon completion. You can watch the progress in the 'AI Assistant' tab.")

    def _create_lab_tab(self):
        """Creates the UI for the LAB / Test Chaining tab."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # Top control bar
        controls_bar = QHBoxLayout()
        self.lab_run_chain_btn = QPushButton(QIcon("icons/play-circle.svg"), " Run Test Chain")
        self.lab_save_chain_btn = QPushButton(QIcon("icons/save.svg"), " Save Chain")
        self.lab_load_chain_btn = QPushButton(QIcon("icons/folder.svg"), " Load Chain")
        self.lab_run_chain_btn.clicked.connect(self.start_lab_chain)
        self.lab_save_chain_btn.clicked.connect(self._lab_save_chain)
        self.lab_load_chain_btn.clicked.connect(self._lab_load_chain)
        controls_bar.addWidget(self.lab_run_chain_btn)
        controls_bar.addWidget(self.lab_save_chain_btn)
        controls_bar.addWidget(self.lab_load_chain_btn)
        controls_bar.addStretch()
        main_layout.addLayout(controls_bar)

        main_splitter = QSplitter(Qt.Orientation.Horizontal)

        left_panel = QGroupBox("Available Tools")
        left_layout = QVBoxLayout(left_panel)
        self.lab_tools_list = QListWidget()
        left_layout.addWidget(self.lab_tools_list)
        main_splitter.addWidget(left_panel)

        center_panel = QWidget()
        center_layout = QVBoxLayout(center_panel)
        center_layout.setContentsMargins(0,0,0,0)
        chain_box = QGroupBox("Test Chain Sequence")
        chain_layout = QVBoxLayout(chain_box)
        self.lab_chain_list = QListWidget()
        chain_layout.addWidget(self.lab_chain_list)

        chain_buttons = QHBoxLayout()
        add_btn = QPushButton("Add ->"); add_btn.clicked.connect(self._lab_add_step)
        remove_btn = QPushButton("<- Remove"); remove_btn.clicked.connect(self._lab_remove_step)
        move_up_btn = QPushButton("Move Up"); move_up_btn.clicked.connect(self._lab_move_step_up)
        move_down_btn = QPushButton("Move Down"); move_down_btn.clicked.connect(self._lab_move_step_down)
        chain_buttons.addWidget(add_btn); chain_buttons.addWidget(remove_btn)
        chain_buttons.addStretch(); chain_buttons.addWidget(move_up_btn); chain_buttons.addWidget(move_down_btn)
        chain_layout.addLayout(chain_buttons)
        center_layout.addWidget(chain_box)
        main_splitter.addWidget(center_panel)

        right_panel = QGroupBox("Step Configuration")
        self.lab_config_stack = QStackedWidget()
        right_panel.setLayout(QVBoxLayout())
        right_panel.layout().addWidget(self.lab_config_stack)

        placeholder_widget = QLabel("Select a step from the chain to configure it.")
        placeholder_widget.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lab_config_stack.addWidget(placeholder_widget)

        # Create new, independent instances of config widgets and their controls for the LAB
        self.lab_tool_configs = {}

        compatible_tools = [
            ("Nmap Scan", self._create_nmap_config_widget),
            ("Subdomain Scanner (Sublist3r)", self._create_subdomain_scanner_config_widget),
            ("Subdomain Scanner (Subfinder)", self._create_subfinder_config_widget),
            ("httpx Probe", self._create_httpx_config_widget),
            ("RustScan", self._create_rustscan_config_widget),
            ("dirsearch", self._create_dirsearch_config_widget),
            ("ffuf", self._create_ffuf_config_widget),
            ("enum4linux-ng", self._create_enum4linux_ng_config_widget),
            ("dnsrecon", self._create_dnsrecon_config_widget),
            ("fierce", self._create_fierce_config_widget),
            ("Nikto Scan", self._create_nikto_config_widget),
            ("Gobuster", self._create_gobuster_config_widget),
            ("WhatWeb", self._create_whatweb_config_widget),
            ("Masscan", self._create_masscan_config_widget),
            ("SQLMap", self._create_sqlmap_config_widget),
            ("theHarvester", self._create_theharvester_config_widget),
            ("Hashcat", self._create_hashcat_config_widget),
            ("Nuclei Scanner", self._create_nuclei_config_widget),
            ("TruffleHog Scanner", self._create_trufflehog_config_widget),
            ("John the Ripper", self._create_jtr_config_widget),
            ("Hydra", self._create_hydra_config_widget),
            ("Sherlock", self._create_sherlock_config_widget),
            ("Spiderfoot", self._create_spiderfoot_config_widget),
            ("ARP Scan (CLI)", self._create_arp_scan_cli_config_widget),
            ("Wifite Auditor", self._create_wifite_config_widget)
        ]

        for name, config_method in compatible_tools:
            self.lab_tools_list.addItem(name)
            config_widget, controls = config_method()
            self.lab_tool_configs[name] = {'widget': config_widget, 'controls': controls}
            self.lab_config_stack.addWidget(config_widget)

        self.lab_chain_list.currentItemChanged.connect(self._lab_on_chain_selection_changed)
        main_splitter.addWidget(right_panel)
        main_splitter.setSizes([200, 300, 500])
        main_layout.addWidget(main_splitter)

        return widget

    def _get_config_from_ui(self, tool_name):
        """Reads the current values from a tool's config UI and returns a dict."""
        config = {}
        if tool_name not in self.lab_tool_configs:
            return config

        controls = self.lab_tool_configs[tool_name]['controls']

        if tool_name == "Nmap Scan":
            for key, widget in controls.items():
                if isinstance(widget, QLineEdit):
                    config[key] = widget.text()
                elif isinstance(widget, QCheckBox):
                    config[key] = widget.isChecked()
                elif isinstance(widget, QComboBox):
                    config[key] = widget.currentText()
        # Add other tools here as they are implemented

        return config

    def _set_config_to_ui(self, tool_name, config):
        """Populates a tool's config UI from a config dict."""
        if tool_name not in self.lab_tool_configs:
            return

        controls = self.lab_tool_configs[tool_name]['controls']

        if tool_name == "Nmap Scan":
            for key, widget in controls.items():
                if key in config:
                    if isinstance(widget, QLineEdit):
                        widget.setText(config[key])
                    elif isinstance(widget, QCheckBox):
                        widget.setChecked(config[key])
                    elif isinstance(widget, QComboBox):
                        widget.setCurrentText(config[key])
        # Add other tools here as they are implemented

    def _lab_add_step(self):
        """Adds a selected tool from the available list to the test chain."""
        selected_item = self.lab_tools_list.currentItem()
        if not selected_item:
            QMessageBox.warning(self, "No Tool Selected", "Please select a tool from the 'Available Tools' list to add.")
            return

        tool_name = selected_item.text()

        # Get the default configuration from the current state of the tool's UI
        default_config = self._get_config_from_ui(tool_name)

        step_data = {
            'tool_name': tool_name,
            'id': str(uuid.uuid4()),
            'config': default_config
        }
        self.lab_test_chain.append(step_data)

        # The list widget item stores the unique ID to link it back to the config
        list_item = QListWidgetItem(f"Step {len(self.lab_test_chain)}: {tool_name}")
        list_item.setData(Qt.ItemDataRole.UserRole, step_data['id'])
        self.lab_chain_list.addItem(list_item)
        self.lab_chain_list.setCurrentItem(list_item)


    def _lab_remove_step(self):
        """Removes the selected step from the test chain."""
        selected_row = self.lab_chain_list.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "No Step Selected", "Please select a step from the chain to remove.")
            return

        # Remove from UI
        item = self.lab_chain_list.takeItem(selected_row)
        item_id = item.data(Qt.ItemDataRole.UserRole)

        # Remove from backend data model
        self.lab_test_chain = [step for step in self.lab_test_chain if step['id'] != item_id]

        # Renumber the remaining steps in the UI for clarity
        for i in range(self.lab_chain_list.count()):
            list_item = self.lab_chain_list.item(i)
            tool_name = list_item.text().split(": ")[1]
            list_item.setText(f"Step {i + 1}: {tool_name}")

    def _lab_move_step_up(self):
        """Moves the selected step up in the test chain."""
        current_row = self.lab_chain_list.currentRow()
        if current_row > 0:
            item = self.lab_chain_list.takeItem(current_row)
            self.lab_chain_list.insertItem(current_row - 1, item)
            self.lab_chain_list.setCurrentRow(current_row - 1)
            # Reorder the backend list as well
            self.lab_test_chain.insert(current_row - 1, self.lab_test_chain.pop(current_row))
            self._renumber_lab_steps()

    def _lab_move_step_down(self):
        """Moves the selected step down in the test chain."""
        current_row = self.lab_chain_list.currentRow()
        if 0 <= current_row < self.lab_chain_list.count() - 1:
            item = self.lab_chain_list.takeItem(current_row)
            self.lab_chain_list.insertItem(current_row + 1, item)
            self.lab_chain_list.setCurrentRow(current_row + 1)
            # Reorder the backend list as well
            self.lab_test_chain.insert(current_row + 1, self.lab_test_chain.pop(current_row))
            self._renumber_lab_steps()

    def _renumber_lab_steps(self):
        """Updates the text of the items in the lab chain list to reflect their new order."""
        for i in range(self.lab_chain_list.count()):
            item = self.lab_chain_list.item(i)
            # The tool name doesn't change, just the step number
            tool_name = self.lab_test_chain[i]['tool_name']
            item.setText(f"Step {i + 1}: {tool_name}")

    def _lab_on_chain_selection_changed(self, current, previous):
        """Shows the correct configuration widget when a step in the chain is selected."""
        # 1. Save the configuration of the previously selected item
        if previous:
            prev_id = previous.data(Qt.ItemDataRole.UserRole)
            # Find the corresponding step in the backend list
            for step in self.lab_test_chain:
                if step['id'] == prev_id:
                    tool_name = step['tool_name']
                    # Get the current UI state and save it to the step's config
                    step['config'] = self._get_config_from_ui(tool_name)
                    logging.info(f"Saved config for step {step['tool_name']} (ID: {prev_id})")
                    break

        # 2. Load the configuration of the currently selected item
        if not current:
            self.lab_config_stack.setCurrentIndex(0) # Show placeholder
            return

        current_id = current.data(Qt.ItemDataRole.UserRole)
        # Find the corresponding step in the backend list
        for step in self.lab_test_chain:
            if step['id'] == current_id:
                tool_name = step['tool_name']
                config = step['config']

                # Use the dedicated lab widgets
                if tool_name in self.lab_tool_configs:
                    # Populate the UI with the stored config
                    self._set_config_to_ui(tool_name, config)

                    # Show the correct widget from the stack
                    widget_to_show = self.lab_tool_configs[tool_name]['widget']
                    self.lab_config_stack.setCurrentWidget(widget_to_show)
                    logging.info(f"Loaded config for step {tool_name} (ID: {current_id})")
                else:
                    self.lab_config_stack.setCurrentIndex(0) # Fallback to placeholder
                return

    def _lab_save_chain(self):
        """Saves the current test chain to a JSON file."""
        if not self.lab_test_chain:
            QMessageBox.information(self, "Empty Chain", "There is nothing to save.")
            return

        file_path, _ = QFileDialog.getSaveFileName(self, "Save Test Chain", "", "Zurvan LAB Files (*.zurvan-lab)", options=QFileDialog.Option.DontUseNativeDialog)
        if not file_path:
            return

        try:
            with open(file_path, 'w') as f:
                json.dump(self.lab_test_chain, f, indent=4)
            self.status_bar.showMessage(f"Test chain saved to {file_path}", 5000)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save test chain: {e}")
            logging.error(f"Failed to save LAB chain: {e}", exc_info=True)

    def _lab_load_chain(self):
        """Loads a test chain from a JSON file."""
        if self.lab_test_chain:
            reply = QMessageBox.question(self, "Confirm Load", "This will overwrite your current test chain. Are you sure?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.No:
                return

        file_path, _ = QFileDialog.getOpenFileName(self, "Load Test Chain", "", "Zurvan LAB Files (*.zurvan-lab);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if not file_path:
            return

        try:
            with open(file_path, 'r') as f:
                loaded_chain = json.load(f)

            # Basic validation
            if not isinstance(loaded_chain, list) or not all('tool_name' in d and 'config' in d for d in loaded_chain):
                raise ValueError("Invalid file format.")

            self.lab_test_chain = loaded_chain
            self.lab_chain_list.clear()

            # Repopulate the UI list
            for i, step in enumerate(self.lab_test_chain):
                # Ensure each step has a unique ID if loading older formats
                if 'id' not in step:
                    step['id'] = str(uuid.uuid4())

                list_item = QListWidgetItem(f"Step {i + 1}: {step['tool_name']}")
                list_item.setData(Qt.ItemDataRole.UserRole, step['id'])
                self.lab_chain_list.addItem(list_item)

            self.status_bar.showMessage(f"Test chain loaded from {file_path}", 5000)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load test chain: {e}")
            logging.error(f"Failed to load LAB chain: {e}", exc_info=True)

    def start_lab_chain(self):
        """Starts the LAB chain execution worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return
        if not self.lab_test_chain:
            QMessageBox.information(self, "Empty Chain", "There are no steps in the test chain to run.")
            return

        self.is_tool_running = True
        self.lab_run_chain_btn.setEnabled(False)
        # In the future, a cancel button will be added
        # self.lab_cancel_chain_btn.setEnabled(True)
        self.tool_stop_event.clear()

        # Deepcopy the chain to avoid race conditions if the user edits it while running
        chain_to_run = copy.deepcopy(self.lab_test_chain)

        self.worker = WorkerThread(self._lab_chain_thread, args=(chain_to_run,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _execute_blocking_command_for_lab(self, command, tool_name):
        """
        Executes a command in a blocking manner for the LAB thread,
        streaming its output to the queue. Returns (success, full_output).
        """
        q = self.tool_results_queue
        full_output = []

        logging.info(f"[LAB] Executing: {' '.join(command)}")
        q.put(('lab_output', f"\n\n$ {' '.join(command)}\n"))

        try:
            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                                     text=True, bufsize=1, encoding='utf-8', errors='replace')

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('lab_output', "\n--- Step Canceled By User ---\n"))
                    return False, "".join(full_output)

                q.put(('lab_output', line))
                full_output.append(line)

            process.stdout.close()
            retcode = process.wait()

            if retcode != 0:
                logging.warning(f"[LAB] {tool_name} exited with non-zero code: {retcode}")

            return True, "".join(full_output)

        except FileNotFoundError:
            q.put(('error', f'{tool_name} Error', f"Command '{command[0]}' not found."))
            return False, ""
        except Exception as e:
            q.put(('error', f'{tool_name} Error', str(e)))
            return False, ""

    def _lab_chain_thread(self, chain):
        """Worker thread that executes each step of the LAB chain sequentially."""
        q = self.tool_results_queue
        logging.info(f"LAB chain execution started with {len(chain)} steps.")
        self.lab_context = {} # This dictionary will hold results passed between steps

        for i, step in enumerate(chain):
            if self.tool_stop_event.is_set():
                logging.info("LAB chain execution cancelled by user.")
                q.put(('lab_status', "Chain cancelled by user."))
                break

            tool_name = step['tool_name']
            config = step['config']
            q.put(('lab_status', f"Executing Step {i+1}/{len(chain)}: {tool_name}"))

            # --- Nmap Scan Execution ---
            if tool_name == "Nmap Scan":
                target = config.get('target_edit', '')
                if not target:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Nmap): Target is missing in configuration."))
                    continue

                command, _, error = self._build_nmap_command(config, target)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Nmap): {error}"))
                    continue

                # Nmap requires a temp file for XML output
                temp_xml_path = None
                try:
                    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".xml", encoding='utf-8') as tmp_xml:
                        temp_xml_path = tmp_xml.name
                    command.extend(["-oX", temp_xml_path])
                except Exception as e:
                    q.put(('error', 'LAB File Error', f"Step {i+1} (Nmap): Could not create temp file: {e}"))
                    continue

                # Execute the command and handle output
                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                if success:
                    try:
                        with open(temp_xml_path, 'r', encoding='utf-8') as f:
                            xml_content = f.read()
                        # Store result in context for next steps
                        self.lab_context['last_nmap_xml'] = xml_content
                        q.put(('lab_status', f"Nmap scan complete. XML results stored in context."))
                    except Exception as e:
                        q.put(('error', 'LAB File Error', f"Step {i+1} (Nmap): Could not read temp file: {e}"))
                    finally:
                        if temp_xml_path and os.path.exists(temp_xml_path):
                            os.remove(temp_xml_path)
                else:
                     q.put(('lab_status', f"Step {i+1} (Nmap) failed or was cancelled."))

            elif tool_name == "Subdomain Scanner (Sublist3r)":
                command, target, error = self._build_sublist3r_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Sublist3r): {error}"))
                    continue

                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                if success and output:
                    # Parse and store results in context
                    subdomains = re.findall(r'^[a-zA-Z0-9.\-]+\.' + re.escape(target), output, re.MULTILINE)
                    unique_subdomains = sorted(list(set(subdomains)))
                    if 'subdomains' not in self.lab_context:
                        self.lab_context['subdomains'] = []
                    self.lab_context['subdomains'].extend(unique_subdomains)
                    self.lab_context['subdomains'] = sorted(list(set(self.lab_context['subdomains'])))
                    q.put(('lab_status', f"Sublist3r found {len(unique_subdomains)} new subdomains. Total unique: {len(self.lab_context['subdomains'])}."))
                else:
                    q.put(('lab_status', f"Step {i+1} (Sublist3r) returned no output or failed."))
            elif tool_name == "Subdomain Scanner (Subfinder)":
                command, target, error = self._build_subfinder_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Subfinder): {error}"))
                    continue

                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                if success and output:
                    subdomains = [line for line in output.splitlines() if target in line and ' ' not in line and not line.startswith('$')]
                    unique_subdomains = sorted(list(set(subdomains)))

                    if 'subdomains' not in self.lab_context:
                        self.lab_context['subdomains'] = []
                    self.lab_context['subdomains'].extend(unique_subdomains)
                    self.lab_context['subdomains'] = sorted(list(set(self.lab_context['subdomains'])))
                    q.put(('lab_status', f"Subfinder found {len(unique_subdomains)} new subdomains. Total unique: {len(self.lab_context['subdomains'])}."))
                else:
                    q.put(('lab_status', f"Step {i+1} (Subfinder) returned no output or failed."))
            elif tool_name == "httpx Probe":
                # httpx can take a file as input. We can create one from the context.
                target_list_from_config = config.get('target_list_edit', '').strip()

                # Prioritize a file from config, otherwise use context
                if not target_list_from_config and 'subdomains' in self.lab_context:
                    try:
                        with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".txt", encoding='utf-8') as tmp_file:
                            tmp_file.write("\n".join(self.lab_context['subdomains']))
                            target_list_from_config = tmp_file.name
                        q.put(('lab_status', f"Created temporary target list for httpx from {len(self.lab_context['subdomains'])} subdomains."))
                        # Update the config for the build command
                        config['target_list_edit'] = target_list_from_config
                    except Exception as e:
                        q.put(('error', 'LAB File Error', f"Step {i+1} (httpx): Could not create temp file: {e}"))
                        continue

                command, target, error = self._build_httpx_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (httpx): {error}"))
                    continue

                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                # Clean up temp file if we created one
                if target_list_from_config and target_list_from_config.startswith(tempfile.gettempdir()):
                    os.remove(target_list_from_config)

                if success and output:
                    # In a real scenario, we might parse the JSON and add live hosts to a new context variable
                    q.put(('lab_status', f"httpx probe finished."))
                else:
                    q.put(('lab_status', f"Step {i+1} (httpx) returned no output or failed."))

            elif tool_name == "RustScan":
                command, target, error = self._build_rustscan_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (RustScan): {error}"))
                    continue

                # RustScan may require sudo, which must be handled by running the entire app with sudo.
                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                if success and output:
                    # TODO: Parse output for open ports and add to context
                    q.put(('lab_status', f"RustScan finished against {target}."))
                else:
                    q.put(('lab_status', f"Step {i+1} (RustScan) returned no output or failed."))
            elif tool_name == "dirsearch":
                command, target, error = self._build_dirsearch_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (dirsearch): {error}"))
                    continue

                temp_json_path = None
                try:
                    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                        temp_json_path = tmp_json.name
                    command = [arg for arg in command if not arg.startswith('--json-report')]
                    command.extend(["--json-report", temp_json_path])
                except Exception as e:
                    q.put(('error', 'LAB File Error', f"Step {i+1} (dirsearch): Could not create temp file: {e}"))
                    continue

                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                if success and temp_json_path and os.path.exists(temp_json_path):
                    try:
                        with open(temp_json_path, 'r', encoding='utf-8') as f:
                            json_data = f.read()
                        if json_data:
                            # TODO: Parse found paths and add to context
                            q.put(('lab_status', f"dirsearch finished against {target}."))
                    except Exception as e:
                         q.put(('error', 'LAB File Error', f"Step {i+1} (dirsearch): Could not read temp file: {e}"))
                    finally:
                        os.remove(temp_json_path)
                else:
                    q.put(('lab_status', f"Step {i+1} (dirsearch) returned no output or failed."))
            elif tool_name == "ffuf":
                command, target, error = self._build_ffuf_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (ffuf): {error}"))
                    continue

                temp_json_path = None
                try:
                    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                        temp_json_path = tmp_json.name
                    command.extend(["-o", temp_json_path, "-of", "json"])
                except Exception as e:
                    q.put(('error', 'LAB File Error', f"Step {i+1} (ffuf): Could not create temp file: {e}"))
                    continue

                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                if success and temp_json_path and os.path.exists(temp_json_path):
                    try:
                        with open(temp_json_path, 'r', encoding='utf-8') as f:
                            json_data = f.read()
                        if json_data:
                            # TODO: Parse found paths and add to context
                            q.put(('lab_status', f"ffuf finished against {target}."))
                    except Exception as e:
                         q.put(('error', 'LAB File Error', f"Step {i+1} (ffuf): Could not read temp file: {e}"))
                    finally:
                        os.remove(temp_json_path)
                else:
                    q.put(('lab_status', f"Step {i+1} (ffuf) returned no output or failed."))

            elif tool_name == "enum4linux-ng":
                command, target, error = self._build_enum4linux_ng_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (enum4linux-ng): {error}"))
                    continue

                temp_json_path = None
                try:
                    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                        temp_json_path = tmp_json.name
                    command.insert(1, "-oJ")
                    command.insert(2, temp_json_path)
                except Exception as e:
                    q.put(('error', 'LAB File Error', f"Step {i+1} (enum4linux-ng): Could not create temp file: {e}"))
                    continue

                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                if success and temp_json_path and os.path.exists(temp_json_path):
                    try:
                        with open(temp_json_path, 'r', encoding='utf-8') as f:
                            json_data = f.read()
                        if json_data:
                            # TODO: Parse found info and add to context
                            q.put(('lab_status', f"enum4linux-ng finished against {target}."))
                    except Exception as e:
                         q.put(('error', 'LAB File Error', f"Step {i+1} (enum4linux-ng): Could not read temp file: {e}"))
                    finally:
                        os.remove(temp_json_path)
                else:
                    q.put(('lab_status', f"Step {i+1} (enum4linux-ng) returned no output or failed."))

            elif tool_name == "dnsrecon":
                command, target, error = self._build_dnsrecon_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (dnsrecon): {error}"))
                    continue

                temp_json_path = None
                try:
                    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                        temp_json_path = tmp_json.name
                    command = [arg for arg in command if not arg.startswith('--json')]
                    command.extend(["--json", temp_json_path])
                except Exception as e:
                    q.put(('error', 'LAB File Error', f"Step {i+1} (dnsrecon): Could not create temp file: {e}"))
                    continue

                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                if success and temp_json_path and os.path.exists(temp_json_path):
                    try:
                        with open(temp_json_path, 'r', encoding='utf-8') as f:
                            json_data = f.read()
                        if json_data:
                            # TODO: Parse found records and add to context
                            q.put(('lab_status', f"dnsrecon finished against {target}."))
                    except Exception as e:
                         q.put(('error', 'LAB File Error', f"Step {i+1} (dnsrecon): Could not read temp file: {e}"))
                    finally:
                        os.remove(temp_json_path)
                else:
                    q.put(('lab_status', f"Step {i+1} (dnsrecon) returned no output or failed."))

            elif tool_name == "fierce":
                command, target, error = self._build_fierce_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (fierce): {error}"))
                    continue

                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                if success:
                    # TODO: Parse output and add to context
                    q.put(('lab_status', f"fierce finished against {target}."))
                else:
                    q.put(('lab_status', f"Step {i+1} (fierce) returned no output or failed."))

            elif tool_name == "Nikto Scan":
                command, target, error = self._build_nikto_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Nikto): {error}"))
                    continue

                success, output = self._execute_blocking_command_for_lab(command, tool_name)

                if success:
                    # TODO: Parse output and add to context
                    q.put(('lab_status', f"Nikto finished against {target}."))
                else:
                    q.put(('lab_status', f"Step {i+1} (Nikto) returned no output or failed."))

            elif tool_name == "Gobuster":
                command, target, error = self._build_gobuster_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Gobuster): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"Gobuster finished against {target}."))

            elif tool_name == "WhatWeb":
                command, target, error = self._build_whatweb_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (WhatWeb): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"WhatWeb finished against {target}."))

            elif tool_name == "Masscan":
                command, target, error = self._build_masscan_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Masscan): {error}"))
                    continue
                # Masscan requires sudo
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"Masscan finished against {target}."))

            elif tool_name == "SQLMap":
                command, target, error = self._build_sqlmap_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (SQLMap): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"SQLMap finished against {target}."))

            elif tool_name == "Hashcat":
                command, target, error = self._build_hashcat_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Hashcat): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"Hashcat finished against {target}."))

            elif tool_name == "Nuclei Scanner":
                command, target, error = self._build_nuclei_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Nuclei): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"Nuclei finished against {target}."))

            elif tool_name == "TruffleHog Scanner":
                command, target, error = self._build_trufflehog_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (TruffleHog): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"TruffleHog finished against {target}."))

            elif tool_name == "John the Ripper":
                # JTR has two modes, we assume crack mode for the lab
                command, target, error = self._build_jtr_command(config, crack_mode=True)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (JTR): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"John the Ripper finished against {target}."))

            elif tool_name == "Hydra":
                command, target, error = self._build_hydra_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Hydra): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"Hydra finished against {target}."))

            elif tool_name == "Sherlock":
                temp_dir = None
                try:
                    temp_dir = tempfile.mkdtemp()
                    command, target, error = self._build_sherlock_command(config, csv_temp_dir=temp_dir)
                    if error:
                        q.put(('error', 'LAB Error', f"Step {i+1} (Sherlock): {error}"))
                        continue
                    success, output = self._execute_blocking_command_for_lab(command, tool_name)
                    q.put(('lab_status', f"Sherlock finished against {target}."))
                finally:
                    if temp_dir and os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)

            elif tool_name == "Spiderfoot":
                command, target, error = self._build_spiderfoot_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Spiderfoot): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"Spiderfoot finished against {target}."))

            elif tool_name == "ARP Scan (CLI)":
                command, target, error = self._build_arp_scan_cli_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (arp-scan): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"arp-scan finished against {target}."))

            elif tool_name == "Wifite Auditor":
                command, target, error = self._build_wifite_command(config)
                if error:
                    q.put(('error', 'LAB Error', f"Step {i+1} (Wifite): {error}"))
                    continue
                success, output = self._execute_blocking_command_for_lab(command, tool_name)
                q.put(('lab_status', f"Wifite finished against {target}."))
            else:
                q.put(('lab_output', f"\n--- Skipping unimplemented tool: {tool_name} ---\n"))
                time.sleep(1)


        q.put(('tool_finished', 'lab_chain'))
        logging.info("LAB chain execution finished.")


import database

def main():
    """Main function to launch the Zurvan application."""
    app = QApplication.instance()
    if app is None:
        app = QApplication(sys.argv)

    try:
        database.initialize_database()
        if 'scapy' not in sys.modules:
            raise ImportError("Scapy is not installed or could not be imported.")

        login_dialog = LoginDialog()

        # Define the custom stylesheet additions for 2026 UI Trends
        extra_qss = {
            'QGroupBox': {
                'border': '1px solid #444;',
                'border-radius': '16px',
                'margin-top': '10px',
                'padding-top': '15px',
            },
            'QGroupBox::title': {
                'subcontrol-origin': 'margin',
                'subcontrol-position': 'top left',
                'padding': '0 10px',
            },
            'QTabWidget::pane': {
                'border-top': '1px solid #444;',
                'margin-top': '-1px',
                'border-radius': '12px',
            },
            'QFrame': {
                'border-radius': '12px',
            },
            'QPushButton': {
                'border-radius': '12px',
                'padding': '8px 16px',
            },
            'QLineEdit': {
                'border-radius': '12px',
                'padding': '6px 12px',
            },
            'QComboBox': {
                'border-radius': '12px',
                'padding': '6px 12px',
            },
            'QTextEdit': {
                'border-radius': '12px',
                'padding': '8px',
            },
            'QPlainTextEdit': {
                'border-radius': '12px',
                'padding': '8px',
            },
            'QListWidget': {
                'border-radius': '12px',
            },
            'QTreeWidget': {
                'border-radius': '12px',
            },
            'QPushButton:hover': {
                'background-color': '{{primaryColor}}',
                'color': '{{secondaryDarkColor}}',
                'border': '2px solid {{primaryLightColor}}',
            },
            'QPushButton:pressed': {
                'background-color': '{{secondaryColor}}',
                'color': '{{primaryColor}}',
                'border': '2px solid {{primaryColor}}',
            }
        }

        # Apply the default theme before showing the login dialog
        apply_stylesheet(app, theme=login_dialog.selected_theme, extra=extra_qss)

        if login_dialog.exec() != QDialog.DialogCode.Accepted:
            sys.exit(0)

        # Re-apply the theme in case the user changed it in the dialog.
        # This ensures the main window gets the final selected theme.
        apply_stylesheet(app, theme=login_dialog.selected_theme, extra=extra_qss)

        window = Zurvan()
        window.current_user = login_dialog.current_user
        # Set window title with username
        if window.current_user and 'username' in window.current_user:
            window.setWindowTitle(f"Welcome, {window.current_user['username']} - Zurvan - Comprehensive AI-Powered Security Platform")
        window._update_menu_bar() # Populate the menu now that we have a user
        window._set_user_avatar() # Set avatar after user is loaded
        window._load_and_apply_user_settings() # Load user-specific settings
        window._setup_app_lock_monitor() # Start the activity monitor
        window.show()

        # For screenshotting - switch to the new tab
        if '--screenshot' in sys.argv:
            window.is_test_mode = True
            QTimer.singleShot(1000, lambda: window.take_screenshot_and_exit('ssh_manager.png'))
        else:
            # Log application start after everything is initialized and shown
            if window.current_user:
                database.log_activity(
                    user_id=window.current_user['id'],
                    category='System',
                    action='Application Started',
                    target=f"v3.0 on {platform.system()}",
                    details=f"User '{window.current_user['username']}' started the application."
                )

        sys.exit(app.exec())

    except ImportError as e:
        QMessageBox.critical(None, "Fatal Error", str(e))
        sys.exit(1)
    except Exception as e:
        logging.critical(f"An unhandled exception occurred: {e}", exc_info=True)
        QMessageBox.critical(None, "Unhandled Exception", f"An unexpected error occurred:\n\n{e}")
        sys.exit(1)

if __name__ == "__main__":
    sys.argv.append("--no-sandbox")
    main()
