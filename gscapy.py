import sys
import logging
import time
from threading import Event, Lock
import queue
import copy
import socket
import random
import os
import csv
import platform
import psutil
import ipaddress
from PyQt6.QtCore import PYQT_VERSION_STR
import subprocess
import numpy as np
import json
import urllib.request
import tempfile
import webbrowser
import shutil
import signal
import uuid
import sqlite3
import gzip
from datetime import datetime

try:
    from lxml import etree
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False
    etree = None
    logging.warning("Optional XML reporting dependency not found. Please run 'pip install lxml'")

import re
from qt_material import apply_stylesheet, list_themes
from PyQt6.QtGui import QActionGroup, QPixmap, QImage, QPalette

def create_themed_icon(icon_path, color_str):
    """Loads an SVG, intelligently replaces its color, and returns a QIcon."""
    try:
        with open(icon_path, 'r', encoding='utf-8') as f:
            svg_data = f.read()

        # First, try to replace a stroke color in a style block (for paper-airplane.svg)
        themed_svg_data, count = re.subn(r'stroke:#[0-9a-fA-F]{6}', f'stroke:{color_str}', svg_data)

        # If no stroke was found in a style, fall back to injecting a fill attribute (for gear.svg)
        if count == 0 and '<svg' in themed_svg_data:
            themed_svg_data = themed_svg_data.replace('<svg', f'<svg fill="{color_str}"')

        image = QImage.fromData(themed_svg_data.encode('utf-8'))
        pixmap = QPixmap.fromImage(image)
        return QIcon(pixmap)
    except Exception as e:
        logging.warning(f"Could not create themed icon for {icon_path}: {e}")
        return QIcon(icon_path) # Fallback to original icon

def get_vendor(mac_address):
    """Retrieves the vendor for a given MAC address from an online API."""
    if not mac_address or mac_address == "N/A":
        return "N/A"
    try:
        # Use a timeout to prevent the application from hanging on network issues
        with urllib.request.urlopen(f"https://api.macvendors.com/{mac_address}", timeout=3) as url:
            data = url.read().decode()
            return data
    except Exception as e:
        logging.warning(f"Could not retrieve vendor for MAC {mac_address}: {e}")
        return "Unknown Vendor"

def _get_random_ip():
    """Generates a random, non-private IP address."""
    while True:
        ip = ".".join(str(random.randint(1, 223)) for _ in range(4))
        if not (ip.startswith('10.') or ip.startswith('192.168.') or (ip.startswith('172.') and 16 <= int(ip.split('.')[1]) <= 31)):
             return ip

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QStatusBar, QMenuBar, QTabWidget, QWidget,
    QVBoxLayout, QLabel, QDockWidget, QPlainTextEdit, QPushButton, QHBoxLayout,
    QTreeWidget, QTreeWidgetItem, QSplitter, QFileDialog, QMessageBox, QComboBox,
    QListWidget, QListWidgetItem, QScrollArea, QLineEdit, QCheckBox, QFrame, QMenu, QTextEdit, QGroupBox,
    QProgressBar, QTextBrowser, QRadioButton, QButtonGroup, QFormLayout, QGridLayout, QDialog,
    QHeaderView, QInputDialog, QGraphicsOpacityEffect, QStackedWidget, QToolButton
)
from ai_tab import AIAssistantTab, AISettingsDialog, AIGuideDialog
from login import LoginDialog
from admin_panel import AdminPanelDialog
from user_profile import UserProfileDialog
from PyQt6.QtCore import QObject, pyqtSignal, Qt, QThread, QTimer, QPropertyAnimation, QEasingCurve, QParallelAnimationGroup, QSequentialAnimationGroup, QPoint, QSize
from PyQt6.QtGui import QAction, QIcon, QFont, QTextCursor, QActionGroup


def sniffer_process_target(queue, iface, bpf_filter):
    """
    This function runs in a separate process. It sniffs packets and puts them
    into a multiprocessing.Queue. This completely isolates the blocking
    sniff() call from the main GUI application.
    """
    try:
        # The packet handler now simply puts the raw packet into the queue
        def packet_handler(packet):
            queue.put(bytes(packet))

        # We don't need a stop_filter anymore, as the process will be terminated directly.
        sniff(prn=packet_handler, iface=iface, filter=bpf_filter, store=False)
    except Exception as e:
        logging.error(f"Critical error in sniffer process: {e}", exc_info=True)


class KrackScanThread(QThread):
    vulnerability_detected = pyqtSignal(str, str) # bssid, client_mac

    def __init__(self, iface, parent=None):
        super().__init__(parent)
        self.iface = iface
        self.stop_event = Event()
        self.eapol_db = {} # { (bssid, client_mac): { replay_counter: count } }

    def _packet_handler(self, pkt):
        if not pkt.haslayer(EAPOL) or not pkt.haslayer(Dot11):
            return

        # Check if frame is going from AP to client (To DS=0, From DS=1)
        if pkt.FCfield & 0x3 != 1:
            return

        try:
            # Key Information field is a good indicator for Message 3
            key_info = pkt[EAPOL].key_info
            # Message 3: Pairwise, Install, Ack, MIC
            # Install = bit 6 (0x40), Ack = bit 7 (0x80), MIC = bit 8 (0x100)
            is_msg3 = (key_info & 0x1c0) == 0x1c0

            if is_msg3:
                bssid = pkt.addr2
                client_mac = pkt.addr1
                replay_counter = pkt[EAPOL].replay_counter

                key = (bssid, client_mac)

                if key not in self.eapol_db:
                    self.eapol_db[key] = {}

                if replay_counter not in self.eapol_db[key]:
                    self.eapol_db[key][replay_counter] = 1
                else:
                    # If we see the same replay counter again, it's a retransmission
                    self.eapol_db[key][replay_counter] += 1
                    if self.eapol_db[key][replay_counter] == 2:
                        logging.info(f"KRACK vulnerability detected! BSSID: {bssid}, Client: {client_mac}")
                        self.vulnerability_detected.emit(bssid, client_mac)
                        # Reset counter to avoid flooding with signals for the same retransmission
                        self.eapol_db[key][replay_counter] = 0


        except (IndexError, AttributeError) as e:
            logging.debug(f"Error processing EAPOL packet for KRACK scan: {e}")

    def run(self):
        logging.info(f"KRACK scanner started on interface {self.iface}")
        while not self.stop_event.is_set():
            try:
                sniff(iface=self.iface, prn=self._packet_handler, filter="ether proto 0x888e", timeout=1)
            except Exception as e:
                logging.error(f"Error in KRACK sniffer loop: {e}", exc_info=True)
                time.sleep(1)

    def stop(self):
        self.stop_event.set()


class AircrackThread(QThread):
    """A thread to run the aircrack-ng process and emit its output."""
    output_received = pyqtSignal(str)
    finished_signal = pyqtSignal(int)

    def __init__(self, pcap_file, wordlist, parent=None, threads=1):
        super().__init__(parent)
        self.pcap_file = pcap_file
        self.wordlist = wordlist
        self.threads = threads
        self.process = None

    def run(self):
        command = ["aircrack-ng", "-w", self.wordlist, "-p", str(self.threads), self.pcap_file]
        try:
            self.process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1)
            for line in iter(self.process.stdout.readline, ''):
                self.output_received.emit(line.strip())
            self.process.stdout.close()
            return_code = self.process.wait()
            self.finished_signal.emit(return_code)
        except FileNotFoundError:
            self.output_received.emit("ERROR: 'aircrack-ng' command not found. Please ensure it is installed and in your system's PATH.")
            self.finished_signal.emit(-1)
        except Exception as e:
            self.output_received.emit(f"An unexpected error occurred: {e}")
            self.finished_signal.emit(-1)

    def stop(self):
        if self.process and self.process.poll() is None:
            self.process.terminate()
            self.process.wait()
            logging.info("Aircrack-ng process terminated.")

try:
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    import docx
except ImportError:
    logging.warning("Optional PDF/DOCX export dependencies not found. Please run 'pip install reportlab python-docx'")

try:
    import pyqtgraph as pg
    PYQTGRAPH_AVAILABLE = True
except ImportError:
    PYQTGRAPH_AVAILABLE = False
    pg = None # Define pg as None to prevent other errors if it's referenced
    logging.warning("Optional graphing dependency not found. Please run 'pip install pyqtgraph'")


try:
    import GPUtil
except ImportError:
    GPUtil = None
    logging.warning("Optional GPU monitoring dependency not found. Please run 'pip install gputil'")

# --- Scapy Imports ---
try:
    from scapy.all import *
    from scapy.utils import hexdump
    from scapy.layers.dns import DNS, DNSQR
    from scapy.layers.dot11 import Dot11, Dot11Beacon, Dot11Elt, Dot11Deauth, RadioTap
    conf.verb = 0
except ImportError:
    logging.critical("Scapy is not installed.")

# --- Constants ---
AVAILABLE_PROTOCOLS = {"Ethernet": Ether, "ARP": ARP, "IP": IP, "IPv6": IPv6, "TCP": TCP, "UDP": UDP, "ICMP": ICMP, "DNS": DNS, "Raw": Raw}
PACKET_TEMPLATES = {
    "ICMP Ping (google.com)": [IP(dst="8.8.8.8"), ICMP()],
    "DNS Query (google.com)": [IP(dst="8.8.8.8"), UDP(dport=53), DNS(rd=1, qd=DNSQR(qname="google.com"))],
    "TCP SYN (localhost:80)": [IP(dst="127.0.0.1"), TCP(dport=80, flags="S")],
    "ARP Request (who-has 192.168.1.1)": [Ether(dst="ff:ff:ff:ff:ff:ff"), ARP(pdst="192.168.1.1")],
    "NTP Query (pool.ntp.org)": [IP(dst="pool.ntp.org"), UDP(sport=123, dport=123), NTP()],
    "SNMP GetRequest (public)": [IP(dst="127.0.0.1"), UDP(), SNMP(community="public", PDU=SNMPget(varbindlist=[SNMPvarbind(oid=ASN1_OID('1.3.6.1.2.1.1.1.0'))]))]
}
FIREWALL_PROBES = {
    "Standard SYN Scan (Top Ports)": [(lambda t: IP(dst=t)/TCP(dport=p, flags="S"), f"TCP SYN to port {p}") for p in [21, 22, 25, 53, 80, 110, 143, 443, 445, 3389, 8080]],
    "Stealthy Scans (FIN, Xmas, Null)": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="F"), f"FIN Scan to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="FPU"), f"Xmas Scan to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags=""), f"Null Scan to port {p}") for p in [80, 443]
    ],
    "ACK Scan (Firewall Detection)": [(lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="A"), f"ACK Scan to port {p}") for p in [22, 80, 443]],
    "Source Port Evasion (DNS)": [(lambda t, p=p: IP(dst=t)/TCP(sport=53, dport=p, flags="S"), f"SYN from port 53 to {p}") for p in [80, 443, 8080]],
    "Fragmented SYN Scan": [(lambda t, p=p: fragment(IP(dst=t)/TCP(dport=p, flags="S")), f"Fragmented SYN to port {p}") for p in [80, 443]],
    "TCP Options Probes (WScale, TS)": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="S", options=[('WScale', 10), ('Timestamp', (12345, 0))]), f"SYN+WScale+TS to port {p}") for p in [80, 443]
    ],
    "ECN Flag Probes": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="SE"), f"SYN+ECE to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="SC"), f"SYN+CWR to port {p}") for p in [80, 443]
    ],
    "HTTP Payload Probe": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="PA")/Raw(load="GET / HTTP/1.0\r\n\r\n"), f"HTTP GET probe to port {p}") for p in [80, 8080, 443]
    ],
    "Common UDP Probes": [(lambda t, p=p: IP(dst=t)/UDP(dport=p), f"UDP Probe to port {p}") for p in [53, 123, 161]],
    "ICMP Probes (Advanced)": [
        (lambda t: IP(dst=t)/ICMP(type=ty), f"ICMP Echo Request (Type 8)") for ty in [8]
    ] + [
        (lambda t: IP(dst=t)/ICMP(type=ty), f"ICMP Timestamp Request (Type 13)") for ty in [13]
    ] + [
        (lambda t: IP(dst=t)/ICMP(type=ty), f"ICMP Address Mask Request (Type 17)") for ty in [17]
    ]
}
SCAN_TYPES = ["TCP SYN Scan", "TCP FIN Scan", "TCP Xmas Scan", "TCP Null Scan", "TCP ACK Scan", "UDP Scan"]
COMMON_FILTERS = [
    "", "tcp", "udp", "arp", "icmp",
    "port 80", "port 443", "udp port 53", "tcp port 22",
    "host 8.8.8.8", "net 192.168.1.0/24", "vlan"
]

COMMUNITY_TOOLS = {
    "Interpreters and REPLs": [
        ("scapy-console", "https://github.com/gpotter2/scapy-console", "A Scapy console with many other tools and features."),
        ("Scapy REPL", "https://github.com/GabrielCama/scapy-repl", "An interactive Scapy REPL with customized commands.")
    ],
    "Networking": [
        ("bettercap", "https://github.com/bettercap/bettercap", "A powerful, flexible and portable tool for network attacks and monitoring."),
        ("Routersploit", "https://github.com/threat9/routersploit", "An open-source exploitation framework dedicated to embedded devices."),
        ("Batfish", "https://www.batfish.org/", "A network configuration analysis tool for validating and verifying network designs.")
    ],
    "Network Scanners & Analyzers": [
        ("Wireshark", "https://www.wireshark.org/", "The world's foremost and widely-used network protocol analyzer."),
        ("Nmap", "https://nmap.org/", "The Network Mapper - a free and open source utility for network discovery and security auditing."),
        ("Zeek", "https://zeek.org/", "A powerful network analysis framework that is much different from a typical IDS."),
        ("BruteShark", "https://github.com/odedshimon/BruteShark", "An open-source, cross-platform network forensic analysis tool (NFAT).")
    ],
    "Wireless": [
        ("Kismet", "https://www.kismetwireless.net/", "A wireless network detector, sniffer, and intrusion detection system."),
        ("Airgeddon", "https://github.com/v1s1t0r1sh3r3/airgeddon", "A multi-use bash script for Linux systems to audit wireless networks."),
        ("wifiphisher", "https://github.com/wifiphisher/wifisher", "A rogue Access Point framework for conducting red team engagements or Wi-Fi security testing."),
        ("Wifite2", "https://github.com/derv82/wifite2", "A complete rewrite of the popular wireless network auditing tool, wifite.")
    ],
    "Password Cracking": [
        ("John the Ripper", "https://www.openwall.com/john/", "A fast password cracker, available for many operating systems."),
        ("Hashcat", "https://hashcat.net/hashcat/", "The world's fastest and most advanced password recovery utility."),
        ("hcxtools", "https://github.com/ZerBea/hcxtools", "Tools to convert Wi-Fi captures into hash formats for Hashcat or John.")
    ],
    "Web & API Security": [
        ("reNgine", "https://github.com/yogeshojha/rengine", "An automated reconnaissance framework for web applications."),
        ("Astra", "https://github.com/flipkart-incubator/Astra", "Automated Security Testing For REST APIs.")
    ],
    "Industrial Control Systems (ICS)": [
        ("Scapy-cip-enip", "https://github.com/scapy-cip/scapy-cip-enip", "An EtherNet/IP and CIP implementation for Scapy."),
        ("Scapy-dnp3", "https://github.com/scapy-dnp3/scapy-dnp3", "A DNP3 implementation for Scapy."),
        ("Scapy-modbus", "https://github.com/scapy-modbus/scapy-modbus", "A Modbus implementation for Scapy.")
    ]
}

class CrunchDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Crunch Wordlist Generator")
        layout = QVBoxLayout(self)
        form_layout = QFormLayout()

        self.min_len = QLineEdit("8")
        self.max_len = QLineEdit("8")
        self.charset = QLineEdit("abcdefghijklmnopqrstuvwxyz0123456789")
        self.output_file = QLineEdit()
        self.output_file.setReadOnly(True)

        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(self.browse_output)

        form_layout.addRow("Min Length:", self.min_len)
        form_layout.addRow("Max Length:", self.max_len)
        form_layout.addRow("Character Set:", self.charset)

        output_layout = QHBoxLayout()
        output_layout.addWidget(self.output_file)
        output_layout.addWidget(browse_btn)
        form_layout.addRow("Output File:", output_layout)

        layout.addLayout(form_layout)

        self.generate_button = QPushButton("Generate")
        self.generate_button.clicked.connect(self.accept)
        layout.addWidget(self.generate_button)

    def browse_output(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Wordlist", "", "Text Files (*.txt)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.output_file.setText(file_path)

    def get_values(self):
        return {
            "min": self.min_len.text(),
            "max": self.max_len.text(),
            "charset": self.charset.text(),
            "outfile": self.output_file.text()
        }

# --- Logging and Threads ---
class QtLogHandler(logging.Handler, QObject):
    """A custom logging handler that emits a Qt signal for each log record."""
    log_updated = pyqtSignal(str)
    def __init__(self): super().__init__(); QObject.__init__(self)
    def emit(self, record): self.log_updated.emit(self.format(record))

class SnifferThread(QThread):
    """
    This QThread does not sniff itself. Instead, it manages a separate
    multiprocessing.Process for sniffing to prevent the GUI from freezing.
    It communicates with the main thread exclusively via thread-safe Qt signals
    that carry raw bytes, not complex objects.
    """
    packet_bytes_received = pyqtSignal(bytes)

    def __init__(self, iface, bpf_filter, parent=None):
        super().__init__(parent)
        self.iface = iface
        self.bpf_filter = bpf_filter
        self.process = None
        self.queue = None
        self.stop_event = Event()

    def run(self):
        from multiprocessing import Process, Queue
        self.queue = Queue()
        self.process = Process(
            target=sniffer_process_target,
            args=(self.queue, self.iface, self.bpf_filter)
        )
        self.process.start()
        logging.info(f"Sniffer process started with PID: {self.process.pid}")

        while not self.stop_event.is_set():
            try:
                # Use a timeout on the queue to remain responsive
                pkt_bytes = self.queue.get(timeout=0.5)
                # Emit the raw bytes. Reconstruction will happen in the main thread.
                self.packet_bytes_received.emit(pkt_bytes)
            except queue.Empty:
                continue
            except Exception as e:
                logging.error(f"Error in SnifferThread queue loop: {e}")

        logging.info("SnifferThread manager loop stopped.")


    def stop(self):
        logging.info("Stopping sniffer manager thread and process...")
        self.stop_event.set()
        if self.process and self.process.is_alive():
            logging.info(f"Terminating sniffer process {self.process.pid}...")
            self.process.terminate()
            self.process.join(timeout=2) # Wait for the process to terminate
            if self.process.is_alive():
                logging.warning(f"Sniffer process {self.process.pid} did not terminate gracefully, killing.")
                self.process.kill()
            logging.info("Sniffer process stopped.")

class ChannelHopperThread(QThread):
    """A thread to automatically hop Wi-Fi channels on Linux for scanning."""
    def __init__(self, iface):
        super().__init__()
        self.iface = iface
        self.stop_event = Event()
    def run(self):
        if sys.platform != "linux":
            logging.warning("Channel hopping is only supported on Linux.")
            return
        logging.info(f"Channel hopper started for interface {self.iface}")
        channels = [1, 6, 11, 2, 7, 3, 8, 4, 9, 5, 10]
        while not self.stop_event.is_set():
            for ch in channels:
                if self.stop_event.is_set(): break
                try:
                    os.system(f"iwconfig {self.iface} channel {ch}")
                    time.sleep(0.5)
                except Exception as e:
                    logging.error(f"Failed to hop channel: {e}")
                    break
        logging.info("Channel hopper stopped.")
    def stop(self): self.stop_event.set()

class WorkerThread(QThread):
    """A generic QThread to run any function in the background."""
    def __init__(self, target, args=()): super().__init__(); self.target = target; self.args = args
    def run(self): self.target(*self.args)

class ResourceMonitorThread(QThread):
    """A thread that monitors and emits system resource usage statistics."""
    stats_updated = pyqtSignal(dict)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.stop_event = Event()
        self.is_paused = False
        self.interval = 1 # default interval

    def run(self):
        """The main loop for monitoring resources."""
        psutil.cpu_percent() # Initial call to prevent first reading from being 0.0
        last_disk_io = psutil.disk_io_counters()
        last_net_io = psutil.net_io_counters()

        while not self.stop_event.is_set():
            if self.is_paused:
                time.sleep(1)
                continue

            time.sleep(self.interval)

            if self.stop_event.is_set():
                break

            cpu_percent = psutil.cpu_percent()
            ram_percent = psutil.virtual_memory().percent

            # GPU Stats
            gpu_percent = 0
            if GPUtil:
                try:
                    gpus = GPUtil.getGPUs()
                    if gpus:
                        gpu = gpus[0] # Use the first GPU
                        gpu_percent = gpu.load * 100
                except Exception as e:
                    logging.debug(f"Could not retrieve GPU stats: {e}")


            disk_io = psutil.disk_io_counters()
            read_mb_s = (disk_io.read_bytes - last_disk_io.read_bytes) / (1024**2) / self.interval
            write_mb_s = (disk_io.write_bytes - last_disk_io.write_bytes) / (1024**2) / self.interval
            last_disk_io = disk_io

            net_io = psutil.net_io_counters()
            sent_kb_s = (net_io.bytes_sent - last_net_io.bytes_sent) / 1024 / self.interval
            recv_kb_s = (net_io.bytes_recv - last_net_io.bytes_recv) / 1024 / self.interval
            last_net_io = net_io

            stats = {
                "cpu_percent": cpu_percent,
                "ram_percent": ram_percent,
                "gpu_percent": gpu_percent,
                "disk_str": f"{read_mb_s:.2f}/{write_mb_s:.2f} MB/s",
                "net_str": f"{sent_kb_s:.2f}/{recv_kb_s:.2f} KB/s"
            }
            self.stats_updated.emit(stats)

    def set_interval(self, interval):
        self.interval = interval
        self.is_paused = False

    def pause(self):
        self.is_paused = True

    def stop(self):
        self.stop_event.set()

class HandshakeSnifferThread(QThread):
    """A specialized thread to capture WPA 4-way handshakes."""
    handshake_captured = pyqtSignal(str, str) # BSSID, file_path
    log_message = pyqtSignal(str)

    def __init__(self, iface, bssid, parent=None):
        super().__init__(parent)
        self.iface = iface
        self.bssid = bssid
        self.packets = []
        self.stop_event = Event()

    def run(self):
        self.log_message.emit(f"Starting handshake capture for BSSID: {self.bssid} on {self.iface}")
        try:
            sniff(iface=self.iface, prn=self._packet_handler, stop_filter=lambda p: self.stop_event.is_set(), filter="ether proto 0x888e")
        except Exception as e:
            self.log_message.emit(f"Handshake sniffer error: {e}")
        self.log_message.emit("Handshake sniffer stopped.")

    def _packet_handler(self, pkt):
        self.packets.append(pkt)
        # Simple check: once we have >= 4 EAPOL packets, save and stop.
        # A more robust implementation would check the actual handshake sequence.
        if len(self.packets) >= 4:
            self.log_message.emit("Potential handshake captured (4 EAPOL packets). Saving to file.")
            file_path = f"handshake_{self.bssid.replace(':', '')}.pcap"
            wrpcap(file_path, self.packets)
            self.handshake_captured.emit(self.bssid, file_path)
            self.stop()

    def stop(self):
        self.stop_event.set()

if PYQTGRAPH_AVAILABLE:
    class ResourceGraph(pg.PlotWidget):
        """A custom PlotWidget for displaying a scrolling resource graph."""
        def __init__(self, parent=None, title="", color='c', text_color=(221, 221, 221)):
            super().__init__(parent)
            self.setMouseEnabled(x=False, y=False)
            self.setMenuEnabled(False)
            self.getPlotItem().hideAxis('bottom')
            self.getPlotItem().hideAxis('left')
            self.setBackground(background=(40, 44, 52)) # Default to dark theme background
            self.setRange(yRange=(0, 100), padding=0)

            self.data = np.zeros(60) # 60 data points for a 1-minute history at 1s refresh
            self.curve = self.plot(self.data, pen=pg.mkPen(color, width=2))

            self.text = pg.TextItem(text="", color=text_color, anchor=(0.5, 0.5))
            self.text.setPos(30, 50) # Position it in the middle of the graph
            self.addItem(self.text)


        def update_data(self, new_value):
            """Shifts the data and adds a new value to the end."""
            self.data[:-1] = self.data[1:]
            self.data[-1] = new_value
            self.curve.setData(self.data)
            self.text.setText(f"{new_value:.0f}%")
else:
    # If pyqtgraph is not available, create a dummy widget to avoid crashing.
    class ResourceGraph(QWidget):
        def __init__(self, parent=None, title="", color='c', text_color=(221, 221, 221)):
            super().__init__(parent)
            layout = QVBoxLayout(self)
            label = QLabel("Graphs disabled\n(pyqtgraph not installed)")
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            label.setStyleSheet("color: #888;")
            layout.addWidget(label)
            self.setMinimumHeight(60)
            # Make the placeholder visible
            self.setStyleSheet("background-color: #2d313a; border: 1px solid #444;")

        def update_data(self, new_value):
            """Dummy method, does nothing."""
            pass

class SubdomainResultsDialog(QDialog):
    """A dialog to show a list of found subdomains with an export option."""
    def __init__(self, subdomains, domain, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Subdomain Scan Results for {domain}")
        self.setMinimumSize(500, 400)
        self.parent = parent # To access the export handler
        self.domain = domain # Store domain for context

        layout = QVBoxLayout(self)

        summary_label = QLabel(f"<b>Found {len(subdomains)} unique subdomains.</b>")
        layout.addWidget(summary_label)

        self.tree = QTreeWidget()
        self.tree.setColumnCount(1)
        self.tree.setHeaderLabels(["Subdomain"])
        for sub in subdomains:
            self.tree.addTopLevelItem(QTreeWidgetItem([sub]))
        self.tree.resizeColumnToContents(0)
        layout.addWidget(self.tree)

        button_layout = QHBoxLayout()
        export_button = self.parent._create_export_button(self.tree)
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(lambda: self.parent._send_to_ai_analyst("subdomain", self.tree, self.domain))
        button_layout.addWidget(export_button)
        button_layout.addWidget(analyze_button)
        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)

        layout.addLayout(button_layout)

class NmapSummaryDialog(QDialog):
    """A dialog to show a summary of Nmap scan results from XML."""
    def __init__(self, xml_data, target_context, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Nmap Scan Summary")
        self.setMinimumSize(700, 500)
        self.xml_data = xml_data
        self.target_context = target_context
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Host / Details", "Port", "Service", "Version"])
        self.tree.header().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.tree.header().setStretchLastSection(False)
        layout.addWidget(self.tree)

        self.parse_and_populate(xml_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("nmap", self.xml_data, self.target_context)
            self.accept() # Close dialog after sending

    def parse_and_populate(self, xml_data):
        if not LXML_AVAILABLE:
            self.tree.addTopLevelItem(QTreeWidgetItem(["LXML library not installed."]))
            return
        if not xml_data:
            self.tree.addTopLevelItem(QTreeWidgetItem(["No XML data to parse."]))
            return

        try:
            parser = etree.XMLParser(recover=True, no_network=True, dtd_validation=False)
            root = etree.fromstring(xml_data.encode('utf-8'), parser=parser)

            for host in root.findall('host'):
                if host.find('status').get('state') != 'up':
                    continue

                address = host.find('address').get('addr')
                hostname_elem = host.find('hostnames/hostname')
                hostname = hostname_elem.get('name') if hostname_elem is not None else ""

                host_text = f"{address} ({hostname})" if hostname else address
                host_item = QTreeWidgetItem([host_text])
                host_item.setExpanded(True)
                self.tree.addTopLevelItem(host_item)

                ports_elem = host.find('ports')
                if ports_elem is None:
                    continue

                for port in ports_elem.findall('port'):
                    if port.find('state').get('state') == 'open':
                        port_id = port.get('portid')
                        protocol = port.get('protocol')

                        service_elem = port.find('service')
                        service = service_elem.get('name', '') if service_elem is not None else ''
                        version_parts = []
                        if service_elem is not None:
                            if service_elem.get('product'): version_parts.append(service_elem.get('product'))
                            if service_elem.get('version'): version_parts.append(service_elem.get('version'))
                        version = " ".join(version_parts)

                        port_item = QTreeWidgetItem(["", f"{port_id}/{protocol}", service, version])
                        host_item.addChild(port_item)

        except Exception as e:
            logging.error(f"Failed to parse Nmap XML for summary: {e}", exc_info=True)
            self.tree.addTopLevelItem(QTreeWidgetItem(["Error parsing XML data."]))

class HttpxResultsDialog(QDialog):
    def __init__(self, json_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("httpx Probe Results")
        self.setMinimumSize(800, 500)
        self.json_data = json_data
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        # Define columns based on common httpx JSON output
        self.tree.setColumnCount(5)
        self.tree.setHeaderLabels(["URL", "Status Code", "Title", "Web Server", "Technologies"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            # The AI can analyze the raw JSON data
            self.parent.ai_assistant_tab.send_to_analyst("httpx", self.json_data, "httpx probe results")
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            # httpx outputs JSON objects separated by newlines
            results = [json.loads(line) for line in json_data.strip().split('\n') if line]
            for res in results:
                url = res.get('url', '')
                status_code = str(res.get('status_code', ''))
                title = res.get('title', '')
                web_server = res.get('webserver', '')
                tech = ", ".join(res.get('tech', []))

                item = QTreeWidgetItem([url, status_code, title, web_server, tech])
                self.tree.addTopLevelItem(item)
        except json.JSONDecodeError:
            # Handle case where output is not JSON
            item = QTreeWidgetItem(["Error parsing JSON output. Displaying raw data in console."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing httpx JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class DirsearchResultsDialog(QDialog):
    def __init__(self, json_data, target_context, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"dirsearch Results for {target_context}")
        self.setMinimumSize(800, 500)
        self.json_data = json_data
        self.target_context = target_context
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Path", "Status Code", "Content-Length", "Redirect"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("dirsearch", self.json_data, self.target_context)
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            # dirsearch report is a dictionary where keys are hostnames
            results = json.loads(json_data)
            for host, findings in results.items():
                host_item = QTreeWidgetItem([f"Host: {host}"])
                self.tree.addTopLevelItem(host_item)
                host_item.setExpanded(True)
                for finding in findings:
                    path = finding.get('path', '')
                    status = str(finding.get('status', ''))
                    length = str(finding.get('content-length', ''))
                    redirect = finding.get('redirect', '')

                    child_item = QTreeWidgetItem([path, status, length, redirect])
                    host_item.addChild(child_item)
        except json.JSONDecodeError:
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing dirsearch JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class FfufResultsDialog(QDialog):
    def __init__(self, json_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("ffuf Scan Results")
        self.setMinimumSize(800, 500)
        self.json_data = json_data
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["URL", "Status", "Length", "Words"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("ffuf", self.json_data, "ffuf scan results")
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            results = json.loads(json_data).get('results', [])
            for res in results:
                url = res.get('url', '')
                status = str(res.get('status', ''))
                length = str(res.get('length', ''))
                words = str(res.get('words', ''))

                item = QTreeWidgetItem([url, status, length, words])
                self.tree.addTopLevelItem(item)
        except (json.JSONDecodeError, AttributeError):
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing ffuf JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class NucleiResultsDialog(QDialog):
    def __init__(self, json_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Nuclei Scan Results")
        self.setMinimumSize(900, 600)
        self.json_data = json_data
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(5)
        self.tree.setHeaderLabels(["Template ID", "Name", "Severity", "Host", "Matched At"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("nuclei", self.json_data, "Nuclei scan results")
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            # Nuclei outputs JSON objects separated by newlines
            results = [json.loads(line) for line in json_data.strip().split('\n') if line]
            for res in results:
                template_id = res.get('template-id', '')
                name = res.get('info', {}).get('name', '')
                severity = res.get('info', {}).get('severity', '')
                host = res.get('host', '')
                matched_at = res.get('matched-at', '')

                item = QTreeWidgetItem([template_id, name, severity, host, matched_at])
                self.tree.addTopLevelItem(item)

                # Add extracted results as children for more detail
                if 'extracted-results' in res:
                    for i, extracted in enumerate(res['extracted-results']):
                        child_item = QTreeWidgetItem([f"  - Extracted {i+1}", str(extracted)])
                        item.addChild(child_item)

                item.setExpanded(True)

        except json.JSONDecodeError:
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing Nuclei JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class TruffleHogResultsDialog(QDialog):
    def __init__(self, json_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("TruffleHog Scan Results")
        self.setMinimumSize(900, 600)
        self.json_data = json_data
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Detector", "Decoder", "File", "Raw Secret"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("trufflehog", self.json_data, "TruffleHog scan results")
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            # TruffleHog outputs JSON objects separated by newlines
            results = [json.loads(line) for line in json_data.strip().split('\n') if line]
            for res in results:
                detector = res.get('DetectorType', '')
                decoder = res.get('DecoderType', '')
                file = res.get('File', '')
                raw = res.get('Raw', '')

                item = QTreeWidgetItem([detector, decoder, file, raw])
                self.tree.addTopLevelItem(item)
        except json.JSONDecodeError:
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing TruffleHog JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class Enum4LinuxNGResultsDialog(QDialog):
    def __init__(self, json_data, target_context, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"enum4linux-ng Results for {target_context}")
        self.setMinimumSize(800, 600)
        self.json_data = json_data
        self.target_context = target_context
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(2)
        self.tree.setHeaderLabels(["Finding", "Details"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("enum4linux-ng", self.json_data, self.target_context)
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            results = json.loads(json_data)
            # The JSON is a list of dictionaries, each representing a finding
            for finding in results:
                method = finding.get('method', 'N/A')
                item = QTreeWidgetItem([method])
                self.tree.addTopLevelItem(item)

                # Add all other keys as children
                for key, value in finding.items():
                    if key != 'method':
                        child_item = QTreeWidgetItem([f"  {key}", str(value)])
                        item.addChild(child_item)
                item.setExpanded(True)
        except json.JSONDecodeError:
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing enum4linux-ng JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class DnsReconResultsDialog(QDialog):
    def __init__(self, json_data, target_context, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"dnsrecon Results for {target_context}")
        self.setMinimumSize(800, 600)
        self.json_data = json_data
        self.target_context = target_context
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Type", "Target", "Address", "Name"])
        layout.addWidget(self.tree)

        self.parse_and_populate(json_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("dnsrecon", self.json_data, self.target_context)
            self.accept()

    def parse_and_populate(self, json_data):
        try:
            results = json.loads(json_data)
            # The JSON is a list of dictionaries
            for res in results:
                rec_type = res.get('type', 'N/A')
                target = res.get('target', 'N/A')
                address = res.get('address', 'N/A')
                name = res.get('name', 'N/A')

                item = QTreeWidgetItem([rec_type, target, address, name])
                self.tree.addTopLevelItem(item)
        except json.JSONDecodeError:
            item = QTreeWidgetItem(["Error parsing JSON output."])
            self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing dnsrecon JSON: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

class SherlockResultsDialog(QDialog):
    def __init__(self, csv_data, target_context, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Sherlock Results for {target_context}")
        self.setMinimumSize(800, 600)
        self.csv_data = csv_data
        self.target_context = target_context
        self.parent = parent

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Username", "Service Name", "URL", "Status"])
        layout.addWidget(self.tree)

        self.parse_and_populate(csv_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        button_layout = QHBoxLayout()
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(self.send_to_ai)
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def send_to_ai(self):
        if self.parent:
            self.parent.ai_assistant_tab.send_to_analyst("sherlock", self.csv_data, self.target_context)
            self.accept()

    def parse_and_populate(self, csv_data):
        try:
            # Use Python's built-in csv module to parse the data
            reader = csv.reader(csv_data.strip().splitlines())
            header = next(reader) # Skip header row
            for row in reader:
                # Assuming standard sherlock csv format: username,name,url,status
                if len(row) >= 4:
                    item = QTreeWidgetItem(row)
                    self.tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Error parsing Sherlock CSV: {e}")
            self.tree.addTopLevelItem(QTreeWidgetItem([f"An unexpected error occurred: {e}"]))

# --- Main Application ---
class GScapy(QMainWindow):
    """The main application window, holding all UI elements and logic."""
    def __init__(self):
        """Initializes the main window, UI components, and internal state."""
        super().__init__()
        self.setWindowTitle("GScapy + AI - The Modern Scapy Interface with AI")
        # Construct path to icon relative to the script's location for robustness
        script_dir = os.path.dirname(os.path.realpath(__file__))
        icon_path = os.path.join(script_dir, "icons", "new_logo.png")
        self.setWindowIcon(QIcon(icon_path))
        self.setGeometry(100, 100, 1200, 800)

        self.current_user = None
        self.packets_data = []; self.sniffer_thread = None; self.channel_hopper = None
        self.packet_layers = []; self.current_field_widgets = []; self.tcp_flag_vars = {}
        self.tool_results_queue = Queue()
        self.is_tool_running = False
        self.loaded_flood_packet = None
        self.found_networks = {}
        self.active_threads = []
        self.thread_finish_lock = Lock()
        self.finished_thread_count = 0
        self.tool_stop_event = Event()
        self.arp_spoof_current_victim = None
        self.arp_spoof_current_target = None
        self.resource_monitor_thread = None
        self.nmap_last_xml = None
        self.nmap_xml_temp_file = None
        self.aircrack_thread = None
        self.ps_thread_lock = Lock()
        self.ps_finished_threads = 0
        self.bf_ssid_list = []
        self.krack_thread = None
        self.sniffer_packet_buffer = []
        self.sniffer_buffer_lock = Lock()
        self.super_scan_active = False
        self.lab_test_chain = []
        self.threat_intel_loaded = False
        self.history_loaded = False
        # self.tool_config_widgets is no longer used.

        self.nmap_script_presets = {
            "HTTP Service Info": ("http-title,http-headers", "", "Gathers the title and headers from web servers."),
            "SMB OS Discovery": ("smb-os-discovery", "", "Attempts to determine the OS, computer name, and domain from SMB."),
            "FTP Anonymous Login": ("ftp-anon", "", "Checks if an FTP server allows anonymous login."),
            "DNS Brute-force": ("dns-brute", "", "Attempts to enumerate DNS hostnames by brute-forcing common subdomain names."),
            "SSL/TLS Certificate Info": ("ssl-cert,sslv2", "", "Retrieves the server's SSL certificate and checks for weak SSLv2 support."),
            "SMTP User Enumeration": ("smtp-enum-users", "smtp-enum-users.methods={VRFY,EXPN,RCPT}", "Attempts to enumerate users on an SMTP server."),
            "Vulnerability Scan (Vulners)": ("vulners", "", "Checks for vulnerabilities based on service versions using Vulners.com. Requires -sV."),
            "SMB Share & User Enum": ("smb-enum-shares,smb-enum-users", "", "Enumerates shared folders and user accounts on an SMB server."),
            "Service Banner Grabbing": ("banner", "", "Connects to open ports and prints the service banner.")
        }

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QVBoxLayout(self.central_widget)

        self._create_resource_bar()
        self.menu_bar = QMenuBar(self)
        self.setMenuBar(self.menu_bar)
        # The menu bar will be populated after login by calling _update_menu_bar()
        self._create_status_bar()
        self._create_header_bar()

        # Config widgets are now created inside their respective tool tabs.

        self._create_main_tabs(); self._create_log_panel(); self._setup_logging()
        self.tab_widget.currentChanged.connect(self._on_main_tab_changed)

        self._setup_result_handlers()
        self.results_processor = QTimer(self); self.results_processor.timeout.connect(self._process_tool_results); self.results_processor.start(100)

        # Setup timer for batching sniffer UI updates
        self.sniffer_ui_update_timer = QTimer(self)
        self.sniffer_ui_update_timer.timeout.connect(self._update_sniffer_display)
        self.sniffer_ui_update_timer.start(500) # Update every 500ms

        # Setup timer for the clock
        self.clock_timer = QTimer(self)
        self.clock_timer.timeout.connect(self._update_clock)
        self.clock_timer.start(1000) # 1000 ms = 1 second
        self._update_clock() # Initial call to show time immediately

        # Start the resource monitor
        self.resource_monitor_thread = ResourceMonitorThread(self)
        self.resource_monitor_thread.stats_updated.connect(self._update_resource_stats)
        self.resource_monitor_thread.start()

        self._update_tool_targets() # Initial population after all widgets are created
        self._set_user_avatar()
        logging.info("GScapy application initialized.")

    def _set_user_avatar(self):
        """Sets the user profile button icon based on the user's avatar."""
        if not self.current_user:
            return

        avatar_data = self.current_user.get('avatar')
        if avatar_data:
            pixmap = QPixmap()
            pixmap.loadFromData(avatar_data)
            self.user_profile_button.setIcon(QIcon(pixmap))
        else:
            # Fallback to a default icon if no avatar is set
            self.user_profile_button.setIcon(QIcon("icons/users.svg"))

    def _show_user_menu(self):
        """Shows a context menu for the user profile button."""
        menu = QMenu(self)
        profile_action = menu.addAction("Profile...")
        logout_action = menu.addAction("Logout")

        action = menu.exec(self.user_profile_button.mapToGlobal(QPoint(0, self.user_profile_button.height())))

        if action == profile_action:
            self._show_user_profile()
        elif action == logout_action:
            self._logout()

    def _show_user_profile(self):
        """Opens the user profile dialog."""
        if not self.current_user:
            return

        # Re-fetch user data to ensure it's up-to-date
        self.current_user = database.get_user_by_id(self.current_user['id'])
        dialog = UserProfileDialog(self.current_user, self)
        if dialog.exec():
            # If changes were saved, re-fetch user data and update avatar
            self.current_user = database.get_user_by_id(self.current_user['id'])
            self._set_user_avatar()

    def _logout(self):
        """Logs the current user out and shows the login screen."""
        self.close()
        # A bit of a hacky way to restart, but effective for a desktop app
        # A more robust solution might use a dedicated controller class to manage windows
        python = sys.executable
        os.execl(python, python, *sys.argv)

    def _update_menu_bar(self):
        """Creates or updates the main menu bar based on the current user's role."""
        self.menu_bar.clear()

        # --- File Menu ---
        file_menu = self.menu_bar.addMenu("&File")
        file_menu.addAction("&Save Captured Packets", self.save_packets)
        file_menu.addAction("&Load Packets from File", self.load_packets)
        file_menu.addSeparator()
        file_menu.addAction("&Exit", self.close)

        # --- Admin Menu (Conditional) ---
        # Use .get() for safer dictionary access, in case current_user is None
        logging.info(f"Updating menu bar for user: {self.current_user}")
        if self.current_user and self.current_user.get('is_admin'):
            admin_menu = self.menu_bar.addMenu("&Admin")
            admin_menu.addAction(QIcon("icons/new_logo.png"), "Admin Panel...", self._show_admin_panel)

        # --- Help Menu ---
        help_menu = self.menu_bar.addMenu("&Help")
        help_menu.addAction("&About GScapy", self._show_about_dialog)
        help_menu.addSeparator()
        help_menu.addAction("&AI Settings...", self._show_ai_settings_dialog)
        help_menu.addAction("AI Guide", self._show_ai_guide_dialog)

    def _show_admin_panel(self):
        """Shows the admin panel dialog."""
        admin_dialog = AdminPanelDialog(self)
        admin_dialog.exec()

    def _show_ai_settings_dialog(self):
        """Shows the AI settings dialog."""
        dialog = AISettingsDialog(self)
        dialog.exec()

    def _show_ai_guide_dialog(self):
        """Shows the AI features user guide."""
        dialog = AIGuideDialog(self)
        dialog.exec()

    def get_ai_settings(self):
        """
        Loads AI settings from the JSON file and returns a dictionary
        containing the active provider's details (endpoint, model, api_key).
        """
        settings_file = "ai_settings.json"
        try:
            if not os.path.exists(settings_file):
                # Show settings dialog if no config exists
                if self._show_ai_settings_dialog() == QDialog.DialogCode.Rejected:
                    return None # User cancelled

            with open(settings_file, 'r') as f:
                settings = json.load(f)

            active_provider_name = settings.get("active_provider")
            active_model_name = settings.get("active_model")

            if not active_provider_name or not active_model_name:
                self.ai_assistant_tab.handle_ai_error("No active AI model selected. Please click the settings icon to choose one.")
                return None

            provider_details = {}
            if active_provider_name == "local_ai":
                local_settings = settings.get("local_ai", {})
                provider_details = {
                    "provider": "local_ai",
                    "endpoint": local_settings.get("endpoint"),
                    "model": active_model_name,
                    "api_key": None
                }
            else: # It's an online service
                online_settings = settings.get("online_ai", {})
                provider_data = online_settings.get(active_provider_name, {})
                api_key = provider_data.get("api_key")

                endpoint = ""
                if active_provider_name == "OpenAI":
                    endpoint = "https://api.openai.com/v1/chat/completions"
                # ... (add other online providers here)

                provider_details = {
                    "provider": active_provider_name,
                    "endpoint": endpoint,
                    "model": active_model_name,
                    "api_key": api_key
                }

            if not provider_details.get("endpoint"):
                 self.ai_assistant_tab.handle_ai_error(f"Endpoint for '{active_provider_name}' is missing or not supported yet.")
                 return None

            return provider_details

        except (IOError, json.JSONDecodeError) as e:
            self.ai_assistant_tab.handle_ai_error(f"Error loading AI settings: {e}")
            return None
        except Exception as e:
            self.ai_assistant_tab.handle_ai_error(f"An unexpected error occurred while getting AI settings: {e}")
            return None


    def _show_about_dialog(self):
        dialog = QMessageBox(self)
        dialog.setWindowTitle("About GScapy + AI")

        # Add the logo
        pixmap = QIcon(os.path.join("icons", "new_logo.png")).pixmap(80, 80)
        dialog.setIconPixmap(pixmap)

        about_text = """
        <b>GScapy + AI v3.0</b>
        <p>The Modern Scapy Interface with AI.</p>
        <p>This application provides tools for sniffing, crafting, and analyzing network packets, with AI-powered analysis and guidance.</p>
        <br>
        <p><b>Developer:</b><br>Mohammadmahdi Farhadianfard (ao ga nai 😁)<br>
        mohammadmahdi.farhadianfard@gmail.com</p>
        """
        dialog.setText(about_text)
        dialog.exec()

    def _create_status_bar(self):
        self.status_bar = QStatusBar(self); self.setStatusBar(self.status_bar); self.status_bar.showMessage("Ready")

    def _create_resource_bar(self):
        """Creates the top resource monitor bar."""
        resource_frame = QFrame(); resource_frame.setFrameShape(QFrame.Shape.StyledPanel)
        resource_layout = QHBoxLayout(resource_frame)
        resource_layout.setContentsMargins(5, 2, 5, 2)

        # Add Logo and Tooltip
        logo_label = QLabel()
        logo_pixmap = QIcon(os.path.join("icons", "new_logo.png")).pixmap(40, 40)
        logo_label.setPixmap(logo_pixmap)
        logo_label.setToolTip("GScapy made by Poorija, Email: mohammadmahdi.farhadianfard@gmail.com")
        resource_layout.addWidget(logo_label)
        resource_layout.addSpacing(15)

        resource_layout.addWidget(QLabel("<b>CPU:</b>"))
        self.cpu_graph = ResourceGraph(color='c')
        self.cpu_graph.setFixedHeight(60)
        self.cpu_graph.setMaximumWidth(250)
        resource_layout.addWidget(self.cpu_graph, 1) # Add stretch factor

        resource_layout.addWidget(QLabel("<b>RAM:</b>"))
        self.ram_graph = ResourceGraph(color='m')
        self.ram_graph.setFixedHeight(60)
        self.ram_graph.setMaximumWidth(250)
        resource_layout.addWidget(self.ram_graph, 1) # Add stretch factor

        if GPUtil:
            resource_layout.addWidget(QLabel("<b>GPU:</b>"))
            self.gpu_graph = ResourceGraph(color='y')
            self.gpu_graph.setFixedHeight(60)
            self.gpu_graph.setMaximumWidth(250)
            resource_layout.addWidget(self.gpu_graph, 1)

        resource_layout.addWidget(QLabel("<b>Disk R/W:</b>"))
        self.disk_label = QLabel("---/--- MB/s"); resource_layout.addWidget(self.disk_label)
        resource_layout.addStretch()

        resource_layout.addWidget(QLabel("<b>Net Sent/Recv:</b>"))
        self.net_label = QLabel("---/--- KB/s"); resource_layout.addWidget(self.net_label)
        resource_layout.addStretch()

        self.time_label = QLabel("Loading...")
        self.time_label.setToolTip("Current system date, time, and timezone")
        self.time_label.setStyleSheet("font-weight: bold;")

        separator = QFrame()
        separator.setFrameShape(QFrame.Shape.VLine)
        separator.setFrameShadow(QFrame.Shadow.Sunken)
        resource_layout.addWidget(separator)

        resource_layout.addWidget(self.time_label)
        resource_layout.addStretch()

        resource_layout.addWidget(QLabel("<b>Refresh:</b>"))
        self.refresh_combo = QComboBox()
        self.refresh_combo.addItems(["1s", "2s", "5s", "Off"])
        resource_layout.addWidget(self.refresh_combo)

        # --- User Profile Button ---
        self.user_profile_button = QPushButton()
        self.user_profile_button.setFlat(True)
        self.user_profile_button.setIconSize(QSize(32, 32))
        self.user_profile_button.setToolTip("User Profile & Logout")
        resource_layout.addWidget(self.user_profile_button)

        self.main_layout.addWidget(resource_frame)
        self.refresh_combo.textActivated.connect(self._handle_refresh_interval_change)
        self.user_profile_button.clicked.connect(self._show_user_menu)


    def _update_clock(self):
        """Updates the time label with the current time and timezone."""
        # Use a try-except block to handle potential time formatting issues on different OSes
        try:
            # This format includes Year-Month-Day, Hour:Minute:Second, and Timezone Name
            current_time = time.strftime("%Y-%m-%d %H:%M:%S %Z")
        except Exception as e:
            logging.warning(f"Could not format time with timezone: {e}")
            # Fallback to a simpler format if the timezone name (%Z) causes issues
            current_time = time.strftime("%H:%M:%S")
        self.time_label.setText(current_time)

    def _update_resource_stats(self, stats):
        """Updates the resource labels with new stats from the monitor thread."""
        self.cpu_graph.update_data(stats["cpu_percent"])
        self.ram_graph.update_data(stats["ram_percent"])
        if hasattr(self, 'gpu_graph'):
            self.gpu_graph.update_data(stats.get("gpu_percent", 0))
        self.disk_label.setText(stats["disk_str"])
        self.net_label.setText(stats["net_str"])

    def _handle_refresh_interval_change(self, text):
        """Updates the resource monitor's refresh interval."""
        if not self.resource_monitor_thread:
            return

        if text == "Off":
            self.resource_monitor_thread.pause()
        else:
            interval = int(text.replace('s', ''))
            self.resource_monitor_thread.set_interval(interval)

    def _create_header_bar(self):
        """Creates the top header bar with interface and theme selectors."""
        header_frame = QWidget()
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(0, 5, 0, 5)

        # Interface Selector
        header_layout.addWidget(QLabel("Network Interface:"))
        try:
            ifaces = ["Automatic"] + [iface.name for iface in get_working_ifaces()]
        except Exception as e:
            logging.error(f"Could not get network interfaces: {e}", exc_info=True)
            ifaces = ["Automatic"]
        self.iface_combo = QComboBox()
        self.iface_combo.addItems(ifaces)
        self.iface_combo.currentTextChanged.connect(self._update_tool_targets)
        header_layout.addWidget(self.iface_combo)

        header_layout.addStretch()

        # Theme Switcher
        header_layout.addWidget(QLabel("Theme:"))
        self.theme_combo = QComboBox()
        self.theme_combo.addItems([theme.replace('.xml', '') for theme in list_themes()])
        self.theme_combo.textActivated.connect(self._handle_theme_change)
        header_layout.addWidget(self.theme_combo)

        self.main_layout.addWidget(header_frame)

    def _handle_theme_change(self, theme_name):
        theme_file = f"{theme_name}.xml"
        invert_secondary = "light" in theme_name

        # This dictionary must be kept in sync with the one in login.py and main()
        extra_qss = {
            'QGroupBox': {
                'border': '1px solid #444;',
                'border-radius': '8px',
                'margin-top': '10px',
            },
            'QGroupBox::title': {
                'subcontrol-origin': 'margin',
                'subcontrol-position': 'top left',
                'padding': '0 10px',
            },
            'QTabWidget::pane': {
                'border-top': '1px solid #444;',
                'margin-top': '-1px',
            },
            'QFrame': {
                'border-radius': '8px',
            },
            'QPushButton': {
                'border-radius': '8px',
            },
            'QLineEdit': {
                'border-radius': '8px',
            },
            'QComboBox': {
                'border-radius': '8px',
            },
            'QTextEdit': {
                'border-radius': '8px',
            },
            'QPlainTextEdit': {
                'border-radius': '8px',
            },
            'QListWidget': {
                'border-radius': '8px',
            },
            'QTreeWidget': {
                'border-radius': '8px',
            }
        }

        apply_stylesheet(QApplication.instance(), theme=theme_file, invert_secondary=invert_secondary, extra=extra_qss)

        # After applying the stylesheet, notify the AI tab to update its themed icons
        if hasattr(self, 'ai_assistant_tab'):
            self.ai_assistant_tab.update_theme()

    def get_selected_iface(self):
        iface = self.iface_combo.currentText()
        return iface if iface != "Automatic" else None

    def _create_main_tabs(self):
        """Creates the main QTabWidget and adds all the tool tabs."""
        self.tab_widget = QTabWidget()
        self.tab_widget.setStyleSheet("""
            QTabWidget::tab-bar {
                alignment: center;
            }
            QTabBar::tab:!selected:!last {
                border-right: 1px solid #444;
            }
        """)
        self.main_layout.addWidget(self.tab_widget)
        self.tab_widget.addTab(self._create_sniffer_tab(), QIcon("icons/search.svg"), "Packet Sniffer")
        self.tab_widget.addTab(self._create_crafter_tab(), QIcon("icons/edit-3.svg"), "Packet Crafter")
        self.tab_widget.addTab(self._create_tools_tab(), QIcon("icons/tool.svg"), "Network Tools")
        self.tab_widget.addTab(self._create_advanced_tools_tab(), QIcon("icons/shield.svg"), "Advanced Tools")
        self.tab_widget.addTab(self._create_wireless_tools_tab(), QIcon("icons/wifi.svg"), "Wireless Tools")
        self.tab_widget.addTab(self._create_reporting_tab(), QIcon("icons/file-text.svg"), "Reporting")
        self.tab_widget.addTab(self._create_lab_tab(), QIcon("icons/layers.svg"), "LAB")

        self.ai_assistant_tab = AIAssistantTab(self)
        self.tab_widget.addTab(self.ai_assistant_tab, QIcon("icons/terminal.svg"), "AI Assistant")

        self.tab_widget.addTab(self._create_threat_intelligence_tab(), QIcon("icons/database.svg"), "Threat Intelligence")
        self.tab_widget.addTab(self._create_history_tab(), QIcon("icons/file-text.svg"), "History")

        self.tab_widget.addTab(self._create_community_tools_tab(), QIcon("icons/users.svg"), "Community Tools")
        self.tab_widget.addTab(self._create_system_info_tab(), QIcon("icons/info.svg"), "System Info")

    def _create_threat_intelligence_tab(self):
        """Creates the tab container for the Threat Intelligence tools."""
        threat_tabs = QTabWidget()
        threat_tabs.addTab(self._create_recent_threats_tab(), "Recent Threats")
        threat_tabs.addTab(self._create_exploit_db_search_tab(), "Exploit-DB Search")
        return threat_tabs

    def _create_history_tab(self):
        """Creates the UI for the user's personal test history."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        controls_layout = QHBoxLayout()
        self.refresh_history_btn = QPushButton(QIcon("icons/refresh-cw.svg"), " Refresh History")
        self.refresh_history_btn.clicked.connect(self._populate_history_tab)
        controls_layout.addWidget(self.refresh_history_btn)
        controls_layout.addStretch()
        layout.addLayout(controls_layout)

        self.history_tree = QTreeWidget()
        self.history_tree.setColumnCount(4)
        self.history_tree.setHeaderLabels(["Timestamp", "Test Type", "Target", "Result Summary"])
        self.history_tree.header().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.history_tree.header().setStretchLastSection(True)
        self.history_tree.header().resizeSection(0, 180)
        self.history_tree.header().resizeSection(1, 150)
        self.history_tree.header().resizeSection(2, 200)
        layout.addWidget(self.history_tree)

        return widget

    def _populate_history_tab(self):
        """Fetches and displays the current user's test history."""
        if not self.current_user:
            return
        self.history_tree.clear()
        try:
            # The user_id is passed to get history for the current user only
            history_records = database.get_test_history(self.current_user['id'])
            for record in history_records:
                # The 'results' column can be long, so we'll show a summary.
                summary = (record['results'] or "").split('\n')[0]
                summary = (summary[:100] + '...') if len(summary) > 100 else summary

                item = QTreeWidgetItem([
                    record['timestamp'],
                    record['test_type'],
                    record['target'],
                    summary
                ])
                # Store the full result in the item for potential future detail view
                item.setData(0, Qt.ItemDataRole.UserRole, record['results'])
                self.history_tree.addTopLevelItem(item)
        except Exception as e:
            logging.error(f"Failed to populate history tab: {e}", exc_info=True)
            QMessageBox.critical(self, "History Error", f"Could not load test history: {e}")

    def _on_main_tab_changed(self, index):
        """Handler for when the main tab is changed, to auto-load data."""
        try:
            tab_text = self.tab_widget.tabText(index)
            if tab_text == "Threat Intelligence" and not self.threat_intel_loaded:
                logging.info("Threat Intelligence tab opened for the first time. Auto-fetching CVEs.")
                self.threat_intel_loaded = True
                self._start_fetch_threats()
            elif tab_text == "History" and not self.history_loaded:
                logging.info("History tab opened for the first time. Auto-populating.")
                self.history_loaded = True
                self._populate_history_tab()
        except Exception as e:
            logging.error(f"Error in _on_main_tab_changed: {e}")

    def _create_recent_threats_tab(self):
        """Creates the UI for displaying recent CVEs with pagination."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Controls ---
        controls_layout = QHBoxLayout()
        self.fetch_threats_btn = QPushButton(QIcon("icons/refresh-cw.svg"), " Refresh")
        self.fetch_threats_btn.setToolTip("Fetch the latest CVEs. This is done automatically on first view.")
        controls_layout.addWidget(self.fetch_threats_btn)
        controls_layout.addStretch()
        controls_layout.addWidget(QLabel("Items per page:"))
        self.threats_per_page_combo = QComboBox()
        self.threats_per_page_combo.addItems(["10", "20", "30", "40"])
        controls_layout.addWidget(self.threats_per_page_combo)
        self.threats_prev_btn = QPushButton("<< Previous")
        self.threats_next_btn = QPushButton("Next >>")
        self.threats_page_label = QLabel("Page 1 / 1")
        controls_layout.addWidget(self.threats_prev_btn)
        controls_layout.addWidget(self.threats_page_label)
        controls_layout.addWidget(self.threats_next_btn)
        layout.addLayout(controls_layout)

        # --- Results Display ---
        results_splitter = QSplitter(Qt.Orientation.Vertical)
        self.threats_table = QTreeWidget()
        self.threats_table.setColumnCount(4)
        self.threats_table.setHeaderLabels(["CVE ID", "Published", "CVSS", "Summary"])
        self.threats_table.header().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.threats_table.header().setStretchLastSection(True)
        self.threats_table.header().resizeSection(0, 150)
        self.threats_table.header().resizeSection(1, 150)
        self.threats_table.header().resizeSection(2, 50)
        results_splitter.addWidget(self.threats_table)

        self.threat_details_view = QTextBrowser()
        self.threat_details_view.setOpenExternalLinks(True)
        results_splitter.addWidget(self.threat_details_view)
        results_splitter.setSizes([300, 200])
        layout.addWidget(results_splitter)

        # --- Instance variables for state ---
        self.recent_threats_data = []
        self.threats_current_page = 0
        self.threats_items_per_page = 10

        # --- Connections ---
        self.fetch_threats_btn.clicked.connect(self._start_fetch_threats)
        self.threats_table.currentItemChanged.connect(self._display_threat_details)
        self.threats_next_btn.clicked.connect(self._go_to_next_threats_page)
        self.threats_prev_btn.clicked.connect(self._go_to_previous_threats_page)
        self.threats_per_page_combo.currentTextChanged.connect(self._change_threats_per_page)

        return widget

    def _start_fetch_threats(self):
        """Initiates the background thread to fetch recent CVEs."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        self.is_tool_running = True
        self.fetch_threats_btn.setEnabled(False)
        self.status_bar.showMessage("Fetching latest CVEs from cve.circl.lu...")
        self.threats_table.clear()
        self.threat_details_view.clear()

        self.worker = WorkerThread(self._recent_threats_thread)
        self.active_threads.append(self.worker)
        self.worker.start()

    def _recent_threats_thread(self):
        """Worker thread to fetch the last 30 CVEs."""
        q = self.tool_results_queue
        try:
            url = "https://cve.circl.lu/api/last"
            with urllib.request.urlopen(url, timeout=15) as response:
                if response.status == 200:
                    data = json.loads(response.read().decode('utf-8'))
                    q.put(('recent_threats_result', data))
                else:
                    raise Exception(f"API request failed with status code {response.status}")
        except Exception as e:
            logging.error(f"Failed to fetch recent threats: {e}", exc_info=True)
            q.put(('error', 'Threat Intelligence Error', str(e)))
        finally:
            q.put(('tool_finished', 'fetch_threats'))

    def _handle_recent_threats_result(self, data):
        """Handles the successful result from the threat fetcher thread."""
        self.recent_threats_data = data
        self.threats_current_page = 0
        self._update_threats_display()
        self.status_bar.showMessage(f"Successfully fetched {len(data)} recent CVEs.", 5000)

    def _update_threats_display(self):
        """Updates the threat table and pagination controls based on the current state."""
        self.threats_table.clear()

        start_index = self.threats_current_page * self.threats_items_per_page
        end_index = start_index + self.threats_items_per_page

        page_data = self.recent_threats_data[start_index:end_index]

        for cve in page_data:
            cve_id = cve.get('id', 'N/A')
            published = cve.get('Published', 'N/A')
            cvss = str(cve.get('cvss', 'N/A'))
            summary = cve.get('summary', '')

            item = QTreeWidgetItem([cve_id, published, cvss, summary])
            item.setData(0, Qt.ItemDataRole.UserRole, cve) # Store full data
            self.threats_table.addTopLevelItem(item)

        total_pages = (len(self.recent_threats_data) + self.threats_items_per_page - 1) // self.threats_items_per_page
        self.threats_page_label.setText(f"Page {self.threats_current_page + 1} / {total_pages}")
        self.threats_prev_btn.setEnabled(self.threats_current_page > 0)
        self.threats_next_btn.setEnabled(end_index < len(self.recent_threats_data))

    def _display_threat_details(self, current_item, previous_item):
        """Displays full details for the selected CVE."""
        if not current_item:
            return
        cve_data = current_item.data(0, Qt.ItemDataRole.UserRole)
        if not cve_data:
            return

        html = f"<h3>{cve_data.get('id', 'N/A')}</h3>"
        html += f"<p><b>Summary:</b><br>{cve_data.get('summary', 'No summary available.')}</p>"
        html += f"<b>Published:</b> {cve_data.get('Published', 'N/A')}<br>"
        html += f"<b>Modified:</b> {cve_data.get('Modified', 'N/A')}<br>"
        html += f"<b>CVSS Score:</b> {cve_data.get('cvss', 'N/A')}<br>"

        refs = cve_data.get('references', [])
        if refs:
            html += "<p><b>References:</b><ul>"
            for ref in refs:
                html += f'<li><a href="{ref}">{ref}</a></li>'
            html += "</ul></p>"

        self.threat_details_view.setHtml(html)

    def _go_to_next_threats_page(self):
        self.threats_current_page += 1
        self._update_threats_display()

    def _go_to_previous_threats_page(self):
        self.threats_current_page -= 1
        self._update_threats_display()

    def _change_threats_per_page(self, text):
        self.threats_items_per_page = int(text)
        self.threats_current_page = 0 # Reset to first page
        self._update_threats_display()

    def _create_exploit_db_search_tab(self):
        """Creates the UI for the Exploit-DB Search tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- API Key Input ---
        api_key_layout = QHBoxLayout()
        api_key_layout.addWidget(QLabel("Vulners API Key:"))
        self.exploitdb_api_key_input = QLineEdit()
        self.exploitdb_api_key_input.setPlaceholderText("Get a free key from vulners.com")
        self.exploitdb_api_key_input.setEchoMode(QLineEdit.EchoMode.Password)
        api_key_layout.addWidget(self.exploitdb_api_key_input)
        save_api_key_btn = QPushButton("Save Key")
        api_key_layout.addWidget(save_api_key_btn)
        layout.addLayout(api_key_layout)

        # --- Search Controls ---
        controls_layout = QHBoxLayout()
        controls_layout.addWidget(QLabel("Search Exploits:"))
        self.exploitdb_search_input = QLineEdit()
        self.exploitdb_search_input.setPlaceholderText("Enter software name, version, etc. (e.g., 'wordpress 4.7.0')")
        controls_layout.addWidget(self.exploitdb_search_input)
        self.exploitdb_search_button = QPushButton("Search")
        controls_layout.addWidget(self.exploitdb_search_button)
        layout.addLayout(controls_layout)

        # --- Results Display ---
        self.exploitdb_results_table = QTreeWidget()
        self.exploitdb_results_table.setColumnCount(3)
        self.exploitdb_results_table.setHeaderLabels(["ID", "Title", "URL"])
        self.exploitdb_results_table.header().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.exploitdb_results_table.header().setStretchLastSection(True)
        layout.addWidget(self.exploitdb_results_table)

        # --- Connections ---
        save_api_key_btn.clicked.connect(self.save_vulners_api_key)
        self.exploitdb_search_button.clicked.connect(self.start_exploit_search)
        self.exploitdb_results_table.itemDoubleClicked.connect(self.open_exploit_url)

        # Load saved API key on startup
        self.load_vulners_api_key()

        return widget

    def save_vulners_api_key(self):
        """Saves the Vulners API key to a file."""
        api_key = self.exploitdb_api_key_input.text()
        if not api_key:
            QMessageBox.warning(self, "Input Error", "Please enter an API key to save.")
            return
        try:
            with open("vulners_api.key", "w") as f:
                f.write(api_key)
            QMessageBox.information(self, "Success", "Vulners API key saved successfully.")
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not save API key: {e}")

    def load_vulners_api_key(self):
        """Loads the Vulners API key from a file."""
        try:
            with open("vulners_api.key", "r") as f:
                api_key = f.read().strip()
                self.exploitdb_api_key_input.setText(api_key)
                logging.info("Loaded Vulners API key.")
        except FileNotFoundError:
            pass # It's okay if the file doesn't exist yet
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not load API key: {e}")

    def start_exploit_search(self):
        """Starts the Exploit-DB search worker thread."""
        if not shutil.which("getsploit"):
            QMessageBox.critical(self, "GetSploit Error", "'getsploit' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        query = self.exploitdb_search_input.text()
        api_key = self.exploitdb_api_key_input.text()

        if not query:
            QMessageBox.critical(self, "Input Error", "Please provide a search query.")
            return
        if not api_key:
            QMessageBox.critical(self, "API Key Error", "Vulners API key is required for searching exploits.")
            return

        self.is_tool_running = True
        self.exploitdb_search_button.setEnabled(False)
        self.exploitdb_results_table.clear()
        self.status_bar.showMessage(f"Searching for exploits related to '{query}'...")

        self.worker = WorkerThread(self._exploit_search_thread, args=(query, api_key))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _exploit_search_thread(self, query, api_key):
        """Worker thread to search for exploits using getsploit."""
        q = self.tool_results_queue
        command = ["getsploit", "--api", api_key, query]

        try:
            # Use CREATE_NO_WINDOW flag on Windows to hide the console
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            output, _ = process.communicate()

            if process.returncode != 0:
                raise Exception(output)

            # Parse the table-formatted output
            lines = output.strip().split('\n')
            # Find the header line to start parsing from
            header_index = -1
            for i, line in enumerate(lines):
                if 'ID' in line and 'Exploit Title' in line and 'URL' in line:
                    header_index = i
                    break

            if header_index == -1:
                q.put(('exploit_search_status', "No results found or could not parse output."))
                return

            results = []
            # Start from 2 lines after the header to skip the header and the '======' line
            for line in lines[header_index + 2:]:
                if line.startswith('+--'): # End of table
                    break
                parts = [p.strip() for p in line.split('|') if p.strip()]
                if len(parts) >= 3:
                    results.append(parts)

            q.put(('exploit_search_results', results))
            q.put(('exploit_search_status', f"Found {len(results)} exploits."))

        except FileNotFoundError:
            q.put(('error', 'GetSploit Error', "'getsploit' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"getsploit search failed: {e}", exc_info=True)
            q.put(('error', 'Exploit Search Error', str(e)))
        finally:
            q.put(('tool_finished', 'exploit_search'))

    def open_exploit_url(self, item, column):
        """Opens the selected exploit URL in the default web browser."""
        url = item.text(2) # URL is in the 3rd column
        if url and url.startswith("http"):
            webbrowser.open(url)
        else:
            QMessageBox.warning(self, "Invalid URL", f"The selected item does not have a valid URL: {url}")

    def _create_community_tools_tab(self):
        """Creates the UI for the Scapy Community Tools tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        text_browser = QTextBrowser()
        text_browser.setOpenExternalLinks(True)

        html_content = "<h1>Scapy Community Tools and Projects</h1>"
        html_content += "<p>This is a curated list of awesome tools, talks, and projects related to Scapy, inspired by the <a href='https://github.com/gpotter2/awesome-scapy'>awesome-scapy</a> repository.</p>"

        for category, tools in COMMUNITY_TOOLS.items():
            html_content += f"<h2>{category}</h2>"
            html_content += "<ul>"
            for name, url, description in tools:
                html_content += f"<li><b><a href='{url}'>{name}</a></b>: {description}</li>"
            html_content += "</ul>"

        text_browser.setHtml(html_content)
        layout.addWidget(text_browser)
        return widget

    def _create_log_panel(self):
        """Creates the dockable logging panel at the bottom of the window."""
        log_dock_widget = QDockWidget("Live Log", self)
        log_dock_widget.setAllowedAreas(Qt.DockWidgetArea.BottomDockWidgetArea)
        self.log_console = QPlainTextEdit(); self.log_console.setReadOnly(True)
        log_dock_widget.setWidget(self.log_console)
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, log_dock_widget)

    def _setup_logging(self):
        """Configures the logging system to output to a file and the UI panel."""
        root_logger = logging.getLogger()
        for handler in root_logger.handlers[:]: root_logger.removeHandler(handler)
        file_handler = logging.FileHandler('gscapy.log', mode='w')
        formatter = logging.Formatter('%(asctime)s - %(threadName)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        root_logger.addHandler(file_handler)
        qt_handler = QtLogHandler()
        qt_handler.log_updated.connect(self.log_console.appendPlainText)
        qt_handler.setFormatter(formatter)
        root_logger.addHandler(qt_handler)
        root_logger.setLevel(logging.INFO)

    def _create_sniffer_tab(self):
        """Creates the UI for the Packet Sniffer tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Create the results widget first
        self.packet_list_widget = QTreeWidget()
        self.packet_list_widget.setColumnCount(6)
        self.packet_list_widget.setHeaderLabels(["No.", "Time", "Source", "Destination", "Protocol", "Length"])
        # Make header columns resizable and stretch the last section
        header = self.packet_list_widget.header()
        header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        header.setStretchLastSection(True)

        # --- Control Panel ---
        control_panel = QFrame()
        control_panel.setObjectName("controlPanel")
        control_panel.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; }")
        control_layout = QHBoxLayout(control_panel)
        control_layout.setContentsMargins(10, 10, 10, 10)
        control_layout.setSpacing(10)

        self.start_sniff_btn = QPushButton(QIcon("icons/search.svg"), " Start Sniffing")
        self.stop_sniff_btn = QPushButton(QIcon("icons/square.svg"), " Stop Sniffing"); self.stop_sniff_btn.setEnabled(False)
        self.clear_sniff_btn = QPushButton("Clear")
        export_btn = self._create_export_button(self.packet_list_widget)

        control_layout.addWidget(self.start_sniff_btn)
        control_layout.addWidget(self.stop_sniff_btn)
        control_layout.addWidget(self.clear_sniff_btn)
        control_layout.addWidget(export_btn)
        control_layout.addStretch(1)

        control_layout.addWidget(QLabel("BPF Filter:"))
        self.filter_input = QLineEdit()
        self.filter_input.setPlaceholderText("e.g., tcp and port 80")
        control_layout.addWidget(self.filter_input, 2) # Give filter more stretch

        control_layout.addWidget(QLabel("Common:"))
        self.common_filter_combo = QComboBox()
        self.common_filter_combo.addItems(COMMON_FILTERS)
        self.common_filter_combo.textActivated.connect(self.filter_input.setText)
        control_layout.addWidget(self.common_filter_combo)
        layout.addWidget(control_panel)

        # Main splitter for top (list) and bottom (details)
        main_splitter = QSplitter(Qt.Orientation.Vertical)
        main_splitter.addWidget(self.packet_list_widget)

        # Bottom splitter for details tree and hex view
        bottom_splitter = QSplitter(Qt.Orientation.Vertical)

        self.packet_details_tree = QTreeWidget()
        self.packet_details_tree.setHeaderLabels(["Field", "Value"])
        # Make header columns resizable and stretch the last section
        details_header = self.packet_details_tree.header()
        details_header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        details_header.setStretchLastSection(True)
        bottom_splitter.addWidget(self.packet_details_tree)

        self.packet_hex_view = QTextBrowser()
        self.packet_hex_view.setReadOnly(True)
        # Use a monospaced font for the hex view for proper alignment
        self.packet_hex_view.setFont(QFont("Courier New", 10))
        bottom_splitter.addWidget(self.packet_hex_view)

        bottom_splitter.setSizes([200, 100]) # Initial sizes for tree and hex view

        main_splitter.addWidget(bottom_splitter)
        main_splitter.setSizes([400, 300]) # Initial sizes for packet list and details area
        layout.addWidget(main_splitter)

        # Connect signals
        self.start_sniff_btn.clicked.connect(self.start_sniffing)
        self.stop_sniff_btn.clicked.connect(self.stop_sniffing)
        self.clear_sniff_btn.clicked.connect(self.clear_sniffer_display)
        self.packet_list_widget.currentItemChanged.connect(self.display_packet_details)
        return widget

    def _create_crafter_tab(self):
        """Creates the UI for the Packet Crafter tab."""
        widget = QWidget(); main_layout = QVBoxLayout(widget)
        top_splitter = QSplitter(Qt.Orientation.Horizontal); main_layout.addWidget(top_splitter)
        left_panel = QWidget(); left_layout = QVBoxLayout(left_panel); top_splitter.addWidget(left_panel)
        controls_layout = QHBoxLayout()
        self.proto_to_add = QComboBox(); self.proto_to_add.addItems(AVAILABLE_PROTOCOLS.keys())
        add_btn = QPushButton("Add"); remove_btn = QPushButton("Remove");
        controls_layout.addWidget(self.proto_to_add); controls_layout.addWidget(add_btn); controls_layout.addWidget(remove_btn)
        left_layout.addLayout(controls_layout)

        layer_actions_layout = QHBoxLayout()
        fuzz_btn = QPushButton("Fuzz/Unfuzz Selected Layer"); layer_actions_layout.addWidget(fuzz_btn)
        templates_btn = QPushButton("Templates"); layer_actions_layout.addWidget(templates_btn)
        left_layout.addLayout(layer_actions_layout)

        self.layer_list_widget = QListWidget(); left_layout.addWidget(self.layer_list_widget)
        left_layout.addWidget(QLabel("Packet Summary:")); self.crafter_summary = QPlainTextEdit(); self.crafter_summary.setReadOnly(True); left_layout.addWidget(self.crafter_summary)
        right_panel = QWidget(); right_layout = QVBoxLayout(right_panel)
        right_layout.addWidget(QLabel("Layer Fields")); self.scroll_area = QScrollArea(); self.scroll_area.setWidgetResizable(True)
        self.fields_widget = QWidget(); self.fields_layout = QVBoxLayout(self.fields_widget); self.scroll_area.setWidget(self.fields_widget)
        right_layout.addWidget(self.scroll_area); top_splitter.addWidget(right_panel); top_splitter.setSizes([300, 400])
        send_frame = QFrame(); send_frame.setFrameShape(QFrame.Shape.StyledPanel); main_layout.addWidget(send_frame)
        send_layout = QVBoxLayout(send_frame)
        send_controls_layout = QHBoxLayout()
        send_controls_layout.addWidget(QLabel("Count:")); self.send_count_edit = QLineEdit("1"); send_controls_layout.addWidget(self.send_count_edit)
        send_controls_layout.addWidget(QLabel("Interval:")); self.send_interval_edit = QLineEdit("0.1"); send_controls_layout.addWidget(self.send_interval_edit)
        self.send_btn = QPushButton("Send Packet(s)")
        self.send_cancel_btn = QPushButton("Cancel"); self.send_cancel_btn.setEnabled(False)
        send_controls_layout.addWidget(self.send_btn)
        send_controls_layout.addWidget(self.send_cancel_btn)
        send_layout.addLayout(send_controls_layout)
        self.send_results_widget = QTreeWidget(); self.send_results_widget.setColumnCount(3); self.send_results_widget.setHeaderLabels(["No.", "Sent", "Received"])
        send_layout.addWidget(self.send_results_widget)
        send_layout.addWidget(self._create_export_button(self.send_results_widget))
        add_btn.clicked.connect(self.crafter_add_layer); remove_btn.clicked.connect(self.crafter_remove_layer)
        self.layer_list_widget.currentRowChanged.connect(self.crafter_display_layer_fields)
        templates_menu = QMenu(self)
        for name in PACKET_TEMPLATES.keys():
            action = QAction(name, self); action.triggered.connect(lambda checked, n=name: self.crafter_load_template(n)); templates_menu.addAction(action)
        templates_btn.setMenu(templates_menu)
        self.send_btn.clicked.connect(self.crafter_send_packet)
        self.send_cancel_btn.clicked.connect(self.cancel_tool)
        fuzz_btn.clicked.connect(self.crafter_toggle_fuzz_layer)
        return widget

    def _create_nmap_scanner_tool(self):
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Configurable options widget ---
        config_widget, self.nmap_controls = self._create_nmap_config_widget()
        main_layout.addWidget(config_widget)

        # Connect signals now that we have the controls dictionary
        controls = self.nmap_controls
        controls['start_btn'].clicked.connect(self.start_nmap_scan)
        controls['cancel_btn'].clicked.connect(self.cancel_tool)
        controls['report_btn'].clicked.connect(self.generate_nmap_report)
        controls['all_ports_btn'].clicked.connect(self._nmap_set_all_ports)
        controls['super_complete_btn'].clicked.connect(self._nmap_toggle_super_complete)
        controls['preset_combo'].textActivated.connect(self._handle_nmap_preset_selected)
        controls['a_check'].toggled.connect(self._nmap_on_aggressive_toggled)
        controls['scan_type_combo'].currentTextChanged.connect(self._nmap_on_ping_scan_toggled)

        # --- Output Console ---
        self.nmap_output_console = QPlainTextEdit()
        self.nmap_output_console.setReadOnly(True)
        self.nmap_output_console.setFont(QFont("Courier New", 10))
        self.nmap_output_console.setPlaceholderText("Nmap command output will be displayed here...")
        main_layout.addWidget(self.nmap_output_console, 1) # Give it stretch factor

        return widget

    def _create_nmap_config_widget(self):
        """Creates a reusable, self-contained widget with all of Nmap's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)

        controls = {}

        top_controls = QFrame()
        top_controls.setObjectName("controlPanel")
        top_controls.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        top_layout = QGridLayout(top_controls)

        top_layout.addWidget(QLabel("Target(s):"), 0, 0)
        controls['target_edit'] = QLineEdit("localhost"); controls['target_edit'].setToolTip("Enter one or more target hosts, separated by spaces.\nExamples:\n- scanme.nmap.org (hostname)\n- 192.168.1.1 (single IP)\n- 192.168.1.0/24 (CIDR block)\n- 10.0.0-5.1-254 (IP range)")
        top_layout.addWidget(controls['target_edit'], 0, 1, 1, 3)

        top_layout.addWidget(QLabel("Ports:"), 1, 0)
        controls['ports_edit'] = QLineEdit(); controls['ports_edit'].setToolTip("Specify ports to scan.\nExamples:\n- 22,80,443 (comma-separated)\n- 1-1024 (range)\n- U:53,T:21-25,80 (specific protocols)\nLeave blank for Nmap's default (top 1000 TCP ports).")
        controls['ports_edit'].setPlaceholderText("Default (top 1000)")
        top_layout.addWidget(controls['ports_edit'], 1, 1, 1, 3)

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['cancel_btn'] = QPushButton("Cancel"); controls['cancel_btn'].setEnabled(False)
        controls['report_btn'] = QPushButton("Generate HTML Report"); controls['report_btn'].setEnabled(False)

        top_layout.addWidget(controls['start_btn'], 0, 4)
        top_layout.addWidget(controls['cancel_btn'], 1, 4)
        top_layout.addWidget(controls['report_btn'], 0, 5, 2, 1)

        presets_layout = QVBoxLayout()
        controls['all_ports_btn'] = QPushButton("All Ports"); controls['all_ports_btn'].setToolTip("Set the port range to all 65535 ports.")
        controls['super_complete_btn'] = QPushButton("Super Complete Scan"); controls['super_complete_btn'].setToolTip("Set options for a highly comprehensive scan (-sS -A -v -T4 -p 1-65535).")
        presets_layout.addWidget(controls['all_ports_btn'])
        presets_layout.addWidget(controls['super_complete_btn'])
        top_layout.addLayout(presets_layout, 0, 6, 2, 1)
        main_layout.addWidget(top_controls)

        options_splitter = QSplitter(Qt.Orientation.Horizontal)
        left_panel = QWidget(); left_layout = QVBoxLayout(left_panel)

        controls['scan_type_box'] = QGroupBox("Scan Type")
        scan_type_layout = QFormLayout(controls['scan_type_box'])
        controls['scan_type_combo'] = QComboBox(); controls['scan_type_combo'].addItems(["SYN Stealth Scan (-sS)", "TCP Connect Scan (-sT)", "UDP Scan (-sU)", "FIN Scan (-sF)", "Xmas Scan (-sX)", "Null Scan (-sN)", "Ping Scan (-sn)"])
        controls['scan_type_combo'].setToolTip("Select the Nmap scan type.\n- SYN Stealth Scan (-sS): Default & most popular. Requires root/admin.\n- TCP Connect Scan (-sT): More reliable, but easily logged. No root needed.\n- UDP Scan (-sU): Scans for open UDP ports.\n- FIN/Xmas/Null: Stealthy scans useful for firewall evasion.\n- Ping Scan (-sn): Disables port scanning. Only discovers if hosts are online.")
        scan_type_layout.addRow(controls['scan_type_combo'])
        left_layout.addWidget(controls['scan_type_box'])

        controls['timing_box'] = QGroupBox("Timing Template")
        timing_layout = QFormLayout(controls['timing_box'])
        controls['timing_combo'] = QComboBox(); controls['timing_combo'].setToolTip("Adjusts timing parameters to be more or less aggressive.\n- T0 (Paranoid): Very slow, for IDS evasion.\n- T1 (Sneaky): Quite slow, for IDS evasion.\n- T2 (Polite): Slows down to consume less bandwidth.\n- T3 (Normal): Default speed.\n- T4 (Aggressive): Assumes a fast and reliable network.\n- T5 (Insane): Extremely aggressive; may sacrifice accuracy.")
        controls['timing_combo'].addItems(["T0 (Paranoid)", "T1 (Sneaky)", "T2 (Polite)", "T3 (Normal)", "T4 (Aggressive)", "T5 (Insane)"])
        controls['timing_combo'].setCurrentIndex(3)
        timing_layout.addRow(controls['timing_combo'])
        left_layout.addWidget(controls['timing_box'])
        left_layout.addStretch()

        right_panel = QWidget(); right_layout = QVBoxLayout(right_panel)
        controls['detection_box'] = QGroupBox("Detection Options")
        detection_grid = QGridLayout(controls['detection_box'])
        controls['sv_check'] = QCheckBox("Service/Version (-sV)"); controls['sv_check'].setToolTip("Probe open ports to determine service and version info (-sV).\nThis is essential for vulnerability scanning and accurate service identification.")
        controls['o_check'] = QCheckBox("OS Detection (-O)"); controls['o_check'].setToolTip("Enable OS detection (-O).\nAttempts to determine the operating system of the target.\nRequires root/administrator privileges.")
        controls['sc_check'] = QCheckBox("Default Scripts (-sC)"); controls['sc_check'].setToolTip("Run a scan using the default set of safe Nmap scripts (-sC).\nEquivalent to --script=default.")
        controls['a_check'] = QCheckBox("Aggressive Scan (-A)"); controls['a_check'].setToolTip("Enable Aggressive scan options (-A).\nThis is a convenient shortcut that enables OS detection (-O), version scanning (-sV), script scanning (-sC), and traceroute (--traceroute).\nIt is a comprehensive but noisy scan.")
        detection_grid.addWidget(controls['sv_check'], 0, 0); detection_grid.addWidget(controls['o_check'], 0, 1)
        detection_grid.addWidget(controls['sc_check'], 1, 0); detection_grid.addWidget(controls['a_check'], 1, 1)
        right_layout.addWidget(controls['detection_box'])

        controls['misc_box'] = QGroupBox("Miscellaneous Options")
        misc_grid = QGridLayout(controls['misc_box'])
        controls['pn_check'] = QCheckBox("No Ping (-Pn)"); controls['pn_check'].setToolTip("Treat all hosts as online (-Pn).\nThis skips the initial host discovery (ping sweep) phase, which is useful for scanning hosts that do not respond to ping probes or are on a firewalled network.")
        controls['v_check'] = QCheckBox("Verbose (-v)"); controls['v_check'].setToolTip("Increase verbosity level (-v).\nProvides more information about the scan in progress.\nCan be used multiple times for more detail (e.g., -vv).")
        controls['traceroute_check'] = QCheckBox("Traceroute (--traceroute)"); controls['traceroute_check'].setToolTip("Trace the network path (hops) to each host.\nUseful for mapping the network route to the target.")
        misc_grid.addWidget(controls['pn_check'], 0, 0); misc_grid.addWidget(controls['v_check'], 0, 1)
        misc_grid.addWidget(controls['traceroute_check'], 1, 0)
        right_layout.addWidget(controls['misc_box'])
        options_splitter.addWidget(left_panel); options_splitter.addWidget(right_panel)
        main_layout.addWidget(options_splitter)

        controls['nse_box'] = QGroupBox("Nmap Scripting Engine (NSE)")
        nse_layout = QFormLayout(controls['nse_box'])
        controls['preset_combo'] = QComboBox(); controls['preset_combo'].setToolTip("Select a common script preset to automatically fill the script fields below.")
        controls['preset_combo'].addItems(["-- Select a Preset --"] + list(self.nmap_script_presets.keys()))
        nse_layout.addRow("Script Presets:", controls['preset_combo'])

        controls['nse_vuln_check'] = QCheckBox("Vulnerability scripts (`vuln`)"); controls['nse_vuln_check'].setToolTip("Run all scripts in the 'vuln' category.\nThese check for specific known vulnerabilities.\nRequires service version detection (-sV).")
        controls['nse_discovery_check'] = QCheckBox("Discovery scripts (`discovery`)"); controls['nse_discovery_check'].setToolTip("Run all scripts in the 'discovery' category.\nThese actively probe for more information, such as registry details or service info.")
        controls['nse_safe_check'] = QCheckBox("Safe scripts (`safe`)"); controls['nse_safe_check'].setToolTip("Run all scripts in the 'safe' category.\nThese are scripts that are not considered intrusive or likely to crash services.")
        category_layout = QHBoxLayout(); category_layout.addWidget(controls['nse_vuln_check']); category_layout.addWidget(controls['nse_discovery_check']); category_layout.addWidget(controls['nse_safe_check'])
        nse_layout.addRow("Categories:", category_layout)

        controls['custom_script_edit'] = QLineEdit(); controls['custom_script_edit'].setPlaceholderText("e.g., http-title,smb-os-discovery"); controls['custom_script_edit'].setToolTip("A comma-separated list of Nmap scripts, directories, or categories to run.\nOverrides category checkboxes if specified.")
        nse_layout.addRow("Custom Scripts:", controls['custom_script_edit'])
        controls['script_args_edit'] = QLineEdit(); controls['script_args_edit'].setPlaceholderText("e.g., http.useragent=MyCustomAgent,user=admin"); controls['script_args_edit'].setToolTip("Provide arguments for your NSE scripts (e.g., http.useragent=MyCustomAgent,user=admin).\nSee Nmap documentation for script-specific arguments.")
        nse_layout.addRow("Script Arguments:", controls['script_args_edit'])
        controls['script_desc_label'] = QLabel("Description: --"); controls['script_desc_label'].setWordWrap(True); controls['script_desc_label'].setStyleSheet("color: #aaa; padding-top: 5px;")
        nse_layout.addRow(controls['script_desc_label'])
        main_layout.addWidget(controls['nse_box'])

        return widget, controls

    def _nmap_on_aggressive_toggled(self, checked):
        controls = self.nmap_controls
        controls['sv_check'].setDisabled(checked)
        controls['o_check'].setDisabled(checked)
        controls['sc_check'].setDisabled(checked)
        controls['traceroute_check'].setDisabled(checked)

    def _nmap_on_ping_scan_toggled(self, text):
        controls = self.nmap_controls
        is_ping_scan = (text == "Ping Scan (-sn)")
        for w_key in ['detection_box', 'misc_box', 'ports_edit', 'timing_box', 'nse_box']:
             controls[w_key].setDisabled(is_ping_scan)

    def _nmap_set_all_ports(self):
        """Sets the Nmap port text field to scan all ports."""
        self.nmap_controls['ports_edit'].setText("1-65535")

    def _handle_nmap_preset_selected(self, preset_name):
        """Populates the script fields based on the selected Nmap preset."""
        controls = self.nmap_controls
        if preset_name == "-- Select a Preset --":
            controls['custom_script_edit'].clear()
            controls['script_args_edit'].clear()
            controls['script_desc_label'].setText("Description: --")
            return

        scripts, args, desc = self.nmap_script_presets.get(preset_name, ("", "", "No description available."))
        controls['custom_script_edit'].setText(scripts)
        controls['script_args_edit'].setText(args)
        controls['script_desc_label'].setText(f"Description: {desc}")

    def _nmap_toggle_super_complete(self):
        """Toggles the 'Super Complete Scan' preset."""
        controls = self.nmap_controls
        if not self.super_scan_active:
            controls['ports_edit'].setText("1-65535")
            controls['scan_type_combo'].setCurrentText("SYN Stealth Scan (-sS)")
            controls['timing_combo'].setCurrentText("T4 (Aggressive)")
            controls['a_check'].setChecked(True)
            controls['v_check'].setChecked(True)

            target = controls['target_edit'].text() or "[target]"
            command_preview = self._build_nmap_command_preview(target)
            self.nmap_output_console.clear()
            self.nmap_output_console.setPlainText(f"# Preset command preview:\n$ {command_preview}")
            QMessageBox.information(self, "Preset Loaded", "Super Complete Scan options have been set.\nClick 'Start Scan' to run, or click the preset button again to cancel.")

            controls['super_complete_btn'].setText("Cancel Super Scan")
            self.super_scan_active = True
        else:
            controls['ports_edit'].setText("")
            controls['scan_type_combo'].setCurrentIndex(0)
            controls['timing_combo'].setCurrentIndex(3)
            controls['a_check'].setChecked(False)
            controls['v_check'].setChecked(False)

            controls['super_complete_btn'].setText("Super Complete Scan")
            self.nmap_output_console.clear()
            self.super_scan_active = False

    def _build_nmap_script_args(self):
        """Builds the --script and --script-args parts of the nmap command."""
        controls = self.nmap_controls
        script_parts = []

        if controls['sc_check'].isChecked():
            script_parts.append("default")
        if controls['nse_vuln_check'].isChecked():
            script_parts.append("vuln")
        if controls['nse_discovery_check'].isChecked():
            script_parts.append("discovery")
        if controls['nse_safe_check'].isChecked():
            script_parts.append("safe")
        if custom_scripts := controls['custom_script_edit'].text().strip():
            script_parts.append(custom_scripts)

        command_args = []
        if script_parts:
            unique_scripts = sorted(list(set(script_parts)))
            command_args.extend(["--script", ",".join(unique_scripts)])
        if script_args := controls['script_args_edit'].text().strip():
            command_args.extend(["--script-args", script_args])

        return command_args

    def _build_nmap_command_preview(self, target):
        """Helper to build a command string for preview purposes."""
        controls = self.nmap_controls
        command = ["nmap"]
        command.append(controls['scan_type_combo'].currentText().split(" ")[-1].strip("()"))
        command.append("-T" + controls['timing_combo'].currentText()[1])
        if controls['a_check'].isChecked(): command.append("-A")
        if controls['v_check'].isChecked(): command.append("-v")
        if controls['pn_check'].isChecked(): command.append("-Pn")
        if ports := controls['ports_edit'].text():
            command.extend(["-p", ports])
        command.extend(self._build_nmap_script_args())
        command.append(target)
        return " ".join(command)

    def start_nmap_scan(self):
        """Starts the Nmap scan worker thread by building a command from the UI."""
        controls = self.nmap_controls
        if not shutil.which("nmap"):
            QMessageBox.critical(self, "Nmap Error", "'nmap' command not found. Please ensure it is installed and in your system's PATH.")
            return
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return
        target = controls['target_edit'].text()
        if not target:
            QMessageBox.critical(self, "Input Error", "Please provide a target for the Nmap scan.")
            return

        command = ["nmap"]
        command.append(controls['scan_type_combo'].currentText().split(" ")[-1].strip("()"))
        command.append("-T" + controls['timing_combo'].currentText()[1])

        if controls['a_check'].isChecked():
            command.append("-A")
        else:
            if controls['sv_check'].isChecked(): command.append("-sV")
            if controls['o_check'].isChecked(): command.append("-O")
            if controls['traceroute_check'].isChecked(): command.append("--traceroute")

        if controls['pn_check'].isChecked(): command.append("-Pn")
        if controls['v_check'].isChecked(): command.append("-v")

        if controls['ports_edit'].isEnabled():
            if ports := controls['ports_edit'].text():
                command.extend(["-p", ports])
            else:
                command.extend(["--top-ports", "1024"])

        command.extend(self._build_nmap_script_args())

        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".xml", encoding='utf-8') as tmp_xml:
                self.nmap_xml_temp_file = tmp_xml.name
            command.extend(["-oX", self.nmap_xml_temp_file])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary file for Nmap report: {e}")
            return

        command.append(target)

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        controls['report_btn'].setEnabled(False)
        self.tool_stop_event.clear()
        self.nmap_output_console.clear()

        self.worker = WorkerThread(self._nmap_scan_thread, args=(command, target))
        self.active_threads.append(self.worker)
        self.worker.start()

    def generate_nmap_report(self):
        """Saves the Nmap XML and generates a styled HTML report using lxml."""
        if not self.nmap_last_xml:
            QMessageBox.information(self, "No Data", "Please run an Nmap scan first to generate data for the report.")
            return

        if not LXML_AVAILABLE:
            QMessageBox.critical(self, "Dependency Error", "The 'lxml' library is required for HTML report generation. Please install it using 'pip install lxml'.")
            return

        save_path, _ = QFileDialog.getSaveFileName(self, "Save Nmap HTML Report", "nmap_report.html", "HTML Files (*.html);;XML Files (*.xml)", options=QFileDialog.Option.DontUseNativeDialog)
        if not save_path:
            return

        try:
            # Always save the raw XML data first
            if save_path.endswith('.html'):
                xml_path = os.path.splitext(save_path)[0] + ".xml"
            else:
                xml_path = save_path

            with open(xml_path, 'w', encoding='utf-8') as f:
                f.write(self.nmap_last_xml)

            # If the user wants HTML, perform the transformation
            if save_path.endswith('.html'):
                # Use a parser that can recover from errors, which can happen with Nmap's XML
                parser = etree.XMLParser(recover=True)
                xml_doc = etree.fromstring(self.nmap_last_xml.encode('utf-8'), parser=parser)

                # Check for the stylesheet file
                xsl_path = "nmap-bootstrap.xsl"
                if not os.path.exists(xsl_path):
                    QMessageBox.critical(self, "File Not Found", f"Stylesheet '{xsl_path}' not found. Make sure it is in the same directory as the application.")
                    return

                xsl_doc = etree.parse(xsl_path)
                transform = etree.XSLT(xsl_doc)
                html_doc = transform(xml_doc)

                with open(save_path, 'wb') as f:
                    f.write(etree.tostring(html_doc, pretty_print=True))

            QMessageBox.information(self, "Report Saved", f"Report successfully saved to:\n{os.path.realpath(save_path)}")

        except Exception as e:
            logging.error(f"Failed to generate or save Nmap report: {e}", exc_info=True)
            QMessageBox.critical(self, "Report Generation Error", f"An unexpected error occurred:\n{e}")

    def _nmap_scan_thread(self, command, target):
        q = self.tool_results_queue
        logging.info(f"Starting Nmap scan with command: {' '.join(command)}")
        q.put(('nmap_output', f"$ {' '.join(command)}\n\n"))
        xml_content = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.nmap_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate() # Terminate the process if cancelled
                    q.put(('nmap_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('nmap_output', line))

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Nmap Error', "'nmap' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            q.put(('error', 'Nmap Error', str(e)))
        finally:
            # After scan, read the XML report from the temp file
            if self.nmap_xml_temp_file and os.path.exists(self.nmap_xml_temp_file):
                try:
                    with open(self.nmap_xml_temp_file, 'r', encoding='utf-8') as f:
                        xml_content = f.read()
                    if xml_content:
                        q.put(('nmap_xml_result', xml_content))
                except Exception as e:
                    logging.error(f"Could not read Nmap XML report: {e}")
                finally:
                    os.remove(self.nmap_xml_temp_file)
                    self.nmap_xml_temp_file = None

            q.put(('tool_finished', 'nmap_scan', target, xml_content))
            with self.thread_finish_lock:
                self.nmap_process = None
            logging.info("Nmap scan thread finished.")

    def _sublist3r_thread(self, domain):
        """Worker thread to run the Sublist3r script."""
        q = self.tool_results_queue
        command = ["python", "tools/sublist3r/sublist3r.py", "-d", domain]
        logging.info(f"Starting Sublist3r scan with command: {' '.join(command)}")
        q.put(('sublist3r_output', f"$ {' '.join(command)}\n\n"))

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')
            with self.thread_finish_lock:
                self.sublist3r_process = process

            full_output = []
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('sublist3r_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('sublist3r_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

            # If scan was not canceled, parse the results and show the popup
            if not self.tool_stop_event.is_set():
                results = []
                try:
                    # New method: Find the last non-empty line which should contain the JSON array
                    json_line = ""
                    for line in reversed(full_output):
                        stripped_line = line.strip()
                        if stripped_line:
                            json_line = stripped_line
                            break

                    if json_line.startswith('[') and json_line.endswith(']'):
                        results = json.loads(json_line)
                        logging.info(f"Successfully parsed {len(results)} subdomains from sublist3r JSON output.")
                    else:
                        # This will trigger the fallback logic
                        raise ValueError("Could not find JSON list in output.")

                except (json.JSONDecodeError, IndexError, ValueError) as e:
                    logging.warning(f"Could not parse JSON from sublist3r output ({e}), falling back to fragile text parsing.")
                    # Fallback to old, fragile parsing method
                    for line in reversed(full_output):
                        if "Total Unique Subdomains Found" in line:
                            break # Stop when we hit the summary line
                        # A simple check to see if the line is likely a subdomain
                        if f'.{domain}' in line and not any(c in '<> ' for c in line):
                             results.append(line.strip())
                    results.reverse()

                q.put(('subdomain_results', domain, results))

        except FileNotFoundError:
            q.put(('error', 'Sublist3r Error', "'python' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            q.put(('error', 'Sublist3r Error', str(e)))
        finally:
            results_str = "\n".join(results) if 'results' in locals() else ""
            q.put(('tool_finished', 'sublist3r_scan', domain, results_str))
            with self.thread_finish_lock:
                self.sublist3r_process = None
            logging.info("Sublist3r scan thread finished.")

    def _create_subfinder_tool(self):
        """Creates the UI for the Subfinder tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.subfinder_controls = self._create_subfinder_config_widget()
        layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.subfinder_controls['start_btn'])
        buttons_layout.addWidget(self.subfinder_controls['cancel_btn'])
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.subfinder_output = QPlainTextEdit()
        self.subfinder_output.setReadOnly(True)
        self.subfinder_output.setFont(QFont("Courier New", 10))
        self.subfinder_output.setPlaceholderText("Subfinder output will be displayed here...")
        layout.addWidget(self.subfinder_output, 1)

        self.subfinder_controls['start_btn'].clicked.connect(self.start_subfinder_scan)
        self.subfinder_controls['cancel_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_subfinder_config_widget(self):
        """Creates a reusable, self-contained widget for the Subfinder scanner's configuration."""
        widget = QFrame()
        widget.setObjectName("controlPanel")
        widget.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        layout = QFormLayout(widget)

        controls = {}

        controls['domain_edit'] = QLineEdit("example.com")
        controls['domain_edit'].setToolTip("Enter the target domain to enumerate subdomains for (-d).")
        layout.addRow("Domain:", controls['domain_edit'])

        controls['recursive_check'] = QCheckBox("Recursive Scan")
        controls['recursive_check'].setToolTip("Enable recursive subdomain discovery (-recursive).")
        layout.addRow(controls['recursive_check'])

        controls['all_sources_check'] = QCheckBox("Use All Sources")
        controls['all_sources_check'].setToolTip("Use all available sources for enumeration (-all).")
        layout.addRow(controls['all_sources_check'])

        output_file_layout = QHBoxLayout()
        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("Optional: path to save results...")
        output_file_layout.addWidget(controls['output_edit'])
        browse_output_btn = QPushButton("Browse...")
        browse_output_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_edit'], "Save Subfinder Results"))
        output_file_layout.addWidget(browse_output_btn)
        layout.addRow("Output File (-o):", output_file_layout)

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['cancel_btn'] = QPushButton("Cancel")
        controls['cancel_btn'].setEnabled(False)

        return widget, controls

    def start_subfinder_scan(self):
        """Starts the Subfinder scan worker thread."""
        controls = self.subfinder_controls
        if not shutil.which("subfinder"):
            QMessageBox.critical(self, "Subfinder Error", "'subfinder' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        domain = controls['domain_edit'].text().strip()
        if not domain:
            QMessageBox.critical(self, "Input Error", "Please provide a domain to scan.")
            return

        command = ["subfinder", "-d", domain, "-silent"]

        if controls['recursive_check'].isChecked():
            command.append("-recursive")

        if controls['all_sources_check'].isChecked():
            command.append("-all")

        if output_file := controls['output_edit'].text().strip():
            command.extend(["-o", output_file])

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.subfinder_output.clear()

        self.worker = WorkerThread(self._subfinder_thread, args=(command, domain))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _subfinder_thread(self, command, domain):
        """Worker thread for running the subfinder command."""
        q = self.tool_results_queue
        logging.info(f"Starting Subfinder with command: {' '.join(command)}")
        q.put(('subfinder_output', f"$ {' '.join(command)}\n\n"))
        results = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.subfinder_process = process

            full_output = []
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('subfinder_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('subfinder_output', line))
                full_output.append(line.strip())

            process.stdout.close()
            process.wait()

            if not self.tool_stop_event.is_set():
                results = [line for line in full_output if domain in line and ' ' not in line and not line.startswith('$')]
                q.put(('subdomain_results', domain, results))

        except FileNotFoundError:
            q.put(('error', 'Subfinder Error', "'subfinder' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Subfinder thread error: {e}", exc_info=True)
            q.put(('error', 'Subfinder Error', str(e)))
        finally:
            q.put(('tool_finished', 'subfinder_scan', domain, "\n".join(results)))
            with self.thread_finish_lock:
                self.subfinder_process = None
            logging.info("Subfinder scan thread finished.")

    def _handle_subfinder_output(self, line):
        self.subfinder_output.insertPlainText(line)
        self.subfinder_output.verticalScrollBar().setValue(self.subfinder_output.verticalScrollBar().maximum())

    def _create_httpx_tool(self):
        """Creates the UI for the httpx tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.httpx_controls = self._create_httpx_config_widget()
        layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.httpx_controls['start_btn'])
        buttons_layout.addWidget(self.httpx_controls['cancel_btn'])
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.httpx_output = QPlainTextEdit()
        self.httpx_output.setReadOnly(True)
        self.httpx_output.setFont(QFont("Courier New", 10))
        self.httpx_output.setPlaceholderText("httpx output will be displayed here...")
        layout.addWidget(self.httpx_output, 1)

        self.httpx_controls['start_btn'].clicked.connect(self.start_httpx_scan)
        self.httpx_controls['cancel_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_httpx_config_widget(self):
        """Creates a reusable, self-contained widget for the httpx scanner's configuration."""
        widget = QFrame()
        widget.setObjectName("controlPanel")
        widget.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        main_layout = QVBoxLayout(widget)
        controls = {}

        # --- Top Row: Target List ---
        top_layout = QFormLayout()
        target_file_layout = QHBoxLayout()
        controls['target_list_edit'] = QLineEdit()
        controls['target_list_edit'].setPlaceholderText("Path to a file with hosts/URLs (one per line)...")
        controls['target_list_edit'].setToolTip("A file containing a list of targets to probe (-l).")
        target_file_layout.addWidget(controls['target_list_edit'])
        browse_target_btn = QPushButton("Browse...")
        browse_target_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['target_list_edit'], "Select Target List File"))
        target_file_layout.addWidget(browse_target_btn)
        top_layout.addRow("Target List (-l):", target_file_layout)
        main_layout.addLayout(top_layout)

        # --- Middle Row: Probes ---
        probes_box = QGroupBox("Probes to Run")
        probes_layout = QGridLayout(probes_box)
        controls['probe_status_code'] = QCheckBox("Status Code"); controls['probe_status_code'].setChecked(True)
        controls['probe_title'] = QCheckBox("Title"); controls['probe_title'].setChecked(True)
        controls['probe_tech_detect'] = QCheckBox("Tech Detect")
        controls['probe_web_server'] = QCheckBox("Web Server")
        controls['probe_cdn'] = QCheckBox("CDN")
        controls['probe_jarm'] = QCheckBox("JARM Hash")
        probes_layout.addWidget(controls['probe_status_code'], 0, 0)
        probes_layout.addWidget(controls['probe_title'], 0, 1)
        probes_layout.addWidget(controls['probe_tech_detect'], 0, 2)
        probes_layout.addWidget(controls['probe_web_server'], 1, 0)
        probes_layout.addWidget(controls['probe_cdn'], 1, 1)
        probes_layout.addWidget(controls['probe_jarm'], 1, 2)
        main_layout.addWidget(probes_box)

        # --- Bottom Row: Other Options ---
        bottom_layout = QFormLayout()
        controls['ports_edit'] = QLineEdit()
        controls['ports_edit'].setPlaceholderText("e.g., 80,443,8080 (optional)")
        controls['ports_edit'].setToolTip("Comma-separated list of ports to scan (-ports).")
        bottom_layout.addRow("Ports (-ports):", controls['ports_edit'])

        controls['json_output_check'] = QCheckBox("JSON Output (-json)")
        controls['json_output_check'].setToolTip("Output results in JSON format.")
        bottom_layout.addRow(controls['json_output_check'])

        main_layout.addLayout(bottom_layout)

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Probing")
        controls['cancel_btn'] = QPushButton("Cancel"); controls['cancel_btn'].setEnabled(False)

        return widget, controls

    def start_httpx_scan(self):
        """Starts the httpx scan worker thread."""
        controls = self.httpx_controls
        if not shutil.which("httpx"):
            QMessageBox.critical(self, "httpx Error", "'httpx' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target_list = controls['target_list_edit'].text().strip()
        if not target_list:
            QMessageBox.critical(self, "Input Error", "Please provide a target list file.")
            return

        command = ["httpx", "-l", target_list, "-silent"]

        probe_flags = {
            'probe_status_code': '-status-code',
            'probe_title': '-title',
            'probe_tech_detect': '-tech-detect',
            'probe_web_server': '-web-server',
            'probe_cdn': '-cdn',
            'probe_jarm': '-jarm',
        }
        for control_name, flag in probe_flags.items():
            if controls[control_name].isChecked():
                command.append(flag)

        if ports := controls['ports_edit'].text().strip():
            command.extend(["-ports", ports])

        if controls['json_output_check'].isChecked():
            command.append("-json")

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.httpx_output.clear()

        self.worker = WorkerThread(self._httpx_thread, args=(command, target_list))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _httpx_thread(self, command, target_list):
        """Worker thread for running the httpx command."""
        q = self.tool_results_queue
        logging.info(f"Starting httpx with command: {' '.join(command)}")
        q.put(('httpx_output', f"$ {' '.join(command)}\n\n"))
        json_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.httpx_process = process

            full_output = []
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('httpx_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('httpx_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

            if not self.tool_stop_event.is_set() and "-json" in command:
                json_data = "".join(full_output)
                q.put(('httpx_results', json_data))

        except FileNotFoundError:
            q.put(('error', 'httpx Error', "'httpx' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"httpx thread error: {e}", exc_info=True)
            q.put(('error', 'httpx Error', str(e)))
        finally:
            q.put(('tool_finished', 'httpx_scan', target_list, json_data))
            with self.thread_finish_lock:
                self.httpx_process = None
            logging.info("httpx scan thread finished.")

    def _handle_httpx_output(self, line):
        self.httpx_output.insertPlainText(line)
        self.httpx_output.verticalScrollBar().setValue(self.httpx_output.verticalScrollBar().maximum())

    def _create_rustscan_tool(self):
        """Creates the UI for the RustScan tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.rustscan_controls = self._create_rustscan_config_widget()
        layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.rustscan_controls['start_btn'])
        buttons_layout.addWidget(self.rustscan_controls['cancel_btn'])
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.rustscan_output = QPlainTextEdit()
        self.rustscan_output.setReadOnly(True)
        self.rustscan_output.setFont(QFont("Courier New", 10))
        self.rustscan_output.setPlaceholderText("RustScan and Nmap output will be displayed here...")
        layout.addWidget(self.rustscan_output, 1)

        self.rustscan_controls['start_btn'].clicked.connect(self.start_rustscan_scan)
        self.rustscan_controls['cancel_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_rustscan_config_widget(self):
        """Creates a reusable, self-contained widget for the RustScan tool's configuration."""
        widget = QFrame()
        widget.setObjectName("controlPanel")
        widget.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        layout = QFormLayout(widget)

        controls = {}

        controls['targets_edit'] = QLineEdit("localhost")
        controls['targets_edit'].setToolTip("A single target or a comma-separated list of targets (e.g., 127.0.0.1, scanme.nmap.org).")
        layout.addRow("Targets (-a):", controls['targets_edit'])

        controls['ports_edit'] = QLineEdit("1-1000")
        controls['ports_edit'].setToolTip("A comma-separated list of ports or a range (e.g., 22,80,443 or 1-1024).")
        layout.addRow("Ports (-p):", controls['ports_edit'])

        controls['batch_size_edit'] = QLineEdit("4500")
        controls['batch_size_edit'].setToolTip("The number of ports to scan at once (-b).")
        layout.addRow("Batch Size (-b):", controls['batch_size_edit'])

        controls['timeout_edit'] = QLineEdit("1500")
        controls['timeout_edit'].setToolTip("The timeout in milliseconds for each port (-T).")
        layout.addRow("Timeout (-T):", controls['timeout_edit'])

        controls['nmap_args_edit'] = QLineEdit("-sV -sC -A")
        controls['nmap_args_edit'].setToolTip("Arguments to pass to Nmap after the port scan (e.g., -sV -A).")
        layout.addRow("Nmap Args:", controls['nmap_args_edit'])

        controls['quiet_check'] = QCheckBox("Quiet Mode (No Nmap)")
        controls['quiet_check'].setToolTip("Only output open ports, do not run Nmap (-q).")
        layout.addRow(controls['quiet_check'])

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['cancel_btn'] = QPushButton("Cancel"); controls['cancel_btn'].setEnabled(False)

        return widget, controls

    def start_rustscan_scan(self):
        """Starts the RustScan worker thread."""
        controls = self.rustscan_controls
        if not shutil.which("rustscan"):
            QMessageBox.critical(self, "RustScan Error", "'rustscan' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        targets = controls['targets_edit'].text().strip()
        if not targets:
            QMessageBox.critical(self, "Input Error", "Please provide at least one target.")
            return

        command = ["rustscan", "-a", targets]

        if ports := controls['ports_edit'].text().strip():
            command.extend(["-p", ports])

        if batch_size := controls['batch_size_edit'].text().strip():
            command.extend(["-b", batch_size])

        if timeout := controls['timeout_edit'].text().strip():
            command.extend(["-T", timeout])

        if controls['quiet_check'].isChecked():
            command.append("-q")
        else:
            if nmap_args := controls['nmap_args_edit'].text().strip():
                command.append("--")
                command.extend(nmap_args.split())

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.rustscan_output.clear()

        self.worker = WorkerThread(self._rustscan_thread, args=(command, targets))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _rustscan_thread(self, command, targets):
        """Worker thread for running the rustscan command."""
        q = self.tool_results_queue
        logging.info(f"Starting RustScan with command: {' '.join(command)}")
        q.put(('rustscan_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.rustscan_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('rustscan_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('rustscan_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'RustScan Error', "'rustscan' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"RustScan thread error: {e}", exc_info=True)
            q.put(('error', 'RustScan Error', str(e)))
        finally:
            q.put(('tool_finished', 'rustscan_scan', targets, "".join(full_output)))
            with self.thread_finish_lock:
                self.rustscan_process = None
            logging.info("RustScan scan thread finished.")

    def _handle_rustscan_output(self, line):
        self.rustscan_output.insertPlainText(line)
        self.rustscan_output.verticalScrollBar().setValue(self.rustscan_output.verticalScrollBar().maximum())

    def _create_dirsearch_tool(self):
        """Creates the UI for the dirsearch tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.dirsearch_controls = self._create_dirsearch_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.dirsearch_controls['start_btn'])
        buttons_layout.addWidget(self.dirsearch_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.dirsearch_output_console = QPlainTextEdit()
        self.dirsearch_output_console.setReadOnly(True)
        self.dirsearch_output_console.setFont(QFont("Courier New", 10))
        self.dirsearch_output_console.setPlaceholderText("dirsearch output will be displayed here...")
        main_layout.addWidget(self.dirsearch_output_console, 1)

        self.dirsearch_controls['start_btn'].clicked.connect(self.start_dirsearch_scan)
        self.dirsearch_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_dirsearch_config_widget(self):
        """Creates a reusable, self-contained widget with all of dirsearch's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['url_edit'] = QLineEdit("http://localhost")
        controls['url_edit'].setToolTip("The full URL of the target application to scan (-u).")
        main_layout.addRow("Target URL (-u):", controls['url_edit'])

        wordlist_layout = QHBoxLayout()
        controls['wordlist_edit'] = QLineEdit()
        controls['wordlist_edit'].setPlaceholderText("Path to wordlist file (required)...")
        wordlist_layout.addWidget(controls['wordlist_edit'])
        browse_wordlist_btn = QPushButton("Browse...")
        browse_wordlist_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['wordlist_edit'], "Select Wordlist File"))
        wordlist_layout.addWidget(browse_wordlist_btn)
        main_layout.addRow("Wordlist (-w):", wordlist_layout)

        controls['extensions_edit'] = QLineEdit("php,html,txt")
        controls['extensions_edit'].setToolTip("Comma-separated list of file extensions to append (-e).")
        main_layout.addRow("Extensions (-e):", controls['extensions_edit'])

        controls['threads_edit'] = QLineEdit("25")
        controls['threads_edit'].setToolTip("Number of concurrent threads to use (-t).")
        main_layout.addRow("Threads (-t):", controls['threads_edit'])

        controls['recursive_check'] = QCheckBox("Recursive Scan")
        controls['recursive_check'].setToolTip("Enable recursive scanning (-r).")
        main_layout.addRow(controls['recursive_check'])

        output_file_layout = QHBoxLayout()
        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("Optional: path to save json report...")
        output_file_layout.addWidget(controls['output_edit'])
        browse_output_btn = QPushButton("Browse...")
        browse_output_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_edit'], "Save dirsearch Report"))
        output_file_layout.addWidget(browse_output_btn)
        main_layout.addRow("JSON Report (--json-report):", output_file_layout)

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_dirsearch_scan(self):
        """Starts the dirsearch scan worker thread."""
        controls = self.dirsearch_controls
        if not shutil.which("dirsearch"):
            QMessageBox.critical(self, "dirsearch Error", "'dirsearch' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        url = controls['url_edit'].text().strip()
        wordlist = controls['wordlist_edit'].text().strip()

        if not url or not wordlist:
            QMessageBox.critical(self, "Input Error", "Target URL and Wordlist are required.")
            return

        command = ["dirsearch", "-u", url, "-w", wordlist]

        if extensions := controls['extensions_edit'].text().strip():
            command.extend(["-e", extensions])
        if threads := controls['threads_edit'].text().strip():
            command.extend(["-t", threads])
        if controls['recursive_check'].isChecked():
            command.append("-r")

        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                self.dirsearch_json_temp_file = tmp_json.name
            command.extend(["--json-report", self.dirsearch_json_temp_file])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary file for dirsearch report: {e}")
            return

        command.append("--no-color")

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.dirsearch_output_console.clear()

        self.worker = WorkerThread(self._dirsearch_thread, args=(command, url))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _dirsearch_thread(self, command, url):
        """Worker thread for running the dirsearch command."""
        q = self.tool_results_queue
        logging.info(f"Starting dirsearch with command: {' '.join(command)}")
        q.put(('dirsearch_output', f"$ {' '.join(command)}\n\n"))
        json_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.dirsearch_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('dirsearch_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('dirsearch_output', line))

            process.stdout.close()
            process.wait()

            if not self.tool_stop_event.is_set():
                try:
                    with open(self.dirsearch_json_temp_file, 'r', encoding='utf-8') as f:
                        json_data = f.read()
                    if json_data:
                        q.put(('dirsearch_results', json_data, url))
                except Exception as e:
                    logging.error(f"Could not read dirsearch JSON report: {e}")
                finally:
                    os.remove(self.dirsearch_json_temp_file)
                    self.dirsearch_json_temp_file = None

        except FileNotFoundError:
            q.put(('error', 'dirsearch Error', "'dirsearch' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"dirsearch thread error: {e}", exc_info=True)
            q.put(('error', 'dirsearch Error', str(e)))
        finally:
            q.put(('tool_finished', 'dirsearch_scan', url, json_data))
            with self.thread_finish_lock:
                self.dirsearch_process = None
            logging.info("dirsearch scan thread finished.")

    def _handle_dirsearch_output(self, line):
        self.dirsearch_output_console.insertPlainText(line)
        self.dirsearch_output_console.verticalScrollBar().setValue(self.dirsearch_output_console.verticalScrollBar().maximum())

    def _create_ffuf_tool(self):
        """Creates the UI for the ffuf tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.ffuf_controls = self._create_ffuf_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.ffuf_controls['start_btn'])
        buttons_layout.addWidget(self.ffuf_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.ffuf_output_console = QPlainTextEdit()
        self.ffuf_output_console.setReadOnly(True)
        self.ffuf_output_console.setFont(QFont("Courier New", 10))
        self.ffuf_output_console.setPlaceholderText("ffuf output will be displayed here...")
        main_layout.addWidget(self.ffuf_output_console, 1)

        self.ffuf_controls['start_btn'].clicked.connect(self.start_ffuf_scan)
        self.ffuf_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_ffuf_config_widget(self):
        """Creates a reusable, self-contained widget with all of ffuf's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['url_edit'] = QLineEdit("http://localhost/FUZZ")
        controls['url_edit'].setToolTip("The full URL to fuzz. Must contain the 'FUZZ' keyword.")
        main_layout.addRow("Target URL (-u):", controls['url_edit'])

        wordlist_layout = QHBoxLayout()
        controls['wordlist_edit'] = QLineEdit()
        controls['wordlist_edit'].setPlaceholderText("Path to wordlist file (required)...")
        wordlist_layout.addWidget(controls['wordlist_edit'])
        browse_wordlist_btn = QPushButton("Browse...")
        browse_wordlist_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['wordlist_edit'], "Select Wordlist File"))
        wordlist_layout.addWidget(browse_wordlist_btn)
        main_layout.addRow("Wordlist (-w):", wordlist_layout)

        controls['extensions_edit'] = QLineEdit(".php,.html,.txt")
        controls['extensions_edit'].setToolTip("Comma-separated list of file extensions to append (-e).")
        main_layout.addRow("Extensions (-e):", controls['extensions_edit'])

        controls['threads_edit'] = QLineEdit("40")
        controls['threads_edit'].setToolTip("Number of concurrent threads to use (-t).")
        main_layout.addRow("Threads (-t):", controls['threads_edit'])

        controls['method_edit'] = QLineEdit("GET")
        controls['method_edit'].setToolTip("HTTP method to use (-X).")
        main_layout.addRow("HTTP Method (-X):", controls['method_edit'])

        controls['match_codes_edit'] = QLineEdit("200,204,301,302,307,401,403,405")
        controls['match_codes_edit'].setToolTip("Match HTTP status codes (-mc).")
        main_layout.addRow("Match Codes (-mc):", controls['match_codes_edit'])

        controls['filter_codes_edit'] = QLineEdit("404")
        controls['filter_codes_edit'].setToolTip("Filter HTTP status codes (-fc).")
        main_layout.addRow("Filter Codes (-fc):", controls['filter_codes_edit'])

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_ffuf_scan(self):
        """Starts the ffuf scan worker thread."""
        controls = self.ffuf_controls
        if not shutil.which("ffuf"):
            QMessageBox.critical(self, "ffuf Error", "'ffuf' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        url = controls['url_edit'].text().strip()
        wordlist = controls['wordlist_edit'].text().strip()

        if not url or not wordlist:
            QMessageBox.critical(self, "Input Error", "Target URL and Wordlist are required.")
            return

        if "FUZZ" not in url:
            QMessageBox.critical(self, "Input Error", "Target URL must contain the 'FUZZ' keyword.")
            return

        command = ["ffuf", "-u", url, "-w", wordlist]

        if extensions := controls['extensions_edit'].text().strip():
            command.extend(["-e", extensions])
        if threads := controls['threads_edit'].text().strip():
            command.extend(["-t", threads])
        if method := controls['method_edit'].text().strip():
            command.extend(["-X", method])
        if mc := controls['match_codes_edit'].text().strip():
            command.extend(["-mc", mc])
        if fc := controls['filter_codes_edit'].text().strip():
            command.extend(["-fc", fc])

        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                self.ffuf_json_temp_file = tmp_json.name
            command.extend(["-o", self.ffuf_json_temp_file, "-of", "json"])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary file for ffuf report: {e}")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.ffuf_output_console.clear()

        self.worker = WorkerThread(self._ffuf_thread, args=(command, url))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _ffuf_thread(self, command, url):
        """Worker thread for running the ffuf command."""
        q = self.tool_results_queue
        logging.info(f"Starting ffuf with command: {' '.join(command)}")
        q.put(('ffuf_output', f"$ {' '.join(command)}\n\n"))
        json_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            # We need to capture stderr because ffuf prints progress there
            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.ffuf_process = process

            # Non-blocking read from stderr for progress
            def read_stderr():
                for line in iter(process.stderr.readline, ''):
                    q.put(('ffuf_output', line))

            stderr_thread = threading.Thread(target=read_stderr)
            stderr_thread.daemon = True
            stderr_thread.start()

            # Wait for the process to finish
            process.wait()
            stderr_thread.join(timeout=1)

            if not self.tool_stop_event.is_set():
                try:
                    with open(self.ffuf_json_temp_file, 'r', encoding='utf-8') as f:
                        json_data = f.read()
                    if json_data:
                        q.put(('ffuf_results', json_data))
                except Exception as e:
                    logging.error(f"Could not read ffuf JSON report: {e}")
                finally:
                    os.remove(self.ffuf_json_temp_file)
                    self.ffuf_json_temp_file = None


        except FileNotFoundError:
            q.put(('error', 'ffuf Error', "'ffuf' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"ffuf thread error: {e}", exc_info=True)
            q.put(('error', 'ffuf Error', str(e)))
        finally:
            q.put(('tool_finished', 'ffuf_scan', url, json_data))
            with self.thread_finish_lock:
                self.ffuf_process = None
            logging.info("ffuf scan thread finished.")

    def _handle_ffuf_output(self, line):
        self.ffuf_output_console.insertPlainText(line)
        self.ffuf_output_console.verticalScrollBar().setValue(self.ffuf_output_console.verticalScrollBar().maximum())

    def _create_enum4linux_ng_tool(self):
        """Creates the UI for the enum4linux-ng tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.enum4linux_ng_controls = self._create_enum4linux_ng_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.enum4linux_ng_controls['start_btn'])
        buttons_layout.addWidget(self.enum4linux_ng_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.enum4linux_ng_output_console = QPlainTextEdit()
        self.enum4linux_ng_output_console.setReadOnly(True)
        self.enum4linux_ng_output_console.setFont(QFont("Courier New", 10))
        self.enum4linux_ng_output_console.setPlaceholderText("enum4linux-ng output will be displayed here...")
        main_layout.addWidget(self.enum4linux_ng_output_console, 1)

        self.enum4linux_ng_controls['start_btn'].clicked.connect(self.start_enum4linux_ng_scan)
        self.enum4linux_ng_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_enum4linux_ng_config_widget(self):
        """Creates a reusable, self-contained widget with enum4linux-ng's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        controls = {}

        # --- Target and Auth ---
        top_box = QGroupBox("Target & Authentication")
        top_layout = QFormLayout(top_box)
        controls['target_edit'] = QLineEdit()
        top_layout.addRow("Target Host:", controls['target_edit'])
        controls['user_edit'] = QLineEdit()
        top_layout.addRow("Username (-u):", controls['user_edit'])
        controls['pass_edit'] = QLineEdit(); controls['pass_edit'].setEchoMode(QLineEdit.EchoMode.Password)
        top_layout.addRow("Password (-p):", controls['pass_edit'])
        main_layout.addWidget(top_box)

        # --- Enumeration Options ---
        enum_box = QGroupBox("Enumeration Options")
        enum_layout = QGridLayout(enum_box)
        controls['all_check'] = QCheckBox("All Simple Enum (-A)"); controls['all_check'].setChecked(True)
        controls['users_check'] = QCheckBox("Users (-U)")
        controls['groups_check'] = QCheckBox("Groups (-G)")
        controls['shares_check'] = QCheckBox("Shares (-S)")
        controls['policy_check'] = QCheckBox("Password Policy (-P)")
        controls['os_check'] = QCheckBox("OS Info (-O)")
        enum_layout.addWidget(controls['all_check'], 0, 0)
        enum_layout.addWidget(controls['users_check'], 1, 0)
        enum_layout.addWidget(controls['groups_check'], 1, 1)
        enum_layout.addWidget(controls['shares_check'], 2, 0)
        enum_layout.addWidget(controls['policy_check'], 2, 1)
        enum_layout.addWidget(controls['os_check'], 3, 0)
        main_layout.addWidget(enum_box)

        # --- UI Logic ---
        def toggle_enum_options(checked):
            for key in ['users_check', 'groups_check', 'shares_check', 'policy_check', 'os_check']:
                controls[key].setDisabled(checked)
        controls['all_check'].toggled.connect(toggle_enum_options)
        toggle_enum_options(True)

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_enum4linux_ng_scan(self):
        """Starts the enum4linux-ng scan worker thread."""
        controls = self.enum4linux_ng_controls
        if not shutil.which("enum4linux-ng"):
            QMessageBox.critical(self, "enum4linux-ng Error", "'enum4linux-ng' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = controls['target_edit'].text().strip()
        if not target:
            QMessageBox.critical(self, "Input Error", "A target host is required.")
            return

        command = ["enum4linux-ng"]

        if controls['all_check'].isChecked():
            command.append("-A")
        else:
            if controls['users_check'].isChecked(): command.append("-U")
            if controls['groups_check'].isChecked(): command.append("-G")
            if controls['shares_check'].isChecked(): command.append("-S")
            if controls['policy_check'].isChecked(): command.append("-P")
            if controls['os_check'].isChecked(): command.append("-O")

        if user := controls['user_edit'].text().strip():
            command.extend(["-u", user])
        if pwd := controls['pass_edit'].text().strip():
            command.extend(["-p", pwd])

        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                self.enum4linux_ng_json_temp_file = tmp_json.name
            command.extend(["-oJ", self.enum4linux_ng_json_temp_file])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary file for enum4linux-ng report: {e}")
            return

        command.append(target)

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.enum4linux_ng_output_console.clear()

        self.worker = WorkerThread(self._enum4linux_ng_thread, args=(command, target))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _enum4linux_ng_thread(self, command, target):
        """Worker thread for running the enum4linux-ng command."""
        q = self.tool_results_queue
        logging.info(f"Starting enum4linux-ng with command: {' '.join(command)}")
        q.put(('enum4linux_ng_output', f"$ {' '.join(command)}\n\n"))
        json_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                q.put(('error', 'Platform Error', 'enum4linux-ng is not supported on Windows.'))
                q.put(('tool_finished', 'enum4linux_ng_scan', target, "Platform not supported"))
                return

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.enum4linux_ng_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('enum4linux_ng_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('enum4linux_ng_output', line))

            process.stdout.close()
            process.wait()

            if not self.tool_stop_event.is_set():
                try:
                    with open(self.enum4linux_ng_json_temp_file, 'r', encoding='utf-8') as f:
                        json_data = f.read()
                    if json_data:
                        q.put(('enum4linux_ng_results', json_data, target))
                except Exception as e:
                    logging.error(f"Could not read enum4linux-ng JSON report: {e}")
                finally:
                    os.remove(self.enum4linux_ng_json_temp_file)
                    self.enum4linux_ng_json_temp_file = None

        except FileNotFoundError:
            q.put(('error', 'enum4linux-ng Error', "'enum4linux-ng' command not found. Please ensure it is installed and in your system's PATH."))
            json_data = "Tool not found"
        except Exception as e:
            logging.error(f"enum4linux-ng thread error: {e}", exc_info=True)
            q.put(('error', 'enum4linux-ng Error', str(e)))
            json_data = f"Error: {e}"
        finally:
            q.put(('tool_finished', 'enum4linux_ng_scan', target, json_data))
            with self.thread_finish_lock:
                self.enum4linux_ng_process = None
            logging.info("enum4linux-ng scan thread finished.")

    def _handle_enum4linux_ng_output(self, line):
        self.enum4linux_ng_output_console.insertPlainText(line)
        self.enum4linux_ng_output_console.verticalScrollBar().setValue(self.enum4linux_ng_output_console.verticalScrollBar().maximum())

    def _create_dnsrecon_tool(self):
        """Creates the UI for the dnsrecon tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.dnsrecon_controls = self._create_dnsrecon_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.dnsrecon_controls['start_btn'])
        buttons_layout.addWidget(self.dnsrecon_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.dnsrecon_output_console = QPlainTextEdit()
        self.dnsrecon_output_console.setReadOnly(True)
        self.dnsrecon_output_console.setFont(QFont("Courier New", 10))
        self.dnsrecon_output_console.setPlaceholderText("dnsrecon output will be displayed here...")
        main_layout.addWidget(self.dnsrecon_output_console, 1)

        self.dnsrecon_controls['start_btn'].clicked.connect(self.start_dnsrecon_scan)
        self.dnsrecon_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_dnsrecon_config_widget(self):
        """Creates a reusable, self-contained widget with dnsrecon's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['domain_edit'] = QLineEdit("example.com")
        main_layout.addRow("Domain (-d):", controls['domain_edit'])

        controls['scan_type_combo'] = QComboBox()
        controls['scan_type_combo'].addItems(["std", "axfr", "brt", "srv", "zonewalk"])
        controls['scan_type_combo'].setToolTip("Select the enumeration type (-t).")
        main_layout.addRow("Scan Type (-t):", controls['scan_type_combo'])

        wordlist_layout = QHBoxLayout()
        controls['dict_edit'] = QLineEdit()
        controls['dict_edit'].setPlaceholderText("Path to wordlist for 'brt' scan...")
        wordlist_layout.addWidget(controls['dict_edit'])
        browse_dict_btn = QPushButton("Browse...")
        browse_dict_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['dict_edit'], "Select Dictionary File"))
        wordlist_layout.addWidget(browse_dict_btn)
        main_layout.addRow("Dictionary (-D):", wordlist_layout)

        controls['json_output_edit'] = QLineEdit()
        controls['json_output_edit'].setPlaceholderText("Optional: path to save JSON report...")
        main_layout.addRow("JSON Output (--json):", controls['json_output_edit'])

        # UI Logic
        def toggle_dict_visibility(text):
            controls['dict_edit'].setVisible(text == 'brt')
            browse_dict_btn.setVisible(text == 'brt')
            # Also find and hide the label
            label = main_layout.labelForField(wordlist_layout)
            if label: label.setVisible(text == 'brt')

        controls['scan_type_combo'].currentTextChanged.connect(toggle_dict_visibility)
        toggle_dict_visibility(controls['scan_type_combo'].currentText())

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_dnsrecon_scan(self):
        """Starts the dnsrecon scan worker thread."""
        controls = self.dnsrecon_controls
        if not shutil.which("dnsrecon"):
            QMessageBox.critical(self, "dnsrecon Error", "'dnsrecon' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        domain = controls['domain_edit'].text().strip()
        if not domain:
            QMessageBox.critical(self, "Input Error", "A domain is required.")
            return

        command = ["dnsrecon", "-d", domain]

        scan_type = controls['scan_type_combo'].currentText()
        command.extend(["-t", scan_type])

        if scan_type == 'brt':
            if dictionary := controls['dict_edit'].text().strip():
                command.extend(["-D", dictionary])
            else:
                QMessageBox.warning(self, "Input Warning", "Brute-force scan selected but no dictionary file was provided. Using default wordlist.")

        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".json", encoding='utf-8') as tmp_json:
                self.dnsrecon_json_temp_file = tmp_json.name
            command.extend(["--json", self.dnsrecon_json_temp_file])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary file for dnsrecon report: {e}")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.dnsrecon_output_console.clear()

        self.worker = WorkerThread(self._dnsrecon_thread, args=(command, domain))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _dnsrecon_thread(self, command, domain):
        """Worker thread for running the dnsrecon command."""
        q = self.tool_results_queue
        logging.info(f"Starting dnsrecon with command: {' '.join(command)}")
        q.put(('dnsrecon_output', f"$ {' '.join(command)}\n\n"))
        json_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.dnsrecon_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('dnsrecon_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('dnsrecon_output', line))

            process.stdout.close()
            process.wait()

            if not self.tool_stop_event.is_set():
                try:
                    with open(self.dnsrecon_json_temp_file, 'r', encoding='utf-8') as f:
                        json_data = f.read()
                    if json_data:
                        q.put(('dnsrecon_results', json_data, domain))
                except Exception as e:
                    logging.error(f"Could not read dnsrecon JSON report: {e}")
                finally:
                    os.remove(self.dnsrecon_json_temp_file)
                    self.dnsrecon_json_temp_file = None

        except FileNotFoundError:
            q.put(('error', 'dnsrecon Error', "'dnsrecon' command not found. Please ensure it is installed and in your system's PATH."))
            json_data = "Tool not found"
        except Exception as e:
            logging.error(f"dnsrecon thread error: {e}", exc_info=True)
            q.put(('error', 'dnsrecon Error', str(e)))
            json_data = f"Error: {e}"
        finally:
            q.put(('tool_finished', 'dnsrecon_scan', domain, json_data))
            with self.thread_finish_lock:
                self.dnsrecon_process = None
            logging.info("dnsrecon scan thread finished.")

    def _handle_dnsrecon_output(self, line):
        self.dnsrecon_output_console.insertPlainText(line)
        self.dnsrecon_output_console.verticalScrollBar().setValue(self.dnsrecon_output_console.verticalScrollBar().maximum())

    def _create_fierce_tool(self):
        """Creates the UI for the fierce DNS scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.fierce_controls = self._create_fierce_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.fierce_controls['start_btn'])
        buttons_layout.addWidget(self.fierce_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.fierce_output_console = QPlainTextEdit()
        self.fierce_output_console.setReadOnly(True)
        self.fierce_output_console.setFont(QFont("Courier New", 10))
        self.fierce_output_console.setPlaceholderText("fierce output will be displayed here...")
        main_layout.addWidget(self.fierce_output_console, 1)

        self.fierce_controls['start_btn'].clicked.connect(self.start_fierce_scan)
        self.fierce_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_fierce_config_widget(self):
        """Creates a reusable, self-contained widget with fierce's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['domain_edit'] = QLineEdit("example.com")
        main_layout.addRow("Domain (--domain):", controls['domain_edit'])

        controls['connect_check'] = QCheckBox("Attempt HTTP Connections")
        controls['connect_check'].setToolTip("Attempt to connect to any non-RFC1918 hosts found via HTTP.")
        main_layout.addRow("--connect", controls['connect_check'])

        controls['wide_check'] = QCheckBox("Wide Scan")
        controls['wide_check'].setToolTip("Scan the entire Class C of any discovered records.")
        main_layout.addRow("--wide", controls['wide_check'])

        controls['traverse_edit'] = QLineEdit("5")
        controls['traverse_edit'].setToolTip("Scan a number of IPs above and below discovered hosts.")
        main_layout.addRow("--traverse", controls['traverse_edit'])

        controls['delay_edit'] = QLineEdit("1")
        controls['delay_edit'].setToolTip("Delay in seconds between lookups.")
        main_layout.addRow("--delay", controls['delay_edit'])

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_fierce_scan(self):
        """Starts the fierce scan worker thread."""
        controls = self.fierce_controls
        if not shutil.which("fierce"):
            QMessageBox.critical(self, "fierce Error", "'fierce' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        domain = controls['domain_edit'].text().strip()
        if not domain:
            QMessageBox.critical(self, "Input Error", "A domain is required.")
            return

        command = ["fierce", "--domain", domain]

        if controls['connect_check'].isChecked():
            command.append("--connect")
        if controls['wide_check'].isChecked():
            command.append("--wide")
        if traverse := controls['traverse_edit'].text().strip():
            command.extend(["--traverse", traverse])
        if delay := controls['delay_edit'].text().strip():
            command.extend(["--delay", delay])

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.fierce_output_console.clear()

        self.worker = WorkerThread(self._fierce_thread, args=(command, domain))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _fierce_thread(self, command, domain):
        """Worker thread for running the fierce command."""
        q = self.tool_results_queue
        logging.info(f"Starting fierce with command: {' '.join(command)}")
        q.put(('fierce_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.fierce_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('fierce_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('fierce_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'fierce Error', "'fierce' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"fierce thread error: {e}", exc_info=True)
            q.put(('error', 'fierce Error', str(e)))
        finally:
            q.put(('tool_finished', 'fierce_scan', domain, "".join(full_output)))
            with self.thread_finish_lock:
                self.fierce_process = None
            logging.info("fierce scan thread finished.")

    def _handle_fierce_output(self, line):
        self.fierce_output_console.insertPlainText(line)
        self.fierce_output_console.verticalScrollBar().setValue(self.fierce_output_console.verticalScrollBar().maximum())

    def _create_tools_tab(self,p=None):
        """Creates the tab container for the standard network tools."""
        tools_tabs = QTabWidget()
        tools_tabs.addTab(self._create_nmap_scanner_tool(), "Nmap Scan")
        tools_tabs.addTab(self._create_subdomain_scanner_tool(), "Subdomain Scanner (Sublist3r)")
        tools_tabs.addTab(self._create_subfinder_tool(), "Subdomain Scanner (Subfinder)")
        tools_tabs.addTab(self._create_httpx_tool(), "httpx Probe")
        tools_tabs.addTab(self._create_rustscan_tool(), "RustScan")
        tools_tabs.addTab(self._create_dirsearch_tool(), "dirsearch")
        tools_tabs.addTab(self._create_ffuf_tool(), "ffuf")
        tools_tabs.addTab(self._create_enum4linux_ng_tool(), "enum4linux-ng")
        tools_tabs.addTab(self._create_dnsrecon_tool(), "dnsrecon")
        tools_tabs.addTab(self._create_fierce_tool(), "fierce")
        tools_tabs.addTab(self._create_nikto_scanner_tool(), "Nikto Scan")
        tools_tabs.addTab(self._create_gobuster_tool(), "Gobuster")
        tools_tabs.addTab(self._create_whatweb_tool(), "WhatWeb")
        tools_tabs.addTab(self._create_masscan_tool(), "Masscan")
        tools_tabs.addTab(self._create_port_scanner_tool(), "Port Scanner (Scapy)")
        tools_tabs.addTab(self._create_arp_scan_tool(), "ARP Scan (Scapy)")
        tools_tabs.addTab(self._create_arp_scan_cli_tool(), "ARP Scan (CLI)")
        tools_tabs.addTab(self._create_ping_sweep_tool(), "Ping Sweep")
        tools_tabs.addTab(self._create_traceroute_tool(), "Traceroute")
        return tools_tabs

    def _create_subdomain_scanner_tool(self):
        """Creates the UI for the Sublist3r Subdomain Scanner tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.subdomain_controls = self._create_subdomain_scanner_config_widget()
        layout.addWidget(config_widget)

        # Connect signals
        self.subdomain_controls['start_btn'].clicked.connect(self.start_sublist3r_scan)
        self.subdomain_controls['cancel_btn'].clicked.connect(self.cancel_tool)

        # --- Output Console ---
        self.sublist3r_output = QPlainTextEdit()
        self.sublist3r_output.setReadOnly(True)
        self.sublist3r_output.setFont(QFont("Courier New", 10))
        self.sublist3r_output.setPlaceholderText("Sublist3r output will be displayed here...")
        layout.addWidget(self.sublist3r_output, 1)

        return widget

    def _create_subdomain_scanner_config_widget(self):
        """Creates a reusable, self-contained widget for the Subdomain scanner's configuration."""
        widget = QFrame()
        widget.setObjectName("controlPanel")
        widget.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        controls_layout = QHBoxLayout(widget)

        controls = {}

        controls_layout.addWidget(QLabel("Domain:"))
        controls['domain_edit'] = QLineEdit("example.com")
        controls['domain_edit'].setToolTip("Enter the target domain to enumerate subdomains for (e.g., example.com).")
        controls_layout.addWidget(controls['domain_edit'], 1) # Add stretch

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['cancel_btn'] = QPushButton("Cancel")
        controls['cancel_btn'].setEnabled(False)

        controls_layout.addWidget(controls['start_btn'])
        controls_layout.addWidget(controls['cancel_btn'])

        return widget, controls

    def start_sublist3r_scan(self):
        """Starts the Sublist3r scan worker thread."""
        controls = self.subdomain_controls
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        domain = controls['domain_edit'].text()
        if not domain:
            QMessageBox.critical(self, "Input Error", "Please provide a domain to scan.")
            return

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['cancel_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.sublist3r_output.clear()

        self.worker = WorkerThread(self._sublist3r_thread, args=(domain,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_nikto_scanner_tool(self):
        """Creates the UI for the Nikto Web Scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.nikto_controls = self._create_nikto_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.nikto_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.nikto_output_console = QPlainTextEdit()
        self.nikto_output_console.setReadOnly(True)
        self.nikto_output_console.setFont(QFont("Courier New", 10))
        self.nikto_output_console.setPlaceholderText("Nikto output will be displayed here...")
        main_layout.addWidget(self.nikto_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_nikto_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_nikto_config_widget(self):
        """Creates a reusable, self-contained widget with all of Nikto's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)

        controls = {}

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Nikto Web Scanner</b></font>
        <p>This tool runs the Nikto web server scanner. Use the tabs below to configure the scan.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        nikto_tabs = QTabWidget()
        main_layout.addWidget(nikto_tabs)

        target_tab = QWidget()
        target_layout = QFormLayout(target_tab)
        controls['target_edit'] = QLineEdit("localhost"); controls['target_edit'].setToolTip("Enter the target host, IP, or CIDR range to scan.")
        target_layout.addRow("Target Host:", controls['target_edit'])
        controls['port_edit'] = QLineEdit("80"); controls['port_edit'].setToolTip("Specify the port(s) to scan. Can be a single port, a comma-separated list (80,443), or a range (80-90).")
        target_layout.addRow("Target Port:", controls['port_edit'])
        controls['ssl_check'] = QCheckBox("Force SSL Mode"); controls['ssl_check'].setToolTip("Force SSL mode on the target port(s) (-ssl). Nikto automatically uses SSL on port 443, but this forces it for other non-standard SSL ports.")
        target_layout.addRow(controls['ssl_check'])
        controls['vhost_edit'] = QLineEdit(); controls['vhost_edit'].setToolTip("Specify the virtual host to use in the HTTP Host header. Useful for testing multiple sites on one IP.")
        target_layout.addRow("Virtual Host:", controls['vhost_edit'])
        nikto_tabs.addTab(target_tab, "Target")

        scan_tab = QWidget()
        scan_layout = QFormLayout(scan_tab)
        controls['tuning_combo'] = QComboBox(); controls['tuning_combo'].setToolTip("Select a scan tuning profile (-Tuning) to focus on specific types of tests.\n'x' can be used to reverse the logic (e.g., -Tuning x12 will exclude 'Interesting File' and 'Misconfiguration').")
        controls['tuning_combo'].addItems(["Default", "0 - File Upload", "1 - Interesting File", "2 - Misconfiguration", "3 - Information Disclosure", "4 - Injection (XSS/Script/HTML)", "5 - Remote File Retrieval", "6 - Denial of Service", "8 - Command Execution", "9 - SQL Injection", "a - Auth Bypass", "b - Software ID", "c - Remote Source Inclusion", "x - Reverse Tuning"])
        scan_layout.addRow("Tuning Profile:", controls['tuning_combo'])
        controls['mutate_combo'] = QComboBox(); controls['mutate_combo'].setToolTip("Perform mutation tests (-mutate) to guess additional file and directory names based on known ones found during the scan.")
        controls['mutate_combo'].addItems(["None", "1 - Test files with root dirs", "2 - Guess password files", "3 - Enumerate users via Apache", "4 - Enumerate users via cgiwrap", "5 - Brute force sub-domains", "6 - Guess directory names"])
        scan_layout.addRow("Mutate:", controls['mutate_combo'])
        controls['plugins_edit'] = QLineEdit(); controls['plugins_edit'].setToolTip("Select specific plugins to run, separated by commas (e.g., apache_users,cgi).\nUse 'list' to see all available plugins in the console output.")
        scan_layout.addRow("Plugins:", controls['plugins_edit'])
        controls['cgidirs_edit'] = QLineEdit(); controls['cgidirs_edit'].setToolTip("Scan these CGI directories. Common values are 'all', 'none', or a specific path like '/cgi-bin/'.")
        scan_layout.addRow("CGI Dirs:", controls['cgidirs_edit'])
        nikto_tabs.addTab(scan_tab, "Scan")

        evasion_tab = QWidget()
        evasion_layout = QFormLayout(evasion_tab)
        controls['evasion_combo'] = QComboBox(); controls['evasion_combo'].setToolTip("Select an IDS evasion technique (-evasion).\nThese techniques attempt to bypass Intrusion Detection Systems by encoding or formatting requests in non-standard ways. Use with caution.")
        controls['evasion_combo'].addItems(["None", "1 - Random URI encoding", "2 - Directory self-reference (/./)", "3 - Premature URL ending", "4 - Prepend long random string", "5 - Fake parameter", "6 - TAB as request spacer", "7 - Change case of URL", "8 - Use Windows directory separator (\\)", "A - Use carriage return (0x0d)", "B - Use binary value 0x0b"])
        evasion_layout.addRow("Evasion Technique:", controls['evasion_combo'])
        nikto_tabs.addTab(evasion_tab, "Evasion")

        config_tab = QWidget()
        config_layout = QFormLayout(config_tab)
        controls['timeout_edit'] = QLineEdit("10"); controls['timeout_edit'].setToolTip("Set the timeout in seconds for each individual HTTP request (default is 10).")
        config_layout.addRow("Timeout (s):", controls['timeout_edit'])
        controls['maxtime_edit'] = QLineEdit(); controls['maxtime_edit'].setToolTip("Set the maximum total testing time for the entire scan per host (e.g., 1h, 60m, 3600s).")
        config_layout.addRow("Max Time:", controls['maxtime_edit'])
        controls['pause_edit'] = QLineEdit(); controls['pause_edit'].setToolTip("Pause in seconds between each test (HTTP request). Useful for reducing scan speed.")
        config_layout.addRow("Pause (s):", controls['pause_edit'])
        controls['id_edit'] = QLineEdit(); controls['id_edit'].setToolTip("Provide HTTP Basic authentication credentials in the format 'id:password' or 'id:password:realm'.")
        config_layout.addRow("Auth (id:pass):", controls['id_edit'])
        controls['root_edit'] = QLineEdit(); controls['root_edit'].setToolTip("Prepend a value to the beginning of every request URI. Useful if the web application is in a subdirectory (e.g., /app).")
        config_layout.addRow("Root Directory:", controls['root_edit'])
        controls['proxy_check'] = QCheckBox("Use proxy from nikto.conf")
        config_layout.addRow(controls['proxy_check'])
        nikto_tabs.addTab(config_tab, "Config")

        output_tab = QWidget()
        output_layout = QFormLayout(output_tab)
        output_file_layout = QHBoxLayout()
        controls['output_file_edit'] = QLineEdit(); controls['output_file_edit'].setPlaceholderText("Optional: path to save report...")
        output_file_layout.addWidget(controls['output_file_edit'])
        browse_out_btn = QPushButton("Browse..."); browse_out_btn.clicked.connect(self.browse_nikto_output)
        output_file_layout.addWidget(browse_out_btn)
        output_layout.addRow("Output File:", output_file_layout)
        controls['format_combo'] = QComboBox(); controls['format_combo'].addItems(["html", "csv", "txt", "xml", "nbe"]); controls['format_combo'].setToolTip("Select the report format (requires an output file to be set).")
        output_layout.addRow("Report Format:", controls['format_combo'])
        save_dir_layout = QHBoxLayout()
        controls['save_dir_edit'] = QLineEdit(); controls['save_dir_edit'].setPlaceholderText("Optional: directory to save positive responses...")
        save_dir_layout.addWidget(controls['save_dir_edit'])
        browse_save_btn = QPushButton("Browse..."); browse_save_btn.clicked.connect(self.browse_nikto_save_dir)
        save_dir_layout.addWidget(browse_save_btn)
        output_layout.addRow("Save Directory:", save_dir_layout)
        nikto_tabs.addTab(output_tab, "Output")

        controls['extra_opts_edit'] = QLineEdit(); controls['extra_opts_edit'].setToolTip("Enter any additional, space-separated Nikto flags here. These will be appended directly to the command.")
        main_layout.addWidget(QLabel("Additional Raw Options:"))
        main_layout.addWidget(controls['extra_opts_edit'])

        # Add buttons to controls dict for external connection
        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Nikto Scan")
        controls['stop_btn'] = QPushButton("Stop Nikto"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_nikto_scan(self):
        """Starts the Nikto scan worker thread by building a command from the extensive UI options."""
        controls = self.nikto_controls
        if not shutil.which("nikto"):
            QMessageBox.critical(self, "Nikto Error", "'nikto' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = controls['target_edit'].text().strip()
        if not target:
            QMessageBox.critical(self, "Input Error", "A target host is required.")
            return
        command = ["nikto", "-host", target]

        if port := controls['port_edit'].text().strip():
            command.extend(["-port", port])
        if controls['ssl_check'].isChecked():
            command.append("-ssl")
        if vhost := controls['vhost_edit'].text().strip():
            command.extend(["-vhost", vhost])
        if (tuning_text := controls['tuning_combo'].currentText()) != "Default":
            command.extend(["-Tuning", tuning_text.split(" ")[0]])
        if (mutate_text := controls['mutate_combo'].currentText()) != "None":
            command.extend(["-mutate", mutate_text.split(" ")[0]])
        if plugins := controls['plugins_edit'].text().strip():
            command.extend(["-Plugins", plugins])
        if cgidirs := controls['cgidirs_edit'].text().strip():
            command.extend(["-Cgidirs", cgidirs])
        if (evasion_text := controls['evasion_combo'].currentText()) != "None":
            command.extend(["-evasion", evasion_text.split(" ")[0]])
        if (timeout := controls['timeout_edit'].text().strip()) != "10":
            command.extend(["-timeout", timeout])
        if maxtime := controls['maxtime_edit'].text().strip():
            command.extend(["-maxtime", maxtime])
        if pause := controls['pause_edit'].text().strip():
            command.extend(["-Pause", pause])
        if auth_id := controls['id_edit'].text().strip():
            command.extend(["-id", auth_id])
        if root := controls['root_edit'].text().strip():
            command.extend(["-root", root])
        if controls['proxy_check'].isChecked():
            command.append("-useproxy")
        if output_file := controls['output_file_edit'].text().strip():
            output_format = controls['format_combo'].currentText()
            command.extend(["-o", output_file, "-Format", output_format])
        if save_dir := controls['save_dir_edit'].text().strip():
            command.extend(["-Save", save_dir])
        if extra_opts := controls['extra_opts_edit'].text().strip():
            command.extend(extra_opts.split())

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.nikto_output_console.clear()

        self.worker = WorkerThread(self._nikto_thread, args=(command, target))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_gobuster_tool(self):
        """Creates the UI for the Gobuster tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.gobuster_controls = self._create_gobuster_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.gobuster_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.gobuster_output_console = QPlainTextEdit()
        self.gobuster_output_console.setReadOnly(True)
        self.gobuster_output_console.setFont(QFont("Courier New", 10))
        self.gobuster_output_console.setPlaceholderText("Gobuster output will be displayed here...")
        main_layout.addWidget(self.gobuster_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_gobuster_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_gobuster_config_widget(self):
        """Creates a reusable, self-contained widget with all of Gobuster's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)

        controls = {}

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Gobuster Directory/File Brute-forcer</b></font>
        <p>This tool uses Gobuster to discover hidden directories and files on web servers. Configure the scan using the tabs below.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        gobuster_tabs = QTabWidget()
        main_layout.addWidget(gobuster_tabs)

        scan_tab = QWidget()
        scan_layout = QFormLayout(scan_tab)
        controls['url_edit'] = QLineEdit("http://localhost"); controls['url_edit'].setToolTip("The full URL of the target application to scan, including the scheme (http/https).")
        scan_layout.addRow("Target URL:", controls['url_edit'])

        wordlist_layout = QHBoxLayout()
        controls['wordlist_edit'] = QLineEdit(); controls['wordlist_edit'].setPlaceholderText("Path to wordlist file (required)..."); controls['wordlist_edit'].setToolTip("The path to the wordlist file to use for brute-forcing directories and files.")
        wordlist_layout.addWidget(controls['wordlist_edit'])
        browse_wordlist_btn = QPushButton("Browse..."); browse_wordlist_btn.clicked.connect(self.browse_gobuster_wordlist)
        wordlist_layout.addWidget(browse_wordlist_btn)
        scan_layout.addRow("Wordlist:", wordlist_layout)

        controls['threads_edit'] = QLineEdit("10"); controls['threads_edit'].setToolTip("Number of concurrent threads to use for the scan. Higher numbers are faster but can overwhelm the target.")
        scan_layout.addRow("Threads:", controls['threads_edit'])
        controls['extensions_edit'] = QLineEdit("php,html,txt"); controls['extensions_edit'].setToolTip("A comma-separated list of file extensions to append to each word in the wordlist (e.g., php,html,txt) (-x).")
        scan_layout.addRow("Extensions (-x):", controls['extensions_edit'])
        controls['status_codes_edit'] = QLineEdit("200,204,301,302,307"); controls['status_codes_edit'].setToolTip("A comma-separated list of HTTP status codes to treat as valid and display in the results (e.g., 200,204,301) (-s).")
        scan_layout.addRow("Status Codes (-s):", controls['status_codes_edit'])
        controls['status_codes_blacklist_edit'] = QLineEdit("404"); controls['status_codes_blacklist_edit'].setToolTip("A comma-separated list of status codes to hide from the output (e.g., 403,404) (-b). This takes precedence over the positive status code list.")
        scan_layout.addRow("Blacklist Status Codes (-b):", controls['status_codes_blacklist_edit'])

        checkbox_layout = QHBoxLayout()
        controls['add_slash_check'] = QCheckBox("Add Slash"); controls['add_slash_check'].setToolTip("Append a forward slash to each directory request (-f). Useful for specifically finding directories.")
        checkbox_layout.addWidget(controls['add_slash_check'])
        controls['follow_redirect_check'] = QCheckBox("Follow Redirect"); controls['follow_redirect_check'].setToolTip("Follow HTTP redirects to their final destination (-r).")
        checkbox_layout.addWidget(controls['follow_redirect_check'])
        scan_layout.addRow(checkbox_layout)
        gobuster_tabs.addTab(scan_tab, "Scan Options")

        request_tab = QWidget()
        request_layout = QFormLayout(request_tab)
        controls['useragent_edit'] = QLineEdit(); controls['useragent_edit'].setToolTip("Set a custom User-Agent string for all requests. Can be used to impersonate different browsers.")
        request_layout.addRow("User-Agent:", controls['useragent_edit'])
        controls['random_agent_check'] = QCheckBox("Use Random User-Agent")
        request_layout.addRow(controls['random_agent_check'])
        controls['cookies_edit'] = QLineEdit(); controls['cookies_edit'].setToolTip("Set cookies for the request. The format is 'name=value; name2=value2'.")
        request_layout.addRow("Cookies:", controls['cookies_edit'])
        controls['proxy_edit'] = QLineEdit(); controls['proxy_edit'].setToolTip("Proxy server to use for requests (e.g., http://127.0.0.1:8080, socks5://127.0.0.1:9050).")
        request_layout.addRow("Proxy:", controls['proxy_edit'])
        controls['timeout_edit'] = QLineEdit("10s"); controls['timeout_edit'].setToolTip("Timeout for each individual HTTP request (e.g., 10s, 1m, 500ms).")
        request_layout.addRow("Timeout:", controls['timeout_edit'])
        controls['username_edit'] = QLineEdit()
        request_layout.addRow("Username (Basic Auth):", controls['username_edit'])
        controls['password_edit'] = QLineEdit(); controls['password_edit'].setEchoMode(QLineEdit.EchoMode.Password)
        request_layout.addRow("Password (Basic Auth):", controls['password_edit'])
        gobuster_tabs.addTab(request_tab, "Request")

        output_tab = QWidget()
        output_layout = QFormLayout(output_tab)
        output_file_layout = QHBoxLayout()
        controls['output_file_edit'] = QLineEdit(); controls['output_file_edit'].setPlaceholderText("Optional: path to save output file...")
        output_file_layout.addWidget(controls['output_file_edit'])
        browse_out_btn = QPushButton("Browse..."); browse_out_btn.clicked.connect(self.browse_gobuster_output)
        output_file_layout.addWidget(browse_out_btn)
        output_layout.addRow("Output File:", output_file_layout)

        output_checkbox_layout = QHBoxLayout()
        controls['no_progress_check'] = QCheckBox("No Progress"); controls['no_progress_check'].setToolTip("Don't display the real-time progress bar during the scan (-z).")
        output_checkbox_layout.addWidget(controls['no_progress_check'])
        controls['quiet_check'] = QCheckBox("Quiet"); controls['quiet_check'].setToolTip("Don't print the startup banner and other non-result information (-q).")
        output_checkbox_layout.addWidget(controls['quiet_check'])
        controls['expanded_check'] = QCheckBox("Expanded View"); controls['expanded_check'].setToolTip("Print the full URL for each result, not just the relative path (-e).")
        output_checkbox_layout.addWidget(controls['expanded_check'])
        output_layout.addRow(output_checkbox_layout)
        gobuster_tabs.addTab(output_tab, "Output")

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Gobuster Scan")
        controls['stop_btn'] = QPushButton("Stop Gobuster"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def browse_gobuster_wordlist(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Wordlist File", "", "Text Files (*.txt);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.gobuster_controls['wordlist_edit'].setText(file_path)

    def browse_gobuster_output(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Gobuster Output", "", "Text Files (*.txt);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.gobuster_controls['output_file_edit'].setText(file_path)

    def start_gobuster_scan(self):
        """Starts the Gobuster scan worker thread."""
        controls = self.gobuster_controls
        if not shutil.which("gobuster"):
            QMessageBox.critical(self, "Gobuster Error", "'gobuster' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        url = controls['url_edit'].text().strip()
        wordlist = controls['wordlist_edit'].text().strip()

        if not url or not wordlist:
            QMessageBox.critical(self, "Input Error", "Target URL and Wordlist are required.")
            return

        command = ["gobuster", "dir", "-u", url, "-w", wordlist]

        if threads := controls['threads_edit'].text().strip():
            command.extend(["-t", threads])
        if extensions := controls['extensions_edit'].text().strip():
            command.extend(["-x", extensions])
        if status_codes := controls['status_codes_edit'].text().strip():
            command.extend(["-s", status_codes])
        if blacklist_codes := controls['status_codes_blacklist_edit'].text().strip():
            command.extend(["-b", blacklist_codes])
        if controls['add_slash_check'].isChecked():
            command.append("-f")
        if controls['follow_redirect_check'].isChecked():
            command.append("-r")
        if useragent := controls['useragent_edit'].text().strip():
            command.extend(["-a", useragent])
        if controls['random_agent_check'].isChecked():
            command.append("--random-agent")
        if cookies := controls['cookies_edit'].text().strip():
            command.extend(["-c", cookies])
        if proxy := controls['proxy_edit'].text().strip():
            command.extend(["--proxy", proxy])
        if timeout := controls['timeout_edit'].text().strip():
            command.extend(["--timeout", timeout])
        if username := controls['username_edit'].text().strip():
            command.extend(["-U", username])
        if password := controls['password_edit'].text().strip():
            command.extend(["-P", password])
        if output_file := controls['output_file_edit'].text().strip():
            command.extend(["-o", output_file])
        if controls['no_progress_check'].isChecked():
            command.append("-z")
        if controls['quiet_check'].isChecked():
            command.append("-q")
        if controls['expanded_check'].isChecked():
            command.append("-e")

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.gobuster_output_console.clear()

        self.worker = WorkerThread(self._gobuster_thread, args=(command, url))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_whatweb_tool(self):
        """Creates the UI for the WhatWeb tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>WhatWeb - Web Scanner</b></font>
        <p>This tool identifies technologies used on websites, including content management systems (CMS), blogging platforms, visitor statistics/analytics packages, JavaScript libraries, web servers, and embedded devices.</p>
        """)
        instructions.setFixedHeight(100)
        main_layout.addWidget(instructions)

        # --- Controls ---
        controls_frame = QGroupBox("Scan Options")
        controls_layout = QFormLayout(controls_frame)

        self.whatweb_target_edit = QLineEdit("http://localhost")
        self.whatweb_target_edit.setToolTip("Enter one or more targets to scan, separated by spaces.\nCan be URLs, hostnames, or IP ranges.")
        controls_layout.addRow("Target(s):", self.whatweb_target_edit)

        self.whatweb_aggression_combo = QComboBox()
        self.whatweb_aggression_combo.addItems(["1 - Stealthy", "3 - Aggressive", "4 - Heavy"])
        self.whatweb_aggression_combo.setToolTip("Set the aggression level (-a).\n- 1 (Stealthy): Light and fast, makes few requests.\n- 3 (Aggressive): Makes more requests, may trigger alerts.\n- 4 (Heavy): Very noisy, runs every single plugin.")
        controls_layout.addRow("Aggression Level (-a):", self.whatweb_aggression_combo)

        self.whatweb_verbose_check = QCheckBox("Enable Verbose Output (-v)")
        self.whatweb_verbose_check.setToolTip("Enable verbose output (-v).\nShows more detail during the scan, including which plugins are running.")
        controls_layout.addRow(self.whatweb_verbose_check)

        self.whatweb_extra_opts_edit = QLineEdit()
        self.whatweb_extra_opts_edit.setToolTip("Enter any additional, space-separated WhatWeb flags here. These will be appended directly to the command.")
        controls_layout.addRow("Additional Raw Options:", self.whatweb_extra_opts_edit)

        main_layout.addWidget(controls_frame)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        self.whatweb_start_btn = QPushButton(QIcon("icons/search.svg"), " Start WhatWeb Scan")
        self.whatweb_stop_btn = QPushButton("Stop WhatWeb"); self.whatweb_stop_btn.setEnabled(False)
        buttons_layout.addWidget(self.whatweb_start_btn)
        buttons_layout.addWidget(self.whatweb_stop_btn)
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.whatweb_output_console = QPlainTextEdit()
        self.whatweb_output_console.setReadOnly(True)
        self.whatweb_output_console.setFont(QFont("Courier New", 10))
        self.whatweb_output_console.setPlaceholderText("WhatWeb output will be displayed here...")
        main_layout.addWidget(self.whatweb_output_console, 1)

        self.whatweb_start_btn.clicked.connect(self.start_whatweb_scan)
        self.whatweb_stop_btn.clicked.connect(self.cancel_tool)

        return widget

    def start_whatweb_scan(self):
        """Starts the WhatWeb scan worker thread."""
        if not shutil.which("whatweb"):
            QMessageBox.critical(self, "WhatWeb Error", "'whatweb' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = self.whatweb_target_edit.text().strip()
        if not target:
            QMessageBox.critical(self, "Input Error", "A target is required.")
            return

        command = ["whatweb"]

        # Aggression
        aggression_level = self.whatweb_aggression_combo.currentText().split(" ")[0]
        command.extend(["-a", aggression_level])

        # Verbose
        if self.whatweb_verbose_check.isChecked():
            command.append("-v")

        # Additional Options
        if extra_opts := self.whatweb_extra_opts_edit.text().strip():
            command.extend(extra_opts.split())

        # Target(s) must be last
        command.extend(target.split())

        self.is_tool_running = True
        self.whatweb_start_btn.setEnabled(False)
        self.whatweb_stop_btn.setEnabled(True)
        self.tool_stop_event.clear()
        self.whatweb_output_console.clear()

        self.worker = WorkerThread(self._whatweb_thread, args=(command,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_sqlmap_tool(self):
        """Creates the UI for the SQLMap tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>SQLMap: The automatic SQL injection and database takeover tool.</b></font>
        <p><b>WARNING:</b> This is a powerful attack tool. You MUST have explicit permission to test the target. Ensure you have selected a target URL or provided a request file.</p>
        """)
        instructions.setFixedHeight(100)
        main_layout.addWidget(instructions)

        # --- Main Tab Widget for Options ---
        sqlmap_tabs = QTabWidget()
        main_layout.addWidget(sqlmap_tabs)

        # --- Target Tab ---
        target_tab = QWidget()
        target_layout = QFormLayout(target_tab)
        self.sqlmap_url_edit = QLineEdit()
        self.sqlmap_url_edit.setToolTip('Target URL to test. Must be a full URL including the vulnerable parameter(s).\nExample: "http://testphp.vulnweb.com/listproducts.php?cat=1"')
        target_layout.addRow("Target URL (-u):", self.sqlmap_url_edit)

        req_file_layout = QHBoxLayout()
        self.sqlmap_reqfile_edit = QLineEdit()
        self.sqlmap_reqfile_edit.setToolTip("Load a raw HTTP request from a file (e.g., from Burp Suite).\nThis is often more reliable than using the URL field, especially for complex requests with headers and POST data.")
        req_file_layout.addWidget(self.sqlmap_reqfile_edit)
        browse_req_btn = QPushButton("Browse...")
        browse_req_btn.clicked.connect(self.browse_sqlmap_request_file)
        req_file_layout.addWidget(browse_req_btn)
        target_layout.addRow("Request File (-r):", req_file_layout)
        sqlmap_tabs.addTab(target_tab, "Target")

        # --- Request Tab ---
        req_tab = QWidget()
        req_layout = QFormLayout(req_tab)
        self.sqlmap_method_edit = QLineEdit()
        req_layout.addRow("Method:", self.sqlmap_method_edit)
        self.sqlmap_data_edit = QLineEdit(); self.sqlmap_data_edit.setToolTip("POST data to send with the request (--data).")
        req_layout.addRow("Data (--data):", self.sqlmap_data_edit)
        self.sqlmap_cookie_edit = QLineEdit(); self.sqlmap_cookie_edit.setToolTip("HTTP Cookie header value (--cookie). Format: 'name=value; name2=value2'.")
        req_layout.addRow("Cookie (--cookie):", self.sqlmap_cookie_edit)
        self.sqlmap_useragent_edit = QLineEdit(); self.sqlmap_useragent_edit.setToolTip("Set a custom User-Agent (--user-agent). 'SQLMAP' is the default.")
        req_layout.addRow("User-Agent (--user-agent):", self.sqlmap_useragent_edit)
        self.sqlmap_referer_edit = QLineEdit(); self.sqlmap_referer_edit.setToolTip("Set a custom HTTP Referer header (--referer).")
        req_layout.addRow("Referer (--referer):", self.sqlmap_referer_edit)
        self.sqlmap_headers_edit = QTextEdit(); self.sqlmap_headers_edit.setToolTip("Extra headers to include in the request (--headers).\nFormat: 'Header1: Value1\\nHeader2: Value2'.")
        req_layout.addRow("Extra Headers (--headers):", self.sqlmap_headers_edit)
        self.sqlmap_auth_type_edit = QLineEdit()
        req_layout.addRow("Auth Type (--auth-type):", self.sqlmap_auth_type_edit)
        self.sqlmap_auth_cred_edit = QLineEdit()
        req_layout.addRow("Auth Creds (--auth-cred):", self.sqlmap_auth_cred_edit)
        self.sqlmap_proxy_edit = QLineEdit()
        req_layout.addRow("Proxy (--proxy):", self.sqlmap_proxy_edit)
        self.sqlmap_random_agent_check = QCheckBox("Random Agent (--random-agent)")
        req_layout.addRow(self.sqlmap_random_agent_check)
        self.sqlmap_force_ssl_check = QCheckBox("Force SSL (--force-ssl)")
        req_layout.addRow(self.sqlmap_force_ssl_check)
        sqlmap_tabs.addTab(req_tab, "Request")

        # --- Injection Tab ---
        inj_tab = QWidget()
        inj_layout = QFormLayout(inj_tab)
        self.sqlmap_test_param_edit = QLineEdit(); self.sqlmap_test_param_edit.setToolTip("Specify a specific parameter to test for SQL injection (-p).")
        inj_layout.addRow("Test Parameter (-p):", self.sqlmap_test_param_edit)
        self.sqlmap_dbms_edit = QLineEdit(); self.sqlmap_dbms_edit.setToolTip("Force sqlmap to test for a specific backend DBMS (e.g., MySQL, MSSQL) (--dbms).")
        inj_layout.addRow("DBMS (--dbms):", self.sqlmap_dbms_edit)
        self.sqlmap_level_combo = QComboBox()
        self.sqlmap_level_combo.addItems(["1","2","3","4","5"]); self.sqlmap_level_combo.setToolTip("Level of tests to perform (1-5). Higher levels perform more tests (--level).")
        inj_layout.addRow("Level (1-5):", self.sqlmap_level_combo)
        self.sqlmap_risk_combo = QComboBox()
        self.sqlmap_risk_combo.addItems(["1","2","3"]); self.sqlmap_risk_combo.setToolTip("Risk of tests to perform (1-3). Higher risk tests are more likely to cause issues but may find more vulnerabilities (--risk).")
        inj_layout.addRow("Risk (1-3):", self.sqlmap_risk_combo)
        self.sqlmap_technique_edit = QLineEdit("BEUSTQ"); self.sqlmap_technique_edit.setToolTip("Specify the injection techniques to use (--technique).\nB: Boolean-based blind\nE: Error-based\nU: Union query-based\nS: Stacked queries\nT: Time-based blind\nQ: Inline queries")
        inj_layout.addRow("Techniques (--technique):", self.sqlmap_technique_edit)
        sqlmap_tabs.addTab(inj_tab, "Injection")

        # --- Enumeration Tab ---
        enum_tab = QWidget()
        enum_layout = QVBoxLayout(enum_tab)
        enum_grid = QGridLayout()
        self.sqlmap_enum_all_check = QCheckBox("All (-a)"); self.sqlmap_enum_all_check.setToolTip("Enumerate everything (-a).")
        self.sqlmap_enum_banner_check = QCheckBox("Banner (-b)"); self.sqlmap_enum_banner_check.setToolTip("Retrieve the DBMS banner (-b).")
        self.sqlmap_enum_current_user_check = QCheckBox("Current User"); self.sqlmap_enum_current_user_check.setToolTip("Retrieve the current DBMS user (--current-user).")
        self.sqlmap_enum_current_db_check = QCheckBox("Current DB"); self.sqlmap_enum_current_db_check.setToolTip("Retrieve the current database name (--current-db).")
        self.sqlmap_enum_is_dba_check = QCheckBox("Is DBA?"); self.sqlmap_enum_is_dba_check.setToolTip("Check if the current user is a Database Administrator (--is-dba).")
        self.sqlmap_enum_passwords_check = QCheckBox("Passwords"); self.sqlmap_enum_passwords_check.setToolTip("Attempt to dump DBMS user password hashes (--passwords).")
        self.sqlmap_enum_dbs_check = QCheckBox("Databases"); self.sqlmap_enum_dbs_check.setToolTip("Enumerate all databases (--dbs).")
        self.sqlmap_enum_tables_check = QCheckBox("Tables"); self.sqlmap_enum_tables_check.setToolTip("Enumerate tables in a specific database (--tables).")
        self.sqlmap_enum_columns_check = QCheckBox("Columns"); self.sqlmap_enum_columns_check.setToolTip("Enumerate columns in a specific table (--columns).")
        self.sqlmap_enum_schema_check = QCheckBox("Schema"); self.sqlmap_enum_schema_check.setToolTip("Enumerate the entire DBMS schema (--schema).")
        self.sqlmap_enum_dump_check = QCheckBox("Dump Table Entries"); self.sqlmap_enum_dump_check.setToolTip("Dump entries from a specific table (--dump).")
        self.sqlmap_enum_dump_all_check = QCheckBox("Dump All"); self.sqlmap_enum_dump_all_check.setToolTip("Dump all table entries from all databases. Warning: this can be very slow.")
        enum_grid.addWidget(self.sqlmap_enum_all_check, 0, 0)
        enum_grid.addWidget(self.sqlmap_enum_banner_check, 0, 1)
        enum_grid.addWidget(self.sqlmap_enum_current_user_check, 0, 2)
        enum_grid.addWidget(self.sqlmap_enum_current_db_check, 1, 0)
        enum_grid.addWidget(self.sqlmap_enum_is_dba_check, 1, 1)
        enum_grid.addWidget(self.sqlmap_enum_passwords_check, 1, 2)
        enum_grid.addWidget(self.sqlmap_enum_dbs_check, 2, 0)
        enum_grid.addWidget(self.sqlmap_enum_tables_check, 2, 1)
        enum_grid.addWidget(self.sqlmap_enum_columns_check, 2, 2)
        enum_grid.addWidget(self.sqlmap_enum_schema_check, 3, 0)
        enum_grid.addWidget(self.sqlmap_enum_dump_check, 3, 1)
        enum_grid.addWidget(self.sqlmap_enum_dump_all_check, 3, 2)
        enum_layout.addLayout(enum_grid)
        enum_form_layout = QFormLayout()
        self.sqlmap_db_edit = QLineEdit(); self.sqlmap_db_edit.setToolTip("Database to use for enumeration (-D).")
        enum_form_layout.addRow("Database (-D):", self.sqlmap_db_edit)
        self.sqlmap_tbl_edit = QLineEdit(); self.sqlmap_tbl_edit.setToolTip("Table to use for enumeration (-T).")
        enum_form_layout.addRow("Table (-T):", self.sqlmap_tbl_edit)
        self.sqlmap_col_edit = QLineEdit(); self.sqlmap_col_edit.setToolTip("Column to use for enumeration (-C).")
        enum_form_layout.addRow("Column (-C):", self.sqlmap_col_edit)
        enum_layout.addLayout(enum_form_layout)
        sqlmap_tabs.addTab(enum_tab, "Enumeration")

        # --- Access Tab ---
        access_tab = QWidget()
        access_layout = QFormLayout(access_tab)
        self.sqlmap_os_shell_check = QCheckBox("OS Shell (--os-shell)"); self.sqlmap_os_shell_check.setToolTip("Attempt to get an interactive OS shell (--os-shell). This requires a successful file write vulnerability.")
        access_layout.addRow(self.sqlmap_os_shell_check)
        self.sqlmap_sql_shell_check = QCheckBox("SQL Shell (--sql-shell)"); self.sqlmap_sql_shell_check.setToolTip("Get an interactive SQL shell (--sql-shell).")
        access_layout.addRow(self.sqlmap_sql_shell_check)
        sqlmap_tabs.addTab(access_tab, "Access")

        # --- General Tab ---
        general_tab = QWidget()
        general_layout = QFormLayout(general_tab)
        self.sqlmap_threads_edit = QLineEdit("1"); self.sqlmap_threads_edit.setToolTip("Number of concurrent threads to use (--threads).")
        general_layout.addRow("Threads (--threads):", self.sqlmap_threads_edit)
        self.sqlmap_batch_check = QCheckBox("Batch Mode (--batch)"); self.sqlmap_batch_check.setToolTip("Run in batch mode (--batch). Never asks for user input, uses default answers.")
        general_layout.addRow(self.sqlmap_batch_check)
        self.sqlmap_flush_session_check = QCheckBox("Flush Session (--flush-session)"); self.sqlmap_flush_session_check.setToolTip("Flush session files for the target, starting fresh (--flush-session).")
        general_layout.addRow(self.sqlmap_flush_session_check)
        sqlmap_tabs.addTab(general_tab, "General")

        # --- Additional Options ---
        main_layout.addWidget(QLabel("Additional Raw Options:"))
        self.sqlmap_extra_opts_edit = QLineEdit()
        self.sqlmap_extra_opts_edit.setToolTip("Enter any additional, space-separated SQLMap flags here. These will be appended to the command.")
        main_layout.addWidget(self.sqlmap_extra_opts_edit)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        self.sqlmap_start_btn = QPushButton(QIcon("icons/search.svg"), " Start SQLMap Scan")
        self.sqlmap_stop_btn = QPushButton("Stop SQLMap"); self.sqlmap_stop_btn.setEnabled(False)
        buttons_layout.addWidget(self.sqlmap_start_btn)
        buttons_layout.addWidget(self.sqlmap_stop_btn)
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.sqlmap_output_console = QPlainTextEdit()
        self.sqlmap_output_console.setReadOnly(True)
        self.sqlmap_output_console.setFont(QFont("Courier New", 10))
        self.sqlmap_output_console.setPlaceholderText("SQLMap output will be displayed here...")
        main_layout.addWidget(self.sqlmap_output_console, 1)

        self.sqlmap_start_btn.clicked.connect(self.start_sqlmap_scan)
        self.sqlmap_stop_btn.clicked.connect(self.cancel_tool)

        return widget

    def browse_sqlmap_request_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Request File", "", "All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.sqlmap_reqfile_edit.setText(file_path)

    def start_sqlmap_scan(self):
        """Starts the SQLMap scan worker thread."""
        if not shutil.which("sqlmap"):
            QMessageBox.critical(self, "SQLMap Error", "'sqlmap' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        command = ["sqlmap"]

        # --- Target Tab ---
        url = self.sqlmap_url_edit.text().strip()
        reqfile = self.sqlmap_reqfile_edit.text().strip()
        if url:
            command.extend(["-u", url])
        elif reqfile:
            command.extend(["-r", reqfile])
        else:
            QMessageBox.critical(self, "Input Error", "A Target URL (-u) or Request File (-r) is required.")
            return

        # --- Request Tab ---
        if method := self.sqlmap_method_edit.text().strip(): command.extend(["--method", method])
        if data := self.sqlmap_data_edit.text().strip(): command.extend(["--data", data])
        if cookie := self.sqlmap_cookie_edit.text().strip(): command.extend(["--cookie", cookie])
        if agent := self.sqlmap_useragent_edit.text().strip(): command.extend(["--user-agent", agent])
        if referer := self.sqlmap_referer_edit.text().strip(): command.extend(["--referer", referer])
        if headers := self.sqlmap_headers_edit.toPlainText().strip(): command.extend(["--headers", headers])
        if auth_type := self.sqlmap_auth_type_edit.text().strip(): command.extend(["--auth-type", auth_type])
        if auth_cred := self.sqlmap_auth_cred_edit.text().strip(): command.extend(["--auth-cred", auth_cred])
        if proxy := self.sqlmap_proxy_edit.text().strip(): command.extend(["--proxy", proxy])
        if self.sqlmap_random_agent_check.isChecked(): command.append("--random-agent")
        if self.sqlmap_force_ssl_check.isChecked(): command.append("--force-ssl")

        # --- Injection Tab ---
        if test_param := self.sqlmap_test_param_edit.text().strip(): command.extend(["-p", test_param])
        if dbms := self.sqlmap_dbms_edit.text().strip(): command.extend(["--dbms", dbms])
        command.extend(["--level", self.sqlmap_level_combo.currentText()])
        command.extend(["--risk", self.sqlmap_risk_combo.currentText()])
        if tech := self.sqlmap_technique_edit.text().strip(): command.extend(["--technique", tech])

        # --- Enumeration Tab ---
        if self.sqlmap_enum_all_check.isChecked(): command.append("-a")
        if self.sqlmap_enum_banner_check.isChecked(): command.append("-b")
        if self.sqlmap_enum_current_user_check.isChecked(): command.append("--current-user")
        if self.sqlmap_enum_current_db_check.isChecked(): command.append("--current-db")
        if self.sqlmap_enum_is_dba_check.isChecked(): command.append("--is-dba")
        if self.sqlmap_enum_passwords_check.isChecked(): command.append("--passwords")
        if self.sqlmap_enum_dbs_check.isChecked(): command.append("--dbs")
        if self.sqlmap_enum_tables_check.isChecked(): command.append("--tables")
        if self.sqlmap_enum_columns_check.isChecked(): command.append("--columns")
        if self.sqlmap_enum_schema_check.isChecked(): command.append("--schema")
        if self.sqlmap_enum_dump_check.isChecked(): command.append("--dump")
        if self.sqlmap_enum_dump_all_check.isChecked(): command.append("--dump-all")
        if db := self.sqlmap_db_edit.text().strip(): command.extend(["-D", db])
        if tbl := self.sqlmap_tbl_edit.text().strip(): command.extend(["-T", tbl])
        if col := self.sqlmap_col_edit.text().strip(): command.extend(["-C", col])

        # --- Access Tab ---
        if self.sqlmap_os_shell_check.isChecked(): command.append("--os-shell")
        if self.sqlmap_sql_shell_check.isChecked(): command.append("--sql-shell")

        # --- General Tab ---
        if threads := self.sqlmap_threads_edit.text().strip(): command.extend(["--threads", threads])
        if self.sqlmap_batch_check.isChecked(): command.append("--batch")
        if self.sqlmap_flush_session_check.isChecked(): command.append("--flush-session")

        # --- Default flags for GUI operation ---
        if "--batch" not in command:
            command.append("--batch")
        command.extend(["--answers", "quit=N"])


        # --- Additional Options ---
        if extra_opts := self.sqlmap_extra_opts_edit.text().strip():
            command.extend(extra_opts.split())

        self.is_tool_running = True
        self.sqlmap_start_btn.setEnabled(False)
        self.sqlmap_stop_btn.setEnabled(True)
        self.tool_stop_event.clear()
        self.sqlmap_output_console.clear()

        target_for_log = url if url else reqfile
        self.worker = WorkerThread(self._sqlmap_thread, args=(command, target_for_log))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _sqlmap_thread(self, command, target):
        """Worker thread for running the sqlmap command."""
        q = self.tool_results_queue
        logging.info(f"Starting SQLMap with command: {' '.join(command)}")
        q.put(('sqlmap_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.sqlmap_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('sqlmap_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('sqlmap_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'SQLMap Error', "'sqlmap' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"SQLMap thread error: {e}", exc_info=True)
            q.put(('error', 'SQLMap Error', str(e)))
        finally:
            q.put(('tool_finished', 'sqlmap_scan', target, "".join(full_output)))
            with self.thread_finish_lock:
                self.sqlmap_process = None
            logging.info("SQLMap scan thread finished.")

    def _create_hashcat_tool(self):
        """Creates the UI for the Hashcat tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Hashcat: The world's fastest password cracker.</b></font>
        <p><b>WARNING:</b> This is a powerful tool. Ensure you have permission to crack the provided hashes. This tool can be very resource-intensive.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Main Tab Widget for Options ---
        hashcat_tabs = QTabWidget()
        main_layout.addWidget(hashcat_tabs)

        # --- Main Config Tab ---
        config_tab = QWidget()
        config_layout = QFormLayout(config_tab)

        hashfile_layout = QHBoxLayout()
        self.hashcat_hashfile_edit = QLineEdit()
        self.hashcat_hashfile_edit.setToolTip("The file containing the hashes to crack.")
        hashfile_layout.addWidget(self.hashcat_hashfile_edit)
        browse_hash_btn = QPushButton("Browse...")
        browse_hash_btn.clicked.connect(self.browse_hashcat_hashfile)
        hashfile_layout.addWidget(browse_hash_btn)
        config_layout.addRow("Hash File:", hashfile_layout)

        hash_type_layout = QVBoxLayout()
        self.hashcat_type_edit = QLineEdit()
        self.hashcat_type_edit.setPlaceholderText("e.g., 0 for MD5, 1000 for NTLM")
        self.hashcat_type_edit.setToolTip("The hash mode code (-m). Click the link below to find the correct mode for your hash type.")
        hash_type_label = QLabel("Hash Mode (-m):")
        hash_type_link = QLabel('<a href="https://hashcat.net/wiki/doku.php?id=example_hashes">Find Hash Mode</a>')
        hash_type_link.setOpenExternalLinks(True)
        hash_type_layout.addWidget(self.hashcat_type_edit)
        hash_type_layout.addWidget(hash_type_link)
        config_layout.addRow(hash_type_label, hash_type_layout)

        self.hashcat_attack_mode_combo = QComboBox()
        self.hashcat_attack_mode_combo.addItems([
            "0 - Straight (Dictionary)",
            "1 - Combination",
            "3 - Brute-force (Mask)",
            "6 - Hybrid (Wordlist + Mask)",
            "7 - Hybrid (Mask + Wordlist)"
        ])
        self.hashcat_attack_mode_combo.setToolTip("The attack mode (-a) to use.\n- Straight: Dictionary attack.\n- Combination: Combines words from two dictionaries.\n- Brute-force: Tries all possible character combinations based on a mask.\n- Hybrid: Combines dictionary words with a mask.")
        config_layout.addRow("Attack Mode (-a):", self.hashcat_attack_mode_combo)

        outfile_layout = QHBoxLayout()
        self.hashcat_outfile_edit = QLineEdit()
        self.hashcat_outfile_edit.setToolTip("The file to save cracked hashes to (-o).")
        outfile_layout.addWidget(self.hashcat_outfile_edit)
        browse_out_btn = QPushButton("Browse...")
        browse_out_btn.clicked.connect(self.browse_hashcat_outfile)
        outfile_layout.addWidget(browse_out_btn)
        config_layout.addRow("Output File (-o):", outfile_layout)

        hashcat_tabs.addTab(config_tab, "Configuration")

        # --- Wordlist Tab ---
        self.hashcat_wordlist_tab = QWidget()
        wordlist_layout = QVBoxLayout(self.hashcat_wordlist_tab)
        self.hashcat_wordlist_list = QListWidget()
        wordlist_layout.addWidget(self.hashcat_wordlist_list)
        wordlist_buttons = QHBoxLayout()
        add_wordlist_btn = QPushButton("Add Wordlist(s)")
        add_wordlist_btn.clicked.connect(self.browse_hashcat_wordlist)
        remove_wordlist_btn = QPushButton("Remove Selected")
        remove_wordlist_btn.clicked.connect(lambda: self.hashcat_wordlist_list.takeItem(self.hashcat_wordlist_list.currentRow()))
        wordlist_buttons.addWidget(add_wordlist_btn)
        wordlist_buttons.addWidget(remove_wordlist_btn)
        wordlist_layout.addLayout(wordlist_buttons)
        hashcat_tabs.addTab(self.hashcat_wordlist_tab, "Wordlists")

        # --- Mask Tab ---
        self.hashcat_mask_tab = QWidget()
        mask_layout = QFormLayout(self.hashcat_mask_tab)
        self.hashcat_mask_edit = QLineEdit()
        self.hashcat_mask_edit.setPlaceholderText("e.g., ?l?l?l?l?l?l?l?l")
        self.hashcat_mask_edit.setToolTip("The mask to use for brute-force or hybrid attacks. Click the link below for syntax details.")
        mask_layout.addRow("Mask:", self.hashcat_mask_edit)
        mask_link = QLabel('<a href="https://hashcat.net/wiki/doku.php?id=mask_attack">Mask Attack Info</a>')
        mask_link.setOpenExternalLinks(True)
        mask_layout.addRow(mask_link)
        hashcat_tabs.addTab(self.hashcat_mask_tab, "Mask")

        # --- Advanced Tab ---
        adv_tab = QWidget()
        adv_layout = QFormLayout(adv_tab)
        self.hashcat_force_check = QCheckBox("Ignore warnings")
        self.hashcat_force_check.setToolTip("Ignore warnings and force the cracking session to start (--force).")
        adv_layout.addRow("Force:", self.hashcat_force_check)
        self.hashcat_extra_opts_edit = QLineEdit()
        self.hashcat_extra_opts_edit.setToolTip("Enter any additional, space-separated Hashcat flags here. These will be appended to the command.")
        adv_layout.addRow("Additional Options:", self.hashcat_extra_opts_edit)
        adv_tab.setLayout(adv_layout)
        hashcat_tabs.addTab(adv_tab, "Advanced")

        # --- UI Logic ---
        def update_tabs(text):
            mode = int(text.split(" ")[0])
            self.hashcat_wordlist_tab.setEnabled(mode in [0, 1, 6, 7])
            self.hashcat_mask_tab.setEnabled(mode in [3, 6, 7])
        self.hashcat_attack_mode_combo.currentTextChanged.connect(update_tabs)
        update_tabs(self.hashcat_attack_mode_combo.currentText()) # Initial state

        # --- Action Buttons & Output ---
        buttons_layout = QHBoxLayout()
        self.hashcat_start_btn = QPushButton(QIcon("icons/tool.svg"), " Start Hashcat")
        self.hashcat_stop_btn = QPushButton("Stop Hashcat"); self.hashcat_stop_btn.setEnabled(False)
        buttons_layout.addWidget(self.hashcat_start_btn)
        buttons_layout.addWidget(self.hashcat_stop_btn)
        main_layout.addLayout(buttons_layout)

        self.hashcat_output_console = QPlainTextEdit()
        self.hashcat_output_console.setReadOnly(True)
        self.hashcat_output_console.setFont(QFont("Courier New", 10))
        main_layout.addWidget(self.hashcat_output_console, 1)

        self.hashcat_start_btn.clicked.connect(self.start_hashcat_scan)
        self.hashcat_stop_btn.clicked.connect(self.cancel_tool)

        return widget

    def browse_hashcat_hashfile(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Hash File", "", "All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.hashcat_hashfile_edit.setText(file_path)

    def browse_hashcat_wordlist(self):
        file_paths, _ = QFileDialog.getOpenFileNames(self, "Select Wordlist(s)", "", "Text Files (*.txt);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        for file_path in file_paths:
            if file_path:
                self.hashcat_wordlist_list.addItem(file_path)

    def browse_hashcat_outfile(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Output File", "", "Text Files (*.txt);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.hashcat_outfile_edit.setText(file_path)

    def start_hashcat_scan(self):
        """Starts the Hashcat worker thread."""
        if not shutil.which("hashcat"):
            QMessageBox.critical(self, "Hashcat Error", "'hashcat' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        command = ["hashcat"]

        # --- Config Tab ---
        hashfile = self.hashcat_hashfile_edit.text().strip()
        if not hashfile:
            QMessageBox.critical(self, "Input Error", "Hash file is required.")
            return
        command.append(hashfile)

        hash_type = self.hashcat_type_edit.text().strip()
        if not hash_type:
            QMessageBox.critical(self, "Input Error", "Hash mode (-m) is required.")
            return
        command.extend(["-m", hash_type])

        attack_mode = self.hashcat_attack_mode_combo.currentText().split(" ")[0]
        command.extend(["-a", attack_mode])

        if outfile := self.hashcat_outfile_edit.text().strip():
            command.extend(["-o", outfile])

        # --- Wordlist/Mask Tabs ---
        if self.hashcat_wordlist_tab.isEnabled():
            wordlists = [self.hashcat_wordlist_list.item(i).text() for i in range(self.hashcat_wordlist_list.count())]
            if not wordlists:
                QMessageBox.critical(self, "Input Error", "This attack mode requires at least one wordlist.")
                return
            command.extend(wordlists)

        if self.hashcat_mask_tab.isEnabled():
            mask = self.hashcat_mask_edit.text().strip()
            if not mask:
                QMessageBox.critical(self, "Input Error", "This attack mode requires a mask.")
                return
            command.append(mask)

        # --- Advanced Tab ---
        if self.hashcat_force_check.isChecked():
            command.append("--force")
        if extra_opts := self.hashcat_extra_opts_edit.text().strip():
            command.extend(extra_opts.split())

        self.is_tool_running = True
        self.hashcat_start_btn.setEnabled(False)
        self.hashcat_stop_btn.setEnabled(True)
        self.tool_stop_event.clear()
        self.hashcat_output_console.clear()

        self.worker = WorkerThread(self._hashcat_thread, args=(command, hashfile))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _hashcat_thread(self, command, hashfile):
        """Worker thread for running the hashcat command."""
        q = self.tool_results_queue
        logging.info(f"Starting Hashcat with command: {' '.join(command)}")
        q.put(('hashcat_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.hashcat_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('hashcat_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('hashcat_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Hashcat Error', "'hashcat' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Hashcat thread error: {e}", exc_info=True)
            q.put(('error', 'Hashcat Error', str(e)))
        finally:
            q.put(('tool_finished', 'hashcat_scan', hashfile, "".join(full_output)))
            with self.thread_finish_lock:
                self.hashcat_process = None
            logging.info("Hashcat scan thread finished.")

    def _create_nuclei_tool(self):
        """Creates the UI for the Nuclei Web Scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.nuclei_controls = self._create_nuclei_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.nuclei_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.nuclei_output_console = QPlainTextEdit()
        self.nuclei_output_console.setReadOnly(True)
        self.nuclei_output_console.setFont(QFont("Courier New", 10))
        self.nuclei_output_console.setPlaceholderText("Nuclei output will be displayed here...")
        main_layout.addWidget(self.nuclei_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_nuclei_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_nuclei_config_widget(self):
        """Creates a reusable, self-contained widget with Nuclei's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)

        controls = {}

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Nuclei - Template-Based Vulnerability Scanner</b></font>
        <p>This tool runs Nuclei, a fast and flexible vulnerability scanner. Configure the scan using the options below.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Main Tab Widget for Options ---
        nuclei_tabs = QTabWidget()
        main_layout.addWidget(nuclei_tabs)

        # --- Target Tab ---
        target_tab = QWidget()
        target_layout = QFormLayout(target_tab)

        target_file_layout = QHBoxLayout()
        controls['target_edit'] = QLineEdit()
        controls['target_edit'].setPlaceholderText("e.g., https://example.com or path to target list file...")
        controls['target_edit'].setToolTip("Enter a single target URL (-u) or path to a file with a list of targets (-l).")
        target_file_layout.addWidget(controls['target_edit'])
        browse_target_btn = QPushButton("Browse...")
        browse_target_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['target_edit'], "Select Target List File"))
        target_file_layout.addWidget(browse_target_btn)
        target_layout.addRow("Target URL/List (-u/-l):", target_file_layout)
        nuclei_tabs.addTab(target_tab, "Target")

        # --- Templates Tab ---
        templates_tab = QWidget()
        templates_layout = QFormLayout(templates_tab)

        template_file_layout = QHBoxLayout()
        controls['templates_edit'] = QLineEdit()
        controls['templates_edit'].setPlaceholderText("e.g., cves/, http/exposures/, path/to/custom/templates/")
        controls['templates_edit'].setToolTip("Comma-separated list of template directories or single template files to run (-t).")
        template_file_layout.addWidget(controls['templates_edit'])
        browse_template_btn = QPushButton("Browse...")
        browse_template_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['templates_edit'], "Select Template File or Directory"))
        template_file_layout.addWidget(browse_template_btn)
        templates_layout.addRow("Templates (-t):", template_file_layout)

        controls['severity_combo'] = QComboBox()
        controls['severity_combo'].addItems(["all", "info", "low", "medium", "high", "critical"])
        controls['severity_combo'].setToolTip("Filter templates by severity (-s).")
        templates_layout.addRow("Severity (-s):", controls['severity_combo'])

        controls['tags_edit'] = QLineEdit()
        controls['tags_edit'].setPlaceholderText("e.g., cve,rce,wordpress")
        controls['tags_edit'].setToolTip("Filter templates by tags (-tags).")
        templates_layout.addRow("Tags (-tags):", controls['tags_edit'])

        nuclei_tabs.addTab(templates_tab, "Templates")

        # --- Output & Config Tab ---
        config_tab = QWidget()
        config_layout = QFormLayout(config_tab)

        output_file_layout = QHBoxLayout()
        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("Optional: path to save report...")
        output_file_layout.addWidget(controls['output_edit'])
        browse_output_btn = QPushButton("Browse...")
        browse_output_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_edit'], "Save Nuclei Report"))
        output_file_layout.addWidget(browse_output_btn)
        config_layout.addRow("Output File (-o):", output_file_layout)

        controls['concurrency_edit'] = QLineEdit("25")
        controls['concurrency_edit'].setToolTip("Number of concurrent requests to send (-c).")
        config_layout.addRow("Concurrency (-c):", controls['concurrency_edit'])

        controls['ratelimit_edit'] = QLineEdit("150")
        controls['ratelimit_edit'].setToolTip("Requests per second (-rl).")
        config_layout.addRow("Rate Limit (-rl):", controls['ratelimit_edit'])

        controls['verbose_check'] = QCheckBox("Verbose Output (-v)")
        config_layout.addRow(controls['verbose_check'])

        nuclei_tabs.addTab(config_tab, "Configuration")

        # --- Action Buttons ---
        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Nuclei Scan")
        controls['stop_btn'] = QPushButton("Stop Nuclei"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_nuclei_scan(self):
        """Starts the Nuclei scan worker thread."""
        controls = self.nuclei_controls
        if not shutil.which("nuclei"):
            QMessageBox.critical(self, "Nuclei Error", "'nuclei' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = controls['target_edit'].text().strip()
        if not target:
            QMessageBox.critical(self, "Input Error", "A target or target list is required.")
            return

        # Determine if target is a file or a URL
        command = ["nuclei"]
        if os.path.exists(target):
            command.extend(["-l", target])
        else:
            command.extend(["-u", target])

        if templates := controls['templates_edit'].text().strip():
            command.extend(["-t", templates])

        if (severity := controls['severity_combo'].currentText()) != "all":
            command.extend(["-s", severity])

        if tags := controls['tags_edit'].text().strip():
            command.extend(["-tags", tags])

        if output_file := controls['output_edit'].text().strip():
            command.extend(["-o", output_file])

        if concurrency := controls['concurrency_edit'].text().strip():
            command.extend(["-c", concurrency])

        if ratelimit := controls['ratelimit_edit'].text().strip():
            command.extend(["-rl", ratelimit])

        if controls['verbose_check'].isChecked():
            command.append("-v")

        command.extend(["-nC", "-json"]) # Always disable color and enable JSON for parsing

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.nuclei_output_console.clear()

        self.worker = WorkerThread(self._nuclei_thread, args=(command, target))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _nuclei_thread(self, command, target):
        """Worker thread for running the nuclei command."""
        q = self.tool_results_queue
        logging.info(f"Starting Nuclei with command: {' '.join(command)}")
        q.put(('nuclei_output', f"$ {' '.join(command)}\n\n"))
        json_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.nuclei_process = process

            full_output = []
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('nuclei_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                # Don't send JSON to the console, just capture it
                full_output.append(line)

            process.stdout.close()
            process.wait()

            json_data = "".join(full_output)
            q.put(('nuclei_output', json_data)) # Put the full JSON blob in the console
            if not self.tool_stop_event.is_set() and json_data.strip():
                q.put(('nuclei_results', json_data))


        except FileNotFoundError:
            q.put(('error', 'Nuclei Error', "'nuclei' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Nuclei thread error: {e}", exc_info=True)
            q.put(('error', 'Nuclei Error', str(e)))
        finally:
            q.put(('tool_finished', 'nuclei_scan', target, json_data))
            with self.thread_finish_lock:
                self.nuclei_process = None
            logging.info("Nuclei scan thread finished.")

    def _handle_nuclei_output(self, line):
        self.nuclei_output_console.insertPlainText(line)
        self.nuclei_output_console.verticalScrollBar().setValue(self.nuclei_output_console.verticalScrollBar().maximum())

    def _create_trufflehog_tool(self):
        """Creates the UI for the TruffleHog Secret Scanner tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.trufflehog_controls = self._create_trufflehog_config_widget()
        main_layout.addWidget(config_widget)

        controls = self.trufflehog_controls

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(controls['start_btn'])
        buttons_layout.addWidget(controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.trufflehog_output_console = QPlainTextEdit()
        self.trufflehog_output_console.setReadOnly(True)
        self.trufflehog_output_console.setFont(QFont("Courier New", 10))
        self.trufflehog_output_console.setPlaceholderText("TruffleHog output will be displayed here...")
        main_layout.addWidget(self.trufflehog_output_console, 1)

        controls['start_btn'].clicked.connect(self.start_trufflehog_scan)
        controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_trufflehog_config_widget(self):
        """Creates a reusable, self-contained widget with TruffleHog's configuration options."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)
        main_layout.setContentsMargins(0,0,0,0)

        controls = {}

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>TruffleHog - Secret Scanner</b></font>
        <p>This tool scans sources like git repositories, GitHub, and filesystems for leaked secrets.</p>
        """)
        instructions.setFixedHeight(80)
        main_layout.addWidget(instructions)

        # --- Source Type and Target ---
        source_box = QGroupBox("Scan Target")
        source_layout = QVBoxLayout(source_box)

        controls['source_type_group'] = QButtonGroup(self)
        rb_layout = QHBoxLayout()
        controls['git_rb'] = QRadioButton("Git Repo"); controls['git_rb'].setChecked(True)
        controls['github_rb'] = QRadioButton("GitHub")
        controls['filesystem_rb'] = QRadioButton("Filesystem")
        controls['source_type_group'].addButton(controls['git_rb'])
        controls['source_type_group'].addButton(controls['github_rb'])
        controls['source_type_group'].addButton(controls['filesystem_rb'])
        rb_layout.addWidget(controls['git_rb']); rb_layout.addWidget(controls['github_rb']); rb_layout.addWidget(controls['filesystem_rb'])
        source_layout.addLayout(rb_layout)

        target_layout = QHBoxLayout()
        controls['target_edit'] = QLineEdit()
        controls['target_edit'].setPlaceholderText("Enter Git URL, GitHub repo, or filesystem path...")
        target_layout.addWidget(controls['target_edit'])
        controls['browse_btn'] = QPushButton("Browse...")
        controls['browse_btn'].clicked.connect(lambda: self._browse_dir_for_lineedit(controls['target_edit'], "Select Directory to Scan"))
        target_layout.addWidget(controls['browse_btn'])
        source_layout.addLayout(target_layout)

        main_layout.addWidget(source_box)

        # --- Options ---
        options_box = QGroupBox("Options")
        options_layout = QFormLayout(options_box)
        controls['only_verified_check'] = QCheckBox("Only Verified Results")
        controls['only_verified_check'].setToolTip("Only output secrets that have been successfully verified against their respective APIs.")
        options_layout.addRow(controls['only_verified_check'])
        main_layout.addWidget(options_box)

        # --- Action Buttons ---
        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        # UI Logic to show/hide browse button
        def toggle_browse_button():
            controls['browse_btn'].setVisible(controls['filesystem_rb'].isChecked())
        controls['source_type_group'].buttonClicked.connect(toggle_browse_button)
        toggle_browse_button() # Set initial state

        return widget, controls

    def _browse_dir_for_lineedit(self, line_edit_widget, dialog_title):
        dir_path = QFileDialog.getExistingDirectory(self, dialog_title, options=QFileDialog.Option.DontUseNativeDialog)
        if dir_path:
            line_edit_widget.setText(dir_path)

    def start_trufflehog_scan(self):
        """Starts the TruffleHog scan worker thread."""
        controls = self.trufflehog_controls
        if not shutil.which("trufflehog"):
            QMessageBox.critical(self, "TruffleHog Error", "'trufflehog' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = controls['target_edit'].text().strip()
        if not target:
            QMessageBox.critical(self, "Input Error", "A target is required.")
            return

        source_type = ""
        if controls['git_rb'].isChecked(): source_type = "git"
        elif controls['github_rb'].isChecked(): source_type = "github"
        elif controls['filesystem_rb'].isChecked(): source_type = "filesystem"

        command = ["trufflehog", source_type, target]

        if controls['only_verified_check'].isChecked():
            command.append("--only-verified")

        # Always add --json for parsing, but hide the checkbox from the user to avoid confusion
        command.append("--json")

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.trufflehog_output_console.clear()

        self.worker = WorkerThread(self._trufflehog_thread, args=(command, target))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _trufflehog_thread(self, command, target):
        """Worker thread for running the trufflehog command."""
        q = self.tool_results_queue
        logging.info(f"Starting TruffleHog with command: {' '.join(command)}")
        q.put(('trufflehog_output', f"$ {' '.join(command)}\n\n"))
        json_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.trufflehog_process = process

            full_output = []
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('trufflehog_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                # Don't send JSON to the console, just capture it
                full_output.append(line)

            process.stdout.close()
            process.wait()

            json_data = "".join(full_output)
            q.put(('trufflehog_output', json_data))
            if not self.tool_stop_event.is_set() and json_data.strip():
                q.put(('trufflehog_results', json_data))

        except FileNotFoundError:
            q.put(('error', 'TruffleHog Error', "'trufflehog' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"TruffleHog thread error: {e}", exc_info=True)
            q.put(('error', 'TruffleHog Error', str(e)))
        finally:
            q.put(('tool_finished', 'trufflehog_scan', target, json_data))
            with self.thread_finish_lock:
                self.trufflehog_process = None
            logging.info("TruffleHog scan thread finished.")

    def _handle_trufflehog_output(self, line):
        self.trufflehog_output_console.insertPlainText(line)
        self.trufflehog_output_console.verticalScrollBar().setValue(self.trufflehog_output_console.verticalScrollBar().maximum())

    def _create_jtr_tool(self):
        """Creates the UI for the John the Ripper tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.jtr_controls = self._create_jtr_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.jtr_controls['start_btn'])
        buttons_layout.addWidget(self.jtr_controls['show_btn'])
        buttons_layout.addWidget(self.jtr_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.jtr_output_console = QPlainTextEdit()
        self.jtr_output_console.setReadOnly(True)
        self.jtr_output_console.setFont(QFont("Courier New", 10))
        self.jtr_output_console.setPlaceholderText("John the Ripper output will be displayed here...")
        main_layout.addWidget(self.jtr_output_console, 1)

        self.jtr_controls['start_btn'].clicked.connect(self.start_jtr_crack)
        self.jtr_controls['show_btn'].clicked.connect(self.show_jtr_cracked)
        self.jtr_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_jtr_config_widget(self):
        """Creates a reusable, self-contained widget with JTR's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        hash_file_layout = QHBoxLayout()
        controls['hash_file_edit'] = QLineEdit()
        controls['hash_file_edit'].setPlaceholderText("Path to hash file (required)...")
        hash_file_layout.addWidget(controls['hash_file_edit'])
        browse_hash_btn = QPushButton("Browse...")
        browse_hash_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['hash_file_edit'], "Select Hash File"))
        hash_file_layout.addWidget(browse_hash_btn)
        main_layout.addRow("Hash File:", hash_file_layout)

        wordlist_layout = QHBoxLayout()
        controls['wordlist_edit'] = QLineEdit()
        controls['wordlist_edit'].setPlaceholderText("Path to wordlist file (optional)...")
        wordlist_layout.addWidget(controls['wordlist_edit'])
        browse_wordlist_btn = QPushButton("Browse...")
        browse_wordlist_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['wordlist_edit'], "Select Wordlist File"))
        wordlist_layout.addWidget(browse_wordlist_btn)
        main_layout.addRow("Wordlist:", wordlist_layout)

        controls['format_edit'] = QLineEdit()
        controls['format_edit'].setPlaceholderText("e.g., raw-md5, nt, sha512crypt")
        controls['format_edit'].setToolTip("Specify the hash format (--format). Leave blank for auto-detection.")
        main_layout.addRow("Format:", controls['format_edit'])

        controls['rules_check'] = QCheckBox("Enable Word Mangling Rules")
        controls['rules_check'].setToolTip("Enable rules for wordlist mode (--rules).")
        main_layout.addRow(controls['rules_check'])

        controls['incremental_check'] = QCheckBox("Incremental Mode (Brute-force)")
        controls['incremental_check'].setToolTip("Enable incremental mode (--incremental). If wordlist is also specified, this will run after.")
        main_layout.addRow(controls['incremental_check'])

        controls['start_btn'] = QPushButton(QIcon("icons/tool.svg"), " Start Cracking")
        controls['show_btn'] = QPushButton("Show Cracked")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_jtr_crack(self):
        """Starts the John the Ripper worker thread."""
        self._start_jtr_generic(crack_mode=True)

    def show_jtr_cracked(self):
        """Shows already cracked passwords."""
        self._start_jtr_generic(crack_mode=False)

    def _start_jtr_generic(self, crack_mode):
        controls = self.jtr_controls
        if not shutil.which("john"):
            QMessageBox.critical(self, "JTR Error", "'john' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        hash_file = controls['hash_file_edit'].text().strip()
        if not hash_file:
            QMessageBox.critical(self, "Input Error", "A hash file is required.")
            return

        command = ["john"]

        if crack_mode:
            wordlist = controls['wordlist_edit'].text().strip()
            if wordlist:
                command.extend([f"--wordlist={wordlist}"])
                if controls['rules_check'].isChecked():
                    command.append("--rules")
            if controls['incremental_check'].isChecked():
                command.append("--incremental")
        else: # Show mode
            command.append("--show")

        if format_type := controls['format_edit'].text().strip():
            command.extend([f"--format={format_type}"])

        command.append(hash_file)

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['show_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.jtr_output_console.clear()

        self.worker = WorkerThread(self._jtr_thread, args=(command, hash_file))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _jtr_thread(self, command, hash_file):
        """Worker thread for running the john command."""
        q = self.tool_results_queue
        logging.info(f"Starting JTR with command: {' '.join(command)}")
        q.put(('jtr_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.jtr_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('jtr_output', "\n\n--- Canceled By User ---\n"))
                    break
                q.put(('jtr_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'JTR Error', "'john' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"JTR thread error: {e}", exc_info=True)
            q.put(('error', 'JTR Error', str(e)))
        finally:
            q.put(('tool_finished', 'jtr_scan', hash_file, "".join(full_output)))
            with self.thread_finish_lock:
                self.jtr_process = None
            logging.info("JTR scan thread finished.")

    def _handle_jtr_output(self, line):
        self.jtr_output_console.insertPlainText(line)
        self.jtr_output_console.verticalScrollBar().setValue(self.jtr_output_console.verticalScrollBar().maximum())

    def _create_hydra_tool(self):
        """Creates the UI for the Hydra network logon cracker."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.hydra_controls = self._create_hydra_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.hydra_controls['start_btn'])
        buttons_layout.addWidget(self.hydra_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.hydra_output_console = QPlainTextEdit()
        self.hydra_output_console.setReadOnly(True)
        self.hydra_output_console.setFont(QFont("Courier New", 10))
        self.hydra_output_console.setPlaceholderText("Hydra output will be displayed here...")
        main_layout.addWidget(self.hydra_output_console, 1)

        self.hydra_controls['start_btn'].clicked.connect(self.start_hydra_attack)
        self.hydra_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_hydra_config_widget(self):
        """Creates a reusable, self-contained widget with Hydra's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['target_edit'] = QLineEdit("localhost")
        main_layout.addRow("Target Host:", controls['target_edit'])

        controls['service_edit'] = QLineEdit("ssh")
        controls['service_edit'].setToolTip("The service to attack (e.g., ssh, ftp, smb, rdp).")
        main_layout.addRow("Service:", controls['service_edit'])

        user_layout = QHBoxLayout()
        controls['user_edit'] = QLineEdit("root")
        controls['user_edit'].setToolTip("A single username to test (-l).")
        user_layout.addWidget(controls['user_edit'])
        user_layout.addWidget(QLabel("OR"))
        controls['user_list_edit'] = QLineEdit()
        controls['user_list_edit'].setToolTip("Path to a file containing a list of usernames (-L).")
        user_layout.addWidget(controls['user_list_edit'])
        browse_user_btn = QPushButton("Browse...")
        browse_user_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['user_list_edit'], "Select User List"))
        user_layout.addWidget(browse_user_btn)
        main_layout.addRow("Username (-l) / List (-L):", user_layout)

        pass_layout = QHBoxLayout()
        controls['pass_edit'] = QLineEdit()
        controls['pass_edit'].setToolTip("A single password to test (-p).")
        pass_layout.addWidget(controls['pass_edit'])
        pass_layout.addWidget(QLabel("OR"))
        controls['pass_list_edit'] = QLineEdit()
        controls['pass_list_edit'].setToolTip("Path to a file containing a list of passwords (-P).")
        pass_layout.addWidget(controls['pass_list_edit'])
        browse_pass_btn = QPushButton("Browse...")
        browse_pass_btn.clicked.connect(lambda: self._browse_file_for_lineedit(controls['pass_list_edit'], "Select Password List"))
        pass_layout.addWidget(browse_pass_btn)
        main_layout.addRow("Password (-p) / List (-P):", pass_layout)

        controls['tasks_edit'] = QLineEdit("16")
        controls['tasks_edit'].setToolTip("Number of parallel tasks/threads (-t).")
        main_layout.addRow("Tasks (-t):", controls['tasks_edit'])

        controls['start_btn'] = QPushButton(QIcon("icons/tool.svg"), " Start Attack")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_hydra_attack(self):
        """Starts the Hydra attack worker thread."""
        controls = self.hydra_controls
        if not shutil.which("hydra"):
            QMessageBox.critical(self, "Hydra Error", "'hydra' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = controls['target_edit'].text().strip()
        service = controls['service_edit'].text().strip()
        if not target or not service:
            QMessageBox.critical(self, "Input Error", "Target and Service are required.")
            return

        command = ["hydra"]

        if user := controls['user_edit'].text().strip():
            command.extend(["-l", user])
        elif user_list := controls['user_list_edit'].text().strip():
            command.extend(["-L", user_list])
        else:
            QMessageBox.critical(self, "Input Error", "A username or user list is required.")
            return

        if pwd := controls['pass_edit'].text().strip():
            command.extend(["-p", pwd])
        elif pass_list := controls['pass_list_edit'].text().strip():
            command.extend(["-P", pass_list])
        else:
            QMessageBox.critical(self, "Input Error", "A password or password list is required.")
            return

        if tasks := controls['tasks_edit'].text().strip():
            command.extend(["-t", tasks])

        command.append(f"{service}://{target}")

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.hydra_output_console.clear()

        target_str = f"{service}://{target}"
        self.worker = WorkerThread(self._hydra_thread, args=(command, target_str))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _hydra_thread(self, command, target):
        """Worker thread for running the hydra command."""
        q = self.tool_results_queue
        logging.info(f"Starting Hydra with command: {' '.join(command)}")
        q.put(('hydra_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.hydra_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('hydra_output', "\n\n--- Canceled By User ---\n"))
                    break
                q.put(('hydra_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Hydra Error', "'hydra' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Hydra thread error: {e}", exc_info=True)
            q.put(('error', 'Hydra Error', str(e)))
        finally:
            q.put(('tool_finished', 'hydra_scan', target, "".join(full_output)))
            with self.thread_finish_lock:
                self.hydra_process = None
            logging.info("Hydra scan thread finished.")

    def _handle_hydra_output(self, line):
        self.hydra_output_console.insertPlainText(line)
        self.hydra_output_console.verticalScrollBar().setValue(self.hydra_output_console.verticalScrollBar().maximum())

    def _create_sherlock_tool(self):
        """Creates the UI for the Sherlock username scanner."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.sherlock_controls = self._create_sherlock_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.sherlock_controls['start_btn'])
        buttons_layout.addWidget(self.sherlock_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.sherlock_output_console = QPlainTextEdit()
        self.sherlock_output_console.setReadOnly(True)
        self.sherlock_output_console.setFont(QFont("Courier New", 10))
        self.sherlock_output_console.setPlaceholderText("Sherlock output will be displayed here...")
        main_layout.addWidget(self.sherlock_output_console, 1)

        self.sherlock_controls['start_btn'].clicked.connect(self.start_sherlock_scan)
        self.sherlock_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_sherlock_config_widget(self):
        """Creates a reusable, self-contained widget with Sherlock's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['usernames_edit'] = QLineEdit()
        controls['usernames_edit'].setToolTip("One or more usernames to check, separated by spaces.")
        main_layout.addRow("Usernames:", controls['usernames_edit'])

        controls['timeout_edit'] = QLineEdit("60")
        controls['timeout_edit'].setToolTip("Timeout in seconds for each request.")
        main_layout.addRow("Timeout (--timeout):", controls['timeout_edit'])

        output_file_layout = QHBoxLayout()
        controls['output_edit'] = QLineEdit()
        controls['output_edit'].setPlaceholderText("Optional: path to save text report...")
        output_file_layout.addWidget(controls['output_edit'])
        browse_output_btn = QPushButton("Browse...")
        browse_output_btn.clicked.connect(lambda: self._browse_save_file_for_lineedit(controls['output_edit'], "Save Sherlock Report"))
        output_file_layout.addWidget(browse_output_btn)
        main_layout.addRow("Output File (-o):", output_file_layout)

        controls['csv_check'] = QCheckBox("Export as CSV")
        main_layout.addRow("--csv:", controls['csv_check'])

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Hunt Usernames")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_sherlock_scan(self):
        """Starts the Sherlock scan worker thread."""
        controls = self.sherlock_controls
        if not shutil.which("sherlock"):
            QMessageBox.critical(self, "Sherlock Error", "'sherlock' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        usernames = controls['usernames_edit'].text().strip()
        if not usernames:
            QMessageBox.critical(self, "Input Error", "At least one username is required.")
            return

        command = ["sherlock", "--no-color"]

        if timeout := controls['timeout_edit'].text().strip():
            command.extend(["--timeout", timeout])

        try:
            # Create a temporary directory for sherlock to save its files
            self.sherlock_temp_dir = tempfile.mkdtemp()
            # Sherlock saves as <username>.csv in the specified folder
            # We don't know the exact filename beforehand if multiple users are scanned
            command.extend(["-fo", self.sherlock_temp_dir, "--csv"])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary directory for Sherlock report: {e}")
            return

        command.extend(usernames.split())

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.sherlock_output_console.clear()

        self.worker = WorkerThread(self._sherlock_thread, args=(command, usernames))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _sherlock_thread(self, command, usernames):
        """Worker thread for running the sherlock command."""
        q = self.tool_results_queue
        logging.info(f"Starting Sherlock with command: {' '.join(command)}")
        q.put(('sherlock_output', f"$ {' '.join(command)}\n\n"))
        csv_data = ""

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.sherlock_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('sherlock_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('sherlock_output', line))

            process.stdout.close()
            process.wait()

            if not self.tool_stop_event.is_set():
                try:
                    # Find the first CSV file in the temp directory
                    for filename in os.listdir(self.sherlock_temp_dir):
                        if filename.endswith(".csv"):
                            with open(os.path.join(self.sherlock_temp_dir, filename), 'r', encoding='utf-8') as f:
                                csv_data = f.read()
                            break # Just read the first one for now
                    if csv_data:
                        q.put(('sherlock_results', csv_data, usernames))
                except Exception as e:
                    logging.error(f"Could not read Sherlock CSV report: {e}")
                finally:
                    shutil.rmtree(self.sherlock_temp_dir)
                    self.sherlock_temp_dir = None


        except FileNotFoundError:
            q.put(('error', 'Sherlock Error', "'sherlock' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Sherlock thread error: {e}", exc_info=True)
            q.put(('error', 'Sherlock Error', str(e)))
        finally:
            q.put(('tool_finished', 'sherlock_scan', usernames, csv_data))
            with self.thread_finish_lock:
                self.sherlock_process = None
            logging.info("Sherlock scan thread finished.")

    def _handle_sherlock_output(self, line):
        self.sherlock_output_console.insertPlainText(line)
        self.sherlock_output_console.verticalScrollBar().setValue(self.sherlock_output_console.verticalScrollBar().maximum())

    def _create_spiderfoot_tool(self):
        """Creates the UI for the Spiderfoot OSINT tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        config_widget, self.spiderfoot_controls = self._create_spiderfoot_config_widget()
        main_layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.spiderfoot_controls['start_btn'])
        buttons_layout.addWidget(self.spiderfoot_controls['stop_btn'])
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.spiderfoot_output_console = QPlainTextEdit()
        self.spiderfoot_output_console.setReadOnly(True)
        self.spiderfoot_output_console.setFont(QFont("Courier New", 10))
        self.spiderfoot_output_console.setPlaceholderText("Spiderfoot output will be displayed here...")
        main_layout.addWidget(self.spiderfoot_output_console, 1)

        self.spiderfoot_controls['start_btn'].clicked.connect(self.start_spiderfoot_scan)
        self.spiderfoot_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_spiderfoot_config_widget(self):
        """Creates a reusable, self-contained widget with Spiderfoot's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['target_edit'] = QLineEdit("example.com")
        main_layout.addRow("Target (-s):", controls['target_edit'])

        controls['types_edit'] = QLineEdit("EMAILADDR,DNS_MX,PHONE_NUMBER")
        controls['types_edit'].setToolTip("Comma-separated list of data types to collect (e.g., EMAILADDR, PHONE_NUMBER).")
        main_layout.addRow("Scan Types (-t):", controls['types_edit'])

        controls['modules_edit'] = QLineEdit()
        controls['modules_edit'].setPlaceholderText("Optional: e.g., sfp_dns,sfp_email")
        controls['modules_edit'].setToolTip("Comma-separated list of specific modules to run (-m).")
        main_layout.addRow("Modules (-m):", controls['modules_edit'])

        controls['silent_check'] = QCheckBox("Silent Output")
        controls['silent_check'].setToolTip("Only report errors (-q).")
        main_layout.addRow("-q:", controls['silent_check'])

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_spiderfoot_scan(self):
        """Starts the Spiderfoot scan worker thread."""
        controls = self.spiderfoot_controls
        if not shutil.which("spiderfoot-cli"):
            QMessageBox.critical(self, "Spiderfoot Error", "'spiderfoot-cli' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = controls['target_edit'].text().strip()
        if not target:
            QMessageBox.critical(self, "Input Error", "A target is required.")
            return

        command = ["spiderfoot-cli", "-s", target]

        if types := controls['types_edit'].text().strip():
            command.extend(["-t", types])
        if modules := controls['modules_edit'].text().strip():
            command.extend(["-m", modules])
        if controls['silent_check'].isChecked():
            command.append("-q")

        command.append("-n") # Disable history logging for non-interactive use

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.spiderfoot_output_console.clear()

        self.worker = WorkerThread(self._spiderfoot_thread, args=(command, target))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _spiderfoot_thread(self, command, target):
        """Worker thread for running the spiderfoot-cli command."""
        q = self.tool_results_queue
        logging.info(f"Starting Spiderfoot with command: {' '.join(command)}")
        q.put(('spiderfoot_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.spiderfoot_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('spiderfoot_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('spiderfoot_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Spiderfoot Error', "'spiderfoot-cli' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Spiderfoot thread error: {e}", exc_info=True)
            q.put(('error', 'Spiderfoot Error', str(e)))
        finally:
            q.put(('tool_finished', 'spiderfoot_scan', target, "".join(full_output)))
            with self.thread_finish_lock:
                self.spiderfoot_process = None
            logging.info("Spiderfoot scan thread finished.")

    def _handle_spiderfoot_output(self, line):
        self.spiderfoot_output_console.insertPlainText(line)
        self.spiderfoot_output_console.verticalScrollBar().setValue(self.spiderfoot_output_console.verticalScrollBar().maximum())

    def _create_masscan_tool(self):
        """Creates the UI for the Masscan tool."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Masscan: The Mass IP Port Scanner</b></font>
        <p>This is an Internet-scale port scanner, capable of scanning the entire Internet in under 5 minutes.</p>
        <p><b>WARNING:</b> Scanning at high rates can cause network disruption and may be detected by network administrators. Use responsibly.</p>
        """)
        instructions.setFixedHeight(100)
        main_layout.addWidget(instructions)

        # --- Controls ---
        controls_frame = QGroupBox("Scan Options")
        controls_layout = QFormLayout(controls_frame)

        self.masscan_target_edit = QLineEdit("0.0.0.0/0")
        self.masscan_target_edit.setToolTip("Enter target IP ranges (e.g., 10.0.0.0/8, 192.168.0.1-192.168.0.254).")
        controls_layout.addRow("Target(s):", self.masscan_target_edit)

        ports_layout = QHBoxLayout()
        self.masscan_ports_edit = QLineEdit("0-65535")
        self.masscan_ports_edit.setToolTip("Specify ports to scan (e.g., 80,443, 0-65535).")
        ports_layout.addWidget(self.masscan_ports_edit)
        common_ports_btn = QPushButton("Common Ports")
        common_ports_btn.clicked.connect(lambda: self.masscan_ports_edit.setText("21,22,23,25,53,80,110,111,135,139,143,443,445,993,995,1723,3306,3389,5900,8080"))
        ports_layout.addWidget(common_ports_btn)
        controls_layout.addRow("Ports:", ports_layout)

        self.masscan_rate_edit = QLineEdit("1000")
        self.masscan_rate_edit.setToolTip("Set the transmission rate in packets/second.")
        controls_layout.addRow("Rate (--rate):", self.masscan_rate_edit)

        outfile_layout = QHBoxLayout()
        self.masscan_outfile_edit = QLineEdit()
        self.masscan_outfile_edit.setPlaceholderText("Optional: path to save report...")
        outfile_layout.addWidget(self.masscan_outfile_edit)
        browse_out_btn = QPushButton("Browse...")
        browse_out_btn.clicked.connect(self.browse_masscan_outfile)
        outfile_layout.addWidget(browse_out_btn)
        controls_layout.addRow("Output File (-oJ):", outfile_layout) # Default to JSON for easy parsing

        self.masscan_extra_opts_edit = QLineEdit()
        self.masscan_extra_opts_edit.setToolTip("Enter any additional, space-separated Masscan flags here.")
        controls_layout.addRow("Additional Options:", self.masscan_extra_opts_edit)

        main_layout.addWidget(controls_frame)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        self.masscan_start_btn = QPushButton(QIcon("icons/search.svg"), " Start Masscan")
        self.masscan_stop_btn = QPushButton("Stop Masscan"); self.masscan_stop_btn.setEnabled(False)
        buttons_layout.addWidget(self.masscan_start_btn)
        buttons_layout.addWidget(self.masscan_stop_btn)
        main_layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.masscan_output_console = QPlainTextEdit()
        self.masscan_output_console.setReadOnly(True)
        self.masscan_output_console.setFont(QFont("Courier New", 10))
        self.masscan_output_console.setPlaceholderText("Masscan output will be displayed here... (Note: Masscan primarily outputs to stderr)")
        main_layout.addWidget(self.masscan_output_console, 1)

        self.masscan_start_btn.clicked.connect(self.start_masscan_scan)
        self.masscan_stop_btn.clicked.connect(self.cancel_tool)

        return widget

    def browse_masscan_outfile(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Masscan Report", "", "JSON files (*.json);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.masscan_outfile_edit.setText(file_path)

    def start_masscan_scan(self):
        """Starts the Masscan worker thread."""
        if not shutil.which("masscan"):
            QMessageBox.critical(self, "Masscan Error", "'masscan' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = self.masscan_target_edit.text().strip()
        ports = self.masscan_ports_edit.text().strip()
        rate = self.masscan_rate_edit.text().strip()

        if not target or not ports or not rate:
            QMessageBox.critical(self, "Input Error", "Target, Ports, and Rate are required.")
            return

        is_root = True
        try:
            is_root = (os.geteuid() == 0)
        except AttributeError: # os.geteuid() does not exist on Windows
            is_root = False

        command_prefix = []
        if not is_root and sys.platform != "win32":
            command_prefix = ["sudo"]
            QMessageBox.warning(self, "Permissions", "Masscan may require root privileges to run correctly. Attempting with 'sudo'. You may be prompted for a password in the terminal where GScapy was launched.")


        command = command_prefix + ["masscan", target, "-p", ports, "--rate", rate]

        if outfile := self.masscan_outfile_edit.text().strip():
            command.extend(["-oJ", outfile]) # JSON output format

        if extra_opts := self.masscan_extra_opts_edit.text().strip():
            command.extend(extra_opts.split())

        self.is_tool_running = True
        self.masscan_start_btn.setEnabled(False)
        self.masscan_stop_btn.setEnabled(True)
        self.tool_stop_event.clear()
        self.masscan_output_console.clear()

        self.worker = WorkerThread(self._masscan_thread, args=(command, target))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _masscan_thread(self, command, target):
        """Worker thread for running the masscan command."""
        q = self.tool_results_queue
        logging.info(f"Starting Masscan with command: {' '.join(command)}")
        q.put(('masscan_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            # Masscan outputs progress to stderr, so we need to capture both
            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.masscan_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('masscan_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('masscan_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Masscan Error', "'masscan' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Masscan thread error: {e}", exc_info=True)
            q.put(('error', 'Masscan Error', str(e)))
        finally:
            q.put(('tool_finished', 'masscan_scan', target, "".join(full_output)))
            with self.thread_finish_lock:
                self.masscan_process = None
            logging.info("Masscan scan thread finished.")

    def _whatweb_thread(self, command):
        """Worker thread for running the whatweb command."""
        q = self.tool_results_queue
        logging.info(f"Starting WhatWeb with command: {' '.join(command)}")
        q.put(('whatweb_output', f"$ {' '.join(command)}\n\n"))

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.whatweb_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('whatweb_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('whatweb_output', line))

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'WhatWeb Error', "'whatweb' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"WhatWeb thread error: {e}", exc_info=True)
            q.put(('error', 'WhatWeb Error', str(e)))
        finally:
            q.put(('tool_finished', 'whatweb_scan'))
            with self.thread_finish_lock:
                self.whatweb_process = None
            logging.info("WhatWeb scan thread finished.")

    def _gobuster_thread(self, command, url):
        """Worker thread for running the gobuster command."""
        q = self.tool_results_queue
        logging.info(f"Starting Gobuster with command: {' '.join(command)}")
        q.put(('gobuster_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.gobuster_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('gobuster_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('gobuster_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Gobuster Error', "'gobuster' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Gobuster thread error: {e}", exc_info=True)
            q.put(('error', 'Gobuster Error', str(e)))
        finally:
            q.put(('tool_finished', 'gobuster_scan', url, "".join(full_output)))
            with self.thread_finish_lock:
                self.gobuster_process = None
            logging.info("Gobuster scan thread finished.")

    def _nikto_thread(self, command, target):
        """Worker thread for running the nikto command."""
        q = self.tool_results_queue
        logging.info(f"Starting Nikto with command: {' '.join(command)}")
        q.put(('nikto_output', f"$ {' '.join(command)}\n\n"))
        full_output = []

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.nikto_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('nikto_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('nikto_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Nikto Error', "'nikto' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"Nikto thread error: {e}", exc_info=True)
            q.put(('error', 'Nikto Error', str(e)))
        finally:
            q.put(('tool_finished', 'nikto_scan', target, "".join(full_output)))
            with self.thread_finish_lock:
                self.nikto_process = None
            logging.info("Nikto scan thread finished.")

    def _create_traceroute_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        self.trace_tree = QTreeWidget(); self.trace_tree.setColumnCount(4); self.trace_tree.setHeaderLabels(["Hop", "IP Address", "Host Name", "Time (ms)"])

        controls = QHBoxLayout()
        controls.addWidget(QLabel("Target:")); self.trace_target = QLineEdit("google.com"); controls.addWidget(self.trace_target)
        self.trace_button = QPushButton("Trace"); controls.addWidget(self.trace_button)
        self.trace_cancel_button = QPushButton("Cancel"); self.trace_cancel_button.setEnabled(False); controls.addWidget(self.trace_cancel_button)
        self.trace_status = QLabel(""); controls.addWidget(self.trace_status); controls.addStretch()

        layout.addLayout(controls)
        layout.addWidget(self.trace_tree)
        layout.addWidget(self._create_export_button(self.trace_tree))
        self.trace_button.clicked.connect(self.start_traceroute)
        self.trace_cancel_button.clicked.connect(self.cancel_tool)
        return widget

    def _update_tcp_scan_options_visibility(self, checked):
        """Shows or hides the TCP scan mode dropdown based on protocol selection."""
        is_tcp_selected = self.scan_proto_tcp_radio.isChecked() or self.scan_proto_both_radio.isChecked()
        self.tcp_scan_type_label.setVisible(is_tcp_selected)
        self.tcp_scan_type_combo.setVisible(is_tcp_selected)

    def _create_port_scanner_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget); controls = QFrame(); clayout = QVBoxLayout(controls)
        row1 = QHBoxLayout(); row1.addWidget(QLabel("Target:")); self.scan_target = QLineEdit("127.0.0.1"); self.scan_target.setToolTip("The IP address of the target machine.")
        row1.addWidget(self.scan_target)
        clayout.addLayout(row1)
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Ports:")); self.scan_ports = QLineEdit("22,80,443"); self.scan_ports.setToolTip("A comma-separated list of ports or port ranges (e.g., 22,80,100-200).")
        row2.addWidget(self.scan_ports)
        all_ports_btn = QPushButton("All"); all_ports_btn.setToolTip("Set the port range to all 65535 ports.")
        all_ports_btn.clicked.connect(lambda: self.scan_ports.setText("1-65535")); row2.addWidget(all_ports_btn)
        clayout.addLayout(row2)

        # Row 3: Protocol Type Radio Buttons
        row3 = QHBoxLayout()
        row3.addWidget(QLabel("Protocol:"))
        self.scan_proto_tcp_radio = QRadioButton("TCP"); self.scan_proto_tcp_radio.setChecked(True)
        self.scan_proto_udp_radio = QRadioButton("UDP")
        self.scan_proto_both_radio = QRadioButton("Both")
        self.scan_proto_group = QButtonGroup(self)
        self.scan_proto_group.addButton(self.scan_proto_tcp_radio)
        self.scan_proto_group.addButton(self.scan_proto_udp_radio)
        self.scan_proto_group.addButton(self.scan_proto_both_radio)
        row3.addWidget(self.scan_proto_tcp_radio)
        row3.addWidget(self.scan_proto_udp_radio)
        row3.addWidget(self.scan_proto_both_radio)
        row3.addStretch()
        clayout.addLayout(row3)

        # Row 4: Advanced TCP Scan Options
        row4 = QHBoxLayout()
        self.tcp_scan_type_label = QLabel("TCP Scan Mode:")
        row4.addWidget(self.tcp_scan_type_label)
        self.tcp_scan_type_combo = QComboBox()
        self.tcp_scan_type_combo.addItems(["SYN Scan", "FIN Scan", "Xmas Scan", "Null Scan", "ACK Scan"])
        self.tcp_scan_type_combo.setToolTip("Select the type of TCP scan to perform for firewall evasion.")
        row4.addWidget(self.tcp_scan_type_combo)
        row4.addStretch()
        self.scan_frag_check = QCheckBox("Use Fragments"); self.scan_frag_check.setToolTip("Send fragmented packets to potentially evade simple firewalls.")
        row4.addWidget(self.scan_frag_check)
        clayout.addLayout(row4)

        # Connect signals for UI logic
        self.scan_proto_tcp_radio.toggled.connect(self._update_tcp_scan_options_visibility)
        self.scan_proto_udp_radio.toggled.connect(self._update_tcp_scan_options_visibility)
        self.scan_proto_both_radio.toggled.connect(self._update_tcp_scan_options_visibility)

        scan_buttons_layout = QHBoxLayout()
        self.scan_button = QPushButton("Scan"); self.scan_button.setToolTip("Start the port scan.")
        scan_buttons_layout.addWidget(self.scan_button)
        self.scan_cancel_button = QPushButton("Cancel"); self.scan_cancel_button.setEnabled(False); self.scan_cancel_button.setToolTip("Stop the current scan.")
        scan_buttons_layout.addWidget(self.scan_cancel_button)
        clayout.addLayout(scan_buttons_layout)
        self.scan_status = QLabel(""); clayout.addWidget(self.scan_status)
        layout.addWidget(controls)
        self.scan_tree = QTreeWidget(); self.scan_tree.setColumnCount(3); self.scan_tree.setHeaderLabels(["Port", "State", "Service"])
        layout.addWidget(self.scan_tree)
        layout.addWidget(self._create_export_button(self.scan_tree))
        self.scan_button.clicked.connect(self.start_port_scan)
        self.scan_cancel_button.clicked.connect(self.cancel_tool)

        self._update_tcp_scan_options_visibility(True) # Initial state
        return widget

    def _create_arp_scan_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget)
        self.arp_tree=QTreeWidget(); self.arp_tree.setColumnCount(3); self.arp_tree.setHeaderLabels(["IP Address","MAC Address", "Status"])

        controls = QHBoxLayout()
        controls.addWidget(QLabel("Target Network:")); self.arp_target=QLineEdit("192.168.1.0/24"); controls.addWidget(self.arp_target)
        self.arp_scan_button=QPushButton("Scan"); controls.addWidget(self.arp_scan_button)
        self.arp_status=QLabel(""); controls.addWidget(self.arp_status); controls.addStretch()

        layout.addLayout(controls)
        layout.addWidget(self.arp_tree)
        layout.addWidget(self._create_export_button(self.arp_tree))
        self.arp_scan_button.clicked.connect(self.start_arp_scan)
        return widget

    def _create_ping_sweep_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Main controls
        controls_frame = QFrame()
        controls_frame.setFrameShape(QFrame.Shape.StyledPanel)
        controls = QVBoxLayout(controls_frame)

        row1 = QHBoxLayout()
        row1.addWidget(QLabel("Target Network (CIDR):"))
        self.ps_target_edit = QLineEdit("192.168.1.0/24")
        self.ps_target_edit.setToolTip("The target network range in CIDR notation (e.g., 192.168.1.0/24).")
        row1.addWidget(self.ps_target_edit)
        controls.addLayout(row1)

        # Options Box
        options_layout = QFormLayout()
        options_layout.setContentsMargins(5, 10, 5, 10)

        # Probe Type
        self.ps_probe_type_combo = QComboBox()
        self.ps_probe_type_combo.addItems(["ICMP Echo", "TCP SYN", "TCP ACK", "UDP Probe"])
        options_layout.addRow("Probe Type:", self.ps_probe_type_combo)

        # Ports
        self.ps_ports_label = QLabel("Target Port(s):")
        self.ps_ports_edit = QLineEdit("80,443,8080")
        self.ps_ports_edit.setToolTip("Comma-separated list of ports for TCP/UDP probes.")
        options_layout.addRow(self.ps_ports_label, self.ps_ports_edit)

        # Timeout
        self.ps_timeout_edit = QLineEdit("1")
        self.ps_timeout_edit.setToolTip("Timeout in seconds for each probe.")
        options_layout.addRow("Timeout (s):", self.ps_timeout_edit)

        # Threads
        self.ps_threads_edit = QLineEdit("10")
        self.ps_threads_edit.setToolTip("Number of concurrent threads to use for scanning.")
        options_layout.addRow("Threads:", self.ps_threads_edit)
        controls.addLayout(options_layout)
        layout.addWidget(controls_frame)

        buttons_layout = QHBoxLayout()
        self.ps_start_button = QPushButton("Start Sweep")
        buttons_layout.addWidget(self.ps_start_button)
        self.ps_cancel_button = QPushButton("Cancel")
        self.ps_cancel_button.setEnabled(False)
        buttons_layout.addWidget(self.ps_cancel_button)
        layout.addLayout(buttons_layout)

        self.ps_status_label = QLabel("Status: Idle")
        layout.addWidget(self.ps_status_label)

        self.ps_tree = QTreeWidget()
        self.ps_tree.setColumnCount(2)
        self.ps_tree.setHeaderLabels(["IP Address", "Status"])
        layout.addWidget(self.ps_tree)
        layout.addWidget(self._create_export_button(self.ps_tree))

        # --- Connections and Logic ---
        def toggle_ports_visibility(text):
            is_tcp_or_udp = "TCP" in text or "UDP" in text
            self.ps_ports_label.setVisible(is_tcp_or_udp)
            self.ps_ports_edit.setVisible(is_tcp_or_udp)

        self.ps_probe_type_combo.currentTextChanged.connect(toggle_ports_visibility)
        # Set initial state
        is_tcp_or_udp_initial = "TCP" in self.ps_probe_type_combo.currentText() or "UDP" in self.ps_probe_type_combo.currentText()
        self.ps_ports_label.setVisible(is_tcp_or_udp_initial)
        self.ps_ports_edit.setVisible(is_tcp_or_udp_initial)


        self.ps_start_button.clicked.connect(self.start_ping_sweep)
        self.ps_cancel_button.clicked.connect(self.cancel_tool)

        return widget

    def _create_advanced_tools_tab(self, p=None):
        """Creates the tab container for advanced tools."""
        adv_tabs = QTabWidget()
        adv_tabs.addTab(self._create_flooder_tool(), "Packet Flooder")
        adv_tabs.addTab(self._create_firewall_tester_tool(), "Firewall Tester")
        adv_tabs.addTab(self._create_arp_spoofer_tool(), "ARP Spoofer")
        adv_tabs.addTab(self._create_sqlmap_tool(), "SQLMap")
        adv_tabs.addTab(self._create_hashcat_tool(), "Hashcat")
        adv_tabs.addTab(self._create_nuclei_tool(), "Nuclei Scanner")
        adv_tabs.addTab(self._create_trufflehog_tool(), "TruffleHog Scanner")
        adv_tabs.addTab(self._create_jtr_tool(), "John the Ripper")
        adv_tabs.addTab(self._create_hydra_tool(), "Hydra")
        adv_tabs.addTab(self._create_sherlock_tool(), "Sherlock")
        adv_tabs.addTab(self._create_spiderfoot_tool(), "Spiderfoot")
        return adv_tabs

    def _create_flooder_tool(self):
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Attack Template Box ---
        template_box = QGroupBox("Attack Configuration")
        template_layout = QFormLayout(template_box)

        self.flood_template_combo = QComboBox()
        self.flood_template_combo.addItems(["Custom (from Crafter)", "TCP SYN Flood", "UDP Flood", "ICMP Echo Flood"])
        template_layout.addRow("Template:", self.flood_template_combo)

        self.flood_target_label = QLabel("Target IP:")
        self.flood_target_edit = QLineEdit("127.0.0.1")
        template_layout.addRow(self.flood_target_label, self.flood_target_edit)

        self.flood_ports_label = QLabel("Target Port(s):")
        self.flood_ports_edit = QLineEdit("80")
        self.flood_ports_edit.setToolTip("A single port for the flood attack.")
        template_layout.addRow(self.flood_ports_label, self.flood_ports_edit)

        self.flood_rand_src_ip_check = QCheckBox()
        self.flood_rand_src_ip_check.setToolTip("Randomize the source IP address for each packet.")
        template_layout.addRow("Randomize Source IP:", self.flood_rand_src_ip_check)
        main_layout.addWidget(template_box)

        # --- Custom Packet Box (for loading) ---
        packet_frame = QGroupBox("Custom Packet Loader")
        packet_layout = QVBoxLayout(packet_frame)
        self.flood_packet_label = QLabel("Packet to send: (Load from Crafter)")
        packet_layout.addWidget(self.flood_packet_label)
        load_btn = QPushButton("Load Packet from Crafter")
        load_btn.clicked.connect(self.load_flood_packet)
        packet_layout.addWidget(load_btn)
        main_layout.addWidget(packet_frame)

        # --- Flood Parameters ---
        controls_frame = QGroupBox("Flood Parameters")
        controls_layout = QFormLayout(controls_frame)
        self.flood_count = QLineEdit("1000")
        self.flood_count.setToolTip("The total number of packets to send.")
        controls_layout.addRow("Count:", self.flood_count)
        self.flood_interval = QLineEdit("0.01")
        self.flood_interval.setToolTip("The time interval (in seconds) between sending each packet.")
        controls_layout.addRow("Interval:", self.flood_interval)
        self.flood_threads = QLineEdit("4")
        self.flood_threads.setToolTip("The number of parallel threads to use for sending packets.")
        controls_layout.addRow("Threads:", self.flood_threads)
        main_layout.addWidget(controls_frame)

        # --- Action Buttons ---
        flood_buttons_layout = QHBoxLayout()
        self.flood_button = QPushButton("Start Flood")
        self.flood_button.setToolTip("Start the packet flood. Warning: This can cause network disruption.")
        flood_buttons_layout.addWidget(self.flood_button)
        self.stop_flood_button = QPushButton("Stop Flood")
        self.stop_flood_button.setEnabled(False)
        self.stop_flood_button.setToolTip("Stop the ongoing flood.")
        flood_buttons_layout.addWidget(self.stop_flood_button)
        main_layout.addLayout(flood_buttons_layout)

        self.flood_status = QLabel("")
        main_layout.addWidget(self.flood_status)
        main_layout.addStretch()

        # --- UI Logic ---
        def update_template_ui(text):
            is_custom = (text == "Custom (from Crafter)")
            is_icmp = (text == "ICMP Echo Flood")

            self.flood_target_label.setVisible(not is_custom)
            self.flood_target_edit.setVisible(not is_custom)
            self.flood_ports_label.setVisible(not is_custom and not is_icmp)
            self.flood_ports_edit.setVisible(not is_custom and not is_icmp)
            self.flood_rand_src_ip_check.setEnabled(not is_custom)
            packet_frame.setVisible(is_custom)

        self.flood_template_combo.currentTextChanged.connect(update_template_ui)
        update_template_ui(self.flood_template_combo.currentText()) # Initial state

        self.flood_button.clicked.connect(self.start_flood)
        self.stop_flood_button.clicked.connect(self.cancel_tool)

        return widget

    def _create_firewall_tester_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget); controls = QHBoxLayout()
        controls.addWidget(QLabel("Target:")); self.fw_target=QLineEdit("127.0.0.1"); controls.addWidget(self.fw_target)
        controls.addWidget(QLabel("Probe Set:")); self.fw_probe_set=QComboBox(); self.fw_probe_set.addItems(FIREWALL_PROBES.keys()); controls.addWidget(self.fw_probe_set)
        self.fw_test_button=QPushButton("Start Test"); controls.addWidget(self.fw_test_button)
        self.fw_cancel_button = QPushButton("Cancel"); self.fw_cancel_button.setEnabled(False); controls.addWidget(self.fw_cancel_button)
        self.fw_status=QLabel(""); controls.addWidget(self.fw_status); controls.addStretch()
        layout.addLayout(controls)
        self.fw_tree=QTreeWidget(); self.fw_tree.setColumnCount(3); self.fw_tree.setHeaderLabels(["Probe Description","Packet Summary","Result"])
        layout.addWidget(self.fw_tree)
        layout.addWidget(self._create_export_button(self.fw_tree))
        self.fw_test_button.clicked.connect(self.start_firewall_test)
        self.fw_cancel_button.clicked.connect(self.cancel_tool)
        return widget

    def _update_tcp_scan_options_visibility(self, checked):
        """Shows or hides the TCP scan mode dropdown based on protocol selection."""
        is_tcp_selected = self.scan_proto_tcp_radio.isChecked() or self.scan_proto_udp_radio.isChecked()
        self.tcp_scan_type_label.setVisible(is_tcp_selected)
        self.tcp_scan_type_combo.setVisible(is_tcp_selected)

    def _create_arp_spoofer_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Ethical Warning
        warning_label = QTextEdit()
        warning_label.setReadOnly(True)
        warning_label.setStyleSheet("background-color: #4c2222; color: #f0f0f0; border: 1px solid #993333;")
        warning_label.setHtml("""
        <font color='#ffcc00'><b>WARNING & ETHICAL NOTICE:</b></font>
        <p>ARP Spoofing is a powerful technique that can intercept and modify network traffic (Man-in-the-Middle attack). Using this tool on networks you do not own or have explicit, written permission to test is <b>illegal</b> and unethical.</p>
        <p>This tool is for educational and authorized security testing purposes only. The developer assumes no liability for misuse.</p>
        """)
        layout.addWidget(warning_label)

        # Controls
        controls = QFrame()
        controls.setFrameShape(QFrame.Shape.StyledPanel)
        clayout = QVBoxLayout(controls)

        # Target Inputs
        row1 = QHBoxLayout()
        row1.addWidget(QLabel("Victim IP:"))
        self.arp_spoof_victim_ip = QLineEdit()
        self.arp_spoof_victim_ip.setPlaceholderText("e.g., 192.168.1.10")
        self.arp_spoof_victim_ip.setToolTip("The IP address of the target (victim) machine on the local network.")
        row1.addWidget(self.arp_spoof_victim_ip)
        clayout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Target IP (Gateway):"))
        self.arp_spoof_target_ip = QLineEdit()
        self.arp_spoof_target_ip.setPlaceholderText("e.g., 192.168.1.1")
        self.arp_spoof_target_ip.setToolTip("The IP address of the machine you want to impersonate (usually the gateway).")
        row2.addWidget(self.arp_spoof_target_ip)
        clayout.addLayout(row2)

        # Buttons
        buttons_layout = QHBoxLayout()
        self.arp_spoof_start_btn = QPushButton("Start Spoofing")
        self.arp_spoof_start_btn.setToolTip("Begin sending malicious ARP packets to poison the cache of the victim and target.")
        buttons_layout.addWidget(self.arp_spoof_start_btn)
        self.arp_spoof_stop_btn = QPushButton("Stop Spoofing")
        self.arp_spoof_stop_btn.setEnabled(False)
        self.arp_spoof_stop_btn.setToolTip("Stop the attack and send corrective ARP packets to restore the network.")
        buttons_layout.addWidget(self.arp_spoof_stop_btn)
        clayout.addLayout(buttons_layout)

        # Status Label
        self.arp_spoof_status = QLabel("Status: Idle")
        clayout.addWidget(self.arp_spoof_status)

        layout.addWidget(controls)
        layout.addStretch()

        self.arp_spoof_start_btn.clicked.connect(self.start_arp_spoof)
        self.arp_spoof_stop_btn.clicked.connect(self.stop_arp_spoof)

        return widget

    def _create_system_info_tab(self):
        """Creates the System Info tab with a redesigned, more modern layout."""
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setStyleSheet("QScrollArea { border: none; }") # Remove scroll area border

        main_widget = QWidget()
        main_layout = QVBoxLayout(main_widget)
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(20, 20, 20, 20)
        scroll_area.setWidget(main_widget)

        # --- Helper for creating styled GroupBoxes ---
        def create_info_box(title):
            box = QGroupBox(title)
            # Basic styling for a modern "card" look
            box.setStyleSheet("""
                QGroupBox {
                    font-size: 14px;
                    font-weight: bold;
                    border: 1px solid #444;
                    border-radius: 8px;
                    margin-top: 10px;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    subcontrol-position: top left;
                    padding: 0 10px;
                }
            """)
            layout = QFormLayout(box)
            layout.setSpacing(10)
            layout.setContentsMargins(15, 25, 15, 15) # Top margin for title
            return box, layout

        # --- Top Row: System, CPU, Memory ---
        top_row_layout = QHBoxLayout()
        top_row_layout.setSpacing(20)

        # System Info Box
        sys_box, sys_layout = create_info_box("System")
        sys_layout.addRow("OS:", QLabel(f"{platform.system()} {platform.release()}"))
        sys_layout.addRow("Architecture:", QLabel(platform.machine()))
        sys_layout.addRow("Hostname:", QLabel(platform.node()))
        sys_layout.addRow("Python Version:", QLabel(platform.python_version()))
        top_row_layout.addWidget(sys_box)

        # CPU Info Box
        cpu_box, cpu_layout = create_info_box("CPU")
        try:
            cpu_freq = psutil.cpu_freq()
            freq_str = f"{cpu_freq.current:.2f} Mhz (Max: {cpu_freq.max:.2f} Mhz)" if cpu_freq else "N/A"
        except Exception:
            freq_str = "N/A (Permission Denied)"
        cpu_layout.addRow("Frequency:", QLabel(freq_str))
        cpu_layout.addRow("Physical Cores:", QLabel(str(psutil.cpu_count(logical=False))))
        cpu_layout.addRow("Logical Cores:", QLabel(str(psutil.cpu_count(logical=True))))
        top_row_layout.addWidget(cpu_box)

        # Memory Info Box
        mem_box, mem_layout = create_info_box("Memory")
        mem = psutil.virtual_memory()
        swap = psutil.swap_memory()
        mem_layout.addRow("Total RAM:", QLabel(f"{mem.total / (1024**3):.2f} GB"))
        mem_layout.addRow("Available RAM:", QLabel(f"{mem.available / (1024**3):.2f} GB"))
        mem_layout.addRow("Swap Total:", QLabel(f"{swap.total / (1024**3):.2f} GB"))
        mem_layout.addRow("Swap Used:", QLabel(f"{swap.used / (1024**3):.2f} GB ({swap.percent}%)"))
        top_row_layout.addWidget(mem_box)

        main_layout.addLayout(top_row_layout)

        # --- Second Row: Libraries and GPU ---
        second_row_layout = QHBoxLayout()
        second_row_layout.setSpacing(20)

        # Library Versions Box
        try:
            scapy_version = scapy.VERSION
        except AttributeError:
            scapy_version = "Unknown"
        lib_box, lib_layout = create_info_box("Library Versions")
        lib_layout.addRow("Scapy:", QLabel(scapy_version))
        lib_layout.addRow("PyQt6:", QLabel(PYQT_VERSION_STR))
        lib_layout.addRow("psutil:", QLabel(psutil.__version__))
        if GPUtil:
            lib_layout.addRow("GPUtil:", QLabel(getattr(GPUtil, '__version__', 'N/A')))
        second_row_layout.addWidget(lib_box)

        # GPU Info Box
        if GPUtil:
            gpu_box, gpu_layout = create_info_box("GPU Information")
            try:
                gpus = GPUtil.getGPUs()
                if not gpus:
                    gpu_layout.addRow(QLabel("No NVIDIA GPU detected."))
                else:
                    for i, gpu in enumerate(gpus):
                        gpu_layout.addRow(f"GPU {i} Name:", QLabel(gpu.name))
                        gpu_layout.addRow("  - Driver:", QLabel(gpu.driver))
                        gpu_layout.addRow("  - Memory:", QLabel(f"{gpu.memoryUsed}MB / {gpu.memoryTotal}MB"))
            except Exception as e:
                gpu_layout.addRow(QLabel(f"Could not retrieve GPU info: {e}"))
            second_row_layout.addWidget(gpu_box)

        second_row_layout.addStretch()
        main_layout.addLayout(second_row_layout)

        # --- Disk Partitions Box ---
        disk_box = QGroupBox("Disk Partitions")
        disk_box.setStyleSheet("""
            QGroupBox {
                font-size: 14px;
                font-weight: bold;
                border: 1px solid #444;
                border-radius: 8px;
                margin-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top left;
                padding: 0 10px;
            }
        """)
        disk_layout = QGridLayout(disk_box)
        disk_layout.setContentsMargins(15, 25, 15, 15)
        try:
            partitions = psutil.disk_partitions()
            if not partitions:
                disk_layout.addWidget(QLabel("No disk partitions found."), 0, 0)
            else:
                row, col = 0, 0
                for part in partitions:
                    try:
                        usage = psutil.disk_usage(part.mountpoint)
                        part_label = QLabel(f"<b>{part.device}</b> on {part.mountpoint} ({part.fstype})<br>"
                                          f"&nbsp;&nbsp;Total: {usage.total / (1024**3):.2f} GB, "
                                          f"Used: {usage.used / (1024**3):.2f} GB ({usage.percent}%)")
                        disk_layout.addWidget(part_label, row, col)
                        col += 1
                        if col >= 2: # 2 columns
                            col = 0
                            row += 1
                    except Exception:
                        continue # Skip inaccessible drives
        except Exception as e:
            disk_layout.addWidget(QLabel(f"Could not retrieve disk partitions: {e}"), 0, 0)
        main_layout.addWidget(disk_box)


        # --- Network Interfaces Box ---
        net_box = QGroupBox("Network Interfaces")
        net_box.setStyleSheet("""
            QGroupBox {
                font-size: 14px; font-weight: bold; border: 1px solid #444;
                border-radius: 8px; margin-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin; subcontrol-position: top left; padding: 0 10px;
            }
        """)
        net_main_layout = QVBoxLayout(net_box)
        net_main_layout.setContentsMargins(15, 25, 15, 15)

        try:
            ifaddrs = psutil.net_if_addrs()
            if not ifaddrs:
                net_main_layout.addWidget(QLabel("No network interfaces found."))
            else:
                for iface, addrs in sorted(ifaddrs.items()):
                    # Skip loopback interfaces unless they have a non-standard address
                    is_loopback = 'loopback' in iface.lower() or iface.startswith('lo')
                    if is_loopback and all(addr.address in ['127.0.0.1', '::1'] for addr in addrs):
                        continue

                    iface_box = QGroupBox(iface)
                    iface_box.setStyleSheet("QGroupBox { border: 1px solid #555; margin-top: 5px; }")
                    iface_layout = QFormLayout(iface_box)

                    addr_map = {'ipv4': [], 'ipv6': [], 'mac': 'N/A'}
                    for addr in addrs:
                        if addr.family == socket.AF_INET:
                            addr_map['ipv4'].append(addr.address)
                        elif addr.family == socket.AF_INET6:
                            # Filter out link-local addresses for cleaner display
                            if not addr.address.startswith('fe80::'):
                                addr_map['ipv6'].append(addr.address)
                        # This logic correctly handles cross-platform MAC address retrieval
                        elif hasattr(psutil, 'AF_LINK') and addr.family == psutil.AF_LINK:
                            addr_map['mac'] = addr.address
                        elif hasattr(socket, 'AF_PACKET') and addr.family == socket.AF_PACKET:
                             addr_map['mac'] = addr.address

                    # Join multiple IPs, or display N/A if none were found
                    ipv4_str = ", ".join(addr_map['ipv4']) or "N/A"
                    ipv6_str = ", ".join(addr_map['ipv6']) or "N/A"

                    iface_layout.addRow(QLabel("<b>IPv4 Address:</b>"), QLabel(ipv4_str))
                    iface_layout.addRow(QLabel("<b>IPv6 Address:</b>"), QLabel(ipv6_str))
                    iface_layout.addRow(QLabel("<b>MAC Address:</b>"), QLabel(addr_map['mac']))

                    net_main_layout.addWidget(iface_box)
        except Exception as e:
            logging.error(f"Could not retrieve network interfaces: {e}", exc_info=True)
            net_main_layout.addWidget(QLabel(f"Could not retrieve interfaces: {e}"))

        main_layout.addWidget(net_box)

        main_layout.addStretch() # Push everything to the top
        return scroll_area

    def _create_wireless_tools_tab(self, p=None):
        """Creates the tab container for 802.11 wireless tools."""
        wireless_tabs = QTabWidget()
        wireless_tabs.addTab(self._create_wifi_scanner_tool(), "Wi-Fi Scanner")
        wireless_tabs.addTab(self._create_deauth_tool(), "Deauthentication Tool")
        wireless_tabs.addTab(self._create_beacon_flood_tool(), "Beacon Flood")
        wireless_tabs.addTab(self._create_wpa_crack_tool(), "WPA Handshake Tool")
        wireless_tabs.addTab(self._create_krack_scanner_tool(), "KRACK Scanner")
        wireless_tabs.addTab(self._create_wifite_tool(), "Wifite Auditor")
        return wireless_tabs

    def _create_krack_scanner_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>KRACK Vulnerability Scanner</b></font>
        <p>This tool passively detects networks vulnerable to Key Reinstallation Attacks (KRACK). It works by listening for retransmitted EAPOL Message 3 packets during a 4-way handshake.</p>
        <p><b>Usage:</b></p>
        <ol>
            <li>Ensure your wireless card is in <b>Monitor Mode</b> and select it at the top.</li>
            <li>Click "Start Scan". The tool will listen indefinitely.</li>
            <li>To trigger a handshake, you can use the Deauthentication Tool to briefly disconnect a client, forcing it to reconnect.</li>
            <li>Any vulnerable networks detected will appear in the results table below.</li>
        </ol>
        """)
        layout.addWidget(instructions)

        controls = QHBoxLayout()
        self.krack_start_btn = QPushButton("Start Scan")
        self.krack_stop_btn = QPushButton("Stop Scan"); self.krack_stop_btn.setEnabled(False)
        controls.addWidget(self.krack_start_btn)
        controls.addWidget(self.krack_stop_btn)
        layout.addLayout(controls)

        self.krack_results_tree = QTreeWidget()
        self.krack_results_tree.setColumnCount(3)
        self.krack_results_tree.setHeaderLabels(["BSSID (AP)", "Client MAC", "Time Detected"])
        layout.addWidget(self.krack_results_tree)

        self.krack_start_btn.clicked.connect(self.start_krack_scan)
        self.krack_stop_btn.clicked.connect(self.stop_krack_scan)

        return widget

    def _create_wifite_tool(self):
        """Creates the UI for the Wifite automated wireless auditor."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Instructions and Warning ---
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>Wifite Automated Auditor</b></font>
        <p>This tool runs the Wifite script to automatically audit WPA/WPA2/WPS encrypted networks.</p>
        <p><b>WARNING:</b> This is an active attack tool. You MUST have explicit permission to test the target network. Ensure your wireless card is in <b>Monitor Mode</b> and select it from the main interface dropdown.</p>
        """)
        layout.addWidget(instructions)

        # --- Controls ---
        controls_frame = QGroupBox("Wifite Options")
        controls_layout = QFormLayout(controls_frame)

        self.wifite_essid_edit = QLineEdit()
        self.wifite_essid_edit.setPlaceholderText("Leave blank for an attack on all networks")
        controls_layout.addRow("Target ESSID:", self.wifite_essid_edit)

        options_layout = QHBoxLayout()
        self.wifite_wps_check = QCheckBox("WPS Attacks (--wps)")
        self.wifite_pmkid_check = QCheckBox("PMKID Attacks (--pmkid)")
        self.wifite_kill_check = QCheckBox("Kill Conflicting Processes (--kill)")
        self.wifite_wps_check.setToolTip("Run only WPS attacks.")
        self.wifite_pmkid_check.setToolTip("Run only PMKID attacks.")
        self.wifite_kill_check.setToolTip("Kill processes that interfere with monitor mode.")
        options_layout.addWidget(self.wifite_wps_check)
        options_layout.addWidget(self.wifite_pmkid_check)
        options_layout.addWidget(self.wifite_kill_check)
        controls_layout.addRow("Attack Types:", options_layout)

        layout.addWidget(controls_frame)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        self.wifite_start_btn = QPushButton(QIcon("icons/wifi.svg"), " Start Wifite Scan")
        self.wifite_stop_btn = QPushButton("Stop Wifite"); self.wifite_stop_btn.setEnabled(False)
        buttons_layout.addWidget(self.wifite_start_btn)
        buttons_layout.addWidget(self.wifite_stop_btn)
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.wifite_output_console = QPlainTextEdit()
        self.wifite_output_console.setReadOnly(True)
        self.wifite_output_console.setFont(QFont("Courier New", 10))
        self.wifite_output_console.setPlaceholderText("Wifite output will be displayed here...")
        layout.addWidget(self.wifite_output_console, 1) # Add stretch factor

        self.wifite_start_btn.clicked.connect(self.start_wifite_scan)
        self.wifite_stop_btn.clicked.connect(self.cancel_tool)

        return widget

    def start_wifite_scan(self):
        """Starts the Wifite scan worker thread."""
        if not shutil.which("wifite"):
            QMessageBox.critical(self, "Wifite Error", "'wifite' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Interface Error", "Please select a monitor-mode wireless interface.")
            return

        # Use pkexec to request root privileges graphically if available and not root
        try:
            is_root = (os.geteuid() == 0)
        except AttributeError: # os.geteuid() does not exist on Windows
            is_root = False

        command_prefix = []
        if not is_root and shutil.which("pkexec"):
             command_prefix = ["pkexec"]
        elif not is_root:
            QMessageBox.critical(self, "Permission Error", "Wifite requires root privileges. Please run GScapy with sudo or ensure pkexec is installed for graphical password prompts.")
            return

        command = command_prefix + ["wifite", "-i", iface]
        essid = self.wifite_essid_edit.text().strip()
        if essid:
            command.extend(["--essid", essid])

        if self.wifite_wps_check.isChecked():
            command.append("--wps")
        if self.wifite_pmkid_check.isChecked():
            command.append("--pmkid")
        if self.wifite_kill_check.isChecked():
            command.append("--kill")

        self.is_tool_running = True
        self.wifite_start_btn.setEnabled(False)
        self.wifite_stop_btn.setEnabled(True)
        self.tool_stop_event.clear()
        self.wifite_output_console.clear()

        self.worker = WorkerThread(self._wifite_thread, args=(command,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _wifite_thread(self, command):
        """Worker thread for running the wifite command."""
        q = self.tool_results_queue
        logging.info(f"Starting Wifite with command: {' '.join(command)}")
        q.put(('wifite_output', f"$ {' '.join(command)}\n\n"))

        try:
            if sys.platform == "win32":
                q.put(('error', 'Platform Error', 'Wifite is not supported on Windows.'))
                q.put(('tool_finished', 'wifite_scan'))
                return

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.wifite_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    # Terminating the parent pkexec/sudo is tricky, we try to kill the child wifite
                    try:
                        # This might require root to kill a root process if not run with sudo
                        os.kill(process.pid, signal.SIGTERM)
                        logging.info(f"Sent SIGTERM to wifite process {process.pid}")
                    except Exception as e:
                        logging.error(f"Could not kill wifite process: {e}")
                    q.put(('wifite_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('wifite_output', line))

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Wifite Error', "'wifite' command not found. Please ensure it is installed and in your system's PATH."))
        except AttributeError: # os.geteuid() doesn't exist on windows
             q.put(('error', 'Platform Error', 'Wifite requires a Linux-based system with root privileges.'))
        except Exception as e:
            logging.error(f"Wifite thread error: {e}", exc_info=True)
            q.put(('error', 'Wifite Error', str(e)))
        finally:
            q.put(('tool_finished', 'wifite_scan'))
            with self.thread_finish_lock:
                self.wifite_process = None
            logging.info("Wifite scan thread finished.")


    def _create_wifi_scanner_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget)
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setStyleSheet("background-color: #3c3c3c; color: #f0f0f0; border: 1px solid #555;")
        instructions.setHtml("""
        <font color='#ffcc00'><b>WARNING:</b> Wireless tools require the selected interface to be in <b>Monitor Mode</b>.</font>
        <p>GScapy cannot enable this mode for you. You must do it manually before scanning.</p>
        <p><b>Example for Linux (using airmon-ng):</b></p>
        <ol>
            <li>Find your interface: <code>iwconfig</code> (e.g., wlan0)</li>
            <li>Start monitor mode: <code>sudo airmon-ng start wlan0</code></li>
            <li>A new interface (e.g., wlan0mon) will be created.</li>
            <li><b>Select the new monitor interface (e.g., wlan0mon) from the dropdown at the top of the GScapy window.</b></li>
        </ol>
        """)
        layout.addWidget(instructions)
        controls = QHBoxLayout()
        self.wifi_scan_button = QPushButton("Scan for Wi-Fi Networks")
        self.wifi_scan_button.setToolTip("Scans for nearby Wi-Fi networks.\nThe selected interface must be in monitor mode.")
        controls.addWidget(self.wifi_scan_button)
        self.wifi_scan_stop_button = QPushButton("Stop Scan")
        self.wifi_scan_stop_button.setToolTip("Stops the current Wi-Fi scan.")
        self.wifi_scan_stop_button.setEnabled(False)
        controls.addWidget(self.wifi_scan_stop_button)
        self.wifi_scan_status = QLabel(""); controls.addWidget(self.wifi_scan_status); controls.addStretch()
        layout.addLayout(controls)
        self.wifi_tree = QTreeWidget(); self.wifi_tree.setColumnCount(4); self.wifi_tree.setHeaderLabels(["SSID", "BSSID", "Channel", "Signal"])
        layout.addWidget(self.wifi_tree)
        layout.addWidget(self._create_export_button(self.wifi_tree))
        self.wifi_scan_button.clicked.connect(self.start_wifi_scan)
        self.wifi_scan_stop_button.clicked.connect(self.stop_wifi_scan)
        return widget

    def _create_wpa_crack_tool(self):
        """Creates the UI for the WPA Handshake and Cracking tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Handshake Capture Section ---
        capture_box = QGroupBox("WPA Handshake Capture")
        capture_layout = QVBoxLayout(capture_box)

        target_layout = QHBoxLayout()
        target_layout.addWidget(QLabel("Target BSSID:"))
        self.wpa_target_combo = QComboBox(); self.wpa_target_combo.setToolTip("Select a target network from the list discovered by the Wi-Fi Scanner.")
        target_layout.addWidget(self.wpa_target_combo)
        refresh_btn = QPushButton("Refresh List"); refresh_btn.setToolTip("Update the list of targets from the Wi-Fi Scanner tab.")
        refresh_btn.clicked.connect(self._refresh_wpa_targets)
        target_layout.addWidget(refresh_btn)
        capture_layout.addLayout(target_layout)

        capture_controls = QHBoxLayout()
        self.wpa_capture_btn = QPushButton("Start Handshake Capture"); self.wpa_capture_btn.setToolTip("Begin sniffing for a WPA handshake from the selected target.")
        capture_controls.addWidget(self.wpa_capture_btn)
        self.wpa_deauth_client_btn = QPushButton("Deauth Client to Speed Up"); self.wpa_deauth_client_btn.setToolTip("Send deauthentication packets to the network to encourage a client to reconnect, speeding up handshake capture.")
        capture_controls.addWidget(self.wpa_deauth_client_btn)
        capture_layout.addLayout(capture_controls)

        self.wpa_capture_status = QLabel("Status: Idle")
        capture_layout.addWidget(self.wpa_capture_status)
        layout.addWidget(capture_box)

        # --- Hash Cracker Section ---
        cracker_box = QGroupBox("WPA Hash Cracker")
        cracker_layout = QVBoxLayout(cracker_box)

        pcap_layout = QHBoxLayout()
        pcap_layout.addWidget(QLabel("Handshake File (.pcap):"))
        self.wpa_pcap_edit = QLineEdit(); self.wpa_pcap_edit.setPlaceholderText("Path to .pcap file containing the handshake...")
        self.wpa_pcap_edit.setToolTip("The .pcap file containing the captured WPA handshake.")
        pcap_layout.addWidget(self.wpa_pcap_edit)
        pcap_browse_btn = QPushButton("Browse...")
        pcap_browse_btn.setToolTip("Browse for a .pcap file containing a WPA handshake.")
        pcap_browse_btn.clicked.connect(self.browse_for_pcap)
        pcap_layout.addWidget(pcap_browse_btn)
        cracker_layout.addLayout(pcap_layout)

        wordlist_layout = QHBoxLayout()
        wordlist_layout.addWidget(QLabel("Wordlist File:"))
        self.wpa_wordlist_edit = QLineEdit(); self.wpa_wordlist_edit.setPlaceholderText("Path to wordlist file (or leave blank for default)...")
        self.wpa_wordlist_edit.setToolTip("The wordlist file to use for the dictionary attack. If left blank, a small internal list will be used.")
        wordlist_layout.addWidget(self.wpa_wordlist_edit)
        wordlist_browse_btn = QPushButton("Browse...")
        wordlist_browse_btn.setToolTip("Browse for a wordlist file (.txt).")
        wordlist_browse_btn.clicked.connect(self.browse_for_wordlist)
        wordlist_layout.addWidget(wordlist_browse_btn)
        crunch_btn = QPushButton("Generate...")
        crunch_btn.setToolTip("Generate a custom wordlist using Crunch (must be installed).")
        crunch_btn.clicked.connect(self.open_crunch_generator)
        wordlist_layout.addWidget(crunch_btn)
        cracker_layout.addLayout(wordlist_layout)

        cpu_layout = QHBoxLayout()
        cpu_layout.addWidget(QLabel("CPU Threads:"))
        self.wpa_threads_edit = QLineEdit("1"); self.wpa_threads_edit.setToolTip("Number of CPU threads for aircrack-ng to use.")
        cpu_layout.addWidget(self.wpa_threads_edit)
        cpu_layout.addStretch()
        cracker_layout.addLayout(cpu_layout)

        self.wpa_crack_btn = QPushButton("Start Cracking"); self.wpa_crack_btn.setToolTip("Begin the cracking process using aircrack-ng.")
        cracker_layout.addWidget(self.wpa_crack_btn)

        self.wpa_crack_output = QPlainTextEdit(); self.wpa_crack_output.setReadOnly(True)
        self.wpa_crack_output.setPlaceholderText("Aircrack-ng output will be shown here...")
        cracker_layout.addWidget(self.wpa_crack_output)
        layout.addWidget(cracker_box)

        self.wpa_capture_btn.clicked.connect(self.start_handshake_capture)
        self.wpa_deauth_client_btn.clicked.connect(self.deauth_for_handshake)
        self.wpa_crack_btn.clicked.connect(self.start_wpa_crack)

        return widget

    def _refresh_wpa_targets(self):
        self.wpa_target_combo.clear()
        if not self.found_networks:
            QMessageBox.information(self, "No Networks", "No networks found. Please run the Wi-Fi Scanner first.")
            return
        for bssid, info in self.found_networks.items():
            ssid = info[0]
            self.wpa_target_combo.addItem(f"{ssid} ({bssid})", bssid)

    def _browse_file_for_lineedit(self, line_edit_widget, dialog_title):
        file_path, _ = QFileDialog.getOpenFileName(self, dialog_title, "", "All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            line_edit_widget.setText(file_path)

    def _browse_save_file_for_lineedit(self, line_edit_widget, dialog_title):
        file_path, _ = QFileDialog.getSaveFileName(self, dialog_title, "", "All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            line_edit_widget.setText(file_path)

    def browse_for_pcap(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Handshake File", "", "Pcap Files (*.pcap *.pcapng);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.wpa_pcap_edit.setText(file_path)

    def browse_for_wordlist(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Wordlist File", "", "Text Files (*.txt);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.wpa_wordlist_edit.setText(file_path)

    def browse_nikto_output(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Nikto Report", "", "HTML Files (*.html);;CSV Files (*.csv);;Text Files (*.txt);;XML Files (*.xml);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            self.nikto_output_file_edit.setText(file_path)

    def browse_nikto_save_dir(self):
        dir_path = QFileDialog.getExistingDirectory(self, "Select Directory to Save Responses", options=QFileDialog.Option.DontUseNativeDialog)
        if dir_path:
            self.nikto_save_dir_edit.setText(dir_path)

    def start_wpa_crack(self):
        if self.aircrack_thread and self.aircrack_thread.isRunning():
            self.aircrack_thread.stop()
            return

        pcap_file = self.wpa_pcap_edit.text()
        wordlist = self.wpa_wordlist_edit.text()
        try:
            threads = int(self.wpa_threads_edit.text())
        except ValueError:
            QMessageBox.warning(self, "Input Error", "CPU threads must be a valid number.")
            return

        if not pcap_file:
            QMessageBox.warning(self, "Input Error", "Please provide a handshake file.")
            return
        if not os.path.exists(pcap_file):
            QMessageBox.warning(self, "File Error", f"Pcap file not found:\n{pcap_file}")
            return

        if not wordlist:
            wordlist = "default_pass.txt"

        if not os.path.exists(wordlist):
            QMessageBox.warning(self, "File Error", f"Wordlist file not found:\n{wordlist}")
            return

        self.wpa_crack_output.clear()
        self.wpa_crack_btn.setText("Stop Cracking")
        self.aircrack_thread = AircrackThread(pcap_file, wordlist, self, threads)
        self.aircrack_thread.output_received.connect(self._process_aircrack_output)
        self.aircrack_thread.finished_signal.connect(self._on_aircrack_finished)
        self.aircrack_thread.start()

    def _process_aircrack_output(self, line):
        self.wpa_crack_output.appendPlainText(line)
        if "KEY FOUND!" in line:
            self.wpa_crack_output.appendPlainText("\n\n---> PASSWORD FOUND! <---")
            self.aircrack_thread.stop()

    def _on_aircrack_finished(self, return_code):
        self.wpa_crack_btn.setText("Start Cracking")
        self.wpa_crack_output.appendPlainText(f"\n--- Process finished with exit code {return_code} ---")

    def open_crunch_generator(self):
        dialog = CrunchDialog(self)
        if dialog.exec():
            values = dialog.get_values()
            min_len, max_len, charset, outfile = values["min"], values["max"], values["charset"], values["outfile"]

            if not all([min_len, max_len, charset, outfile]):
                QMessageBox.warning(self, "Input Error", "All fields are required to generate a wordlist.")
                return

            command = ["crunch", min_len, max_len, charset, "-o", outfile]

            try:
                self.wpa_crack_output.appendPlainText(f"Starting crunch: {' '.join(command)}")

                def run_crunch():
                    try:
                        # Use CREATE_NO_WINDOW flag on Windows to hide the console
                        startupinfo = None
                        if sys.platform == "win32":
                            startupinfo = subprocess.STARTUPINFO()
                            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

                        process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, startupinfo=startupinfo)
                        for line in iter(process.stdout.readline, ''):
                            logging.info(f"[crunch] {line.strip()}")
                        process.wait()
                        self.tool_results_queue.put(('crunch_finished', outfile, process.returncode))
                    except FileNotFoundError:
                        self.tool_results_queue.put(('error', 'Crunch Error', "'crunch' command not found. Please ensure it is installed and in your system's PATH."))
                    except Exception as e:
                        self.tool_results_queue.put(('error', 'Crunch Error', str(e)))

                self.worker = WorkerThread(target=run_crunch)
                self.worker.start()

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to start crunch: {e}")

    def start_handshake_capture(self):
        if self.is_tool_running:
            self.stop_handshake_capture()
            return

        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Interface Error", "Please select a monitor-mode interface.")
            return

        bssid = self.wpa_target_combo.currentData()
        if not bssid:
            QMessageBox.warning(self, "Target Error", "Please select a target network.")
            return

        self.is_tool_running = True
        self.wpa_capture_btn.setText("Stop Capture")
        self.handshake_sniffer_thread = HandshakeSnifferThread(iface, bssid)
        self.handshake_sniffer_thread.log_message.connect(lambda msg: self.wpa_capture_status.setText(f"Status: {msg}"))
        self.handshake_sniffer_thread.handshake_captured.connect(self._on_handshake_captured)
        self.handshake_sniffer_thread.start()

    def stop_handshake_capture(self):
        if self.handshake_sniffer_thread and self.handshake_sniffer_thread.isRunning():
            self.handshake_sniffer_thread.stop()
            self.handshake_sniffer_thread.wait()
        self.is_tool_running = False
        self.wpa_capture_btn.setText("Start Handshake Capture")
        self.wpa_capture_status.setText("Status: Idle")

    def _on_handshake_captured(self, bssid, file_path):
        self.wpa_capture_status.setText(f"Status: Handshake for {bssid} captured and saved to {file_path}!")
        self.stop_handshake_capture()
        QMessageBox.information(self, "Success", f"Handshake captured and saved to {file_path}")
        self.wpa_pcap_edit.setText(file_path)

    def deauth_for_handshake(self):
        bssid = self.wpa_target_combo.currentData()
        if not bssid:
            QMessageBox.warning(self, "Target Error", "Please select a target network to deauthenticate.")
            return
        args = (bssid, "ff:ff:ff:ff:ff:ff", 5)
        self.worker = WorkerThread(self._deauth_thread, args=args)
        self.worker.start()
        QMessageBox.information(self, "Deauth Sent", f"Sent 5 deauth packets to the network {bssid} to encourage re-association.")

    def _create_deauth_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget)
        warning_label = QLabel("WARNING: Sending deauthentication packets can disrupt networks you do not own. Use responsibly and only on your own network for testing purposes.")
        warning_label.setStyleSheet("color: #ffcc00;")
        layout.addWidget(warning_label)
        controls = QFrame(); clayout = QVBoxLayout(controls)
        row1 = QHBoxLayout()
        row1.addWidget(QLabel("AP BSSID (MAC):"))
        self.deauth_bssid = QLineEdit("ff:ff:ff:ff:ff:ff")
        self.deauth_bssid.setToolTip("The MAC address (BSSID) of the target Access Point.")
        row1.addWidget(self.deauth_bssid)
        clayout.addLayout(row1)
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Client MAC:"))
        self.deauth_client = QLineEdit("ff:ff:ff:ff:ff:ff")
        self.deauth_client.setToolTip("The MAC address of the client to deauthenticate.\nUse 'ff:ff:ff:ff:ff:ff' to deauthenticate all clients.")
        row2.addWidget(self.deauth_client)
        clayout.addLayout(row2)
        row3 = QHBoxLayout()
        row3.addWidget(QLabel("Count:"))
        self.deauth_count = QLineEdit("10")
        self.deauth_count.setToolTip("The number of deauthentication packets to send.")
        row3.addWidget(self.deauth_count)
        clayout.addLayout(row3)
        self.deauth_button = QPushButton(QIcon("icons/user-minus.svg"), " Send Deauth Packets")
        self.deauth_button.setToolTip("Start sending deauthentication packets.\nWARNING: This will disrupt the target's connection.")
        clayout.addWidget(self.deauth_button)
        self.deauth_status = QLabel(""); clayout.addWidget(self.deauth_status)
        layout.addWidget(controls)
        self.deauth_button.clicked.connect(self.start_deauth)
        return widget

    def _create_beacon_flood_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        warning_label = QLabel("WARNING: Flooding the air with beacon frames can disrupt Wi-Fi networks in the area. Use this tool responsibly and only for legitimate testing purposes.")
        warning_label.setStyleSheet("color: #ffcc00;")
        warning_label.setWordWrap(True)
        layout.addWidget(warning_label)

        controls = QFrame()
        controls.setFrameShape(QFrame.Shape.StyledPanel)
        clayout = QFormLayout(controls)

        # SSID controls
        ssid_layout = QHBoxLayout()
        self.bf_ssid_edit = QLineEdit("TestNet")
        self.bf_ssid_edit.setToolTip("A single SSID, or load multiple from a file.")
        ssid_layout.addWidget(self.bf_ssid_edit)
        self.bf_ssid_from_file_btn = QPushButton("Load from File")
        self.bf_ssid_from_file_btn.setToolTip("Load a list of SSIDs from a .txt file (one per line).")
        ssid_layout.addWidget(self.bf_ssid_from_file_btn)
        clayout.addRow("SSID(s):", ssid_layout)

        self.bf_bssid_edit = QLineEdit("random")
        self.bf_bssid_edit.setToolTip("The BSSID (MAC address) of the fake AP. 'random' will generate a new MAC for each packet.")
        clayout.addRow("BSSID:", self.bf_bssid_edit)

        # Encryption
        self.bf_enc_combo = QComboBox()
        self.bf_enc_combo.addItems(["Open", "WEP", "WPA2-PSK", "WPA3-SAE"])
        self.bf_enc_combo.setToolTip("Select the advertised encryption type for the fake network(s).")
        clayout.addRow("Encryption:", self.bf_enc_combo)

        # Channel
        self.bf_channel_edit = QLineEdit("1")
        self.bf_channel_edit.setToolTip("The 802.11 channel to broadcast the beacons on.")
        clayout.addRow("Channel:", self.bf_channel_edit)

        self.bf_count_edit = QLineEdit("1000")
        self.bf_count_edit.setToolTip("The number of beacon frames to send. Use '0' for an infinite flood.")
        clayout.addRow("Count:", self.bf_count_edit)

        self.bf_interval_edit = QLineEdit("0.1")
        self.bf_interval_edit.setToolTip("The time interval (in seconds) between sending each beacon frame.")
        clayout.addRow("Interval:", self.bf_interval_edit)

        layout.addWidget(controls)

        buttons_layout = QHBoxLayout()
        self.bf_start_button = QPushButton("Start Beacon Flood")
        self.bf_start_button.setToolTip("Begin sending fake beacon frames.")
        buttons_layout.addWidget(self.bf_start_button)

        self.bf_stop_button = QPushButton("Stop Flood")
        self.bf_stop_button.setEnabled(False)
        self.bf_stop_button.setToolTip("Stop the ongoing beacon flood.")
        buttons_layout.addWidget(self.bf_stop_button)
        layout.addLayout(buttons_layout)

        self.bf_status_label = QLabel("Status: Idle")
        layout.addWidget(self.bf_status_label)
        layout.addStretch()

        self.bf_start_button.clicked.connect(self.start_beacon_flood)
        self.bf_stop_button.clicked.connect(self.cancel_tool)
        self.bf_ssid_from_file_btn.clicked.connect(self.load_ssids_for_beacon_flood)

        return widget

    def load_ssids_for_beacon_flood(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select SSID List File", "", "Text Files (*.txt);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    self.bf_ssid_list = [line.strip() for line in f if line.strip()]
                if self.bf_ssid_list:
                    self.bf_ssid_edit.setText(f"Loaded {len(self.bf_ssid_list)} SSIDs from file")
                    self.bf_ssid_edit.setReadOnly(True)
                    logging.info(f"Loaded {len(self.bf_ssid_list)} SSIDs for beacon flood.")
                else:
                    self.bf_ssid_edit.setText("")
                    self.bf_ssid_edit.setReadOnly(False)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load SSID file: {e}")


    # --- Backend Methods: Sniffer ---
    def start_sniffing(self):
        """Starts the packet sniffer thread."""
        self.start_sniff_btn.setEnabled(False)
        self.stop_sniff_btn.setEnabled(True)
        self.clear_sniffer_display()
        iface = self.get_selected_iface()
        bpf_filter = self.filter_input.text()
        # Create the thread manager
        self.sniffer_thread = SnifferThread(iface=iface, bpf_filter=bpf_filter, parent=self)
        # Connect the new thread-safe signal to the reconstruction slot
        self.sniffer_thread.packet_bytes_received.connect(self._handle_packet_bytes)
        self.sniffer_thread.finished.connect(self._on_sniffer_finished)
        self.sniffer_thread.start()
        self.status_bar.showMessage(f"Sniffing on interface: {iface or 'default'}...")

    def _handle_packet_bytes(self, pkt_bytes):
        """Reconstructs a packet from bytes and adds it to a buffer for batch updating."""
        try:
            # Reconstruction is quick, so it's safe to do here.
            packet = Ether(pkt_bytes)
            with self.sniffer_buffer_lock:
                self.sniffer_packet_buffer.append(packet)
        except Exception as e:
            logging.error(f"Failed to reconstruct or buffer packet: {e}")

    def stop_sniffing(self):
        """Signals the packet sniffer thread to stop."""
        if self.sniffer_thread and self.sniffer_thread.isRunning():
            self.stop_sniff_btn.setEnabled(False) # Prevent multiple clicks
            self.status_bar.showMessage("Stopping sniffer...")
            self.sniffer_thread.stop()

    def _on_sniffer_finished(self):
        """Handles cleanup after the sniffer thread has terminated."""
        self.start_sniff_btn.setEnabled(True)
        # The stop button is already disabled in stop_sniffing, but let's ensure it here too
        self.stop_sniff_btn.setEnabled(False)
        self.status_bar.showMessage("Sniffing stopped.")
        self.sniffer_thread = None # Clear the reference to the finished thread

    def _update_sniffer_display(self):
        """Periodically called by a timer to batch-update the sniffer GUI."""
        with self.sniffer_buffer_lock:
            if not self.sniffer_packet_buffer:
                return
            # Quickly swap the buffer and release the lock
            packets_to_add = self.sniffer_packet_buffer
            self.sniffer_packet_buffer = []

        # Now process the packets without holding the lock
        items_to_add = []
        for packet in packets_to_add:
            self.packets_data.append(packet)
            n = len(self.packets_data)
            try:
                pt = f"{time.strftime('%H:%M:%S', time.localtime(packet.time))}.{int(packet.time * 1000) % 1000}"
                src = packet[IP].src if packet.haslayer(IP) else (packet[ARP].psrc if packet.haslayer(ARP) else "N/A")
                dst = packet[IP].dst if packet.haslayer(IP) else (packet[ARP].pdst if packet.haslayer(ARP) else "N/A")
                proto = packet.summary().split('/')[1].strip() if '/' in packet.summary() else "N/A"
                length = len(packet)
                item_data = [str(n), pt, src, dst, proto, str(length)]
            except Exception:
                item_data = [str(n), "Parse Error", "N/A", "N/A", "N/A", "N/A"]

            items_to_add.append(QTreeWidgetItem(item_data))

        self.packet_list_widget.addTopLevelItems(items_to_add)
        self.packet_list_widget.scrollToBottom()


    def add_packet_to_list(self, packet):
        """Callback function to add a sniffed packet to the UI list."""
        self.packets_data.append(packet); n = len(self.packets_data)
        try:
            pt = f"{time.strftime('%H:%M:%S', time.localtime(packet.time))}.{int(packet.time * 1000) % 1000}"
            src = packet[IP].src if packet.haslayer(IP) else (packet[ARP].psrc if packet.haslayer(ARP) else "N/A")
            dst = packet[IP].dst if packet.haslayer(IP) else (packet[ARP].pdst if packet.haslayer(ARP) else "N/A")
            proto = packet.summary().split('/')[1].strip() if '/' in packet.summary() else "N/A"
            length = len(packet)
        except Exception: pt, src, dst, proto, length = "Parse Error", "N/A", "N/A", "N/A", "N/A"
        item = QTreeWidgetItem([str(n), pt, src, dst, proto, str(length)]); self.packet_list_widget.addTopLevelItem(item); self.packet_list_widget.scrollToBottom()

    def display_packet_details(self, current_item, previous_item):
        """Displays the selected packet's details in the tree and hex views."""
        self.packet_details_tree.clear()
        self.packet_hex_view.clear()

        if not current_item:
            return

        try:
            packet_index = int(current_item.text(0)) - 1
            if not (0 <= packet_index < len(self.packets_data)):
                return

            packet = self.packets_data[packet_index]

            # Populate the hex view
            hex_dump = hexdump(packet, dump=True)
            self.packet_hex_view.setText(hex_dump)

            # Populate the details tree
            # We need to keep track of layer names to avoid duplicates from scapy's perspective
            layer_counts = {}
            current_layer = packet
            while current_layer:
                layer_name_raw = current_layer.name
                if layer_name_raw in layer_counts:
                    layer_counts[layer_name_raw] += 1
                    layer_name = f"{layer_name_raw} #{layer_counts[layer_name_raw]}"
                else:
                    layer_counts[layer_name_raw] = 1
                    layer_name = layer_name_raw

                layer_item = QTreeWidgetItem([layer_name])
                self.packet_details_tree.addTopLevelItem(layer_item)

                for field in current_layer.fields_desc:
                    field_name = field.name
                    try:
                        val = current_layer.getfieldval(field_name)
                        # i2repr is the standard Scapy way to get a display-friendly representation
                        display_value = field.i2repr(current_layer, val)
                    except Exception as e:
                        # Log the actual error for debugging, but still show a user-friendly message
                        logging.warning(f"Could not display field '{field_name}': {e}")
                        display_value = "Error reading value"

                    field_item = QTreeWidgetItem([field_name, display_value])
                    layer_item.addChild(field_item)

                layer_item.setExpanded(True)
                current_layer = current_layer.payload

            self.packet_details_tree.resizeColumnToContents(0)

        except (ValueError, IndexError):
            self.packet_details_tree.addTopLevelItem(QTreeWidgetItem(["Error displaying packet details."]))
        except Exception as e:
            logging.error(f"Unexpected error in display_packet_details: {e}", exc_info=True)
            self.packet_details_tree.addTopLevelItem(QTreeWidgetItem([f"Error: {e}"]))

    def clear_sniffer_display(self):
        self.packet_list_widget.clear(); self.packet_details_tree.clear(); self.packet_hex_view.clear(); self.packets_data.clear(); logging.info("Sniffer display cleared.")

    def save_packets(self):
        """Saves captured packets to a pcap file."""
        if not self.packets_data: QMessageBox.information(self, "Info", "There are no packets to save."); return
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Packets", "", "Pcap Files (*.pcap *.pcapng);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            try: wrpcap(file_path, self.packets_data); self.status_bar.showMessage(f"Saved {len(self.packets_data)} packets to {file_path}")
            except Exception as e: QMessageBox.critical(self, "Error", f"Failed to save packets: {e}")

    def load_packets(self):
        """Loads packets from a pcap file into the sniffer view."""
        if self.packets_data and QMessageBox.question(self, "Confirm", "Clear captured packets?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.No: return
        self.clear_sniffer_display()
        file_path, _ = QFileDialog.getOpenFileName(self, "Load Packets", "", "Pcap Files (*.pcap *.pcapng);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if file_path:
            try:
                loaded_packets = rdpcap(file_path)
                for packet in loaded_packets: self.add_packet_to_list(packet)
                self.status_bar.showMessage(f"Loaded {len(loaded_packets)} packets from {file_path}")
            except Exception as e: QMessageBox.critical(self, "Error", f"Failed to load packets: {e}")

    def crafter_add_layer(self):
        """Adds a new protocol layer to the packet being crafted."""
        proto_name = self.proto_to_add.currentText()
        if proto_name in AVAILABLE_PROTOCOLS:
            self.packet_layers.append(AVAILABLE_PROTOCOLS[proto_name]())
            self.crafter_rebuild_layer_list(); self.layer_list_widget.setCurrentRow(len(self.packet_layers) - 1)

    def crafter_remove_layer(self):
        """Removes the selected protocol layer from the packet."""
        if (row := self.layer_list_widget.currentRow()) >= 0:
            del self.packet_layers[row]; self.crafter_rebuild_layer_list(); self.crafter_clear_fields_display()

    def crafter_toggle_fuzz_layer(self):
        """Toggles fuzzing on the selected layer."""
        row = self.layer_list_widget.currentRow()
        if row < 0:
            QMessageBox.information(self, "Info", "Please select a layer to fuzz/unfuzz.")
            return

        layer = self.packet_layers[row]

        # Use hasattr to reliably check for fuzzed layers (duck typing)
        if hasattr(layer, 'obj'):
            # It's already fuzzed, so unfuzz it by replacing it with its original object
            self.packet_layers[row] = layer.obj
        else:
            # It's a normal layer, so wrap it with fuzz()
            self.packet_layers[row] = fuzz(layer)

        self.crafter_rebuild_layer_list()
        self.layer_list_widget.setCurrentRow(row)

    def crafter_rebuild_layer_list(self):
        """Updates the UI list of layers from the internal self.packet_layers."""
        self.layer_list_widget.clear()
        for i, layer in enumerate(self.packet_layers):
            if hasattr(layer, 'obj'):
                # Display fuzzed layers differently
                self.layer_list_widget.addItem(f"{i}: Fuzzed({layer.obj.name})")
            else:
                self.layer_list_widget.addItem(f"{i}: {layer.name}")
        self.crafter_update_packet_summary()

    def crafter_load_template(self, name):
        if self.packet_layers and QMessageBox.question(self, "Confirm", "Clear current packet stack?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.No: return
        self.packet_layers = [copy.deepcopy(l) for l in PACKET_TEMPLATES[name]]
        self.crafter_rebuild_layer_list()
        if self.packet_layers: self.layer_list_widget.setCurrentRow(0)

    def crafter_clear_fields_display(self):
        for widget in self.current_field_widgets: widget.deleteLater()
        self.current_field_widgets = []

    def crafter_display_layer_fields(self, row):
        self.crafter_clear_fields_display()
        if not (0 <= row < len(self.packet_layers)): return

        layer = self.packet_layers[row]

        if hasattr(layer, 'obj'):
            self.scroll_area.setEnabled(False)
            label = QLabel("Fields are not editable for fuzzed layers.")
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.fields_layout.addWidget(label)
            self.current_field_widgets.append(label)
            return

        self.scroll_area.setEnabled(True)
        for field in layer.fields_desc:
            container = QWidget(); hbox = QHBoxLayout(container); hbox.setContentsMargins(0,0,0,0)
            hbox.addWidget(QLabel(f"{field.name}:"))
            if isinstance(layer, TCP) and field.name == "flags":
                flags_widget = QWidget(); flags_layout = QHBoxLayout(flags_widget)
                self.tcp_flag_vars = {}
                for flag in "FSRPAUEC":
                    var = QCheckBox(flag); self.tcp_flag_vars[flag] = var
                    if flag in str(layer.flags): var.setChecked(True)
                    var.stateChanged.connect(lambda state, l=layer: self.crafter_update_tcp_flags(l))
                    flags_layout.addWidget(var)
                hbox.addWidget(flags_widget)
            else:
                le = QLineEdit(str(getattr(layer, field.name, ''))); le.editingFinished.connect(lambda l=layer, f=field.name, w=le: self.crafter_update_field(l, f, w.text()))
                hbox.addWidget(le)
            self.fields_layout.addWidget(container); self.current_field_widgets.append(container)

    def crafter_update_tcp_flags(self, layer):
        layer.flags = "".join([f for f, v in self.tcp_flag_vars.items() if v.isChecked()])
        self.crafter_update_packet_summary()

    def crafter_update_field(self, layer, field_name, text):
        try: setattr(layer, field_name, text)
        except: pass
        self.crafter_update_packet_summary()

    def build_packet(self):
        if not self.packet_layers: return None

        # Avoid deepcopying fuzz objects, as it can cause crashes.
        layers = []
        for l in self.packet_layers:
            if hasattr(l, 'obj'):
                layers.append(l)  # Use the fuzz object directly
            else:
                layers.append(copy.deepcopy(l))  # Deepcopy standard layers

        if not layers: return None

        pkt = layers[0]
        for i in range(1, len(layers)):
            pkt /= layers[i]
        return pkt

    def crafter_update_packet_summary(self):
        try: pkt = self.build_packet(); summary = pkt.summary() if pkt else "No layers."
        except Exception as e: summary = f"Error: {e}"
        self.crafter_summary.setPlainText(summary)

    def crafter_send_packet(self):
        """Starts the thread to send the crafted packet(s)."""
        if not self.packet_layers:
            QMessageBox.critical(self, "Error", "No packet layers to build a packet from.")
            return
        try:
            count, interval = int(self.send_count_edit.text()), float(self.send_interval_edit.text())
        except ValueError:
            QMessageBox.critical(self, "Error", "Invalid count or interval.")
            return
        self.send_results_widget.clear()
        self.send_btn.setEnabled(False)
        self.send_cancel_btn.setEnabled(True)
        self.tool_stop_event.clear()
        self.worker = WorkerThread(self._send_thread, args=(count, interval)); self.worker.start()

    def _send_thread(self, c, i):
        iface = self.get_selected_iface()
        q = self.tool_results_queue
        try:
            ans_list = []
            unans_list = []
            for pkt_num in range(c):
                if self.tool_stop_event.is_set():
                    logging.info("Packet sending cancelled.")
                    break

                pkt = self.build_packet()
                if not pkt:
                    logging.error("Failed to build packet in send thread.")
                    break

                send_receive_func = srp1 if pkt.haslayer(Ether) else sr1
                reply = send_receive_func(pkt, timeout=2, iface=iface, verbose=0)
                if reply:
                    ans_list.append((pkt, reply))
                else:
                    unans_list.append(pkt)
                time.sleep(i)
            q.put(('send_results', ans_list, unans_list))
        except Exception as e:
            logging.error("Send packet failed", exc_info=True)
            q.put(('error', 'Send Error', str(e)))
        finally:
            q.put(('send_finished',))

    def start_traceroute(self):
        """Starts the traceroute worker thread."""
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t = self.trace_target.text()
        if not t: QMessageBox.critical(self, "Error", "Please enter a target."); return
        self.trace_button.setEnabled(False)
        self.trace_cancel_button.setEnabled(True)
        self.is_tool_running = True
        self.tool_stop_event.clear()
        self.worker = WorkerThread(self._traceroute_thread, args=(t,)); self.worker.start()

    def _traceroute_thread(self,t):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"Traceroute thread started for target: {t} on iface: {iface}")
        results = []
        try:
            q.put(('trace_status',f"Resolving {t}...")); dest_ip=socket.gethostbyname(t)
            q.put(('trace_clear',));
            initial_msg = ("",f"Traceroute to {t} ({dest_ip})","","")
            q.put(('trace_result', initial_msg))
            results.append(initial_msg)

            for i in range(1,30):
                if self.tool_stop_event.is_set():
                    q.put(('trace_status', "Traceroute Canceled."))
                    break
                q.put(('trace_status',f"Sending probe to TTL {i}"))
                pkt=IP(dst=dest_ip,ttl=i)/UDP(dport=33434)
                st=time.time(); reply=sr1(pkt,timeout=2,iface=iface); rtt=(time.time()-st)*1000

                if reply is None:
                    result_tuple = (i,"* * *","Timeout","")
                else:
                    h_ip=reply.src
                    try: h_name,_,_=socket.gethostbyaddr(h_ip)
                    except socket.herror: h_name="Unknown"
                    result_tuple = (i,h_ip,h_name,f"{rtt:.2f}")

                q.put(('trace_result', result_tuple))
                results.append(result_tuple)
                if reply and (reply.type==3 or h_ip==dest_ip):
                    q.put(('trace_status',"Trace Complete."))
                    break
            else: q.put(('trace_status',"Trace Finished (Max hops reached)."))
        except Exception as e: logging.error("Exception in traceroute thread",exc_info=True); q.put(('error',"Traceroute Error",str(e)))
        finally:
            results_str = "\n".join([f"{hop} - {ip} - {name} - {rtt}ms" for hop, ip, name, rtt in results])
            q.put(('tool_finished','traceroute', t, results_str))
            logging.info("Traceroute thread finished.")

    def start_port_scan(self):
        """Starts the port scanner worker thread."""
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t=self.scan_target.text(); ps=self.scan_ports.text(); use_frags=self.scan_frag_check.isChecked()

        scan_protocols = []
        if self.scan_proto_tcp_radio.isChecked(): scan_protocols.append("TCP")
        if self.scan_proto_udp_radio.isChecked(): scan_protocols.append("UDP")
        if self.scan_proto_both_radio.isChecked(): scan_protocols.extend(["TCP", "UDP"])

        tcp_scan_type = self.tcp_scan_type_combo.currentText() if self.tcp_scan_type_combo.isVisible() else "SYN Scan"

        if not t or not ps: QMessageBox.critical(self, "Error", "Target and ports required."); return
        try: ports=sorted(list(set(self._parse_ports(ps))))
        except ValueError: QMessageBox.critical(self, "Error","Invalid port format. Use '22, 80, 100-200'."); return

        self.scan_button.setEnabled(False)
        self.scan_cancel_button.setEnabled(True)
        self.is_tool_running=True
        self.tool_stop_event.clear()

        args = (t, ports, scan_protocols, tcp_scan_type, use_frags)
        self.worker = WorkerThread(self._port_scan_thread, args=args); self.worker.start()

    def _parse_ports(self,ps):
        ports=[]
        for part in ps.split(','):
            part=part.strip()
            if '-' in part: start,end=map(int,part.split('-')); ports.extend(range(start,end+1))
            else: ports.append(int(part))
        return ports

    def _port_scan_thread(self,t,ports,scan_protocols,tcp_scan_type,use_frags):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"Port scan started: T={t}, P={ports}, Protocols={scan_protocols}, TCP_Mode={tcp_scan_type}, Frags={use_frags}")
        scan_results = []
        try:
            q.put(('scan_clear',))
            total_ports = len(ports) * len(scan_protocols)
            ports_scanned = 0

            tcp_scan_flags = {
                "SYN Scan": "S", "FIN Scan": "F", "Xmas Scan": "FPU",
                "Null Scan": "", "ACK Scan": "A"
            }

            for protocol in scan_protocols:
                if self.tool_stop_event.is_set(): break
                for port in ports:
                    if self.tool_stop_event.is_set(): break

                    ports_scanned += 1
                    status_msg = f"Scanning {t}:{port} ({protocol}"
                    if protocol == "TCP": status_msg += f"/{tcp_scan_type}"
                    status_msg += f") - {ports_scanned}/{total_ports}"
                    q.put(('scan_status', status_msg))

                    pkt = None
                    if protocol == "TCP":
                        flags = tcp_scan_flags.get(tcp_scan_type, "S")
                        pkt = IP(dst=t)/TCP(dport=port, flags=flags)
                    elif protocol == "UDP":
                        pkt = IP(dst=t)/UDP(dport=port)

                    if not pkt: continue

                    probes = fragment(pkt) if use_frags else [pkt]
                    # Only need one response, not for every fragment
                    resp=sr1(probes[0] if len(probes) == 1 else probes, timeout=1, iface=iface, verbose=0)
                    state = "No Response / Filtered"
                    if resp:
                        if resp.haslayer(TCP):
                            if resp.getlayer(TCP).flags == 0x12: state = "Open" # SYN-ACK
                            elif resp.getlayer(TCP).flags == 0x14: state = "Closed" # RST-ACK
                            elif resp.getlayer(TCP).flags == 0x4: state = "Unfiltered (RST)" # RST from ACK scan
                        elif resp.haslayer(UDP):
                            state = "Open | Filtered" # UDP is connectionless, open might not respond
                        elif resp.haslayer(ICMP) and resp.getlayer(ICMP).type == 3:
                            if resp.getlayer(ICMP).code in [1, 2, 3, 9, 10, 13]:
                                state = "Filtered"
                            else:
                                state = "Closed (ICMP)"

                    service = "Unknown"
                    if state.startswith("Open"):
                        try: service=socket.getservbyport(port, protocol.lower())
                        except OSError: pass

                    # Add to list for final popup and also to queue for live view
                    result_tuple = (f"{port}/{protocol.lower()}", state, service)
                    scan_results.append(result_tuple)
                    q.put(('scan_result', result_tuple))

            if self.tool_stop_event.is_set():
                q.put(('scan_status', "Scan Canceled."))
            else:
                q.put(('scan_status',"Scan Complete."))
                q.put(('show_port_scan_popup', scan_results, t)) # New message for popup
        except Exception as e: logging.error("Exception in port scan thread",exc_info=True); q.put(('error',"Scan Error",str(e)))
        finally:
            results_str = "\n".join([f"{p} - {s} - {svc}" for p, s, svc in scan_results])
            q.put(('tool_finished','scanner', t, results_str));
            logging.info("Port scan thread finished.")

    def start_arp_scan(self):
        """Starts the ARP scan worker thread."""
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t=self.arp_target.text()
        if not t: QMessageBox.critical(self, "Error", "Please enter a target network."); return
        self.arp_scan_button.setEnabled(False); self.is_tool_running=True
        self.worker = WorkerThread(self._arp_scan_thread, args=(t,)); self.worker.start()

    def _arp_scan_thread(self,t):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"ARP scan thread started for target: {t} on iface: {iface}")
        try:
            q.put(('arp_status', f"Scanning {t}...")); q.put(('arp_clear',))
            pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/ARP(pdst=t)
            ans,unans=srp(pkt,timeout=2,iface=iface,verbose=0)

            # Keep adding to tree for live results
            answered_results_for_tree = [{'ip': r.psrc, 'mac': r.hwsrc, 'status': 'Responded'} for s, r in ans]
            if answered_results_for_tree:
                q.put(('arp_results', answered_results_for_tree))

            # Now prepare results for popup
            popup_results = []
            q.put(('arp_status', f"Found {len(ans)} hosts. Resolving vendors..."))
            for i, (s, r) in enumerate(ans):
                q.put(('arp_status', f"Resolving vendor for {r.hwsrc} ({i+1}/{len(ans)})"))
                vendor = get_vendor(r.hwsrc)
                popup_results.append({'ip': r.psrc, 'mac': r.hwsrc, 'vendor': vendor})

            total_found = len(ans)
            q.put(('arp_status',f"Scan Complete. Found {total_found} active hosts."))
            q.put(('show_arp_scan_popup', popup_results, t)) # New message for popup

        except Exception as e: logging.error("Exception in ARP scan thread",exc_info=True); q.put(('error',"ARP Scan Error",str(e)))
        finally:
            results_str = "\n".join([f"{res['ip']} - {res['mac']} - {res['vendor']}" for res in popup_results])
            q.put(('tool_finished','arp_scan', t, results_str))
            logging.info("ARP scan thread finished.")

    def _create_arp_scan_cli_tool(self):
        """Creates the UI for the arp-scan CLI tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        config_widget, self.arp_scan_cli_controls = self._create_arp_scan_cli_config_widget()
        layout.addWidget(config_widget)

        # --- Action Buttons ---
        buttons_layout = QHBoxLayout()
        buttons_layout.addWidget(self.arp_scan_cli_controls['start_btn'])
        buttons_layout.addWidget(self.arp_scan_cli_controls['stop_btn'])
        layout.addLayout(buttons_layout)

        # --- Output Console ---
        self.arp_scan_cli_output_console = QPlainTextEdit()
        self.arp_scan_cli_output_console.setReadOnly(True)
        self.arp_scan_cli_output_console.setFont(QFont("Courier New", 10))
        self.arp_scan_cli_output_console.setPlaceholderText("arp-scan output will be displayed here...")
        layout.addWidget(self.arp_scan_cli_output_console, 1)

        self.arp_scan_cli_controls['start_btn'].clicked.connect(self.start_arp_scan_cli)
        self.arp_scan_cli_controls['stop_btn'].clicked.connect(self.cancel_tool)

        return widget

    def _create_arp_scan_cli_config_widget(self):
        """Creates a reusable, self-contained widget with arp-scan's configuration options."""
        widget = QWidget()
        main_layout = QFormLayout(widget)
        controls = {}

        controls['localnet_check'] = QCheckBox("Scan Local Network (--localnet)")
        controls['localnet_check'].setChecked(True)
        controls['localnet_check'].setToolTip("Automatically scan the network of the selected interface.")
        main_layout.addRow(controls['localnet_check'])

        controls['target_edit'] = QLineEdit()
        controls['target_edit'].setPlaceholderText("e.g., 192.168.1.0/24 (optional if --localnet is checked)")
        controls['target_edit'].setToolTip("Specify a custom target network or host if not using --localnet.")
        main_layout.addRow("Custom Target:", controls['target_edit'])

        controls['verbose_check'] = QCheckBox("Verbose Output (-v)")
        main_layout.addRow("--verbose:", controls['verbose_check'])

        # UI Logic
        def toggle_target_edit(checked):
            controls['target_edit'].setDisabled(checked)
        controls['localnet_check'].toggled.connect(toggle_target_edit)
        toggle_target_edit(True) # Initial state

        controls['start_btn'] = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls['stop_btn'] = QPushButton("Stop Scan"); controls['stop_btn'].setEnabled(False)

        return widget, controls

    def start_arp_scan_cli(self):
        """Starts the arp-scan CLI worker thread."""
        controls = self.arp_scan_cli_controls
        if not shutil.which("arp-scan"):
            QMessageBox.critical(self, "arp-scan Error", "'arp-scan' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        command = ["arp-scan"]

        iface = self.get_selected_iface()
        if iface:
            command.extend(["--interface", iface])

        if controls['localnet_check'].isChecked():
            command.append("--localnet")
        else:
            target = controls['target_edit'].text().strip()
            if not target:
                QMessageBox.critical(self, "Input Error", "A custom target is required if --localnet is not checked.")
                return
            command.append(target)

        if controls['verbose_check'].isChecked():
            command.append("--verbose")

        self.is_tool_running = True
        controls['start_btn'].setEnabled(False)
        controls['stop_btn'].setEnabled(True)
        self.tool_stop_event.clear()
        self.arp_scan_cli_output_console.clear()

        self.worker = WorkerThread(self._arp_scan_cli_thread, args=(command,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _arp_scan_cli_thread(self, command):
        """Worker thread for running the arp-scan command."""
        q = self.tool_results_queue
        logging.info(f"Starting arp-scan with command: {' '.join(command)}")
        q.put(('arp_scan_cli_output', f"$ {' '.join(command)}\n\n"))

        try:
            startupinfo = None
            if sys.platform == "win32":
                q.put(('error', 'Platform Error', 'arp-scan is not supported on Windows.'))
                q.put(('tool_finished', 'arp_scan_cli_scan'))
                return

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.arp_scan_cli_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('arp_scan_cli_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('arp_scan_cli_output', line))

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'arp-scan Error', "'arp-scan' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            logging.error(f"arp-scan thread error: {e}", exc_info=True)
            q.put(('error', 'arp-scan Error', str(e)))
        finally:
            q.put(('tool_finished', 'arp_scan_cli_scan'))
            with self.thread_finish_lock:
                self.arp_scan_cli_process = None
            logging.info("arp-scan scan thread finished.")

    def _handle_arp_scan_cli_output(self, line):
        self.arp_scan_cli_output_console.insertPlainText(line)
        self.arp_scan_cli_output_console.verticalScrollBar().setValue(self.arp_scan_cli_output_console.verticalScrollBar().maximum())

    def start_ping_sweep(self):
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target_network = self.ps_target_edit.text()
        probe_type = self.ps_probe_type_combo.currentText()
        ports_str = self.ps_ports_edit.text()

        try:
            net = ipaddress.ip_network(target_network, strict=False)
            timeout = float(self.ps_timeout_edit.text())
            num_threads = int(self.ps_threads_edit.text())
            ports = [int(p.strip()) for p in ports_str.split(',')] if ports_str else []
        except ValueError as e:
            QMessageBox.critical(self, "Input Error", f"Invalid input: {e}")
            return

        if ("TCP" in probe_type or "UDP" in probe_type) and not ports:
            QMessageBox.critical(self, "Input Error", "Please specify at least one port for TCP/UDP probes.")
            return

        self.is_tool_running = True
        self.ps_start_button.setEnabled(False)
        self.ps_cancel_button.setEnabled(True)
        self.tool_stop_event.clear()
        self.ps_tree.clear()

        args = (net, probe_type, ports, timeout, num_threads)
        self.worker = WorkerThread(self._ping_sweep_thread, args=args)
        self.worker.start()

    def _ping_sweep_thread(self, net, probe_type, ports, timeout, num_threads):
        """Master thread that populates a queue and starts worker threads."""
        q = self.tool_results_queue
        logging.info(f"Ping sweep started for {net} with {probe_type} on ports {ports}")

        hosts_queue = queue.Queue()
        for host in net.hosts():
            hosts_queue.put(str(host))

        if hosts_queue.qsize() == 0:
            q.put(('ps_status', "Sweep Complete (No hosts in range)."))
            q.put(('tool_finished', 'ping_sweep'))
            return

        self.ps_finished_threads = 0
        self.active_threads = []

        for i in range(num_threads):
            worker = WorkerThread(target=self._ping_sweep_worker, args=(hosts_queue, probe_type, ports, timeout, num_threads))
            self.active_threads.append(worker)
            worker.start()

    def _ping_sweep_worker(self, hosts_queue, probe_type, ports, timeout, num_threads):
        """Worker function that each ping sweep thread executes."""
        q = self.tool_results_queue
        while not self.tool_stop_event.is_set():
            try:
                host_str = hosts_queue.get_nowait()
            except queue.Empty:
                break # Queue is empty, this thread is done

            q.put(('ps_status', f"Pinging {host_str}..."))

            reply = None
            try:
                if probe_type == "ICMP Echo":
                    pkt = IP(dst=host_str)/ICMP()
                    reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                elif probe_type == "TCP SYN":
                    for port in ports:
                        pkt = IP(dst=host_str)/TCP(dport=port, flags="S")
                        reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                        if reply and reply.haslayer(TCP) and reply.getlayer(TCP).flags == 0x12: # SYN-ACK
                            break # Host is up, no need to check other ports
                elif probe_type == "TCP ACK":
                    for port in ports:
                        pkt = IP(dst=host_str)/TCP(dport=port, flags="A")
                        reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                        if reply and reply.haslayer(TCP) and reply.getlayer(TCP).flags == 0x4: # RST
                            break # Host is up, no need to check other ports
                elif probe_type == "UDP Probe":
                    for port in ports:
                        pkt = IP(dst=host_str)/UDP(dport=port)
                        reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                        if reply and reply.haslayer(ICMP) and reply.getlayer(ICMP).type == 3: # Dest Unreachable
                            break # Port is closed, but host is up.
            except Exception as e:
                logging.warning(f"Probe to {host_str} failed: {e}")


            if reply:
                q.put(('ps_result', (host_str, "Host is up")))

        # Signal that this worker is done
        q.put(('ps_worker_finished', num_threads))

    def load_flood_packet(self):
        packet=self.build_packet()
        if not packet: QMessageBox.critical(self, "Error", "Please craft a packet in the Packet Crafter tab first."); return
        self.loaded_flood_packet=packet
        self.flood_packet_label.setText(f"Loaded: {self.loaded_flood_packet.summary()}")
        logging.info(f"Loaded flood packet: {self.loaded_flood_packet.summary()}")

    def start_flood(self):
        """Starts the packet flooder worker threads."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        template = self.flood_template_combo.currentText()
        if template == "Custom (from Crafter)" and not self.loaded_flood_packet:
            QMessageBox.critical(self, "Error", "Please load a packet from the crafter first.")
            return

        warning_msg = "WARNING: This tool sends a high volume of packets and can disrupt network services. Only use this tool on networks you own or have explicit permission to test. Misuse of this tool may be illegal.\n\nDo you accept responsibility and wish to continue?"
        if not QMessageBox.question(self, "Ethical Use Warning", warning_msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.Yes:
            return

        try:
            count = int(self.flood_count.text())
            interval = float(self.flood_interval.text())
            num_threads = int(self.flood_threads.text())
            target_ip = self.flood_target_edit.text()
            target_port = int(self.flood_ports_edit.text()) if self.flood_ports_edit.text() else 80
            random_source = self.flood_rand_src_ip_check.isChecked()
        except ValueError:
            QMessageBox.critical(self, "Error", "Invalid count, interval, or thread number.")
            return

        self.flood_button.setEnabled(False)
        self.stop_flood_button.setEnabled(True)
        self.is_tool_running = True
        self.finished_thread_count = 0
        self.active_threads = []
        self.tool_stop_event.clear()

        packets_per_thread = count // num_threads
        extra_packets = count % num_threads

        flood_params = {
            "template": template, "target_ip": target_ip, "target_port": target_port,
            "random_source": random_source, "custom_packet": self.loaded_flood_packet
        }

        for i in range(num_threads):
            count_for_this_thread = packets_per_thread + (1 if i < extra_packets else 0)
            if count_for_this_thread == 0:
                continue

            worker = WorkerThread(self._flood_thread, args=(flood_params, count_for_this_thread, interval, num_threads))
            self.active_threads.append(worker)
            worker.start()

    def _flood_thread(self, params, count, interval, total_threads):
        q = self.tool_results_queue
        iface = self.get_selected_iface()
        logging.info(f"Flood thread started. Params: {params}, Count: {count}")
        try:
            q.put(('flood_status', f"Flooding with {count} packets..."))
            send_func = sendp # Assume Layer 2 for templates for now

            for i in range(count):
                if self.tool_stop_event.is_set():
                    logging.info("Flood thread detected stop event.")
                    break

                pkt = None
                template = params["template"]

                if template == "Custom (from Crafter)":
                    pkt = params["custom_packet"]
                    send_func = sendp if pkt.haslayer(Ether) else send
                else:
                    # On-the-fly packet creation for templates
                    src_ip = _get_random_ip() if params["random_source"] else "1.2.3.4" # Dummy IP if not random
                    target_ip = params["target_ip"]
                    target_port = params["target_port"]

                    if template == "TCP SYN Flood":
                        pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/IP(src=src_ip, dst=target_ip) / TCP(sport=RandShort(), dport=target_port, flags="S")
                    elif template == "UDP Flood":
                        pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/IP(src=src_ip, dst=target_ip) / UDP(sport=RandShort(), dport=target_port) / Raw(load=b"X"*1024)
                    elif template == "ICMP Echo Flood":
                        pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/IP(src=src_ip, dst=target_ip) / ICMP()

                if pkt:
                    send_func(pkt, iface=iface, verbose=0)

                time.sleep(interval)

        except Exception as e:
            logging.error("Exception in flood thread", exc_info=True)
            q.put(('error', "Flood Error", str(e)))
        finally:
            q.put(('flood_thread_finished', total_threads))
            logging.info("A flood thread finished.")

    def start_krack_scan(self):
        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Interface Error", "Please select a monitor-mode interface.")
            return

        self.krack_start_btn.setEnabled(False)
        self.krack_stop_btn.setEnabled(True)
        self.krack_results_tree.clear()

        self.krack_thread = KrackScanThread(iface, self)
        self.krack_thread.vulnerability_detected.connect(self.add_krack_result)
        self.krack_thread.start()

    def stop_krack_scan(self):
        if self.krack_thread and self.krack_thread.isRunning():
            self.krack_thread.stop()
            self.krack_thread.wait()
        self.krack_start_btn.setEnabled(True)
        self.krack_stop_btn.setEnabled(False)

    def add_krack_result(self, bssid, client_mac):
        # Avoid adding duplicates
        for i in range(self.krack_results_tree.topLevelItemCount()):
            item = self.krack_results_tree.topLevelItem(i)
            if item.text(0) == bssid and item.text(1) == client_mac:
                return # Already exists

        timestamp = time.strftime('%H:%M:%S')
        item = QTreeWidgetItem([bssid, client_mac, timestamp])
        self.krack_results_tree.addTopLevelItem(item)

    def start_firewall_test(self):
        """Starts the firewall testing worker thread."""
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t=self.fw_target.text(); ps_name=self.fw_probe_set.currentText()
        if not t: QMessageBox.critical(self, "Error", "Please enter a target."); return
        self.fw_test_button.setEnabled(False); self.is_tool_running=True
        self.worker = WorkerThread(self._firewall_test_thread, args=(t,ps_name)); self.worker.start()

    def _firewall_test_thread(self,t,ps_name):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"Firewall test thread started for target {t}, probe set {ps_name}")
        results = []
        try:
            q.put(('fw_clear',)); q.put(('fw_status',f"Testing {ps_name}..."))
            probe_set = FIREWALL_PROBES[ps_name]
            for i, (pkt_builder, desc) in enumerate(probe_set):
                q.put(('fw_status',f"Sending probe {i+1}/{len(probe_set)}: {desc}"))

                pkt = pkt_builder(t)
                pkt_summary = ""

                if isinstance(pkt, list): # It's a fragmented packet
                    pkt_summary = f"{len(pkt)} fragments"
                    ans, unans = sr(pkt, timeout=2, iface=iface, verbose=0)
                    resp = ans[0][1] if ans else None # Take the first response as representative
                else: # It's a single packet
                    pkt_summary = pkt.summary()
                    resp = sr1(pkt, timeout=2, iface=iface, verbose=0)

                result = "Responded" if resp is not None else "No Response / Blocked"
                result_tuple = (desc, pkt_summary, result)
                q.put(('fw_result', result_tuple))
                results.append(result_tuple)
            q.put(('fw_status',"Firewall Test Complete."))
        except Exception as e: logging.error("Exception in firewall test thread",exc_info=True); q.put(('error',"Firewall Test Error",str(e)))
        finally:
            results_str = "\n".join([f"{desc} - {summary} - {res}" for desc, summary, res in results])
            q.put(('tool_finished','fw_tester', t, results_str))
            logging.info("Firewall test thread finished.")

    def start_wifi_scan(self):
        """Starts the Wi-Fi scanner and channel hopper threads."""
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return

        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Warning", "Please select a wireless interface for scanning.")
            return

        self.wifi_scan_button.setEnabled(False)
        self.wifi_scan_stop_button.setEnabled(True)
        self.is_tool_running = True
        self.found_networks = {}
        self.wifi_tree.clear()

        self.sniffer_thread = SnifferThread(iface=iface, handler=self._wifi_scan_handler, bpf_filter="type mgt subtype beacon or type mgt subtype probe-resp")
        self.sniffer_thread.start()
        self.channel_hopper = ChannelHopperThread(iface)
        self.channel_hopper.start()

        self.tool_results_queue.put(('wifi_scan_status', 'Scanning... Press Stop to finish.'))

        # We can still have a timeout as a safeguard, but the user can now stop it.
        self.scan_timer = QTimer(self)
        self.scan_timer.setSingleShot(True)
        self.scan_timer.timeout.connect(self.stop_wifi_scan)
        self.scan_timer.start(30000) # 30 second safeguard timer

    def _wifi_scan_handler(self, pkt):
        if pkt.haslayer(Dot11Beacon) or pkt.haslayer(Dot11ProbeResp):
            bssid = pkt[Dot11].addr2
            if bssid not in self.found_networks:
                try: ssid = pkt[Dot11Elt].info.decode(errors="ignore")
                except: ssid = "<Hidden>"
                if not ssid: ssid = "<Hidden>"

                channel = "N/A"
                try:
                    elt = pkt.getlayer(Dot11Elt, ID=3)
                    if elt: channel = ord(elt.info)
                except: pass
                signal = "N/A"
                try: signal = pkt[RadioTap].dbm_antsignal
                except: pass
                self.found_networks[bssid] = (ssid, bssid, channel, signal)
                self.tool_results_queue.put(('wifi_scan_update', self.found_networks[bssid]))

    def stop_wifi_scan(self):
        if hasattr(self, 'scan_timer') and self.scan_timer.isActive():
            self.scan_timer.stop()

        if self.sniffer_thread and self.sniffer_thread.isRunning():
            self.sniffer_thread.stop()
            self.sniffer_thread.wait()
        if self.channel_hopper and self.channel_hopper.isRunning():
            self.channel_hopper.stop()
            self.channel_hopper.wait()

        self.tool_results_queue.put(('wifi_scan_status', 'Scan Finished.'))
        self.tool_results_queue.put(('tool_finished', 'wifi_scan'))

    def start_deauth(self):
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        bssid = self.deauth_bssid.text(); client = self.deauth_client.text()
        try: count = int(self.deauth_count.text())
        except ValueError: QMessageBox.critical(self, "Error", "Count must be an integer."); return
        warning_msg="This will send deauthentication packets which can disrupt a network. Are you sure you want to continue?"
        if QMessageBox.question(self, "Confirm Deauth", warning_msg) == QMessageBox.StandardButton.No: return
        self.deauth_button.setEnabled(False); self.is_tool_running = True
        args = (bssid, client, count)
        self.worker = WorkerThread(self._deauth_thread, args=args); self.worker.start()

    def _deauth_thread(self, bssid, client, count):
        q = self.tool_results_queue; iface = self.get_selected_iface()
        logging.info(f"Deauth thread started: BSSID={bssid}, Client={client}, Count={count}")
        try:
            pkt = RadioTap()/Dot11(type=0, subtype=12, addr1=client, addr2=bssid, addr3=bssid)/Dot11Deauth(reason=7)
            q.put(('deauth_status', f"Sending {count} deauth packets..."))
            sendp(pkt, iface=iface, count=count, inter=0.1, verbose=0)
            q.put(('deauth_status', "Deauth packets sent."))
        except Exception as e: logging.error("Exception in deauth thread", exc_info=True); q.put(('error',"Deauth Error",str(e)))
        finally: q.put(('tool_finished','deauth')); logging.info("Deauth thread finished.")

    def start_beacon_flood(self):
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Interface Error", "Please select a monitor-mode interface.")
            return

        bssid = self.bf_bssid_edit.text()
        enc_type = self.bf_enc_combo.currentText()

        # Handle SSIDs
        ssids = []
        if self.bf_ssid_edit.isReadOnly(): # Loaded from file
            ssids = self.bf_ssid_list
        else:
            ssids = [self.bf_ssid_edit.text().strip()]

        if not ssids or not ssids[0]:
            QMessageBox.critical(self, "Input Error", "Please provide at least one SSID.")
            return

        try:
            count = int(self.bf_count_edit.text())
            interval = float(self.bf_interval_edit.text())
            channel = int(self.bf_channel_edit.text())
            if not (1 <= channel <= 14):
                raise ValueError("Channel must be between 1 and 14.")
        except ValueError as e:
            QMessageBox.critical(self, "Input Error", f"Invalid input for Count, Interval, or Channel: {e}")
            return

        self.is_tool_running = True
        self.bf_start_button.setEnabled(False)
        self.bf_stop_button.setEnabled(True)
        self.tool_stop_event.clear()

        args = (iface, ssids, bssid, count, interval, enc_type, channel)
        self.worker = WorkerThread(self._beacon_flood_thread, args=args)
        self.worker.start()

    def _build_beacon_frame(self, ssid, bssid, channel, enc_type):
        dot11 = Dot11(type=0, subtype=8, addr1='ff:ff:ff:ff:ff:ff', addr2=bssid, addr3=bssid)

        cap = 'ESS'
        if enc_type != "Open":
            cap += '+privacy'

        beacon = Dot11Beacon(cap=cap)
        essid = Dot11Elt(ID='SSID', info=ssid)
        ds_param = Dot11Elt(ID='DSset', info=chr(channel))

        frame = RadioTap() / dot11 / beacon / essid / ds_param

        if enc_type == "WEP":
            # WEP is signaled by the privacy bit in the capability field alone.
            pass
        elif enc_type == "WPA2-PSK":
            rsn_info = Dot11Elt(ID='RSNinfo', info=(
                b'\x01\x00'      # RSN Version 1
                b'\x00\x0f\xac\x04'  # Group Cipher Suite: AES (CCMP)
                b'\x01\x00'      # 1 Pairwise Cipher Suite
                b'\x00\x0f\xac\x04'  # AES (CCMP)
                b'\x01\x00'      # 1 Authentication Key Management Suite (AKM)
                b'\x00\x0f\xac\x02'  # PSK
                b'\x00\x00'      # RSN Capabilities
            ))
            frame /= rsn_info
        elif enc_type == "WPA3-SAE":
            rsn_info = Dot11Elt(ID='RSNinfo', info=(
                b'\x01\x00'      # RSN Version 1
                b'\x00\x0f\xac\x04'  # Group Cipher Suite: AES (CCMP)
                b'\x01\x00'      # 1 Pairwise Cipher Suite
                b'\x00\x0f\xac\x04'  # AES (CCMP)
                b'\x01\x00'      # 1 Authentication Key Management Suite (AKM)
                b'\x00\x0f\xac\x08'  # SAE
                b'\x8c\x00'      # RSN Capabilities (MFPC, MFPR)
            ))
            frame /= rsn_info

        return frame

    def _beacon_flood_thread(self, iface, ssids, bssid, count, interval, enc_type, channel):
        q = self.tool_results_queue
        logging.info(f"Beacon flood started: SSIDs={len(ssids)}, BSSID={bssid}, Count={count}, Enc={enc_type}")

        sent_count = 0
        ssid_index = 0
        infinite_mode = (count == 0)

        try:
            while not self.tool_stop_event.is_set():
                if not infinite_mode and sent_count >= count:
                    break

                current_bssid = RandMAC() if bssid.lower() == 'random' else bssid
                current_ssid = ssids[ssid_index]

                beacon_frame = self._build_beacon_frame(current_ssid, current_bssid, channel, enc_type)

                sendp(beacon_frame, iface=iface, verbose=0)
                sent_count += 1
                ssid_index = (ssid_index + 1) % len(ssids) # Cycle through SSIDs

                status_msg = f"Flooding {current_ssid}... (Packets sent: {sent_count})"
                if not infinite_mode:
                    status_msg += f" / {count}"
                q.put(('bf_status', status_msg))

                time.sleep(interval)

            if self.tool_stop_event.is_set():
                q.put(('bf_status', "Beacon flood canceled."))
            else:
                q.put(('bf_status', "Beacon flood complete."))

        except Exception as e:
            logging.error("Exception in beacon flood thread", exc_info=True)
            q.put(('error', "Beacon Flood Error", str(e)))
        finally:
            q.put(('tool_finished', 'beacon_flood'))
            logging.info("Beacon flood thread finished.")

    def _arp_spoof_thread(self, victim_ip, target_ip):
        q = self.tool_results_queue
        iface = self.get_selected_iface()
        logging.info(f"ARP spoof thread started for Victim={victim_ip}, Target={target_ip}")

        try:
            q.put(('arp_spoof_status', "Resolving MAC addresses..."))
            victim_mac = getmacbyip(victim_ip)
            target_mac = getmacbyip(target_ip)

            if not victim_mac or not target_mac:
                raise Exception("Could not resolve MAC address for one or both targets. Are they online?")

            q.put(('arp_spoof_status', f"Victim: {victim_mac} | Target: {target_mac}"))
            logging.info(f"Resolved MACs -> Victim: {victim_mac}, Target: {target_mac}")

            victim_packet = Ether(dst=victim_mac)/ARP(op=2, pdst=victim_ip, hwdst=victim_mac, psrc=target_ip)
            target_packet = Ether(dst=target_mac)/ARP(op=2, pdst=target_ip, hwdst=target_mac, psrc=victim_ip)

            sent_count = 0
            while not self.tool_stop_event.is_set():
                sendp(victim_packet, iface=iface, verbose=0)
                sendp(target_packet, iface=iface, verbose=0)
                sent_count += 2
                q.put(('arp_spoof_status', f"Spoofing active... (Packets sent: {sent_count})"))
                time.sleep(2)

        except Exception as e:
            logging.error("Exception in ARP spoof thread", exc_info=True)
            q.put(('error', "ARP Spoof Error", str(e)))
        finally:
            q.put(('tool_finished', 'arp_spoof'))
            logging.info("ARP spoof thread finished.")

    def start_arp_spoof(self):
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        victim_ip = self.arp_spoof_victim_ip.text()
        target_ip = self.arp_spoof_target_ip.text()

        if not victim_ip or not target_ip:
            QMessageBox.critical(self, "Error", "Victim IP and Target IP are required.")
            return

        warning_msg = """
        <p>You are about to perform an ARP Spoofing attack. This will intercept traffic between the two targets and constitutes a Man-in-the-Middle attack.</p>
        <p>Ensure you have <b>explicit, written permission</b> to test on this network. Misuse of this tool is illegal.</p>
        <p><b>Do you accept full responsibility and wish to continue?</b></p>
        """
        if QMessageBox.question(self, "Ethical Use Confirmation", warning_msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.No:
            return

        self.arp_spoof_current_victim = victim_ip
        self.arp_spoof_current_target = target_ip

        self.is_tool_running = True
        self.arp_spoof_start_btn.setEnabled(False)
        self.arp_spoof_stop_btn.setEnabled(True)
        self.tool_stop_event.clear()

        args = (victim_ip, target_ip)
        self.worker = WorkerThread(self._arp_spoof_thread, args=args)
        self.worker.start()

    def stop_arp_spoof(self):
        if self.is_tool_running:
            logging.info("User requested to stop ARP spoofing.")
            self.arp_spoof_status.setText("Stopping...")
            self.tool_stop_event.set()

    def _restore_arp(self, victim_ip, target_ip):
        iface = self.get_selected_iface()
        logging.info(f"Attempting to restore ARP tables for {victim_ip} and {target_ip}")
        try:
            victim_mac = getmacbyip(victim_ip)
            target_mac = getmacbyip(target_ip)

            if not victim_mac or not target_mac:
                raise Exception("Could not resolve MACs for restoration. Manual correction may be needed.")

            # Create the legitimate ARP packets
            restore_victim_packet = Ether(dst=victim_mac)/ARP(op=2, pdst=victim_ip, hwdst=victim_mac, psrc=target_ip, hwsrc=target_mac)
            restore_target_packet = Ether(dst=target_mac)/ARP(op=2, pdst=target_ip, hwdst=target_mac, psrc=victim_ip, hwsrc=victim_mac)

            # Send them multiple times to ensure the cache is corrected
            sendp([restore_victim_packet, restore_target_packet], count=5, inter=0.2, iface=iface, verbose=0)

            logging.info("ARP restoration packets sent.")
            self.arp_spoof_status.setText("ARP tables restored. Attack stopped.")

        except Exception as e:
            logging.error(f"Failed to restore ARP tables: {e}", exc_info=True)
            QMessageBox.critical(self, "Restore Error", f"Could not restore ARP tables: {e}")

    def cancel_tool(self):
        if self.is_tool_running:
            logging.info("User requested to cancel the current tool.")
            self.tool_stop_event.set()

            # Special handling for subprocesses that need to be terminated directly
            # This is more robust than only relying on the thread's check loop.
            with self.thread_finish_lock:
                if hasattr(self, 'nmap_process') and self.nmap_process and self.nmap_process.poll() is None:
                    try:
                        self.nmap_process.terminate()
                        logging.info("Nmap process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Nmap process: {e}")

                if hasattr(self, 'arp_scan_cli_process') and self.arp_scan_cli_process and self.arp_scan_cli_process.poll() is None:
                    try:
                        self.arp_scan_cli_process.terminate()
                        logging.info("arp-scan process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating arp-scan process: {e}")
                if hasattr(self, 'spiderfoot_process') and self.spiderfoot_process and self.spiderfoot_process.poll() is None:
                    try:
                        self.spiderfoot_process.terminate()
                        logging.info("Spiderfoot process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Spiderfoot process: {e}")
                if hasattr(self, 'sherlock_process') and self.sherlock_process and self.sherlock_process.poll() is None:
                    try:
                        self.sherlock_process.terminate()
                        logging.info("Sherlock process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Sherlock process: {e}")
                if hasattr(self, 'fierce_process') and self.fierce_process and self.fierce_process.poll() is None:
                    try:
                        self.fierce_process.terminate()
                        logging.info("fierce process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating fierce process: {e}")
                if hasattr(self, 'dnsrecon_process') and self.dnsrecon_process and self.dnsrecon_process.poll() is None:
                    try:
                        self.dnsrecon_process.terminate()
                        logging.info("dnsrecon process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating dnsrecon process: {e}")
                if hasattr(self, 'enum4linux_ng_process') and self.enum4linux_ng_process and self.enum4linux_ng_process.poll() is None:
                    try:
                        self.enum4linux_ng_process.terminate()
                        logging.info("enum4linux-ng process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating enum4linux-ng process: {e}")
                if hasattr(self, 'hydra_process') and self.hydra_process and self.hydra_process.poll() is None:
                    try:
                        self.hydra_process.terminate()
                        logging.info("Hydra process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Hydra process: {e}")
                if hasattr(self, 'jtr_process') and self.jtr_process and self.jtr_process.poll() is None:
                    try:
                        self.jtr_process.terminate()
                        logging.info("JTR process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating JTR process: {e}")
                if hasattr(self, 'ffuf_process') and self.ffuf_process and self.ffuf_process.poll() is None:
                    try:
                        self.ffuf_process.terminate()
                        logging.info("ffuf process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating ffuf process: {e}")
                if hasattr(self, 'dirsearch_process') and self.dirsearch_process and self.dirsearch_process.poll() is None:
                    try:
                        self.dirsearch_process.terminate()
                        logging.info("dirsearch process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating dirsearch process: {e}")
                if hasattr(self, 'rustscan_process') and self.rustscan_process and self.rustscan_process.poll() is None:
                    try:
                        self.rustscan_process.terminate()
                        logging.info("RustScan process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating RustScan process: {e}")
                if hasattr(self, 'trufflehog_process') and self.trufflehog_process and self.trufflehog_process.poll() is None:
                    try:
                        self.trufflehog_process.terminate()
                        logging.info("TruffleHog process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating TruffleHog process: {e}")
                if hasattr(self, 'httpx_process') and self.httpx_process and self.httpx_process.poll() is None:
                    try:
                        self.httpx_process.terminate()
                        logging.info("httpx process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating httpx process: {e}")
                if hasattr(self, 'subfinder_process') and self.subfinder_process and self.subfinder_process.poll() is None:
                    try:
                        self.subfinder_process.terminate()
                        logging.info("Subfinder process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Subfinder process: {e}")

                if hasattr(self, 'sublist3r_process') and self.sublist3r_process and self.sublist3r_process.poll() is None:
                    try:
                        self.sublist3r_process.terminate()
                        logging.info("Sublist3r process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Sublist3r process: {e}")

                if hasattr(self, 'wifite_process') and self.wifite_process and self.wifite_process.poll() is None:
                    try:
                        os.kill(self.wifite_process.pid, signal.SIGTERM)
                        logging.info("Wifite process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating wifite process: {e}")

                if hasattr(self, 'nikto_process') and self.nikto_process and self.nikto_process.poll() is None:
                    try:
                        self.nikto_process.terminate()
                        logging.info("Nikto process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Nikto process: {e}")

                if hasattr(self, 'gobuster_process') and self.gobuster_process and self.gobuster_process.poll() is None:
                    try:
                        self.gobuster_process.terminate()
                        logging.info("Gobuster process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Gobuster process: {e}")

                if hasattr(self, 'sqlmap_process') and self.sqlmap_process and self.sqlmap_process.poll() is None:
                    try:
                        self.sqlmap_process.terminate()
                        logging.info("SQLMap process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating SQLMap process: {e}")

                if hasattr(self, 'whatweb_process') and self.whatweb_process and self.whatweb_process.poll() is None:
                    try:
                        self.whatweb_process.terminate()
                        logging.info("WhatWeb process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating WhatWeb process: {e}")

                if hasattr(self, 'hashcat_process') and self.hashcat_process and self.hashcat_process.poll() is None:
                    try:
                        self.hashcat_process.terminate()
                        logging.info("Hashcat process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Hashcat process: {e}")

                if hasattr(self, 'nuclei_process') and self.nuclei_process and self.nuclei_process.poll() is None:
                    try:
                        self.nuclei_process.terminate()
                        logging.info("Nuclei process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Nuclei process: {e}")
                if hasattr(self, 'masscan_process') and self.masscan_process and self.masscan_process.poll() is None:
                    try:
                        self.masscan_process.terminate()
                        logging.info("Masscan process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Masscan process: {e}")

                if hasattr(self, 'masscan_process') and self.masscan_process and self.masscan_process.poll() is None:
                    try:
                        self.masscan_process.terminate()
                        logging.info("Masscan process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Masscan process: {e}")

                if hasattr(self, 'masscan_process') and self.masscan_process and self.masscan_process.poll() is None:
                    try:
                        self.masscan_process.terminate()
                        logging.info("Masscan process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Masscan process: {e}")

    def _show_port_scan_summary_popup(self, results, target):
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Port Scan Results for {target}")
        dialog.setMinimumSize(400, 300)
        layout = QVBoxLayout(dialog)

        categorized = {"Open": [], "Open | Filtered": [], "Closed": [], "Filtered": [], "Unfiltered (RST)": [], "No Response / Filtered": []}
        for port, state, service in results:
            # Normalize states
            normalized_state = state
            if "No Response" in state:
                normalized_state = "No Response / Filtered"

            if normalized_state in categorized:
                categorized[normalized_state].append(f"{port} ({service})")
            else: # Fallback for any unexpected state
                if "Other" not in categorized: categorized["Other"] = []
                categorized["Other"].append(f"{port} ({state}, {service})")


        text_browser = QTextBrowser()
        text_browser.setOpenExternalLinks(False)
        html = f"<h1>Scan Report: {target}</h1>"
        # Display open ports first
        if categorized["Open"]:
            html += f"<h2>Open Ports ({len(categorized['Open'])})</h2>"
            html += "<ul>" + "".join(f"<li>{p}</li>" for p in sorted(categorized['Open'])) + "</ul>"

        for state, ports in categorized.items():
            if state != "Open" and ports:
                html += f"<h2>{state} ({len(ports)})</h2>"
                html += "<ul>" + "".join(f"<li>{p}</li>" for p in sorted(ports)) + "</ul>"

        text_browser.setHtml(html)
        layout.addWidget(text_browser)

        button_layout = QHBoxLayout()

        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(lambda: self._send_to_ai_analyst("port_scanner", results, context=target))
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(dialog.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

        dialog.exec()

    def _show_arp_scan_summary_popup(self, results, target):
        dialog = QDialog(self)
        dialog.setWindowTitle(f"ARP Scan Results for {target}")
        dialog.setMinimumSize(500, 400)
        layout = QVBoxLayout(dialog)

        summary_label = QLabel(f"<b>Found {len(results)} active hosts on network {target}.</b>")
        layout.addWidget(summary_label)

        tree = QTreeWidget()
        tree.setColumnCount(3)
        tree.setHeaderLabels(["IP Address", "MAC Address", "Vendor"])
        for res in results:
            item = QTreeWidgetItem([res['ip'], res['mac'], res['vendor']])
            tree.addTopLevelItem(item)
        tree.resizeColumnToContents(0)
        tree.resizeColumnToContents(1)
        layout.addWidget(tree)

        export_button = self._create_export_button(tree) # Reuse export functionality
        layout.addWidget(export_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(dialog.accept)
        layout.addWidget(ok_button)

        dialog.exec()

    def _process_tool_results(self):
        """Processes results from worker threads via a queue using a handler dictionary."""
        while not self.tool_results_queue.empty():
            msg = self.tool_results_queue.get()
            msg_type = msg[0]

            # Prioritize exact matches
            if msg_type in self.result_handlers:
                # Unpack arguments; msg[1:] creates a tuple of the remaining elements
                self.result_handlers[msg_type](*msg[1:])
                continue

            # Check for suffix-based dynamic handlers
            matched = False
            for suffix, handler in self.dynamic_handlers.items():
                if msg_type.endswith(suffix):
                    tool_name = msg_type.rsplit(suffix, 1)[0]
                    handler(tool_name, *msg[1:])
                    matched = True
                    break

            if not matched:
                logging.warning(f"No handler found for message type: {msg_type}")

    def _setup_result_handlers(self):
        """Initializes the dictionary mapping result queue messages to handler functions."""
        self.result_handlers = {
            # Exact message matches
            'send_results': self._handle_send_results,
            'send_finished': self._handle_send_finished,
            'tool_finished': self._handle_tool_finished,
            'report_finished': self._handle_report_finished,
            'flood_thread_finished': self._handle_flood_thread_finished,
            'ps_worker_finished': self._handle_ps_worker_finished,
            'crunch_finished': self._handle_crunch_finished,
            'show_port_scan_popup': self._show_port_scan_summary_popup,
            'show_arp_scan_popup': self._show_arp_scan_summary_popup,
            'arp_results': self._handle_arp_results,
            'error': self._handle_error,
        }
        # Handlers for dynamic message types that end with a specific suffix
        self.dynamic_handlers = {
            'lab_status': self._handle_lab_status,
            'cve_search_status': self._handle_cve_search_status,
            'cve_result': self._handle_cve_result,
            'exploit_search_status': self._handle_exploit_search_status,
            'exploit_search_results': self._handle_exploit_search_results,
            '_status': self._handle_status_update,
            '_clear': self._handle_clear_update,
            '_result': self._handle_result_update,
            '_update': self._handle_result_update, # Catches 'wifi_scan_update'
        }
        self.result_handlers['nmap_output'] = self._handle_nmap_output
        self.result_handlers['nmap_xml_result'] = self._handle_nmap_xml_result
        self.result_handlers['sublist3r_output'] = self._handle_sublist3r_output
        self.result_handlers['sublist3r_results'] = self._show_subdomain_results_popup
        self.result_handlers['subfinder_output'] = self._handle_subfinder_output
        self.result_handlers['subdomain_results'] = self._show_subdomain_results_popup
        self.result_handlers['httpx_output'] = self._handle_httpx_output
        self.result_handlers['httpx_results'] = self._show_httpx_results_popup
        self.result_handlers['trufflehog_output'] = self._handle_trufflehog_output
        self.result_handlers['trufflehog_results'] = self._show_trufflehog_results_popup
        self.result_handlers['rustscan_output'] = self._handle_rustscan_output
        self.result_handlers['dirsearch_output'] = self._handle_dirsearch_output
        self.result_handlers['dirsearch_results'] = self._show_dirsearch_results_popup
        self.result_handlers['ffuf_output'] = self._handle_ffuf_output
        self.result_handlers['ffuf_results'] = self._show_ffuf_results_popup
        self.result_handlers['jtr_output'] = self._handle_jtr_output
        self.result_handlers['hydra_output'] = self._handle_hydra_output
        self.result_handlers['enum4linux_ng_output'] = self._handle_enum4linux_ng_output
        self.result_handlers['enum4linux_ng_results'] = self._show_enum4linux_ng_results_popup
        self.result_handlers['dnsrecon_output'] = self._handle_dnsrecon_output
        self.result_handlers['dnsrecon_results'] = self._show_dnsrecon_results_popup
        self.result_handlers['sherlock_output'] = self._handle_sherlock_output
        self.result_handlers['sherlock_results'] = self._show_sherlock_results_popup
        self.result_handlers['spiderfoot_output'] = self._handle_spiderfoot_output
        self.result_handlers['arp_scan_cli_output'] = self._handle_arp_scan_cli_output
        self.result_handlers['fierce_output'] = self._handle_fierce_output
        self.result_handlers['wifite_output'] = self._handle_wifite_output
        self.result_handlers['nikto_output'] = self._handle_nikto_output
        self.result_handlers['gobuster_output'] = self._handle_gobuster_output
        self.result_handlers['sqlmap_output'] = self._handle_sqlmap_output
        self.result_handlers['whatweb_output'] = self._handle_whatweb_output
        self.result_handlers['hashcat_output'] = self._handle_hashcat_output
        self.result_handlers['nuclei_output'] = self._handle_nuclei_output
        self.result_handlers['nuclei_results'] = self._show_nuclei_results_popup
        self.result_handlers['masscan_output'] = self._handle_masscan_output
        self.result_handlers['report_finding'] = self._handle_report_finding
        self.result_handlers['recent_threats_result'] = self._handle_recent_threats_result

    def _handle_report_finished(self, success, message):
        """Handles the result of the report generation thread."""
        self.report_generate_btn.setEnabled(True)
        if success:
            QMessageBox.information(self, "Report Generated", f"Report successfully saved to:\n{message}")
            self.status_bar.showMessage("Report generation successful.", 5000)
        else:
            QMessageBox.critical(self, "Report Generation Error", f"Failed to generate report:\n{message}")
            self.status_bar.showMessage("Report generation failed.", 5000)

    def _handle_report_finding(self, finding_data):
        """Adds a finding to the report tree. finding_data is a tuple."""
        # The data should be (host, port_service, finding, details)
        item = QTreeWidgetItem([str(col) for col in finding_data])
        self.report_findings_tree.addTopLevelItem(item)

    def _handle_aggregation(self):
        """Starts the background thread for aggregating and enriching tool results."""
        if not self.nmap_last_xml:
            QMessageBox.warning(self, "No Data", "No Nmap scan data is available to analyze. Please run an Nmap scan with XML output first.")
            return

        self.report_aggregate_btn.setEnabled(False)
        self.report_findings_tree.clear()
        self.status_bar.showMessage("Aggregating and enriching results...")

        self.worker = WorkerThread(self._aggregation_thread)
        self.active_threads.append(self.worker)
        self.worker.start()

    def _aggregation_thread(self):
        """Parses tool outputs, enriches data, and sends it to the reporting tab."""
        q = self.tool_results_queue
        use_offline_db = self.offline_cve_check.isChecked()
        try:
            if not LXML_AVAILABLE:
                q.put(('error', 'Dependency Error', "The 'lxml' library is required for report generation. Please install it."))
                return

            parser = etree.XMLParser(recover=True, no_network=True, dtd_validation=False)
            root = etree.fromstring(self.nmap_last_xml.encode('utf-8'), parser=parser)

            for host in root.findall('host'):
                if host.find('status').get('state') != 'up':
                    continue

                address = host.find('address').get('addr')
                ports_elem = host.find('ports')
                if ports_elem is None:
                    continue

                for port in ports_elem.findall('port'):
                    if port.find('state').get('state') != 'open':
                        continue

                    port_id = port.get('portid')
                    protocol = port.get('protocol')
                    port_service_str = f"{port_id}/{protocol}"

                    service_elem = port.find('service')
                    if service_elem is not None:
                        service_name = service_elem.get('name', 'N/A')
                        product = service_elem.get('product', '')
                        version = service_elem.get('version', '')

                        # Construct a search term, prioritizing product and version
                        search_term = f"{product} {version}".strip()
                        if not search_term:
                            search_term = service_name

                        full_service_str = f"{service_name} ({search_term})"

                        # Enrich data with CVEs and Exploits
                        if use_offline_db:
                            cve_details = self._query_local_cve_db(search_term)
                        else:
                            cve_details = self._query_cve_api(search_term)
                        exploit_details = self._query_searchsploit(search_term)

                        q.put(('report_finding', (address, port_service_str, full_service_str, f"{cve_details}\n{exploit_details}")))

        except Exception as e:
            logging.error(f"Error during result aggregation: {e}", exc_info=True)
            q.put(('error', 'Aggregation Error', f'An error occurred: {e}'))
        finally:
            q.put(('tool_finished', 'aggregation'))

    def _handle_generate_report(self):
        """Handles the report generation process by gathering data and starting a worker."""
        # 1. Gather all data from the UI
        report_data = {
            "client": self.report_client_name.text(),
            "dates": self.report_assessment_dates.text(),
            "objectives": self.report_objectives.toPlainText(),
            "in_scope": self.report_in_scope.toPlainText(),
            "out_of_scope": self.report_out_of_scope.toPlainText(),
            "summary": self.report_summary_text.toPlainText(),
            "findings": []
        }

        if self.report_findings_tree.topLevelItemCount() == 0:
            QMessageBox.warning(self, "No Data", "There are no findings to report. Please run the 'Aggregate & Enrich Results' tool first.")
            return

        for i in range(self.report_findings_tree.topLevelItemCount()):
            item = self.report_findings_tree.topLevelItem(i)
            report_data["findings"].append({
                "host": item.text(0),
                "service": item.text(1),
                "finding": item.text(2),
                "details": item.text(3)
            })

        # 2. Prompt user for save location
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Report", "GScapy_Report.html", "HTML Files (*.html)", options=QFileDialog.Option.DontUseNativeDialog)
        if not file_path:
            return

        template_name = self.report_template_combo.currentText()

        # 3. Start worker thread
        self.status_bar.showMessage("Generating report...")
        self.report_generate_html_btn.setEnabled(False)
        self.worker = WorkerThread(self._report_generation_thread, args=(report_data, file_path, template_name))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _start_cve_db_update(self):
        """Starts the background thread to download/update the offline CVE database."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        self.is_tool_running = True
        self.update_cve_db_btn.setEnabled(False)
        self.status_bar.showMessage("Starting offline CVE database update...")

        self.worker = WorkerThread(self._update_cve_db_thread)
        self.active_threads.append(self.worker)
        self.worker.start()

    def _update_cve_db_thread(self):
        """Worker thread to download and process NVD data feeds into a local SQLite DB."""
        q = self.tool_results_queue
        db_path = "cve.db"
        try:
            # Step 1: Setup Database
            con = sqlite3.connect(db_path)
            cur = con.cursor()
            cur.execute("""
                CREATE TABLE IF NOT EXISTS vulnerabilities (
                    cve_id TEXT PRIMARY KEY,
                    description TEXT,
                    cvss_v3_score REAL,
                    cvss_v2_score REAL,
                    keywords TEXT,
                    published_date TEXT
                )
            """)
            cur.execute("CREATE INDEX IF NOT EXISTS idx_keywords ON vulnerabilities(keywords);")
            con.commit()

            # Step 2: Define URLs
            current_year = datetime.now().year
            base_url = "https://nvd.nist.gov/feeds/json/cve/1.1/nvdcve-1.1-{year}.json.gz"
            urls = [base_url.format(year=y) for y in range(2002, current_year + 1)]
            urls.append("https://nvd.nist.gov/feeds/json/cve/1.1/nvdcve-1.1-modified.json.gz")

            total_files = len(urls)
            # Step 3: Download and Process each file
            for i, url in enumerate(urls):
                if self.tool_stop_event.is_set():
                    q.put(('status', "CVE DB update cancelled."))
                    break

                filename = url.split('/')[-1]
                q.put(('status', f"Downloading and processing file {i+1}/{total_files}: {filename}"))

                req = urllib.request.Request(url, headers={'User-Agent': 'GScapy/1.0'})
                with urllib.request.urlopen(req, timeout=60) as response:
                    compressed_file = response.read()

                json_data = gzip.decompress(compressed_file)
                data = json.loads(json_data)

                cve_items = data.get('CVE_Items', [])
                cves_to_insert = []
                for cve_item in cve_items:
                    cve_id = cve_item['cve']['CVE_data_meta']['ID']
                    description = cve_item['cve']['description']['description_data'][0]['value']

                    # Extract keywords from CPEs
                    keywords = set()
                    nodes = cve_item.get('configurations', {}).get('nodes', [])
                    for node in nodes:
                        cpe_matches = node.get('cpe_match', [])
                        for cpe_match in cpe_matches:
                            cpe_uri = cpe_match.get('cpe23Uri', '')
                            # cpe:2.3:a:vendor:product:version:update:edition:language:sw_edition:target_sw:target_hw:other
                            parts = cpe_uri.split(':')
                            if len(parts) > 4:
                                keywords.add(parts[3]) # vendor
                                keywords.add(parts[4]) # product

                    keywords_str = " ".join(sorted(list(keywords)))

                    # Extract CVSS scores
                    metrics_v3 = cve_item.get('impact', {}).get('baseMetricV3', {})
                    cvss_v3_score = metrics_v3.get('cvssV3', {}).get('baseScore')
                    metrics_v2 = cve_item.get('impact', {}).get('baseMetricV2', {})
                    cvss_v2_score = metrics_v2.get('cvssV2', {}).get('baseScore')

                    published_date = cve_item.get('publishedDate', '')

                    cves_to_insert.append((cve_id, description, cvss_v3_score, cvss_v2_score, keywords_str, published_date))

                # Step 4: Insert into DB
                cur.executemany("INSERT OR REPLACE INTO vulnerabilities VALUES (?, ?, ?, ?, ?, ?)", cves_to_insert)
                con.commit()
                q.put(('status', f"Finished processing {filename}. {len(cves_to_insert)} records updated."))

            if not self.tool_stop_event.is_set():
                q.put(('status', "CVE Database update complete!"))

        except Exception as e:
            logging.error(f"Failed to update offline CVE DB: {e}", exc_info=True)
            q.put(('error', 'CVE DB Error', str(e)))
        finally:
            if 'con' in locals() and con:
                con.close()
            q.put(('tool_finished', 'cve_db_update'))

    def _query_local_cve_db(self, keyword):
        """Queries the local SQLite CVE database for a given keyword."""
        db_path = "cve.db"
        if not os.path.exists(db_path):
            return "Offline CVE database (cve.db) not found. Please update it first."

        try:
            con = sqlite3.connect(db_path)
            cur = con.cursor()

            search_terms = keyword.split()
            query = "SELECT cve_id, cvss_v3_score, description FROM vulnerabilities WHERE "
            query += " AND ".join(["keywords LIKE ?"] * len(search_terms))
            query += " ORDER BY cvss_v3_score DESC LIMIT 5"

            params = [f"%{term}%" for term in search_terms]

            cur.execute(query, params)
            results = cur.fetchall()
            con.close()

            if not results:
                return "No CVEs found in offline DB."

            output = ["--- CVEs (Offline) ---"]
            for cve_id, score, description in results:
                output.append(f"{cve_id} (Score: {score or 'N/A'}): {description[:100]}...")
            return "\n".join(output)

        except Exception as e:
            logging.error(f"Offline CVE DB query for '{keyword}' failed: {e}")
            return f"Offline CVE lookup failed: {e}"

    def _report_generation_thread(self, report_data, file_path, template_name):
        """Worker thread to generate the final HTML report."""
        q = self.tool_results_queue
        try:
            template_path = os.path.join("report_templates", template_name)
            with open(template_path, 'r', encoding='utf-8') as f:
                template_html = f.read()

            findings_html = ""
            for finding in report_data['findings']:
                findings_html += "<tr>\n"
                findings_html += f"    <td>{finding['host']}</td>\n"
                findings_html += f"    <td>{finding['service']}</td>\n"
                findings_html += f"    <td>{finding['finding']}</td>\n"
                findings_html += f"    <td><pre>{finding['details']}</pre></td>\n"
                findings_html += "</tr>\n"

            # Sanitize and format the data
            sanitized_data = {k: v.replace('<', '&lt;').replace('>', '&gt;') if isinstance(v, str) else v for k, v in report_data.items()}

            final_html = template_html.format(
                client=sanitized_data['client'],
                dates=sanitized_data['dates'],
                objectives=sanitized_data['objectives'],
                in_scope=sanitized_data['in_scope'],
                out_of_scope=sanitized_data['out_of_scope'],
                summary=sanitized_data['summary'].replace('\n', '<br>'),
                findings_loop=findings_html
            )

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(final_html)

            q.put(('report_finished', True, file_path))

        except Exception as e:
            logging.error(f"Error during report generation: {e}", exc_info=True)
            q.put(('report_finished', False, str(e)))

    def _gather_report_data(self):
        """Gathers all data from the reporting UI fields into a dictionary."""
        if self.report_findings_tree.topLevelItemCount() == 0:
            QMessageBox.warning(self, "No Data", "There are no findings to report. Please run the 'Aggregate & Enrich Results' tool first.")
            return None

        report_data = {
            "client_name": self.report_client_name.text(),
            "assessment_dates": self.report_assessment_dates.text(),
            "objectives": self.report_objectives.toPlainText(),
            "in_scope": self.report_in_scope.toPlainText(),
            "out_of_scope": self.report_out_of_scope.toPlainText(),
            "summary": self.report_summary_text.toPlainText(),
            "findings": []
        }

        for i in range(self.report_findings_tree.topLevelItemCount()):
            item = self.report_findings_tree.topLevelItem(i)
            report_data["findings"].append({
                "host": item.text(0),
                "service": item.text(1),
                "finding": item.text(2),
                "details": item.text(3)
            })
        return report_data

    def _handle_generate_doc_report(self, file_format):
        """Handles the generation of reports in different document formats."""
        report_data = self._gather_report_data()
        if not report_data:
            return

        file_path, _ = QFileDialog.getSaveFileName(self, f"Save Report as {file_format.upper()}", f"GScapy_Report.{file_format}", f"{file_format.upper()} Files (*.{file_format})", options=QFileDialog.Option.DontUseNativeDialog)
        if not file_path:
            return

        self.status_bar.showMessage(f"Generating {file_format.upper()} report...")

        try:
            if file_format == 'docx':
                self._export_report_to_docx(report_data, file_path)
            elif file_format == 'pdf':
                self._export_report_to_pdf(report_data, file_path)

            QMessageBox.information(self, "Success", f"Report successfully saved to:\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate {file_format.upper()} report: {e}")
        finally:
            self.status_bar.showMessage("Report generation finished.", 5000)

    def _export_report_to_docx(self, data, file_path):
        """Exports the structured report data to a DOCX file."""
        document = docx.Document()
        document.add_heading(f"Penetration Test Report: {data['client_name']}", 0)

        document.add_heading("Executive Summary", level=1)
        document.add_paragraph(data['summary'])

        document.add_heading("Scope", level=1)
        document.add_paragraph(f"Assessment Dates: {data['assessment_dates']}")
        document.add_heading("Objectives", level=2)
        document.add_paragraph(data['objectives'])
        document.add_heading("In-Scope Targets", level=2)
        document.add_paragraph(data['in_scope'])
        document.add_heading("Out-of-Scope Targets", level=2)
        document.add_paragraph(data['out_of_scope'])

        document.add_heading("Detailed Findings", level=1)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Host'
        hdr_cells[1].text = 'Service'
        hdr_cells[2].text = 'Finding'
        hdr_cells[3].text = 'Details'

        for finding in data['findings']:
            row_cells = table.add_row().cells
            row_cells[0].text = finding['host']
            row_cells[1].text = finding['service']
            row_cells[2].text = finding['finding']
            row_cells[3].text = finding['details']

        document.save(file_path)

    def _export_report_to_pdf(self, data, file_path):
        """Exports the structured report data to a PDF file."""
        doc = SimpleDocTemplate(file_path)
        styles = getSampleStyleSheet()
        elements = [Paragraph(f"Penetration Test Report: {data['client_name']}", styles['h1'])]

        elements.append(Paragraph("Executive Summary", styles['h2']))
        elements.append(Paragraph(data['summary'].replace('\n', '<br/>'), styles['BodyText']))

        elements.append(Paragraph("Scope", styles['h2']))
        elements.append(Paragraph(f"<b>Assessment Dates:</b> {data['assessment_dates']}", styles['BodyText']))
        elements.append(Paragraph("Objectives", styles['h3']))
        elements.append(Paragraph(data['objectives'].replace('\n', '<br/>'), styles['BodyText']))
        elements.append(Paragraph("In-Scope Targets", styles['h3']))
        elements.append(Paragraph(data['in_scope'].replace('\n', '<br/>'), styles['BodyText']))
        elements.append(Paragraph("Out-of-Scope Targets", styles['h3']))
        elements.append(Paragraph(data['out_of_scope'].replace('\n', '<br/>'), styles['BodyText']))

        elements.append(Paragraph("Detailed Findings", styles['h2']))

        table_data = [['Host', 'Service', 'Finding', 'Details']]
        for f in data['findings']:
            details_p = Paragraph(f['details'], styles['BodyText'])
            table_data.append([f['host'], f['service'], f['finding'], details_p])

        table = Table(table_data, colWidths=[1.5*inch, 1.5*inch, 2*inch, 2.5*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        elements.append(table)

        doc.build(elements)

    def _handle_ai_report_generation(self):
        """Gathers findings and instructions, and sends them to the AI for analysis."""
        report_data = self._gather_report_data()
        if not report_data:
            return

        persona = self.ai_persona_combo.currentText()
        custom_instructions = self.ai_instructions_edit.toPlainText()

        findings_text = ""
        for i, finding in enumerate(report_data['findings']):
            findings_text += f"Finding {i+1}:\n"
            findings_text += f"  Host: {finding['host']}\n"
            findings_text += f"  Service: {finding['service']}\n"
            findings_text += f"  Vulnerability: {finding['finding']}\n"
            findings_text += f"  Details: {finding['details']}\n\n"

        prompt = (
            f"You are an AI assistant. Please analyze the following penetration testing findings from the perspective of a **{persona}**. "
            "Your goal is to provide a detailed analysis and recommendations based on these findings.\n\n"
        )

        if custom_instructions:
            prompt += f"Please follow these custom instructions: '{custom_instructions}'\n\n"

        prompt += (
            "--- FINDINGS ---\n"
            f"{findings_text}"
            "--- END FINDINGS ---\n\n"
            "Please provide your analysis."
        )

        self.ai_assistant_tab.send_message(prompt)
        self.tab_widget.setCurrentWidget(self.ai_assistant_tab)
        QMessageBox.information(self, "AI Task Started", "The AI is analyzing the findings. You can see the results in the 'AI Assistant' tab.")

    def _query_cve_api(self, keyword):
        """Queries the NVD CVE API for a given keyword and returns a formatted string."""
        if not keyword:
            return "No service info to query CVEs."
        try:
            # URL-encode the keyword to handle spaces and special characters
            encoded_keyword = urllib.parse.quote(keyword)
            url = f"https://services.nvd.nist.gov/rest/json/cves/2.0?keywordSearch={encoded_keyword}&resultsPerPage=5"

            # Add a user-agent to be compliant with API usage policies
            req = urllib.request.Request(url, headers={'User-Agent': 'GScapy/1.0'})

            with urllib.request.urlopen(req, timeout=10) as response:
                data = json.load(response)

            vulnerabilities = data.get('vulnerabilities', [])
            if not vulnerabilities:
                return "No CVEs found."

            output = ["--- CVEs ---"]
            for item in vulnerabilities:
                cve = item.get('cve', {})
                cve_id = cve.get('id', 'N/A')

                # Get description
                description = "No description available."
                for desc in cve.get('descriptions', []):
                    if desc.get('lang') == 'en':
                        description = desc.get('value', '')
                        break

                # Get CVSS V3 score if available, otherwise V2
                cvss_score = "N/A"
                if 'cvssMetricV31' in cve.get('metrics', {}):
                    cvss_score = cve['metrics']['cvssMetricV31'][0]['cvssData']['baseScore']
                elif 'cvssMetricV2' in cve.get('metrics', {}):
                    cvss_score = cve['metrics']['cvssMetricV2'][0]['cvssData']['baseScore']

                output.append(f"{cve_id} (Score: {cvss_score}): {description[:100]}...")
            return "\n".join(output)

        except Exception as e:
            logging.error(f"CVE API query for '{keyword}' failed: {e}")
            return f"CVE lookup failed: {e}"

    def _query_searchsploit(self, keyword):
        """Queries the local searchsploit database and returns a formatted string."""
        if not keyword:
            return "No service info to query exploits."

        if not shutil.which("searchsploit"):
            return "searchsploit command not found. Please install Exploit-DB."

        try:
            command = ["searchsploit", "--json", keyword]
            process = subprocess.run(command, capture_output=True, text=True, timeout=15)

            if process.returncode != 0:
                return f"Searchsploit error: {process.stderr}"

            data = json.loads(process.stdout)
            results = data.get('RESULTS_EXPLOIT', [])

            if not results:
                return "No exploits found."

            output = ["--- Exploits ---"]
            for item in results[:5]: # Limit to top 5 results
                title = item.get('Title', 'N/A')
                edb_id = item.get('EDB-ID', 'N/A')
                output.append(f"EDB-ID: {edb_id} - {title}")
            return "\n".join(output)

        except Exception as e:
            logging.error(f"Searchsploit query for '{keyword}' failed: {e}")
            return f"Exploit lookup failed: {e}"

    def _show_subdomain_results_popup(self, domain, subdomains):
        """Shows the results of a subdomain scan in a dedicated dialog."""
        if not subdomains:
            QMessageBox.information(self, "No Results", f"No subdomains were found for {domain}.")
            return
        dialog = SubdomainResultsDialog(subdomains, domain, self)
        dialog.exec()

    def _handle_sublist3r_output(self, line):
        self.sublist3r_output.insertPlainText(line)
        self.sublist3r_output.verticalScrollBar().setValue(self.sublist3r_output.verticalScrollBar().maximum())

    def _handle_wifite_output(self, line):
        self.wifite_output_console.insertPlainText(line)
        self.wifite_output_console.verticalScrollBar().setValue(self.wifite_output_console.verticalScrollBar().maximum())

    def _handle_nikto_output(self, line):
        self.nikto_output_console.insertPlainText(line)
        self.nikto_output_console.verticalScrollBar().setValue(self.nikto_output_console.verticalScrollBar().maximum())

    def _handle_gobuster_output(self, line):
        self.gobuster_output_console.insertPlainText(line)
        self.gobuster_output_console.verticalScrollBar().setValue(self.gobuster_output_console.verticalScrollBar().maximum())

    def _handle_sqlmap_output(self, line):
        self.sqlmap_output_console.insertPlainText(line)
        self.sqlmap_output_console.verticalScrollBar().setValue(self.sqlmap_output_console.verticalScrollBar().maximum())

    def _handle_whatweb_output(self, line):
        self.whatweb_output_console.insertPlainText(line)
        self.whatweb_output_console.verticalScrollBar().setValue(self.whatweb_output_console.verticalScrollBar().maximum())

    def _handle_hashcat_output(self, line):
        self.hashcat_output_console.insertPlainText(line)
        self.hashcat_output_console.verticalScrollBar().setValue(self.hashcat_output_console.verticalScrollBar().maximum())

    def _handle_masscan_output(self, line):
        self.masscan_output_console.insertPlainText(line)
        self.masscan_output_console.verticalScrollBar().setValue(self.masscan_output_console.verticalScrollBar().maximum())

    def _handle_nmap_output(self, line):
        self.nmap_output_console.insertPlainText(line)
        self.nmap_output_console.verticalScrollBar().setValue(self.nmap_output_console.verticalScrollBar().maximum())

    def _handle_nmap_xml_result(self, xml_content):
        """Stores the captured Nmap XML report and shows a summary dialog."""
        self.nmap_last_xml = xml_content
        logging.info(f"Captured Nmap XML report ({len(xml_content)} bytes).")
        self.status_bar.showMessage("Nmap scan complete. XML report captured.", 5000)
        self.nmap_controls['report_btn'].setEnabled(True)

        # Automatically show the summary dialog
        target_context = self.nmap_controls['target_edit'].text()
        summary_dialog = NmapSummaryDialog(xml_content, target_context, self)
        summary_dialog.exec()

    def _show_subdomain_results_popup(self, domain, subdomains):
        """Shows the results of a subdomain scan in a dedicated dialog."""
        if not subdomains:
            QMessageBox.information(self, "No Results", f"No subdomains were found for {domain}.")
            return
        dialog = SubdomainResultsDialog(subdomains, domain, self)
        dialog.exec()

    def _show_httpx_results_popup(self, json_data):
        """Shows the results of an httpx scan in a dedicated dialog."""
        dialog = HttpxResultsDialog(json_data, self)
        dialog.exec()

    def _show_dirsearch_results_popup(self, json_data, target_context):
        """Shows the results of a dirsearch scan in a dedicated dialog."""
        dialog = DirsearchResultsDialog(json_data, target_context, self)
        dialog.exec()

    def _show_ffuf_results_popup(self, json_data):
        """Shows the results of an ffuf scan in a dedicated dialog."""
        dialog = FfufResultsDialog(json_data, self)
        dialog.exec()

    def _show_nuclei_results_popup(self, json_data):
        """Shows the results of a nuclei scan in a dedicated dialog."""
        dialog = NucleiResultsDialog(json_data, self)
        dialog.exec()

    def _show_trufflehog_results_popup(self, json_data):
        """Shows the results of a trufflehog scan in a dedicated dialog."""
        dialog = TruffleHogResultsDialog(json_data, self)
        dialog.exec()

    def _show_enum4linux_ng_results_popup(self, json_data, target_context):
        """Shows the results of an enum4linux-ng scan in a dedicated dialog."""
        dialog = Enum4LinuxNGResultsDialog(json_data, target_context, self)
        dialog.exec()

    def _show_dnsrecon_results_popup(self, json_data, target_context):
        """Shows the results of a dnsrecon scan in a dedicated dialog."""
        dialog = DnsReconResultsDialog(json_data, target_context, self)
        dialog.exec()

    def _show_sherlock_results_popup(self, csv_data, target_context):
        """Shows the results of a sherlock scan in a dedicated dialog."""
        dialog = SherlockResultsDialog(csv_data, target_context, self)
        dialog.exec()

    def _handle_send_results(self, ans, unans):
        self.send_results_widget.clear()
        for i, (s, r) in enumerate(ans):
            self.send_results_widget.addTopLevelItem(QTreeWidgetItem([str(i+1), s.summary(), r.summary()]))
        start_num = len(ans)
        for i, s in enumerate(unans):
            self.send_results_widget.addTopLevelItem(QTreeWidgetItem([str(start_num+i+1), s.summary(), "No response"]))

    def _handle_send_finished(self):
        self.send_btn.setEnabled(True)
        self.send_cancel_btn.setEnabled(False)

    def _handle_status_update(self, tool_name, status_text):
        widgets = {'trace': self.trace_status, 'scan': self.scan_status, 'arp': self.arp_status,
                   'flood': self.flood_status, 'fw': self.fw_status, 'wifi_scan': self.wifi_scan_status,
                   'deauth': self.deauth_status, 'arp_spoof': self.arp_spoof_status,
                   'bf': self.bf_status_label, 'ps': self.ps_status_label}
        if tool_name in widgets:
            widgets[tool_name].setText(status_text)

    def _handle_clear_update(self, tool_name):
        widgets = {'trace': self.trace_tree, 'scan': self.scan_tree, 'arp': self.arp_tree,
                   'fw': self.fw_tree, 'wifi_scan': self.wifi_tree}
        if tool_name in widgets:
            widgets[tool_name].clear()

    def _handle_result_update(self, tool_name, result_data):
        widgets = {'trace': self.trace_tree, 'scan': self.scan_tree, 'fw': self.fw_tree,
                   'wifi_scan': self.wifi_tree, 'ps': self.ps_tree}
        if tool_name in widgets:
            widgets[tool_name].addTopLevelItem(QTreeWidgetItem([str(x) for x in result_data]))

    def _handle_arp_results(self, results):
        for res in results:
            self.arp_tree.addTopLevelItem(QTreeWidgetItem([res['ip'], res['mac'], res['status']]))

    def _handle_crunch_finished(self, outfile, returncode):
        if returncode == 0:
            self.wpa_crack_output.appendPlainText(f"Crunch finished successfully. Wordlist saved to:\n{outfile}")
            self.wpa_wordlist_edit.setText(outfile)
        else:
            self.wpa_crack_output.appendPlainText(f"Crunch finished with an error (code: {returncode}). Check gscapy.log for details.")

    def _handle_ps_worker_finished(self, total_threads):
        with self.ps_thread_lock:
            self.ps_finished_threads += 1
            if self.ps_finished_threads >= total_threads:
                if self.tool_stop_event.is_set():
                    self.ps_status_label.setText("Ping sweep canceled.")
                else:
                    self.ps_status_label.setText("Ping sweep complete.")
                self.tool_results_queue.put(('tool_finished', 'ping_sweep'))

    def _handle_flood_thread_finished(self, total_threads):
        with self.thread_finish_lock:
            self.finished_thread_count += 1
            if self.finished_thread_count >= total_threads:
                self.is_tool_running = False
                self.flood_button.setEnabled(True)
                self.stop_flood_button.setEnabled(False)
                if self.tool_stop_event.is_set():
                    self.flood_status.setText("Flood Canceled.")
                else:
                    self.flood_status.setText("Flood complete.")

    def _handle_tool_finished(self, tool, target=None, results=""):
        """
        Handles the 'tool_finished' signal from any worker thread.
        Resets UI state and logs the test to history if applicable.
        """
        # Log the action to the history database
        if self.current_user and target is not None:
            # Ensure results are a string and not too long for the DB
            if not isinstance(results, str):
                results = json.dumps(results, indent=2) # Pretty print JSON results

            # Truncate long results to avoid massive DB entries
            if len(results) > 10000: # 10KB limit
                results = results[:10000] + "\n... (truncated)"

            database.log_test_to_history(self.current_user['id'], tool, target, results)

        if tool == 'aggregation':
            self.report_aggregate_btn.setEnabled(True)
            self.status_bar.showMessage("Result aggregation and enrichment complete.", 5000)
            return

        if tool == 'cve_db_update':
            self.update_cve_db_btn.setEnabled(True)
            self.status_bar.showMessage("CVE DB update finished.", 5000)
            self.is_tool_running = False # Explicitly set this as it's a special case
            return

        self.is_tool_running = False
        buttons = {'traceroute': self.trace_button, 'scanner': self.scan_button, 'arp_scan': self.arp_scan_button,
                   'flooder': self.flood_button, 'fw_tester': self.fw_test_button, 'wifi_scan': self.wifi_scan_button,
                   'deauth': self.deauth_button, 'arp_spoof': self.arp_spoof_start_btn,
                   'beacon_flood': self.bf_start_button, 'ping_sweep': self.ps_start_button, 'nmap_scan': self.nmap_controls['start_btn'],
                          'sublist3r_scan': self.subdomain_controls['start_btn'], 'subfinder_scan': self.subfinder_controls['start_btn'], 'httpx_scan': self.httpx_controls['start_btn'], 'trufflehog_scan': self.trufflehog_controls['start_btn'], 'rustscan_scan': self.rustscan_controls['start_btn'], 'dirsearch_scan': self.dirsearch_controls['start_btn'], 'ffuf_scan': self.ffuf_controls['start_btn'], 'jtr_scan': self.jtr_controls['start_btn'], 'hydra_scan': self.hydra_controls['start_btn'], 'enum4linux_ng_scan': self.enum4linux_ng_controls['start_btn'], 'dnsrecon_scan': self.dnsrecon_controls['start_btn'], 'fierce_scan': self.fierce_controls['start_btn'], 'sherlock_scan': self.sherlock_controls['start_btn'], 'spiderfoot_scan': self.spiderfoot_controls['start_btn'], 'arp_scan_cli_scan': self.arp_scan_cli_controls['start_btn'], 'wifite_scan': self.wifite_start_btn, 'nikto_scan': self.nikto_controls['start_btn'], 'gobuster_scan': self.gobuster_controls['start_btn'], 'sqlmap_scan': self.sqlmap_start_btn, 'whatweb_scan': self.whatweb_start_btn, 'hashcat_scan': self.hashcat_start_btn, 'masscan_scan': self.masscan_start_btn, 'nuclei_scan': self.nuclei_controls['start_btn'],
                           'cve_search': self.cve_search_button, 'exploit_search': self.exploitdb_search_button, 'cve_db_update': self.update_cve_db_btn}
        cancel_buttons = {'scanner': self.scan_cancel_button, 'flooder': self.stop_flood_button,
                          'arp_spoof': self.arp_spoof_stop_btn, 'beacon_flood': self.bf_stop_button,
                          'ping_sweep': self.ps_cancel_button, 'fw_tester': self.fw_cancel_button,
                          'traceroute': self.trace_cancel_button, 'wifi_scan': self.wifi_scan_stop_button, 'nmap_scan': self.nmap_controls['cancel_btn'],
                          'sublist3r_scan': self.subdomain_controls['cancel_btn'], 'subfinder_scan': self.subfinder_controls['cancel_btn'], 'httpx_scan': self.httpx_controls['cancel_btn'], 'trufflehog_scan': self.trufflehog_controls['stop_btn'], 'rustscan_scan': self.rustscan_controls['cancel_btn'], 'dirsearch_scan': self.dirsearch_controls['stop_btn'], 'ffuf_scan': self.ffuf_controls['stop_btn'], 'jtr_scan': self.jtr_controls['stop_btn'], 'hydra_scan': self.hydra_controls['stop_btn'], 'enum4linux_ng_scan': self.enum4linux_ng_controls['stop_btn'], 'dnsrecon_scan': self.dnsrecon_controls['stop_btn'], 'fierce_scan': self.fierce_controls['stop_btn'], 'sherlock_scan': self.sherlock_controls['stop_btn'], 'spiderfoot_scan': self.spiderfoot_controls['stop_btn'], 'arp_scan_cli_scan': self.arp_scan_cli_controls['stop_btn'], 'wifite_scan': self.wifite_stop_btn, 'nikto_scan': self.nikto_controls['stop_btn'], 'gobuster_scan': self.gobuster_controls['stop_btn'], 'sqlmap_scan': self.sqlmap_stop_btn, 'whatweb_scan': self.whatweb_stop_btn, 'hashcat_scan': self.hashcat_stop_btn, 'masscan_scan': self.masscan_stop_btn, 'nuclei_scan': self.nuclei_controls['stop_btn']}

        if tool == 'lab_chain':
            self.lab_run_chain_btn.setEnabled(True)
            self.status_bar.showMessage("LAB chain finished.", 5000)
            return

        if tool == 'jtr_scan':
            self.jtr_controls['start_btn'].setEnabled(True)
            self.jtr_controls['show_btn'].setEnabled(True)
            self.jtr_controls['stop_btn'].setEnabled(False)
            return

        if tool == 'arp_spoof':
            if self.arp_spoof_current_victim and self.arp_spoof_current_target:
                self._restore_arp(self.arp_spoof_current_victim, self.arp_spoof_current_target)
                self.arp_spoof_current_victim = None
                self.arp_spoof_current_target = None

        if tool in buttons:
            buttons[tool].setEnabled(True)
        if tool in cancel_buttons:
            cancel_buttons[tool].setEnabled(False)

        if self.tool_stop_event.is_set():
            status_labels = {'scanner': self.scan_status, 'traceroute': self.trace_status}
            if tool in status_labels:
                status_labels[tool].setText("Canceled by user.")

    def _handle_cve_search_status(self, status_text):
        self.status_bar.showMessage(status_text, 5000)

    def _handle_cve_result(self, result_data, cve_object):
        """Adds a CVE result to the table and stores the full object."""
        item = QTreeWidgetItem(result_data)
        item.setData(0, Qt.ItemDataRole.UserRole, cve_object) # Store the full object
        self.cve_results_table.addTopLevelItem(item)

    def _handle_exploit_search_status(self, status_text):
        self.status_bar.showMessage(status_text, 5000)

    def _handle_exploit_search_results(self, results):
        """Adds exploit search results to the table."""
        self.exploitdb_results_table.clear()
        for result in results:
            item = QTreeWidgetItem(result)
            self.exploitdb_results_table.addTopLevelItem(item)

    def _handle_lab_status(self, status_text):
        self.status_bar.showMessage(status_text, 0) # 0 means it stays until changed

    def _handle_error(self, title, text):
        QMessageBox.critical(self, title, text)

    def _create_export_button(self, source_widget):
        button = QPushButton("Export Results")
        button.setToolTip("Export the results to a file (CSV, HTML, PDF, DOCX).")
        button.clicked.connect(lambda: self._handle_export(source_widget))
        return button

    def _handle_export(self, source_widget):
        if source_widget.topLevelItemCount() == 0:
            QMessageBox.information(self, "No Data", "There is no data to export.")
            return

        formats = "HTML (*.html);;CSV (*.csv);;PDF (*.pdf);;Word Document (*.docx)"
        file_path, selected_format = QFileDialog.getSaveFileName(self, "Export Results", "", formats, options=QFileDialog.Option.DontUseNativeDialog)

        if not file_path:
            return

        try:
            if 'html' in selected_format:
                self._export_to_html(source_widget, file_path)
            elif 'csv' in selected_format:
                self._export_to_csv(source_widget, file_path)
            elif 'pdf' in selected_format:
                self._export_to_pdf(source_widget, file_path)
            elif 'docx' in selected_format:
                self._export_to_docx(source_widget, file_path)
            else:
                QMessageBox.warning(self, "Unsupported Format", "Selected file format is not supported.")
                return
            self.status_bar.showMessage(f"Successfully exported results to {file_path}")
        except NameError:
            logging.error("Export failed due to missing optional dependencies.", exc_info=True)
            QMessageBox.critical(self, "Dependency Error", "Optional libraries for PDF/DOCX export are not installed.\nPlease run: pip install reportlab python-docx")
        except Exception as e:
            logging.error(f"Failed to export results: {e}", exc_info=True)
            QMessageBox.critical(self, "Export Error", f"An error occurred during export:\n{e}")

    def _export_to_csv(self, tree_widget, file_path):
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]
            writer.writerow(header)
            for i in range(tree_widget.topLevelItemCount()):
                item = tree_widget.topLevelItem(i)
                row = [item.text(j) for j in range(tree_widget.columnCount())]
                writer.writerow(row)

    def _export_to_html(self, tree_widget, file_path):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("<html><head><title>Exported Results</title>")
            f.write("<style>body { font-family: sans-serif; } table { border-collapse: collapse; width: 100%; }")
            f.write("th, td { border: 1px solid #dddddd; text-align: left; padding: 8px; }")
            f.write("tr:nth-child(even) { background-color: #f2f2f2; }</style></head><body>")
            f.write("<h2>Exported Results</h2><table><tr>")
            header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]
            for h in header:
                f.write(f"<th>{h}</th>")
            f.write("</tr>")
            for i in range(tree_widget.topLevelItemCount()):
                f.write("<tr>")
                item = tree_widget.topLevelItem(i)
                for j in range(tree_widget.columnCount()):
                    f.write(f"<td>{item.text(j)}</td>")
                f.write("</tr>")
            f.write("</table></body></html>")

    def _export_to_pdf(self, tree_widget, file_path):
        doc = SimpleDocTemplate(file_path)
        elements = []
        styles = getSampleStyleSheet()
        elements.append(Paragraph("GScapy Exported Results", styles['h1']))

        header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]
        data = [header]
        for i in range(tree_widget.topLevelItemCount()):
            row = [tree_widget.topLevelItem(i).text(j) for j in range(tree_widget.columnCount())]
            data.append(row)

        table = Table(data)
        style = TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ])
        table.setStyle(style)
        elements.append(table)
        doc.build(elements)

    def _export_to_docx(self, tree_widget, file_path):
        document = docx.Document()
        document.add_heading('GScapy Exported Results', 0)

        header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]

        table = document.add_table(rows=1, cols=len(header))
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(header):
            hdr_cells[i].text = h

        for i in range(tree_widget.topLevelItemCount()):
            row_cells = table.add_row().cells
            item = tree_widget.topLevelItem(i)
            for j in range(tree_widget.columnCount()):
                row_cells[j].text = item.text(j)

        document.save(file_path)

    def _update_tool_targets(self):
        """Automatically updates tool target fields based on the selected interface."""
        iface_name = self.get_selected_iface()

        network_cidr = "192.168.1.0/24" # Default fallback
        if iface_name and iface_name != "Automatic":
            try:
                addrs = psutil.net_if_addrs().get(iface_name, [])
                for addr in addrs:
                    if addr.family == socket.AF_INET:
                        ip = addr.address
                        netmask = addr.netmask
                        if ip and netmask:
                            # Use ipaddress module to calculate network CIDR
                            host_iface = ipaddress.IPv4Interface(f"{ip}/{netmask}")
                            network_cidr = host_iface.network.with_prefixlen
                            logging.info(f"Updated tool targets for interface {iface_name} to {network_cidr}")
                            break # Found the IPv4 addr, no need to continue
            except Exception as e:
                logging.error(f"Could not auto-populate tool targets for {iface_name}: {e}")
                # Keep the default fallback

        # Update all relevant tool fields
        if hasattr(self, 'arp_target'):
            self.arp_target.setText(network_cidr)
        if hasattr(self, 'ps_target_edit'):
            self.ps_target_edit.setText(network_cidr)

    def closeEvent(self, event):
        """Shows a confirmation dialog and ensures background threads are stopped on exit."""
        reply = QMessageBox.question(self, 'Exit Confirmation',
                                     "Are you sure you want to exit?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            logging.info("User confirmed exit. Stopping background threads.")
            if self.sniffer_thread and self.sniffer_thread.isRunning(): self.sniffer_thread.stop()
            if self.channel_hopper and self.channel_hopper.isRunning(): self.channel_hopper.stop()
            if self.resource_monitor_thread and self.resource_monitor_thread.isRunning():
                self.resource_monitor_thread.stop()
            logging.info("GScapy application closing.")
            event.accept()
        else:
            logging.info("User canceled exit.")
            event.ignore()

    def _create_reporting_tab(self):
        """Creates the UI for the Reporting & Analysis tab."""
        widget = QWidget()
        main_layout = QHBoxLayout(widget)
        main_splitter = QSplitter(Qt.Orientation.Horizontal)

        # --- Left Panel: Configuration & Narrative ---
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(0,0,0,0)

        # Rules of Engagement Box
        roe_box = QGroupBox("Rules of Engagement (ROE)")
        roe_layout = QFormLayout(roe_box)
        self.report_client_name = QLineEdit()
        self.report_assessment_dates = QLineEdit()
        self.report_objectives = QTextEdit()
        self.report_objectives.setPlaceholderText(
"""Example:
- Objective 1: Determine the ability of a threat actor to compromise critical customer transactional data.
- Objective 2: Evaluate the integrity of the customer's order database.
- Objective 3: Assess the effectiveness of Incident Response procedures."""
        )
        self.report_in_scope = QTextEdit()
        self.report_in_scope.setPlaceholderText(
"""Example:

--- Authorized Target Space ---
- IP Range(s): 10.10.12.0/24, 10.10.13.0/24
- Domains: *.example.com
- URLs: https://www.example.com/login
- Network Segments: Corporate LAN, Guest Wi-Fi

--- Authorized Hosts ---
- All hosts not expressly restricted."""
        )
        self.report_out_of_scope = QTextEdit()
        self.report_out_of_scope.setPlaceholderText(
"""Example:

--- Explicit Restrictions ---
- No Denial of Service (DoS) attacks.
- No testing outside of business hours (9am-5pm Local Time).
- Social engineering of staff is not permitted.

--- Restricted IP Addresses ---
- 10.10.10.0/24 (HR Department)
- 10.10.11.0/24 (Accounting)

--- Restricted Hosts ---
- CRITICAL_DB_SERVER_01"""
        )
        roe_layout.addRow("Client Name:", self.report_client_name)
        roe_layout.addRow("Assessment Dates:", self.report_assessment_dates)
        roe_layout.addRow("Objectives:", self.report_objectives)
        roe_layout.addRow("In-Scope Targets:", self.report_in_scope)
        roe_layout.addRow("Out-of-Scope & Restrictions:", self.report_out_of_scope)
        left_layout.addWidget(roe_box)

        # Executive Summary Box
        summary_box = QGroupBox("Executive Summary")
        summary_layout = QVBoxLayout(summary_box)
        self.report_summary_text = QTextEdit()
        self.report_summary_text.setPlaceholderText("Write a high-level summary of the assessment's findings and recommendations for a non-technical audience, or generate one with AI.")

        ai_summary_btn = QPushButton(QIcon("icons/terminal.svg"), " Generate Summary with AI")
        ai_summary_btn.clicked.connect(self._handle_ai_summary_generation)

        summary_layout.addWidget(self.report_summary_text)
        summary_layout.addWidget(ai_summary_btn)
        left_layout.addWidget(summary_box)

        left_panel.setLayout(left_layout)

        # --- Right Panel: Findings & Generation ---
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(0,0,0,0)

        # Findings Box
        findings_box = QGroupBox("Aggregated Findings")
        findings_layout = QVBoxLayout(findings_box)

        aggregation_controls = QHBoxLayout()
        self.report_aggregate_btn = QPushButton(QIcon("icons/search.svg"), " Aggregate & Enrich Results")
        self.report_aggregate_btn.setToolTip("Scan the results from all tool outputs in the current session and enrich them with CVE and Exploit-DB information.")
        self.report_aggregate_btn.clicked.connect(self._handle_aggregation)
        aggregation_controls.addWidget(self.report_aggregate_btn)

        self.update_cve_db_btn = QPushButton(QIcon("icons/download-cloud.svg"), "Update Offline DB")
        self.update_cve_db_btn.setToolTip("Download or update the offline CVE database from NVD. This may take a while.")
        self.update_cve_db_btn.clicked.connect(self._start_cve_db_update)
        aggregation_controls.addWidget(self.update_cve_db_btn)

        aggregation_controls.addStretch()
        self.offline_cve_check = QCheckBox("Use offline CVE_DB")
        self.offline_cve_check.setToolTip("Use a local copy of CVE data for enrichment. Requires initial download.")
        aggregation_controls.addWidget(self.offline_cve_check)
        findings_layout.addLayout(aggregation_controls)

        self.report_findings_tree = QTreeWidget()
        self.report_findings_tree.setColumnCount(4)
        self.report_findings_tree.setHeaderLabels(["Host", "Port/Service", "Vulnerability/Finding", "Details (CVE/Exploit)"])
        self.report_findings_tree.header().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.report_findings_tree.header().setStretchLastSection(True)
        findings_layout.addWidget(self.report_findings_tree)
        right_layout.addWidget(findings_box)

        # Generation Box
        generation_box = QGroupBox("Report Generation")
        generation_layout = QVBoxLayout(generation_box)

        # --- AI Generation ---
        ai_box = QGroupBox("AI-Powered Generation")
        ai_layout = QFormLayout(ai_box)
        self.ai_persona_combo = QComboBox()
        self.ai_persona_combo.addItems(["Default", "Technical Manager", "C-Suite Executive", "Lead Developer"])
        self.ai_persona_combo.setToolTip("Select a persona for the AI to adopt when generating report sections.")
        ai_layout.addRow("AI Persona:", self.ai_persona_combo)

        self.ai_instructions_edit = QTextEdit()
        self.ai_instructions_edit.setPlaceholderText("Optional: Provide custom instructions for the AI. For example, 'Focus on the financial impact of the SQL injection vulnerability'.")
        self.ai_instructions_edit.setFixedHeight(80)
        ai_layout.addRow("AI Instructions:", self.ai_instructions_edit)

        self.report_generate_ai_btn = QPushButton(QIcon("icons/terminal.svg"), "Generate with AI")
        self.report_generate_ai_btn.setToolTip("Use the AI Assistant to generate sections of the report based on findings.")
        self.report_generate_ai_btn.clicked.connect(self._handle_ai_report_generation)
        ai_layout.addRow(self.report_generate_ai_btn)
        generation_layout.addWidget(ai_box)

        # --- Final Report Generation ---
        final_report_box = QGroupBox("Final Export")
        final_report_layout = QFormLayout(final_report_box)

        self.report_template_combo = QComboBox()
        try:
            templates = [f for f in os.listdir("report_templates") if f.endswith('.html')]
            self.report_template_combo.addItems(templates)
        except FileNotFoundError:
            logging.error("report_templates directory not found. Report generation may fail.")
            self.report_template_combo.addItem("default_report.html")

        final_report_layout.addRow("HTML Template:", self.report_template_combo)

        self.report_generate_html_btn = QPushButton(QIcon("icons/file-text.svg"), "Generate Final HTML Report")
        self.report_generate_html_btn.setToolTip("Compile all the information above into a final HTML report document.")
        self.report_generate_html_btn.clicked.connect(self._handle_generate_report) # Connect to existing handler for now

        self.report_generate_doc_btn = QToolButton()
        self.report_generate_doc_btn.setText("Generate Document")
        self.report_generate_doc_btn.setIcon(QIcon("icons/file.svg"))
        self.report_generate_doc_btn.setPopupMode(QToolButton.ToolButtonPopupMode.MenuButtonPopup)
        self.report_generate_doc_btn.setToolTip("Generate the report in various document formats (e.g., DOCX, PDF).")

        doc_menu = QMenu(self)
        doc_menu.addAction("Export as DOCX", lambda: self._handle_generate_doc_report('docx'))
        doc_menu.addAction("Export as PDF", lambda: self._handle_generate_doc_report('pdf'))
        self.report_generate_doc_btn.setMenu(doc_menu)

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.report_generate_html_btn)
        button_layout.addWidget(self.report_generate_doc_btn)
        final_report_layout.addRow(button_layout)

        generation_layout.addWidget(final_report_box)
        right_layout.addWidget(generation_box)

        right_panel.setLayout(right_layout)

        # --- Add panels to splitter ---
        main_splitter.addWidget(left_panel)
        main_splitter.addWidget(right_panel)
        main_splitter.setSizes([400, 600]) # Initial sizing
        main_layout.addWidget(main_splitter)

        return widget

    def _handle_ai_summary_generation(self):
        """Gathers findings, sends them to the AI, and sets a callback to populate the summary."""
        if self.report_findings_tree.topLevelItemCount() == 0:
            QMessageBox.warning(self, "No Data", "There are no findings to summarize. Please run the 'Aggregate & Enrich Results' tool first.")
            return

        findings_text = ""
        for i in range(self.report_findings_tree.topLevelItemCount()):
            item = self.report_findings_tree.topLevelItem(i)
            host = item.text(0)
            service = item.text(1)
            finding = item.text(2)
            details = item.text(3)
            findings_text += f"- Host: {host}, Service: {service}, Finding: {finding}\n  Details: {details}\n\n"

        prompt = (
            "Based on the following list of penetration testing findings, please write a concise executive summary "
            "suitable for a non-technical audience. Focus on the overall risk posture, key areas of weakness, "
            "and high-level recommendations. The summary should be a few paragraphs long.\n\n"
            f"--- FINDINGS ---\n{findings_text}--- END FINDINGS ---"
        )

        def _populate_summary(generated_text):
            self.report_summary_text.setPlainText(generated_text)
            QMessageBox.information(self, "Success", "AI-generated summary has been populated.")

        self.ai_assistant_tab.set_completion_callback(_populate_summary)
        self.ai_assistant_tab.send_message(prompt)
        self.tab_widget.setCurrentWidget(self.ai_assistant_tab)
        QMessageBox.information(self, "AI Task Started", "The AI is generating the summary. You will be notified upon completion. You can watch the progress in the 'AI Assistant' tab.")

    def _create_lab_tab(self):
        """Creates the UI for the LAB / Test Chaining tab."""
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # Top control bar
        controls_bar = QHBoxLayout()
        self.lab_run_chain_btn = QPushButton(QIcon("icons/play-circle.svg"), " Run Test Chain")
        self.lab_save_chain_btn = QPushButton(QIcon("icons/save.svg"), " Save Chain")
        self.lab_load_chain_btn = QPushButton(QIcon("icons/folder.svg"), " Load Chain")
        self.lab_run_chain_btn.clicked.connect(self.start_lab_chain)
        self.lab_save_chain_btn.clicked.connect(self._lab_save_chain)
        self.lab_load_chain_btn.clicked.connect(self._lab_load_chain)
        controls_bar.addWidget(self.lab_run_chain_btn)
        controls_bar.addWidget(self.lab_save_chain_btn)
        controls_bar.addWidget(self.lab_load_chain_btn)
        controls_bar.addStretch()
        main_layout.addLayout(controls_bar)

        main_splitter = QSplitter(Qt.Orientation.Horizontal)

        left_panel = QGroupBox("Available Tools")
        left_layout = QVBoxLayout(left_panel)
        self.lab_tools_list = QListWidget()
        left_layout.addWidget(self.lab_tools_list)
        main_splitter.addWidget(left_panel)

        center_panel = QWidget()
        center_layout = QVBoxLayout(center_panel)
        center_layout.setContentsMargins(0,0,0,0)
        chain_box = QGroupBox("Test Chain Sequence")
        chain_layout = QVBoxLayout(chain_box)
        self.lab_chain_list = QListWidget()
        chain_layout.addWidget(self.lab_chain_list)

        chain_buttons = QHBoxLayout()
        add_btn = QPushButton("Add ->"); add_btn.clicked.connect(self._lab_add_step)
        remove_btn = QPushButton("<- Remove"); remove_btn.clicked.connect(self._lab_remove_step)
        move_up_btn = QPushButton("Move Up"); move_up_btn.clicked.connect(self._lab_move_step_up)
        move_down_btn = QPushButton("Move Down"); move_down_btn.clicked.connect(self._lab_move_step_down)
        chain_buttons.addWidget(add_btn); chain_buttons.addWidget(remove_btn)
        chain_buttons.addStretch(); chain_buttons.addWidget(move_up_btn); chain_buttons.addWidget(move_down_btn)
        chain_layout.addLayout(chain_buttons)
        center_layout.addWidget(chain_box)
        main_splitter.addWidget(center_panel)

        right_panel = QGroupBox("Step Configuration")
        self.lab_config_stack = QStackedWidget()
        right_panel.setLayout(QVBoxLayout())
        right_panel.layout().addWidget(self.lab_config_stack)

        placeholder_widget = QLabel("Select a step from the chain to configure it.")
        placeholder_widget.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lab_config_stack.addWidget(placeholder_widget)

        # Create new, independent instances of config widgets and their controls for the LAB
        self.lab_tool_configs = {}

        compatible_tools = [
            ("Nmap Scan", self._create_nmap_config_widget),
            ("Subdomain Scanner (Sublist3r)", self._create_subdomain_scanner_config_widget),
            ("Subdomain Scanner (Subfinder)", self._create_subfinder_config_widget),
            ("httpx Probe", self._create_httpx_config_widget),
            ("RustScan", self._create_rustscan_config_widget),
            ("dirsearch", self._create_dirsearch_config_widget),
            ("ffuf", self._create_ffuf_config_widget),
            ("enum4linux-ng", self._create_enum4linux_ng_config_widget),
            ("dnsrecon", self._create_dnsrecon_config_widget),
            ("fierce", self._create_fierce_config_widget),
            ("Nikto Scan", self._create_nikto_config_widget),
            ("Gobuster", self._create_gobuster_config_widget),
            ("Nuclei Scanner", self._create_nuclei_config_widget),
            ("TruffleHog Scanner", self._create_trufflehog_config_widget),
            ("John the Ripper", self._create_jtr_config_widget),
            ("Hydra", self._create_hydra_config_widget),
            ("Sherlock", self._create_sherlock_config_widget),
            ("Spiderfoot", self._create_spiderfoot_config_widget),
            ("ARP Scan (CLI)", self._create_arp_scan_cli_config_widget)
        ]

        for name, config_method in compatible_tools:
            self.lab_tools_list.addItem(name)
            config_widget, controls = config_method()
            self.lab_tool_configs[name] = {'widget': config_widget, 'controls': controls}
            self.lab_config_stack.addWidget(config_widget)

        self.lab_chain_list.currentItemChanged.connect(self._lab_on_chain_selection_changed)
        main_splitter.addWidget(right_panel)
        main_splitter.setSizes([200, 300, 500])
        main_layout.addWidget(main_splitter)

        return widget

    def _get_config_from_ui(self, tool_name):
        """Reads the current values from a tool's config UI and returns a dict."""
        config = {}
        if tool_name not in self.lab_tool_configs:
            return config

        controls = self.lab_tool_configs[tool_name]['controls']

        if tool_name == "Nmap Scan":
            for key, widget in controls.items():
                if isinstance(widget, QLineEdit):
                    config[key] = widget.text()
                elif isinstance(widget, QCheckBox):
                    config[key] = widget.isChecked()
                elif isinstance(widget, QComboBox):
                    config[key] = widget.currentText()
        # Add other tools here as they are implemented

        return config

    def _set_config_to_ui(self, tool_name, config):
        """Populates a tool's config UI from a config dict."""
        if tool_name not in self.lab_tool_configs:
            return

        controls = self.lab_tool_configs[tool_name]['controls']

        if tool_name == "Nmap Scan":
            for key, widget in controls.items():
                if key in config:
                    if isinstance(widget, QLineEdit):
                        widget.setText(config[key])
                    elif isinstance(widget, QCheckBox):
                        widget.setChecked(config[key])
                    elif isinstance(widget, QComboBox):
                        widget.setCurrentText(config[key])
        # Add other tools here as they are implemented

    def _lab_add_step(self):
        """Adds a selected tool from the available list to the test chain."""
        selected_item = self.lab_tools_list.currentItem()
        if not selected_item:
            QMessageBox.warning(self, "No Tool Selected", "Please select a tool from the 'Available Tools' list to add.")
            return

        tool_name = selected_item.text()

        # Get the default configuration from the current state of the tool's UI
        default_config = self._get_config_from_ui(tool_name)

        step_data = {
            'tool_name': tool_name,
            'id': str(uuid.uuid4()),
            'config': default_config
        }
        self.lab_test_chain.append(step_data)

        # The list widget item stores the unique ID to link it back to the config
        list_item = QListWidgetItem(f"Step {len(self.lab_test_chain)}: {tool_name}")
        list_item.setData(Qt.ItemDataRole.UserRole, step_data['id'])
        self.lab_chain_list.addItem(list_item)
        self.lab_chain_list.setCurrentItem(list_item)


    def _lab_remove_step(self):
        """Removes the selected step from the test chain."""
        selected_row = self.lab_chain_list.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "No Step Selected", "Please select a step from the chain to remove.")
            return

        # Remove from UI
        item = self.lab_chain_list.takeItem(selected_row)
        item_id = item.data(Qt.ItemDataRole.UserRole)

        # Remove from backend data model
        self.lab_test_chain = [step for step in self.lab_test_chain if step['id'] != item_id]

        # Renumber the remaining steps in the UI for clarity
        for i in range(self.lab_chain_list.count()):
            list_item = self.lab_chain_list.item(i)
            tool_name = list_item.text().split(": ")[1]
            list_item.setText(f"Step {i + 1}: {tool_name}")

    def _lab_move_step_up(self):
        """Moves the selected step up in the test chain."""
        current_row = self.lab_chain_list.currentRow()
        if current_row > 0:
            item = self.lab_chain_list.takeItem(current_row)
            self.lab_chain_list.insertItem(current_row - 1, item)
            self.lab_chain_list.setCurrentRow(current_row - 1)
            # Reorder the backend list as well
            self.lab_test_chain.insert(current_row - 1, self.lab_test_chain.pop(current_row))
            self._renumber_lab_steps()

    def _lab_move_step_down(self):
        """Moves the selected step down in the test chain."""
        current_row = self.lab_chain_list.currentRow()
        if 0 <= current_row < self.lab_chain_list.count() - 1:
            item = self.lab_chain_list.takeItem(current_row)
            self.lab_chain_list.insertItem(current_row + 1, item)
            self.lab_chain_list.setCurrentRow(current_row + 1)
            # Reorder the backend list as well
            self.lab_test_chain.insert(current_row + 1, self.lab_test_chain.pop(current_row))
            self._renumber_lab_steps()

    def _renumber_lab_steps(self):
        """Updates the text of the items in the lab chain list to reflect their new order."""
        for i in range(self.lab_chain_list.count()):
            item = self.lab_chain_list.item(i)
            # The tool name doesn't change, just the step number
            tool_name = self.lab_test_chain[i]['tool_name']
            item.setText(f"Step {i + 1}: {tool_name}")

    def _lab_on_chain_selection_changed(self, current, previous):
        """Shows the correct configuration widget when a step in the chain is selected."""
        # 1. Save the configuration of the previously selected item
        if previous:
            prev_id = previous.data(Qt.ItemDataRole.UserRole)
            # Find the corresponding step in the backend list
            for step in self.lab_test_chain:
                if step['id'] == prev_id:
                    tool_name = step['tool_name']
                    # Get the current UI state and save it to the step's config
                    step['config'] = self._get_config_from_ui(tool_name)
                    logging.info(f"Saved config for step {step['tool_name']} (ID: {prev_id})")
                    break

        # 2. Load the configuration of the currently selected item
        if not current:
            self.lab_config_stack.setCurrentIndex(0) # Show placeholder
            return

        current_id = current.data(Qt.ItemDataRole.UserRole)
        # Find the corresponding step in the backend list
        for step in self.lab_test_chain:
            if step['id'] == current_id:
                tool_name = step['tool_name']
                config = step['config']

                # Use the dedicated lab widgets
                if tool_name in self.lab_tool_configs:
                    # Populate the UI with the stored config
                    self._set_config_to_ui(tool_name, config)

                    # Show the correct widget from the stack
                    widget_to_show = self.lab_tool_configs[tool_name]['widget']
                    self.lab_config_stack.setCurrentWidget(widget_to_show)
                    logging.info(f"Loaded config for step {tool_name} (ID: {current_id})")
                else:
                    self.lab_config_stack.setCurrentIndex(0) # Fallback to placeholder
                return

    def _lab_save_chain(self):
        """Saves the current test chain to a JSON file."""
        if not self.lab_test_chain:
            QMessageBox.information(self, "Empty Chain", "There is nothing to save.")
            return

        file_path, _ = QFileDialog.getSaveFileName(self, "Save Test Chain", "", "GScapy LAB Files (*.gscapy-lab)", options=QFileDialog.Option.DontUseNativeDialog)
        if not file_path:
            return

        try:
            with open(file_path, 'w') as f:
                json.dump(self.lab_test_chain, f, indent=4)
            self.status_bar.showMessage(f"Test chain saved to {file_path}", 5000)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save test chain: {e}")
            logging.error(f"Failed to save LAB chain: {e}", exc_info=True)

    def _lab_load_chain(self):
        """Loads a test chain from a JSON file."""
        if self.lab_test_chain:
            reply = QMessageBox.question(self, "Confirm Load", "This will overwrite your current test chain. Are you sure?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.No:
                return

        file_path, _ = QFileDialog.getOpenFileName(self, "Load Test Chain", "", "GScapy LAB Files (*.gscapy-lab);;All Files (*)", options=QFileDialog.Option.DontUseNativeDialog)
        if not file_path:
            return

        try:
            with open(file_path, 'r') as f:
                loaded_chain = json.load(f)

            # Basic validation
            if not isinstance(loaded_chain, list) or not all('tool_name' in d and 'config' in d for d in loaded_chain):
                raise ValueError("Invalid file format.")

            self.lab_test_chain = loaded_chain
            self.lab_chain_list.clear()

            # Repopulate the UI list
            for i, step in enumerate(self.lab_test_chain):
                # Ensure each step has a unique ID if loading older formats
                if 'id' not in step:
                    step['id'] = str(uuid.uuid4())

                list_item = QListWidgetItem(f"Step {i + 1}: {step['tool_name']}")
                list_item.setData(Qt.ItemDataRole.UserRole, step['id'])
                self.lab_chain_list.addItem(list_item)

            self.status_bar.showMessage(f"Test chain loaded from {file_path}", 5000)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load test chain: {e}")
            logging.error(f"Failed to load LAB chain: {e}", exc_info=True)

    def start_lab_chain(self):
        """Starts the LAB chain execution worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return
        if not self.lab_test_chain:
            QMessageBox.information(self, "Empty Chain", "There are no steps in the test chain to run.")
            return

        self.is_tool_running = True
        self.lab_run_chain_btn.setEnabled(False)
        # In the future, a cancel button will be added
        # self.lab_cancel_chain_btn.setEnabled(True)
        self.tool_stop_event.clear()

        # Deepcopy the chain to avoid race conditions if the user edits it while running
        chain_to_run = copy.deepcopy(self.lab_test_chain)

        self.worker = WorkerThread(self._lab_chain_thread, args=(chain_to_run,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _lab_chain_thread(self, chain):
        """Worker thread that executes each step of the LAB chain."""
        q = self.tool_results_queue
        logging.info(f"LAB chain execution started with {len(chain)} steps.")
        lab_context = {} # This dictionary will hold results passed between steps

        for i, step in enumerate(chain):
            if self.tool_stop_event.is_set():
                logging.info("LAB chain execution cancelled by user.")
                break

            tool_name = step['tool_name']
            config = step['config']
            q.put(('lab_status', f"Executing Step {i+1}/{len(chain)}: {tool_name}"))

            # This is where the magic happens.
            # For now, just log the config.
            logging.info(f"--- Running LAB Step {i+1}: {tool_name} ---")
            logging.info(f"Config: {json.dumps(config, indent=2)}")

            # Placeholder for actual execution
            time.sleep(2) # Simulate work


        q.put(('tool_finished', 'lab_chain'))
        logging.info("LAB chain execution finished.")


import database

def main():
    """Main function to launch the GScapy application."""
    try:
        database.initialize_database()
        if 'scapy' not in sys.modules: raise ImportError

        app = QApplication(sys.argv)

        login_dialog = LoginDialog()

        # Define the custom stylesheet additions
        extra_qss = {
            'QGroupBox': {
                'border': '1px solid #444;',
                'border-radius': '8px',
                'margin-top': '10px',
            },
            'QGroupBox::title': {
                'subcontrol-origin': 'margin',
                'subcontrol-position': 'top left',
                'padding': '0 10px',
            },
            'QTabWidget::pane': {
                'border-top': '1px solid #444;',
                'margin-top': '-1px',
            },
            'QFrame': {
                'border-radius': '8px',
            },
            'QPushButton': {
                'border-radius': '8px',
            },
            'QLineEdit': {
                'border-radius': '8px',
            },
            'QComboBox': {
                'border-radius': '8px',
            },
            'QTextEdit': {
                'border-radius': '8px',
            },
            'QPlainTextEdit': {
                'border-radius': '8px',
            },
            'QListWidget': {
                'border-radius': '8px',
            },
            'QTreeWidget': {
                'border-radius': '8px',
            }
        }

        # Apply the default theme before showing the login dialog
        apply_stylesheet(app, theme=login_dialog.selected_theme, extra=extra_qss)

        if login_dialog.exec() != QDialog.DialogCode.Accepted:
            sys.exit(0)

        # Re-apply the theme in case the user changed it in the dialog.
        # This ensures the main window gets the final selected theme.
        apply_stylesheet(app, theme=login_dialog.selected_theme, extra=extra_qss)

        window = GScapy()
        window.current_user = login_dialog.current_user
        # Set window title with username
        if window.current_user and 'username' in window.current_user:
            window.setWindowTitle(f"Welcome, {window.current_user['username']} - GScapy + AI - The Modern Scapy Interface with AI")
        window._update_menu_bar() # Populate the menu now that we have a user
        window.show()
        sys.exit(app.exec())

    except ImportError:
        app = QApplication(sys.argv)
        QMessageBox.critical(None, "Fatal Error", "Scapy is not installed.")
        sys.exit(1)
    except Exception as e:
        logging.critical(f"An unhandled exception occurred: {e}", exc_info=True)
        app = QApplication(sys.argv)
        QMessageBox.critical(None, "Unhandled Exception", f"An unexpected error occurred:\n\n{e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
